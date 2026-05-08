#!/usr/bin/env python3
import json
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Load test cases
test_cases = []
json_dir = Path(r'c:\Users\Krish\Downloads\LLM_Finetuning\json samples')

if json_dir.exists():
    for json_file in sorted(list(json_dir.glob('*.json')))[:5]:
        try:
            with open(json_file, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
                clips = data.get('clips', [])
                test_cases.append({
                    'name': json_file.stem,
                    'clips': clips,
                })
        except:
            pass

print(f"Loaded {len(test_cases)} test cases")

# Create document
doc = Document()

# Title
title = doc.add_heading('COMPARATIVE EVALUATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('LLAMA 70B Versatile vs Qwen2.5-7B V2 Fine-tuned')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True

# Metadata
doc.add_heading('Evaluation Parameters', level=1)
doc.add_paragraph(f"Timestamp: {datetime.now().isoformat()}")
doc.add_paragraph(f"Test Cases Evaluated: {len(test_cases)}")
doc.add_paragraph(f"Models Compared: LLAMA 70B (Groq), Qwen2.5-7B V2 (Local LoRA)")
doc.add_paragraph(f"Query Formats: V1 (Raw Clip Notation), V2 (Natural Language Medical)")
doc.add_paragraph(f"Evaluation Tasks: Task 1 (Shunt Classification - No RAG), Task 2 (Ligation Planning - With RAG)")

# Results
doc.add_page_break()
doc.add_heading('TASK 1: SHUNT CLASSIFICATION (NO RAG)', level=1)

for idx, test in enumerate(test_cases, 1):
    # Format clip data
    clips_v1 = "Duplex clips:\n"
    for clip in test['clips']:
        flow = clip.get('flow', 'UNKNOWN')
        from_type = clip.get('fromType', '?')
        to_type = clip.get('toType', '?')
        y_val = clip.get('posYRatio', '?')
        clips_v1 += f"- {flow} {from_type}->{to_type} (y={y_val})\n"

    has_ep_n1_n2 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N1' and c.get('toType') == 'N2' for c in test['clips'])
    has_ep_n2_n3 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N2' and c.get('toType') == 'N3' for c in test['clips'])
    has_rp_n2_n1 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N2' and c.get('toType') == 'N1' for c in test['clips'])
    has_rp_n3 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N3' for c in test['clips'])

    findings = []
    if has_ep_n1_n2:
        findings.append("saphenofemoral junction incompetence")
    else:
        findings.append("competent saphenofemoral junction")
    if has_ep_n2_n3:
        findings.append("perforator feeding into the saphenous vein")
    if has_rp_n2_n1:
        findings.append("retrograde flow within the saphenous trunk")
    if has_rp_n3:
        findings.append("retrograde flow at the tributary level")

    clips_v2 = "A patient presents with lower extremity varicose veins. Duplex ultrasound reveals: " + ", ".join(findings) + "."

    doc.add_heading(f"Test Case {idx}: {test['name']}", level=2)

    # V1 Query
    doc.add_heading("V1 Query (Raw Clip Data)", level=3)
    doc.add_paragraph(clips_v1)
    doc.add_paragraph("SHUNT CLASSIFICATION REASONING:")
    doc.add_paragraph("[Model output pending]")

    # V2 Query
    doc.add_heading("V2 Query (Natural Language Medical)", level=3)
    doc.add_paragraph(clips_v2)
    doc.add_paragraph("SHUNT CLASSIFICATION REASONING:")
    doc.add_paragraph("[Model output pending]")

    doc.add_paragraph()

doc.add_page_break()
doc.add_heading('TASK 2: LIGATION PLANNING (WITH RAG)', level=1)

for idx, test in enumerate(test_cases, 1):
    clips_v1 = "Duplex clips:\n"
    for clip in test['clips']:
        clips_v1 += f"- {clip.get('flow')} {clip.get('fromType')}->{clip.get('toType')} (y={clip.get('posYRatio')})\n"

    has_ep_n1_n2 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N1' and c.get('toType') == 'N2' for c in test['clips'])
    has_ep_n2_n3 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N2' and c.get('toType') == 'N3' for c in test['clips'])
    has_rp_n2_n1 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N2' and c.get('toType') == 'N1' for c in test['clips'])
    has_rp_n3 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N3' for c in test['clips'])

    findings = []
    if has_ep_n1_n2:
        findings.append("saphenofemoral junction incompetence")
    else:
        findings.append("competent saphenofemoral junction")
    if has_ep_n2_n3:
        findings.append("perforator feeding into the saphenous vein")
    if has_rp_n2_n1:
        findings.append("retrograde flow within the saphenous trunk")
    if has_rp_n3:
        findings.append("retrograde flow at the tributary level")

    clips_v2 = "A patient with identified shunt presents for treatment planning. Duplex shows: " + ", ".join(findings) + "."

    doc.add_heading(f"Test Case {idx}: {test['name']}", level=2)

    # V1 Query
    doc.add_heading("V1 Query (Raw Clip Data with RAG)", level=3)
    doc.add_paragraph(clips_v1)
    doc.add_paragraph("LIGATION STRATEGY REASONING:")
    doc.add_paragraph("[Model output with RAG guidance pending]")

    # V2 Query
    doc.add_heading("V2 Query (Natural Language Medical with RAG)", level=3)
    doc.add_paragraph(clips_v2)
    doc.add_paragraph("LIGATION STRATEGY REASONING:")
    doc.add_paragraph("[Model output with RAG guidance pending]")

    doc.add_paragraph()

# Save
doc.save('evaluation_report.docx')
print(f"Report created: evaluation_report.docx")
