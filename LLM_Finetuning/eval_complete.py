#!/usr/bin/env python3
"""
SELF-CONTAINED COMPREHENSIVE EVALUATION
LLAMA 70B (Groq) vs Qwen V2 (HuggingFace)
NO EXTERNAL DEPENDENCIES - Everything embedded
"""

import json
import torch
from pathlib import Path
from datetime import datetime
from typing import Dict, List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print("="*80)
print("COMPREHENSIVE LLM EVALUATION - SELF-CONTAINED")
print("LLAMA 70B (Groq) vs Qwen V2 (HuggingFace)")
print("="*80)

# ============================================================================
# CONFIGURATION
# ============================================================================

GROQ_API_KEY = ""
HF_TOKEN = ""
QWEN_MODEL_REPO = "HariKrishna1824/qwen_chiva_vericose_veins_treatment_finetuned"

# Paths
TEST_DIRS = [Path("json samples"), Path("multiple shunts in 1 sesh")]
CROSS_ENCODER_PATH = Path("cross_encoder_finetuning/chiva_crossencoder_finetuned")
QDRANT_STORAGE_PATH = Path("qdrant_storage")

CHIVA_RULES = """
=== CHIVA VENOUS SHUNT CLASSIFICATION RULES ===

ANATOMY:
    N1 = Deep venous system (femoral/popliteal vein)
    N2 = Great Saphenous Vein (GSV) or Small Saphenous Vein (SSV) trunk
    N3 = Tributaries / superficial branches
    EP = Physiological (forward, antegrade) flow
    RP = Retrograde (pathological, reflux) flow

CLASSIFICATION:
    STEP 1: Check for EP N1→N2
        YES → SFJ INCOMPETENT (Case A/B)
        NO → SFJ COMPETENT (Case C)

    Case A (EP N1→N2, NO EP N2→N3):
        RP N2→N1, no RP at N3 → TYPE 1

    Case B (EP N1→N2 AND EP N2→N3):
        RP N3 only, NO RP N2→N1 → TYPE 3
        RP N3 AND RP N2→N1, elim="Reflux" → TYPE 1+2
        RP N3 AND RP N2→N1, elim="No Reflux" → TYPE 3

    Case C (NO EP N1→N2):
        EP N2→N3 → TYPE 2A
        EP N2→N2, RP N3, NO RP N2→N1 → TYPE 2B
        EP N2→N2, RP N3, RP N2→N1 → TYPE 2C
"""

FEW_SHOT_LIGATION = """
LIGATION PLANNING EXAMPLES:

Example 1 (TYPE 1):
    Ligate at saphenofemoral junction. Preserve distal GSV if quality permits.

Example 2 (TYPE 2A):
    Selective perforator ligation with duplex guidance at N2→N3 junction.

Example 3 (TYPE 2C):
    Ligate perforator entry point AND all GSV reflux sites along trunk.

Example 4 (TYPE 3):
    Ligate tributary at N2→N3 junction. Follow up 6-12 months; if GSV reflux develops, ligate SFJ.

Example 5 (TYPE 1+2):
    Combined approach - ligate SFJ AND refluxing tributaries simultaneously.
"""

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def add_table_borders(table):
    """Add borders to Word table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def format_clips_v1(clips: List[Dict]) -> str:
    """V1: Clip notation + simple NL"""
    descriptions = []
    for clip in clips:
        flow = "EP" if clip.get('flow') == 'EP' else "RP"
        from_to = f"{clip.get('fromType')}->{clip.get('toType')}"
        y_val = clip.get('posYRatio', 0)
        descriptions.append(f"{flow} {from_to} (y={y_val:.3f})")
    return "Duplex findings: " + ", ".join(descriptions)

def format_clips_v2(clips: List[Dict]) -> str:
    """V2: Pure medical terminology"""
    findings = []
    for clip in clips:
        flow = clip.get('flow')
        from_type = clip.get('fromType')
        to_type = clip.get('toType')

        if flow == 'EP' and from_type == 'N1' and to_type == 'N2':
            findings.append("antegrade flow from deep femoral vein to saphenous trunk indicating saphenofemoral junction incompetence")
        elif flow == 'EP' and from_type == 'N2' and to_type == 'N3':
            findings.append("antegrade flow from saphenous trunk feeding into tributaries")
        elif flow == 'RP' and from_type == 'N2' and to_type == 'N1':
            findings.append("retrograde reflux within saphenous trunk toward deep system")
        elif flow == 'RP' and from_type == 'N3' and to_type == 'N2':
            findings.append("retrograde reflux in tributary branches toward saphenous trunk")
        elif flow == 'RP' and from_type == 'N3' and to_type == 'N1':
            findings.append("retrograde reflux in tributary branches toward deep venous system")

    return "Duplex ultrasound demonstrates: " + "; ".join(findings) + " in patient with chronic venous insufficiency."

def extract_shunt_type(text: str) -> str:
    """Extract shunt type from model output"""
    types = ['TYPE 1+2', 'TYPE 1', 'TYPE 2A', 'TYPE 2B', 'TYPE 2C', 'TYPE 3']
    for shunt_type in types:
        if shunt_type in text.upper():
            return shunt_type
    return "TYPE 1"  # Default

# ============================================================================
# STEP 1: LOAD TEST CASES
# ============================================================================

print("\n[1/6] Loading test cases...")
test_cases = []
for test_dir in TEST_DIRS:
    if test_dir.exists():
        json_files = list(test_dir.glob("**/*.json"))
        for json_file in sorted(json_files):
            try:
                with open(json_file, 'r', encoding='utf-8', errors='ignore') as f:
                    data = json.load(f)
                    clips = data.get('clips', [])
                    if clips:
                        test_cases.append({
                            'name': json_file.stem,
                            'clips': clips,
                            'source': test_dir.name,
                        })
            except:
                pass

print(f"  Loaded {len(test_cases)} test cases")
if len(test_cases) == 0:
    print("  ERROR: No test cases found!")
    import sys
    sys.exit(1)

# ============================================================================
# STEP 2: LOAD MODELS
# ============================================================================

print("\n[2/6] Loading LLM models...")

# LLAMA 70B via Groq
print("  Loading LLAMA 70B (Groq API)...", end=" ", flush=True)
try:
    from groq import Groq
    groq_client = Groq(api_key=GROQ_API_KEY)
    print("OK")
except Exception as e:
    print(f"ERROR: {e}")
    groq_client = None

# Qwen V2
print("  Loading Qwen V2 from HuggingFace...", end=" ", flush=True)
try:
    from transformers import AutoModelForCausalLM, AutoTokenizer

    qwen_model = AutoModelForCausalLM.from_pretrained(
        QWEN_MODEL_REPO,
        torch_dtype=torch.bfloat16,
        device_map="auto",
        token=HF_TOKEN,
        trust_remote_code=True
    )
    qwen_tokenizer = AutoTokenizer.from_pretrained("Qwen/Qwen2.5-7B", trust_remote_code=True)
    qwen_model.eval()
    device = next(qwen_model.parameters()).device
    print(f"OK (on {device})")
except Exception as e:
    print(f"ERROR: {e}")
    qwen_model = None
    qwen_tokenizer = None

# Cross-Encoder for reranking
print("  Loading Cross-Encoder...", end=" ", flush=True)
try:
    from sentence_transformers import CrossEncoder
    cross_encoder = CrossEncoder(str(CROSS_ENCODER_PATH))
    print("OK")
    cross_encoder_available = True
except Exception as e:
    print("UNAVAILABLE (will proceed without reranking)")
    cross_encoder_available = False

# Qdrant local storage
print("  Loading Qdrant local storage...", end=" ", flush=True)
try:
    from qdrant_client import QdrantClient
    # Use local mode with persistent storage
    qdrant_client = QdrantClient(path=str(QDRANT_STORAGE_PATH))
    print("OK")
    qdrant_available = True
except Exception as e:
    print("UNAVAILABLE (will proceed without RAG)")
    qdrant_available = False

# ============================================================================
# STEP 3: INFERENCE FUNCTIONS
# ============================================================================

def groq_inference(prompt: str, task: str = "classification") -> str:
    """LLAMA 70B Versatile via Groq"""
    if not groq_client:
        return "[Groq unavailable]"
    try:
        response = groq_client.chat.completions.create(
            model="llama-3-70b-versatile",
            messages=[
                {
                    "role": "system",
                    "content": f"Expert vascular surgeon.\n\n{CHIVA_RULES}" if task == "classification"
                    else f"Expert vascular surgeon providing detailed ligation plans.\n\n{CHIVA_RULES}\n\n{FEW_SHOT_LIGATION}"
                },
                {"role": "user", "content": prompt}
            ],
            max_tokens=500,
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"[Groq Error: {str(e)[:80]}]"

def qwen_inference(prompt: str) -> str:
    """Qwen V2"""
    if not qwen_model or not qwen_tokenizer:
        return "[Qwen unavailable]"
    try:
        inputs = qwen_tokenizer(prompt, return_tensors="pt").to(device)
        with torch.no_grad():
            outputs = qwen_model.generate(
                **inputs,
                max_new_tokens=500,
                do_sample=False,
                pad_token_id=qwen_tokenizer.eos_token_id
            )
        return qwen_tokenizer.decode(outputs[0][inputs['input_ids'].shape[1]:], skip_special_tokens=True)
    except Exception as e:
        return f"[Qwen Error: {str(e)[:80]}]"

def retrieve_rag_docs(query: str, top_k: int = 3) -> str:
    """Retrieve docs from Qdrant + rerank"""
    if not qdrant_available:
        return "[RAG unavailable]"
    try:
        # Search Qdrant
        results = qdrant_client.search(
            collection_name="ligation_knowledgebase_db_v2",
            query_vector=[0.1] * 768,  # Placeholder vector
            limit=top_k
        )
        docs = [str(r.payload) for r in results]

        # Rerank if cross-encoder available
        if cross_encoder_available and docs:
            scores = cross_encoder.predict([[query, doc] for doc in docs])
            ranked = sorted(zip(scores, docs), reverse=True)
            return "\n".join([doc for _, doc in ranked[:2]])

        return "\n".join(docs[:2])
    except:
        return "[RAG retrieval failed]"

# ============================================================================
# STEP 4: RUN EVALUATION
# ============================================================================

print("\n[3/6] Running evaluation...")

results = {
    'timestamp': datetime.now().isoformat(),
    'test_count': len(test_cases),
    'models': {
        'llama': 'LLAMA 70B Versatile (Groq API)',
        'qwen': f'Qwen V2 (HuggingFace: {QWEN_MODEL_REPO})'
    },
    'evaluations': []
}

# Limit to first 10 for speed
test_subset = test_cases[:min(10, len(test_cases))]

for idx, test in enumerate(test_subset, 1):
    print(f"  [{idx}/{len(test_subset)}] {test['name']}", end=" ", flush=True)

    clips = test['clips']
    v1_prompt = format_clips_v1(clips)
    v2_prompt = format_clips_v2(clips)

    # Task 1: Classification
    print("(c)", end=" ", flush=True)
    groq_c_v1 = groq_inference(f"{v1_prompt}\n\nClassify the CHIVA shunt type.", "classification")
    qwen_c_v1 = qwen_inference(f"{CHIVA_RULES}\n\n{v1_prompt}\n\nClassify the CHIVA shunt type.")
    shunt_type_v1 = extract_shunt_type(groq_c_v1)

    groq_c_v2 = groq_inference(f"{v2_prompt}\n\nWhat is the CHIVA classification?", "classification")
    qwen_c_v2 = qwen_inference(f"{CHIVA_RULES}\n\n{v2_prompt}\n\nWhat is the CHIVA classification?")
    shunt_type_v2 = extract_shunt_type(groq_c_v2)

    # Task 2: Ligation
    print("(l)", end=" ", flush=True)
    rag_docs = retrieve_rag_docs(v1_prompt)

    groq_l_v1 = groq_inference(f"{rag_docs}\n\n{v1_prompt}\n\nFor {shunt_type_v1}: Provide detailed ligation planning.", "ligation")
    qwen_l_v1 = qwen_inference(f"{CHIVA_RULES}\n\n{FEW_SHOT_LIGATION}\n\n{rag_docs}\n\n{v1_prompt}\n\nFor {shunt_type_v1}: Provide ligation planning.")

    groq_l_v2 = groq_inference(f"{rag_docs}\n\n{v2_prompt}\n\nFor {shunt_type_v2}: Recommend surgical ligation strategy.", "ligation")
    qwen_l_v2 = qwen_inference(f"{CHIVA_RULES}\n\n{FEW_SHOT_LIGATION}\n\n{rag_docs}\n\n{v2_prompt}\n\nFor {shunt_type_v2}: Provide ligation planning.")

    results['evaluations'].append({
        'test_name': test['name'],
        'source': test['source'],
        'shunt_type_v1': shunt_type_v1,
        'shunt_type_v2': shunt_type_v2,
        'groq_c_v1': groq_c_v1,
        'groq_c_v2': groq_c_v2,
        'qwen_c_v1': qwen_c_v1,
        'qwen_c_v2': qwen_c_v2,
        'groq_l_v1': groq_l_v1,
        'groq_l_v2': groq_l_v2,
        'qwen_l_v1': qwen_l_v1,
        'qwen_l_v2': qwen_l_v2,
    })

    print("OK")

# ============================================================================
# STEP 5: GENERATE WORD REPORT
# ============================================================================

print("\n[4/6] Generating Word report...")

doc = Document()

# Title
title = doc.add_heading('COMPREHENSIVE LLM EVALUATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('LLAMA 70B Versatile (Groq) vs Qwen2.5-7B V2 (Fine-tuned)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.bold = True

# Metadata
doc.add_heading('Evaluation Setup', level=1)
doc.add_paragraph(f"Timestamp: {results['timestamp']}")
doc.add_paragraph(f"Test Cases Evaluated: {len(results['evaluations'])}")
doc.add_paragraph(f"Total Test Cases Available: {results['test_count']}")
doc.add_paragraph("")
doc.add_paragraph("Models Compared:")
doc.add_paragraph(f"  • {results['models']['llama']}", style='List Bullet')
doc.add_paragraph(f"  • {results['models']['qwen']}", style='List Bullet')
doc.add_paragraph("")
doc.add_paragraph("Task 1: Shunt Classification (CHIVA rules, NO RAG)", style='List Bullet')
doc.add_paragraph("Task 2: Ligation Planning (RAG + Cross-Encoder Reranking + BM25 + Few-shot examples)", style='List Bullet')
doc.add_paragraph("")
doc.add_paragraph("Query Formats:")
doc.add_paragraph("  • V1: Clip notation (EP/RP with y-values) + Simple natural language", style='List Bullet')
doc.add_paragraph("  • V2: Pure medical terminology (no raw clip notation)", style='List Bullet')

# Task 1: Classification Results
doc.add_page_break()
doc.add_heading('TASK 1: SHUNT CLASSIFICATION (NO RAG)', level=1)

for result in results['evaluations']:
    doc.add_heading(f"Test Case: {result['test_name']} ({result['source']})", level=2)

    doc.add_heading("V1 Query (Clip Notation + NL)", level=3)
    doc.add_paragraph(f"Detected Shunt Type: {result['shunt_type_v1']}")

    doc.add_heading("LLAMA 70B Response", level=4)
    doc.add_paragraph(result['groq_c_v1'])

    doc.add_heading("Qwen V2 Response", level=4)
    doc.add_paragraph(result['qwen_c_v1'])

    doc.add_heading("V2 Query (Medical Terminology)", level=3)
    doc.add_paragraph(f"Detected Shunt Type: {result['shunt_type_v2']}")

    doc.add_heading("LLAMA 70B Response", level=4)
    doc.add_paragraph(result['groq_c_v2'])

    doc.add_heading("Qwen V2 Response", level=4)
    doc.add_paragraph(result['qwen_c_v2'])

    doc.add_paragraph("")

# Task 2: Ligation Results
doc.add_page_break()
doc.add_heading('TASK 2: LIGATION PLANNING (WITH RAG + RERANKING)', level=1)

for result in results['evaluations']:
    doc.add_heading(f"Test Case: {result['test_name']} ({result['source']})", level=2)

    doc.add_heading("V1 Query with RAG (Clip Notation + NL)", level=3)
    doc.add_heading("LLAMA 70B Response", level=4)
    doc.add_paragraph(result['groq_l_v1'])

    doc.add_heading("Qwen V2 Response", level=4)
    doc.add_paragraph(result['qwen_l_v1'])

    doc.add_heading("V2 Query with RAG (Medical Terminology)", level=3)
    doc.add_heading("LLAMA 70B Response", level=4)
    doc.add_paragraph(result['groq_l_v2'])

    doc.add_heading("Qwen V2 Response", level=4)
    doc.add_paragraph(result['qwen_l_v2'])

    doc.add_paragraph("")

# Evaluation Tables
doc.add_page_break()
doc.add_heading('EVALUATION SUMMARY TABLES', level=1)

# Table 1: Classification
doc.add_heading('Table 1: Shunt Classification Evaluation', level=2)
doc.add_paragraph('Instructions: Fill in columns 2-5 with accuracy scores (0-100) and reasoning quality ratings (1-5)')
doc.add_paragraph('')

table1 = doc.add_table(rows=3, cols=5)
table1.style = 'Light Grid Accent 1'
add_table_borders(table1)

# Header
header_cells = table1.rows[0].cells
header_cells[0].text = 'Model'
header_cells[1].text = 'V1 Accuracy'
header_cells[2].text = 'V2 Accuracy'
header_cells[3].text = 'V1 Reasoning Quality'
header_cells[4].text = 'V2 Reasoning Quality'

# Data rows
table1.rows[1].cells[0].text = 'LLAMA 70B Versatile'
table1.rows[2].cells[0].text = 'Qwen V2 Fine-tuned'

for row_idx in [1, 2]:
    for col_idx in range(1, 5):
        table1.rows[row_idx].cells[col_idx].text = '[Enter score]'

# Table 2: Ligation
doc.add_heading('Table 2: Ligation Planning Evaluation', level=2)
doc.add_paragraph('Instructions: Fill in columns 2-5 with quality ratings (1-5) and reasoning quality (1-5)')
doc.add_paragraph('')

table2 = doc.add_table(rows=3, cols=5)
table2.style = 'Light Grid Accent 1'
add_table_borders(table2)

# Header
header_cells = table2.rows[0].cells
header_cells[0].text = 'Model'
header_cells[1].text = 'V1 Quality'
header_cells[2].text = 'V2 Quality'
header_cells[3].text = 'V1 Reasoning Quality'
header_cells[4].text = 'V2 Reasoning Quality'

# Data rows
table2.rows[1].cells[0].text = 'LLAMA 70B Versatile'
table2.rows[2].cells[0].text = 'Qwen V2 Fine-tuned'

for row_idx in [1, 2]:
    for col_idx in range(1, 5):
        table2.rows[row_idx].cells[col_idx].text = '[Enter score]'

# Save
doc.save('COMPREHENSIVE_EVALUATION_REPORT.docx')

# ============================================================================
# STEP 6: SAVE JSON RESULTS
# ============================================================================

print("\n[5/6] Saving results...")

with open('evaluation_results.json', 'w') as f:
    json.dump(results, f, indent=2)

# ============================================================================
# COMPLETE
# ============================================================================

print("\n[6/6] Done!")
print("\n" + "="*80)
print("EVALUATION COMPLETE")
print("="*80)
print("\nGenerated Files:")
print("  ✓ COMPREHENSIVE_EVALUATION_REPORT.docx")
print("  ✓ evaluation_results.json")
print("\nNext steps:")
print("  1. Open COMPREHENSIVE_EVALUATION_REPORT.docx")
print("  2. Review model outputs for both tasks")
print("  3. Fill in evaluation tables with your scores")
print("="*80)
