#!/usr/bin/env python3
"""
Evaluate merged model on CHIVA classification using real test data
"""

import os
import sys
import json
import torch
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("="*80)
print("EVALUATING MERGED MODEL ON CHIVA CLASSIFICATION")
print("="*80)

# Configuration
HF_TOKEN = ""
MODEL_REPO = "HariKrishna1824/qwen_chiva_vericose_veins_treatment_finetuned"
TEST_CASES_DIR = Path("json samples")

CHIVA_RULES = """
=== CHIVA CLASSIFICATION RULES ===

ANATOMY:
    N1 = Deep venous system (femoral / popliteal vein)
    N2 = Great Saphenous Vein (GSV) or Small Saphenous Vein (SSV) trunk
    N3 = Tributaries / superficial branches
    EP = Physiological (forward, antegrade) flow
    RP = Retrograde (pathological, reflux) flow
    SFJ = Saphenofemoral Junction

CLASSIFICATION RULES:
    STEP 1 — CHECK FOR EP N1→N2:
        YES → SFJ/Hunterian INCOMPETENT (Case A or B)
        NO  → SFJ COMPETENT (Case C)

    Case A — EP N1→N2 EXISTS, NO EP N2→N3:
        If RP N2→N1, no RP at N3 → TYPE 1

    Case B — EP N1→N2 EXISTS AND EP N2→N3 EXISTS:
        B1: RP N3→N2 or RP N3→N1, NO RP N2→N1 → TYPE 3
        B4: RP N3→N1 AND RP N2→N1, elim="Reflux" → TYPE 1+2
        B5: RP N3→N1 AND RP N2→N1, elim="No Reflux" → TYPE 3

    Case C — NO EP N1→N2 (SFJ COMPETENT):
        TYPE 2A: EP N2→N3, NO EP N1→N2
        TYPE 2B: EP N2→N2, NO EP N1→N2, RP at N3, NO RP N2→N1
        TYPE 2C: EP N2→N2, NO EP N1→N2, RP at N3, RP N2→N1 present
"""

RAG_DOCS = [
    "For TYPE 1 shunt: Ligate GSV at saphenofemoral junction, preserve distal GSV if quality permits",
    "For TYPE 2A shunt: Selective perforator ligation with duplex guidance",
    "For TYPE 2B shunt: Conservative therapy with compression; consider staged intervention",
    "For TYPE 2C shunt: Evaluate SFJ status; may require combined intervention",
    "For TYPE 3 shunt: Ligate incompetent tributaries at their junction with GSV/SSV",
    "For TYPE 1+2 shunt: Combined approach - SFJ ligation plus selective perforator ligation",
]

def format_clips_v1(clips: List[Dict]) -> str:
    """V1: Natural language with numerical description"""
    descriptions = []
    for clip in clips:
        flow = "physiological forward flow" if clip.get('flow') == 'EP' else "retrograde reflux"
        from_to = f"from {clip.get('fromType')} to {clip.get('toType')}"
        y_val = clip.get('posYRatio', 0)
        descriptions.append(f"There is a {flow} (y={y_val:.3f}) {from_to}")
    return "Duplex findings:\n" + "\n".join(descriptions)

def format_clips_v2(clips: List[Dict]) -> str:
    """V2: Full medical terminology"""
    findings = []

    for clip in clips:
        flow = clip.get('flow')
        from_type = clip.get('fromType')
        to_type = clip.get('toType')
        y_val = clip.get('posYRatio', 0)

        if flow == 'EP' and from_type == 'N1' and to_type == 'N2':
            findings.append("saphenofemoral junction incompetence with antegrade flow from deep venous system")
        elif flow == 'EP' and from_type == 'N2' and to_type == 'N3':
            findings.append("perforator draining from saphenous trunk into tributaries")
        elif flow == 'RP' and from_type == 'N2' and to_type == 'N1':
            findings.append("retrograde reflux within saphenous trunk to deep system")
        elif flow == 'RP' and from_type == 'N3' and to_type in ['N2', 'N1']:
            findings.append("retrograde flow from tributaries back toward proximal system")
        elif flow == 'EP' and from_type == 'N2':
            findings.append("antegrade flow within saphenous trunk")

    if not findings:
        findings.append("venous incompetence pattern identified on duplex")

    return "Patient with varicose veins and venous insufficiency. Duplex demonstrates: " + "; ".join(findings) + "."

print("\n[1/4] Loading test cases...")

test_cases = []
if TEST_CASES_DIR.exists():
    for json_file in sorted(TEST_CASES_DIR.glob("*.json")):
        try:
            with open(json_file, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
                clips = data.get('clips', [])
                if clips:
                    test_cases.append({
                        'name': json_file.stem,
                        'clips': clips,
                    })
        except Exception as e:
            print(f"  Warning: Could not load {json_file.name}: {str(e)[:50]}")

print(f"  Loaded {len(test_cases)} test cases")

# Load model
print("\n[2/4] Loading merged model...")
print(f"  Model: {MODEL_REPO}")
print(f"  Device: {torch.cuda.get_device_name(0) if torch.cuda.is_available() else 'CPU'}")

try:
    from transformers import AutoModelForCausalLM, AutoTokenizer

    print("  Loading base model and weights...", end=" ", flush=True)
    model = AutoModelForCausalLM.from_pretrained(
        MODEL_REPO,
        torch_dtype=torch.bfloat16 if torch.cuda.is_available() else torch.float32,
        device_map="auto" if torch.cuda.is_available() else "cpu",
        token=HF_TOKEN,
        trust_remote_code=True
    )
    print("OK")

    print("  Loading tokenizer...", end=" ", flush=True)
    tokenizer = AutoTokenizer.from_pretrained(MODEL_REPO, token=HF_TOKEN, trust_remote_code=True)
    print("OK")

    model.eval()
    device = next(model.parameters()).device
    print(f"  Loaded successfully")

except Exception as e:
    print(f"\n  ERROR: {str(e)}")
    sys.exit(1)

# Inference function
def infer(prompt: str, max_tokens: int = 500) -> str:
    try:
        messages = [{"role": "user", "content": f"{CHIVA_RULES}\n\n{prompt}"}]
        text = tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
        inputs = tokenizer(text, return_tensors="pt").to(device)

        with torch.no_grad():
            outputs = model.generate(
                **inputs,
                max_new_tokens=max_tokens,
                do_sample=False,
                pad_token_id=tokenizer.eos_token_id
            )

        return tokenizer.decode(outputs[0][inputs['input_ids'].shape[1]:], skip_special_tokens=True)
    except Exception as e:
        return f"[Inference Error: {str(e)[:100]}]"

# Run evaluation
print("\n[3/4] Running evaluation...")

results = {
    'timestamp': datetime.now().isoformat(),
    'model': MODEL_REPO,
    'test_count': len(test_cases),
    'evaluations': []
}

for idx, test in enumerate(test_cases, 1):
    print(f"  [{idx}/{len(test_cases)}] {test['name']}", end=" ", flush=True)

    clips = test['clips']
    v1_prompt = format_clips_v1(clips)
    v2_prompt = format_clips_v2(clips)

    # Classification
    print("(classify)", end=" ", flush=True)
    c_v1 = infer(v1_prompt + "\n\nClassify the CHIVA shunt type.")
    c_v2 = infer(v2_prompt + "\n\nWhat is the CHIVA shunt type?")

    # Ligation
    print("(ligate)", end=" ", flush=True)
    rag_context = "\n".join(RAG_DOCS)
    l_v1 = infer(f"{rag_context}\n\n{v1_prompt}\n\nProvide ligation planning strategy.")
    l_v2 = infer(f"{rag_context}\n\n{v2_prompt}\n\nProvide ligation planning strategy.")

    results['evaluations'].append({
        'test_name': test['name'],
        'clips': clips,
        'v1_classification': c_v1,
        'v2_classification': c_v2,
        'v1_ligation': l_v1,
        'v2_ligation': l_v2,
    })

    print("OK")

# Generate report
print("\n[4/4] Generating Word report...")

doc = Document()

# Title
title = doc.add_heading('CHIVA SHUNT CLASSIFICATION EVALUATION', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph(f'Model: {MODEL_REPO}')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)

# Metadata
doc.add_heading('Evaluation Metadata', level=1)
doc.add_paragraph(f"Timestamp: {results['timestamp']}")
doc.add_paragraph(f"Test Cases: {results['test_count']}")
doc.add_paragraph(f"Model: Merged Qwen2.5-7B + CHIVA LoRA")
doc.add_paragraph("Query Formats: V1 (Natural Language Numeric), V2 (Medical Terminology)")
doc.add_paragraph("Tasks: Task 1 (Shunt Classification - No RAG), Task 2 (Ligation Planning - With RAG)")

# Results
doc.add_page_break()
doc.add_heading('TASK 1: SHUNT CLASSIFICATION (NO RAG)', level=1)

for result in results['evaluations']:
    doc.add_heading(f"Test Case: {result['test_name']}", level=2)

    doc.add_heading("V1 Query (Natural Language with Y-values)", level=3)
    clips_desc = format_clips_v1(result['clips'])
    doc.add_paragraph(clips_desc)
    doc.add_paragraph("Classification Output:")
    doc.add_paragraph(result['v1_classification'], style='List Bullet')

    doc.add_heading("V2 Query (Full Medical Terminology)", level=3)
    clips_med = format_clips_v2(result['clips'])
    doc.add_paragraph(clips_med)
    doc.add_paragraph("Classification Output:")
    doc.add_paragraph(result['v2_classification'], style='List Bullet')

# Ligation planning
doc.add_page_break()
doc.add_heading('TASK 2: LIGATION PLANNING (WITH RAG)', level=1)

for result in results['evaluations']:
    doc.add_heading(f"Test Case: {result['test_name']}", level=2)

    doc.add_heading("V1 Query with RAG Context", level=3)
    doc.add_paragraph("Ligation Planning Output:")
    doc.add_paragraph(result['v1_ligation'], style='List Bullet')

    doc.add_heading("V2 Query with RAG Context", level=3)
    doc.add_paragraph("Ligation Planning Output:")
    doc.add_paragraph(result['v2_ligation'], style='List Bullet')

# Save
doc.save('evaluation_results_merged_model.docx')

# Save JSON
with open('evaluation_results_merged_model.json', 'w') as f:
    json.dump(results, f, indent=2)

print("\n" + "="*80)
print("EVALUATION COMPLETE")
print("="*80)
print("\nFiles saved:")
print("  - evaluation_results_merged_model.docx")
print("  - evaluation_results_merged_model.json")
print(f"\nTest cases evaluated: {len(test_cases)}")
print(f"Timestamp: {results['timestamp']}")
