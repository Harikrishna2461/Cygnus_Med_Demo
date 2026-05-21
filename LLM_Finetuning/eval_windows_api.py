#!/usr/bin/env python3
"""
WINDOWS ONLY - Evaluate merged model via HuggingFace Inference API
No GPU required - works on Windows with just requests library
"""

import requests
import json
import time
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("="*80)
print("CHIVA EVALUATION - WINDOWS (HF API)")
print("="*80)

# Configuration
HF_TOKEN = ""
MODEL_REPO = "HariKrishna1824/qwen_chiva_vericose_veins_treatment_finetuned"
HF_API_URL = f"https://api-inference.huggingface.co/models/{MODEL_REPO}"
TEST_CASES_DIR = Path("json samples")

CHIVA_RULES = """
CHIVA_RULES :
=== CHIVA VENOUS SHUNT CLASSIFICATION RULES ===

ANATOMY:
    N1 = Deep venous system (femoral / popliteal vein)
    N2 = Great Saphenous Vein (GSV) or Small Saphenous Vein (SSV) trunk
    N3 = Tributaries / superficial branches
    EP = Physiological (forward, antegrade) flow — NORMAL clip
    RP = Retrograde (pathological, reflux) flow — ABNORMAL clip
    SFJ = Saphenofemoral Junction  →  posYRatio ≤ 0.098
    Hunterian Perforator            →  0.098 < posYRatio ≤ 0.353

═══════════════════════════════════════════════════════════
CRITICAL RULE — SFJ COMPETENCE (read before classifying):
    SFJ is INCOMPETENT if and only if a clip has fromType=N1 AND toType=N2 (EP N1→N2).
    EP N2→N2 means blood circulates within the saphenous trunk via a perforator — SFJ REMAINS COMPETENT.
    This is true regardless of posYRatio or step label. Even posYRatio=0.05 with step=SFJ-Knee
    is a perforator entry if the clip reads EP N2→N2, NOT EP N1→N2.
═══════════════════════════════════════════════════════════

STEP 1 — CHECK FOR EP N1→N2:
    Scan ALL clips. Does any clip have flow=EP, fromType=N1, toType=N2?
    YES → SFJ/Hunterian INCOMPETENT → go to Case A or B.
    NO  → SFJ COMPETENT → go to Case C.

─────────────────────────────────────────────────────────
Case A — EP N1→N2 EXISTS (SFJ or Hunterian), NO EP N2→N3
─────────────────────────────────────────────────────────
    If RP N2→N1 present AND no RP at N3 (no RP N3→N2, no RP N3→N1) → TYPE 1
    Ligation: Ligate at SFJ (y≤0.098) or Hunterian (y≤0.353).
            If multiple RP N2→N1: ligate below each except the most distal.

─────────────────────────────────────────────────────────
Case B — EP N1→N2 EXISTS (SFJ or Hunterian) AND EP N2→N3 EXISTS
─────────────────────────────────────────────────────────
    B1: RP N3→N2 or RP N3→N1, NO RP N2→N1               → TYPE 3
    B2: RP N3→N2 AND RP N2→N1                             → TYPE 3
    B3: RP N3→N1 AND RP N2→N1, eliminationTest absent    → UNDETERMINED (set needs_elim_test=true)
    B4: RP N3→N1 AND RP N2→N1, eliminationTest="Reflux"  → TYPE 1+2
    B5: RP N3→N1 AND RP N2→N1, eliminationTest="No Reflux" → TYPE 3

    TYPE 3 Ligation:
        Single RP at N3: Ligate EP at N2→N3. Follow up 6–12 months; if N2 reflux develops, ligate SFJ.
        Multiple RP at N3: Ligate every refluxing tributary at N2 junction (CHIVA 2 step 1). Same follow-up.

    TYPE 1+2 Ligation — depends on RP N2→N1 calibre (set ask_diameter=true):
        Small RP N2→N1: Apply CHIVA 2 (ligate EP N2→N3 first, then SFJ/Hunterian).
                        OR ligate SFJ first + all tributaries except one; once N2 normalises ligate last.
        Large / multiple RP N2→N1: Ligate SFJ/Hunterian + every refluxing tributary simultaneously.
                                    Ligate below each RP N2→N1 except the most distal.

─────────────────────────────────────────────────────────
Case C — NO EP N1→N2 ANYWHERE (SFJ COMPETENT)
─────────────────────────────────────────────────────────
    C-Sub-check: what type of EP clip exists?

    ── TYPE 2A ── EP N2→N3 present, NO EP N1→N2
        The defining feature is EP N2→N3 (GSV feeding a tributary) without any SFJ entry.
        RP may or may not be present in early/developing cases.
        Typical pattern: EP N2→N3 + RP N3→N2 or N3→N1. No RP N2→N1.
        Key signal: EP N2→N3 clip exists + NO EP N1→N2 clip exists anywhere.
        If multiple RP at N3 → set ask_branching=true (need calibre/distance/drainage info).
        Ligation: Ligate highest EP at N2→N3 junction.
                    If multiple branching at N3: ligate based on calibre, distance to perforator, drainage.

    ── TYPE 2B ── EP N2→N2 present, NO EP N1→N2, RP at N3, NO RP N2→N1
        Entry is via perforator (fromType=N2, toType=N2 — NOT N1→N2).
        IMPORTANT: EP N2→N2 at ANY posYRatio (even 0.05, SFJ-Knee step) = perforator, NOT SFJ.
        Key signal: EP N2→N2 clip + RP N3→N2 or N3→N1 + NO EP N1→N2 + NO RP N2→N1.
        If multiple RP at N3 → set ask_branching=true.
        Ligation: Ligate the highest EP N2→N2 (perforator entry point).

    ── TYPE 2C ── EP N2→N2 present, NO EP N1→N2, RP at N3, RP N2→N1 ALSO present
        Perforator entry (EP N2→N2) with secondary GSV reflux (RP N2→N1). SFJ still competent.
        IMPORTANT: 2C has EP N2→N2 (perforator), while Type 1+2 has EP N1→N2 (SFJ entry).
        If NO EP N1→N2 but RP N2→N1 exists with EP N2→N2 → TYPE 2C, not Type 1+2.
        Key signal: EP N2→N2 + RP N3 + RP N2→N1 + NO EP N1→N2.
        Ligation: Ligate perforator entry (highest EP N2→N2) AND all RP N2→N1 sites along GSV.

    Case C — NO SHUNT:
        If EP N2→N2 exists but NO RP clips of any kind → NO SHUNT DETECTED.

─────────────────────────────────────────────────────────
Case D — No RP in any clip → NO SHUNT DETECTED. No ligation needed.
─────────────────────────────────────────────────────────

QUICK DECISION TABLE:
    Has EP N1→N2? YES + no EP N2→N3 + RP N2→N1           → TYPE 1
    Has EP N1→N2? YES + EP N2→N3 + RP N3 only             → TYPE 3
    Has EP N1→N2? YES + EP N2→N3 + RP N3 + RP N2→N1 + eliminationTest absent → UNDETERMINED
    Has EP N1→N2? YES + EP N2→N3 + RP N3 + RP N2→N1 + elim="Reflux"          → TYPE 1+2
    Has EP N1→N2? YES + EP N2→N3 + RP N3 + RP N2→N1 + elim="No Reflux"       → TYPE 3
    No EP N1→N2  + EP N2→N3                                → TYPE 2A
    No EP N1→N2  + EP N2→N2 + RP N3 + NO RP N2→N1         → TYPE 2B
    No EP N1→N2  + EP N2→N2 + RP N3 + RP N2→N1            → TYPE 2C
    No EP N1→N2  + EP N2→N2 + NO RP                        → NO SHUNT
    No RP at all                                            → NO SHUNT

CONCRETE EXAMPLES (match these patterns exactly):
    Type 1:  [EP N1→N2 y=0.06 SFJ-ENTRY, RP N2→N1 y=0.25]
            → EP N1→N2 present, RP N2→N1, no EP N2→N3, no N3 reflux → TYPE 1
    Type 2A: [EP N2→N3 y=0.20]  OR  [EP N2→N3 y=0.20, RP N3→N2 y=0.47]
            → No EP N1→N2, EP N2→N3 present → TYPE 2A
    Type 2B: [EP N2→N2 y=0.050 step=SFJ-Knee ligation-point-marker, RP N3→N1 y=0.132]
            → No EP N1→N2, EP N2→N2 = perforator, RP N3 only → TYPE 2B
    Type 2C: [EP N2→N2 y=0.050 step=SFJ-Knee ligation-point-marker, RP N3→N1 y=0.132, RP N2→N1 y=0.212]
            → No EP N1→N2, EP N2→N2 = perforator, RP N3 + RP N2→N1 → TYPE 2C
    Type 3:  [EP N1→N2 y=0.05 SFJ-ENTRY, EP N2→N3 y=0.132 ligation-point-marker, RP N3→N1 y=0.212]
            → EP N1→N2 + EP N2→N3 + RP N3→N1, no RP N2→N1 → TYPE 3
    Type 3 variant 2 (no elim test):
            [EP N1→N2, EP N2→N3, RP N3→N1, RP N2→N1, no eliminationTest] → UNDETERMINED
    Type 1+2:[EP N1→N2, EP N2→N3 eliminationTest="Reflux", RP N3→N1, RP N2→N1] → TYPE 1+2
    No shunt:[EP N1→N2 only, no RP]  OR  [EP N2→N2 only, no RP] → NO SHUNT

TYPE 2 BRANCHING — ask_branching flag:
    Set ask_branching=true when there are MULTIPLE RP at N3 tributaries in a Type 2A, 2B, or 2C case.
    The ligation choice among multiple N3 branches depends on:
        • Calibre of branches (equal or unequal)
        • Distance of each branch to its perforator
        • Whether drainage through the thinner vessel is possible
    If unequal calibre with drainage possible → ligate the larger vessel.
    If unequal calibre, no drainage → ligate the smaller vessel.
    If equal calibre, unequal distance → ligate the branch with longer distance to perforator.
"""

RAG_DOCS = [
    "For TYPE 1 shunt: Ligate GSV at saphenofemoral junction, preserve distal GSV if quality permits",
    "For TYPE 2A shunt: Selective perforator ligation with duplex guidance",
    "For TYPE 2B shunt: Conservative therapy with compression; consider staged intervention",
    "For TYPE 2C shunt: Evaluate SFJ status; may require combined intervention",
    "For TYPE 3 shunt: Ligate incompetent tributaries at their junction with GSV/SSV",
    "For TYPE 1+2 shunt: Combined approach - SFJ ligation plus selective perforator ligation",
]

def format_clips_v1(clips: List[Dict]) -> str:
    """V1: Natural language with numerical description"""
    descriptions = []
    for clip in clips:
        flow = "physiological forward flow" if clip.get('flow') == 'EP' else "retrograde reflux"
        from_to = f"from {clip.get('fromType')} to {clip.get('toType')}"
        y_val = clip.get('posYRatio', 0)
        descriptions.append(f"There is a {flow} (y={y_val:.3f}) {from_to}")
    return "Duplex findings:\n" + "\n".join(descriptions)

def format_clips_v2(clips: List[Dict]) -> str:
    """V2: Full medical terminology"""
    findings = []
    for clip in clips:
        flow = clip.get('flow')
        from_type = clip.get('fromType')
        to_type = clip.get('toType')

        if flow == 'EP' and from_type == 'N1' and to_type == 'N2':
            findings.append("saphenofemoral junction incompetence with antegrade flow from deep venous system")
        elif flow == 'EP' and from_type == 'N2' and to_type == 'N3':
            findings.append("perforator draining from saphenous trunk into tributaries")
        elif flow == 'RP' and from_type == 'N2' and to_type == 'N1':
            findings.append("retrograde reflux within saphenous trunk to deep system")
        elif flow == 'RP' and from_type == 'N3' and to_type in ['N2', 'N1']:
            findings.append("retrograde flow from tributaries back toward proximal system")
        elif flow == 'EP' and from_type == 'N2':
            findings.append("antegrade flow within saphenous trunk")

    if not findings:
        findings.append("venous incompetence pattern identified")

    return "Patient with varicose veins. Duplex shows: " + ", ".join(findings) + "."

def query_api(prompt: str, max_tokens: int = 500) -> str:
    """Query model via HF Inference API"""
    headers = {"Authorization": f"Bearer {HF_TOKEN}"}
    payload = {
        "inputs": f"{CHIVA_RULES}\n\n{prompt}",
        "parameters": {"max_new_tokens": max_tokens}
    }

    try:
        response = requests.post(HF_API_URL, headers=headers, json=payload, timeout=120)

        if response.status_code == 200:
            result = response.json()
            if isinstance(result, list) and len(result) > 0:
                output = result[0].get('generated_text', '[No output]')
                if prompt in output:
                    output = output.replace(prompt, "").strip()
                return output
            return "[Empty response]"
        elif response.status_code == 503:
            return "[Model loading, please wait...]"
        else:
            return f"[HTTP {response.status_code}]"
    except requests.exceptions.Timeout:
        return "[Timeout - model may be loading]"
    except Exception as e:
        return f"[Error: {str(e)[:100]}]"

# Load test cases
print("\n[1/4] Loading test cases...")
test_cases = []
if TEST_CASES_DIR.exists():
    for json_file in sorted(TEST_CASES_DIR.glob("*.json")):
        try:
            with open(json_file, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
                clips = data.get('clips', [])
                if clips:
                    test_cases.append({
                        'name': json_file.stem,
                        'clips': clips,
                    })
        except Exception as e:
            print(f"  Warning: {json_file.name}: {str(e)[:50]}")

print(f"  Loaded {len(test_cases)} test cases")

# Test API connection
print("\n[2/4] Testing HF API connection...")
test_response = query_api("What is TYPE 1 shunt?")
if "[" in test_response and "]" in test_response:
    print(f"  API Status: {test_response}")
    print("  Waiting 30 seconds for model to fully load...")
    time.sleep(30)
else:
    print(f"  API Ready: OK")

# Run evaluation
print("\n[3/4] Running evaluation...")

results = {
    'timestamp': datetime.now().isoformat(),
    'model': MODEL_REPO,
    'test_count': len(test_cases),
    'evaluations': []
}

for idx, test in enumerate(test_cases, 1):
    print(f"  [{idx}/{len(test_cases)}] {test['name']}", end=" ", flush=True)

    clips = test['clips']
    v1_prompt = format_clips_v1(clips)
    v2_prompt = format_clips_v2(clips)

    # Classification
    print("(c)", end=" ", flush=True)
    c_v1 = query_api(v1_prompt + "\n\nClassify the CHIVA shunt type.")
    time.sleep(2)
    c_v2 = query_api(v2_prompt + "\n\nWhat is the CHIVA shunt type?")

    # Ligation
    print("(l)", end=" ", flush=True)
    rag_context = "\n".join(RAG_DOCS)
    l_v1 = query_api(f"{rag_context}\n\n{v1_prompt}\n\nProvide ligation planning strategy.")
    time.sleep(2)
    l_v2 = query_api(f"{rag_context}\n\n{v2_prompt}\n\nProvide ligation planning strategy.")

    results['evaluations'].append({
        'test_name': test['name'],
        'v1_classification': c_v1,
        'v2_classification': c_v2,
        'v1_ligation': l_v1,
        'v2_ligation': l_v2,
    })

    print("OK")

# Generate report
print("\n[4/4] Generating report...")

doc = Document()

title = doc.add_heading('CHIVA EVALUATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph(f'Model: {MODEL_REPO}')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)

doc.add_heading('Metadata', level=1)
doc.add_paragraph(f"Timestamp: {results['timestamp']}")
doc.add_paragraph(f"Test Cases: {results['test_count']}")
doc.add_paragraph("Query Formats: V1 (Natural Language), V2 (Medical Terminology)")

doc.add_page_break()
doc.add_heading('TASK 1: SHUNT CLASSIFICATION', level=1)

for result in results['evaluations']:
    doc.add_heading(f"Test: {result['test_name']}", level=2)

    doc.add_heading("V1 Query", level=3)
    doc.add_paragraph(result['v1_classification'])

    doc.add_heading("V2 Query", level=3)
    doc.add_paragraph(result['v2_classification'])

doc.add_page_break()
doc.add_heading('TASK 2: LIGATION PLANNING (WITH RAG)', level=1)

for result in results['evaluations']:
    doc.add_heading(f"Test: {result['test_name']}", level=2)

    doc.add_heading("V1 with RAG", level=3)
    doc.add_paragraph(result['v1_ligation'])

    doc.add_heading("V2 with RAG", level=3)
    doc.add_paragraph(result['v2_ligation'])

doc.save('EVALUATION_REPORT.docx')

with open('evaluation_results.json', 'w') as f:
    json.dump(results, f, indent=2)

print("\n" + "="*80)
print("COMPLETE")
print("="*80)
print("\nFiles saved:")
print("  - EVALUATION_REPORT.docx")
print("  - evaluation_results.json")
print(f"\nTest cases: {len(test_cases)}")
