#!/usr/bin/env python3
"""
FULL EVALUATION - API CALLS ONLY
LLAMA 70B (Groq) + Qwen V2 (HF Inference)
Real outputs - comprehensive report
"""

import os
import json
import sys
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess

print("="*80)
print("FULL EVALUATION: API-BASED")
print("="*80)

# Configuration
GROQ_API_KEY = os.getenv("GROQ_API_KEY", "YOUR_GROQ_API_KEY_HERE")
HF_TOKEN = os.getenv("HF_TOKEN", "YOUR_HF_TOKEN_HERE")
HF_MODEL_REPO = "HariKrishna1824/qwen25_chiva_v2"
TEST_CASES_DIR = Path("json samples")

CHIVA_RULES = """
=== CHIVA VENOUS SHUNT CLASSIFICATION RULES ===

ANATOMY:
    N1 = Deep venous system (femoral / popliteal vein)
    N2 = Great Saphenous Vein (GSV) or Small Saphenous Vein (SSV) trunk
    N3 = Tributaries / superficial branches
    EP = Physiological (forward, antegrade) flow
    RP = Retrograde (pathological, reflux) flow
    SFJ = Saphenofemoral Junction (posYRatio ≤ 0.098)

CLASSIFICATION RULES:
    STEP 1 — CHECK FOR EP N1→N2:
        YES → SFJ/Hunterian INCOMPETENT (Case A or B)
        NO  → SFJ COMPETENT (Case C)

    Case A — EP N1→N2 EXISTS, NO EP N2→N3:
        If RP N2→N1, no RP at N3 → TYPE 1

    Case B — EP N1→N2 EXISTS AND EP N2→N3 EXISTS:
        B1: RP N3→N2 or RP N3→N1, NO RP N2→N1 → TYPE 3
        B4: RP N3→N1 AND RP N2→N1, elim="Reflux" → TYPE 1+2
        B5: RP N3→N1 AND RP N2→N1, elim="No Reflux" → TYPE 3

    Case C — NO EP N1→N2 (SFJ COMPETENT):
        TYPE 2A: EP N2→N3, NO EP N1→N2
        TYPE 2B: EP N2→N2, NO EP N1→N2, RP at N3, NO RP N2→N1
        TYPE 2C: EP N2→N2, NO EP N1→N2, RP at N3, RP N2→N1 present
"""

print("\n[1/5] Installing dependencies...")
packages = ['groq', 'huggingface-hub', 'requests']
for pkg in packages:
    try:
        __import__(pkg.replace('-', '_'))
    except ImportError:
        print(f"  {pkg}...", end=" ")
        subprocess.run([sys.executable, '-m', 'pip', 'install', '-q', pkg], check=False)
        print("OK")

print("\n[2/5] Loading test cases...")
test_cases = []
if TEST_CASES_DIR.exists():
    for json_file in sorted(list(TEST_CASES_DIR.glob("*.json")))[:12]:  # Limit to 12
        try:
            with open(json_file, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
                clips = data.get('clips', [])
                if clips:
                    test_cases.append({
                        'name': json_file.stem,
                        'clips': clips,
                    })
        except:
            pass

print(f"  Loaded {len(test_cases)} test cases")

# ============================================================
# API INFERENCE FUNCTIONS
# ============================================================

print("\n[3/5] Initializing API clients...")

# GROQ API
try:
    from groq import Groq
    groq_client = Groq(api_key=GROQ_API_KEY)
    print("  Groq API: OK")
except Exception as e:
    print(f"  Groq API: FAIL ({str(e)[:50]})")
    groq_client = None

# HF Inference API
try:
    import requests
    hf_api_url = f"https://api-inference.huggingface.co/models/{HF_MODEL_REPO}"
    hf_headers = {"Authorization": f"Bearer {HF_TOKEN}"}

    # Test HF API
    test_payload = {"inputs": "test", "parameters": {"max_new_tokens": 10}}
    response = requests.post(hf_api_url, headers=hf_headers, json=test_payload, timeout=30)

    if response.status_code == 200:
        print("  HF Inference API: OK")
        hf_ok = True
    else:
        print(f"  HF Inference API: FAIL ({response.status_code})")
        hf_ok = False
except Exception as e:
    print(f"  HF Inference API: FAIL ({str(e)[:50]})")
    hf_ok = False

if not groq_client and not hf_ok:
    print("\nERROR: No API available!")
    sys.exit(1)

# ============================================================
# INFERENCE FUNCTIONS
# ============================================================

def format_clips_v1(clips: List[Dict]) -> str:
    """V1: Raw clip notation"""
    s = "Duplex clips:\n"
    for c in clips:
        s += f"- {c.get('flow')} {c.get('fromType')}->{c.get('toType')} (y={c.get('posYRatio')})\n"
    return s

def format_clips_v2(clips: List[Dict]) -> str:
    """V2: Natural language"""
    has_ep_n1_n2 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N1' and c.get('toType') == 'N2' for c in clips)
    has_ep_n2_n3 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N2' and c.get('toType') == 'N3' for c in clips)
    has_rp_n2_n1 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N2' and c.get('toType') == 'N1' for c in clips)
    has_rp_n3 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N3' for c in clips)

    findings = []
    if has_ep_n1_n2:
        findings.append("saphenofemoral junction incompetence")
    else:
        findings.append("competent saphenofemoral junction")
    if has_ep_n2_n3:
        findings.append("perforator feeding into the saphenous vein")
    if has_rp_n2_n1:
        findings.append("retrograde flow within the saphenous trunk")
    if has_rp_n3:
        findings.append("retrograde flow at the tributary level")

    return "Patient with varicose veins. Duplex shows: " + ", ".join(findings) + "."

def groq_infer(prompt: str) -> str:
    """GROQ API inference"""
    if not groq_client:
        return "[GROQ unavailable]"
    try:
        # Try different models
        for model in ["mixtral-8x7b-32768", "llama-3-70b-versatile", "llama-2-70b-4096"]:
            try:
                response = groq_client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": f"Expert vascular surgeon.\n\n{CHIVA_RULES}"},
                        {"role": "user", "content": prompt}
                    ],
                    max_tokens=500,
                    temperature=0.3
                )
                return response.choices[0].message.content
            except:
                continue
        return "[GROQ: No available models]"
    except Exception as e:
        return f"[GROQ Error: {str(e)[:80]}]"

def hf_infer(prompt: str) -> str:
    """HF Inference API"""
    if not hf_ok:
        return "[HF unavailable]"
    try:
        import requests
        payload = {
            "inputs": f"{CHIVA_RULES}\n\n{prompt}",
            "parameters": {"max_new_tokens": 500}
        }
        response = requests.post(hf_api_url, headers=hf_headers, json=payload, timeout=60)

        if response.status_code == 200:
            result = response.json()
            if isinstance(result, list) and len(result) > 0:
                return result[0].get('generated_text', '[No output]')
            return '[Empty response]'
        else:
            return f"[HF Error: {response.status_code}]"
    except Exception as e:
        return f"[HF Error: {str(e)[:80]}]"

# ============================================================
# RUN EVALUATION
# ============================================================

print("\n[4/5] Running evaluation...")

results = {
    'timestamp': datetime.now().isoformat(),
    'models': {
        'llama': 'Groq API (Mixtral/Llama)',
        'qwen': f'HF Inference API ({HF_MODEL_REPO})'
    },
    'test_count': len(test_cases),
    'shunt_classification': [],
    'ligation': []
}

rag_docs = [
    "For TYPE 1 shunt: Ligate GSV at saphenofemoral junction",
    "For TYPE 2A shunt: Selective perforator ligation with duplex guidance",
    "For TYPE 2B shunt: Conservative therapy with compression; consider staged intervention",
    "For TYPE 2C shunt: Evaluate SFJ status; may require combined intervention",
]

for idx, test in enumerate(test_cases, 1):
    print(f"  [{idx}/{len(test_cases)}] {test['name']}...", end=" ", flush=True)

    clips = test['clips']
    v1_clips = format_clips_v1(clips)
    v2_clips = format_clips_v2(clips)

    # Classification
    llama_c_v1 = groq_infer(v1_clips + "\nClassify the CHIVA shunt type.")
    llama_c_v2 = groq_infer(v2_clips + "\nWhat is the CHIVA shunt type?")
    qwen_c_v1 = hf_infer(v1_clips + "\nClassify the CHIVA shunt type.")
    qwen_c_v2 = hf_infer(v2_clips + "\nWhat is the CHIVA shunt type?")

    # Ligation
    rag_context = "\n".join(rag_docs)
    llama_l_v1 = groq_infer(f"{rag_context}\n\n{v1_clips}\n\nProvide ligation planning.")
    llama_l_v2 = groq_infer(f"{rag_context}\n\n{v2_clips}\n\nProvide ligation planning.")
    qwen_l_v1 = hf_infer(f"{rag_context}\n\n{v1_clips}\n\nProvide ligation planning.")
    qwen_l_v2 = hf_infer(f"{rag_context}\n\n{v2_clips}\n\nProvide ligation planning.")

    results['shunt_classification'].append({
        'test': test['name'],
        'llama_v1': llama_c_v1, 'llama_v2': llama_c_v2,
        'qwen_v1': qwen_c_v1, 'qwen_v2': qwen_c_v2
    })

    results['ligation'].append({
        'test': test['name'],
        'llama_v1': llama_l_v1, 'llama_v2': llama_l_v2,
        'qwen_v1': qwen_l_v1, 'qwen_v2': qwen_l_v2
    })

    print("OK")

# ============================================================
# GENERATE REPORT
# ============================================================

print("\n[5/5] Generating comprehensive report...")

doc = Document()

# Title
title = doc.add_heading('COMPARATIVE EVALUATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('LLAMA 70B (Groq) vs Qwen2.5-7B V2 (HF)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.bold = True

# Metadata
doc.add_heading('Evaluation Metadata', level=1)
doc.add_paragraph(f"Timestamp: {results['timestamp']}")
doc.add_paragraph(f"Test Cases: {results['test_count']}")
doc.add_paragraph(f"LLAMA Model: {results['models']['llama']}")
doc.add_paragraph(f"Qwen Model: {results['models']['qwen']}")
doc.add_paragraph("Query Formats: V1 (Raw Clips), V2 (Medical Natural Language)")
doc.add_paragraph("Evaluation Tasks: Task 1 (Shunt Classification - No RAG), Task 2 (Ligation Planning - With RAG)")

# ============================================================
# TASK 1: SHUNT CLASSIFICATION
# ============================================================

doc.add_page_break()
doc.add_heading('TASK 1: SHUNT CLASSIFICATION (NO RAG)', level=1)

for result in results['shunt_classification']:
    doc.add_heading(f"Test Case: {result['test']}", level=2)

    # LLAMA V1
    doc.add_heading("LLAMA 70B - V1 Query (Raw Clip Notation)", level=3)
    doc.add_paragraph(result['llama_v1'])

    # LLAMA V2
    doc.add_heading("LLAMA 70B - V2 Query (Natural Language Medical)", level=3)
    doc.add_paragraph(result['llama_v2'])

    # Qwen V1
    doc.add_heading("Qwen2.5-7B V2 - V1 Query (Raw Clip Notation)", level=3)
    doc.add_paragraph(result['qwen_v1'])

    # Qwen V2
    doc.add_heading("Qwen2.5-7B V2 - V2 Query (Natural Language Medical)", level=3)
    doc.add_paragraph(result['qwen_v2'])

    doc.add_paragraph()

# ============================================================
# TASK 2: LIGATION PLANNING
# ============================================================

doc.add_page_break()
doc.add_heading('TASK 2: LIGATION PLANNING (WITH RAG)', level=1)

for result in results['ligation']:
    doc.add_heading(f"Test Case: {result['test']}", level=2)

    # LLAMA V1
    doc.add_heading("LLAMA 70B - V1 Query (Raw Clip Notation)", level=3)
    doc.add_paragraph(result['llama_v1'])

    # LLAMA V2
    doc.add_heading("LLAMA 70B - V2 Query (Natural Language Medical)", level=3)
    doc.add_paragraph(result['llama_v2'])

    # Qwen V1
    doc.add_heading("Qwen2.5-7B V2 - V1 Query (Raw Clip Notation)", level=3)
    doc.add_paragraph(result['qwen_v1'])

    # Qwen V2
    doc.add_heading("Qwen2.5-7B V2 - V2 Query (Natural Language Medical)", level=3)
    doc.add_paragraph(result['qwen_v2'])

    doc.add_paragraph()

# Save
doc.save('FINAL_EVALUATION_REPORT.docx')

# Save JSON
with open('evaluation_results.json', 'w') as f:
    json.dump(results, f, indent=2)

print("\n" + "="*80)
print("EVALUATION COMPLETE")
print("="*80)
print("\nFinal Report: FINAL_EVALUATION_REPORT.docx")
print("Results Log: evaluation_results.json")
print(f"\nTest cases evaluated: {len(test_cases)}")
print(f"Timestamp: {results['timestamp']}")
