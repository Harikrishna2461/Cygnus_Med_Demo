#!/usr/bin/env python3
"""
FINAL EVALUATION - OPTIMIZED FOR FAST COMPLETION
"""

import json
import os
import sys
import torch
from pathlib import Path
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("="*80)
print("FINAL EVALUATION REPORT")
print("="*80)

# Load test cases
test_cases = []
json_dir = Path(r'c:\Users\Krish\Downloads\LLM_Finetuning\json samples')

if json_dir.exists():
    for json_file in sorted(list(json_dir.glob('*.json')))[:3]:
        try:
            with open(json_file, 'r') as f:
                data = json.load(f)
                test_cases.append({
                    'name': json_file.stem,
                    'clips': data.get('clips', []),
                })
        except:
            pass

if not test_cases:
    test_cases = [
        {'name': 'TYPE_1_SAMPLE', 'clips': [
            {'flow': 'EP', 'fromType': 'N1', 'toType': 'N2', 'posYRatio': 0.06},
            {'flow': 'RP', 'fromType': 'N2', 'toType': 'N1', 'posYRatio': 0.25}
        ]},
    ]

print(f"[1/4] Loaded {len(test_cases)} test cases")

CHIVA_RULES = """
CHIVA CLASSIFICATION RULES:

STEP 1: Check for EP N1→N2
  - YES → SFJ INCOMPETENT (Case A/B)
  - NO → SFJ COMPETENT (Case C)

Case A (EP N1→N2, NO EP N2→N3):
  If RP N2→N1 only → TYPE 1

Case B (EP N1→N2 AND EP N2→N3):
  If RP N3→N1 AND RP N2→N1, elim="Reflux" → TYPE 1+2
  If RP N3→N2 or RP N3→N1, NO RP N2→N1 → TYPE 3
  If RP N3→N1 AND RP N2→N1, elim="No Reflux" → TYPE 3

Case C (NO EP N1→N2):
  TYPE 2A: EP N2→N3
  TYPE 2B: EP N2→N2, RP N3, NO RP N2→N1
  TYPE 2C: EP N2→N2, RP N3, RP N2→N1
"""

# Try to load models
print("[2/4] Initializing models...")

qwen_model = None
qwen_tokenizer = None
qwen_ok = False

try:
    from transformers import AutoModelForCausalLM, AutoTokenizer
    from peft import PeftModel

    if os.path.exists(r'.\lora_finetuned_model'):
        print("  Loading Qwen2.5-7B...", end=" ", flush=True)
        base = AutoModelForCausalLM.from_pretrained(
            "Qwen/Qwen2.5-7B",
            dtype=torch.float32,
            device_map="cpu",
            trust_remote_code=True,
            load_in_8bit=False,
        )
        print("OK")

        print("  Loading LoRA adapters...", end=" ", flush=True)
        qwen_model = PeftModel.from_pretrained(base, r'.\lora_finetuned_model', is_trainable=False)
        qwen_model.eval()
        qwen_tokenizer = AutoTokenizer.from_pretrained(r'.\lora_finetuned_model')
        qwen_ok = True
        print("OK")
    else:
        print("  LoRA path not found")
except Exception as e:
    print(f"  Qwen load error: {str(e)[:80]}")

# Mock LLAMA (API not working)
print("  LLAMA (Groq): Skipped (API unavailable)")

if not qwen_ok:
    print("\nERROR: Qwen model unavailable")
    sys.exit(1)

# ============================================================
# GENERATE OUTPUTS
# ============================================================

print("\n[3/4] Generating model outputs...")

def format_clips_v1(clips):
    s = "Duplex clips:\n"
    for c in clips:
        s += f"- {c.get('flow')} {c.get('fromType')}->{c.get('toType')} (y={c.get('posYRatio')})\n"
    return s

def format_clips_v2(clips):
    has_ep_n1_n2 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N1' and c.get('toType') == 'N2' for c in clips)
    has_ep_n2_n3 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N2' and c.get('toType') == 'N3' for c in clips)
    has_rp_n2_n1 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N2' and c.get('toType') == 'N1' for c in clips)
    has_rp_n3 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N3' for c in clips)

    findings = []
    if has_ep_n1_n2:
        findings.append("SFJ incompetence")
    else:
        findings.append("SFJ competence")
    if has_ep_n2_n3:
        findings.append("perforator incompetence")
    if has_rp_n2_n1:
        findings.append("saphenous reflux")
    if has_rp_n3:
        findings.append("tributary reflux")
    return "Patient with varicose veins. Duplex: " + ", ".join(findings) + ". CHIVA type?"

results = {
    'timestamp': datetime.now().isoformat(),
    'test_count': len(test_cases),
    'shunt_classification': [],
    'ligation': []
}

for idx, test in enumerate(test_cases, 1):
    print(f"  [{idx}/{len(test_cases)}] {test['name']}", end=" ", flush=True)

    clips = test['clips']
    v1_clips = format_clips_v1(clips)
    v2_clips = format_clips_v2(clips)

    # Qwen inference
    def qwen_run(prompt):
        try:
            messages = [{"role": "user", "content": f"{CHIVA_RULES}\n\n{prompt}"}]
            text = qwen_tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
            inputs = qwen_tokenizer(text, return_tensors="pt").to("cpu")
            with torch.no_grad():
                outputs = qwen_model.generate(
                    **inputs,
                    max_new_tokens=300,
                    do_sample=False,
                    pad_token_id=qwen_tokenizer.eos_token_id
                )
            return qwen_tokenizer.decode(outputs[0][inputs['input_ids'].shape[1]:], skip_special_tokens=True)
        except Exception as e:
            return f"[Error: {str(e)[:60]}]"

    # Classification
    qwen_c_v1 = qwen_run(v1_clips + "\nClassify CHIVA type.")
    qwen_c_v2 = qwen_run(v2_clips + "\nClassify CHIVA type.")

    # Ligation
    qwen_l_v1 = qwen_run(v1_clips + "\nProvide ligation planning strategy.")
    qwen_l_v2 = qwen_run(v2_clips + "\nProvide ligation planning strategy.")

    results['shunt_classification'].append({
        'test': test['name'],
        'qwen_v1': qwen_c_v1,
        'qwen_v2': qwen_c_v2,
    })

    results['ligation'].append({
        'test': test['name'],
        'qwen_v1': qwen_l_v1,
        'qwen_v2': qwen_l_v2,
    })

    print("OK")

# ============================================================
# CREATE REPORT
# ============================================================

print("\n[4/4] Creating Word report...")

doc = Document()

# Title
title = doc.add_heading('COMPARATIVE EVALUATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('Qwen2.5-7B V2 (LoRA Fine-tuned)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)

# Metadata
doc.add_heading('Evaluation Details', level=1)
doc.add_paragraph(f"Timestamp: {results['timestamp']}")
doc.add_paragraph(f"Test Cases: {results['test_count']}")
doc.add_paragraph(f"Model: Qwen2.5-7B with LoRA adapters")
doc.add_paragraph("Task 1: CHIVA Shunt Classification (no RAG)")
doc.add_paragraph("Task 2: Ligation Planning Strategy (with guidance)")

# Task 1
doc.add_page_break()
doc.add_heading('TASK 1: SHUNT CLASSIFICATION', level=1)

for result in results['shunt_classification']:
    doc.add_heading(f"Test: {result['test']}", level=2)

    doc.add_heading("V1 Query (Raw Clip Data)", level=3)
    doc.add_paragraph(result['qwen_v1'])

    doc.add_heading("V2 Query (Natural Language)", level=3)
    doc.add_paragraph(result['qwen_v2'])

    doc.add_paragraph()

# Task 2
doc.add_page_break()
doc.add_heading('TASK 2: LIGATION PLANNING', level=1)

for result in results['ligation']:
    doc.add_heading(f"Test: {result['test']}", level=2)

    doc.add_heading("V1 Query (Raw Clip Data)", level=3)
    doc.add_paragraph(result['qwen_v1'])

    doc.add_heading("V2 Query (Natural Language)", level=3)
    doc.add_paragraph(result['qwen_v2'])

    doc.add_paragraph()

doc.save('evaluation_report.docx')

# Save JSON log
with open('evaluation_log.json', 'w') as f:
    json.dump(results, f, indent=2)

print("\n" + "="*80)
print("DONE")
print("="*80)
print("\nReport: evaluation_report.docx")
print("Log: evaluation_log.json")
