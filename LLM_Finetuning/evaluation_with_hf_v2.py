#!/usr/bin/env python3
"""
EVALUATION WITH HF V2 MODEL: HariKrishna1824/qwen25_chiva_v2
Real model outputs - no mocking
"""

import os
import json
import torch
from pathlib import Path
from datetime import datetime
from typing import Dict, List, Any
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("="*80)
print("EVALUATION WITH V2 MODEL FROM HUGGINGFACE")
print("="*80)

# Configuration
HF_TOKEN = "YOUR_HF_TOKEN_HERE"
HF_MODEL_REPO = "HariKrishna1824/qwen25_chiva_v2"
TEST_CASES_DIR = Path("json samples")

CHIVA_RULES = """
=== CHIVA CLASSIFICATION RULES ===

ANATOMY:
    N1 = Deep venous system
    N2 = GSV/SSV trunk
    N3 = Tributaries
    EP = Physiological (antegrade) flow
    RP = Retrograde (reflux) flow

CLASSIFICATION:
    STEP 1: Check EP N1→N2
        YES → SFJ INCOMPETENT (Case A/B)
        NO → SFJ COMPETENT (Case C)

    Case A (EP N1→N2, NO EP N2→N3):
        RP N2→N1 only → TYPE 1

    Case B (EP N1→N2 AND EP N2→N3):
        RP N3→N1 AND RP N2→N1, elim="Reflux" → TYPE 1+2
        RP N3→N2 or RP N3→N1, NO RP N2→N1 → TYPE 3

    Case C (NO EP N1→N2):
        TYPE 2A: EP N2→N3
        TYPE 2B: EP N2→N2, RP N3, NO RP N2→N1
        TYPE 2C: EP N2→N2, RP N3, RP N2→N1
"""

# ============================================================
# LOAD TEST CASES
# ============================================================

print("\n[1/4] Loading test cases...")

test_cases = []
if TEST_CASES_DIR.exists():
    for json_file in sorted(list(TEST_CASES_DIR.glob("*.json"))):
        try:
            with open(json_file, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
                clips = data.get('clips', [])
                if clips:
                    test_cases.append({
                        'name': json_file.stem,
                        'clips': clips,
                    })
        except:
            pass

if not test_cases:
    # Demo cases
    test_cases = [
        {'name': 'TYPE_1', 'clips': [
            {'flow': 'EP', 'fromType': 'N1', 'toType': 'N2', 'posYRatio': 0.06},
            {'flow': 'RP', 'fromType': 'N2', 'toType': 'N1', 'posYRatio': 0.25}
        ]},
    ]

print(f"  Loaded {len(test_cases)} test cases")

# ============================================================
# LOAD MODEL FROM HF
# ============================================================

print("\n[2/4] Loading model from Hugging Face...")
print(f"  Repository: {HF_MODEL_REPO}")

try:
    from transformers import AutoModelForCausalLM, AutoTokenizer
    from peft import PeftModel

    print("  Loading base model (Qwen2.5-7B)...", end=" ", flush=True)
    base_model = AutoModelForCausalLM.from_pretrained(
        "Qwen/Qwen2.5-7B",
        torch_dtype=torch.bfloat16 if torch.cuda.is_available() else torch.float32,
        device_map="auto" if torch.cuda.is_available() else "cpu",
        trust_remote_code=True,
        token=HF_TOKEN,
    )
    print("OK")

    print("  Loading LoRA from HF...", end=" ", flush=True)
    model = PeftModel.from_pretrained(
        base_model,
        HF_MODEL_REPO,
        token=HF_TOKEN,
        is_trainable=False
    )
    print("OK")

    print("  Loading tokenizer...", end=" ", flush=True)
    tokenizer = AutoTokenizer.from_pretrained(
        HF_MODEL_REPO,
        token=HF_TOKEN,
        trust_remote_code=True,
    )
    print("OK")

    model.eval()
    device = next(model.parameters()).device
    print(f"  Device: {device}")

except Exception as e:
    print(f"\n  ERROR: {str(e)}")
    print("\n  Make sure:")
    print(f"    1. Model exists: {HF_MODEL_REPO}")
    print(f"    2. Token is valid: {HF_TOKEN[:20]}...")
    print(f"    3. Internet connection works")
    import sys
    sys.exit(1)

# ============================================================
# INFERENCE FUNCTIONS
# ============================================================

def format_clips_v1(clips: List[Dict]) -> str:
    """V1: Raw clip notation"""
    s = "Duplex clips:\n"
    for c in clips:
        s += f"- {c.get('flow')} {c.get('fromType')}->{c.get('toType')} (y={c.get('posYRatio')})\n"
    return s

def format_clips_v2(clips: List[Dict]) -> str:
    """V2: Natural language medical description"""
    has_ep_n1_n2 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N1' and c.get('toType') == 'N2' for c in clips)
    has_ep_n2_n3 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N2' and c.get('toType') == 'N3' for c in clips)
    has_rp_n2_n1 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N2' and c.get('toType') == 'N1' for c in clips)
    has_rp_n3 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N3' for c in clips)

    findings = []
    if has_ep_n1_n2:
        findings.append("saphenofemoral junction incompetence")
    else:
        findings.append("competent saphenofemoral junction")
    if has_ep_n2_n3:
        findings.append("perforator feeding into the saphenous vein")
    if has_rp_n2_n1:
        findings.append("retrograde flow within the saphenous trunk")
    if has_rp_n3:
        findings.append("retrograde flow at the tributary level")

    return "Patient with varicose veins. Duplex shows: " + ", ".join(findings) + "."

def infer(prompt: str, max_tokens: int = 400) -> str:
    """Run inference on model"""
    try:
        messages = [{"role": "user", "content": f"{CHIVA_RULES}\n\n{prompt}"}]
        text = tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
        inputs = tokenizer(text, return_tensors="pt").to(device)

        with torch.no_grad():
            outputs = model.generate(
                **inputs,
                max_new_tokens=max_tokens,
                do_sample=False,
                pad_token_id=tokenizer.eos_token_id
            )

        return tokenizer.decode(outputs[0][inputs['input_ids'].shape[1]:], skip_special_tokens=True)
    except Exception as e:
        return f"[Error: {str(e)[:100]}]"

# ============================================================
# RUN EVALUATION
# ============================================================

print("\n[3/4] Running inference...")

results = {
    'timestamp': datetime.now().isoformat(),
    'model': HF_MODEL_REPO,
    'test_count': len(test_cases),
    'shunt_classification': [],
    'ligation': []
}

rag_docs = [
    "For TYPE 1 shunt: Ligate GSV at saphenofemoral junction, preserve distal GSV if quality permits",
    "For TYPE 2A shunt: Perform selective perforator ligation with duplex guidance",
    "For TYPE 2B shunt: Conservative therapy with compression; consider staged perforator intervention",
    "For TYPE 2C shunt: Evaluate SFJ status; may require SFJ ligation with perforator management",
    "Post-operative duplex surveillance recommended for all interventions",
]

for idx, test in enumerate(test_cases, 1):
    print(f"  [{idx}/{len(test_cases)}] {test['name']}", end=" ", flush=True)

    clips = test['clips']
    v1_clips = format_clips_v1(clips)
    v2_clips = format_clips_v2(clips)

    # Task 1: Classification
    print("(classify)", end=" ", flush=True)
    c_v1 = infer(v1_clips + "\nClassify the CHIVA shunt type.")
    c_v2 = infer(v2_clips + "\nWhat is the CHIVA shunt type?")

    # Task 2: Ligation
    print("(ligate)", end=" ", flush=True)
    rag_context = "\n".join(rag_docs)
    l_v1 = infer(f"{rag_context}\n\n{v1_clips}\n\nProvide ligation planning strategy.")
    l_v2 = infer(f"{rag_context}\n\n{v2_clips}\n\nProvide ligation planning strategy.")

    results['shunt_classification'].append({
        'test': test['name'],
        'v1': c_v1,
        'v2': c_v2,
    })

    results['ligation'].append({
        'test': test['name'],
        'v1': l_v1,
        'v2': l_v2,
    })

    print("OK")

# ============================================================
# GENERATE REPORT
# ============================================================

print("\n[4/4] Generating Word report...")

doc = Document()

# Title
title = doc.add_heading('EVALUATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph(f'Model: {HF_MODEL_REPO}')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)

# Metadata
doc.add_heading('Evaluation Details', level=1)
doc.add_paragraph(f"Timestamp: {results['timestamp']}")
doc.add_paragraph(f"Model: {results['model']}")
doc.add_paragraph(f"Test Cases: {results['test_count']}")
doc.add_paragraph("Query Formats: V1 (Raw Clips), V2 (Natural Language)")
doc.add_paragraph("Tasks: 1) Shunt Classification, 2) Ligation Planning")

# Task 1: Classification
doc.add_page_break()
doc.add_heading('TASK 1: SHUNT CLASSIFICATION', level=1)

for result in results['shunt_classification']:
    doc.add_heading(f"Test Case: {result['test']}", level=2)

    doc.add_heading("V1 Query (Raw Clip Data)", level=3)
    doc.add_paragraph(result['v1'])

    doc.add_heading("V2 Query (Natural Language)", level=3)
    doc.add_paragraph(result['v2'])

    doc.add_paragraph()

# Task 2: Ligation
doc.add_page_break()
doc.add_heading('TASK 2: LIGATION PLANNING', level=1)

for result in results['ligation']:
    doc.add_heading(f"Test Case: {result['test']}", level=2)

    doc.add_heading("V1 Query (Raw Clip Data with RAG)", level=3)
    doc.add_paragraph(result['v1'])

    doc.add_heading("V2 Query (Natural Language with RAG)", level=3)
    doc.add_paragraph(result['v2'])

    doc.add_paragraph()

# Save
doc.save('evaluation_report_v2.docx')

# Save JSON log
with open('evaluation_log_v2.json', 'w') as f:
    json.dump(results, f, indent=2)

print("\n" + "="*80)
print("EVALUATION COMPLETE")
print("="*80)
print("\nFiles saved:")
print("  - evaluation_report_v2.docx")
print("  - evaluation_log_v2.json")
print(f"\nModel: {HF_MODEL_REPO}")
print(f"Test cases: {len(test_cases)}")
print(f"Timestamp: {results['timestamp']}")
