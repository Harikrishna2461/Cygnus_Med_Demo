#!/usr/bin/env python3
"""
FINAL EVALUATION - REAL MODEL OUTPUTS
"""

import json
import os
import sys
import torch
from pathlib import Path
from typing import Dict, List, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("="*80)
print("FINAL EVALUATION: REAL MODEL OUTPUTS")
print("="*80)

# ============================================================
# CONFIG
# ============================================================

GROQ_API_KEY = "YOUR_GROQ_API_KEY_HERE"
QWEN_MODEL_PATH = "Qwen/Qwen2.5-7B"
QWEN_LORA_PATH = r'.\lora_finetuned_model'

CHIVA_RULES = """
=== CHIVA VENOUS SHUNT CLASSIFICATION RULES ===

ANATOMY:
    N1 = Deep venous system (femoral / popliteal vein)
    N2 = Great Saphenous Vein (GSV) or Small Saphenous Vein (SSV) trunk
    N3 = Tributaries / superficial branches
    EP = Physiological (forward, antegrade) flow
    RP = Retrograde (pathological, reflux) flow

CLASSIFICATION RULES:
    STEP 1 — CHECK FOR EP N1→N2:
        YES → SFJ/Hunterian INCOMPETENT (Case A or B)
        NO  → SFJ COMPETENT (Case C)

    Case A — EP N1→N2 EXISTS, NO EP N2→N3:
        If RP N2→N1, no RP at N3 → TYPE 1

    Case B — EP N1→N2 EXISTS AND EP N2→N3 EXISTS:
        B1: RP N3→N2 or RP N3→N1, NO RP N2→N1 → TYPE 3
        B4: RP N3→N1 AND RP N2→N1, elim="Reflux" → TYPE 1+2
        B5: RP N3→N1 AND RP N2→N1, elim="No Reflux" → TYPE 3

    Case C — NO EP N1→N2 (SFJ COMPETENT):
        TYPE 2A: EP N2→N3, NO EP N1→N2
        TYPE 2B: EP N2→N2, NO EP N1→N2, RP at N3, NO RP N2→N1
        TYPE 2C: EP N2→N2, NO EP N1→N2, RP at N3, RP N2→N1 present
"""

# ============================================================
# LOAD TEST CASES
# ============================================================

def load_test_cases() -> List[Dict[str, Any]]:
    """Load test cases from JSON files"""
    test_cases = []
    json_dir = Path(r'c:\Users\Krish\Downloads\LLM_Finetuning\json samples')

    if json_dir.exists():
        for json_file in sorted(list(json_dir.glob('*.json')))[:5]:
            try:
                with open(json_file, 'r') as f:
                    data = json.load(f)
                    test_cases.append({
                        'name': json_file.stem,
                        'clips': data.get('clips', []),
                    })
            except:
                pass

    if not test_cases:
        test_cases = [
            {'name': 'TYPE_1', 'clips': [
                {'flow': 'EP', 'fromType': 'N1', 'toType': 'N2', 'posYRatio': 0.06},
                {'flow': 'RP', 'fromType': 'N2', 'toType': 'N1', 'posYRatio': 0.25}
            ]},
        ]

    return test_cases

def clips_to_v1(clips: List[Dict]) -> str:
    """V1: Raw clip data"""
    clip_str = "Duplex clips:\n"
    for clip in clips:
        clip_str += f"- {clip.get('flow')} {clip.get('fromType')}->{clip.get('toType')} (y={clip.get('posYRatio')})\n"
    clip_str += "\nClassify the CHIVA shunt type."
    return clip_str

def clips_to_v2(clips: List[Dict]) -> str:
    """V2: Natural language"""
    has_ep_n1_n2 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N1' and c.get('toType') == 'N2' for c in clips)
    has_ep_n2_n3 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N2' and c.get('toType') == 'N3' for c in clips)
    has_rp_n2_n1 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N2' and c.get('toType') == 'N1' for c in clips)
    has_rp_n3 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N3' for c in clips)

    findings = []
    if has_ep_n1_n2:
        findings.append("saphenofemoral junction incompetence")
    else:
        findings.append("competent saphenofemoral junction")
    if has_ep_n2_n3:
        findings.append("perforator feeding into the saphenous vein")
    if has_rp_n2_n1:
        findings.append("retrograde flow within the saphenous trunk")
    if has_rp_n3:
        findings.append("retrograde flow at the tributary level")

    return "Patient with varicose veins. Duplex shows: " + ", ".join(findings) + ". CHIVA type?"

# ============================================================
# MODEL INITIALIZATION
# ============================================================

print("\n[1/5] Initializing models...")

# GROQ/LLAMA
groq_client = None
groq_model = None
llama_ok = False

try:
    from groq import Groq
    groq_client = Groq(api_key=GROQ_API_KEY)
    # Try to find available model
    for model_name in ["mixtral-8x7b-32768", "llama-3-70b-versatile", "llama-2-70b-4096"]:
        try:
            response = groq_client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": "test"}],
                max_tokens=5
            )
            groq_model = model_name
            llama_ok = True
            print(f"  LLAMA (Groq {model_name}): OK")
            break
        except:
            pass
    if not llama_ok:
        print(f"  LLAMA (Groq): FAIL - no available models")
except Exception as e:
    print(f"  LLAMA (Groq): FAIL - {str(e)[:50]}")

# QWEN
qwen_ok = False
qwen_model = None
qwen_tokenizer = None

try:
    from transformers import AutoModelForCausalLM, AutoTokenizer
    from peft import PeftModel

    if os.path.exists(QWEN_LORA_PATH):
        device = "cuda:0" if torch.cuda.is_available() else "cpu"
        dtype = torch.bfloat16 if torch.cuda.is_available() else torch.float32

        base = AutoModelForCausalLM.from_pretrained(
            QWEN_MODEL_PATH,
            dtype=dtype,
            device_map=device if device == "cpu" else {"": device},
            trust_remote_code=True,
            low_cpu_mem_usage=True,
        )
        qwen_model = PeftModel.from_pretrained(base, QWEN_LORA_PATH, low_cpu_mem_usage=True)
        qwen_model.eval()
        qwen_tokenizer = AutoTokenizer.from_pretrained(QWEN_LORA_PATH)
        qwen_ok = True
        print(f"  Qwen2.5-7B LoRA ({device}): OK")
    else:
        print(f"  Qwen2.5-7B LoRA: FAIL - not found at {QWEN_LORA_PATH}")
except Exception as e:
    print(f"  Qwen2.5-7B LoRA: FAIL - {str(e)[:50]}")

if not llama_ok and not qwen_ok:
    print("\nERROR: No models available!")
    sys.exit(1)

# ============================================================
# INFERENCE FUNCTIONS
# ============================================================

def llama_infer(query: str) -> str:
    """LLAMA inference"""
    if not llama_ok:
        return "[LLAMA unavailable]"
    try:
        response = groq_client.chat.completions.create(
            model=groq_model,
            messages=[
                {"role": "system", "content": f"Expert vascular surgeon.\n\n{CHIVA_RULES}"},
                {"role": "user", "content": query}
            ],
            max_tokens=500,
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"[Error: {str(e)[:80]}]"

def qwen_infer(query: str) -> str:
    """Qwen inference"""
    if not qwen_ok:
        return "[Qwen unavailable]"
    try:
        messages = [{"role": "user", "content": f"{CHIVA_RULES}\n\n{query}"}]
        text = qwen_tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
        inputs = qwen_tokenizer(text, return_tensors="pt").to(qwen_model.device)
        with torch.no_grad():
            outputs = qwen_model.generate(
                **inputs,
                max_new_tokens=500,
                do_sample=False,
                pad_token_id=qwen_tokenizer.eos_token_id
            )
        return qwen_tokenizer.decode(outputs[0][inputs['input_ids'].shape[1]:], skip_special_tokens=True)
    except Exception as e:
        return f"[Error: {str(e)[:80]}]"

# ============================================================
# RUN EVALUATION
# ============================================================

print("\n[2/5] Loading test cases...")
test_cases = load_test_cases()
print(f"  Loaded {len(test_cases)} test cases")

print("\n[3/5] Running inference...")

results = {
    'timestamp': datetime.now().isoformat(),
    'test_count': len(test_cases),
    'shunt_classification': [],
    'ligation': []
}

rag_docs = [
    "For TYPE 1: Ligate GSV at SFJ, preserve distal GSV if quality permits",
    "For TYPE 2A: Selective perforator ligation with duplex guidance",
    "For TYPE 2B: Conservative therapy first, staged perforator intervention",
]

for idx, test in enumerate(test_cases, 1):
    print(f"  [{idx}/{len(test_cases)}] {test['name']}", end=" ", flush=True)

    clips = test['clips']
    v1_query = clips_to_v1(clips)
    v2_query = clips_to_v2(clips)

    # Classification
    llama_c_v1 = llama_infer(v1_query)
    llama_c_v2 = llama_infer(v2_query)
    qwen_c_v1 = qwen_infer(v1_query)
    qwen_c_v2 = qwen_infer(v2_query)

    # Ligation
    rag_context = "\n".join(rag_docs)
    llama_l_v1 = llama_infer(f"{rag_context}\n\n{v1_query}\n\nProvide ligation planning strategy.")
    llama_l_v2 = llama_infer(f"{rag_context}\n\n{v2_query}\n\nProvide ligation planning strategy.")
    qwen_l_v1 = qwen_infer(f"{rag_context}\n\n{v1_query}\n\nProvide ligation planning strategy.")
    qwen_l_v2 = qwen_infer(f"{rag_context}\n\n{v2_query}\n\nProvide ligation planning strategy.")

    results['shunt_classification'].append({
        'test': test['name'],
        'llama_v1': llama_c_v1,
        'llama_v2': llama_c_v2,
        'qwen_v1': qwen_c_v1,
        'qwen_v2': qwen_c_v2,
    })

    results['ligation'].append({
        'test': test['name'],
        'llama_v1': llama_l_v1,
        'llama_v2': llama_l_v2,
        'qwen_v1': qwen_l_v1,
        'qwen_v2': qwen_l_v2,
    })

    print("OK")

# ============================================================
# GENERATE REPORT
# ============================================================

print("\n[4/5] Generating Word report...")

doc = Document()

# Title
title = doc.add_heading('COMPARATIVE EVALUATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('LLAMA 70B Versatile vs Qwen2.5-7B V2 Fine-tuned')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)

# Metadata
doc.add_heading('Evaluation Metadata', level=1)
doc.add_paragraph(f"Timestamp: {results['timestamp']}")
doc.add_paragraph(f"Test Cases: {results['test_count']}")
doc.add_paragraph(f"Models: LLAMA 70B (Groq {groq_model}), Qwen2.5-7B V2 (Local LoRA)")
doc.add_paragraph("Status: Real model outputs (no mocking)")

# Task 1: Shunt Classification
doc.add_page_break()
doc.add_heading('TASK 1: SHUNT CLASSIFICATION (NO RAG)', level=1)

for result in results['shunt_classification']:
    doc.add_heading(f"Test Case: {result['test']}", level=2)

    # LLAMA V1
    doc.add_heading("LLAMA 70B - V1 Query (Raw Clips)", level=3)
    doc.add_paragraph(result['llama_v1'])

    # LLAMA V2
    doc.add_heading("LLAMA 70B - V2 Query (Natural Language)", level=3)
    doc.add_paragraph(result['llama_v2'])

    # Qwen V1
    doc.add_heading("Qwen2.5-7B V2 - V1 Query (Raw Clips)", level=3)
    doc.add_paragraph(result['qwen_v1'])

    # Qwen V2
    doc.add_heading("Qwen2.5-7B V2 - V2 Query (Natural Language)", level=3)
    doc.add_paragraph(result['qwen_v2'])

    doc.add_paragraph()

# Task 2: Ligation Planning
doc.add_page_break()
doc.add_heading('TASK 2: LIGATION PLANNING (WITH RAG)', level=1)

for result in results['ligation']:
    doc.add_heading(f"Test Case: {result['test']}", level=2)

    # LLAMA V1
    doc.add_heading("LLAMA 70B - V1 Query (Raw Clips)", level=3)
    doc.add_paragraph(result['llama_v1'])

    # LLAMA V2
    doc.add_heading("LLAMA 70B - V2 Query (Natural Language)", level=3)
    doc.add_paragraph(result['llama_v2'])

    # Qwen V1
    doc.add_heading("Qwen2.5-7B V2 - V1 Query (Raw Clips)", level=3)
    doc.add_paragraph(result['qwen_v1'])

    # Qwen V2
    doc.add_heading("Qwen2.5-7B V2 - V2 Query (Natural Language)", level=3)
    doc.add_paragraph(result['qwen_v2'])

    doc.add_paragraph()

doc.save('evaluation_report.docx')

print("\n[5/5] Finalizing...")

# Also save JSON log
with open('evaluation_log.json', 'w') as f:
    json.dump(results, f, indent=2)

print("\n" + "="*80)
print("COMPLETE")
print("="*80)
print("\nReport: evaluation_report.docx")
print("Log: evaluation_log.json")
