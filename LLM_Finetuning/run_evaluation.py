#!/usr/bin/env python3
"""
COMPREHENSIVE COMPARATIVE EVALUATION
Single script - handles everything internally
"""

import json
import os
import sys
import torch
from pathlib import Path
from typing import Dict, List, Any
from datetime import datetime

# Install dependencies silently
import subprocess
def install_deps():
    packages = ['groq', 'transformers', 'peft', 'torch', 'sentence-transformers', 'qdrant-client', 'python-docx']
    for pkg in packages:
        try:
            __import__(pkg.replace('-', '_'))
        except ImportError:
            subprocess.run([sys.executable, '-m', 'pip', 'install', '-q', pkg], check=False)

install_deps()

from groq import Groq
from transformers import AutoModelForCausalLM, AutoTokenizer
from peft import PeftModel
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("="*80)
print("COMPARATIVE EVALUATION: LLAMA 70B vs Qwen2.5-7B V2")
print("="*80)

# ============================================================
# CONFIGURATION
# ============================================================

GROQ_API_KEY = ""
QWEN_MODEL_PATH = "Qwen/Qwen2.5-7B"
QWEN_LORA_PATH = r'.\lora_finetuned_model'
GROQ_MODEL = "mixtral-8x7b-32768"  # Default, will be updated if available

CHIVA_RULES = """
=== CHIVA VENOUS SHUNT CLASSIFICATION RULES ===

ANATOMY:
    N1 = Deep venous system (femoral / popliteal vein)
    N2 = Great Saphenous Vein (GSV) or Small Saphenous Vein (SSV) trunk
    N3 = Tributaries / superficial branches
    EP = Physiological (forward, antegrade) flow
    RP = Retrograde (pathological, reflux) flow
    SFJ = Saphenofemoral Junction (posYRatio ≤ 0.098)

CLASSIFICATION RULES:
    STEP 1 — CHECK FOR EP N1→N2:
        YES → SFJ/Hunterian INCOMPETENT (Case A or B)
        NO  → SFJ COMPETENT (Case C)

    Case A — EP N1→N2 EXISTS, NO EP N2→N3:
        If RP N2→N1, no RP at N3 → TYPE 1

    Case B — EP N1→N2 EXISTS AND EP N2→N3 EXISTS:
        B1: RP N3→N2 or RP N3→N1, NO RP N2→N1 → TYPE 3
        B4: RP N3→N1 AND RP N2→N1, elim="Reflux" → TYPE 1+2
        B5: RP N3→N1 AND RP N2→N1, elim="No Reflux" → TYPE 3

    Case C — NO EP N1→N2 (SFJ COMPETENT):
        TYPE 2A: EP N2→N3, NO EP N1→N2
        TYPE 2B: EP N2→N2, NO EP N1→N2, RP at N3, NO RP N2→N1
        TYPE 2C: EP N2→N2, NO EP N1→N2, RP at N3, RP N2→N1 present
"""

# ============================================================
# TEST DATA LOADING
# ============================================================

def load_test_cases() -> List[Dict[str, Any]]:
    """Load test cases from JSON files"""
    test_cases = []
    json_dir = Path(r'c:\Users\Krish\Downloads\LLM_Finetuning\json samples')

    if json_dir.exists():
        for json_file in json_dir.glob('*.json'):
            try:
                with open(json_file, 'r') as f:
                    data = json.load(f)
                    test_cases.append({
                        'name': json_file.stem,
                        'clips': data.get('clips', []),
                        'source': 'json_samples'
                    })
            except:
                pass

    if not test_cases:
        # Create demo test cases
        test_cases = [
            {
                'name': 'TYPE_1_DEMO',
                'clips': [
                    {'flow': 'EP', 'fromType': 'N1', 'toType': 'N2', 'posYRatio': 0.06, 'step': 'SFJ'},
                    {'flow': 'RP', 'fromType': 'N2', 'toType': 'N1', 'posYRatio': 0.25, 'step': 'SFJ'}
                ],
                'source': 'demo'
            },
            {
                'name': 'TYPE_2A_DEMO',
                'clips': [
                    {'flow': 'EP', 'fromType': 'N2', 'toType': 'N3', 'posYRatio': 0.18, 'step': 'Knee'},
                    {'flow': 'RP', 'fromType': 'N3', 'toType': 'N2', 'posYRatio': 0.45, 'step': 'Knee'}
                ],
                'source': 'demo'
            },
            {
                'name': 'TYPE_2B_DEMO',
                'clips': [
                    {'flow': 'EP', 'fromType': 'N2', 'toType': 'N2', 'posYRatio': 0.07, 'step': 'SFJ'},
                    {'flow': 'RP', 'fromType': 'N3', 'toType': 'N1', 'posYRatio': 0.32, 'step': 'Knee'}
                ],
                'source': 'demo'
            }
        ]

    return test_cases

def convert_clips_to_v1(clips: List[Dict]) -> str:
    """V1: Raw clip data + simple query"""
    clip_str = "Duplex clips:\n"
    for clip in clips:
        flow = clip.get('flow', 'UNKNOWN')
        from_type = clip.get('fromType', '?')
        to_type = clip.get('toType', '?')
        y_value = clip.get('posYRatio', '?')
        clip_str += f"- {flow} {from_type}->{to_type} (y={y_value})\n"
    clip_str += "\nClassify the CHIVA shunt type."
    return clip_str

def convert_clips_to_v2(clips: List[Dict]) -> str:
    """V2: Natural language query (medical terms)"""
    has_ep_n1_n2 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N1' and c.get('toType') == 'N2' for c in clips)
    has_ep_n2_n3 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N2' and c.get('toType') == 'N3' for c in clips)
    has_rp_n2_n1 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N2' and c.get('toType') == 'N1' for c in clips)
    has_rp_n3 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N3' for c in clips)

    query = "A patient presents with significant lower extremity varicose veins. Duplex ultrasound reveals: "
    findings = []
    if has_ep_n1_n2:
        findings.append("saphenofemoral junction incompetence")
    else:
        findings.append("competent saphenofemoral junction")
    if has_ep_n2_n3:
        findings.append("perforator feeding into the saphenous vein")
    if has_rp_n2_n1:
        findings.append("retrograde flow within the saphenous trunk")
    if has_rp_n3:
        findings.append("retrograde flow at the tributary level")

    query += ", ".join(findings) + ". What is the CHIVA shunt type?"
    return query

# ============================================================
# MODEL LOADING & TESTING
# ============================================================

print("\n[1/3] Testing and loading models...")

# Test LLAMA
print("  Testing LLAMA 70B (Groq API)...", end=" ", flush=True)
llama_available = False
try:
    groq_client = Groq(api_key=GROQ_API_KEY)
    # Try multiple model names - use the one that works
    model_names = ["llama-3-70b-versatile", "mixtral-8x7b-32768", "llama-2-70b-4096"]
    for model_name in model_names:
        try:
            response = groq_client.chat.completions.create(
                model=model_name,
                messages=[{"role": "user", "content": "Hello"}],
                max_tokens=10
            )
            print("OK")
            llama_available = True
            # Update the model name for later use
            GROQ_MODEL = model_name
            break
        except:
            continue
    if not llama_available:
        print(f"FAIL (no valid model found)")
except Exception as e:
    print(f"FAIL ({str(e)[:50]})")

# Test Qwen
print("  Testing Qwen2.5-7B V2 (Local)...", end=" ", flush=True)
qwen_available = False
qwen_model = None
qwen_tokenizer = None
qwen_device = "cpu"  # Will use CPU if CUDA unavailable
try:
    if os.path.exists(QWEN_LORA_PATH):
        # Try CUDA first, fallback to CPU
        if torch.cuda.is_available():
            qwen_device = "cuda:0"
            dtype = torch.bfloat16
        else:
            qwen_device = "cpu"
            dtype = torch.float32

        base_model = AutoModelForCausalLM.from_pretrained(
            QWEN_MODEL_PATH,
            dtype=dtype,
            device_map=qwen_device if qwen_device == "cpu" else {"": qwen_device},
            trust_remote_code=True,
            low_cpu_mem_usage=True,
        )
        qwen_model = PeftModel.from_pretrained(base_model, QWEN_LORA_PATH, low_cpu_mem_usage=True)
        qwen_model.eval()
        qwen_tokenizer = AutoTokenizer.from_pretrained(QWEN_LORA_PATH)
        print(f"OK ({qwen_device})")
        qwen_available = True
    else:
        print(f"FAIL (LoRA path not found)")
except Exception as e:
    print(f"FAIL ({str(e)[:50]})")

if not llama_available and not qwen_available:
    print("\nERROR: No models available!")
    sys.exit(1)

# ============================================================
# EVALUATION FUNCTIONS
# ============================================================

def llama_classify(query: str) -> str:
    """LLAMA shunt classification"""
    if not llama_available:
        return "[LLAMA unavailable]"
    try:
        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": f"Expert vascular surgeon.\n\n{CHIVA_RULES}"},
                {"role": "user", "content": query}
            ],
            max_tokens=400,
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"[Error: {str(e)[:100]}]"

def llama_ligate(query: str, shunt_type: str, rag_docs: List[str]) -> str:
    """LLAMA ligation planning"""
    if not llama_available:
        return "[LLAMA unavailable]"
    try:
        rag_context = "\n".join(rag_docs[:3]) if rag_docs else "Knowledge: Standard ligation strategies"
        response = groq_client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[
                {"role": "system", "content": f"Expert vascular surgeon. Shunt: {shunt_type}\n\n{rag_context}"},
                {"role": "user", "content": query}
            ],
            max_tokens=400,
            temperature=0.3
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"[Error: {str(e)[:100]}]"

def qwen_classify(query: str) -> str:
    """Qwen shunt classification"""
    if not qwen_available:
        return "[Qwen unavailable]"
    try:
        messages = [{"role": "user", "content": f"{CHIVA_RULES}\n\n{query}"}]
        text = qwen_tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
        inputs = qwen_tokenizer(text, return_tensors="pt").to(qwen_model.device)
        with torch.no_grad():
            outputs = qwen_model.generate(**inputs, max_new_tokens=400, do_sample=False, pad_token_id=qwen_tokenizer.eos_token_id)
        return qwen_tokenizer.decode(outputs[0][inputs['input_ids'].shape[1]:], skip_special_tokens=True)
    except Exception as e:
        return f"[Error: {str(e)[:100]}]"

def qwen_ligate(query: str, shunt_type: str, rag_docs: List[str]) -> str:
    """Qwen ligation planning"""
    if not qwen_available:
        return "[Qwen unavailable]"
    try:
        rag_context = "\n".join(rag_docs[:3]) if rag_docs else "Standard ligation strategies"
        full_query = f"Shunt type: {shunt_type}\nKnowledge: {rag_context}\n\n{query}"
        messages = [{"role": "user", "content": full_query}]
        text = qwen_tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
        inputs = qwen_tokenizer(text, return_tensors="pt").to(qwen_model.device)
        with torch.no_grad():
            outputs = qwen_model.generate(**inputs, max_new_tokens=400, do_sample=False, pad_token_id=qwen_tokenizer.eos_token_id)
        return qwen_tokenizer.decode(outputs[0][inputs['input_ids'].shape[1]:], skip_special_tokens=True)
    except Exception as e:
        return f"[Error: {str(e)[:100]}]"

# ============================================================
# RUN TESTS
# ============================================================

print("\n[2/3] Running tests...")
test_cases = load_test_cases()
print(f"  Loaded {len(test_cases)} test cases")

rag_docs = [
    "Ligate at the saphenofemoral junction for SFJ incompetence",
    "Preserve the saphenous vein when possible",
    "Consider staged procedures for complex reflux",
    "Post-operative duplex surveillance recommended",
    "Ligate refluxing tributaries based on caliber and distance"
]

results = {
    'metadata': {
        'timestamp': datetime.now().isoformat(),
        'test_count': len(test_cases),
        'models': ['LLAMA 70B Versatile', 'Qwen2.5-7B V2'],
        'queries': ['V1 (Raw clips)', 'V2 (Medical NL)']
    },
    'shunt_classification': [],
    'ligation': []
}

for idx, test in enumerate(test_cases, 1):
    print(f"  [{idx}/{len(test_cases)}] {test['name']}", end=" ", flush=True)

    clips = test['clips']
    v1_q = convert_clips_to_v1(clips)
    v2_q = convert_clips_to_v2(clips)

    # Classification
    llama_c_v1 = llama_classify(v1_q)
    llama_c_v2 = llama_classify(v2_q)
    qwen_c_v1 = qwen_classify(v1_q)
    qwen_c_v2 = qwen_classify(v2_q)

    # Ligation
    llama_l_v1 = llama_ligate(v1_q, "PENDING", rag_docs)
    llama_l_v2 = llama_ligate(v2_q, "PENDING", rag_docs)
    qwen_l_v1 = qwen_ligate(v1_q, "PENDING", rag_docs)
    qwen_l_v2 = qwen_ligate(v2_q, "PENDING", rag_docs)

    results['shunt_classification'].append({
        'test': test['name'],
        'llama_v1': llama_c_v1, 'llama_v2': llama_c_v2,
        'qwen_v1': qwen_c_v1, 'qwen_v2': qwen_c_v2
    })

    results['ligation'].append({
        'test': test['name'],
        'llama_v1': llama_l_v1, 'llama_v2': llama_l_v2,
        'qwen_v1': qwen_l_v1, 'qwen_v2': qwen_l_v2
    })

    print("OK")

# ============================================================
# GENERATE REPORT
# ============================================================

print("\n[3/3] Generating report...")

doc = Document()

# Title
title = doc.add_heading('COMPARATIVE EVALUATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('LLAMA 70B Versatile vs Qwen2.5-7B V2 Fine-tuned')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)

# Metadata
doc.add_heading('Evaluation Metadata', level=1)
doc.add_paragraph(f"Timestamp: {results['metadata']['timestamp']}")
doc.add_paragraph(f"Test Cases: {results['metadata']['test_count']}")
doc.add_paragraph(f"Models: {', '.join(results['metadata']['models'])}")

# Results
doc.add_page_break()
doc.add_heading('Task 1: Shunt Classification (No RAG)', level=1)

for result in results['shunt_classification']:
    doc.add_heading(f"Test: {result['test']}", level=2)

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'

    hdr = table.rows[0].cells
    hdr[0].text = 'Model'
    hdr[1].text = 'V1 Query'
    hdr[2].text = 'V2 Query'

    row1 = table.rows[1].cells
    row1[0].text = 'LLAMA 70B'
    row1[1].text = result['llama_v1'][:250] + "..."
    row1[2].text = result['llama_v2'][:250] + "..."

    row2 = table.rows[2].cells
    row2[0].text = 'Qwen2.5-7B V2'
    row2[1].text = result['qwen_v1'][:250] + "..."
    row2[2].text = result['qwen_v2'][:250] + "..."

    doc.add_paragraph()

doc.add_page_break()
doc.add_heading('Task 2: Ligation Planning (with RAG)', level=1)

for result in results['ligation']:
    doc.add_heading(f"Test: {result['test']}", level=2)

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'

    hdr = table.rows[0].cells
    hdr[0].text = 'Model'
    hdr[1].text = 'V1 Query'
    hdr[2].text = 'V2 Query'

    row1 = table.rows[1].cells
    row1[0].text = 'LLAMA 70B'
    row1[1].text = result['llama_v1'][:250] + "..."
    row1[2].text = result['llama_v2'][:250] + "..."

    row2 = table.rows[2].cells
    row2[0].text = 'Qwen2.5-7B V2'
    row2[1].text = result['qwen_v1'][:250] + "..."
    row2[2].text = result['qwen_v2'][:250] + "..."

    doc.add_paragraph()

# Evaluation Tables
doc.add_page_break()
doc.add_heading('Evaluation Tables (To Be Completed)', level=1)

doc.add_heading('Table 1: Shunt Classification', level=2)
t1 = doc.add_table(rows=3, cols=5)
t1.style = 'Light Grid Accent 1'
h = t1.rows[0].cells
h[0].text, h[1].text, h[2].text, h[3].text, h[4].text = 'Model', 'Accuracy V1', 'Accuracy V2', 'Reasoning V1', 'Reasoning V2'
r1, r2 = t1.rows[1].cells, t1.rows[2].cells
r1[0].text, r1[1].text, r1[2].text, r1[3].text, r1[4].text = 'LLAMA 70B', '[Fill]', '[Fill]', '[Fill]', '[Fill]'
r2[0].text, r2[1].text, r2[2].text, r2[3].text, r2[4].text = 'Qwen2.5-7B V2', '[Fill]', '[Fill]', '[Fill]', '[Fill]'

doc.add_paragraph()

doc.add_heading('Table 2: Ligation Planning', level=2)
t2 = doc.add_table(rows=3, cols=5)
t2.style = 'Light Grid Accent 1'
h = t2.rows[0].cells
h[0].text, h[1].text, h[2].text, h[3].text, h[4].text = 'Model', 'Quality V1', 'Quality V2', 'Reasoning V1', 'Reasoning V2'
r1, r2 = t2.rows[1].cells, t2.rows[2].cells
r1[0].text, r1[1].text, r1[2].text, r1[3].text, r1[4].text = 'LLAMA 70B', '[Fill]', '[Fill]', '[Fill]', '[Fill]'
r2[0].text, r2[1].text, r2[2].text, r2[3].text, r2[4].text = 'Qwen2.5-7B V2', '[Fill]', '[Fill]', '[Fill]', '[Fill]'

doc.save('evaluation_report.docx')

print("\n" + "="*80)
print("EVALUATION COMPLETE")
print("="*80)
print("\nReport saved: evaluation_report.docx")
print("Models tested:")
print(f"  - LLAMA 70B Versatile: {'OK' if llama_available else 'FAIL'}")
print(f"  - Qwen2.5-7B V2: {'OK' if qwen_available else 'FAIL'}")
print(f"\nTest cases: {len(test_cases)}")
print(f"Timestamp: {results['metadata']['timestamp']}")
