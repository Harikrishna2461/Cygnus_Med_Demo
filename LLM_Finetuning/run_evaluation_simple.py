#!/usr/bin/env python3
"""
SIMPLIFIED COMPARATIVE EVALUATION - DEMO MODE
Generates report with sample data (models not actually loaded)
"""

import json
import os
import sys
from pathlib import Path
from typing import Dict, List, Any
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("="*80)
print("COMPARATIVE EVALUATION: LLAMA 70B vs Qwen2.5-7B V2 (DEMO MODE)")
print("="*80)

# ============================================================
# TEST DATA LOADING
# ============================================================

def load_test_cases() -> List[Dict[str, Any]]:
    """Load test cases from JSON files"""
    test_cases = []
    json_dir = Path(r'c:\Users\Krish\Downloads\LLM_Finetuning\json samples')

    if json_dir.exists():
        for json_file in list(json_dir.glob('*.json'))[:5]:  # Limit to 5 files
            try:
                with open(json_file, 'r') as f:
                    data = json.load(f)
                    test_cases.append({
                        'name': json_file.stem,
                        'clips': data.get('clips', []),
                        'source': 'json_samples'
                    })
            except:
                pass

    if not test_cases:
        test_cases = [
            {
                'name': 'TYPE_1_DEMO',
                'clips': [
                    {'flow': 'EP', 'fromType': 'N1', 'toType': 'N2', 'posYRatio': 0.06, 'step': 'SFJ'},
                    {'flow': 'RP', 'fromType': 'N2', 'toType': 'N1', 'posYRatio': 0.25, 'step': 'SFJ'}
                ],
                'source': 'demo'
            },
            {
                'name': 'TYPE_2A_DEMO',
                'clips': [
                    {'flow': 'EP', 'fromType': 'N2', 'toType': 'N3', 'posYRatio': 0.18, 'step': 'Knee'},
                    {'flow': 'RP', 'fromType': 'N3', 'toType': 'N2', 'posYRatio': 0.45, 'step': 'Knee'}
                ],
                'source': 'demo'
            },
            {
                'name': 'TYPE_2B_DEMO',
                'clips': [
                    {'flow': 'EP', 'fromType': 'N2', 'toType': 'N2', 'posYRatio': 0.07, 'step': 'SFJ'},
                    {'flow': 'RP', 'fromType': 'N3', 'toType': 'N1', 'posYRatio': 0.32, 'step': 'Knee'}
                ],
                'source': 'demo'
            }
        ]

    return test_cases

def convert_clips_to_v1(clips: List[Dict]) -> str:
    """V1: Raw clip data + simple query"""
    clip_str = "Duplex clips:\n"
    for clip in clips:
        flow = clip.get('flow', 'UNKNOWN')
        from_type = clip.get('fromType', '?')
        to_type = clip.get('toType', '?')
        y_value = clip.get('posYRatio', '?')
        clip_str += f"- {flow} {from_type}->{to_type} (y={y_value})\n"
    clip_str += "\nClassify the CHIVA shunt type."
    return clip_str

def convert_clips_to_v2(clips: List[Dict]) -> str:
    """V2: Natural language query (medical terms)"""
    has_ep_n1_n2 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N1' and c.get('toType') == 'N2' for c in clips)
    has_ep_n2_n3 = any(c.get('flow') == 'EP' and c.get('fromType') == 'N2' and c.get('toType') == 'N3' for c in clips)
    has_rp_n2_n1 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N2' and c.get('toType') == 'N1' for c in clips)
    has_rp_n3 = any(c.get('flow') == 'RP' and c.get('fromType') == 'N3' for c in clips)

    query = "A patient presents with significant lower extremity varicose veins. Duplex ultrasound reveals: "
    findings = []
    if has_ep_n1_n2:
        findings.append("saphenofemoral junction incompetence")
    else:
        findings.append("competent saphenofemoral junction")
    if has_ep_n2_n3:
        findings.append("perforator feeding into the saphenous vein")
    if has_rp_n2_n1:
        findings.append("retrograde flow within the saphenous trunk")
    if has_rp_n3:
        findings.append("retrograde flow at the tributary level")

    query += ", ".join(findings) + ". What is the CHIVA shunt type?"
    return query

# ============================================================
# MOCK RESPONSES
# ============================================================

LLAMA_RESPONSES_V1 = [
    "Based on the duplex findings showing antegrade flow from N1 to N2 with retrograde flow N2 to N1, this is classified as CHIVA TYPE 1. The saphenofemoral junction is incompetent with primary varicosity in the great saphenous vein trunk. Treatment strategy focuses on proximal ligation of the GSV at the saphenofemoral junction to eliminate the incompetent junction.",
    "The absence of antegrade flow N1→N2 indicates saphenofemoral junction competence (Case C). With antegrade N2→N3 and retrograde N3→N2, this is CHIVA TYPE 2A. The varicosities are secondary to incompetent perforating veins. Treatment involves perforator ligation or ablation at the site of incompetence.",
    "With antegrade N2→N2 flow, absence of N1→N2, and retrograde N3→N1, this represents CHIVA TYPE 2B. Incompetent perforators feed the saphenous trunk without SFJ involvement. Conservative management with sclerotherapy or selective perforator ligation is appropriate.",
]

LLAMA_RESPONSES_V2 = [
    "This patient demonstrates saphenofemoral junction incompetence with primary GSV reflux, indicative of CHIVA TYPE 1 shunt. The hemodynamic pattern of retrograde flow returning through the saphenous trunk to the deep system is pathognomonic for SFJ-driven varicosity. Surgical correction requires saphenofemoral junction ligation to interrupt the reflux source.",
    "The clinical presentation with competent SFJ but evidence of perforator incompetence feeding tributary varicosities is consistent with CHIVA TYPE 2A. This secondary varicosity pattern requires identification and treatment of the incompetent perforating vein. Duplex-guided sclerotherapy or minimally invasive perforator ablation provides effective treatment.",
    "The examination findings indicate isolated saphenous trunk reflux without proximal junction involvement, consistent with CHIVA TYPE 2B. This pattern suggests incompetent perforators with secondary saphenous involvement. Treatment targeting the primary incompetent perforators typically resolves symptoms.",
]

QWEN_RESPONSES_V1 = [
    "Analysis of the duplex data: EP N1→N2 present indicates SFJ incompetence. RP N2→N1 without distal RP flows points to TYPE 1 classification. In TYPE 1, the GSV serves as the primary reflux conduit with a competent lower extremity. Recommended management: ligation of the GSV at the SFJ level to eliminate the proximal incompetence source.",
    "The duplex pattern with NO EP N1→N2 confirms SFJ competence. EP N2→N3 with RP N3→N2 indicates secondary varicosity from incompetent tributaries. This satisfies criteria for TYPE 2A. Management approach: selective tributary ligation or endovenous ablation at sites of perforator incompetence.",
    "Diagnostic criteria met: No EP N1→N2, EP N2→N2 present, RP N3 without proximal RP N2→N1. Consistent with TYPE 2B classification. The varicosity is driven by incompetent perforators. Conservative management with graduated compression and selective intervention on symptomatic perforators is indicated.",
]

QWEN_RESPONSES_V2 = [
    "Clinical examination reveals saphenofemoral junction compromise permitting hemodynamic reflux into the GSV. The retrograde flow pattern documents reflux transit through the saphenous system back to deep veins. This physiologic pattern defines CHIVA TYPE 1. Therapeutic goal: eliminate the reflux source through targeted SFJ intervention while preserving saphenous function where possible.",
    "Patient's hemodynamic picture shows intact proximal drainage pathways with distal venous incompetence manifesting as varicose tributaries. The perforation through incompetent communicating veins drives this secondary pattern characteristic of TYPE 2A. Intervention strategy: address incompetent perforating veins through minimally invasive techniques.",
    "The vascular anatomy demonstrates retained proximal competence with selective saphenous involvement secondary to incompetent perforating veins. This configuration defines TYPE 2B disease. Management approach: conservative initial therapy with progression to targeted perforator intervention based on clinical response.",
]

LIGATION_RESPONSES_V1 = [
    "For TYPE 1 management: Perform saphenofemoral junction ligation with careful dissection preserving the saphenous vein distal to the ligation point if vein quality permits. Division of all tributaries within 5cm of the junction is mandatory. Post-operative compression for 2-3 weeks optimizes outcomes. Consider postoperative duplex surveillance to confirm hemodynamic correction.",
    "For TYPE 2A management: Identify the incompetent perforating vein using intraoperative duplex guidance. Perform selective fasciotomy or subfascial ligation at the perforator site. Preserve the saphenous vein. GSV can be left in situ if hemodynamically favorable. Compression stockings for 3-4 weeks post-operatively.",
    "For TYPE 2B management: Conservative approach recommended initially with compression therapy. If intervention needed, perform selective perforator ligation using mini-invasive technique with fasciotomy. Preserve the saphenous system. Consider staged approach with future intervention based on symptom progression.",
]

LIGATION_RESPONSES_V2 = [
    "TYPE 1 ligation strategy: Expose the SFJ through an oblique groin incision. Identify and preserve the epigastric vessels and external pudendal vein. Ligate and divide all tributaries entering within 5cm of the junction. Ligate the GSV proximal to its junction but preserve the distal GSV for potential future use. Post-operative duplex surveillance at 1-2 weeks.",
    "TYPE 2A perforator management: Use intraoperative B-mode ultrasound to localize the incompetent perforator precisely. Make a small incision directly over the perforator site. Perform subfascial ligation of the perforator vein. Minimize injury to surrounding tissues. Early ambulation and compression therapy facilitate recovery.",
    "TYPE 2B approach: Begin with maximal compression therapy and lifestyle modifications. Reserve surgical intervention for refractory symptoms. If surgery indicated: selective perforator ligation using minimally invasive subfascial endoscopic perforator surgery (SEPS) technique. Avoid aggressive saphenous vein manipulation. Staged approach with reassessment at 3-6 months.",
]

# ============================================================
# RUN EVALUATION
# ============================================================

print("\n[1/3] Loading test cases...")
test_cases = load_test_cases()
print(f"  Loaded {len(test_cases)} test cases")

print("\n[2/3] Generating mock responses...")

results = {
    'metadata': {
        'timestamp': datetime.now().isoformat(),
        'test_count': len(test_cases),
        'models': ['LLAMA 70B (via Groq API)', 'Qwen2.5-7B V2 (Local LoRA)'],
        'queries': ['V1 (Raw clips)', 'V2 (Medical NL)'],
        'mode': 'DEMO - Mock responses for evaluation framework'
    },
    'shunt_classification': [],
    'ligation': []
}

for idx, test in enumerate(test_cases, 1):
    print(f"  [{idx}/{len(test_cases)}] {test['name']}", end=" ", flush=True)

    clips = test['clips']
    v1_q = convert_clips_to_v1(clips)
    v2_q = convert_clips_to_v2(clips)

    # Use cyclic responses
    resp_idx = (idx - 1) % 3

    # Classification
    llama_c_v1 = LLAMA_RESPONSES_V1[resp_idx]
    llama_c_v2 = LLAMA_RESPONSES_V2[resp_idx]
    qwen_c_v1 = QWEN_RESPONSES_V1[resp_idx]
    qwen_c_v2 = QWEN_RESPONSES_V2[resp_idx]

    # Ligation
    llama_l_v1 = LIGATION_RESPONSES_V1[resp_idx]
    llama_l_v2 = LIGATION_RESPONSES_V2[resp_idx]
    qwen_l_v1 = LIGATION_RESPONSES_V1[(resp_idx + 1) % 3]
    qwen_l_v2 = LIGATION_RESPONSES_V2[(resp_idx + 1) % 3]

    results['shunt_classification'].append({
        'test': test['name'],
        'llama_v1': llama_c_v1, 'llama_v2': llama_c_v2,
        'qwen_v1': qwen_c_v1, 'qwen_v2': qwen_c_v2
    })

    results['ligation'].append({
        'test': test['name'],
        'llama_v1': llama_l_v1, 'llama_v2': llama_l_v2,
        'qwen_v1': qwen_l_v1, 'qwen_v2': qwen_l_v2
    })

    print("OK")

# ============================================================
# GENERATE REPORT
# ============================================================

print("\n[3/3] Generating Word report...")

doc = Document()

# Title
title = doc.add_heading('COMPARATIVE EVALUATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('LLAMA 70B Versatile vs Qwen2.5-7B V2 Fine-tuned')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)

# Mode note
mode_para = doc.add_paragraph(f"Mode: {results['metadata']['mode']}")
mode_para.runs[0].font.italic = True

# Metadata
doc.add_heading('Evaluation Metadata', level=1)
doc.add_paragraph(f"Timestamp: {results['metadata']['timestamp']}")
doc.add_paragraph(f"Test Cases: {results['metadata']['test_count']}")
doc.add_paragraph(f"Models: {', '.join(results['metadata']['models'])}")
doc.add_paragraph(f"Query Formats: {', '.join(results['metadata']['queries'])}")

# Results - Shunt Classification
doc.add_page_break()
doc.add_heading('Task 1: Shunt Classification (No RAG)', level=1)

for result in results['shunt_classification']:
    doc.add_heading(f"Test Case: {result['test']}", level=2)

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'

    hdr = table.rows[0].cells
    hdr[0].text = 'Model'
    hdr[1].text = 'V1 Query (Raw Clips)'
    hdr[2].text = 'V2 Query (Medical NL)'

    row1 = table.rows[1].cells
    row1[0].text = 'LLAMA 70B'
    row1[1].text = result['llama_v1'][:300] + "..." if len(result['llama_v1']) > 300 else result['llama_v1']
    row1[2].text = result['llama_v2'][:300] + "..." if len(result['llama_v2']) > 300 else result['llama_v2']

    row2 = table.rows[2].cells
    row2[0].text = 'Qwen2.5-7B V2'
    row2[1].text = result['qwen_v1'][:300] + "..." if len(result['qwen_v1']) > 300 else result['qwen_v1']
    row2[2].text = result['qwen_v2'][:300] + "..." if len(result['qwen_v2']) > 300 else result['qwen_v2']

    doc.add_paragraph()

# Results - Ligation Planning
doc.add_page_break()
doc.add_heading('Task 2: Ligation Planning (with RAG)', level=1)

for result in results['ligation']:
    doc.add_heading(f"Test Case: {result['test']}", level=2)

    table = doc.add_table(rows=3, cols=3)
    table.style = 'Light Grid Accent 1'

    hdr = table.rows[0].cells
    hdr[0].text = 'Model'
    hdr[1].text = 'V1 Query (Raw Clips)'
    hdr[2].text = 'V2 Query (Medical NL)'

    row1 = table.rows[1].cells
    row1[0].text = 'LLAMA 70B'
    row1[1].text = result['llama_v1'][:300] + "..." if len(result['llama_v1']) > 300 else result['llama_v1']
    row1[2].text = result['llama_v2'][:300] + "..." if len(result['llama_v2']) > 300 else result['llama_v2']

    row2 = table.rows[2].cells
    row2[0].text = 'Qwen2.5-7B V2'
    row2[1].text = result['qwen_v1'][:300] + "..." if len(result['qwen_v1']) > 300 else result['qwen_v1']
    row2[2].text = result['qwen_v2'][:300] + "..." if len(result['qwen_v2']) > 300 else result['qwen_v2']

    doc.add_paragraph()

# Evaluation Tables (Manual completion)
doc.add_page_break()
doc.add_heading('Evaluation Scoring (To Be Completed)', level=1)

doc.add_paragraph("Score each response on a scale of 1-5 for accuracy and reasoning quality.")

doc.add_heading('Table 1: Shunt Classification', level=2)
t1 = doc.add_table(rows=3, cols=5)
t1.style = 'Light Grid Accent 1'
h = t1.rows[0].cells
h[0].text, h[1].text, h[2].text, h[3].text, h[4].text = 'Model', 'Accuracy V1', 'Accuracy V2', 'Reasoning V1', 'Reasoning V2'
r1, r2 = t1.rows[1].cells, t1.rows[2].cells
r1[0].text, r1[1].text, r1[2].text, r1[3].text, r1[4].text = 'LLAMA 70B', '[FILL]', '[FILL]', '[FILL]', '[FILL]'
r2[0].text, r2[1].text, r2[2].text, r2[3].text, r2[4].text = 'Qwen2.5-7B V2', '[FILL]', '[FILL]', '[FILL]', '[FILL]'

doc.add_paragraph()

doc.add_heading('Table 2: Ligation Planning', level=2)
t2 = doc.add_table(rows=3, cols=5)
t2.style = 'Light Grid Accent 1'
h = t2.rows[0].cells
h[0].text, h[1].text, h[2].text, h[3].text, h[4].text = 'Model', 'Quality V1', 'Quality V2', 'Reasoning V1', 'Reasoning V2'
r1, r2 = t2.rows[1].cells, t2.rows[2].cells
r1[0].text, r1[1].text, r1[2].text, r1[3].text, r1[4].text = 'LLAMA 70B', '[FILL]', '[FILL]', '[FILL]', '[FILL]'
r2[0].text, r2[1].text, r2[2].text, r2[3].text, r2[4].text = 'Qwen2.5-7B V2', '[FILL]', '[FILL]', '[FILL]', '[FILL]'

# Summary
doc.add_page_break()
doc.add_heading('Summary', level=1)
doc.add_paragraph(f"Total test cases evaluated: {len(test_cases)}")
doc.add_paragraph(f"Models compared: LLAMA 70B (via Groq), Qwen2.5-7B V2 (Local LoRA)")
doc.add_paragraph(f"Query formats: V1 (raw clip notation) and V2 (natural language medical descriptions)")
doc.add_paragraph("This report includes full model responses for manual scoring and comparison.")

doc.save('evaluation_report.docx')

print("\n" + "="*80)
print("EVALUATION COMPLETE")
print("="*80)
print("\nReport saved: evaluation_report.docx")
print(f"Test cases: {len(test_cases)}")
print(f"Timestamp: {results['metadata']['timestamp']}")
print("\nNOTE: This is a DEMO VERSION with mock responses.")
print("For actual evaluation, ensure:")
print("  1. Groq API key is valid and models are available")
print("  2. PyTorch with CUDA support is installed: pip install torch torchvision torchaudio --index-url https://download.pytorch.org/whl/cu118")
print("  3. Sufficient system memory for model loading")
