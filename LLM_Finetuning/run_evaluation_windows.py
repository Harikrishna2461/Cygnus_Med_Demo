#!/usr/bin/env python3
"""
WINDOWS EVALUATION - Load model from HuggingFace and run full evaluation
"""

import json
import torch
from pathlib import Path
from datetime import datetime
from typing import Dict, List
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

print("="*80)
print("LOADING MODEL FROM HUGGINGFACE (WINDOWS)")
print("="*80)

# Configuration
HF_TOKEN = ""
MODEL_REPO = "HariKrishna1824/qwen_chiva_vericose_veins_treatment_finetuned"
TEST_CASES_DIR = Path("json samples")

CHIVA_RULES = """
CHIVA_RULES :
=== CHIVA VENOUS SHUNT CLASSIFICATION RULES ===

ANATOMY:
    N1 = Deep venous system (femoral / popliteal vein)
    N2 = Great Saphenous Vein (GSV) or Small Saphenous Vein (SSV) trunk
    N3 = Tributaries / superficial branches
    EP = Physiological (forward, antegrade) flow — NORMAL clip
    RP = Retrograde (pathological, reflux) flow — ABNORMAL clip
    SFJ = Saphenofemoral Junction  →  posYRatio ≤ 0.098
    Hunterian Perforator            →  0.098 < posYRatio ≤ 0.353

CRITICAL RULE — SFJ COMPETENCE:
    SFJ is INCOMPETENT if and only if a clip has fromType=N1 AND toType=N2 (EP N1→N2).
    EP N2→N2 means blood circulates within the saphenous trunk via a perforator — SFJ REMAINS COMPETENT.
    This is true regardless of posYRatio or step label.

STEP 1 — CHECK FOR EP N1→N2:
    Scan ALL clips. Does any clip have flow=EP, fromType=N1, toType=N2?
    YES → SFJ/Hunterian INCOMPETENT → go to Case A or B.
    NO  → SFJ COMPETENT → go to Case C.

Case A — EP N1→N2 EXISTS, NO EP N2→N3:
    If RP N2→N1 present AND no RP at N3 → TYPE 1

Case B — EP N1→N2 EXISTS AND EP N2→N3 EXISTS:
    B1: RP N3→N2 or RP N3→N1, NO RP N2→N1 → TYPE 3
    B4: RP N3→N1 AND RP N2→N1, elim="Reflux" → TYPE 1+2
    B5: RP N3→N1 AND RP N2→N1, elim="No Reflux" → TYPE 3

Case C — NO EP N1→N2 (SFJ COMPETENT):
    TYPE 2A: EP N2→N3, NO EP N1→N2
    TYPE 2B: EP N2→N2, NO EP N1→N2, RP at N3, NO RP N2→N1
    TYPE 2C: EP N2→N2, NO EP N1→N2, RP at N3, RP N2→N1 present

QUICK DECISION TABLE:
    Has EP N1→N2? YES + no EP N2→N3 + RP N2→N1 → TYPE 1
    Has EP N1→N2? YES + EP N2→N3 + RP N3 only → TYPE 3
    No EP N1→N2 + EP N2→N3 → TYPE 2A
    No EP N1→N2 + EP N2→N2 + RP N3 + NO RP N2→N1 → TYPE 2B
    No EP N1→N2 + EP N2→N2 + RP N3 + RP N2→N1 → TYPE 2C
"""

RAG_DOCS = [
    "For TYPE 1 shunt: Ligate GSV at saphenofemoral junction, preserve distal GSV if quality permits",
    "For TYPE 2A shunt: Selective perforator ligation with duplex guidance",
    "For TYPE 2B shunt: Conservative therapy with compression; consider staged intervention",
    "For TYPE 2C shunt: Evaluate SFJ status; may require combined intervention",
    "For TYPE 3 shunt: Ligate incompetent tributaries at their junction with GSV/SSV",
    "For TYPE 1+2 shunt: Combined approach - SFJ ligation plus selective perforator ligation",
]

def format_clips_v1(clips: List[Dict]) -> str:
    """V1: Natural language with y-values"""
    descriptions = []
    for clip in clips:
        flow = "physiological forward flow" if clip.get('flow') == 'EP' else "retrograde reflux"
        from_to = f"from {clip.get('fromType')} to {clip.get('toType')}"
        y_val = clip.get('posYRatio', 0)
        descriptions.append(f"There is a {flow} (y={y_val:.3f}) {from_to}")
    return "Duplex findings:\n" + "\n".join(descriptions)

def format_clips_v2(clips: List[Dict]) -> str:
    """V2: Full medical terminology"""
    findings = []
    for clip in clips:
        flow = clip.get('flow')
        from_type = clip.get('fromType')
        to_type = clip.get('toType')

        if flow == 'EP' and from_type == 'N1' and to_type == 'N2':
            findings.append("saphenofemoral junction incompetence")
        elif flow == 'EP' and from_type == 'N2' and to_type == 'N3':
            findings.append("perforator draining from saphenous trunk")
        elif flow == 'RP' and from_type == 'N2' and to_type == 'N1':
            findings.append("retrograde reflux in saphenous trunk")
        elif flow == 'RP' and from_type == 'N3':
            findings.append("retrograde flow in tributaries")

    return "Patient with varicose veins. Duplex findings: " + ", ".join(findings) + "."

print("\n[1/4] Loading model from HuggingFace...")
print(f"  Model: {MODEL_REPO}")
print(f"  Device: CPU (will auto-convert if CUDA available)")

try:
    from transformers import AutoModelForCausalLM, AutoTokenizer

    print("  Downloading model weights...", end=" ", flush=True)
    model = AutoModelForCausalLM.from_pretrained(
        MODEL_REPO,
        torch_dtype="auto",
        device_map="auto",
        token=HF_TOKEN,
        trust_remote_code=True
    )
    print("OK")

    print("  Loading tokenizer...", end=" ", flush=True)
    # Use base model tokenizer to avoid config issues
    tokenizer = AutoTokenizer.from_pretrained(
        "Qwen/Qwen2.5-7B",
        trust_remote_code=True
    )
    print("OK")

    model.eval()
    device = next(model.parameters()).device
    print(f"  Model loaded on: {device}")

except Exception as e:
    print(f"\n  ERROR: {str(e)}")
    import sys
    sys.exit(1)

print("\n[2/4] Loading test cases...")
test_cases = []
if TEST_CASES_DIR.exists():
    for json_file in sorted(TEST_CASES_DIR.glob("*.json")):
        try:
            with open(json_file, 'r', encoding='utf-8', errors='ignore') as f:
                data = json.load(f)
                clips = data.get('clips', [])
                if clips:
                    test_cases.append({
                        'name': json_file.stem,
                        'clips': clips,
                    })
        except Exception as e:
            print(f"  Warning: {json_file.name}")

print(f"  Loaded {len(test_cases)} test cases")

def infer(prompt: str, max_tokens: int = 500) -> str:
    try:
        inputs = tokenizer(prompt, return_tensors="pt").to(device)
        with torch.no_grad():
            outputs = model.generate(
                **inputs,
                max_new_tokens=max_tokens,
                do_sample=False,
                pad_token_id=tokenizer.eos_token_id
            )
        return tokenizer.decode(outputs[0][inputs['input_ids'].shape[1]:], skip_special_tokens=True)
    except Exception as e:
        return f"[Error: {str(e)[:100]}]"

print("\n[3/4] Running evaluation...")

results = {
    'timestamp': datetime.now().isoformat(),
    'model': MODEL_REPO,
    'test_count': len(test_cases),
    'evaluations': []
}

for idx, test in enumerate(test_cases, 1):
    print(f"  [{idx}/{len(test_cases)}] {test['name']}", end=" ", flush=True)

    clips = test['clips']
    v1_prompt = format_clips_v1(clips)
    v2_prompt = format_clips_v2(clips)

    print("(c)", end=" ", flush=True)
    c_v1 = infer(f"{CHIVA_RULES}\n\n{v1_prompt}\n\nClassify the CHIVA shunt type.")
    c_v2 = infer(f"{CHIVA_RULES}\n\n{v2_prompt}\n\nWhat is the CHIVA shunt type?")

    print("(l)", end=" ", flush=True)
    rag_context = "\n".join(RAG_DOCS)
    l_v1 = infer(f"{CHIVA_RULES}\n\n{rag_context}\n\n{v1_prompt}\n\nProvide ligation planning.")
    l_v2 = infer(f"{CHIVA_RULES}\n\n{rag_context}\n\n{v2_prompt}\n\nProvide ligation planning.")

    results['evaluations'].append({
        'test_name': test['name'],
        'v1_classification': c_v1,
        'v2_classification': c_v2,
        'v1_ligation': l_v1,
        'v2_ligation': l_v2,
    })

    print("OK")

print("\n[4/4] Generating Word report...")

doc = Document()

title = doc.add_heading('CHIVA EVALUATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph(f'Model: {MODEL_REPO}')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)

doc.add_heading('Evaluation Details', level=1)
doc.add_paragraph(f"Timestamp: {results['timestamp']}")
doc.add_paragraph(f"Test Cases: {results['test_count']}")
doc.add_paragraph(f"Model: Qwen2.5-7B + CHIVA LoRA (Merged)")

doc.add_page_break()
doc.add_heading('TASK 1: SHUNT CLASSIFICATION (NO RAG)', level=1)

for result in results['evaluations']:
    doc.add_heading(f"Test Case: {result['test_name']}", level=2)

    doc.add_heading("V1 Query (Natural Language + Y-values)", level=3)
    doc.add_paragraph(result['v1_classification'])

    doc.add_heading("V2 Query (Medical Terminology)", level=3)
    doc.add_paragraph(result['v2_classification'])

    doc.add_paragraph()

doc.add_page_break()
doc.add_heading('TASK 2: LIGATION PLANNING (WITH RAG)', level=1)

for result in results['evaluations']:
    doc.add_heading(f"Test Case: {result['test_name']}", level=2)

    doc.add_heading("V1 with RAG Context", level=3)
    doc.add_paragraph(result['v1_ligation'])

    doc.add_heading("V2 with RAG Context", level=3)
    doc.add_paragraph(result['v2_ligation'])

    doc.add_paragraph()

doc.save('CHIVA_EVALUATION_FINAL_REPORT.docx')

with open('evaluation_results_final.json', 'w') as f:
    json.dump(results, f, indent=2)

print("\n" + "="*80)
print("EVALUATION COMPLETE")
print("="*80)
print("\nFiles generated:")
print("  - CHIVA_EVALUATION_FINAL_REPORT.docx")
print("  - evaluation_results_final.json")
print(f"\nTest cases evaluated: {len(test_cases)}")
print(f"Timestamp: {results['timestamp']}")
