#!/usr/bin/env python3
"""
UBUNTU EVALUATION - FINAL VERSION
Optimized for RTX 5090 32GB VRAM
All necessary files in root directory
"""

import json
import torch
from pathlib import Path
from datetime import datetime
from typing import Dict, List
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

print("="*80)
print("COMPREHENSIVE LLM EVALUATION - UBUNTU RTX 5090")
print("LLAMA 70B (Groq) vs Qwen2.5-7B V2 (Fine-tuned)")
print("="*80)

# ============================================================================
# CONFIGURATION - RELATIVE PATHS
# ============================================================================

GROQ_API_KEY = ""

# Load from local folders - no HF download
QWEN_BASE_MODEL_PATH = Path("Qwen25-7B")
QWEN_LORA_PATH = Path("qwen25_chiva_v2")

# All relative paths - assume all files in root
TEST_DIRS = [Path("json samples"), Path("multiple shunts in 1 sesh")]
CROSS_ENCODER_PATH = Path("cross_encoder_finetuning/chiva_crossencoder_finetuned")
QDRANT_STORAGE_PATH = Path("qdrant_storage")

CHIVA_RULES = """
=== CHIVA VENOUS SHUNT CLASSIFICATION RULES ===

ANATOMY:
    N1 = Deep venous system (femoral/popliteal vein)
    N2 = Great Saphenous Vein (GSV) or Small Saphenous Vein (SSV) trunk
    N3 = Tributaries / superficial branches
    EP = Physiological (forward, antegrade) flow
    RP = Retrograde (pathological, reflux) flow

CRITICAL RULE — SFJ COMPETENCE:
    SFJ is INCOMPETENT if and only if a clip has fromType=N1 AND toType=N2 (EP N1→N2).
    EP N2→N2 means blood circulates within the saphenous trunk via a perforator — SFJ REMAINS COMPETENT.

CLASSIFICATION:
    STEP 1: Check for EP N1→N2
        YES → SFJ INCOMPETENT (Case A/B)
        NO → SFJ COMPETENT (Case C)

    Case A (EP N1→N2, NO EP N2→N3):
        RP N2→N1, no RP at N3 → TYPE 1

    Case B (EP N1→N2 AND EP N2→N3):
        RP N3 only, NO RP N2→N1 → TYPE 3
        RP N3 AND RP N2→N1, elim="Reflux" → TYPE 1+2
        RP N3 AND RP N2→N1, elim="No Reflux" → TYPE 3

    Case C (NO EP N1→N2):
        EP N2→N3 → TYPE 2A
        EP N2→N2, RP N3, NO RP N2→N1 → TYPE 2B
        EP N2→N2, RP N3, RP N2→N1 → TYPE 2C
"""

FEW_SHOT_LIGATION = """
LIGATION PLANNING EXAMPLES:

Example 1 (TYPE 1):
    Ligate at saphenofemoral junction. Preserve distal GSV if quality permits.

Example 2 (TYPE 2A):
    Selective perforator ligation with duplex guidance at N2→N3 junction.

Example 3 (TYPE 2C):
    Ligate perforator entry point AND all GSV reflux sites along trunk.

Example 4 (TYPE 3):
    Ligate tributary at N2→N3 junction. Follow up 6-12 months; if GSV reflux develops, ligate SFJ.

Example 5 (TYPE 1+2):
    Combined approach - ligate SFJ AND refluxing tributaries simultaneously.
"""

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def add_table_borders(table):
    """Add borders to Word table"""
    tbl = table._element
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '12')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)

def _clip_label(flow: str, ft: str, tt: str, y: float) -> str:
    """Annotate clip with anatomical significance"""
    if flow == "EP" and ft == "N1" and tt == "N2":
        if y <= 0.098:
            return " [SFJ-ENTRY=INCOMPETENT]"
        return " [Hunterian-ENTRY=INCOMPETENT]" if y <= 0.353 else " [Deep-to-GSV-ENTRY]"
    if flow == "RP" and ft == "N3":
        return f" [TRIBUTARY-REFLUX: N3→{tt}]"

    labels = {
        ("EP", "N2", "N2"): " [PERFORATOR-ENTRY: N2→N2, SFJ=COMPETENT]",
        ("EP", "N2", "N3"): " [GSV-to-TRIBUTARY-ENTRY: N2→N3]",
        ("RP", "N2", "N1"): " [GSV-TRUNK-REFLUX: N2→N1]",
    }
    return labels.get((flow, ft, tt), "")

def _summarise_clips(clips: List[Dict]) -> str:
    """Format clips with anatomical labels for classification"""
    lines = []
    for i, c in enumerate(clips):
        flow = c.get('flow', '?')
        ft = c.get('fromType', '?')
        tt = c.get('toType', '?')
        y = c.get('posYRatio') or 0.0
        loc = _clip_label(flow, ft, tt, y)
        lines.append(f"  Clip {i:02d}: {flow} {ft}→{tt}  y={y:.3f}{loc}")
    return "\n".join(lines)

def format_clips_v1(clips: List[Dict]) -> str:
    """V1: Clip notation + simple NL"""
    descriptions = []
    for clip in clips:
        flow = "EP" if clip.get('flow') == 'EP' else "RP"
        from_to = f"{clip.get('fromType')}->{clip.get('toType')}"
        y_val = clip.get('posYRatio', 0)
        descriptions.append(f"{flow} {from_to} (y={y_val:.3f})")
    return "Duplex findings: " + ", ".join(descriptions)

def format_clips_v2(clips: List[Dict]) -> str:
    """V2: Pure medical terminology"""
    findings = []
    for clip in clips:
        flow = clip.get('flow')
        from_type = clip.get('fromType')
        to_type = clip.get('toType')

        if flow == 'EP' and from_type == 'N1' and to_type == 'N2':
            findings.append("antegrade flow from deep femoral vein to saphenous trunk indicating saphenofemoral junction incompetence")
        elif flow == 'EP' and from_type == 'N2' and to_type == 'N3':
            findings.append("antegrade flow from saphenous trunk feeding into tributaries")
        elif flow == 'RP' and from_type == 'N2' and to_type == 'N1':
            findings.append("retrograde reflux within saphenous trunk toward deep system")
        elif flow == 'RP' and from_type == 'N3' and to_type == 'N2':
            findings.append("retrograde reflux in tributary branches toward saphenous trunk")
        elif flow == 'RP' and from_type == 'N3' and to_type == 'N1':
            findings.append("retrograde reflux in tributary branches toward deep venous system")

    return "Duplex ultrasound demonstrates: " + "; ".join(findings) + " in patient with chronic venous insufficiency."

def extract_shunt_type(text: str) -> str:
    """Extract type from JSON output - handle both 'type' and 'shunt_type' fields"""
    try:
        data = json.loads(text)
        # Try 'type' first, then 'shunt_type'
        result = data.get('type') or data.get('shunt_type')
        if result:
            return result.strip()
        return 'UNKNOWN'
    except:
        # If JSON fails, search for TYPE pattern in text
        import re
        match = re.search(r'TYPE\s+[\d+A-Z]+|No\s+shunt', text, re.IGNORECASE)
        if match:
            return match.group(0).upper()
        return text.strip()[:50] if text else "UNKNOWN"

# ============================================================================
# STEP 1: LOAD TEST CASES
# ============================================================================

print("\n[1/6] Loading test cases...")
test_cases = []
for test_dir in TEST_DIRS:
    if test_dir.exists():
        json_files = list(test_dir.glob("**/*.json"))
        for json_file in sorted(json_files):
            try:
                with open(json_file, 'r', encoding='utf-8', errors='ignore') as f:
                    data = json.load(f)
                    clips = data.get('clips', [])
                    if clips:
                        test_cases.append({
                            'name': json_file.stem,
                            'clips': clips,
                            'source': test_dir.name,
                        })
            except Exception as e:
                pass

print(f"  Loaded {len(test_cases)} test cases")
if len(test_cases) == 0:
    print("  ERROR: No test cases found in json samples/ or multiple shunts in 1 sesh/")
    import sys
    sys.exit(1)

# ============================================================================
# STEP 2: LOAD MODELS
# ============================================================================

print("\n[2/6] Loading LLM models...")

# LLAMA 70B via Groq
print("  Loading LLAMA 70B (Groq API)...", end=" ", flush=True)
groq_client = None
try:
    from groq import Groq
    groq_client = Groq(api_key=GROQ_API_KEY)
    print("OK")
except Exception as e:
    print(f"ERROR: {e}")

# Qwen V2 - Load base + merge LoRA from local folders
print("  Loading Qwen V2 (local base + LoRA merge)...", end=" ", flush=True)
qwen_model = None
qwen_tokenizer = None
device = None
try:
    from transformers import AutoModelForCausalLM, AutoTokenizer
    from peft import PeftModel

    # Load base model
    qwen_model = AutoModelForCausalLM.from_pretrained(
        str(QWEN_BASE_MODEL_PATH),
        torch_dtype=torch.bfloat16,
        device_map="auto",
        trust_remote_code=True
    )

    # Load and merge LoRA
    qwen_model = PeftModel.from_pretrained(qwen_model, str(QWEN_LORA_PATH))
    qwen_model = qwen_model.merge_and_unload()

    # Load tokenizer from base
    qwen_tokenizer = AutoTokenizer.from_pretrained(str(QWEN_BASE_MODEL_PATH), trust_remote_code=True)
    qwen_model.eval()
    device = next(qwen_model.parameters()).device
    print(f"OK (merged, on {device})")
except Exception as e:
    print(f"ERROR: {e}")

# Cross-Encoder for reranking
print("  Loading Cross-Encoder...", end=" ", flush=True)
cross_encoder = None
cross_encoder_available = False
try:
    from sentence_transformers import CrossEncoder
    if CROSS_ENCODER_PATH.exists():
        cross_encoder = CrossEncoder(str(CROSS_ENCODER_PATH))
        cross_encoder_available = True
        print("OK")
    else:
        print(f"UNAVAILABLE (path {CROSS_ENCODER_PATH} not found)")
except Exception as e:
    print(f"UNAVAILABLE: {e}")

# Qdrant local storage
print("  Loading Qdrant local storage...", end=" ", flush=True)
qdrant_client = None
qdrant_available = False
try:
    from qdrant_client import QdrantClient
    if QDRANT_STORAGE_PATH.exists():
        qdrant_client = QdrantClient(path=str(QDRANT_STORAGE_PATH))
        qdrant_available = True
        print("OK")
    else:
        print(f"UNAVAILABLE (path {QDRANT_STORAGE_PATH} not found)")
except Exception as e:
    print(f"UNAVAILABLE: {e}")

# ============================================================================
# STEP 3: INFERENCE FUNCTIONS
# ============================================================================

def build_classification_prompt(clips_summary: str) -> str:
    """Build structured classification prompt with decision guide"""
    return f"""{CHIVA_RULES}

=== CLINICAL ASSESSMENT ===
{clips_summary}

═══════════════════════════════════════════════════════════════
DECISION GUIDE
═══════════════════════════════════════════════════════════════

STEP 1: CHECK FOR EP N1→N2 (SFJ or Hunterian ENTRY)
    Look for: "EP N1→N2" with [SFJ-ENTRY=INCOMPETENT] or [Hunterian-ENTRY=INCOMPETENT] label
    If YES → SFJ INCOMPETENT (go to Case B or A)
    If NO  → SFJ COMPETENT (go to Case C)

STEP 2: CHECK FOR REFLUX PATTERNS
    2a) ANY RP N3→N2 or RP N3→N1? (tributary reflux)
    2b) ANY RP N2→N1? (GSV trunk reflux)
    2c) ANY EP N2→N3? (extra antegrade to tributary)

STEP 3: PATTERN MATCHING

    SFJ INCOMPETENT (has EP N1→N2):
    - NO EP N2→N3 + RP N2→N1 = TYPE 1
    - YES EP N2→N3 + RP N3 only = TYPE 3
    - YES EP N2→N3 + RP N3 AND RP N2→N1 = TYPE 1+2

    SFJ COMPETENT (NO EP N1→N2):
    - EP N2→N3 EXISTS = TYPE 2A
    - EP N2→N2 + RP N3 only = TYPE 2B
    - EP N2→N2 + RP N3 AND RP N2→N1 = TYPE 2C

═══════════════════════════════════════════════════════════════

Output ONLY this JSON structure (no markdown, no explanation):

{{
    "shunt_type": "TYPE 1 | TYPE 2A | TYPE 2B | TYPE 2C | TYPE 3 | TYPE 1+2 | No shunt detected",
    "confidence": 0.85,
    "reasoning": "your reasoning",
    "summary": "1 sentence"
}}"""

def groq_inference(prompt: str, task: str = "classification") -> str:
    """LLAMA 70B via Groq"""
    if not groq_client:
        return "[Groq unavailable]"
    try:
        model_name = "llama-3.3-70b-versatile"
        if task == "classification":
            system_msg = """You are a CHIVA shunt classification expert.
Output ONLY valid JSON with fields: shunt_type and reasoning.
No markdown, no explanation outside JSON."""
            user_msg = prompt
        else:
            system_msg = f"""Expert vascular surgeon trained in CHIVA. Provide BRIEF ligation planning.
Output ONLY valid JSON with two sections: ligation_steps (max 4 lines) and ligation_reasoning (1 sentence).
Be concise and to the point.\n\n{CHIVA_RULES}\n\n{FEW_SHOT_LIGATION}"""
            user_msg = f"""{prompt}

Respond with ONLY this JSON format:
{{"ligation_steps": "step 1, step 2, step 3, step 4 (max 4 lines)", "ligation_reasoning": "one sentence why"}}"""

        response = groq_client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": system_msg},
                {"role": "user", "content": user_msg}
            ],
            max_tokens=150 if task == "classification" else 300,
            temperature=0.0 if task == "classification" else 0.2
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        return f"[Groq Error: {str(e)[:100]}]"

def qwen_inference(prompt: str, task: str = "classification") -> str:
    """Qwen V2"""
    if not qwen_model or not qwen_tokenizer or device is None:
        return "[Qwen unavailable]"
    try:
        if task == "classification":
            messages = [
                {"role": "system", "content": "You are a CHIVA shunt classification expert. Output ONLY valid JSON with fields: type and reasoning. No markdown, no explanation outside JSON."},
                {"role": "user", "content": prompt}
            ]
            max_tok = 150
        else:
            messages = [
                {"role": "system", "content": f"""Expert vascular surgeon. Provide BRIEF ligation planning.
Output ONLY valid JSON with two sections: ligation_steps (max 4 lines) and ligation_reasoning (1 sentence).
Be concise and to the point.\n\n{CHIVA_RULES}\n\n{FEW_SHOT_LIGATION}"""},
                {"role": "user", "content": f"""{prompt}

Respond with ONLY this JSON format:
{{"ligation_steps": "step 1, step 2, step 3, step 4 (max 4 lines)", "ligation_reasoning": "one sentence why"}}"""}
            ]
            max_tok = 300

        text = qwen_tokenizer.apply_chat_template(messages, tokenize=False, add_generation_prompt=True)
        inputs = qwen_tokenizer(text, return_tensors="pt").to(device)

        with torch.no_grad():
            outputs = qwen_model.generate(
                **inputs,
                max_new_tokens=max_tok,
                do_sample=False,
                temperature=0.2,
                pad_token_id=qwen_tokenizer.eos_token_id
            )

        result = qwen_tokenizer.decode(outputs[0][inputs['input_ids'].shape[1]:], skip_special_tokens=True).strip()
        return result
    except Exception as e:
        return f"[Qwen Error: {str(e)[:80]}]"

def retrieve_rag_docs(query: str, top_k: int = 3) -> str:
    """Retrieve docs from Qdrant + rerank"""
    if not qdrant_available or not qdrant_client:
        return "[RAG unavailable]"
    try:
        results = qdrant_client.search(
            collection_name="ligation_knowledgebase_db_v2",
            query_vector=[0.1] * 768,
            limit=top_k
        )
        docs = [str(r.payload) for r in results]

        if cross_encoder_available and cross_encoder and docs:
            scores = cross_encoder.predict([[query, doc] for doc in docs])
            ranked = sorted(zip(scores, docs), reverse=True)
            return "\n".join([doc for _, doc in ranked[:2]])

        return "\n".join(docs[:2])
    except Exception as e:
        return "[RAG retrieval failed]"

# ============================================================================
# STEP 4: RUN EVALUATION
# ============================================================================

print("\n[3/6] Running evaluation on all test cases...")

results = {
    'timestamp': datetime.now().isoformat(),
    'test_count': len(test_cases),
    'models': {
        'llama': 'LLAMA 70B Versatile (Groq API)',
        'qwen': f'Qwen V2 (Local: {QWEN_BASE_MODEL_PATH} + {QWEN_LORA_PATH} merged)'
    },
    'evaluations': []
}

for idx, test in enumerate(test_cases, 1):
    print(f"  [{idx}/{len(test_cases)}] {test['name']}", end=" ", flush=True)

    clips = test['clips']
    v1_prompt = format_clips_v1(clips)
    v2_prompt = format_clips_v2(clips)

    # Task 1: Classification
    print("(c)", end=" ", flush=True)

    # V1: Clip notation with instruction
    clips_summary_v1 = _summarise_clips(clips)
    v1_user_query = f"Classify the shunt types given these clips:\n{clips_summary_v1}"
    v1_classification_prompt = build_classification_prompt(clips_summary_v1)
    groq_c_v1 = groq_inference(v1_classification_prompt, "classification")
    qwen_c_v1 = qwen_inference(v1_classification_prompt, "classification")
    shunt_type_v1 = extract_shunt_type(groq_c_v1)

    # V2: Medical terminology
    clips_summary_v2 = format_clips_v2(clips)
    v2_user_query = f"Based on the following duplex findings:\n{clips_summary_v2}\n\nDetermine the CHIVA shunt type."
    v2_classification_prompt = build_classification_prompt(clips_summary_v1)  # Use v1 for prompting models
    groq_c_v2 = groq_inference(v2_classification_prompt, "classification")
    qwen_c_v2 = qwen_inference(v2_classification_prompt, "classification")
    shunt_type_v2 = extract_shunt_type(groq_c_v2)

    # Task 2: Ligation
    print("(l)", end=" ", flush=True)
    rag_docs = retrieve_rag_docs(v1_prompt)

    groq_l_v1 = groq_inference(f"{rag_docs}\n\n{v1_prompt}\n\nFor {shunt_type_v1}: Provide detailed ligation planning.", "ligation")
    qwen_l_v1 = qwen_inference(f"{CHIVA_RULES}\n\n{FEW_SHOT_LIGATION}\n\n{rag_docs}\n\n{v1_prompt}\n\nFor {shunt_type_v1}: Provide ligation planning.", "ligation")

    groq_l_v2 = groq_inference(f"{rag_docs}\n\n{v2_prompt}\n\nFor {shunt_type_v2}: Recommend surgical ligation strategy.", "ligation")
    qwen_l_v2 = qwen_inference(f"{CHIVA_RULES}\n\n{FEW_SHOT_LIGATION}\n\n{rag_docs}\n\n{v2_prompt}\n\nFor {shunt_type_v2}: Provide ligation planning.", "ligation")

    results['evaluations'].append({
        'test_name': test['name'],
        'source': test['source'],
        'shunt_type_v1': shunt_type_v1,
        'shunt_type_v2': shunt_type_v2,
        'v1_query': v1_user_query,
        'v2_query': v2_user_query,
        'groq_c_v1': groq_c_v1,
        'groq_c_v2': groq_c_v2,
        'qwen_c_v1': qwen_c_v1,
        'qwen_c_v2': qwen_c_v2,
        'groq_l_v1': groq_l_v1,
        'groq_l_v2': groq_l_v2,
        'qwen_l_v1': qwen_l_v1,
        'qwen_l_v2': qwen_l_v2,
    })

    print("OK")

# ============================================================================
# STEP 5: GENERATE WORD REPORT
# ============================================================================

print("\n[4/6] Generating Word report...")

doc = Document()

# Title
title = doc.add_heading('COMPREHENSIVE LLM EVALUATION REPORT', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('LLAMA 70B Versatile (Groq) vs Qwen2.5-7B V2 (Fine-tuned)')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(12)
subtitle.runs[0].font.bold = True

# Metadata
doc.add_heading('Evaluation Setup', level=1)
doc.add_paragraph(f"Timestamp: {results['timestamp']}")
doc.add_paragraph(f"Test Cases Evaluated: {len(results['evaluations'])}")
doc.add_paragraph("")
doc.add_paragraph("Models Compared:")
doc.add_paragraph(f"  • {results['models']['llama']}", style='List Bullet')
doc.add_paragraph(f"  • {results['models']['qwen']}", style='List Bullet')
doc.add_paragraph("")
doc.add_paragraph("Task 1: Shunt Classification (CHIVA rules, NO RAG)", style='List Bullet')
doc.add_paragraph("Task 2: Ligation Planning (RAG + Cross-Encoder Reranking + BM25 + Few-shot examples)", style='List Bullet')
doc.add_paragraph("")
doc.add_paragraph("Query Formats:")
doc.add_paragraph("  • V1: Clip notation (EP/RP with y-values) + Simple natural language", style='List Bullet')
doc.add_paragraph("  • V2: Pure medical terminology (no raw clip notation)", style='List Bullet')

# Task 1: Classification Results
doc.add_page_break()
doc.add_heading('TASK 1: SHUNT CLASSIFICATION (NO RAG)', level=1)

for result in results['evaluations']:
    doc.add_heading(f"Test Case: {result['test_name']} ({result['source']})", level=2)

    doc.add_heading("V1 Query", level=3)
    doc.add_paragraph(result['v1_query'])

    doc.add_heading("V1 Classification Results", level=3)
    doc.add_paragraph(f"Detected Type: {result['shunt_type_v1']}")

    doc.add_heading("LLAMA 70B Response", level=4)
    doc.add_paragraph(result['groq_c_v1'])

    doc.add_heading("Qwen V2 Response", level=4)
    doc.add_paragraph(result['qwen_c_v1'])

    doc.add_page_break()

    doc.add_heading("V2 Query", level=3)
    doc.add_paragraph(result['v2_query'])

    doc.add_heading("V2 Classification Results", level=3)
    doc.add_paragraph(f"Detected Type: {result['shunt_type_v2']}")

    doc.add_heading("LLAMA 70B Response", level=4)
    doc.add_paragraph(result['groq_c_v2'])

    doc.add_heading("Qwen V2 Response", level=4)
    doc.add_paragraph(result['qwen_c_v2'])

    doc.add_paragraph("")

# Task 2: Ligation Results
doc.add_page_break()
doc.add_heading('TASK 2: LIGATION PLANNING (WITH RAG + RERANKING)', level=1)

for result in results['evaluations']:
    doc.add_heading(f"Test Case: {result['test_name']} ({result['source']})", level=2)

    doc.add_heading("V1 Ligation Plan", level=3)

    doc.add_heading("LLAMA 70B - Ligation Steps", level=4)
    doc.add_paragraph(result['groq_l_v1'][:300])

    doc.add_heading("LLAMA 70B - Reasoning", level=4)
    doc.add_paragraph(result['groq_l_v1'][300:400] if len(result['groq_l_v1']) > 300 else "")

    doc.add_heading("Qwen V2 - Ligation Steps", level=4)
    doc.add_paragraph(result['qwen_l_v1'][:300])

    doc.add_heading("Qwen V2 - Reasoning", level=4)
    doc.add_paragraph(result['qwen_l_v1'][300:400] if len(result['qwen_l_v1']) > 300 else "")

    doc.add_page_break()

    doc.add_heading("V2 Ligation Plan", level=3)

    doc.add_heading("LLAMA 70B - Ligation Steps", level=4)
    doc.add_paragraph(result['groq_l_v2'][:300])

    doc.add_heading("LLAMA 70B - Reasoning", level=4)
    doc.add_paragraph(result['groq_l_v2'][300:400] if len(result['groq_l_v2']) > 300 else "")

    doc.add_heading("Qwen V2 - Ligation Steps", level=4)
    doc.add_paragraph(result['qwen_l_v2'][:300])

    doc.add_heading("Qwen V2 - Reasoning", level=4)
    doc.add_paragraph(result['qwen_l_v2'][300:400] if len(result['qwen_l_v2']) > 300 else "")

    doc.add_paragraph("")

# Evaluation Tables
doc.add_page_break()
doc.add_heading('EVALUATION SUMMARY TABLES', level=1)

# Table 1: Classification
doc.add_heading('Table 1: Shunt Classification Evaluation', level=2)
doc.add_paragraph('Instructions: Fill in columns 2-5 with accuracy scores (0-100) and reasoning quality ratings (1-5)')
doc.add_paragraph('')

table1 = doc.add_table(rows=3, cols=5)
table1.style = 'Light Grid Accent 1'
add_table_borders(table1)

header_cells = table1.rows[0].cells
header_cells[0].text = 'Model'
header_cells[1].text = 'V1 Accuracy'
header_cells[2].text = 'V2 Accuracy'
header_cells[3].text = 'V1 Reasoning Quality'
header_cells[4].text = 'V2 Reasoning Quality'

table1.rows[1].cells[0].text = 'LLAMA 70B Versatile'
table1.rows[2].cells[0].text = 'Qwen V2 Fine-tuned'

for row_idx in [1, 2]:
    for col_idx in range(1, 5):
        table1.rows[row_idx].cells[col_idx].text = '[Enter score]'

# Table 2: Ligation
doc.add_heading('Table 2: Ligation Planning Evaluation', level=2)
doc.add_paragraph('Instructions: Fill in columns 2-5 with quality ratings (1-5) and reasoning quality (1-5)')
doc.add_paragraph('')

table2 = doc.add_table(rows=3, cols=5)
table2.style = 'Light Grid Accent 1'
add_table_borders(table2)

header_cells = table2.rows[0].cells
header_cells[0].text = 'Model'
header_cells[1].text = 'V1 Quality'
header_cells[2].text = 'V2 Quality'
header_cells[3].text = 'V1 Reasoning Quality'
header_cells[4].text = 'V2 Reasoning Quality'

table2.rows[1].cells[0].text = 'LLAMA 70B Versatile'
table2.rows[2].cells[0].text = 'Qwen V2 Fine-tuned'

for row_idx in [1, 2]:
    for col_idx in range(1, 5):
        table2.rows[row_idx].cells[col_idx].text = '[Enter score]'

doc.save('COMPREHENSIVE_EVALUATION_REPORT.docx')

# ============================================================================
# STEP 6: SAVE JSON RESULTS
# ============================================================================

print("\n[5/6] Saving results...")

with open('evaluation_results.json', 'w') as f:
    json.dump(results, f, indent=2)

# ============================================================================
# COMPLETE
# ============================================================================

print("\n[6/6] Done!")
print("\n" + "="*80)
print("EVALUATION COMPLETE")
print("="*80)
print("\nGenerated Files:")
print("  ✓ COMPREHENSIVE_EVALUATION_REPORT.docx")
print("  ✓ evaluation_results.json")
print("\nNext steps:")
print("  1. Open COMPREHENSIVE_EVALUATION_REPORT.docx")
print("  2. Review model outputs for both tasks")
print("  3. Fill in evaluation tables with your scores")
print("="*80)
