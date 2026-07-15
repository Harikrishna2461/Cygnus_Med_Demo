"""
Generates Notes/Manager_Agent_Plan.docx
Run from Task_2_App root: python Notes/gen_manager_agent_plan.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def h3(text):
    doc.add_heading(text, level=3)

def body(text):
    doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def numbered(text):
    doc.add_paragraph(text, style="List Number")

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()

# ── Title ─────────────────────────────────────────────────────────────────────
title = doc.add_heading("Manager Agent Architecture Plan", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    "Contingency plan for upgrading the Task-2 guidance pipeline from the current "
    "tool-use / sequential CrewAI architecture to a hierarchical manager-agent design. "
    "Activate this plan if the tool-use approach reaches a quality ceiling that cannot "
    "be resolved by prompt tuning or adding more tools."
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — WHEN TO SWITCH
# ══════════════════════════════════════════════════════════════════════════════
h1("1. When to Switch to a Manager Agent")
body(
    "The current sequential + tool-use architecture handles the Task-2 guidance "
    "pipeline correctly for a linear CHIVA examination. Switch to a manager agent "
    "when one or more of these conditions arise:"
)
add_table(
    ["Trigger Condition", "Why Sequential Fails Here", "Manager Advantage"],
    [
        ["Clinical workflow becomes branching",
         "Sequential always runs all 5 agents in the same order regardless of circuit state. "
         "For a confirmed Type 1 circuit, ShuntAnalyst and CircuitAnalyst still run unnecessarily.",
         "Manager skips agents not relevant to the current state — faster and more focused."],
        ["Agent count grows beyond 7",
         "Sequential context chains become too long; downstream agents receive diluted context "
         "from too many upstream summaries.",
         "Manager selects which 2-3 agents are relevant per probe_move and calls only those."],
        ["Parallel examination tracks needed",
         "E.g. simultaneously tracking a Type 1 on the right leg and a Type 2A on the left. "
         "Sequential has no concept of parallel sub-tasks.",
         "Manager can run sub-crews for each leg independently and merge results."],
        ["Dynamic protocol selection required",
         "Tool-use lets NavigationPlanner look up one target zone. If the manager needs to "
         "compare protocols across multiple candidate zones before deciding, tool calls per agent "
         "become unwieldy.",
         "Manager reasons over all zones at once and delegates protocol lookup centrally."],
        ["Quality plateau with tool-use approach",
         "If NavigationPlanner guidance quality stops improving despite protocol grounding, "
         "it may need the manager's broader clinical picture to make better routing decisions.",
         "Manager has full CHIVA knowledge and can override routing decisions made by CircuitAnalyst."],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — ARCHITECTURE OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Architecture Overview")

h2("2.1 Current Architecture (Sequential + Tool-Use)")
body("For reference — what the manager architecture replaces:")
add_table(
    ["Layer", "Component", "Role"],
    [
        ["Pre-processing", "protocol_agent, history_agent, q_state_agent",
         "Python lookups; results injected into state_message before crew runs."],
        ["Agent 1", "ClinicalInterpreter", "Reads state_message; assesses clip validity."],
        ["Agent 2", "ShuntAnalyst", "Classifies shunt type from Task1 output."],
        ["Agent 3", "CircuitAnalyst", "Identifies target zone from Task1+Task2 output."],
        ["Agent 4", "NavigationPlanner", "Issues move command; calls get_zone_protocol tool."],
        ["Agent 5", "GuidanceSpecialist", "Formats Task4 output as JSON."],
    ]
)

h2("2.2 Proposed Manager Agent Architecture")
body(
    "Process: CrewAI Process.hierarchical. One manager LLM orchestrates a pool of "
    "sub-agents. The manager decides which sub-agents to invoke, in what order, and "
    "whether to re-invoke any agent for a follow-up."
)
add_table(
    ["Component", "Model Tier", "Role"],
    [
        ["ExaminationDirector (manager)", "heavy (120B)",
         "Receives full state. Decides which sub-agents to call and in what order. "
         "Has complete CHIVA knowledge to make routing decisions without delegating everything."],
        ["ClinicalInterpreter (sub-agent)", "mid (70B)",
         "Same as current. Called by manager when clip assessment is needed."],
        ["ShuntAnalyst (sub-agent)", "heavy (120B)",
         "Same as current. Called when enough clips exist to classify."],
        ["CircuitAnalyst (sub-agent)", "heavy (120B)",
         "Same as current. Called when zone routing decision is needed."],
        ["NavigationPlanner (sub-agent)", "mid (70B)",
         "Same as current. Always called last before GuidanceSpecialist. "
         "Retains get_zone_protocol and get_vein_examination_objectives tools."],
        ["GuidanceSpecialist (sub-agent)", "heavy (120B)",
         "Same as current. Always the final call — formats JSON output. "
         "Manager MUST call this last to guarantee structured output."],
    ]
)

h2("2.3 Manager Agent Definition")
body("The ExaminationDirector system prompt must cover:")
bullet("Full CHIVA shunt type logic (Types 1, 2A, 2B, 2C, 3, 1+2, 4, 5, 6) — "
       "so the manager can decide whether ShuntAnalyst is even needed on this turn.")
bullet("Q1-Q4 state machine — so the manager knows which agents are relevant given the current circuit state.")
bullet("Delegation rules: when to call ClinicalInterpreter (always first), when to skip ShuntAnalyst "
       "(no clips yet), when to call CircuitAnalyst (Q1-Q3 still open), when to skip CircuitAnalyst "
       "(circuit complete, just need JSON formatter).")
bullet("Hard rule: GuidanceSpecialist MUST always be the last agent called — it is the only "
       "agent that outputs valid JSON for the streaming UI.")
bullet("Fast-path triggers: if Q1-Q4 all confirmed and elim test done, manager calls "
       "GuidanceSpecialist directly with 'Circuit complete — classification confirmed' instruction.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — LATENCY ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Latency Analysis")

h2("3.1 Current Latency Budget (Sequential + Tool-Use)")
add_table(
    ["Step", "Component", "Typical Duration"],
    [
        ["VLM call", "Gemini Vision", "2-4 s"],
        ["Python pre-processing", "history/q_state/protocol agents", "~0 ms"],
        ["Task1 — ClinicalInterpreter", "mid model (Groq)", "3-5 s"],
        ["Task2 — ShuntAnalyst", "heavy model (Groq)", "4-8 s"],
        ["Task3 — CircuitAnalyst", "heavy model (Groq)", "4-8 s"],
        ["Task4 — NavigationPlanner + tool call", "mid model + Python tool", "4-7 s"],
        ["Task5 — GuidanceSpecialist", "heavy model (Groq)", "4-8 s"],
        ["TOTAL per probe_move (typical)", "", "21-40 s"],
        ["Hard timeout", "", "90 s"],
    ]
)

h2("3.2 Manager Agent Latency Budget")
add_table(
    ["Step", "Component", "Typical Duration"],
    [
        ["VLM call", "Gemini Vision", "2-4 s"],
        ["Manager initial reasoning", "ExaminationDirector (heavy model)", "5-10 s"],
        ["Manager delegation round 1", "Calls ClinicalInterpreter", "3-5 s"],
        ["Manager delegation round 2", "Calls ShuntAnalyst (if needed)", "4-8 s"],
        ["Manager delegation round 3", "Calls CircuitAnalyst (if needed)", "4-8 s"],
        ["Manager delegation round 4", "Calls NavigationPlanner", "4-7 s"],
        ["Manager delegation round 5", "Calls GuidanceSpecialist", "4-8 s"],
        ["Manager inter-round reasoning", "Overhead between delegations", "2-4 s each"],
        ["TOTAL per probe_move (worst case)", "", "35-65 s"],
        ["TOTAL per probe_move (fast-path, 3 agents)", "", "18-28 s"],
        ["Hard timeout risk", "Close to 90 s on Groq under load", "HIGH"],
    ]
)
body(
    "Key insight: the manager adds its own reasoning time BETWEEN every agent delegation. "
    "With 5 sub-agents, there are 5 manager reasoning cycles, each costing 2-4 s on "
    "the heavy model. This is the dominant latency source — not the sub-agents themselves."
)

h2("3.3 Latency Mitigation Strategies")
bullet("Fast-path detection: manager checks circuit state BEFORE delegating. "
       "If Q1-Q4 confirmed, skip ClinicalInterpreter/ShuntAnalyst/CircuitAnalyst and "
       "call NavigationPlanner + GuidanceSpecialist directly. Expected saving: 10-20 s.")
bullet("Manager model tier: use mid-tier model (70B) for the manager if its role is "
       "purely orchestration (no clinical reasoning). Use heavy model only if the manager "
       "must reason clinically. Expected saving: 2-4 s per manager cycle.")
bullet("Reduce max_iter for sub-agents from 3 to 2 under manager control — manager "
       "handles retry logic itself. Expected saving: 0-5 s per agent.")
bullet("Separate timeout budgets: manager gets 15 s; each sub-agent gets 12 s; "
       "GuidanceSpecialist always gets 10 s reserved. Abort and return fallback if "
       "any budget is exceeded.")
bullet("Cache manager decisions: if probe posY and clip set unchanged from last call, "
       "return the cached guidance without running the manager. Same as current "
       "STREAM_LLM_THRESHOLD logic.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — COMPLEXITY ANALYSIS
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Complexity Analysis")

h2("4.1 New Code Required")
add_table(
    ["File", "Change", "Estimated Lines"],
    [
        ["backend/agents/crew_agents.py",
         "Add make_examination_director() factory with full CHIVA system prompt",
         "80-120"],
        ["backend/agents/crew_pipeline.py",
         "Replace run_guidance_crew() sequential logic with hierarchical Process; "
         "add manager agent; restructure task pool",
         "150-200 (rewrite of ~120 existing lines)"],
        ["backend/agents/crew_tools.py",
         "Add fast-path detection tool and any new tools needed by manager",
         "20-40"],
        ["backend/streaming_guidance_engine.py",
         "Add separate timeout budget logic; adjust fallback behaviour",
         "20-30"],
    ]
)

h2("4.2 Debugging Complexity")
body(
    "The manager agent introduces non-deterministic agent invocation order. "
    "This significantly increases debugging difficulty:"
)
bullet("With sequential process: if output is wrong, you know exactly which task in "
       "the fixed chain produced the error. The thinking_log already records each task output.")
bullet("With hierarchical process: the manager may call agents in a different order "
       "each time. The same probe state might produce different agent invocation sequences "
       "on consecutive runs.")
bullet("Required mitigation: extend thinking_log to record (a) the order of manager "
       "delegations, (b) the manager's reasoning between each delegation, and "
       "(c) which sub-agents were SKIPPED. Without this, debugging is blind.")
bullet("CrewAI hierarchical process is less battle-tested than sequential — edge cases "
       "around context passing between non-adjacent agents may need workarounds.")

h2("4.3 JSON Output Guarantee")
body(
    "The current architecture guarantees GuidanceSpecialist (Task5) always runs last "
    "and always produces JSON, because the sequential chain is fixed. "
    "With a manager, this guarantee must be enforced explicitly:"
)
bullet("The ExaminationDirector system prompt must contain a hard rule: "
       "'Your final delegation MUST always be GuidanceSpecialist. '")
bullet("Add a post-processing check in run_guidance_crew(): if the last task output "
       "is not valid JSON, fall back to the last navigation command from NavigationPlanner "
       "and wrap it in the standard JSON format.")
bullet("Add a sentinel: if manager does not call GuidanceSpecialist within the timeout "
       "budget, force-call it with the last available NavigationPlanner output.")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — QUALITY IMPROVEMENT ASSESSMENT
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Expected Quality Improvement")

add_table(
    ["Scenario", "Sequential + Tools", "Manager Agent", "Delta"],
    [
        ["Probe at calf, CircuitAnalyst routes to SPJ",
         "NavigationPlanner calls get_zone_protocol('popliteal_spj') — correct protocol.",
         "Manager calls NavigationPlanner with SPJ context already resolved.",
         "Neutral — tools already solve this."],
        ["Circuit complete (Q1-Q4 confirmed)",
         "Still runs all 5 agents including ClinicalInterpreter and ShuntAnalyst unnecessarily.",
         "Manager fast-paths to NavigationPlanner + GuidanceSpecialist only. Faster + more focused.",
         "Moderate improvement in speed and output focus."],
        ["Ambiguous clip set — ShuntAnalyst uncertain",
         "ShuntAnalyst returns 'undetermined'; CircuitAnalyst still runs with no classification.",
         "Manager can re-invoke ClinicalInterpreter with a specific follow-up question before re-calling ShuntAnalyst.",
         "Potential quality improvement — better handling of ambiguity."],
        ["Mixed shunt (Type 1+2 vs Type 3 disambiguation)",
         "Elimination test prompt must come from ShuntAnalyst output; no retry if ambiguous.",
         "Manager can escalate to ShuntAnalyst twice with different framing.",
         "Potential quality improvement — but adds 5-10 s latency."],
        ["First probe_move with no clips",
         "ClinicalInterpreter, ShuntAnalyst both output minimal content; runs wastefully.",
         "Manager skips ShuntAnalyst when clip count is 0. Saves 4-8 s.",
         "Latency improvement, neutral quality."],
    ]
)

body(
    "Overall quality improvement estimate: LOW to MODERATE, with HIGH variance. "
    "The cases where the manager provides meaningfully better output (ambiguous clips, "
    "circuit complete fast-path) are relatively rare in a normal examination flow. "
    "For the common case (probe scanning, 1-3 clips confirmed), the manager adds "
    "latency without improving guidance quality over the sequential + tools approach."
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — IMPLEMENTATION STEPS
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Step-by-Step Implementation Plan")

h2("Phase 1 — Manager Agent Definition (Day 1)")
numbered("Write make_examination_director() in crew_agents.py. "
         "Role: 'CHIVA Examination Director'. "
         "System prompt: full CHIVA shunt type logic + Q1-Q4 state machine + delegation rules. "
         "Model: heavy. allow_delegation=True.")
numbered("Define the agent pool: all 5 current agents + ExaminationDirector.")
numbered("Write unit test: given a known clip set, does the manager call the right sub-agents?")

h2("Phase 2 — Pipeline Restructure (Day 2)")
numbered("Rewrite run_guidance_crew() to use Process.hierarchical with manager=examination_director.")
numbered("Remove the fixed task chain (task1-task5). Replace with a task pool — "
         "one Task object per sub-agent, all available to the manager.")
numbered("Add the JSON output guarantee logic: post-processing fallback + sentinel.")
numbered("Extend thinking_log schema to record manager delegation order.")

h2("Phase 3 — Timeout Budget (Day 2-3)")
numbered("Implement per-sub-agent timeout budgets (12 s each) separate from the crew-level 90 s limit.")
numbered("Implement fast-path: if Q1-Q4 confirmed, inject fast-path flag that manager checks first.")
numbered("Test timeout behaviour under Groq load: simulate slow responses and verify fallback fires.")

h2("Phase 4 — Validation (Day 3-4)")
numbered("Run the existing test case library against the manager architecture.")
numbered("Compare guidance output quality against sequential + tools baseline.")
numbered("Measure actual latency per probe_move across 20 representative sessions.")
numbered("Tune manager system prompt based on failure modes observed.")

h2("Rollback Plan")
body(
    "Keep the sequential + tool-use implementation in a separate git branch. "
    "If the manager architecture degrades guidance quality or causes timeout issues "
    "in production, revert to the sequential branch in under 5 minutes. "
    "Feature flag in config.py (USE_MANAGER_AGENT = True/False) recommended "
    "to enable A/B testing without a full deploy."
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DECISION CHECKLIST
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Go / No-Go Decision Checklist")
body("Before activating this plan, confirm all of the following:")
add_table(
    ["Check", "Threshold", "Status"],
    [
        ["Tool-use quality plateau confirmed",
         "Guidance quality not improving despite 2+ weeks of prompt tuning", "TBD"],
        ["Groq capacity sufficient",
         "P95 latency per agent call < 8 s under production load", "TBD"],
        ["Test coverage for sequential exists",
         "At least 20 labelled probe_move sequences with expected outputs", "TBD"],
        ["Debugging tooling ready",
         "thinking_log extended to record manager delegation order", "TBD"],
        ["Fallback branch tagged",
         "Sequential + tool-use commit tagged and deployable in < 5 min", "TBD"],
        ["Feature flag implemented",
         "USE_MANAGER_AGENT in config.py, tested both values", "TBD"],
    ]
)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "Manager_Agent_Plan.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
