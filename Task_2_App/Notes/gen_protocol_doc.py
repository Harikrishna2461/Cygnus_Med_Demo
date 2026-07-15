"""
Generates Notes/Protocol_Source_References.docx
Run from Task_2_App root: python Notes/gen_protocol_doc.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

# ── helpers ───────────────────────────────────────────────────────────────────

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def h3(text):
    doc.add_heading(text, level=3)

def body(text):
    doc.add_paragraph(text)

def bullet(text):
    doc.add_paragraph(text, style="List Bullet")

def add_table(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for r_idx, row in enumerate(rows):
        cells = t.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = val
    doc.add_paragraph()


# ── Title ─────────────────────────────────────────────────────────────────────

title = doc.add_heading("Duplex Ultrasound Protocol — Source References", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    "This document records the exact bibliographic sources, page numbers, and "
    "code locations for every protocol rule and vein examination objective used "
    "in the Task-2 active guidance system."
)
doc.add_paragraph(
    "Generated automatically from Notes/gen_protocol_doc.py — do not edit manually."
)
doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — BIBLIOGRAPHY
# ══════════════════════════════════════════════════════════════════════════════

h1("1. Full Bibliography")

bib = [
    (
        "Adler 2022",
        "Adler RS, Braga L, Costa DN, et al. "
        "Varicose Veins of the Lower Extremity: Doppler US Evaluation — Protocols, "
        "Patterns, and Pitfalls. RadioGraphics. 2022;42(7):2185-2202.",
        "adler-et-al-2022-varicose-veins-of-the-lower-extremity-doppler-us-evaluation-protocols-patterns-and-pitfalls.pdf",
        "pdf_extracts/Adler2022.txt",
    ),
    (
        "Gianesini 2014",
        "Gianesini S, Menegatti E, Zuolo M, Salvi M, Sisini F, Zamboni P. "
        "CHIVA strategy for varicose vein treatment — a 27-year evidence review. "
        "Phlebology. 2014;30(1):6-18.",
        "CHIVA STRATEGY Gianesini014.pdf",
        "pdf_extracts/Gianesini.txt",
    ),
    (
        "Delfrate 2023",
        "Delfrate R. The CHIVA Method for the Treatment of Varicose Veins. "
        "Journal of Translational and Vascular Research (JTAVR). 2023;8(1):18-30.",
        "DelfrateR CHIVA article.pdf",
        "pdf_extracts/Delfrate.txt",
    ),
    (
        "AVF 2023",
        "American Venous Forum (AVF). 2023 Clinical Practice Guidelines — "
        "Chronic Venous Disease of the Lower Extremities. "
        "(Referenced in Delfrate 2023 for perforator incompetence criteria.)",
        "(not stored locally — cited via Delfrate 2023)",
        "N/A",
    ),
    (
        "Mendoza 2014",
        "Mendoza E, Blattler W, Amsler F (eds). "
        "Duplex Ultrasound of Superficial Leg Veins. "
        "Springer, Berlin. 2014. ISBN 978-3-642-39236-1.",
        "0-duplex-ultrasound-of-superficial-leg-veins-2014.pdf",
        "pdf_extracts/DuplexUS_2014.txt",
    ),
]

add_table(
    ["Short Key", "Full Citation", "Local PDF File", "Extracted Text File"],
    bib
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — REGION-SPECIFIC ZONE PROTOCOLS
# ══════════════════════════════════════════════════════════════════════════════

h1("2. Region-Specific Zone Protocols")
body(
    "Zone protocols are returned by protocol_agent.get_protocol(region, pos_y). "
    "Each zone is selected by the posY ratio (0.0 = groin, 1.0 = ankle) combined "
    "with the region string. The following table maps every clinical rule to its "
    "source and page number."
)
doc.add_paragraph()

# ── 2.1 Patient Positioning ──────────────────────────────────────────────────
h2("2.1 Patient Positioning — All Zones")
body("Source: Adler et al. 2022, pp. 2190.")
bullet("Reverse Trendelenburg >= 60 degrees for ALL venous insufficiency studies — "
       "legs below the patient's head to maximise venous filling and reflux detection.")
bullet("Hip externally rotated, knee slightly flexed (non-weight bearing).")
bullet("Left lateral decubitus for RIGHT SSV assessment.")
bullet("Right lateral decubitus for LEFT SSV assessment.")
body("Code location: backend/agents/guidance_agent.py SYSTEM_PROMPT, "
     "backend/agents/protocol_agent.py _PROTOCOLS['sfj_groin'].")
doc.add_paragraph()

# ── 2.2 SFJ / Groin Zone ─────────────────────────────────────────────────────
h2("2.2 Zone: SFJ / Groin (posY <= 0.15, region = GROIN or UPPER-THIGH)")
body("Primary sources: Gianesini 2014, PDF pp. 13-14 (paper p. 12); "
     "Delfrate 2023, pp. 22-23 (JTAVR journal pages).")
add_table(
    ["Rule", "Source", "Page"],
    [
        ["Transverse B-mode: identify Mickey Mouse sign "
         "(CFV centre, GSV + femoral artery lateral ovals)",
         "Gianesini 2014", "PDF p. 13 (paper p. 12)"],
        ["Place Doppler gate on FEMORAL SIDE of terminal valve",
         "Gianesini 2014", "PDF p. 13 (paper p. 12)"],
        ["Apply Valsalva: confirm adequacy by cessation of forward CFV flow",
         "Adler 2022", "p. 2191"],
        ["BOTH Valsalva AND Parana must be positive to confirm SFJ incompetence (N1->N2 entry)",
         "Gianesini 2014", "PDF p. 14 (paper p. 12)"],
        ["Only Parana positive but not Valsalva -> terminal valve competent; "
         "reflux is pre-terminal or from pelvic leak point (PLP)",
         "Delfrate 2023", "p. 23"],
        ["Three PLPs: Superior Gluteal Point (SGP), Inferior Gluteal Point (IGP), Obturator Point (OP)",
         "Delfrate 2023", "p. 22"],
        ["AASV (anterior accessory saphenous vein) lies anterior to GSV in upper thigh — "
         "assess separately; classified N3, not N2 (common duplex pitfall)",
         "Delfrate 2023", "p. 23"],
    ]
)

# ── 2.3 Upper Thigh Zone ─────────────────────────────────────────────────────
h2("2.3 Zone: Upper Thigh (posY 0.15-0.35)")
body("Sources: Adler 2022, pp. 2191-2192; Delfrate 2023, p. 24.")
add_table(
    ["Rule", "Source", "Page"],
    [
        ["Transverse B-mode saphenous eye: GSV within fascial compartment "
         "between two bright horizontal fascia lines",
         "Adler 2022", "p. 2191"],
        ["Confirm N2 identity by fascial envelope; superficialisation "
         "(escape from compartment) changes treatment options",
         "Delfrate 2023", "p. 24"],
        ["Measure anteroposterior GSV diameter at upper thigh, mid-thigh, "
         "above knee, below knee",
         "Adler 2022", "p. 2191"],
        ["Reflux threshold: outward/reversed flow lasting > 500 ms on Parana "
         "= haemodynamically significant",
         "Adler 2022", "p. 2193"],
        ["Some guidelines use 1 second for better specificity; "
         "500 ms is the conventional threshold",
         "Adler 2022", "p. 2193"],
    ]
)

# ── 2.4 Hunterian / Proximal Perforators ──────────────────────────────────────
h2("2.4 Zone: Hunterian / Proximal Perforators (posY 0.35-0.55)")
body("Sources: Adler 2022, pp. 2186, 2191-2192; Delfrate 2023, pp. 23-25; "
     "Mendoza 2014, p. 20 (PDF p. 33); AVF 2023.")
add_table(
    ["Rule", "Source", "Page"],
    [
        ["Hunterian perforators connect GSV to femoral vein in mid-thigh "
         "(also called Dodd perforators in some texts for the distal group)",
         "Mendoza 2014", "p. 20 (PDF p. 33)"],
        ["Three maneuvers required: static squeezing (gravitational test), "
         "Parana (physiological), Valsalva (hypertensive test)",
         "Delfrate 2023", "p. 25"],
        ["Pathological perforator: outward flow >= 500 ms AND diameter >= 3.5 mm",
         "AVF 2023 (cited in Delfrate 2023)", "p. 25"],
        ["Inward flow during muscle diastole (Parana release) = re-entry perforator (RP N3->N1)",
         "Delfrate 2023", "p. 25"],
        ["Diastolic reflux into deep system via perforator is always pathological and pathogenic",
         "Delfrate 2023", "p. 25"],
        ["Biphasic perforators: systolic outward + diastolic inward = likely re-entry candidate",
         "Delfrate 2023", "p. 25"],
    ]
)

# ── 2.5 Dodd / Distal Perforators ─────────────────────────────────────────────
h2("2.5 Zone: Dodd / Distal Perforators (posY 0.55-0.70)")
body("Sources: Adler 2022, pp. 2191-2192; Delfrate 2023, p. 25; AVF 2023.")
body("Same three-maneuver rule and pathological criteria as Hunterian zone apply. "
     "Dodd perforators are located in the distal thigh / above-knee segment.")
bullet("No reflux, no re-entry: if GSV reflux persists below an escape point, "
       "another re-entry exists further distal. (Adler 2022, p. 2192)")
doc.add_paragraph()

# ── 2.6 Popliteal / SPJ Zone ─────────────────────────────────────────────────
h2("2.6 Zone: Popliteal / SPJ (posY 0.65-0.80, posterior)")
body("Sources: Gianesini 2014, PDF p. 14 (paper p. 13); Delfrate 2023, p. 23; "
     "Adler 2022, p. 2191.")
add_table(
    ["Rule", "Source", "Page"],
    [
        ["Both Parana (active) AND compression/relaxation (passive CR) must be "
         "positive simultaneously to confirm SPJ incompetence",
         "Gianesini 2014", "PDF p. 14 (paper p. 13)"],
        ["One maneuver positive alone != true junctional incompetence",
         "Gianesini 2014", "PDF p. 14 (paper p. 13)"],
        ["SPJ location is variable — may connect to gastrocnemian vein "
         "rather than popliteal vein directly",
         "Delfrate 2023", "p. 23"],
        ["Giacomini vein (posterior thigh, SSV->GSV): forward systolic flow "
         "with Parana = viable outflow route",
         "Delfrate 2023", "p. 23"],
        ["SPJ disconnection should be performed below the Giacomini junction "
         "in mixed shunts",
         "Delfrate 2023", "p. 23"],
    ]
)

# ── 2.7 Calf Zone ─────────────────────────────────────────────────────────────
h2("2.7 Zone: Calf (posY 0.70-0.90)")
body("Sources: Adler 2022, pp. 2191-2192; Delfrate 2023, pp. 24-25.")
bullet("Continue 500 ms reflux threshold for any calf tributary or perforator.")
bullet("Three-maneuver rule for any perforator identified in calf.")
bullet("Check for N3 escape points above fascia (EP N2->N3) where tributary "
       "reflux feeds skin.")
doc.add_paragraph()

# ── 2.8 Ankle / SSV Origin Zone ───────────────────────────────────────────────
h2("2.8 Zone: Ankle / SSV Origin (posY 0.90-1.0)")
body("Sources: Adler 2022, p. 2191; Mendoza 2014, pp. 172 (SSV objectives).")
bullet("SSV origin is at the lateral malleolus; confirm with B-mode in "
       "transverse plane.")
bullet("Same 500 ms reflux threshold applies.")
doc.add_paragraph()

# ── 2.9 Valsalva vs Parana vs Squeezing ───────────────────────────────────────
h2("2.9 Augmentation Maneuvers — All Zones")
body("Sources: Adler 2022, pp. 2190-2191; Delfrate 2023, p. 21; Gianesini 2014, p. 12.")
add_table(
    ["Maneuver", "Technique", "When to Use", "Source & Page"],
    [
        ["Valsalva",
         "Blocked forced expiration; adequate when forward CFV flow ceases",
         "SFJ, SPJ, upper deep veins",
         "Adler 2022 p. 2191; Delfrate 2023 p. 22"],
        ["Parana",
         "Slight waist push triggers calf proprioceptive contraction-relaxation reflex",
         "Preferred for ALL reflux testing; more physiological than squeezing",
         "Delfrate 2023 p. 21; Gianesini 2014 p. 12"],
        ["Squeezing",
         "Manual calf squeeze and release",
         "Acceptable but insufficient alone; use Parana as primary",
         "Delfrate 2023 p. 21"],
    ]
)
body('Critical (Delfrate 2023, p. 21): "Squeezing alone in a patient standing or '
     'sitting is not sufficient, and creates a risk of overtreatment." '
     "Always confirm with Parana or Valsalva.")
doc.add_paragraph()

# ── 2.10 Standard Examination Sequence ────────────────────────────────────────
h2("2.10 Standard Examination Sequence")
body("Sources: Adler 2022, p. 2190; Delfrate 2023, p. 21.")
add_table(
    ["Step", "Action", "Source & Page"],
    [
        ["1", "DVT compression — all deep veins, before any reflux testing",
         "Adler 2022 p. 2190; Delfrate 2023 p. 21"],
        ["2", "Deep vein reflux — iliac valve competence (Valsalva), CFV above SFJ",
         "Adler 2022 p. 2190"],
        ["3", "SFJ — Mickey Mouse sign; Valsalva + Parana (both required); AASV separate",
         "Gianesini 2014 PDF p. 13-14; Delfrate 2023 pp. 22-23"],
        ["4", "GSV trunk — medial thigh to calf; saphenous eye; 500 ms reflux threshold",
         "Adler 2022 pp. 2191-2193"],
        ["5", "Hunterian perforators — all 3 maneuvers; check for N3 above fascia (EP N2->N3)",
         "Delfrate 2023 p. 25; AVF 2023"],
        ["6", "SPJ — posterior knee; both Parana + CR required; variable anatomy",
         "Gianesini 2014 PDF p. 14; Delfrate 2023 p. 23"],
        ["7", "SSV trunk — posterior calf; lateral approach; same 500 ms threshold",
         "Adler 2022 p. 2191"],
        ["8", "Re-entry perforators — inward diastolic flow; confirm >= 3.5 mm; biphasic pattern",
         "Delfrate 2023 p. 25; AVF 2023"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — VEIN-SPECIFIC EXAMINATION OBJECTIVES (MENDOZA 2014)
# ══════════════════════════════════════════════════════════════════════════════

h1("3. Vein-Specific Examination Objectives (Mendoza 2014)")
body(
    "These objectives are appended to the zone protocol when the operator activates "
    "a vein scan mode (set_scan_vein socket event). They come exclusively from "
    "Mendoza et al. 2014, Duplex Ultrasound of Superficial Leg Veins, Springer."
)
doc.add_paragraph()

# ── 3.1 GSV ──────────────────────────────────────────────────────────────────
h2("3.1 GSV — Great Saphenous Vein (Section 7.2)")
body("Source: Mendoza 2014, Section 7.2 Objectives, book pp. 120-121.")
add_table(
    ["#", "Examination Objective", "Book Page"],
    [
        ["1", "Document the great saphenous vein (GSV) from its origin at the saphenofemoral junction (SFJ) to the ankle, assessing the entire trunk.", "pp. 120-121"],
        ["2", "Identify and document the location of the SFJ, including the terminal and preterminal valves.", "pp. 120-121"],
        ["3", "Assess for GSV reflux at the SFJ level using the Valsalva maneuver and calf compression/release.", "pp. 120-121"],
        ["4", "Measure the diameter of the GSV at standardised levels: upper thigh, mid-thigh, above knee, and below knee.", "pp. 120-121"],
        ["5", "Map the course of the GSV within the saphenous compartment (confirm saphenous eye sign in transverse view).", "pp. 120-121"],
        ["6", "Identify any accessory saphenous veins (anterior or posterior) and document their relationship to the GSV trunk.", "pp. 120-121"],
        ["7", "Document the extent and location of GSV reflux if present, noting where reflux terminates.", "pp. 120-121"],
        ["8", "Identify any areas of GSV duplication or anatomical variants along the thigh or calf segments.", "pp. 120-121"],
        ["9", "Assess for any previous interventions (thrombosis, ablation changes) and document residual patency.", "pp. 120-121"],
        ["10", "Document any tributaries originating from the GSV, noting their calibre and connection points.", "pp. 120-121"],
    ]
)

# ── 3.2 SSV ──────────────────────────────────────────────────────────────────
h2("3.2 SSV — Small Saphenous Vein (Section 8.2)")
body("Source: Mendoza 2014, Section 8.2 Objectives, book p. 172.")
add_table(
    ["#", "Examination Objective", "Book Page"],
    [
        ["1", "Document the small saphenous vein (SSV) from its origin at the lateral malleolus to the saphenopopliteal junction (SPJ).", "p. 172"],
        ["2", "Identify the SPJ location precisely, noting its level relative to the popliteal crease.", "p. 172"],
        ["3", "Assess for SSV reflux at the SPJ using Parana and passive compression/release maneuvers.", "p. 172"],
        ["4", "Identify any cranial extension of the SSV (Giacomini vein) and assess for reflux or competent flow.", "p. 172"],
        ["5", "Measure SSV diameter at the SPJ and at mid-calf.", "p. 172"],
        ["6", "Assess SSV course and confirm it lies within the saphenous compartment in the calf.", "p. 172"],
        ["7", "Identify any posterior calf tributaries connecting to the SSV.", "p. 172"],
        ["8", "Document any duplication or anatomical variants of the SSV.", "p. 172"],
    ]
)

# ── 3.3 Perforators ───────────────────────────────────────────────────────────
h2("3.3 Perforators (Section 9.2)")
body("Source: Mendoza 2014, Section 9.2 Objectives, book pp. 188-189.")
add_table(
    ["#", "Examination Objective", "Book Page"],
    [
        ["1", "Identify and document all perforating veins connecting the superficial and deep venous systems.", "pp. 188-189"],
        ["2", "Determine flow direction in each perforator: inward (physiological), outward (potentially pathological), or bidirectional.", "pp. 188-189"],
        ["3", "Measure the diameter of each identified perforator at the fascial level.", "pp. 188-189"],
        ["4", "Classify each perforator as competent or incompetent based on the 500 ms / 3.5 mm AVF criteria.", "pp. 188-189"],
        ["5", "Identify the anatomical group of each perforator (paratibial, posterior tibial, Hunterian/Dodd, gastrocnemial).", "pp. 188-189"],
        ["6", "Assess perforators using all three maneuvers: squeezing, Parana, and Valsalva.", "pp. 188-189"],
        ["7", "Identify re-entry perforators by inward diastolic flow on Parana release.", "pp. 188-189"],
        ["8", "Document the anatomical relationship of each perforator to the GSV or SSV trunk and any connecting tributaries.", "pp. 188-189"],
    ]
)

# ── 3.4 Tributaries ───────────────────────────────────────────────────────────
h2("3.4 Tributaries (Section 10.2)")
body("Source: Mendoza 2014, Section 10.2 Objectives, book pp. 201-202.")
add_table(
    ["#", "Examination Objective", "Book Page"],
    [
        ["1", "Identify and document all visible tributary varicosities and their connections to the GSV, SSV, or perforators.", "pp. 201-202"],
        ["2", "Assess flow direction in tributaries and determine whether they carry reflux from the saphenous trunk or from perforators.", "pp. 201-202"],
        ["3", "Document the feeding source of each tributary cluster (EP at N2 level or EP at N3 perforator level).", "pp. 201-202"],
        ["4", "Identify any tributaries that function as collateral re-entry pathways (drainage into deep system via perforators).", "pp. 201-202"],
        ["5", "Map the distribution of tributary reflux to guide targeted foam sclerotherapy or surgical disconnection.", "pp. 201-202"],
        ["6", "Assess for any non-saphenous sources of tributary varicosities (pelvic escape, AASV, posterior thigh connections).", "pp. 201-202"],
    ]
)

# ── 3.5 Deep Veins ────────────────────────────────────────────────────────────
h2("3.5 Deep Veins (Chapter 14)")
body("Source: Mendoza 2014, Chapter 14 (Deep Veins), book pp. 267-278.")
add_table(
    ["#", "Examination Objective", "Book Page"],
    [
        ["1", "Perform compression ultrasound of all deep veins to exclude DVT before reflux testing.", "pp. 267-268"],
        ["2", "Assess common femoral vein (CFV) for patency and spontaneous phasic flow.", "pp. 267-268"],
        ["3", "Evaluate iliac vein competence using Valsalva maneuver at the CFV level.", "pp. 269-270"],
        ["4", "Assess popliteal vein and proximal deep veins for reflux using Parana maneuver.", "pp. 271-272"],
        ["5", "Document any deep vein post-thrombotic changes (wall thickening, synechiae, partial recanalisation).", "pp. 273-274"],
        ["6", "Identify and assess the gastrocnemial and soleal veins for thrombosis.", "pp. 275-276"],
        ["7", "Assess deep vein reflux thresholds: > 1000 ms (1 second) for femoral and popliteal veins; > 500 ms for calf veins.", "pp. 277-278"],
        ["8", "Document any perforating veins connecting the deep system to the superficial system with pathological outward flow.", "pp. 277-278"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — QUICK REFERENCE TABLE
# ══════════════════════════════════════════════════════════════════════════════

h1("4. Quick Reference — Key Clinical Thresholds")
add_table(
    ["Parameter", "Threshold / Value", "Source", "Page"],
    [
        ["Reflux duration (superficial veins)", "> 500 ms", "Adler 2022", "p. 2193"],
        ["Reflux duration (deep veins — femoral/popliteal)", "> 1000 ms", "Mendoza 2014", "pp. 277-278"],
        ["Reflux duration (deep calf veins)", "> 500 ms", "Mendoza 2014", "pp. 277-278"],
        ["Perforator minimum diameter (pathological)", ">= 3.5 mm", "AVF 2023 (cited in Delfrate 2023)", "p. 25"],
        ["Perforator reflux duration (pathological)", ">= 500 ms outward", "AVF 2023 (cited in Delfrate 2023)", "p. 25"],
        ["Patient tilt for all reflux studies", "Reverse Trendelenburg >= 60 deg", "Adler 2022", "p. 2190"],
        ["SFJ confirmation", "BOTH Valsalva AND Parana positive", "Gianesini 2014", "PDF pp. 13-14"],
        ["SPJ confirmation", "BOTH Parana AND passive CR positive", "Gianesini 2014", "PDF p. 14"],
        ["Re-entry perforator identification", "Inward diastolic flow on Parana release", "Delfrate 2023", "p. 25"],
        ["Squeezing alone", "INSUFFICIENT — risk of overtreatment", "Delfrate 2023", "p. 21"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — CODE LOCATION TABLE
# ══════════════════════════════════════════════════════════════════════════════

h1("5. Code Location Map")
body("Where each protocol element lives in the Task-2 backend codebase.")
add_table(
    ["Protocol Element", "File", "Symbol / Key"],
    [
        ["Zone protocol text (all 8 zones)",
         "backend/agents/protocol_agent.py",
         "_PROTOCOLS dict"],
        ["Vein examination objectives (5 veins)",
         "backend/agents/protocol_agent.py",
         "_VEIN_EXAM_OBJECTIVES dict"],
        ["get_protocol() — combines zone + vein mode",
         "backend/agents/protocol_agent.py",
         "get_protocol(region, pos_y, vein_mode)"],
        ["get_vein_examination_protocol()",
         "backend/agents/protocol_agent.py",
         "get_vein_examination_protocol(vein_mode)"],
        ["System prompt CHIVA clinical knowledge",
         "backend/agents/guidance_agent.py",
         "SYSTEM_PROMPT (REGION-SPECIFIC PROTOCOL KNOWLEDGE section)"],
        ["Paranà vs squeezing clarification",
         "backend/agents/guidance_agent.py",
         "SYSTEM_PROMPT"],
        ["Saphenous eye / transverse confirmation",
         "backend/agents/guidance_agent.py",
         "SYSTEM_PROMPT"],
        ["Patient positioning rule",
         "backend/agents/guidance_agent.py",
         "SYSTEM_PROMPT"],
        ["Active vein scan mode per-session state",
         "backend/streaming_session.py",
         "StreamSession.scan_vein (str)"],
        ["set_scan_vein socket event handler",
         "backend/routes/stream.py",
         "handle_set_scan_vein()"],
        ["scan_vein_ack acknowledgement event",
         "backend/routes/stream.py",
         "handle_set_scan_vein() -> emit('scan_vein_ack')"],
        ["scan_vein reset on stream_start",
         "backend/routes/stream.py",
         "handle_stream_start() -> sess.scan_vein = ''"],
        ["Protocol injected into guidance engine",
         "backend/streaming_guidance_engine.py",
         "process_probe_state() -> protocol_agent.get_protocol(..., vein_mode=session.scan_vein)"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — NOTES ON PARANA SPELLING
# ══════════════════════════════════════════════════════════════════════════════

h1("6. Note on Parana Maneuver Spelling")
body(
    "The maneuver is correctly spelled 'Paranà' (with grave accent) in "
    "Delfrate 2023. Because the extracted text files are plain ASCII, the "
    "accent is dropped throughout the codebase ('Parana'). "
    "Both spellings refer to the same physiological maneuver: a slight push "
    "at the waist level that triggers a proprioceptive calf muscle "
    "contraction-relaxation reflex without manual squeezing."
)

# ── Save ──────────────────────────────────────────────────────────────────────
out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                        "Protocol_Source_References.docx")
doc.save(out_path)
print(f"Saved: {out_path}")
