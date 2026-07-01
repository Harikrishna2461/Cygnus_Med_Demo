"""Generates Task2_Methodology_Report.docx using python-docx."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1F, 0x39, 0x64)
TEAL   = RGBColor(0x1F, 0x74, 0x7D)
BLACK  = RGBColor(0x00, 0x00, 0x00)
GREY   = RGBColor(0x44, 0x44, 0x44)
LGREY  = RGBColor(0xF2, 0xF2, 0xF2)


def set_run_font(run, size_pt, bold=False, italic=False, color=None):
    run.font.size  = Pt(size_pt)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    set_run_font(run, 16, bold=True, color=NAVY)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1F3964')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    set_run_font(run, 13, bold=True, color=TEAL)
    return p


def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_run_font(run, 11, bold=True, color=GREY)
    return p


def body(text, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_run_font(run, 11, color=BLACK)
    return p


def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    set_run_font(run, 11, color=BLACK)
    return p


def code_block(lines):
    """Light-grey code paragraph."""
    text = '\n'.join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # shading
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    rPr.append(shd)
    return p


def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # navy fill
        tc  = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F3964')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10)
            if ri % 2 == 1:
                tc  = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'),   'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'EBF3F8')
                tcPr.append(shd)
    if col_widths:
        for ri2, row in enumerate(t.rows):
            for ci2, cell in enumerate(row.cells):
                cell.width = Inches(col_widths[ci2])
    doc.add_paragraph()
    return t


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('TASK 2 — PROBE LOCALISATION\nAND ACTIVE GUIDANCE SYSTEM')
r.font.size  = Pt(22)
r.font.bold  = True
r.font.color.rgb = NAVY

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_before = Pt(6)
r2 = p2.add_run('Methodology Report')
r2.font.size = Pt(15)
r2.font.color.rgb = TEAL

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(20)
r3 = p3.add_run('Cygnus Med Demo  |  CHIVA Duplex Ultrasound AI Assistant  |  July 2026')
r3.font.size  = Pt(11)
r3.font.italic = True
r3.font.color.rgb = GREY

doc.add_page_break()

# ═══════════════════════════════════════════════════════════════════════════════
# 1. INTRODUCTION
# ═══════════════════════════════════════════════════════════════════════════════
heading1('1. Introduction')

heading2('1.1  What This Report Covers')
body('This report describes the design, architecture, and implementation of Task 2 — a real-time AI system '
     'that guides a surgeon through a CHIVA duplex ultrasound examination of varicose veins. It explains '
     'every component of the pipeline, from the moment the surgeon clicks on a leg diagram to the moment '
     'a 12-word guidance instruction appears on screen.')

heading2('1.2  What is CHIVA?')
body('CHIVA (Conservatrice et Hémodynamique de l\'Insuffisance Veineuse en Ambulatoire) is a minimally '
     'invasive surgical approach to varicose vein treatment. Rather than removing all varicose veins by '
     'stripping, CHIVA works by identifying and cutting only the specific abnormal blood flow circuits '
     'that feed the varices — leaving healthy veins intact and reducing the risk of recurrence.')
body('To plan CHIVA surgery, the surgeon first maps the patient\'s abnormal veins using duplex ultrasound. '
     'This examination follows a specific sequence of questions and requires the surgeon to correctly '
     'identify the direction of blood flow at multiple points along the leg. Task 2 is an AI assistant '
     'that guides this examination in real time.')

heading2('1.3  What Task 2 Does')
body('Task 2 watches the examination as it happens. At every probe position, it knows:')
bullet('Where the probe currently is on the leg')
bullet('What blood flow events the surgeon has confirmed so far')
bullet('What the ultrasound image currently shows (via AI vision analysis)')
bullet('What zones of the leg have already been scanned')
body('From all of this, it produces a short instruction — 12 words or fewer — telling the surgeon exactly '
     'where to move the probe next. This is not a chatbot or a reference guide. It is a live, adaptive '
     'navigation assistant that adapts to the specific findings of the current patient.')

# ═══════════════════════════════════════════════════════════════════════════════
# 2. CLINICAL BACKGROUND
# ═══════════════════════════════════════════════════════════════════════════════
heading1('2. Clinical Background')

heading2('2.1  The Three Venous Compartments (N1, N2, N3)')
body('The CHIVA framework divides the venous system of the leg into three zones:')
add_table(
    ['Compartment', 'Name', 'Contents'],
    [
        ['N1', 'Deep system', 'Common femoral vein (CFV), femoral vein (FV), popliteal vein (PV), deep calf veins. Blood is supposed to flow upward toward the heart through N1.'],
        ['N2', 'Saphenous trunks', 'Great Saphenous Vein (GSV) — groin to ankle, medial side. Small Saphenous Vein (SSV) — popliteal to ankle, posterior side. Identified on ultrasound by the "saphenous eye" — N2 sits inside the fascial compartment.'],
        ['N3', 'Superficial layer', 'Everything above the fascia: tributaries, varicosities, AASV (Anterior Accessory Saphenous Vein), and perforating veins exiting through the fascia.'],
    ],
    col_widths=[1.2, 1.5, 3.8]
)
body('Under normal conditions, blood flows inward: N3 → N2 → N1. In varicose vein disease, this flow '
     'reverses at specific points, creating abnormal circuits that feed the varicosities.')

heading2('2.2  Flow Direction Notation — EP and RP')
body('The system records abnormal blood flow events as "clips":')
add_table(
    ['Term', 'Full name', 'Direction', 'Examples'],
    [
        ['EP', 'Entry Point', 'Deep → Superficial (abnormal)', 'EP N1→N2 (deep enters trunk)\nEP N1→N3 (deep enters tributary)\nEP N2→N3 (trunk escapes to tributary)'],
        ['RP', 'Re-entry Point', 'Superficial → Deep (return path)', 'RP N2→N1 (trunk reflux back to deep)\nRP N3→N1 (tributary re-enters deep)\nRP N3→N2 (tributary re-enters trunk)'],
    ],
    col_widths=[0.8, 1.5, 2.0, 2.2]
)
body('A "clip" is a confirmed duplex finding at a specific location on the leg. Position is recorded as '
     'posY — a decimal from 0.0 (groin) to 1.0 (ankle).')

heading2('2.3  CHIVA Shunt Types')
body('Different combinations of EP and RP clips define the six CHIVA shunt types:')
add_table(
    ['Type', 'Required Clips', 'Clinical Meaning'],
    [
        ['Type 1', 'EP N1→N2 + RP N2→N1\n(No N3)', 'Blood enters GSV at SFJ or perforator, refluxes down the trunk, re-enters deep system lower.'],
        ['Type 2', 'EP N2→N3 + RP N3→N1\n(No EP N1→N2, no RP N2→N1)', 'Blood escapes trunk into tributary; tributary returns to deep via perforator. SFJ competent.'],
        ['Type 3', 'Type 1+2 clips + elimTest = No Reflux', 'Complex combined circuit. Elimination test MANDATORY to confirm and distinguish from Type 1+2.'],
        ['Type 1+2', 'Same clips as Type 3 + elimTest = Reflux', 'Types 1 and 2 coexist. Elimination test shows GSV still refluxes after tributary compression.'],
        ['Type 4', 'EP N1→N3 + RP N2→N1', 'Blood enters N3 directly (bypasses trunk); GSV trunk carries it back to deep.'],
        ['Type 5', 'EP N1→N3 + RP N3→N2 + EP N2→N3 + RP N3→N1\n(No RP N2→N1)', 'Complex loop: N1→N3→N2→N3→N1. Trunk is in the middle of the circuit.'],
        ['Type 6', 'EP N1→N3 + RP N3→N1\n(No N2 involvement)', 'Pure perforator-to-perforator circuit in N3. GSV and SSV trunks completely uninvolved.'],
    ],
    col_widths=[0.8, 2.2, 3.5]
)

heading2('2.4  The Four Diagnostic Questions (Q1–Q4)')
body('Every CHIVA examination follows four sequential questions. The AI uses these as its primary '
     'navigation logic — it always tracks which Q is open and directs the surgeon accordingly:')
add_table(
    ['Question', 'What it asks', 'Looking for'],
    [
        ['Q1', 'Where does blood ENTER the superficial system?', 'EP N1→N2 or EP N1→N3 — typically starts at SFJ (posY 0.04–0.07)'],
        ['Q2', 'Does the trunk REFLUX?', 'RP N2→N1 — does the GSV carry blood backward down the leg?'],
        ['Q3', 'Where does blood ESCAPE the trunk into a tributary?', 'EP N2→N3 — the escape point where a tributary branches off and becomes varicose'],
        ['Q4', 'Where does the tributary RE-ENTER the deep system?', 'RP N3→N1 or RP N3→N2 — closes the circuit'],
    ],
    col_widths=[0.7, 2.3, 3.5]
)

heading2('2.5  Anatomical Positions (posY Landmarks)')
body('The entire guidance system is built around a 0.0–1.0 position scale from groin to ankle:')
add_table(
    ['posY Range', 'Anatomical Zone'],
    [
        ['0.04 – 0.07', 'SFJ — Great Saphenous Vein meets Common Femoral Vein at the groin crease'],
        ['0.08 – 0.20', 'Upper thigh — GSV medial surface'],
        ['0.21 – 0.33', 'Mid-thigh / Dodd perforator zone (medial)'],
        ['0.34 – 0.47', 'Lower thigh / Hunterian perforator zone (medial)'],
        ['0.48 – 0.57', 'Popliteal fossa — SPJ, SSV origin, Giacomini vein (posterior)'],
        ['0.58 – 0.88', 'Calf — GSV medial, SSV posterior'],
        ['0.89 – 1.00', 'Ankle region'],
    ],
    col_widths=[1.3, 5.2]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 3. SYSTEM ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════
heading1('3. System Architecture Overview')

heading2('3.1  Technology Stack')
add_table(
    ['Component', 'Technology'],
    [
        ['Backend web framework', 'Flask 3.0 (Python)'],
        ['Real-time communication', 'Flask-SocketIO (WebSocket / Socket.IO)'],
        ['Text LLM for guidance', 'Groq API — Llama 3.3 70B Versatile'],
        ['Vision LLM (VLM)', 'Groq API — Llama 4 Scout 17B'],
        ['Multi-agent framework', 'CrewAI 0.28+'],
        ['LLM routing layer', 'LiteLLM'],
        ['Video frame extraction', 'OpenCV (headless)'],
        ['Frontend', 'Vanilla HTML / CSS / JavaScript (no build step)'],
        ['Production server', 'Gunicorn (gthread workers) + Nginx'],
    ],
    col_widths=[2.5, 4.0]
)

heading2('3.2  Three Operating Modes')
add_table(
    ['Mode', 'URL', 'Purpose'],
    [
        ['Single-frame', '/', 'Surgeon positions probe on leg diagram and gets a single guidance result. Used for testing individual positions.'],
        ['Streaming', '/stream', 'Full real-time guidance experience. Live leg diagram, clip marking, continuously updated guidance. Primary clinical interface.'],
        ['Test', '/test', 'Developer interface for running pre-defined scenarios and validating the pipeline.'],
    ],
    col_widths=[1.3, 1.0, 4.2]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 4. THE STREAMING PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════
heading1('4. The Streaming Pipeline — How Real-Time Guidance Works')

body('When the surgeon clicks on the leg diagram, a chain of events produces a guidance response '
     'in approximately 5–10 seconds. The pipeline runs entirely in the background, and the surgeon '
     'receives the result as a live WebSocket update.')

heading2('4.1  Step-by-Step Pipeline')
add_table(
    ['Step', 'What happens'],
    [
        ['1. Click', 'Surgeon clicks on the leg SVG diagram in stream.html. coordToState() computes posY, region, surface, leg, and segment values from the click coordinates.'],
        ['2. WebSocket event', 'sendProbeMove() emits a probe_move Socket.IO event containing posY, region, surface, leg, and segment_dist to the Flask server.'],
        ['3. Rate-limit check', 'streaming_guidance_engine.py checks whether posY has changed enough to warrant a new analysis (≥0.06 for LLM, ≥0.05 for VLM). If not, skip.'],
        ['4. Pre-processing sub-agents', 'Four Python modules run in parallel: history_agent (scan coverage map), q_state_agent (which Q is open), protocol_agent (exam protocol lookup), vlm_agent (extract frame, call VLM).'],
        ['5. State message assembly', 'build_state_message() in guidance_agent.py assembles all 7 signals into a single structured text block (see Section 6).'],
        ['6. CrewAI 5-agent pipeline', 'run_guidance_crew() runs five agents in sequence. Each agent sees the outputs of the agents before it. Agent 5 produces the final guidance JSON.'],
        ['7. WebSocket response', 'guidance_update event is emitted to the browser with guidance text, action type, VLM result, and thinking log entry.'],
        ['8. Frontend update', 'Browser renders guidance in the Active Guidance box, updates VLM Frame Analysis grid, and adds a Thinking Log entry.'],
    ],
    col_widths=[1.5, 5.0]
)

heading2('4.2  Rate Limiting')
body('To avoid flooding the AI API on every small mouse movement, the engine only fires when the '
     'probe has moved enough to be clinically meaningful:')
bullet('VLM (vision model): re-runs only if posY changed by ≥ 0.05 since the last VLM call')
bullet('CrewAI crew (guidance): re-runs only if posY changed by ≥ 0.06 since the last call')
bullet('First click of a session: both always fire (last positions start at -1.0)')

heading2('4.3  Stale Response Prevention')
body('Because the crew pipeline takes several seconds, the surgeon may click multiple times before '
     'a result arrives. The system uses a generation counter (sess.bump()) to discard stale results: '
     'if the surgeon has clicked 4 or more times since a computation started, its result is silently '
     'dropped. Only the result from the most recent relevant click is displayed.')

# ═══════════════════════════════════════════════════════════════════════════════
# 5. THE SEVEN SIGNALS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('5. The Seven Signals — What the AI Sees at Each Turn')

body('At each probe position, the system assembles a rich state message with seven distinct '
     'pieces of information. This is the complete input to the AI pipeline:')

add_table(
    ['Signal', 'Name', 'Source', 'What it provides'],
    [
        ['A', 'PROBE STATE', 'Frontend click event', 'Current region, probe surface (anterior-medial / posterior), which leg, and posY value'],
        ['B', 'CONFIRMED FINDINGS', 'Session clip list', 'All EP/RP findings the surgeon has confirmed this session — the growing evidence base'],
        ['C', 'SCAN HISTORY SUMMARY', 'history_agent.py', 'Band-by-band coverage map of the leg ([DONE] / blank) plus a 2-sentence narrative summary of findings per zone'],
        ['D', 'VLM FRAME ANNOTATION', 'vlm_agent.py + Groq Llama 4 Scout', 'Which vessels are visible in the current annotated video frame and whether they are inside/outside the fascial compartment'],
        ['E', 'Q1-Q4 STATUS', 'q_state_agent.py', 'Which of Q1–Q4 is currently open (derived from clips by pure Python logic) plus the "next step" navigation hint'],
        ['F', 'EXAMINATION PROTOCOL', 'protocol_agent.py', 'Book-sourced maneuver instructions for the current probe zone (Adler 2022, Gianesini 2014, Delfrate 2023)'],
        ['G', 'posX', 'Frontend click event (optional)', 'Exact lateral position on the leg diagram, disambiguating medial vs lateral surface at the same posY level'],
    ],
    col_widths=[0.6, 1.7, 1.8, 2.4]
)

body('The state message is assembled by build_state_message() in guidance_agent.py and fed '
     'as the first task description to Agent 1 of the crew. A complete real-world example is '
     'shown below:')

code_block([
    'PROBE STATE',
    'Region: GSV-THI | Surface: anterior-medial | Leg: right | posY: 0.28',
    '',
    'CONFIRMED FINDINGS',
    '  • EP N1→N2  posY=0.06  right leg',
    '  • RP N2→N1  posY=0.31  right leg',
    '',
    'VLM FRAME ANNOTATION',
    'Fascial layer visible. N2 (saphenous trunk/GSV) within fascial compartment.',
    '',
    'SCAN HISTORY SUMMARY',
    '[DONE] SFJ/groin (0.00–0.07): Entry point at SFJ confirmed.',
    '[DONE] Mid-thigh/Dodd (0.21–0.33): Trunk reflux confirmed.',
    '[    ] Hunterian/lower thigh (0.34–0.47)',
    '[    ] Popliteal/SPJ (0.48–0.57)',
    '',
    'Q1-Q4 STATUS',
    'Q1 ANSWERED: EP N1→N2 at SFJ (right leg, posY 0.06).',
    'Q2 ANSWERED: RP N2→N1 (mid-thigh, posY 0.31).',
    'Q3 OPEN: No EP N2→N3 escape point yet.',
    'Next step: Scan Hunterian zone (posY 0.34–0.47) for escape.',
    '',
    'EXAMINATION PROTOCOL',
    'Hunterian perforators (posY 0.34–0.47, medial). Three maneuvers required:',
    'squeezing → Paranà → Valsalva. Pathological: ≥500 ms AND ≥3.5 mm.',
])

# ═══════════════════════════════════════════════════════════════════════════════
# 6. PRE-PROCESSING SUB-AGENTS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('6. Pre-Processing Sub-Agents')

body('Before the CrewAI pipeline runs, four lightweight modules prepare the context blocks. '
     'These are called sub-agents but run as plain Python code — fast and deterministic. '
     'Only the VLM sub-agent makes an external API call.')

heading2('6.1  History Agent (history_agent.py)')
body('Purpose: Track which leg zones have been visited and summarise what was found.')
body('The agent maintains a 7-band coverage map — SFJ/groin, upper thigh, mid-thigh/Dodd, '
     'Hunterian, popliteal/SPJ, calf, and ankle. Each band is marked [DONE] once the probe '
     'has been in that zone. It then calls Groq LLM (temperature=0) to write a 2-sentence '
     'narrative summary of the full session so far. This gives the main crew agents contextual '
     'memory without repeating all raw coordinates.')
body('Why plain Python for the map, LLM only for the summary? The band marking is deterministic '
     '(a posY number either falls in a range or it does not). The summary requires natural language '
     'generation, so a brief LLM call is appropriate. This is the only sub-agent that calls an LLM.')

heading2('6.2  Q-State Agent (q_state_agent.py)')
body('Purpose: Determine exactly which of Q1–Q4 is currently open from the confirmed clip list.')
body('This is pure Python logic — no LLM call. The agent walks the clip list and checks each Q '
     'criterion in sequence: has any EP been confirmed (Q1)? Has RP N2→N1 been found (Q2)? '
     'Has EP N2→N3 been found (Q3)? Has RP N3→N1 or RP N3→N2 been found (Q4)?')
body('It also handles special cases:')
bullet('Elimination test trigger: if EP N2→N3 + RP N3→N1 + RP N2→N1 are all confirmed but no '
       'elimination test result is recorded, it flags this as the next required step.')
bullet('Complex shunt detection: EP N1→N3 triggers the Type 4/5/6 branch, which requires different '
       'navigation logic from the standard Q1–Q4 sequence.')
body('The Q-state is always logically consistent because it is derived by rule, not inferred by AI.')

heading2('6.3  Protocol Agent (protocol_agent.py)')
body('Purpose: Supply book-sourced examination instructions for the current probe zone.')
body('A pure lookup function — no LLM. Based on the current posY value, it returns the '
     'appropriate protocol text from a pre-built dictionary. All protocol content is sourced '
     'from published clinical literature with specific page references:')
bullet('Adler et al. 2022 (RadioGraphics): standard duplex sequence, reflux thresholds (>500 ms), '
       'patient positioning (Reverse Trendelenburg ≥60°), perforator criteria')
bullet('Gianesini et al. 2014 (Phlebology): CHIVA strategy, SFJ protocol — BOTH Valsalva AND Paranà '
       'required, SPJ protocol (both Paranà + compression/relaxation simultaneously required)')
bullet('Delfrate 2023 (JTAVR): three-maneuver perforator protocol, squeezing-alone limitations, '
       'pelvic leak points, AASV pitfall')
bullet('AVF 2023 guidelines: pathological perforator definition (≥500 ms AND ≥3.5 mm)')

heading2('6.4  VLM Agent (vlm_agent.py)')
body('Purpose: Extract the current video frame and call the Vision Language Model.')
body('A thin wrapper around vlm_analyzer.analyze_frame(). It receives the current posY, '
     'extracts the corresponding frame from the annotated duplex video using OpenCV, and '
     'sends it to the Groq VLM. It returns both the raw JSON result and a plain English '
     'summary string for inclusion in Signal D of the state message.')

# ═══════════════════════════════════════════════════════════════════════════════
# 7. THE VLM
# ═══════════════════════════════════════════════════════════════════════════════
heading1('7. The VLM — Vision Language Model')

heading2('7.1  Purpose')
body('The VLM (Vision Language Model) reads annotated ultrasound video frames and reports '
     'which vessels are visible and how they relate to the fascial layer. This adds real-time '
     'visual confirmation that the probe is in a position suitable for clinical assessment.')

heading2('7.2  The Annotated Video')
body('The system uses a pre-recorded annotated duplex ultrasound video of a real leg examination. '
     'The video has a colour overlay scheme that marks each vessel compartment:')
bullet('Yellow polygon (N3): superficial tributary, above the fascia')
bullet('Green polygon (N2): saphenous trunk inside the "saphenous eye" fascial compartment')
bullet('Blue polygon (N1): deep vein, below the fascia')
bullet('Bright yellow horizontal lines: the fascial layer boundary')
body('The video is synchronised with the leg diagram: posY=0.35 shows the frame at 35% through '
     'the video duration. OpenCV extracts the frame, encodes it as a 70%-quality JPEG '
     '(~16 KB), and converts it to base64 for the VLM API call.')

heading2('7.3  VLM Model — Llama 4 Scout 17B')
body('Model: meta-llama/llama-4-scout-17b-16e-instruct (Groq API)')
body('This multimodal model processes both images and text simultaneously. It receives the '
     'annotated frame plus a context prompt: the current region, leg, and expected anatomy '
     'at this posY position. It returns a structured JSON assessment:')
code_block([
    '{',
    '  "image_quality":             "good" | "poor" | "off-axis",',
    '  "fascial_layer_visible":     true | false,',
    '  "n2_in_fascial_compartment": true | false,',
    '  "n3_superficial_to_fascia":  true | false,',
    '  "n1_deep_to_fascia":         true | false,',
    '  "label_n2_visible":          true | false,',
    '  "frame_note":                "N2 in fascia"',
    '}',
])
body('The VLM only re-runs when posY has changed by ≥ 0.05 since the last call, avoiding '
     'unnecessary API usage on small probe movements where the frame is nearly identical.')

# ═══════════════════════════════════════════════════════════════════════════════
# 8. THE CREWAI 5-AGENT PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════
heading1('8. The CrewAI 5-Agent Pipeline')

heading2('8.1  Why Five Agents Instead of One?')
body('In earlier versions of this system, a single LLM received all seven signals and produced '
     'a guidance instruction directly. This worked but had a fundamental limitation: one model '
     'was simultaneously required to be a clinical interpreter, a classification specialist, '
     'a protocol expert, a navigation planner, and a concise writer — all in one call.')
body('The 5-agent CrewAI design separates these concerns. Each agent has a single focused job '
     'with its own role, goal, and backstory. Each reads only what it needs and passes a '
     'distilled, specialised output to the next agent. The result is more reliable clinical '
     'reasoning and clearer, more accurate guidance.')

heading2('8.2  How Context Flows Through the Crew')
body('CrewAI\'s sequential context mechanism links tasks so each agent automatically sees the '
     'outputs of all upstream agents:')
code_block([
    'Task 1 (Clinical Interpreter) → output: clinical picture (100 words)',
    '  └─ context ──→ Task 2 (Shunt Analyst) → output: shunt type + missing clips (80 words)',
    '                   └─ context ──→ Task 3 (Circuit Analyst) → output: open Q + target zone (60 words)',
    '                                   └─ context ──→ Task 4 (Navigation Planner) → output: scan plan (50 words)',
    '                                                   └─ context ──→ Task 5 (Guidance Specialist)',
    '                                                                   → output: final ≤12-word JSON',
])
body('Agent 5 does not re-read the raw state message — it reads three layers of processed, '
     'increasingly specific output from the agents before it. Each agent narrows the problem '
     'space for the next.')

heading2('8.3  The Five Agents in Detail')

heading3('Agent 1 — Clinical Interpreter  (ported from Task 1)')
add_table(
    ['Property', 'Value'],
    [
        ['Role', 'CHIVA Clinical Interpreter'],
        ['Sees', 'Full state message — all 7 signals'],
        ['Produces', 'Clinical interpretation of the current evidence (≤100 words)'],
        ['Key constraint', 'Never fabricates flow events. If a clip is not in the confirmed list, the agent treats it as absent.'],
    ],
    col_widths=[1.8, 4.7]
)
body('This agent reads the confirmed clips, compares them to the VLM frame annotation, and '
     'assesses the overall clinical picture. It answers four questions:')
bullet('Which clips are unambiguous, and what does each establish?')
bullet('Is there anything suspicious? (e.g., AASV wrongly labelled as GSV)')
bullet('Does the VLM frame support or contradict the confirmed clip list?')
bullet('Which clips are expected but missing given the developing circuit?')
body('Example output: "EP N1→N2 at SFJ (posY 0.06) unambiguous — Mickey Mouse sign confirmed, '
     'both Valsalva and Paranà positive. RP N2→N1 at posY 0.31 confirms trunk reflux. VLM '
     'at current position shows N2 within fascial compartment — consistent with GSV trunk '
     'scanning. EP N2→N3 escape point not yet found — expected in Hunterian zone."')

heading3('Agent 2 — Shunt Analyst  (ported from Task 1)')
add_table(
    ['Property', 'Value'],
    [
        ['Role', 'CHIVA Shunt Classification Specialist'],
        ['Sees', 'Agent 1\'s clinical interpretation'],
        ['Produces', 'Shunt type (or undetermined) + clips still needed (≤80 words)'],
        ['Key constraint', 'Never declares a type confirmed without its full minimum clip set. Always flags when an elimination test is mandatory.'],
    ],
    col_widths=[1.8, 4.7]
)
body('This agent applies formal CHIVA classification rules to the clinical picture. It knows '
     'the exact minimum clip set for every shunt type and explains what still needs to be found. '
     'In Task 2 (unlike Task 1), its output is an intermediate step that feeds navigation '
     'planning — not a final deliverable.')
body('Example output: "Developing Type 1 circuit. EP N1→N2 (SFJ, posY 0.06) + RP N2→N1 '
     '(mid-thigh, posY 0.31) confirmed. Satisfies Type 1 minimum. Missing: EP N2→N3 escape '
     'point. If found, circuit upgrades to Type 3 or 1+2, requiring elimination test."')

heading3('Agent 3 — Circuit Analyst  (new in Task 2)')
add_table(
    ['Property', 'Value'],
    [
        ['Role', 'CHIVA Circuit Analyst'],
        ['Sees', 'Agent 2\'s classification + Q1–Q4 status (Signal E)'],
        ['Produces', 'The specific open Q and the exact anatomical zone to examine (≤60 words)'],
        ['Key constraint', 'Handles edge cases, e.g., after EP N1→N3 at Giacomini, must direct to medial upper thigh (not distally) to check for RP N2→N1.'],
    ],
    col_widths=[1.8, 4.7]
)
body('This agent converts the classification narrative into a concrete navigation target. It '
     'answers: "Given these clips and this developing shunt type, which diagnostic question is '
     'open, and exactly where on the leg (with posY band) should the probe go next?"')
body('Example output: "Q3 open: no escape from GSV trunk confirmed yet. Target Hunterian zone, '
     'posY 0.34–0.47, medial surface — scan transversely for N3 above fascia. Also check '
     'Dodd zone (0.21–0.33). No elimination test trigger yet."')

heading3('Agent 4 — Navigation Planner  (new in Task 2)')
add_table(
    ['Property', 'Value'],
    [
        ['Role', 'CHIVA Navigation Planner'],
        ['Sees', 'Agent 3\'s circuit analysis + examination protocol (Signal F)'],
        ['Produces', 'Specific posY band, probe surface, anatomical target, direction word, and maneuver (≤50 words)'],
        ['Key constraint', 'Applies clinical protocol rules: BOTH Valsalva AND Paranà at SFJ/SPJ; three maneuvers at perforators; 500 ms + 3.5 mm pathological thresholds.'],
    ],
    col_widths=[1.8, 4.7]
)
body('This agent converts the abstract navigation target from Agent 3 into a concrete scanning '
     'plan, adding the clinical protocol layer — which examination maneuvers apply at this zone '
     'and what criteria define a pathological finding.')
body('Example output: "Target posY 0.34–0.47, medial surface. Scan distally from current '
     'position. Target: N3 perforator exiting fascia above GSV. Three maneuvers: squeezing → '
     'Paranà → Valsalva. Pathological if outward flow ≥500 ms AND ≥3.5 mm (Delfrate 2023)."')

heading3('Agent 5 — Guidance Specialist  (new in Task 2)')
add_table(
    ['Property', 'Value'],
    [
        ['Role', 'CHIVA Real-Time Guidance Specialist'],
        ['Sees', 'Agents 2, 3, and 4\'s outputs'],
        ['Produces', 'Single JSON object: {"guidance": "<≤12 words>", "action": "move|maneuver|complete"}'],
        ['Key constraint', 'FORBIDDEN in guidance text: EP, RP, N1, N2, N3, reflux, shunt type, Q1–Q4, confirmed, findings, diagnostic.'],
    ],
    col_widths=[1.8, 4.7]
)
body('This is the final output agent. It synthesises everything upstream and produces the '
     'instruction that appears on screen for the surgeon. The three action types are:')
add_table(
    ['Action', 'When used', 'Example guidance text'],
    [
        ['"move"', 'Default probe navigation (≥95% of turns)', '"Scan distally along medial thigh toward Hunterian perforator"'],
        ['"maneuver"', 'Surgeon must perform the elimination test (compress a tributary and record whether GSV Doppler changes)', '"Compress tributary at mid-thigh and record Doppler response"'],
        ['"complete"', 'Full CHIVA circuit is fully mapped', '"Circuit mapped — sufficient findings for classification"'],
    ],
    col_widths=[1.0, 2.3, 3.2]
)

heading2('8.4  Shunt Detection — By-Product of the Pipeline')
body('As a by-product of every crew run, Agent 2\'s output text is parsed for shunt '
     'confirmation. When a type is fully confirmed:')
bullet('The backend emits a shunt_confirmed WebSocket event to the browser')
bullet('A confirmation modal appears showing the shunt type and clip evidence')
bullet('The confirmed_shunts set in StreamSession is updated — the same type is never '
       'announced twice in the same session')
body('If the surgeon dismisses the modal (disagreeing with the classification), the system '
     'records a rejection note. On the next crew run, this is prepended to Agent 1\'s task: '
     '"SURGEON FEEDBACK — PRIOR CLASSIFICATION REJECTED. Re-evaluate more critically — do '
     'not re-confirm without meaningfully different clip evidence."')

heading2('8.5  The Groq + LiteLLM Compatibility Fix')
body('CrewAI uses LiteLLM internally to route API calls. LiteLLM injects a "cache_breakpoint" '
     'field into message dictionaries, which Groq\'s API rejects with a 400 error (Groq '
     'validates message structure strictly). The pipeline applies a monkey-patch at module '
     'load time to strip this field before every API call:')
code_block([
    '_orig_completion = litellm.completion',
    '',
    'def _patched_completion(*args, **kwargs):',
    '    for msg in kwargs.get("messages", []):',
    '        msg.pop("cache_breakpoint", None)',
    '    return _orig_completion(*args, **kwargs)',
    '',
    'litellm.completion = _patched_completion',
])
body('This adds zero latency, requires no configuration, and is entirely invisible to the '
     'rest of the application. It was a necessary fix to make the full CrewAI pipeline '
     'run stably on Groq infrastructure.')

# ═══════════════════════════════════════════════════════════════════════════════
# 9. POSITION ALERTS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('9. Position Alerts — Hard-Coded Safety Rules')

body('In addition to the seven signals, three hard-coded position alerts are appended to the '
     'state message when the probe is at a critical decision zone. These override any tendency '
     'for the AI to give incorrect guidance at the most clinically important positions.')

add_table(
    ['Alert', 'Trigger condition', 'What is added to state message', 'Forced behaviour'],
    [
        ['SFJ anchoring', 'posY 0.04–0.07 AND no clips yet', '"Probe IS at SFJ zone. Apply Mickey Mouse transverse scan. Output transverse-scan instruction — do NOT navigate away."', 'Forces transverse scan instruction. Prevents directing surgeon away from SFJ before junction is assessed.'],
        ['Escape search', 'posY 0.08–0.47 AND RP N2→N1 confirmed AND no EP N2→N3', '"Trunk reflux confirmed; no escape found. Scan for N3 above fascia at this level — do NOT output complete."', 'Prevents declaring circuit complete when escape point has not been found.'],
        ['SPJ assessment', 'posY 0.48–0.57, posterior, no SPJ entry confirmed', '"Probe in POPLITEAL zone. Apply Paranà + CR maneuvers to assess SPJ."', 'Ensures both active (Paranà) and passive (compression/relaxation) maneuvers are applied at the SPJ.'],
    ],
    col_widths=[1.2, 1.8, 2.0, 1.5]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 10. FRONTEND
# ═══════════════════════════════════════════════════════════════════════════════
heading1('10. Frontend — The Streaming UI')

heading2('10.1  Three-Panel Layout')
add_table(
    ['Panel', 'Contents'],
    [
        ['Left — Probe Control', 'Interactive SVG leg diagram (surgeon clicks to move probe, red dot tracks position). Mark Finding form (records EP/RP clips with flow type, compartments, posY, leg, optional elimination test result). Confirmed clips list with delete buttons.'],
        ['Centre — Visual', 'Current annotated video frame (synced to posY). VLM Frame Analysis grid showing fascial layer, N1/N2/N3 visibility, and image quality. Elimination test banner (appears when action = "maneuver").'],
        ['Right — Guidance', 'Active Guidance box — current ≤12-word instruction in large text, updated in real time. Model Thinking Log — expandable entries for every turn, showing the exact state message and raw crew output.'],
    ],
    col_widths=[1.5, 5.0]
)

heading2('10.2  Shunt Confirmed Modal')
body('When Agent 2 detects a confirmed CHIVA circuit, a modal dialog appears showing the shunt '
     'type, the clip evidence, and two buttons: Confirm (accepts the classification) or Dismiss '
     '(records a rejection note and triggers the feedback loop described in Section 8.4). '
     'The modal fires only once per type per session.')

heading2('10.3  Real-Time Communication')
body('All guidance updates, clip marks, and session events travel over WebSocket (Socket.IO), '
     'providing live updates without page reloads. This required Flask-SocketIO with threading '
     'async_mode on the backend, and Gunicorn configured with gthread workers (1 worker, '
     '4 threads) to support persistent WebSocket connections.')

# ═══════════════════════════════════════════════════════════════════════════════
# 11. COMPARISON WITH TASK 1
# ═══════════════════════════════════════════════════════════════════════════════
heading1('11. Comparison with Task 1')

add_table(
    ['Aspect', 'Task 1', 'Task 2'],
    [
        ['Timing', 'After examination (retrospective)', 'During examination (real-time)'],
        ['Input', 'Text description of completed case', 'Live probe clicks + video frames'],
        ['Knowledge base', 'Qdrant vector database (RAG pipeline)', 'Embedded in agent system prompts'],
        ['Output', 'Full clinical report and surgical plan', '≤12-word probe movement instruction'],
        ['AI architecture', 'Single LLM + RAG retrieval + cross-encoder reranker', '5-agent CrewAI sequential crew'],
        ['Communication', 'HTTP REST (synchronous)', 'WebSocket / Socket.IO (real-time)'],
        ['Vision model', 'Not used', 'Llama 4 Scout 17B on annotated video frames'],
        ['Agents used', '3 (Interpreter, Shunt Analyst, General Assistant)', '5 (3 ported from Task 1 + 2 new navigation agents)'],
    ],
    col_widths=[1.8, 2.3, 2.4]
)

body('Task 2 inherits Agents 1 and 2 (Clinical Interpreter and Shunt Analyst) from Task 1. '
     'Agents 3, 4, and 5 (Circuit Analyst, Navigation Planner, Guidance Specialist) are new '
     'additions specific to the real-time navigation problem.')

# ═══════════════════════════════════════════════════════════════════════════════
# 12. KEY DESIGN DECISIONS
# ═══════════════════════════════════════════════════════════════════════════════
heading1('12. Key Design Decisions and Rationale')

heading2('12.1  Why ≤12 Words for Guidance?')
body('The surgeon is actively performing an examination. They cannot read a paragraph while '
     'holding an ultrasound probe over a patient. 12 words is long enough to be specific '
     '("Scan distally along medial thigh toward Hunterian perforator") but short enough to '
     'process in a single glance. The constraint also prevents the model from adding clinical '
     'explanations that the surgeon does not need in real time.')

heading2('12.2  Why Sequential Agents and Not Parallel?')
body('The five agents must run sequentially because each depends on the previous one. Agent 5 '
     'cannot produce "scan toward Hunterian perforator" without Agent 3 first determining '
     'that the Hunterian zone is the correct target, which Agent 3 cannot do without Agent 2\'s '
     'classification, which Agent 2 cannot do without Agent 1\'s clinical interpretation. '
     'Parallelising the agents would break the reasoning chain.')

heading2('12.3  Why Sub-Agents as Plain Python Instead of LLM?')
body('The history coverage map, Q-state derivation, and protocol lookup were initially designed '
     'as LLM calls but converted to deterministic Python for three reasons:')
bullet('Speed: Plain Python runs in milliseconds; an LLM call adds 1–3 seconds. With 5 agents '
       'already in the crew, avoiding 3 additional LLM calls keeps the pipeline usable in '
       'real time.')
bullet('Reliability: The Q-state logic (which Q is open) must be correct 100% of the time. '
       'Pure rule-based logic cannot hallucinate.')
bullet('Cost: Three fewer Groq API calls per guidance turn significantly reduces API usage '
       'over a full examination session.')

heading2('12.4  Why Not Use Conversation History (as in the Earlier Version)?')
body('An earlier streaming engine version used a rolling conversation window (last 8 LLM turns), '
     'similar to a chatbot. This was replaced with the CrewAI approach because:')
bullet('Rolling history accumulates outdated context. What the LLM said 8 clicks ago may no '
       'longer be clinically relevant at the current probe position.')
bullet('Each new state message now contains explicit, structured summaries (scan coverage from '
       'history_agent, Q-state from q_state_agent) rather than requiring the LLM to infer '
       'session context from conversation history.')
bullet('Specialised agents with focused roles produce more reliable reasoning than a single '
       'LLM simultaneously tracking all dimensions of the examination.')

heading2('12.5  Why opencv-python-headless Instead of opencv-python?')
body('The standard opencv-python package links against OpenGL (libGL.so.1) — a display library '
     'that is not present on headless Linux servers. The headless variant provides identical '
     'image processing and video decoding capabilities but without the display dependency. '
     'This is the standard approach for any server that uses OpenCV only for image/video '
     'processing, not GUI windows.')

# ═══════════════════════════════════════════════════════════════════════════════
# 13. CLINICAL KNOWLEDGE SOURCES
# ═══════════════════════════════════════════════════════════════════════════════
heading1('13. Clinical Knowledge Sources')

body('All clinical protocol rules, thresholds, and anatomical facts used in the system are '
     'derived from published, peer-reviewed literature. Every fact in the agent system prompts '
     'and protocol lookup table is traceable to a specific reference:')

add_table(
    ['Reference', 'Journal / Publisher', 'Used for'],
    [
        ['Adler et al. 2022', 'RadioGraphics, 42(7), pp. 2184–2202', 'Standard duplex examination sequence, patient positioning (Reverse Trendelenburg ≥60°), reflux threshold (>500 ms), Dodd and Hunterian anatomy, perforator criteria, posY landmarks'],
        ['Gianesini et al. 2014', 'Phlebology, 29(1)', 'CHIVA haemodynamic classification framework, SFJ protocol (Mickey Mouse sign, BOTH Valsalva AND Paranà required), SPJ protocol (both Paranà + compression/relaxation simultaneously), Giacomini vein evaluation'],
        ['Delfrate R. 2023', 'JTAVR (Journal of Translational Anatomy and Vascular Research)', 'Three-maneuver perforator protocol (squeezing → Paranà → Valsalva), squeezing-alone limitations, pelvic leak points (SGP, IGP, OP), biphasic perforator interpretation, AASV pitfall'],
        ['AVF 2023', 'American Venous Forum Clinical Practice Guidelines', 'Pathological perforator definition: outward flow ≥500 ms AND diameter ≥3.5 mm'],
        ['DuplexUS 2014', 'Textbook, Chapter 3 (pp. 33–41)', 'Dodd perforators (middle thigh, posY 0.21–0.33), Hunterian perforators (distal thigh, posY 0.34–0.47), SPJ anatomical variability'],
        ['Lee et al. 2017', 'Textbook, p. 129', 'Additional Hunterian perforator posY landmark references'],
    ],
    col_widths=[1.5, 1.9, 3.1]
)

# ═══════════════════════════════════════════════════════════════════════════════
# 14. DEPLOYMENT
# ═══════════════════════════════════════════════════════════════════════════════
heading1('14. Deployment Architecture')

body('Task 2 is deployed on an Alibaba Cloud ECS (Linux) server, running alongside Task 1 '
     'on the same machine. The two applications do not interfere — they use separate Unix '
     'sockets, separate systemd services, and separate Nginx server blocks on different ports.')

add_table(
    ['Component', 'Task 1', 'Task 2'],
    [
        ['Port (Nginx)', '80 (HTTP)', '8080'],
        ['Systemd service', 'shunt', 'task2'],
        ['Gunicorn worker', 'sync (1 worker)', 'gthread (1 worker, 4 threads)'],
        ['Unix socket', '/var/www/shunt_classification_and_ligation/shunt.sock', '/var/www/task2/task2.sock'],
        ['Log directory', '/var/log/shunt_classification_and_ligation/', '/var/log/task2_probe_guidance/'],
        ['CI/CD', 'GitHub Actions → deploy.yml → SSH → update.sh', 'Same pattern — separate deploy.yml and update.sh'],
    ],
    col_widths=[1.5, 2.3, 2.7]
)

body('The deployment chain on every git push to main:')
code_block([
    'Developer pushes to GitHub (main branch)',
    '  → GitHub Actions triggers Task_2_App/.github/workflows/deploy.yml',
    '    → SSH into Alibaba Cloud ECS server',
    '      → git pull origin main',
    '        → pip install -r requirements.txt  (in Task_2_App venv)',
    '          → systemctl restart task2',
    '            → Gunicorn starts, creates Unix socket',
    '              → Nginx proxies port 8080 → socket (with WebSocket headers)',
])

body('Nginx requires specific WebSocket proxy headers that are not needed for Task 1\'s '
     'standard HTTP REST traffic:')
code_block([
    'proxy_http_version 1.1;',
    'proxy_set_header Upgrade $http_upgrade;',
    'proxy_set_header Connection "upgrade";',
])
body('These headers enable the Socket.IO WebSocket upgrade handshake through Nginx.')

# ── Save ──────────────────────────────────────────────────────────────────────
out = r'C:\Users\Krish\Downloads\Cygnus_Med_Demo\Task_2_App\Task2_Methodology_Report.docx'
doc.save(out)
print(f'Saved: {out}')
