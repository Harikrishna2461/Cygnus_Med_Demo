"""
Generates Task2_Supplement_ReproducibilityAndLessonsLearned.docx
Covers: How to reproduce the system + What worked / What did not work
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import copy

OUTPUT = r"C:\Users\Krish\Downloads\Cygnus_Med_Demo\Task_2_App\Task2_Supplement_ReproducibilityAndLessonsLearned.docx"

NAVY   = RGBColor(0x1F, 0x39, 0x64)
TEAL   = RGBColor(0x1F, 0x7A, 0x8C)
GREY   = RGBColor(0x44, 0x44, 0x44)
GREEN  = RGBColor(0x1E, 0x6B, 0x2E)
RED    = RGBColor(0x8B, 0x00, 0x00)
AMBER  = RGBColor(0x7B, 0x4F, 0x00)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BLUE = RGBColor(0xDB, 0xE5, 0xF1)
LIGHT_GREEN = RGBColor(0xE2, 0xEF, 0xDA)
LIGHT_RED   = RGBColor(0xFC, 0xE4, 0xD6)
LIGHT_AMBER = RGBColor(0xFF, 0xF2, 0xCC)
SHADE_GREY  = RGBColor(0xF2, 0xF2, 0xF2)

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def add_bottom_border(para):
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "1F3964")
    pBdr.append(bot)
    pPr.append(pBdr)


def heading1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    add_bottom_border(p)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = NAVY
    return p


def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = TEAL
    return p


def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = GREY
    return p


def body(text, bold=False, color=None):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.bold       = bold
    if color:
        run.font.color.rgb = color
    return p


def bullet(text, level=0, color=None, prefix="•"):
    p   = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.3 * (level + 1))
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(f"{prefix}  {text}" if level == 0 else f"    {prefix}  {text}")
    run.font.size = Pt(10.5)
    if color:
        run.font.color.rgb = color
    return p


def tick(text):   return bullet(text, color=GREEN,  prefix="✔")
def cross(text):  return bullet(text, color=RED,    prefix="✘")
def warn(text):   return bullet(text, color=AMBER,  prefix="⚠")


def code_block(text):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1F, 0x39, 0x64)
    return p


def add_table(headers, rows, col_widths=None, row_colors=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for j, h in enumerate(headers):
        cell = tbl.cell(0, j)
        set_cell_bg(cell, NAVY)
        run = cell.paragraphs[0].add_run(h)
        run.bold            = True
        run.font.color.rgb  = WHITE
        run.font.size       = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for i, row in enumerate(rows):
        bg = (row_colors[i] if row_colors and i < len(row_colors)
              else (LIGHT_BLUE if i % 2 == 0 else RGBColor(0xFF,0xFF,0xFF)))
        for j, val in enumerate(row):
            cell = tbl.cell(i + 1, j)
            set_cell_bg(cell, bg)
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)

    if col_widths:
        for i, row in enumerate(tbl.rows):
            for j, cell in enumerate(row.cells):
                if j < len(col_widths):
                    cell.width = Inches(col_widths[j])

    doc.add_paragraph()
    return tbl


def page_break():
    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Task 2 — Supplementary Document")
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Reproducibility Guide & Lessons Learned")
sr.font.size = Pt(15)
sr.font.color.rgb = TEAL

doc.add_paragraph()
meta_lines = [
    ("System", "AI-Guided CHIVA Ultrasound Examination — Streaming Guidance Engine"),
    ("Supplement Type", "Reproducibility Steps + What Worked / What Did Not Work"),
    ("Stack", "Flask · SocketIO · CrewAI · Groq LLM/VLM · Nginx · Gunicorn · systemd"),
    ("Server", "Alibaba Cloud ECS — Ubuntu 22.04 LTS  |  IP: 8.219.69.76  |  Port: 8080"),
    ("Date", "July 2026"),
]
for label, val in meta_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = GREY
    r2 = p.add_run(val)
    r2.font.size = Pt(11)

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PREREQUISITES
# ══════════════════════════════════════════════════════════════════════════════

heading1("1.  Prerequisites")
body("Before you can reproduce the system you need the following accounts, keys, and infrastructure in place.")

heading2("1.1  Server Infrastructure")
add_table(
    ["Requirement", "Specification", "Notes"],
    [
        ["Cloud provider", "Alibaba Cloud ECS (or any Ubuntu 22.04 VPS)", "Other clouds work — only the security-group UI differs"],
        ["vCPU / RAM", "2 vCPU / 4 GB minimum", "LLM calls go to Groq cloud — no local GPU needed"],
        ["OS", "Ubuntu 22.04 LTS", "Tested on this version only"],
        ["Open ports", "8080 (HTTP app)  |  22 (SSH)", "Port 80 is used by Task-1 — Task-2 runs on 8080"],
        ["SSH key pair", ".pem file downloaded at instance creation", "Keep it — needed for GitHub Actions secret"],
    ],
    col_widths=[1.6, 2.8, 2.8],
)

heading2("1.2  API Keys")
add_table(
    ["Key", "Where to get it", "Used for"],
    [
        ["GROQ_API_KEY", "console.groq.com → API Keys", "All LLM calls (llama-3.3-70b) + VLM calls (llama-4-scout-17b)"],
        ["GitHub PAT", "github.com → Settings → Developer settings → Tokens", "Git push from server; stored in deploy/update.sh remote URL"],
        ["GitHub Actions SSH secret", "Repository → Settings → Secrets → TASK2_SSH_KEY", "Paste full contents of .pem file"],
    ],
    col_widths=[1.8, 2.5, 2.9],
)

heading2("1.3  Local Machine")
bullet("Python 3.11+ with python-docx installed (only needed to regenerate this document)")
bullet("Git configured with push access to the repository")
bullet("SSH client (e.g. OpenSSH or PuTTY) to connect to the server for first-time setup")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — HOW TO REPRODUCE (Step by Step)
# ══════════════════════════════════════════════════════════════════════════════

heading1("2.  How to Reproduce the System End-to-End")

heading2("Step 1 — Clone the Repository")
body("On your local machine or directly on the server:")
code_block("git clone https://github.com/Harikrishna2461/Cygnus_Med_Demo.git\ncd Cygnus_Med_Demo/Task_2_App")

heading2("Step 2 — Create the .env File")
body("Copy the example environment file and fill in your Groq API key:")
code_block("cp .env.example .env\nnano .env     # or any editor")
body("The .env file must contain:")
code_block("GROQ_API_KEY=gsk_<your-key-here>\nSTREAM_VIDEO_PATH=/var/www/Cygnus_Med_Demo/Task_2_App/assets/stream_video.mp4")
warn("Never commit the .env file to git — it is already listed in .gitignore.")

heading2("Step 3 — SSH Into the Server")
body("From your local machine (where your .pem file lives):")
code_block('ssh -i "Alibaba_Cloud_Key.pem" root@8.219.69.76')
body("If permission is denied, fix the key permissions first:")
code_block("chmod 400 Alibaba_Cloud_Key.pem   # Linux/macOS\n# Windows: right-click .pem → Properties → Security → remove all except your user")

heading2("Step 4 — Pull the Latest Code on the Server")
code_block("cd /var/www/Cygnus_Med_Demo\ngit pull origin main")

heading2("Step 5 — Run the One-Time Setup Script")
body("This script creates directories, a Python virtual environment, installs all dependencies, registers the systemd service, and configures Nginx.")
code_block("bash /var/www/Cygnus_Med_Demo/Task_2_App/deploy/setup.sh")
body("The script performs the following operations in order:")
add_table(
    ["Step inside setup.sh", "What it does"],
    [
        ["Create directories", "mkdir -p /var/www/task2  /var/log/task2_probe_guidance"],
        ["Install system libs", "apt-get install -y libgl1  (fixes OpenCV on headless servers — critical)"],
        ["Create Python venv", "python3 -m venv /var/www/task2/venv"],
        ["Install Python deps", "pip install -r Task_2_App/backend/requirements.txt"],
        ["Copy .env", "cp Task_2_App/.env.example Task_2_App/.env  (edit manually after)"],
        ["Install systemd service", "cp task2.service /etc/systemd/system/ && systemctl enable task2"],
        ["Configure Nginx", "cp nginx.conf /etc/nginx/sites-available/task2  (port 8080)"],
        ["Reload Nginx", "nginx -t && systemctl reload nginx"],
    ],
    col_widths=[2.5, 4.7],
)

heading2("Step 6 — Add the Real Groq API Key to .env")
code_block("nano /var/www/Cygnus_Med_Demo/Task_2_App/.env\n# Set: GROQ_API_KEY=gsk_<your-real-key>")

heading2("Step 7 — Start the Service")
code_block("systemctl start task2\nsystemctl status task2   # should show 'active (running)'")

heading2("Step 8 — Verify the App is Running")
code_block("# Local check (on server):\ncurl http://localhost:8080/api/status\n\n# Remote check (from your machine):\ncurl http://8.219.69.76:8080/api/status")
body("Expected response: HTTP 200 with a JSON status payload. If you get 502 Bad Gateway, see Section 3.2.")

heading2("Step 9 — Upload the Video File")
body("The streaming demo requires a sample ultrasound video. Copy it from your local machine to the server:")
code_block('scp -i "Alibaba_Cloud_Key.pem" stream_video.mp4 root@8.219.69.76:/var/www/Cygnus_Med_Demo/Task_2_App/assets/')
body("Then verify the STREAM_VIDEO_PATH in .env matches this path.")

heading2("Step 10 — Configure GitHub Actions for CI/CD")
body("Every push to the main branch should auto-deploy. For this to work:")
bullet("Go to the GitHub repository → Settings → Secrets and variables → Actions")
bullet("Add secret: TASK2_SSH_KEY → paste the full contents of your .pem file")
bullet("The workflow file at .github/workflows/deploy.yml is already committed")
body("From this point, git push origin main triggers SSH into the server, git pull, pip install, and service restart automatically.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — WHAT WORKED
# ══════════════════════════════════════════════════════════════════════════════

heading1("3.  What Worked")

heading2("3.1  Core Application Architecture")
tick("Flask + Flask-SocketIO with async_mode='threading' — stable under concurrent WebSocket connections and does not require an async event loop (asyncio/gevent), which would conflict with CrewAI's synchronous Groq calls.")
tick("simple-websocket transport — lighter than gevent-websocket and compatible with the threading mode; no C-extension compile step needed on the server.")
tick("CrewAI 5-agent sequential pipeline — agents reliably pass context to each other via the task context= parameter; agent 5 (guidance specialist) consistently produces valid JSON from the combined output of agents 1-4.")
tick("LiteLLM monkey-patch for Groq — stripping the cache_breakpoint field from LiteLLM's request before it reaches Groq resolved a silent 400 error that caused the crew to silently fail. This one-line patch in crew_pipeline.py was robust across all test scenarios.")
tick("Sub-agents (history, q_state, protocol) as plain Python functions — keeping these outside CrewAI avoids unnecessary LLM calls for logic that is deterministic (Q-state derivation, protocol lookup) or needs to run before the crew.")
tick("Groq llama-3.3-70b-versatile for text agents — fast inference (<3 s per crew run) with consistent JSON output when given strict output schemas in the task prompts.")
tick("Groq llama-4-scout-17b for VLM — correctly identifies fascial layer visibility and N1/N2/N3 needle positions from annotated ultrasound frames.")

heading2("3.2  Deployment Infrastructure")
tick("systemd service with ExecStartPre for git pull + pip install — ensures the server is always on the latest committed code before Gunicorn starts, without needing a separate deployment pipeline.")
tick("Gunicorn with worker_class='gthread' and threads=4 — the only Gunicorn configuration that works with Flask-SocketIO threading mode; sync and gevent workers both failed.")
tick("Unix socket for Gunicorn↔Nginx IPC (task2.sock) — lower latency than TCP localhost:port and avoids port conflicts when multiple apps run on the same server.")
tick("Nginx WebSocket proxy with Upgrade + Connection headers — without these two headers the WebSocket handshake fails silently; adding them resolved all WebSocket connectivity issues.")
tick("Port 8080 for Task-2 alongside Task-1 on port 80 — two independent systemd services (task1 and task2), two Unix sockets, two Nginx server blocks coexist cleanly with no interference.")
tick("GitHub Actions SSH deploy (.github/workflows/deploy.yml) — push-to-deploy pipeline works reliably; the update.sh script is idempotent so repeated pushes do not break anything.")
tick("libgl1 system package (apt-get install -y libgl1) — once installed, OpenCV imports without error on headless Ubuntu; the fix persists across reboots and service restarts.")
tick("opencv-python-headless in requirements.txt — the headless variant avoids pulling in GTK/Qt display libraries that are meaningless on a server and would cause import errors.")

heading2("3.3  Frontend")
tick("Interactive leg SVG with mousemove → probe_move WebSocket events — smooth real-time position updates with no perceptible lag.")
tick("Collapsible thinking log — the per-turn expandable entries (posY, region, raw state message, raw LLM response) are essential for debugging crew behaviour without cluttering the UI.")
tick("Shunt confirmed modal (fires once per type per session) — the confirmed_shunts set in StreamSession correctly prevents re-firing even when the LLM repeats a confirmed finding.")
tick("Elimination test banner (action='maneuver') — visible yellow overlay gives clear clinical cue during the elimination test step.")

heading2("3.4  Clinical Logic")
tick("posY coordinate system (0.0 = groin → 1.0 = ankle) — a single continuous float cleanly encodes probe position along the leg without region-specific edge cases.")
tick("7-band posY landmark zones in history_agent — maps position to anatomical region (SFJ, thigh, Dodd, popliteal, calf, SPJ, ankle) for concise history summaries.")
tick("Q1–Q4 state derivation from clip types (q_state_agent.py) — pure Python logic with no LLM call; produces a reliable boolean status that the crew can reason over.")
tick("Sliding-window segment stabilisation (10 readings, 60% threshold) in probe_localizer.py — prevents spurious region changes from single noisy probe-move events.")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — WHAT DID NOT WORK
# ══════════════════════════════════════════════════════════════════════════════

heading1("4.  What Did Not Work (and How It Was Resolved)")

heading2("4.1  Deployment Blockers")

heading3("4.1.1  OpenCV ImportError — Missing libGL.so.1  [CRITICAL]")
body("This was the primary deployment blocker. The service crashed on every start attempt with:")
code_block("ImportError: libGL.so.1: cannot open shared object file: No such file or directory")
body("Root cause: The standard opencv-python package links against OpenGL display libraries (libGL, libGLX, etc.) that are only present on desktop Linux. A minimal cloud server image (Ubuntu 22.04 ECS) does not install these by default.")

add_table(
    ["Fix Attempted", "Result"],
    [
        ["pip install opencv-python-headless (without first removing opencv-python)",
         "Failed — both packages were simultaneously installed (pip list showed opencv-python 4.13 AND opencv-python-headless 4.13). Python loaded whichever it found first in sys.path — still the non-headless one."],
        ["pip uninstall opencv-python (remove non-headless only, keep headless)",
         "Failed — even the headless build on this specific Ubuntu image still linked against libGL due to how the wheel was compiled."],
        ["apt-get install -y libgl1 (install the missing system library)",
         "SUCCESS — resolves the missing library at the OS level. Works for both opencv-python and opencv-python-headless variants. Persists across reboots."],
    ],
    col_widths=[2.8, 4.4],
    row_colors=[LIGHT_RED, LIGHT_RED, LIGHT_GREEN],
)
body("Lesson: Always add 'apt-get install -y libgl1' to your server setup script before starting any service that imports cv2, regardless of which OpenCV pip variant you use.")

heading3("4.1.2  Service Showing 'active' Briefly Then Crashing")
body("systemctl status showed 'active (running)' for approximately 2 seconds before switching to 'activating (auto-restart)'. This led to initial confusion about whether the fix had worked.")
code_block("# Misleading:\nsystemctl status task2   # shows 'active' briefly\n\n# Reliable diagnostic:\ncat /var/log/task2_probe_guidance/gunicorn-error.log")
body("Lesson: Always check the Gunicorn error log, not just systemctl status. systemd marks a service active as soon as the process starts — it takes a moment to detect a crash.")

heading3("4.1.3  502 Bad Gateway From Nginx")
body("Nginx returned 502 when Gunicorn was crashing before creating the Unix socket file. Nginx had nothing to proxy to.")
code_block("ls -la /var/www/task2/   # socket file missing → confirms Gunicorn never started cleanly\nps aux | grep gunicorn  # only Task-1 Gunicorn processes visible")
body("Resolution: Same as 4.1.1 — fix the Gunicorn crash first; 502 resolves automatically once Gunicorn creates the socket.")

heading3("4.1.4  Gunicorn Worker Class — sync vs gthread")
body("Using the default Gunicorn sync worker class caused Flask-SocketIO WebSocket connections to hang indefinitely. The sync worker handles one request per worker with no threading, which cannot serve the long-lived WebSocket connection while also processing events.")
code_block("# Wrong (default sync):\nbind = \"unix:/var/www/task2/task2.sock\"\nworkers = 4\n\n# Correct (gthread for SocketIO threading mode):\nbind = \"unix:/var/www/task2/task2.sock\"\nworkers = 1\nworker_class = \"gthread\"\nthreads = 4")
body("Lesson: Flask-SocketIO with async_mode='threading' requires Gunicorn's gthread worker. Do not use sync, gevent, or eventlet.")

heading2("4.2  LLM / Agent Issues")

heading3("4.2.1  Groq Rejecting cache_breakpoint Field")
body("CrewAI uses LiteLLM to route LLM calls. LiteLLM added a cache_breakpoint field to requests in a newer version; Groq's API does not recognise this field and returns a 400 error, causing the crew to silently fail with no guidance output.")
code_block("# Fix in crew_pipeline.py (before calling crew.kickoff):\nimport litellm\n_orig = litellm.utils.get_optional_params\ndef _patched(*args, **kwargs):\n    result = _orig(*args, **kwargs)\n    result.pop(\"cache_breakpoint\", None)\n    return result\nlitellm.utils.get_optional_params = _patched")
body("Lesson: When integrating CrewAI + Groq, always check whether LiteLLM is sending fields that Groq does not support. The monkey-patch pattern above is a reliable workaround until upstream fixes the issue.")

heading3("4.2.2  Single LLM Fallback Was Insufficient for Complex Shunt Classification")
body("The original architecture used a single LLM call with a large system prompt to produce guidance + shunt classification simultaneously. Testing showed that the single-call approach frequently:")
cross("Confused shunt type rules when multiple clip types were present")
cross("Missed elimination test logic (Type 3 vs Type 1+2 disambiguation)")
cross("Produced inconsistent JSON when the reasoning was complex")
body("Resolution: Replaced with the 5-agent CrewAI pipeline where each agent has a focused role. Agent 2 (shunt analyst) has sole responsibility for shunt classification and outputs a dedicated JSON block, completely decoupled from guidance text generation.")

heading3("4.2.3  LLM Temperature Tuning")
body("Early tests with temperature=0.7 produced variable guidance phrasing that was clinically imprecise. Tests with temperature=0.0 caused the model to occasionally produce identical guidance for different positions ('frozen' outputs).")
add_table(
    ["Temperature", "Behaviour", "Verdict"],
    [
        ["0.7", "Too variable — inconsistent phrasing, occasional hallucinations in shunt classification", "Rejected"],
        ["0.0", "Too rigid — occasional repeated identical outputs across different states", "Rejected"],
        ["0.3 (final)", "Consistent clinical language with appropriate variation in navigation instructions", "Adopted"],
    ],
    col_widths=[1.5, 4.2, 1.5],
    row_colors=[LIGHT_RED, LIGHT_RED, LIGHT_GREEN],
)

heading3("4.2.4  VLM Frame Analysis Latency")
body("The Groq llama-4-scout-17b VLM call adds approximately 2–4 seconds per crew run when an annotated frame is available. Setting the VLM position threshold too low (STREAM_VLM_THRESHOLD=0.03) caused a VLM call on almost every probe movement, making the UI feel sluggish.")
body("Resolution: Raised STREAM_VLM_THRESHOLD to 0.05 and STREAM_LLM_THRESHOLD to 0.06 in config.py. The VLM now fires only when the probe has moved meaningfully, not on every mouse event.")

heading2("4.3  Development / Tooling Issues")

heading3("4.3.1  Both opencv Variants Installed Simultaneously")
body("Running pip install opencv-python-headless without first uninstalling opencv-python resulted in both packages coexisting in the virtual environment. pip list showed both at the same version. Python loaded whichever package provided the cv2 module first in the import path — not necessarily the headless one.")
code_block("# Diagnostic:\npip list | grep opencv\n# Bad output example:\n# opencv-python             4.13.0.92\n# opencv-python-headless   4.13.0.92\n\n# Fix:\npip uninstall opencv-python -y\n# Then verify:\npip list | grep opencv\n# Should show only: opencv-python-headless")
body("Lesson: When switching between opencv variants, always explicitly uninstall the old one first. The two packages provide the same cv2 module — having both is ambiguous.")

heading3("4.3.2  Alibaba Cloud Security Group Port Configuration")
body("The security group rule creation form had the Destination port range field locked/greyed out in some UI states. This appeared to block adding port 8080.")
body("Resolution: The Destination field showing 'All (1/65535)' already allows all ports including 8080. Submitting the rule with that value was correct — no change needed. The confusion arose from mistaking a display-only field for an editable one.")

heading3("4.3.3  SSH Key Permissions on Windows")
body("On Windows, the .pem file retains broad permissions by default. The SSH client rejects keys that are readable by other users:")
code_block("# Error:\n@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@\n@ WARNING: UNPROTECTED PRIVATE KEY FILE! @\n@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@@\n\n# Fix (Windows):\n# Right-click .pem file → Properties → Security tab\n# Click Advanced → Disable inheritance → Remove all inherited permissions\n# Add only your own user account with Read permission")
body("On Linux/macOS the fix is simply: chmod 400 <key>.pem")

page_break()

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — SUMMARY TABLE
# ══════════════════════════════════════════════════════════════════════════════

heading1("5.  Summary: Worked vs. Did Not Work")

add_table(
    ["Area", "Item", "Status", "Resolution"],
    [
        ["Deployment", "libGL.so.1 missing", "Failed", "apt-get install -y libgl1"],
        ["Deployment", "Both OpenCV variants installed", "Failed", "pip uninstall opencv-python (keep headless only)"],
        ["Deployment", "Gunicorn sync worker + SocketIO", "Failed", "Switch to worker_class='gthread', threads=4"],
        ["Deployment", "502 Bad Gateway from Nginx", "Failed", "Fix Gunicorn crash — 502 resolves automatically"],
        ["Deployment", "systemd shows 'active' then crashes", "Misleading", "Check Gunicorn error log, not systemctl status"],
        ["Deployment", "Alibaba Cloud security group UI", "Confusing", "All-ports rule already covers 8080 — submit as-is"],
        ["Deployment", "SSH key permissions on Windows", "Failed", "Remove all except owner Read via Security tab"],
        ["LLM/Agents", "Groq rejecting cache_breakpoint", "Failed", "Monkey-patch LiteLLM to strip the field"],
        ["LLM/Agents", "Single LLM call for complex shunt classification", "Insufficient", "Replaced with 5-agent CrewAI pipeline"],
        ["LLM/Agents", "LLM temperature 0.7 — too variable", "Failed", "Lowered to 0.3"],
        ["LLM/Agents", "LLM temperature 0.0 — frozen outputs", "Failed", "Raised to 0.3"],
        ["LLM/Agents", "VLM threshold too low — UI sluggish", "Poor UX", "Raised thresholds (0.05 / 0.06)"],
        ["Architecture", "Flask-SocketIO threading mode", "Worked", "—"],
        ["Architecture", "CrewAI 5-agent sequential crew", "Worked", "—"],
        ["Architecture", "Unix socket Gunicorn↔Nginx", "Worked", "—"],
        ["Architecture", "WebSocket Nginx proxy headers", "Worked", "—"],
        ["Architecture", "Two apps on same server (ports 80 + 8080)", "Worked", "—"],
        ["Architecture", "GitHub Actions SSH CI/CD", "Worked", "—"],
    ],
    col_widths=[1.4, 2.5, 1.2, 2.1],
    row_colors=[
        LIGHT_RED, LIGHT_RED, LIGHT_RED, LIGHT_RED, LIGHT_AMBER, LIGHT_AMBER, LIGHT_RED,
        LIGHT_RED, LIGHT_RED, LIGHT_RED, LIGHT_RED, LIGHT_AMBER,
        LIGHT_GREEN, LIGHT_GREEN, LIGHT_GREEN, LIGHT_GREEN, LIGHT_GREEN, LIGHT_GREEN,
    ],
)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — QUICK-REFERENCE COMMANDS
# ══════════════════════════════════════════════════════════════════════════════

heading1("6.  Quick-Reference Commands")
body("Commands you will run most often when managing the deployed service.")

heading2("Service Management")
add_table(
    ["Command", "Purpose"],
    [
        ["systemctl status task2", "Check if the service is running"],
        ["systemctl start task2", "Start the service"],
        ["systemctl stop task2", "Stop the service"],
        ["systemctl restart task2", "Restart after config/code changes"],
        ["systemctl enable task2", "Auto-start on server reboot"],
        ["journalctl -u task2 -n 50 --no-pager", "Last 50 systemd log lines"],
        ["cat /var/log/task2_probe_guidance/gunicorn-error.log", "Full Gunicorn error log (most useful)"],
        ["cat /var/log/task2_probe_guidance/gunicorn-access.log", "HTTP access log"],
    ],
    col_widths=[3.8, 3.4],
)

heading2("Deployment")
add_table(
    ["Command", "Purpose"],
    [
        ["cd /var/www/Cygnus_Med_Demo && git pull origin main", "Pull latest code"],
        ["bash Task_2_App/deploy/update.sh", "Full update: pull + pip install + restart"],
        ["bash Task_2_App/deploy/setup.sh", "First-time setup only (re-running is safe)"],
        ["nginx -t && systemctl reload nginx", "Validate and reload Nginx config"],
    ],
    col_widths=[3.8, 3.4],
)

heading2("Diagnostics")
add_table(
    ["Command", "Purpose"],
    [
        ["curl http://localhost:8080/api/status", "Verify app responds locally"],
        ["curl http://8.219.69.76:8080/api/status", "Verify app responds from internet"],
        ["ls -la /var/www/task2/", "Check if task2.sock exists (confirms Gunicorn started)"],
        ["ps aux | grep gunicorn", "List all Gunicorn processes"],
        ["pip list | grep opencv", "Confirm only headless OpenCV is installed"],
        ["python3 -c \"import cv2; print(cv2.__version__)\"", "Quick OpenCV import test"],
    ],
    col_widths=[3.8, 3.4],
)

# ══════════════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════════════

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
