"""
Socket.IO test runner for stream-mode CHIVA scenarios.

Usage:
    python tests/run_stream_scenario.py [--api http://localhost:7861] [--scenario <id>] [--all]

Scenarios:
    type1   — EP N1→N2 at SFJ + RP N2→N1 in trunk. Tests basic Type 1 complete.
    type2a  — EP N1→N2 at SPJ + RP N2→N1 in SSV. Tests posterior/SPJ navigation.
    type3   — SFJ EP + trunk RP + tributary escape + No Reflux elimination test.
    type1p2 — SFJ EP + trunk RP + tributary escape + Reflux elimination test.
    type4   — EP N1→N3 perforator + RP N2→N1 trunk. Tests direct perforator path.
    type6   — EP N1→N3 + RP N3→N1 only. Tests pure perforator circuit, no trunk.

Results are saved to:  tests/results/test_log_<timestamp>.docx

Requirements:
    pip install python-socketio[client] python-docx
"""
from __future__ import annotations

import argparse
import io
import json
import os
import sys
import time

# Force UTF-8 output on Windows (avoids cp1252 UnicodeEncodeError for box-drawing chars)
if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

try:
    import socketio
except ImportError:
    print("ERROR: python-socketio not installed. Run: pip install python-socketio[client]")
    sys.exit(1)


# ─────────────────────────────────────────────────────────────────────────────
# Step definition
# ─────────────────────────────────────────────────────────────────────────────

@dataclass
class Step:
    label: str
    event: str
    data: dict
    expected_action: Optional[str] = None        # "move" | "maneuver" | "complete" | None (skip)
    guidance_must_contain: list[str] = field(default_factory=list)  # ANY of these (case-insensitive)
    forbidden_action: Optional[str] = None        # this action must NOT appear


def _pm(sid, region, pos_y, surface, leg="right"):
    """Shorthand for a probe_move dict."""
    return {"session_id": sid, "region": region, "pos_y_ratio": pos_y,
            "surface": surface, "leg": leg}

def _cm(sid, flow, ft, tt, pos_y, region, surface, elim="", leg="right"):
    """Shorthand for a clip_mark dict."""
    return {"session_id": sid, "flow": flow, "from_type": ft, "to_type": tt,
            "pos_y_ratio": pos_y, "leg": leg, "region": region,
            "surface": surface, "elimination_test": elim}


# ─────────────────────────────────────────────────────────────────────────────
# Scenario 1 — Type 1
# EP N1→N2 at SFJ + RP N2→N1 in GSV trunk.  No EP N2→N3.
# Tests: basic navigation, Type 1 complete recognition.
# ─────────────────────────────────────────────────────────────────────────────

def _type1_scenario(sid="test_type1") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P1 probe at SFJ — no clips yet",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction", "saphenofemoral"]),

        Step("clip_mark EP N1→N2 at SFJ",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.06, "SFJ", "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["distal", "thigh", "trunk", "GSV", "medial"]),

        Step("P2 probe upper thigh — 1 EP clip",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.18, "medial"),
             expected_action="move",
             guidance_must_contain=["trunk", "reflux", "squeeze", "retrograde", "saphenous", "thigh"]),

        Step("clip_mark RP N2→N1 upper thigh",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.18, "UPPER_THIGH", "medial"),
             expected_action="move",
             guidance_must_contain=["Hunterian", "distal", "mid-thigh", "tributary", "escape", "perforator"]),

        Step("P3 probe Hunterian — EP+RP present",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="move",
             guidance_must_contain=["Hunterian", "tributary", "escape", "perforator", "mid-thigh"]),

        Step("P4 probe popliteal — should check SPJ competence",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "SSV", "posterior", "junction"]),

        Step("P5 mid-thigh after EP+RP — must NOT fire maneuver (no EP N2→N3)",
             "probe_move", _pm(sid, "MID_THIGH", 0.36, "medial"),
             forbidden_action="maneuver"),

        Step("P6 KEY — EP+RP confirmed, no N2→N3 → must fire complete",
             "probe_move", _pm(sid, "MID_THIGH", 0.25, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Scenario 2 — Type 2A
# EP N1→N2 at SPJ (popliteal, posterior) + RP N2→N1 in SSV (calf, posterior).
# SFJ is competent. Tests: posterior surface navigation, SPJ recognition.
# ─────────────────────────────────────────────────────────────────────────────

def _type2a_scenario(sid="test_type2a") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P1 probe at SFJ — competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P2 probe at popliteal — Q1 still open, should assess SPJ",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior", "SSV", "junction"]),

        Step("clip_mark EP N1→N2 at SPJ (posY=0.47 posterior)",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.47, "POPLITEAL", "posterior"),
             expected_action="move",
             guidance_must_contain=["distal", "SSV", "posterior", "calf", "trunk"]),

        Step("P3 probe calf posterior — should trace SSV for RP",
             "probe_move", _pm(sid, "CALF", 0.62, "posterior"),
             expected_action="move",
             guidance_must_contain=["SSV", "posterior", "calf", "reflux", "trunk"]),

        Step("clip_mark RP N2→N1 SSV calf (posY=0.62 posterior)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.62, "CALF", "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P4 KEY — EP at SPJ + RP in SSV → complete",
             "probe_move", _pm(sid, "CALF", 0.62, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Scenario 3 — Type 3  (elimination test → No Reflux)
# EP N1→N2 SFJ + RP N2→N1 trunk + EP N2→N3 tributary escape + RP N3→N1 re-entry
# → maneuver fires → surgeon records "No Reflux" → complete (Type 3).
# Tests: maneuver trigger, No Reflux path.
# ─────────────────────────────────────────────────────────────────────────────

def _type3_scenario(sid="test_type3") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P1 probe at SFJ", "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("clip_mark EP N1→N2 SFJ",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.06, "SFJ", "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["distal", "thigh", "trunk", "GSV"]),

        Step("clip_mark RP N2→N1 upper thigh",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.18, "UPPER_THIGH", "medial"),
             expected_action="move",
             guidance_must_contain=["Hunterian", "distal", "tributary", "escape"]),

        Step("P2 probe Hunterian — looking for EP N2→N3",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="move",
             guidance_must_contain=["tributary", "escape", "Hunterian", "perforator", "mid-thigh"]),

        Step("clip_mark EP N2→N3 Hunterian (no elimTest yet)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.28, "HUNTERIAN", "medial"),
             # NOT maneuver yet — need RP N3→N1 before maneuver fires
             expected_action="move",
             guidance_must_contain=["tributary", "distal", "calf", "re-entry", "follow"],
             forbidden_action="maneuver"),

        Step("P3 probe calf — follow tributary for re-entry",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="move",
             guidance_must_contain=["tributary", "distal", "calf", "perforator", "re-entry"]),

        Step("clip_mark RP N3→N1 calf (re-entry found)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.65, "CALF", "medial"),
             # NOW: EP N2→N3 + RP N3→N1 + RP N2→N1 present, no elimTest → MANEUVER
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary", "Doppler", "GSV", "record"]),

        Step("P4 KEY — probe after all 3 clips + no elimTest → must fire maneuver",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary", "record", "Doppler"]),

        Step("clip_mark EP N2→N3 with elimTest=No Reflux (Type 3 confirmed)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.28, "HUNTERIAN", "medial", elim="No Reflux"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P5 KEY — No Reflux recorded → must fire complete",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Scenario 4 — Type 1+2  (elimination test → Reflux)
# Identical path to Type 3 but surgeon records "Reflux" instead of "No Reflux".
# Tests: Reflux path after maneuver, Type 1+2 complete recognition.
# ─────────────────────────────────────────────────────────────────────────────

def _type1p2_scenario(sid="test_type1p2") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("clip_mark EP N1→N2 SFJ",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.06, "SFJ", "anterior-medial")),

        Step("clip_mark RP N2→N1 upper thigh",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.18, "UPPER_THIGH", "medial")),

        Step("clip_mark EP N2→N3 Hunterian",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.28, "HUNTERIAN", "medial"),
             forbidden_action="maneuver"),

        Step("clip_mark RP N3→N1 calf — maneuver must fire",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.65, "CALF", "medial"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary", "Doppler", "record"]),

        Step("P1 KEY — probe after all 3 clips + no elimTest → maneuver",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary", "record"]),

        Step("clip_mark EP N2→N3 with elimTest=Reflux (Type 1+2 confirmed)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.28, "HUNTERIAN", "medial", elim="Reflux"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P2 KEY — Reflux recorded → complete",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Scenario 5 — Type 4
# EP N1→N3 (perforator direct to tributary) + RP N2→N1 (GSV trunk retrograde).
# No SFJ entry. Tests: perforator-direct pathway.
# ─────────────────────────────────────────────────────────────────────────────

def _type4_scenario(sid="test_type4") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # SFJ visited but competent — no clip
        Step("P1 probe at SFJ — competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral"]),

        # SPJ visited — competent
        Step("P2 probe at popliteal — SPJ competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        # Find EP N1→N3 at calf perforator
        Step("P3 probe calf — looking for perforators",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="move",
             guidance_must_contain=["perforator", "calf", "tributary"]),

        Step("clip_mark EP N1→N3 calf perforator",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.65, "CALF", "medial"),
             expected_action="move",
             guidance_must_contain=["tributary", "trunk", "GSV", "trace", "distal"]),

        # Find RP N2→N1 in GSV trunk (retrograde trunk secondary to perforator circuit)
        Step("clip_mark RP N2→N1 GSV trunk",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.30, "HUNTERIAN", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P4 KEY — EP N1→N3 + RP N2→N1 → complete",
             "probe_move", _pm(sid, "HUNTERIAN", 0.30, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Scenario 6 — Type 6
# Pure perforator circuit: EP N1→N3 + RP N3→N1. No trunk involved.
# No RP N2→N1 and no RP N3→N2. Tests: pure N3 circuit complete recognition.
# ─────────────────────────────────────────────────────────────────────────────

def _type6_scenario(sid="test_type6") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # SFJ and SPJ both competent — no clips
        Step("P1 probe at SFJ — competent",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral"]),

        Step("P2 probe at popliteal — SPJ competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        # EP N1→N3 at calf perforator
        Step("clip_mark EP N1→N3 calf perforator (posY=0.68)",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.68, "CALF", "lateral"),
             expected_action="move",
             guidance_must_contain=["perforator", "tributary", "N3", "distal", "re-entry"]),

        # RP N3→N1 — tributary re-enters deep at another perforator distally
        Step("clip_mark RP N3→N1 distal perforator (posY=0.78)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.78, "CALF", "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P3 KEY — EP N1→N3 + RP N3→N1 only → complete (no trunk needed)",
             "probe_move", _pm(sid, "CALF", 0.78, "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),
    ]


# ─────────────────────────────────────────────────────────────────────────────
# Registry
# ─────────────────────────────────────────────────────────────────────────────

_SCENARIOS: dict[str, callable] = {
    "type1":   _type1_scenario,
    "type2a":  _type2a_scenario,
    "type3":   _type3_scenario,
    "type1p2": _type1p2_scenario,
    "type4":   _type4_scenario,
    "type6":   _type6_scenario,
}


# ─────────────────────────────────────────────────────────────────────────────
# Human-readable movement description
# ─────────────────────────────────────────────────────────────────────────────

def _describe_movement(event: str, data: dict) -> str:
    """Turn a raw socket event + data into plain English surgeon action text."""
    if event == "stream_start":
        return f"Start session  (id: {data.get('session_id', '?')})"

    if event == "probe_move":
        region  = data.get("region", "?")
        pos_y   = data.get("pos_y_ratio", 0.0)
        surface = data.get("surface", "?")
        leg     = data.get("leg", "right")
        return (
            f"Move probe to {region}\n"
            f"posY = {pos_y:.2f}  |  {surface} surface  |  {leg} leg"
        )

    if event == "clip_mark":
        flow   = data.get("flow", "?")
        ft     = data.get("from_type", "?")
        tt     = data.get("to_type",   "?")
        pos_y  = data.get("pos_y_ratio", 0.0)
        region = data.get("region", "?")
        leg    = data.get("leg", "right")
        elim   = data.get("elimination_test", "")
        base   = (
            f"Mark clip: {flow} {ft}→{tt}\n"
            f"posY = {pos_y:.2f}  |  {region}  |  {leg} leg"
        )
        if elim:
            base += f"\nElimination test: {elim}"
        return base

    return event


def _describe_expectation(step: "Step") -> str:
    """Format what the test expects from the AI for this step."""
    if step.event == "stream_start":
        return "session_ready event"
    parts = []
    if step.expected_action:
        parts.append(f"Action = \"{step.expected_action}\"")
    if step.forbidden_action:
        parts.append(f"Must NOT fire action = \"{step.forbidden_action}\"")
    if step.guidance_must_contain:
        parts.append("Guidance must mention any of:\n" +
                     ",  ".join(step.guidance_must_contain))
    return "\n".join(parts) if parts else "No assertion (observe only)"


# ─────────────────────────────────────────────────────────────────────────────
# Runner
# ─────────────────────────────────────────────────────────────────────────────

CYAN   = "\033[96m"
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

_SKIP_EVENTS = {"stream_start"}


class ScenarioRunner:
    def __init__(self, api_base: str):
        self.api_base = api_base.rstrip("/")
        self._sio = socketio.SimpleClient()

    def connect(self) -> None:
        print(f"Connecting to {self.api_base} …")
        self._sio.connect(self.api_base, wait_timeout=10)
        print("Connected.\n")

    def disconnect(self) -> None:
        self._sio.disconnect()

    def _emit_and_wait(self, event: str, data: dict, timeout: float = 60.0):
        self._sio.emit(event, data)
        if event in _SKIP_EVENTS:
            deadline = time.time() + timeout
            while time.time() < deadline:
                try:
                    ev = self._sio.receive(timeout=2)
                    if ev and ev[0] == "session_ready":
                        return ev[1]
                except Exception:
                    pass
            return None
        deadline = time.time() + timeout
        while time.time() < deadline:
            try:
                ev = self._sio.receive(timeout=2)
                if ev and ev[0] == "guidance_update":
                    return ev[1]
            except Exception:
                pass
        return None

    def run(self, steps: list[Step]) -> list[dict]:
        results = []
        for step in steps:
            print(f"\n{CYAN}─── {step.label} ───{RESET}")
            print(f"  event: {step.event}")

            movement    = _describe_movement(step.event, step.data)
            expectation = _describe_expectation(step)

            response = self._emit_and_wait(step.event, step.data)

            # ── stream_start ──
            if step.event in _SKIP_EVENTS:
                if response:
                    print(f"  session_ready: {response}")
                    results.append({
                        "step": step.label, "status": "SKIP",
                        "movement": movement, "expectation": expectation,
                        "llm_action": "—", "llm_guidance": "—", "failures": [],
                    })
                else:
                    print(f"  {YELLOW}⚠ No session_ready received{RESET}")
                    results.append({
                        "step": step.label, "status": "TIMEOUT",
                        "movement": movement, "expectation": expectation,
                        "llm_action": "—", "llm_guidance": "—", "failures": [],
                    })
                continue

            # ── observe-only steps (no assertions) ──
            if (step.expected_action is None
                    and not step.guidance_must_contain
                    and step.forbidden_action is None):
                if response:
                    llm_g = response.get("guidance", "")
                    llm_a = response.get("action", "?")
                    print(f"  guidance: {llm_g!r}  action: {llm_a}")
                    results.append({
                        "step": step.label, "status": "SKIP",
                        "movement": movement, "expectation": expectation,
                        "llm_action": llm_a, "llm_guidance": llm_g, "failures": [],
                    })
                else:
                    print(f"  {YELLOW}⚠ timeout{RESET}")
                    results.append({
                        "step": step.label, "status": "TIMEOUT",
                        "movement": movement, "expectation": expectation,
                        "llm_action": "—", "llm_guidance": "—", "failures": [],
                    })
                continue

            # ── timeout ──
            if response is None:
                print(f"  {YELLOW}⚠ No response received (timeout){RESET}")
                results.append({
                    "step": step.label, "status": "TIMEOUT",
                    "movement": movement, "expectation": expectation,
                    "llm_action": "—", "llm_guidance": "—", "failures": [],
                })
                continue

            llm_guidance = response.get("guidance", "")
            llm_action   = response.get("action", "move")
            print(f"  guidance: {llm_guidance!r}")
            print(f"  action  : {llm_action}")

            passed   = True
            failures = []

            if step.expected_action and llm_action != step.expected_action:
                passed = False
                failures.append(
                    f"Action: got \"{llm_action}\", expected \"{step.expected_action}\""
                )

            if step.guidance_must_contain:
                if not any(kw.lower() in llm_guidance.lower()
                           for kw in step.guidance_must_contain):
                    passed = False
                    failures.append(
                        f"Guidance missing any of: {', '.join(step.guidance_must_contain)}"
                    )

            if step.forbidden_action and llm_action == step.forbidden_action:
                passed = False
                failures.append(
                    f"Action \"{llm_action}\" must NOT fire at this step"
                )

            if passed:
                print(f"  {GREEN}{BOLD}PASS{RESET}")
            else:
                for f in failures:
                    print(f"  {RED}FAIL: {f}{RESET}")

            results.append({
                "step":        step.label,
                "status":      "PASS" if passed else "FAIL",
                "movement":    movement,
                "expectation": expectation,
                "llm_action":  llm_action,
                "llm_guidance": llm_guidance,
                "failures":    failures,
            })
            time.sleep(0.4)

        return results

    def print_summary(self, label: str, results: list[dict]) -> bool:
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        passed  = [r for r in checked if r["status"] == "PASS"]
        failed  = [r for r in checked if r["status"] == "FAIL"]
        timeout = [r for r in results if r["status"] == "TIMEOUT"]

        print(f"\n{'─'*55}")
        print(f"{BOLD}{label}{RESET}  "
              f"{GREEN}{len(passed)} passed{RESET}  "
              f"{RED}{len(failed)} failed{RESET}  "
              f"{YELLOW}{len(timeout)} timeout{RESET}  "
              f"/ {len(checked)} checked")

        for r in failed:
            print(f"  {RED}✗ {r['step']}{RESET}")
            for f in r.get("failures", []):
                print(f"      {f}")

        return len(failed) == 0 and len(timeout) == 0


# ─────────────────────────────────────────────────────────────────────────────
# Word document writer
# ─────────────────────────────────────────────────────────────────────────────

_SCENARIO_DESCRIPTIONS = {
    "type1":   "Type 1 — EP N1→N2 at SFJ + RP N2→N1 in GSV trunk. No tributary escape.",
    "type2a":  "Type 2A — EP N1→N2 at SPJ (posterior) + RP N2→N1 in SSV. SFJ competent.",
    "type3":   "Type 3 — SFJ EP + trunk RP + tributary escape. Elimination test: No Reflux.",
    "type1p2": "Type 1+2 — SFJ EP + trunk RP + tributary escape. Elimination test: Reflux.",
    "type4":   "Type 4 — EP N1→N3 perforator direct + RP N2→N1 trunk. No SFJ entry.",
    "type6":   "Type 6 — EP N1→N3 + RP N3→N1 only. Pure perforator circuit, no trunk.",
}

_STATUS_COLORS = {
    "PASS":    "00B050",
    "FAIL":    "FF0000",
    "TIMEOUT": "FF8C00",
    "SKIP":    "808080",
}


def _colored_cell(cell, text: str, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    from docx.shared import RGBColor
    cell.text = text
    run = (cell.paragraphs[0].runs[0]
           if cell.paragraphs[0].runs
           else cell.paragraphs[0].add_run(text))
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, pt: int = 9, bold: bool = False,
                   italic: bool = False, color_rgb=None) -> None:
    """Clear cell and write multi-line text with optional formatting."""
    cell.text = ""
    lines = text.split("\n")
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run  = para.add_run(line)
        run.font.size   = __import__("docx").shared.Pt(pt)
        run.font.bold   = bold
        run.font.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb


def write_word_report(
    all_scenario_results: dict[str, list[dict]],
    api_base: str,
    out_path: str,
) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # ── wider page margins so table fits ──
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    section = doc.sections[0]
    section.left_margin   = Inches(0.6)
    section.right_margin  = Inches(0.6)
    section.top_margin    = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    # ── Title ──
    title = doc.add_heading("CHIVA Streaming Guidance — Test Run Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Date: {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}    |    "
        f"API: {api_base}    |    "
        f"Scenarios: {len(all_scenario_results)}"
    )
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # ── Overall summary ──
    doc.add_heading("Overall Summary", level=1)
    stbl = doc.add_table(rows=1, cols=5)
    stbl.style = "Table Grid"
    for i, h in enumerate(["Scenario", "Description", "Passed", "Failed", "Timeouts"]):
        c = stbl.rows[0].cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True

    overall_pass = True
    for name, results in all_scenario_results.items():
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")
        row = stbl.add_row().cells
        row[0].text = name
        row[1].text = _SCENARIO_DESCRIPTIONS.get(name, "")
        row[2].text = str(p)
        row[3].text = str(f)
        row[4].text = str(t)
        if f or t:
            overall_pass = False

    doc.add_paragraph()
    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run("OVERALL RESULT: " + ("ALL PASS ✓" if overall_pass else "FAILURES PRESENT ✗"))
    vr.font.bold = True
    vr.font.size = Pt(13)
    vr.font.color.rgb = (RGBColor(0x00, 0xB0, 0x50) if overall_pass else RGBColor(0xFF, 0x00, 0x00))

    # ── Per-scenario detail pages ──
    for name, results in all_scenario_results.items():
        doc.add_page_break()
        doc.add_heading(f"Scenario: {name.upper()}", level=1)
        dp = doc.add_paragraph(_SCENARIO_DESCRIPTIONS.get(name, ""))
        if dp.runs:
            dp.runs[0].font.italic = True

        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")

        sl = doc.add_paragraph()
        sl.add_run(f"Steps checked: {len(checked)}   |   ")
        pr = sl.add_run(f"Passed: {p}")
        pr.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
        sl.add_run("   |   ")
        fr = sl.add_run(f"Failed: {f}")
        fr.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        sl.add_run("   |   ")
        tr = sl.add_run(f"Timeouts: {t}")
        tr.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        doc.add_paragraph()

        # ── 6-column detail table ──
        # Columns: # | What Surgeon Did | Expected Guidance | LLM Guidance | LLM Action | Status
        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        headers = ["#", "What Surgeon Did", "Expected Guidance", "LLM Guidance", "LLM Action", "Status"]
        col_w   = [0.25, 1.85, 2.10, 2.10, 0.85, 0.65]

        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            c.text = h
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.bold = True

        for i, w in enumerate(col_w):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)

        for idx, r in enumerate(results, 1):
            status = r.get("status", "SKIP")
            row    = tbl.add_row().cells

            # Col 0 — step number
            row[0].text = str(idx)

            # Col 1 — What Surgeon Did (movement)
            _set_cell_text(row[1], r.get("movement", "—"), pt=8)

            # Col 2 — Expected Guidance
            _set_cell_text(row[2], r.get("expectation", "—"), pt=8, italic=True)

            # Col 3 — LLM Guidance (actual output)
            llm_g = r.get("llm_guidance", "—")
            _set_cell_text(row[3], llm_g if llm_g else "—", pt=8)

            # Col 4 — LLM Action
            llm_a = r.get("llm_action", "—")
            _set_cell_text(row[4], llm_a if llm_a else "—", pt=9, bold=True)

            # Col 5 — Status (colored background)
            _colored_cell(row[5], status, _STATUS_COLORS.get(status, "808080"))

            # Failure note row (merged across all columns)
            if status == "FAIL" and r.get("failures"):
                fr_row  = tbl.add_row().cells
                merged  = fr_row[0].merge(fr_row[5])
                merged.text = ""
                for fn in r["failures"]:
                    pp = merged.add_paragraph(f"  ✗  {fn}")
                    if pp.runs:
                        pp.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        pp.runs[0].font.size      = Pt(8)
                        pp.runs[0].font.italic    = True

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"\nReport saved: {out_path}")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="CHIVA stream scenario test runner")
    parser.add_argument("scenario", nargs="?", default="type1",
                        choices=list(_SCENARIOS) + ["all"],
                        help="Scenario to run (default: type1)")
    parser.add_argument("--api",  default="http://localhost:7861")
    parser.add_argument("--all",  action="store_true", dest="run_all")
    args = parser.parse_args()

    to_run = list(_SCENARIOS.items()) if (args.run_all or args.scenario == "all") \
             else [(args.scenario, _SCENARIOS[args.scenario])]

    runner = ScenarioRunner(args.api)
    all_ok              = True
    all_scenario_results: dict[str, list[dict]] = {}

    try:
        runner.connect()
        for name, builder in to_run:
            print(f"\n{'='*55}")
            print(f"{BOLD}SCENARIO: {name.upper()}{RESET}")
            print(f"{'='*55}")
            steps   = builder()
            results = runner.run(steps)
            ok      = runner.print_summary(name, results)
            all_ok  = all_ok and ok
            all_scenario_results[name] = results
    except KeyboardInterrupt:
        print("\nInterrupted.")
    finally:
        runner.disconnect()

    if len(to_run) > 1:
        print(f"\n{'='*55}")
        print(f"{BOLD}OVERALL: {'ALL PASS' if all_ok else 'FAILURES PRESENT'}{RESET}")

    # Write Word report
    if all_scenario_results:
        ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
        here     = os.path.dirname(os.path.abspath(__file__))
        out_path = os.path.join(here, "results", f"test_log_{ts}.docx")
        try:
            write_word_report(all_scenario_results, args.api, out_path)
        except Exception as exc:
            print(f"{YELLOW}Warning: could not write Word report — {exc}{RESET}")

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
