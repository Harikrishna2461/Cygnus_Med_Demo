"""
test_set_4 — Six factually correct CHIVA shunt-type scenarios (Mendoza classification).

Each scenario represents one canonical shunt type with anatomically accurate
EP/RP clip patterns derived from the Mendoza duplex ultrasound classification.

Mendoza CHIVA Shunt Types
--------------------------
Type I   (R1→R2→R1):  Deep system (N1) enters saphenous trunk (N2) at SFJ/Hunterian.
                       Trunk refluxes distally. Blood re-enters deep via perforator
                       directly FROM THE TRUNK. Clips: EP N1→N2 + RP N2→N1.

Type II  (R2→R3→R1):  Both junctions (SFJ, SPJ) competent. Saphenous trunk is the
                       source of reflux — no deep junction entry. Trunk overflows into
                       tributary (EP N2→N3) and re-enters deep via tributary perforator
                       (RP N3→N1). No EP N1→N2; no RP N2→N1.

Type III (R1→R2→R3→R1): Deep entry at SFJ (EP N1→N2). ALL reflux volume is conducted
                       through the trunk to an escape tributary (EP N2→N3). Tributary
                       re-enters deep (RP N3→N1). Trunk appears to reflux (RP N2→N1)
                       but it is conducted flow — abolished on compression (No Reflux).

Type IV  (R1→R3→R2→R1): Pelvic vein or perforator exits directly into tributary
                       (EP N1→N3) — NOT at SFJ or SPJ. Tributary fills trunk
                       retrogradely. Trunk re-enters deep via perforator (RP N2→N1).
                       SFJ and SPJ both competent.

Type V   (R1→R3→R2→R3→R1): Source perforator fills tributary (EP N1→N3). Tributary
                       drains into saphenous trunk (RP N3→N2). Trunk escapes into a
                       second tributary (EP N2→N3). That tributary re-enters deep
                       (RP N3→N1). Four-clip perforator-mediated circuit.

Type VI  (R1→R3→R1):  Perforator exits deep into tributary (EP N1→N3). Blood travels
                       through tributary ONLY — saphenous trunk is NOT involved. Second
                       perforator re-enters deep (RP N3→N1). No RP N2→N1, no RP N3→N2.

Usage
-----
    python tests/run_stream_scenario_4.py [--api http://localhost:7861] [--scenario <id>] [--all]
"""
from __future__ import annotations

import argparse
import io
import os
import sys
import time

if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

try:
    import socketio
except ImportError:
    print("ERROR: python-socketio not installed.  Run: pip install python-socketio[client]")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Step definition
# ---------------------------------------------------------------------------

@dataclass
class Step:
    label: str
    event: str
    data: dict
    expected_action: Optional[str] = None
    guidance_must_contain: list[str] = field(default_factory=list)
    forbidden_action: Optional[str] = None


def _pm(sid, region, pos_y, surface, leg="right"):
    return {"session_id": sid, "region": region, "pos_y_ratio": pos_y,
            "surface": surface, "leg": leg}

def _cm(sid, flow, ft, tt, pos_y, region, surface, elim="", leg="right"):
    return {"session_id": sid, "flow": flow, "from_type": ft, "to_type": tt,
            "pos_y_ratio": pos_y, "leg": leg, "region": region,
            "surface": surface, "elimination_test": elim}


# ---------------------------------------------------------------------------
# Scenario 1 — mendoza_type1  (Type I: R1→R2→R1)  22 checked steps
# ---------------------------------------------------------------------------
# Anatomy (Mendoza):
#   Deep system (N1) enters GSV trunk (N2) at SFJ — incompetent terminal SFJ valve.
#   Trunk refluxes distally (retrograde flow in N2).
#   Blood re-enters deep via a PERFORATING VEIN directly from the saphenous trunk.
#   No tributary (N3) involvement at all.
#
# Clip pattern: EP N1→N2 (SFJ) + RP N2→N1 (calf perforator from trunk)
# Engine rule: Rule 6 — ep_n1_n2 + rp_n2_n1 + NOT ep_n2_n3 + max_visited ≥ 0.44

def _mendoza_type1_scenario(sid="ts4_type1") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # SFJ: incompetent — both Valsalva AND squeeze positive
        Step("P01 SFJ approach — assess terminal valve",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 SFJ deeper transverse view",
             "probe_move", _pm(sid, "SFJ", 0.09, "anterior-medial"),
             expected_action="move"),

        # EP N1→N2: SFJ incompetent — deep blood enters GSV trunk here
        Step("CM-1  EP N1→N2 at SFJ — SFJ incompetent, trunk entry confirmed",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.07, "SFJ", "anterior-medial"),
             expected_action="move"),

        # Sweep GSV trunk distally looking for tributary escape
        Step("P03 upper thigh — tracing GSV trunk distally",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.13, "medial"),
             expected_action="move"),

        Step("P04 upper thigh mid — no escape tributary",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.19, "medial"),
             expected_action="move"),

        Step("P05 Hunterian proximal — scan for escape perforator",
             "probe_move", _pm(sid, "HUNTERIAN", 0.26, "medial"),
             expected_action="move"),

        Step("P06 Hunterian mid — no EP N2→N3 found",
             "probe_move", _pm(sid, "HUNTERIAN", 0.32, "medial"),
             expected_action="move"),

        Step("P07 Hunterian distal — continue toward popliteal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.38, "medial"),
             expected_action="move"),

        # Popliteal — SPJ competent (one positive only → terminal competent)
        Step("P08 popliteal — confirm SPJ competence",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        # max_visited now 0.47 ≥ 0.44 — but rp_n2_n1 still absent → no complete yet
        Step("P09 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.53, "posterior"),
             expected_action="move"),

        # Calf — medial GSV, no tributaries filling
        Step("P10 calf upper medial — GSV trunk clean",
             "probe_move", _pm(sid, "CALF", 0.60, "medial"),
             expected_action="move"),

        Step("P11 calf mid medial",
             "probe_move", _pm(sid, "CALF", 0.66, "medial"),
             expected_action="move"),

        Step("P12 calf lower medial",
             "probe_move", _pm(sid, "CALF", 0.72, "medial"),
             expected_action="move"),

        Step("P13 ankle — distal GSV",
             "probe_move", _pm(sid, "ANKLE", 0.85, "medial"),
             expected_action="move"),

        # Return sweep — confirm no tributary escape at any level
        Step("P14 calf return lower",
             "probe_move", _pm(sid, "CALF", 0.79, "medial"),
             expected_action="move"),

        Step("P15 calf return mid",
             "probe_move", _pm(sid, "CALF", 0.73, "medial"),
             expected_action="move"),

        Step("P16 calf return upper",
             "probe_move", _pm(sid, "CALF", 0.67, "medial"),
             expected_action="move"),

        Step("P17 Hunterian return",
             "probe_move", _pm(sid, "HUNTERIAN", 0.30, "medial"),
             expected_action="move"),

        Step("P18 Hunterian proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.24, "medial"),
             expected_action="move"),

        # RP N2→N1: perforator from trunk directly back to deep.
        # This is the TYPE I re-entry — blood exits trunk into deep via perforator.
        # old_max_visited=0.85 ≥ 0.44; ep_n1_n2=T; ep_n2_n3=F → Rule 6 COMPLETE
        Step("CM-2  RP N2→N1 calf perforator — re-entry from trunk — COMPLETE (Type I)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.64, "CALF", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P19 verify complete — popliteal",
             "probe_move", _pm(sid, "POPLITEAL", 0.50, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P20 verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 2 — mendoza_type2  (Type II: R2→R3→R1)  21 checked steps
# ---------------------------------------------------------------------------
# Anatomy (Mendoza):
#   Both SFJ and SPJ terminal valves COMPETENT — no deep (N1) to trunk (N2) entry.
#   The saphenous trunk itself is the source of reflux (e.g. pre-terminal valve
#   incompetence, or Giacomini-mediated segment of local reflux in N2).
#   Trunk overflows into a tributary: EP N2→N3.
#   Tributary re-enters deep system: RP N3→N1.
#   No RP N2→N1 (trunk has no perforator back to deep — all exits via tributary).
#
# Clip pattern: EP N2→N3 (trunk→tributary) + RP N3→N1 (tributary→deep)
# Engine rule: Rule 7 — ep_n2_n3 + rp_n3_n1 + NOT ep_n1_n2 + NOT ep_n1_n3 + NOT rp_n2_n1

def _mendoza_type2_scenario(sid="ts4_type2") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # SFJ competent — terminal valve holds, no deep-to-trunk entry
        Step("P01 SFJ — terminal valve competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 upper thigh — tracing GSV trunk",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.12, "medial"),
             expected_action="move"),

        Step("P03 upper thigh mid — antegrade flow in trunk",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.18, "medial"),
             expected_action="move"),

        Step("P04 Hunterian proximal — scanning for reflux origin in trunk",
             "probe_move", _pm(sid, "HUNTERIAN", 0.24, "medial"),
             expected_action="move"),

        Step("P05 Hunterian mid — local trunk segment reflux detected",
             "probe_move", _pm(sid, "HUNTERIAN", 0.30, "medial"),
             expected_action="move"),

        Step("P06 Hunterian distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.36, "medial"),
             expected_action="move"),

        # Popliteal — SPJ competent (terminal valve holds)
        Step("P07 popliteal — SPJ terminal valve competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P08 calf posterior — SSV competent",
             "probe_move", _pm(sid, "CALF", 0.55, "posterior"),
             expected_action="move"),

        Step("P09 calf medial — trunk reflux with tributary overflow",
             "probe_move", _pm(sid, "CALF", 0.61, "medial"),
             expected_action="move"),

        # EP N2→N3: trunk overflows into tributary at Hunterian zone.
        # This is the TYPE II pattern — source is trunk (N2), not deep (N1).
        # No EP N1→N2 ever confirmed.
        Step("CM-1  EP N2→N3 at Hunterian — trunk→tributary escape (Type II source)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.29, "HUNTERIAN", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Follow tributary distally toward re-entry perforator
        Step("P10 distal thigh — following N3 tributary",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.40, "medial"),
             expected_action="move"),

        Step("P11 popliteal — tracking tributary into popliteal region",
             "probe_move", _pm(sid, "POPLITEAL", 0.49, "posterior"),
             expected_action="move"),

        Step("P12 calf posterior upper — tracing N3 distally",
             "probe_move", _pm(sid, "CALF", 0.56, "posterior"),
             expected_action="move"),

        Step("P13 calf posterior mid",
             "probe_move", _pm(sid, "CALF", 0.63, "posterior"),
             expected_action="move"),

        Step("P14 calf posterior lower",
             "probe_move", _pm(sid, "CALF", 0.69, "posterior"),
             expected_action="move"),

        Step("P15 calf lateral — scanning for re-entry perforator",
             "probe_move", _pm(sid, "CALF", 0.75, "lateral"),
             expected_action="move"),

        Step("P16 ankle lateral — tracking tributary to ankle",
             "probe_move", _pm(sid, "ANKLE", 0.85, "lateral"),
             expected_action="move"),

        Step("P17 calf lateral return — re-tracing for perforator",
             "probe_move", _pm(sid, "CALF", 0.78, "lateral"),
             expected_action="move"),

        Step("P18 calf lateral mid return",
             "probe_move", _pm(sid, "CALF", 0.71, "lateral"),
             expected_action="move"),

        # RP N3→N1: tributary re-enters deep system.
        # No RP N2→N1 found anywhere (trunk has no direct perforator re-entry).
        # Rule 7 fires: ep_n2_n3 + rp_n3_n1 + NOT ep_n1_n2 + NOT ep_n1_n3 + NOT rp_n2_n1
        Step("CM-2  RP N3→N1 at calf — tributary re-enters deep — COMPLETE (Type II)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.68, "CALF", "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P19 verify complete — Hunterian",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P20 verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 3 — mendoza_type3  (Type III: R1→R2→R3→R1)  23 checked steps
# ---------------------------------------------------------------------------
# Anatomy (Mendoza):
#   Deep system enters GSV trunk at SFJ (EP N1→N2) — identical to Type I start.
#   However, ALL reflux volume is conducted through the trunk to a tributary escape
#   (EP N2→N3). The trunk acts as a CONDUIT — not a final re-entry point.
#   Tributary re-enters deep (RP N3→N1). Trunk appears to reflux (RP N2→N1) but
#   this is conducted flow that disappears on tributary compression (No Reflux).
#
# IMPORTANT clinical distinction from Type I+2:
#   In Type III, RP N2→N1 is present but is conducted (not independent reflux).
#   Compression of the tributary (elimination test) abolishes RP N2→N1 → "No Reflux".
#   In Type I+2, RP N2→N1 persists because the trunk has INDEPENDENT deep reflux.
#
# Clip pattern: EP N1→N2 + EP N2→N3 (before RP N2→N1 to block Rule 6) +
#               RP N3→N1 + RP N2→N1 → MANEUVER → elim "No Reflux" → COMPLETE
# Engine rules: Rule 1 (maneuver) → Rule 2 (No Reflux → complete)

def _mendoza_type3_scenario(sid="ts4_type3") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01 SFJ approach — assess terminal valve",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        # SFJ incompetent: deep blood enters trunk (same as Type I start)
        Step("CM-1  EP N1→N2 at SFJ — deep enters trunk (Type III source)",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.07, "SFJ", "anterior-medial"),
             expected_action="move"),

        Step("P02 upper thigh — tracing GSV trunk distally",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.13, "medial"),
             expected_action="move"),

        Step("P03 upper thigh mid",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.20, "medial"),
             expected_action="move"),

        Step("P04 Hunterian proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.26, "medial"),
             expected_action="move"),

        # EP N2→N3 marked BEFORE RP N2→N1 — this is CRITICAL.
        # It establishes escape before we find trunk reflux, blocking Rule 6 (Type I).
        # In true Type III, ALL blood goes through this escape point.
        Step("CM-2  EP N2→N3 Hunterian — trunk conducts ALL volume to tributary",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.28, "HUNTERIAN", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Popliteal — SPJ competent
        Step("P05 popliteal — SPJ competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P06 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.53, "posterior"),
             expected_action="move"),

        # Follow tributary and trunk distally to find RP N3→N1 and RP N2→N1
        Step("P07 calf posterior — following N3 tributary distally",
             "probe_move", _pm(sid, "CALF", 0.60, "posterior"),
             expected_action="move"),

        Step("P08 calf medial — tracing N2 trunk alongside",
             "probe_move", _pm(sid, "CALF", 0.66, "medial"),
             expected_action="move"),

        Step("P09 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.72, "medial"),
             expected_action="move"),

        Step("P10 ankle",
             "probe_move", _pm(sid, "ANKLE", 0.85, "medial"),
             expected_action="move"),

        Step("P11 calf return lower",
             "probe_move", _pm(sid, "CALF", 0.79, "medial"),
             expected_action="move"),

        Step("P12 calf return mid",
             "probe_move", _pm(sid, "CALF", 0.73, "medial"),
             expected_action="move"),

        # RP N3→N1: tributary re-enters deep. Maneuver needs RP N2→N1 too → not yet.
        Step("CM-3  RP N3→N1 calf — tributary re-enters deep (conduit circuit partial)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.69, "CALF", "medial"),
             expected_action="move",
             forbidden_action="maneuver"),

        Step("P13 Hunterian return — sweep toward conducted RP N2→N1",
             "probe_move", _pm(sid, "HUNTERIAN", 0.32, "medial"),
             expected_action="move"),

        Step("P14 Hunterian proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.27, "medial"),
             expected_action="move"),

        # RP N2→N1: conducted trunk reflux present (NOT independent — will be abolished
        # on compression). With ep_n2_n3_no_elim + rp_n3_n1 + rp_n2_n1 → Rule 1 MANEUVER
        Step("CM-4  RP N2→N1 Hunterian — conducted trunk reflux — MANEUVER fires",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.25, "HUNTERIAN", "medial"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary"]),

        Step("P15 calf — maneuver state persists (awaiting elimination result)",
             "probe_move", _pm(sid, "CALF", 0.64, "medial"),
             expected_action="maneuver"),

        Step("P16 Hunterian — return to escape point for compression test",
             "probe_move", _pm(sid, "HUNTERIAN", 0.30, "medial"),
             expected_action="maneuver"),

        # Elimination test: compress tributary at Hunterian escape point.
        # RP N2→N1 disappears → "No Reflux" → Type III confirmed (trunk was conduit only).
        Step("CM-5  EP N2→N3 elim='No Reflux' — trunk was conduit — COMPLETE (Type III)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.28, "HUNTERIAN", "medial",
                               elim="No Reflux"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P17 verify complete — calf",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P18 verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 4 — mendoza_type4  (Type IV: R1→R3→R2→R1)  23 checked steps
# ---------------------------------------------------------------------------
# Anatomy (Mendoza):
#   Pelvic vein or perforator exits deep and enters a TRIBUTARY directly (EP N1→N3).
#   NOT at SFJ or SPJ — both junctions are competent.
#   The tributary fills the saphenous trunk (N2) retrogradely.
#   The trunk then re-enters deep via a proximal perforator (RP N2→N1).
#   CRITICAL DISTINCTION from Type I: entry is via N3 tributary, not at N1→N2 junction.
#
# Clip pattern: EP N1→N3 (perforator/pelvic into tributary) + RP N2→N1 (trunk→deep)
# Engine rule: Rule 4 — ep_n1_n3 + rp_n2_n1 (fires immediately, no max_visited gate)

def _mendoza_type4_scenario(sid="ts4_type4") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # Both junctions competent
        Step("P01 SFJ — terminal valve competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 popliteal — SPJ terminal valve competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        # Systematic distal sweep searching for tributary entry perforator
        Step("P03 calf posterior — searching calf perforators",
             "probe_move", _pm(sid, "CALF", 0.55, "posterior"),
             expected_action="move"),

        Step("P04 calf medial upper",
             "probe_move", _pm(sid, "CALF", 0.61, "medial"),
             expected_action="move"),

        Step("P05 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.67, "medial"),
             expected_action="move"),

        Step("P06 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.73, "medial"),
             expected_action="move"),

        Step("P07 ankle medial",
             "probe_move", _pm(sid, "ANKLE", 0.85, "medial"),
             expected_action="move"),

        Step("P08 ankle lateral",
             "probe_move", _pm(sid, "ANKLE", 0.91, "lateral"),
             expected_action="move"),

        Step("P09 ankle return medial",
             "probe_move", _pm(sid, "ANKLE", 0.85, "medial"),
             expected_action="move"),

        Step("P10 calf lower return",
             "probe_move", _pm(sid, "CALF", 0.78, "medial"),
             expected_action="move"),

        # EP N1→N3: perforator exits deep DIRECTLY into tributary.
        # This is the TYPE IV entry — NOT via SFJ/SPJ — SFJ is competent.
        # Tributary then fills GSV trunk retrogradely (no clip for this junction).
        Step("CM-1  EP N1→N3 at calf — pelvic/perf exits into tributary (Type IV entry)",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.71, "CALF", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Sweep trunk proximally to find where blood re-enters deep from trunk
        Step("P11 popliteal — sweep trunk proximally",
             "probe_move", _pm(sid, "POPLITEAL", 0.50, "posterior"),
             expected_action="move"),

        Step("P12 distal thigh",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.42, "medial"),
             expected_action="move"),

        Step("P13 Hunterian distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.36, "medial"),
             expected_action="move"),

        Step("P14 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.30, "medial"),
             expected_action="move"),

        Step("P15 Hunterian proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.24, "medial"),
             expected_action="move"),

        Step("P16 upper thigh",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.17, "medial"),
             expected_action="move"),

        Step("P17 upper thigh proximal",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.13, "medial"),
             expected_action="move"),

        Step("P18 Hunterian return sweep",
             "probe_move", _pm(sid, "HUNTERIAN", 0.20, "medial"),
             expected_action="move"),

        Step("P19 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.27, "medial"),
             expected_action="move"),

        # RP N2→N1: trunk blood re-enters deep via Hunterian perforator.
        # This completes R1→R3→R2→R1. Rule 4 fires IMMEDIATELY (no max_visited gate).
        Step("CM-2  RP N2→N1 Hunterian — trunk re-enters deep — COMPLETE (Type IV)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.25, "HUNTERIAN", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P20 verify complete — calf",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P21 verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 5 — mendoza_type5  (Type V: R1→R3→R2→R3→R1)  24 checked steps
# ---------------------------------------------------------------------------
# Anatomy (Mendoza):
#   Source perforator exits deep into tributary: EP N1→N3.
#   Tributary fills saphenous trunk (N2) via junction: RP N3→N2.
#   Trunk re-escapes into a SECOND tributary: EP N2→N3.
#   Second tributary re-enters deep: RP N3→N1.
#   No direct trunk-to-deep re-entry (no RP N2→N1) — all volume goes trunk→N3→deep.
#   Both junctions (SFJ, SPJ) are competent.
#
# CRITICAL distinction from Type IV: in Type IV, trunk re-enters deep directly
# (RP N2→N1). In Type V, trunk escapes via a SECOND tributary before re-entering.
#
# Clip pattern: EP N1→N3 + RP N3→N2 + EP N2→N3 + RP N3→N1
# Engine rule: Rule 5b — ep_n1_n3 + rp_n3_n2 + ep_n2_n3 + rp_n3_n1

def _mendoza_type5_scenario(sid="ts4_type5") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01 SFJ — competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 popliteal — SPJ competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P03 calf posterior — searching source perforator",
             "probe_move", _pm(sid, "CALF", 0.55, "posterior"),
             expected_action="move"),

        Step("P04 calf medial upper",
             "probe_move", _pm(sid, "CALF", 0.61, "medial"),
             expected_action="move"),

        Step("P05 calf medial mid — scanning for perforator",
             "probe_move", _pm(sid, "CALF", 0.67, "medial"),
             expected_action="move"),

        # EP N1→N3: perforator exits deep directly into tributary (Type V source).
        # Tributary then fills trunk (N2) via RP N3→N2 connection.
        Step("CM-1  EP N1→N3 at calf — source perforator exits deep into N3 tributary",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.64, "CALF", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Trace N3 tributary proximally to where it joins GSV trunk
        Step("P06 distal thigh — tracing N3 tributary toward trunk junction",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.39, "medial"),
             expected_action="move"),

        Step("P07 Hunterian distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.33, "medial"),
             expected_action="move"),

        Step("P08 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.27, "medial"),
             expected_action="move"),

        Step("P09 Hunterian proximal — scanning tributary-to-trunk junction",
             "probe_move", _pm(sid, "HUNTERIAN", 0.21, "medial"),
             expected_action="move"),

        Step("P10 Hunterian return",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="move"),

        # RP N3→N2: tributary drains INTO saphenous trunk at this junction.
        # This is where N3 feeds N2 — distinguishes Type V from Type VI (which has no trunk).
        # ep_n1_n3 + rp_n3_n2 now present, but ep_n2_n3 absent → Rule 5b not yet complete.
        Step("CM-2  RP N3→N2 Hunterian — N3 tributary drains into N2 trunk",
             "clip_mark", _cm(sid, "RP", "N3", "N2", 0.30, "HUNTERIAN", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Now follow trunk distally to find where it escapes into a second tributary
        Step("P11 calf upper — trunk fills from Hunterian, sweeping distally",
             "probe_move", _pm(sid, "CALF", 0.57, "medial"),
             expected_action="move"),

        Step("P12 calf medial",
             "probe_move", _pm(sid, "CALF", 0.63, "medial"),
             expected_action="move"),

        Step("P13 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.69, "medial"),
             expected_action="move"),

        Step("P14 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.75, "medial"),
             expected_action="move"),

        # EP N2→N3: trunk escapes into a SECOND tributary at distal calf.
        # ep_n1_n3 + rp_n3_n2 + ep_n2_n3 now present, but rp_n3_n1 absent → not complete yet.
        Step("CM-3  EP N2→N3 at calf — trunk escapes into second N3 tributary",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.72, "CALF", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Follow second tributary to its re-entry perforator
        Step("P15 calf lower — tracing second tributary to re-entry",
             "probe_move", _pm(sid, "CALF", 0.78, "medial"),
             expected_action="move"),

        Step("P16 ankle approach",
             "probe_move", _pm(sid, "ANKLE", 0.86, "medial"),
             expected_action="move"),

        Step("P17 ankle distal",
             "probe_move", _pm(sid, "ANKLE", 0.92, "medial"),
             expected_action="move"),

        Step("P18 ankle return",
             "probe_move", _pm(sid, "ANKLE", 0.86, "medial"),
             expected_action="move"),

        # RP N3→N1: second tributary re-enters deep system.
        # Rule 5b fires: ep_n1_n3 + rp_n3_n2 + ep_n2_n3 + rp_n3_n1 → Type V complete.
        # No RP N2→N1 → Rule 4 (Type IV) cannot fire. rp_n3_n2 present → Rule 5 (Type VI) blocked.
        Step("CM-4  RP N3→N1 ankle — second tributary re-enters deep — COMPLETE (Type V)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.89, "ANKLE", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P19 verify complete — calf",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P20 verify complete — Hunterian",
             "probe_move", _pm(sid, "HUNTERIAN", 0.29, "medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 6 — mendoza_type6  (Type VI: R1→R3→R1)  21 checked steps
# ---------------------------------------------------------------------------
# Anatomy (Mendoza):
#   Perforator exits deep and enters tributary: EP N1→N3.
#   Blood travels through tributary ONLY — saphenous trunk (N2) NOT involved.
#   Second perforator re-enters deep: RP N3→N1.
#   No RP N2→N1 (trunk not involved), no RP N3→N2 (no tributary-to-trunk connection).
#   Both SFJ and SPJ confirmed competent.
#
# CRITICAL distinction from Type V: Type VI has NO trunk involvement at all.
# If RP N3→N2 is found, circuit is Type V not VI.
#
# Clip pattern: EP N1→N3 + RP N3→N1 (NO RP N2→N1, NO RP N3→N2)
# Engine rule: Rule 5 — ep_n1_n3 + rp_n3_n1 + NOT rp_n2_n1 + NOT rp_n3_n2

def _mendoza_type6_scenario(sid="ts4_type6") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # Both junctions competent — confirms N2 trunk is not involved
        Step("P01 SFJ — terminal valve competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 popliteal — SPJ terminal valve competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        # Systematic search across calf and ankle for perforators
        Step("P03 calf posterior upper",
             "probe_move", _pm(sid, "CALF", 0.55, "posterior"),
             expected_action="move"),

        Step("P04 calf medial upper",
             "probe_move", _pm(sid, "CALF", 0.61, "medial"),
             expected_action="move"),

        Step("P05 calf lateral upper",
             "probe_move", _pm(sid, "CALF", 0.64, "lateral"),
             expected_action="move"),

        Step("P06 calf lateral mid",
             "probe_move", _pm(sid, "CALF", 0.70, "lateral"),
             expected_action="move"),

        Step("P07 calf lateral lower",
             "probe_move", _pm(sid, "CALF", 0.76, "lateral"),
             expected_action="move"),

        Step("P08 ankle lateral",
             "probe_move", _pm(sid, "ANKLE", 0.85, "lateral"),
             expected_action="move"),

        Step("P09 ankle lateral distal",
             "probe_move", _pm(sid, "ANKLE", 0.91, "lateral"),
             expected_action="move"),

        Step("P10 ankle return lateral",
             "probe_move", _pm(sid, "ANKLE", 0.85, "lateral"),
             expected_action="move"),

        # EP N1→N3: source perforator exits deep into tributary.
        # No SFJ/SPJ involvement — purely perforator-to-tributary entry.
        # ep_n1_n3=T, but rp_n3_n1=F → Rule 5 not yet complete.
        Step("CM-1  EP N1→N3 ankle lateral — source perforator exits deep into N3",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.87, "ANKLE", "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        # Trace N3 tributary proximally — looking for re-entry perforator
        # MUST NOT find RP N3→N2 (that would mean trunk involvement → Type V)
        Step("P11 calf lateral return — tracing N3 tributary proximally",
             "probe_move", _pm(sid, "CALF", 0.80, "lateral"),
             expected_action="move"),

        Step("P12 calf lateral",
             "probe_move", _pm(sid, "CALF", 0.74, "lateral"),
             expected_action="move"),

        Step("P13 calf lateral mid",
             "probe_move", _pm(sid, "CALF", 0.68, "lateral"),
             expected_action="move"),

        Step("P14 calf posterior — check trunk not involved",
             "probe_move", _pm(sid, "CALF", 0.62, "posterior"),
             expected_action="move"),

        Step("P15 calf medial — GSV trunk shows NO reflux (confirming Type VI)",
             "probe_move", _pm(sid, "CALF", 0.66, "medial"),
             expected_action="move"),

        Step("P16 calf lateral — returning to find re-entry",
             "probe_move", _pm(sid, "CALF", 0.73, "lateral"),
             expected_action="move"),

        Step("P17 calf lateral proximal",
             "probe_move", _pm(sid, "CALF", 0.67, "lateral"),
             expected_action="move"),

        # RP N3→N1: tributary re-enters deep via second perforator.
        # NO RP N2→N1 (trunk not involved). NO RP N3→N2 (no trunk connection).
        # Rule 5 fires: ep_n1_n3 + rp_n3_n1 + NOT rp_n2_n1 + NOT rp_n3_n2 → Type VI COMPLETE.
        Step("CM-2  RP N3→N1 calf lateral — tributary re-enters deep — COMPLETE (Type VI)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.71, "CALF", "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P18 verify complete — calf",
             "probe_move", _pm(sid, "CALF", 0.67, "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P19 verify complete — popliteal",
             "probe_move", _pm(sid, "POPLITEAL", 0.48, "posterior"),
             expected_action="complete"),

        Step("P20 verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

_SCENARIOS: dict[str, callable] = {
    "mendoza_type1": _mendoza_type1_scenario,
    "mendoza_type2": _mendoza_type2_scenario,
    "mendoza_type3": _mendoza_type3_scenario,
    "mendoza_type4": _mendoza_type4_scenario,
    "mendoza_type5": _mendoza_type5_scenario,
    "mendoza_type6": _mendoza_type6_scenario,
}

_SCENARIO_DESCRIPTIONS = {
    "mendoza_type1": (
        "Type I (R1→R2→R1): SFJ incompetent entry (EP N1→N2); GSV trunk reflux; "
        "no tributary escape; re-entry from trunk directly via calf perforator (RP N2→N1). "
        "Rule 6 complete (max_visited ≥ 0.44)."
    ),
    "mendoza_type2": (
        "Type II (R2→R3→R1): Both junctions competent; saphenous trunk is reflux source; "
        "trunk overflows to tributary (EP N2→N3); tributary re-enters deep (RP N3→N1). "
        "No EP N1→N2; no RP N2→N1. Rule 7 complete."
    ),
    "mendoza_type3": (
        "Type III (R1→R2→R3→R1): SFJ entry (EP N1→N2); ALL volume conducted via tributary "
        "(EP N2→N3 before RP N2→N1); RP N3→N1 + RP N2→N1 → Maneuver; "
        "elim='No Reflux' → Type III complete (trunk was conduit)."
    ),
    "mendoza_type4": (
        "Type IV (R1→R3→R2→R1): Both junctions competent; pelvic/perf exits into tributary "
        "(EP N1→N3); tributary fills trunk; trunk re-enters deep (RP N2→N1). "
        "Rule 4 complete immediately (no max_visited gate)."
    ),
    "mendoza_type5": (
        "Type V (R1→R3→R2→R3→R1): Source perforator into tributary (EP N1→N3); "
        "tributary drains into trunk (RP N3→N2); trunk escapes to second tributary (EP N2→N3); "
        "second tributary re-enters deep (RP N3→N1). Rule 5b complete."
    ),
    "mendoza_type6": (
        "Type VI (R1→R3→R1): Both junctions competent; perforator exits into tributary "
        "(EP N1→N3); PURELY superficial N3 path; re-entry perforator (RP N3→N1). "
        "No trunk (N2) involvement. Rule 5 complete."
    ),
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _describe_movement(event: str, data: dict) -> str:
    if event == "stream_start":
        return f"Start session  (id: {data.get('session_id', '?')})"
    if event == "probe_move":
        return (f"Move probe → {data.get('region', '?')}\n"
                f"posY={data.get('pos_y_ratio', 0.0):.2f}  "
                f"{data.get('surface', '?')}  {data.get('leg', 'right')} leg")
    if event == "clip_mark":
        base = (f"Mark: {data.get('flow', '?')} "
                f"{data.get('from_type', '?')}→{data.get('to_type', '?')}  "
                f"posY={data.get('pos_y_ratio', 0.0):.2f}  {data.get('region', '?')}  "
                f"{data.get('leg', 'right')} leg")
        if data.get("elimination_test"):
            base += f"\nElim: {data['elimination_test']}"
        return base
    return event


def _describe_expectation(step: Step) -> str:
    if step.event == "stream_start":
        return "session_ready event"
    parts = []
    if step.expected_action:
        parts.append(f'Action = "{step.expected_action}"')
    if step.forbidden_action:
        parts.append(f'Must NOT = "{step.forbidden_action}"')
    if step.guidance_must_contain:
        parts.append("Guidance contains any of: " + ", ".join(step.guidance_must_contain))
    return "\n".join(parts) if parts else "observe only"


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

CYAN   = "\033[96m"
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

_SKIP_EVENTS = {"stream_start"}


class ScenarioRunner:
    def __init__(self, api_base: str):
        self.api_base = api_base.rstrip("/")
        self._sio = socketio.SimpleClient()

    def connect(self) -> None:
        print(f"Connecting to {self.api_base} ...")
        self._sio.connect(self.api_base, wait_timeout=10)
        print("Connected.\n")

    def disconnect(self) -> None:
        self._sio.disconnect()

    def _emit_and_wait(self, event: str, data: dict, timeout: float = 60.0):
        self._sio.emit(event, data)
        if event in _SKIP_EVENTS:
            deadline = time.time() + timeout
            while time.time() < deadline:
                try:
                    ev = self._sio.receive(timeout=2)
                    if ev and ev[0] == "session_ready":
                        return ev[1]
                except Exception:
                    pass
            return None
        deadline = time.time() + timeout
        while time.time() < deadline:
            try:
                ev = self._sio.receive(timeout=2)
                if ev and ev[0] == "guidance_update":
                    return ev[1]
            except Exception:
                pass
        return None

    def run(self, steps: list[Step]) -> list[dict]:
        results = []
        for step in steps:
            print(f"\n{CYAN}--- {step.label} ---{RESET}")
            print(f"  event: {step.event}")

            movement    = _describe_movement(step.event, step.data)
            expectation = _describe_expectation(step)
            response    = self._emit_and_wait(step.event, step.data)

            if step.event in _SKIP_EVENTS:
                if response:
                    print(f"  session_ready: {response}")
                    results.append({"step": step.label, "status": "SKIP",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                else:
                    print(f"  {YELLOW}No session_ready{RESET}")
                    results.append({"step": step.label, "status": "TIMEOUT",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                continue

            if (step.expected_action is None
                    and not step.guidance_must_contain
                    and step.forbidden_action is None):
                if response:
                    g = response.get("guidance", "")
                    a = response.get("action", "?")
                    print(f"  guidance: {g!r}  action: {a}")
                    results.append({"step": step.label, "status": "SKIP",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": a, "llm_guidance": g, "failures": []})
                else:
                    print(f"  {YELLOW}timeout{RESET}")
                    results.append({"step": step.label, "status": "TIMEOUT",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                continue

            if response is None:
                print(f"  {YELLOW}No response (timeout){RESET}")
                results.append({"step": step.label, "status": "TIMEOUT",
                                "movement": movement, "expectation": expectation,
                                "llm_action": "--", "llm_guidance": "--",
                                "failures": []})
                continue

            llm_guidance = response.get("guidance", "")
            llm_action   = response.get("action", "move")
            print(f"  guidance: {llm_guidance!r}")
            print(f"  action  : {llm_action}")

            passed   = True
            failures = []

            if step.expected_action and llm_action != step.expected_action:
                passed = False
                failures.append(
                    f'Action: got "{llm_action}", expected "{step.expected_action}"')

            if step.guidance_must_contain:
                if not any(kw.lower() in llm_guidance.lower()
                           for kw in step.guidance_must_contain):
                    passed = False
                    failures.append(
                        "Guidance missing any of: "
                        + ", ".join(step.guidance_must_contain))

            if step.forbidden_action and llm_action == step.forbidden_action:
                passed = False
                failures.append(
                    f'Action "{llm_action}" must NOT fire here')

            if passed:
                print(f"  {GREEN}{BOLD}PASS{RESET}")
            else:
                for f in failures:
                    print(f"  {RED}FAIL: {f}{RESET}")

            results.append({"step": step.label,
                            "status": "PASS" if passed else "FAIL",
                            "movement": movement, "expectation": expectation,
                            "llm_action": llm_action, "llm_guidance": llm_guidance,
                            "failures": failures})
            time.sleep(0.4)

        return results

    def print_summary(self, label: str, results: list[dict]) -> bool:
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        passed  = [r for r in checked if r["status"] == "PASS"]
        failed  = [r for r in checked if r["status"] == "FAIL"]
        timeout = [r for r in results if r["status"] == "TIMEOUT"]

        print(f"\n{'─'*60}")
        print(f"{BOLD}{label}{RESET}  "
              f"{GREEN}{len(passed)} passed{RESET}  "
              f"{RED}{len(failed)} failed{RESET}  "
              f"{YELLOW}{len(timeout)} timeout{RESET}  "
              f"/ {len(checked)} checked")
        for r in failed:
            print(f"  {RED}x {r['step']}{RESET}")
            for f in r.get("failures", []):
                print(f"      {f}")
        return len(failed) == 0 and len(timeout) == 0


# ---------------------------------------------------------------------------
# Word report
# ---------------------------------------------------------------------------

_STATUS_COLORS = {"PASS": "00B050", "FAIL": "FF0000",
                  "TIMEOUT": "FF8C00", "SKIP": "808080"}


def _colored_cell(cell, text: str, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    from docx.shared import RGBColor
    cell.text = text
    run = (cell.paragraphs[0].runs[0]
           if cell.paragraphs[0].runs
           else cell.paragraphs[0].add_run(text))
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, pt: int = 9, bold=False,
                   italic=False, color_rgb=None) -> None:
    cell.text = ""
    for i, line in enumerate(text.split("\n")):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run  = para.add_run(line)
        run.font.size   = __import__("docx").shared.Pt(pt)
        run.font.bold   = bold
        run.font.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb


def write_word_report(all_results: dict[str, list[dict]],
                      api_base: str, out_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin  = sec.bottom_margin = Inches(0.7)

    t = doc.add_heading(
        "CHIVA Streaming — Test Set 4 Report (Mendoza Shunt Types I–VI)", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = m.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}    |    "
                   f"API: {api_base}    |    Scenarios: {len(all_results)}")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    doc.add_heading("Overall Summary", level=1)
    stbl = doc.add_table(rows=1, cols=5)
    stbl.style = "Table Grid"
    for i, h in enumerate(["Scenario", "Description", "Passed", "Failed", "Timeouts"]):
        c = stbl.rows[0].cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True

    overall_pass = True
    for name, results in all_results.items():
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")
        row = stbl.add_row().cells
        row[0].text = name
        row[1].text = _SCENARIO_DESCRIPTIONS.get(name, "")
        row[2].text = str(p); row[3].text = str(f); row[4].text = str(t)
        if f or t:
            overall_pass = False

    doc.add_paragraph()
    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run("OVERALL: " + ("ALL PASS" if overall_pass else "FAILURES PRESENT"))
    vr.font.bold = True; vr.font.size = Pt(13)
    vr.font.color.rgb = (RGBColor(0x00, 0xB0, 0x50) if overall_pass
                         else RGBColor(0xFF, 0x00, 0x00))

    for name, results in all_results.items():
        doc.add_page_break()
        doc.add_heading(f"Scenario: {name.upper()}", level=1)
        dp = doc.add_paragraph(_SCENARIO_DESCRIPTIONS.get(name, ""))
        if dp.runs:
            dp.runs[0].font.italic = True

        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")

        sl = doc.add_paragraph()
        sl.add_run(f"Checked: {len(checked)}   |   ")
        pr = sl.add_run(f"Passed: {p}"); pr.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
        sl.add_run("   |   ")
        fr = sl.add_run(f"Failed: {f}"); fr.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        sl.add_run("   |   ")
        tr = sl.add_run(f"Timeouts: {t}"); tr.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        doc.add_paragraph()

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        for i, h in enumerate(["#", "Surgeon Action", "Expected",
                                "LLM Guidance", "Action", "Status"]):
            c = tbl.rows[0].cells[i]
            c.text = h
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.bold = True
        for i, w in enumerate([0.25, 1.80, 2.10, 2.10, 0.80, 0.65]):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)

        for idx, r in enumerate(results, 1):
            status = r.get("status", "SKIP")
            row    = tbl.add_row().cells
            row[0].text = str(idx)
            _set_cell_text(row[1], r.get("movement", "--"), pt=8)
            _set_cell_text(row[2], r.get("expectation", "--"), pt=8, italic=True)
            _set_cell_text(row[3], r.get("llm_guidance", "--") or "--", pt=8)
            _set_cell_text(row[4], r.get("llm_action", "--") or "--", pt=9, bold=True)
            _colored_cell(row[5], status, _STATUS_COLORS.get(status, "808080"))
            if status == "FAIL" and r.get("failures"):
                fr_row = tbl.add_row().cells
                merged = fr_row[0].merge(fr_row[5])
                merged.text = ""
                for fn in r["failures"]:
                    pp = merged.add_paragraph(f"  x  {fn}")
                    if pp.runs:
                        pp.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        pp.runs[0].font.size      = Pt(8)
                        pp.runs[0].font.italic    = True

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"\nReport saved: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="CHIVA stream test set 4 — Mendoza shunt types I–VI")
    parser.add_argument("scenario", nargs="?", default="all",
                        choices=list(_SCENARIOS.keys()) + ["all"],
                        help="Scenario to run (default: all)")
    parser.add_argument("--api",  default="http://localhost:7861")
    parser.add_argument("--all",  action="store_true", dest="run_all")
    args = parser.parse_args()

    to_run = list(_SCENARIOS.items()) if (args.run_all or args.scenario == "all") \
             else [(args.scenario, _SCENARIOS[args.scenario])]

    runner = ScenarioRunner(args.api)
    all_ok = True
    all_results: dict[str, list[dict]] = {}

    try:
        runner.connect()
        for name, builder in to_run:
            print(f"\n{'='*60}")
            print(f"{BOLD}SCENARIO: {name.upper()}{RESET}")
            print(f"{'='*60}")
            steps   = builder()
            results = runner.run(steps)
            ok      = runner.print_summary(name, results)
            all_ok  = all_ok and ok
            all_results[name] = results
    except KeyboardInterrupt:
        print("\nInterrupted.")
    finally:
        runner.disconnect()

    if len(to_run) > 1:
        print(f"\n{'='*60}")
        print(f"{BOLD}OVERALL: {'ALL PASS' if all_ok else 'FAILURES PRESENT'}{RESET}")

    if all_results:
        ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
        here     = os.path.dirname(os.path.abspath(__file__))
        out_path = os.path.join(here, "results", f"test_set4_mendoza_{ts}.docx")
        try:
            write_word_report(all_results, args.api, out_path)
        except Exception as exc:
            print(f"{YELLOW}Warning: could not write Word report — {exc}{RESET}")

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
