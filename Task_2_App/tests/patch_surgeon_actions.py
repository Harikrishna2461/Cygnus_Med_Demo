"""
Patch the Word report: rewrite the 'Surgeon Action' column (col index 2
of the step table) in plain English with short sentences.
Also rewrites the same field in the Clinical Deep-Dive section.
"""
from __future__ import annotations
import os, sys, copy
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Simplified surgeon action texts (22 steps) ────────────────────────────────
SIMPLE_ACTIONS = [
    # Step 1 – SFJ Arrival
    "Probe placed at the groin crease. Two vessels visible: CFV (N1, larger) and GSV (N2, smaller). Both at the junction on the same plane.",

    # Step 2 – SFJ Orientation
    "Probe rotated to get a clear view of the SFJ. Terminal valve spotted just below the junction. Probe now centred on the saphenofemoral junction.",

    # Step 3 – SFJ Maneuvers
    "Valsalva and Paranà tests done at the SFJ. Both show reversed blood flow from CFV into GSV. SFJ is incompetent.",

    # Step 4 – EP N1→N2 Marked
    "EP N1→N2 clip marked at the SFJ. Q1 is answered. Probe stays at the groin before moving down the thigh.",

    # Step 5 – Proximal Thigh
    "Probe moved to the upper thigh just below the groin. GSV visible inside the saphenous compartment. No perforator found at this level.",

    # Step 6 – Hunterian Zone
    "Probe at mid-thigh (Hunterian canal). GSV still clear in the fascial compartment. Looking for a perforator junction — none found yet.",

    # Step 7 – Mid-Thigh Reflux Test
    "Probe held at mid-thigh. Paranà maneuver done. GSV shows reflux over 500 ms. RP N2→N1 clip marked. Type 1 circuit confirmed.",

    # Step 8 – Dodd Zone Transition
    "Probe moved to the lower thigh near the knee (Dodd zone). GSV still in compartment. No escape to a tributary here. Moving into the calf next.",

    # Step 9 – Upper Calf Entry
    "Probe moved to the upper medial calf, just below the knee. GSV entering the calf is visible. No tributary junction found yet.",

    # Step 10 – Mid-Calf Tracking
    "Probe at mid-calf. GSV is narrower here. A small N3 structure appears next to the GSV — possible tributary junction.",

    # Step 11 – EP N2→N3 Confirmed
    "Paranà at mid-calf confirms flow from GSV into the N3 tributary. EP N2→N3 clip marked at posY 0.68. Q3 answered. Three clips confirmed.",

    # Step 12 – Tributary Tracking
    "Probe moved down the medial calf following the N3 tributary. Tributary widens near a perforator site. Re-entry not confirmed yet.",

    # Step 13 – Re-entry Perforator Zone
    "Probe near a fascial defect in the lower calf. The tributary passes through the fascia. Paranà shows inward diastolic flow — re-entry perforator pattern.",

    # Step 14 – RP N3→N1 Confirmed
    "RP N3→N1 clip marked at the lower calf. All four circuit clips confirmed. Circuit is complete. Elimination test now required.",

    # Step 15 – Maneuver Requested
    "Probe still at the re-entry site. Surgeon confirms elimination test is needed. No new clips added.",

    # Step 16 – SPJ Approach
    "Probe moved to the popliteal fossa (back of knee). Popliteal vein (N1) and SSV (N2) visible. Checking if SSV has a separate reflux circuit.",

    # Step 17 – SPJ Assessed
    "Paranà at SPJ: SSV shows no reflux. SPJ is competent. No SSV circuit found. Surgeon preparing to return to the mid-calf.",

    # Step 18 – Return to Escape Site
    "Probe back at mid-calf over the EP N2→N3 escape site. This is the compression point for the elimination test.",

    # Step 19 – Elim Test Compression
    "Surgeon compresses the N3 tributary at the escape site. GSV reflux continues at SFJ during compression. Trunk reflux is independent — confirms Type 1+2, not Type 3.",

    # Step 20 – Elim Test Result Recorded
    "Surgeon records elimTest=Reflux on the EP N2→N3 clip. This is the final data point. Type 1+2 classification is confirmed.",

    # Step 21 – SFJ Return Post-Diagnosis
    "Probe back at SFJ for documentation. Circuit fully classified. SFJ confirmed as the primary ligation site.",

    # Step 22 – Final State Check
    "Final probe position at SFJ. Examination complete. Type 1+2 diagnosis finalised. SFJ ligation and escape tributary disconnection to be planned.",
]


def _set_cell_text(cell, text: str, size: int = 7):
    """Clear cell and write plain text at given font size."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    # Use the first paragraph; clear all others
    first_para = cell.paragraphs[0]
    # Remove extra paragraphs
    for extra in cell.paragraphs[1:]:
        p = extra._element
        p.getparent().remove(p)
    first_para.clear()
    run = first_para.add_run(text)
    run.font.size = Pt(size)


def patch(src_path: str, dst_path: str):
    doc = Document(src_path)

    tables = doc.tables
    # The step table is the 2nd table (index 1; index 0 is the scoreboard)
    step_table = tables[1]

    data_rows = step_table.rows[1:]  # skip header row
    assert len(data_rows) == 22, f"Expected 22 data rows, got {len(data_rows)}"

    for i, row in enumerate(data_rows):
        cell = row.cells[2]   # column index 2 = "Surgeon Action"
        _set_cell_text(cell, SIMPLE_ACTIONS[i], size=7)

    # ── Also patch the Clinical Deep-Dive section (Section 3) ─────────────────
    # The deep-dive has headings + paragraphs. Each step paragraph has runs
    # with bold labels. We look for the run after "Surgeon action: " bold label.
    in_deep_dive = False
    step_counter = 0

    for para in doc.paragraphs:
        # Detect deep-dive section start
        if "Clinical Deep-Dive" in para.text:
            in_deep_dive = True
            step_counter = 0
            continue

        if not in_deep_dive:
            continue

        # Each step heading starts "Step N:"
        if para.style.name.startswith("Heading") and "Step" in para.text:
            step_counter += 1
            continue

        # Within a step paragraph, find "Surgeon action:" run and replace
        # the very next run's text
        if step_counter < 1:
            continue

        runs = para.runs
        for j, run in enumerate(runs):
            if run.bold and "Surgeon action" in run.text:
                # The next run holds the value
                if j + 1 < len(runs):
                    runs[j + 1].text = SIMPLE_ACTIONS[step_counter - 1] + "\n"
                break

    doc.save(dst_path)
    print(f"[OK] Patched report saved -> {dst_path}")


if __name__ == "__main__":
    import glob as _glob
    # Find the most recent report
    pattern = r"C:\Users\Krish\Downloads\Task2_Type12_LLMEval_*.docx"
    files = sorted(_glob.glob(pattern))
    if not files:
        print("No report found.")
        sys.exit(1)
    src = files[-1]
    # Save with _patched suffix
    dst = src.replace(".docx", "_patched.docx")
    print(f"Patching: {src}")
    patch(src, dst)
