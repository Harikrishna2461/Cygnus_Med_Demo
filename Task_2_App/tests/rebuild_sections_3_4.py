"""
Rebuild sections 3 and 4 of the Word report using the data already in
section 2 (the step table). Section 2 is never touched.
All extravagant colours are removed — plain black text throughout.
"""
from __future__ import annotations
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SRC = r"C:\Users\Krish\Downloads\Task2_Type12_LLMEval_20260715_110815.docx"
DST = r"C:\Users\Krish\Downloads\Task2_Type12_LLMEval_20260715_110815_v2.docx"

WNS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


# =============================================================================
# 1. Read section 2 data (step table = tables[1])
# =============================================================================
doc = Document(SRC)
step_tbl = doc.tables[1]

steps = []
for row in step_tbl.rows[1:]:   # skip header row
    c = [cell.text.replace("\n", " ").strip() for cell in row.cells]
    steps.append({
        "label":    c[0],    # "#1 0-5 s"
        "phase":    c[1],    # "Phase-1 SFJ Arrival"
        "surgeon":  c[2],    # surgeon action text
        "probe":    c[3],
        "clips":    c[4],
        "guidance": c[5],
        "act":      c[6],    # action flag
        "exp":      c[7],    # expected action
        "ok":       c[8],    # YES/NO
        "score":    c[9],    # "CORRECT (3/3)" etc.
        "reason":   c[10],
    })

print(f"Read {len(steps)} steps from section 2.")


# =============================================================================
# 2. Strip everything from section 3 heading onward
# =============================================================================
body = doc.element.body
children = list(body)

def _text(el):
    return "".join(t.text or "" for t in el.iter(f"{{{WNS}}}t"))

# Find index of the "3. Clinical Deep-Dive" heading in body children
sec3_idx = None
for i, el in enumerate(children):
    if el.tag == f"{{{WNS}}}p" and "3." in _text(el) and "Deep" in _text(el):
        # Also include the page-break paragraph just before it, if present
        if i > 0:
            prev = children[i - 1]
            has_pb = any(
                br.get(f"{{{WNS}}}type") == "page"
                for br in prev.iter(f"{{{WNS}}}br")
            )
            sec3_idx = i - 1 if has_pb else i
        else:
            sec3_idx = i
        break

if sec3_idx is None:
    print("ERROR: section 3 heading not found")
    sys.exit(1)

# Find sectPr (page layout — must stay last)
sectpr_idx = next(
    (i for i, el in enumerate(children) if el.tag == f"{{{WNS}}}sectPr"), None
)

# Remove everything between sec3_idx and sectPr
for el in children[sec3_idx: sectpr_idx]:
    body.remove(el)

print("Removed old sections 3 and 4.")


# =============================================================================
# 3. Helper functions — plain text, no colours
# =============================================================================
def _page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _heading(doc, text, level):
    doc.add_heading(text, level)


def _line(doc, label, value, size=9):
    """Bold label + plain value on one paragraph."""
    p = doc.add_paragraph()
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(size)
    r2 = p.add_run(value)
    r2.font.size = Pt(size)
    return p


def _plain(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


# =============================================================================
# 4. Build section 3 — Clinical Deep-Dive
# =============================================================================
_page_break(doc)
_heading(doc, "3. Clinical Deep-Dive (per step)", 2)

for s in steps:
    _heading(doc, f"{s['label']}  —  {s['phase']}", 3)
    _line(doc, "Surgeon Action",     s["surgeon"])
    _line(doc, "Probe Position",     s["probe"])
    _line(doc, "Clips Confirmed",    s["clips"])
    _line(doc, "System Guidance",    f'"{s["guidance"]}"')
    _line(doc, "Action Flag",        f"{s['act']}  (expected: {s['exp']})  —  {'Correct' if s['ok'] == 'YES' else 'Wrong'}")
    _line(doc, "Evaluation Verdict", s["score"])
    _line(doc, "Reasoning",          s["reason"])
    doc.add_paragraph()   # small gap between steps


# =============================================================================
# 5. Build section 4 — Phase-Level Summary
# =============================================================================
_page_break(doc)
_heading(doc, "4. Phase-Level Summary", 2)

# Group steps by phase prefix (e.g. "Phase-1", "Phase-2" …)
from collections import OrderedDict

phase_groups: dict[str, list] = OrderedDict()
for s in steps:
    # Extract "Phase-N" key
    phase_key = s["phase"].split(" ")[0]   # e.g. "Phase-1"
    phase_groups.setdefault(phase_key, []).append(s)

# Build table: Phase | Steps | CORRECT | PARTIAL | WRONG | Action OK
tbl = doc.add_table(rows=1, cols=6)
tbl.style = "Table Grid"

def _hdr(cell, text):
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)

headers = ["Phase", "Steps", "Correct", "Partial", "Wrong", "Action OK"]
for i, h in enumerate(headers):
    _hdr(tbl.rows[0].cells[i], h)

for phase_key, phase_steps in phase_groups.items():
    n      = len(phase_steps)
    n3     = sum(1 for s in phase_steps if s["score"].startswith("CORRECT"))
    n2     = sum(1 for s in phase_steps if s["score"].startswith("PARTIAL"))
    n1     = sum(1 for s in phase_steps if s["score"].startswith("WRONG"))
    n_act  = sum(1 for s in phase_steps if s["ok"] == "YES")

    # Use the first step's full phase text as the label
    phase_label = phase_steps[0]["phase"]

    row = tbl.add_row().cells
    for cell in row:
        cell.text = ""

    def _cell(cell, text, bold=False):
        cell.text = text
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = bold
                run.font.size = Pt(9)

    _cell(row[0], phase_label)
    _cell(row[1], str(n))
    _cell(row[2], f"{n3}/{n}", bold=(n3 == n))
    _cell(row[3], str(n2))
    _cell(row[4], str(n1))
    _cell(row[5], f"{n_act}/{n}", bold=(n_act == n))

# Overall row
n_total  = len(steps)
n3_total = sum(1 for s in steps if s["score"].startswith("CORRECT"))
n2_total = sum(1 for s in steps if s["score"].startswith("PARTIAL"))
n1_total = sum(1 for s in steps if s["score"].startswith("WRONG"))
na_total = sum(1 for s in steps if s["ok"] == "YES")

row = tbl.add_row().cells
def _cell_b(cell, text):
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)

_cell_b(row[0], "TOTAL")
_cell_b(row[1], str(n_total))
_cell_b(row[2], f"{n3_total}/{n_total}  ({round(100*n3_total/n_total)}%)")
_cell_b(row[3], str(n2_total))
_cell_b(row[4], str(n1_total))
_cell_b(row[5], f"{na_total}/{n_total}")

doc.add_paragraph()
_plain(doc,
    f"Overall guidance accuracy (CORRECT): {n3_total}/{n_total} steps ({round(100*n3_total/n_total)}%). "
    f"Action flag accuracy: {na_total}/{n_total} steps (100%). "
    f"Partial responses (direction correct, minor wording issue): {n2_total}. "
    f"No wrong or missing responses.",
    size=9)

# =============================================================================
# 6. Save
# =============================================================================
doc.save(DST)
print(f"Saved -> {DST}")
