"""
CHIVA Streaming Guidance — Consolidated Test Scenarios (v2)

All 6 Mendoza shunt types + Reflux elimination variant + SPJ-entry (Type 2A) + left-leg Type I.
~290 checked steps across 9 scenarios.

Region names match stream.html mapReg() exactly:
  SFJ          posY 0.00–0.07  surface "anterior-medial"
  Upper Thigh  posY 0.08–0.20  surface "medial"
  Dodd         posY 0.21–0.33  surface "medial"          (mid-thigh Dodd perforator zone)
  Hunterian    posY 0.34–0.47  surface "medial"          (lower thigh Hunterian perforator zone)
  Giacomini    posY 0.08–0.47  surface "posterior"       (posterior thigh / Giacomini vein)
  SPJ          posY 0.48–0.57  surface "posterior"
  Calf         posY 0.58–0.88  surface "medial"
  SSV          posY 0.58–0.88  surface "posterior"
  Ankle        posY 0.89–1.00  surface "medial" / "lateral"

Complete-rule gate:  old_max_visited_pos_y >= 0.48  (probe must reach SPJ zone BEFORE RP N2->N1)

Engine rules (_rule_based_action in streaming_guidance_engine.py):
  Rule 1  maneuver : ep_n2_n3 + rp_n3_n1 + rp_n2_n1 + ep_n2_n3 has no elimTest
  Rule 2  complete : elimTest on EP N2->N3 == "No Reflux"                    (Type III)
  Rule 3  complete : elimTest on EP N2->N3 == "Reflux"                       (Type I+2 / III-b)
  Rule 4  complete : ep_n1_n3 + rp_n2_n1                                     (Type IV)
  Rule 5  complete : ep_n1_n3 + rp_n3_n1 + NOT rp_n2_n1 + NOT rp_n3_n2     (Type VI)
  Rule 5b complete : ep_n1_n3 + rp_n3_n2 + ep_n2_n3 + rp_n3_n1             (Type V)
  Rule 6  complete : ep_n1_n2 + rp_n2_n1 + NOT ep_n2_n3 + old_max >= 0.48  (Type I / 2A)
  Rule 7  complete : ep_n2_n3 + rp_n3_n1 + NOT ep_n1_n2 + NOT ep_n1_n3 + NOT rp_n2_n1  (Type II)

Usage:
    python tests/run_scenarios.py [--api http://localhost:7861] [<scenario>|all]

Output: tests/results/scenarios_<timestamp>.docx
"""
from __future__ import annotations

import argparse
import io
import os
import sys
import time

if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

try:
    import socketio
except ImportError:
    print("ERROR: python-socketio not installed.  Run: pip install python-socketio[client]")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Step definition
# ---------------------------------------------------------------------------

@dataclass
class Step:
    label: str
    event: str
    data: dict
    expected_action: Optional[str] = None
    guidance_must_contain: list[str] = field(default_factory=list)
    forbidden_action: Optional[str] = None


def _pm(sid, region, pos_y, surface, leg="right") -> dict:
    return {"session_id": sid, "region": region, "pos_y_ratio": pos_y,
            "surface": surface, "leg": leg}

def _cm(sid, flow, ft, tt, pos_y, region, surface, elim="", leg="right") -> dict:
    return {"session_id": sid, "flow": flow, "from_type": ft, "to_type": tt,
            "pos_y_ratio": pos_y, "leg": leg, "region": region,
            "surface": surface, "elimination_test": elim}


# ---------------------------------------------------------------------------
# Scenario 1 — mendoza_type1  (Type I: R1→R2→R1)  35 steps
# ---------------------------------------------------------------------------
# Deep system enters GSV trunk at SFJ (EP N1→N2 — incompetent terminal valve).
# Trunk refluxes distally.  Blood re-enters deep directly from the trunk
# via a calf perforator (RP N2→N1).  No tributary (N3) involvement.
#
# Rule 6 gate: old_max_visited >= 0.48.
# Strategy: systematic sweep groin → ankle → return sweep → mark RP N2→N1.
# ---------------------------------------------------------------------------

def _type1_right(sid="s1_type1_r") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # --- SFJ assessment (posY 0.04-0.07) ---
        Step("P01  SFJ approach — transverse B-mode, Mickey Mouse sign",
             "probe_move", _pm(sid, "SFJ", 0.04, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction", "saphenofemoral"]),

        Step("P02  SFJ — Doppler gate on femoral side of terminal valve",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move"),

        Step("P03  SFJ — Valsalva + Paranà both positive: SFJ incompetent",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move"),

        # EP N1->N2: deep blood enters GSV trunk at SFJ
        Step("CM-1  EP N1->N2 at SFJ (posY=0.06) — SFJ incompetent confirmed",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.06, "SFJ", "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["distal", "thigh", "trunk", "GSV", "medial"]),

        # --- Upper thigh sweep (posY 0.08-0.20) ---
        Step("P04  Upper Thigh — saphenous eye in transverse, measure GSV diameter",
             "probe_move", _pm(sid, "Upper Thigh", 0.10, "medial"),
             expected_action="move"),

        Step("P05  Upper Thigh mid — Paranà squeeze: retrograde trunk flow confirmed",
             "probe_move", _pm(sid, "Upper Thigh", 0.14, "medial"),
             expected_action="move"),

        Step("P06  Upper Thigh distal",
             "probe_move", _pm(sid, "Upper Thigh", 0.18, "medial"),
             expected_action="move"),

        # --- Dodd zone (posY 0.21-0.33) ---
        Step("P07  Dodd proximal — scanning mid-thigh for escape perforators",
             "probe_move", _pm(sid, "Dodd", 0.23, "medial"),
             expected_action="move",
             guidance_must_contain=["Hunterian", "Dodd", "perforator", "escape", "tributary"]),

        Step("P08  Dodd mid — no N3 tributary filling",
             "probe_move", _pm(sid, "Dodd", 0.27, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P09  Dodd distal",
             "probe_move", _pm(sid, "Dodd", 0.31, "medial"),
             expected_action="move"),

        # --- Hunterian zone (posY 0.34-0.47) ---
        Step("P10  Hunterian proximal",
             "probe_move", _pm(sid, "Hunterian", 0.35, "medial"),
             expected_action="move"),

        Step("P11  Hunterian mid — all 3 maneuvers: no EP N2->N3 found",
             "probe_move", _pm(sid, "Hunterian", 0.39, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P12  Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.44, "medial"),
             expected_action="move"),

        # --- Popliteal / SPJ zone (posY 0.48-0.57) max_visited now >=0.48 ---
        Step("P13  SPJ — lateral decubitus, both Paranà + CR required",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior", "junction"]),

        Step("P14  SPJ mid — only one maneuver positive: SPJ competent",
             "probe_move", _pm(sid, "SPJ", 0.52, "posterior"),
             expected_action="move"),

        Step("P15  SPJ distal — Giacomini vein checked: no significant connection",
             "probe_move", _pm(sid, "SPJ", 0.56, "posterior"),
             expected_action="move"),

        # --- Calf medial (posY 0.58-0.88): GSV trunk ---
        Step("P16  Calf upper — entering calf, posterior view first",
             "probe_move", _pm(sid, "SSV", 0.59, "posterior"),
             expected_action="move"),

        Step("P17  Calf medial upper — GSV trunk in saphenous eye",
             "probe_move", _pm(sid, "Calf", 0.62, "medial"),
             expected_action="move"),

        Step("P18  Calf medial mid",
             "probe_move", _pm(sid, "Calf", 0.66, "medial"),
             expected_action="move"),

        Step("P19  Calf medial — Paranà maneuver: retrograde flow in GSV trunk",
             "probe_move", _pm(sid, "Calf", 0.70, "medial"),
             expected_action="move"),

        Step("P20  Calf medial lower",
             "probe_move", _pm(sid, "Calf", 0.74, "medial"),
             expected_action="move"),

        Step("P21  Calf lower",
             "probe_move", _pm(sid, "Calf", 0.78, "medial"),
             expected_action="move"),

        Step("P22  Calf lower — no tributary escape at any calf level",
             "probe_move", _pm(sid, "Calf", 0.82, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # --- Ankle (posY 0.89-1.00) ---
        Step("P23  Ankle medial — GSV at malleolus, confirm trunk continuity",
             "probe_move", _pm(sid, "Ankle", 0.90, "medial"),
             expected_action="move"),

        Step("P24  Ankle distal",
             "probe_move", _pm(sid, "Ankle", 0.94, "medial"),
             expected_action="move"),

        # --- Return sweep proximal — confirming no escape at any level ---
        Step("P25  Return: ankle to lower calf",
             "probe_move", _pm(sid, "Calf", 0.87, "medial"),
             expected_action="move"),

        Step("P26  Return: lower calf",
             "probe_move", _pm(sid, "Calf", 0.80, "medial"),
             expected_action="move"),

        Step("P27  Return: calf mid",
             "probe_move", _pm(sid, "Calf", 0.74, "medial"),
             expected_action="move"),

        Step("P28  Return: calf upper — scanning for inward perforator flow on relaxation",
             "probe_move", _pm(sid, "Calf", 0.68, "medial"),
             expected_action="move"),

        # RP N2->N1: trunk blood re-enters deep via Boyd/paratibial perforator.
        # old_max_visited=0.94 >= 0.48.  ep_n1_n2=T, ep_n2_n3=F -> Rule 6 fires.
        Step("CM-2  RP N2->N1 calf (posY=0.68) — trunk re-enters deep via perforator — COMPLETE (Type I)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.68, "Calf", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P29  Verify complete — calf",
             "probe_move", _pm(sid, "Calf", 0.66, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P30  Verify complete — SPJ zone",
             "probe_move", _pm(sid, "SPJ", 0.51, "posterior"),
             expected_action="complete"),

        Step("P31  Verify complete — Hunterian",
             "probe_move", _pm(sid, "Hunterian", 0.40, "medial"),
             expected_action="complete"),

        Step("P32  Verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 2 — mendoza_type2  (Type II: R2→R3→R1)  33 steps
# ---------------------------------------------------------------------------
# Both SFJ and SPJ terminal valves COMPETENT — no deep (N1) to trunk (N2) entry.
# The saphenous trunk itself is the source of reflux (pre-terminal incompetence).
# Trunk overflows into tributary at Dodd zone (EP N2->N3).
# Tributary re-enters deep (RP N3->N1).  No RP N2->N1.
#
# Rule 7: ep_n2_n3 + rp_n3_n1 + NOT ep_n1_n2 + NOT ep_n1_n3 + NOT rp_n2_n1
# ---------------------------------------------------------------------------

def _type2_right(sid="s2_type2_r") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # --- SFJ: competent ---
        Step("P01  SFJ — Valsalva negative, Paranà negative: terminal valve competent",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02  SFJ deeper — pre-terminal valve: competent",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move"),

        # --- Upper thigh: local trunk reflux detected ---
        Step("P03  Upper Thigh — saphenous eye confirmed; measuring diameter",
             "probe_move", _pm(sid, "Upper Thigh", 0.11, "medial"),
             expected_action="move"),

        Step("P04  Upper Thigh mid — Paranà: trunk shows local segment reflux",
             "probe_move", _pm(sid, "Upper Thigh", 0.16, "medial"),
             expected_action="move"),

        Step("P05  Upper Thigh distal",
             "probe_move", _pm(sid, "Upper Thigh", 0.20, "medial"),
             expected_action="move"),

        # --- Dodd zone: trunk overflow into tributary ---
        Step("P06  Dodd proximal — trunk reflux intensifies; watching for N3 tributary",
             "probe_move", _pm(sid, "Dodd", 0.24, "medial"),
             expected_action="move"),

        Step("P07  Dodd mid — colour flow shows tributaries filling from trunk",
             "probe_move", _pm(sid, "Dodd", 0.28, "medial"),
             expected_action="move"),

        # EP N2->N3: trunk overflows into tributary.  No EP N1->N2 ever confirmed.
        Step("CM-1  EP N2->N3 at Dodd (posY=0.28) — trunk overflows to tributary (Type II source)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.28, "Dodd", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P08  Dodd distal — following N3 tributary",
             "probe_move", _pm(sid, "Dodd", 0.33, "medial"),
             expected_action="move"),

        # --- Hunterian: no RP N2->N1 on trunk ---
        Step("P09  Hunterian proximal — trunk clean; no perforator re-entry from trunk",
             "probe_move", _pm(sid, "Hunterian", 0.36, "medial"),
             expected_action="move"),

        Step("P10  Hunterian mid",
             "probe_move", _pm(sid, "Hunterian", 0.41, "medial"),
             expected_action="move"),

        Step("P11  Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.46, "medial"),
             expected_action="move"),

        # --- Popliteal/SPJ: competent ---
        Step("P12  SPJ — both maneuvers: terminal valve competent",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P13  SPJ mid",
             "probe_move", _pm(sid, "SPJ", 0.53, "posterior"),
             expected_action="move"),

        # --- Following N3 tributary distally ---
        Step("P14  SSV upper — tracing N3 tributary on posterior calf",
             "probe_move", _pm(sid, "SSV", 0.58, "posterior"),
             expected_action="move"),

        Step("P15  SSV mid",
             "probe_move", _pm(sid, "SSV", 0.63, "posterior"),
             expected_action="move"),

        Step("P16  SSV mid-lower",
             "probe_move", _pm(sid, "SSV", 0.68, "posterior"),
             expected_action="move"),

        Step("P17  SSV lower",
             "probe_move", _pm(sid, "SSV", 0.73, "posterior"),
             expected_action="move"),

        Step("P18  Calf lateral — tributary turns lateral",
             "probe_move", _pm(sid, "Calf", 0.76, "lateral"),
             expected_action="move"),

        Step("P19  Calf lateral lower",
             "probe_move", _pm(sid, "Calf", 0.81, "lateral"),
             expected_action="move"),

        Step("P20  Ankle lateral — tracking tributary to re-entry",
             "probe_move", _pm(sid, "Ankle", 0.86, "lateral"),
             expected_action="move"),

        Step("P21  Ankle lateral distal",
             "probe_move", _pm(sid, "Ankle", 0.92, "lateral"),
             expected_action="move"),

        # Return sweep toward re-entry perforator
        Step("P22  Return: ankle to lower calf lateral",
             "probe_move", _pm(sid, "Calf", 0.87, "lateral"),
             expected_action="move"),

        Step("P23  Return: lower calf lateral",
             "probe_move", _pm(sid, "Calf", 0.81, "lateral"),
             expected_action="move"),

        Step("P24  Return: calf lateral mid — biphasic perforator detected",
             "probe_move", _pm(sid, "Calf", 0.75, "lateral"),
             expected_action="move"),

        Step("P25  Return: calf lateral — diastolic inward flow confirmed",
             "probe_move", _pm(sid, "Calf", 0.70, "lateral"),
             expected_action="move"),

        # RP N3->N1: N3 tributary re-enters deep.
        # NOT ep_n1_n2, NOT ep_n1_n3, NOT rp_n2_n1 -> Rule 7 fires.
        Step("CM-2  RP N3->N1 at calf lateral (posY=0.70) — tributary re-enters deep — COMPLETE (Type II)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.70, "Calf", "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P26  Verify complete — calf lateral",
             "probe_move", _pm(sid, "Calf", 0.68, "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P27  Verify complete — Dodd zone",
             "probe_move", _pm(sid, "Dodd", 0.28, "medial"),
             expected_action="complete"),

        Step("P28  Verify complete — SPJ",
             "probe_move", _pm(sid, "SPJ", 0.50, "posterior"),
             expected_action="complete"),

        Step("P29  Verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 3 — mendoza_type3  (Type III: R1→R2→R3→R1, elimination "No Reflux")
# ---------------------------------------------------------------------------
# Deep entry at SFJ (EP N1->N2).  ALL reflux volume conducted via trunk to
# tributary escape (EP N2->N3).  Trunk acts as conduit — not final re-entry.
# Tributary re-enters deep (RP N3->N1).  Trunk appears to reflux (RP N2->N1)
# but compression test abolishes it: "No Reflux" -> Type III.
#
# ORDERING CRITICAL: EP N2->N3 must be marked BEFORE RP N2->N1 to prevent
# Rule 6 (Type I) from firing.  With ep_n2_n3+rp_n2_n1+rp_n3_n1 -> Rule 1 maneuver.
# Then elimTest "No Reflux" on EP N2->N3 -> Rule 2 complete.
# ---------------------------------------------------------------------------

def _type3_no_reflux(sid="s3_type3_nr") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # --- SFJ: incompetent ---
        Step("P01  SFJ — both Valsalva AND Paranà positive: SFJ incompetent",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02  SFJ mid",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move"),

        # EP N1->N2: SFJ incompetent
        Step("CM-1  EP N1->N2 at SFJ (posY=0.06)",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.06, "SFJ", "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["distal", "trunk", "GSV", "thigh"]),

        # --- Upper thigh ---
        Step("P03  Upper Thigh — tracing trunk distally",
             "probe_move", _pm(sid, "Upper Thigh", 0.11, "medial"),
             expected_action="move"),

        Step("P04  Upper Thigh mid",
             "probe_move", _pm(sid, "Upper Thigh", 0.16, "medial"),
             expected_action="move"),

        Step("P05  Upper Thigh distal",
             "probe_move", _pm(sid, "Upper Thigh", 0.20, "medial"),
             expected_action="move"),

        # --- Dodd zone: find EP N2->N3 escape BEFORE marking RP N2->N1 ---
        Step("P06  Dodd proximal — ALL volume exits here; large tributary filling",
             "probe_move", _pm(sid, "Dodd", 0.24, "medial"),
             expected_action="move"),

        Step("P07  Dodd mid — tributary expands distal to this escape point",
             "probe_move", _pm(sid, "Dodd", 0.27, "medial"),
             expected_action="move"),

        # CRITICAL: EP N2->N3 marked BEFORE RP N2->N1.
        # This blocks Rule 6 from firing (would declare Type I complete).
        Step("CM-2  EP N2->N3 at Dodd (posY=0.26) — trunk conducts ALL volume here (no elimTest yet)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.26, "Dodd", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P08  Dodd distal — trunk still present but thin (most volume escaped)",
             "probe_move", _pm(sid, "Dodd", 0.32, "medial"),
             expected_action="move"),

        # --- Hunterian zone ---
        Step("P09  Hunterian proximal",
             "probe_move", _pm(sid, "Hunterian", 0.36, "medial"),
             expected_action="move"),

        Step("P10  Hunterian mid — trunk thin, no further escape",
             "probe_move", _pm(sid, "Hunterian", 0.41, "medial"),
             expected_action="move"),

        Step("P11  Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.46, "medial"),
             expected_action="move"),

        # --- Popliteal / SPJ: competent ---
        Step("P12  SPJ — only one maneuver positive: SPJ competent",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P13  SPJ mid",
             "probe_move", _pm(sid, "SPJ", 0.53, "posterior"),
             expected_action="move"),

        # --- Calf: follow N3 tributary and N2 trunk ---
        Step("P14  SSV upper — N3 tributary on posterior surface",
             "probe_move", _pm(sid, "SSV", 0.58, "posterior"),
             expected_action="move"),

        Step("P15  Calf medial upper — N2 trunk alongside",
             "probe_move", _pm(sid, "Calf", 0.61, "medial"),
             expected_action="move"),

        Step("P16  Calf medial mid",
             "probe_move", _pm(sid, "Calf", 0.66, "medial"),
             expected_action="move"),

        Step("P17  Calf medial — Paranà: trunk still shows retrograde flow (conducted)",
             "probe_move", _pm(sid, "Calf", 0.71, "medial"),
             expected_action="move"),

        Step("P18  Calf medial lower",
             "probe_move", _pm(sid, "Calf", 0.76, "medial"),
             expected_action="move"),

        # --- Ankle ---
        Step("P19  Ankle medial",
             "probe_move", _pm(sid, "Ankle", 0.89, "medial"),
             expected_action="move"),

        Step("P20  Ankle distal",
             "probe_move", _pm(sid, "Ankle", 0.94, "medial"),
             expected_action="move"),

        # --- Return sweep ---
        Step("P21  Return: ankle to lower calf",
             "probe_move", _pm(sid, "Calf", 0.87, "medial"),
             expected_action="move"),

        Step("P22  Return: lower calf",
             "probe_move", _pm(sid, "Calf", 0.80, "medial"),
             expected_action="move"),

        # RP N3->N1: N3 tributary re-enters deep.
        # ep_n2_n3_no_elim + rp_n3_n1, but rp_n2_n1 still absent -> Rule 1 NOT YET.
        Step("CM-3  RP N3->N1 at calf (posY=0.72) — N3 tributary re-enters deep (maneuver not yet — need RP N2->N1)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.72, "Calf", "medial"),
             expected_action="move",
             forbidden_action="maneuver"),

        # Return toward Dodd zone to find conducted RP N2->N1
        Step("P23  Return: Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.45, "medial"),
             expected_action="move"),

        Step("P24  Return: Hunterian mid",
             "probe_move", _pm(sid, "Hunterian", 0.40, "medial"),
             expected_action="move"),

        Step("P25  Return: Hunterian proximal",
             "probe_move", _pm(sid, "Hunterian", 0.35, "medial"),
             expected_action="move"),

        Step("P26  Return: Dodd distal",
             "probe_move", _pm(sid, "Dodd", 0.32, "medial"),
             expected_action="move"),

        Step("P27  Return: Dodd mid — trunk shows conducted retrograde flow",
             "probe_move", _pm(sid, "Dodd", 0.28, "medial"),
             expected_action="move"),

        # RP N2->N1: conducted trunk reflux (NOT independent — will be abolished on compression).
        # ep_n2_n3_no_elim + rp_n3_n1 + rp_n2_n1 -> Rule 1 MANEUVER fires.
        Step("CM-4  RP N2->N1 Dodd (posY=0.25) — conducted trunk reflux — MANEUVER fires",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.25, "Dodd", "medial"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary", "Doppler", "record"]),

        Step("P28  Maneuver state — return to escape point for compression test",
             "probe_move", _pm(sid, "Dodd", 0.26, "medial"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary"]),

        Step("P29  Maneuver state — calf trunk also checked: disappears on compression",
             "probe_move", _pm(sid, "Calf", 0.63, "medial"),
             expected_action="maneuver"),

        Step("P30  Return to Dodd for elimination test",
             "probe_move", _pm(sid, "Dodd", 0.26, "medial"),
             expected_action="maneuver"),

        # Elimination test: compress tributary at Dodd escape.
        # RP N2->N1 disappears -> "No Reflux" -> trunk was conduit only -> Type III.
        # Rule 2 fires.
        Step("CM-5  EP N2->N3 elim='No Reflux' at Dodd — trunk was conduit — COMPLETE (Type III)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.26, "Dodd", "medial", elim="No Reflux"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P31  Verify complete — Dodd",
             "probe_move", _pm(sid, "Dodd", 0.26, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P32  Verify complete — calf",
             "probe_move", _pm(sid, "Calf", 0.65, "medial"),
             expected_action="complete"),

        Step("P33  Verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 4 — mendoza_type3b  (Type III path, elimination "Reflux")  30 steps
# ---------------------------------------------------------------------------
# Identical probe/clip sequence to Type III but surgeon records "Reflux" on
# the elimination test — RP N2->N1 persists because trunk has independent reflux.
# This is the Type I+2 classification.  Rule 3 fires.
# ---------------------------------------------------------------------------

def _type3_reflux(sid="s4_type3_rf") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01  SFJ — incompetent",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("CM-1  EP N1->N2 at SFJ (posY=0.06)",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.06, "SFJ", "anterior-medial"),
             expected_action="move"),

        Step("P02  Upper Thigh",
             "probe_move", _pm(sid, "Upper Thigh", 0.12, "medial"),
             expected_action="move"),

        Step("P03  Upper Thigh mid",
             "probe_move", _pm(sid, "Upper Thigh", 0.18, "medial"),
             expected_action="move"),

        Step("P04  Dodd proximal",
             "probe_move", _pm(sid, "Dodd", 0.23, "medial"),
             expected_action="move"),

        # EP N2->N3 marked BEFORE RP N2->N1 (same as Type III).
        Step("CM-2  EP N2->N3 at Dodd (posY=0.26) — tributary escape (no elimTest)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.26, "Dodd", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P05  Hunterian",
             "probe_move", _pm(sid, "Hunterian", 0.36, "medial"),
             expected_action="move"),

        Step("P06  Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.44, "medial"),
             expected_action="move"),

        Step("P07  SPJ — competent",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P08  Calf medial — following trunk",
             "probe_move", _pm(sid, "Calf", 0.62, "medial"),
             expected_action="move"),

        Step("P09  Calf medial mid",
             "probe_move", _pm(sid, "Calf", 0.68, "medial"),
             expected_action="move"),

        Step("P10  Calf medial lower",
             "probe_move", _pm(sid, "Calf", 0.74, "medial"),
             expected_action="move"),

        Step("P11  Ankle",
             "probe_move", _pm(sid, "Ankle", 0.90, "medial"),
             expected_action="move"),

        Step("P12  Return: lower calf",
             "probe_move", _pm(sid, "Calf", 0.81, "medial"),
             expected_action="move"),

        Step("CM-3  RP N3->N1 at calf (posY=0.71) — N3 re-enters deep",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.71, "Calf", "medial"),
             expected_action="move",
             forbidden_action="maneuver"),

        Step("P13  Return: Hunterian",
             "probe_move", _pm(sid, "Hunterian", 0.38, "medial"),
             expected_action="move"),

        Step("P14  Return: Dodd",
             "probe_move", _pm(sid, "Dodd", 0.29, "medial"),
             expected_action="move"),

        # RP N2->N1: trunk conducted reflux (may be independent here).
        # ep_n2_n3_no_elim + rp_n3_n1 + rp_n2_n1 -> Rule 1 MANEUVER.
        Step("CM-4  RP N2->N1 Dodd (posY=0.25) — trunk reflux — MANEUVER fires",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.25, "Dodd", "medial"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary", "Doppler", "record"]),

        Step("P15  Maneuver state — at escape point for compression",
             "probe_move", _pm(sid, "Dodd", 0.26, "medial"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary"]),

        Step("P16  Maneuver — check trunk at calf: still refluxing during compression",
             "probe_move", _pm(sid, "Calf", 0.64, "medial"),
             expected_action="maneuver"),

        Step("P17  Return to escape point",
             "probe_move", _pm(sid, "Dodd", 0.26, "medial"),
             expected_action="maneuver"),

        # Elimination test: compress tributary.
        # RP N2->N1 PERSISTS -> "Reflux" -> trunk has independent reflux -> Type I+2.
        # Rule 3 fires.
        Step("CM-5  EP N2->N3 elim='Reflux' at Dodd — trunk has independent reflux — COMPLETE (Type I+2)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.26, "Dodd", "medial", elim="Reflux"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P18  Verify complete — Dodd",
             "probe_move", _pm(sid, "Dodd", 0.26, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P19  Verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),

        Step("P20  Verify complete — calf",
             "probe_move", _pm(sid, "Calf", 0.62, "medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 5 — mendoza_type4  (Type IV: R1→R3→R2→R1)  33 steps
# ---------------------------------------------------------------------------
# SFJ and SPJ both competent.  Pelvic vein or perforator exits deep directly
# into a tributary (EP N1->N3) — NOT via SFJ/SPJ.
# Tributary fills GSV trunk (N2) retrogradely.
# Trunk re-enters deep via proximal perforator (RP N2->N1).
#
# Rule 4: ep_n1_n3 + rp_n2_n1 (no max_visited gate).
# ---------------------------------------------------------------------------

def _type4_right(sid="s5_type4_r") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # Both junctions competent
        Step("P01  SFJ — both Valsalva AND Paranà negative: SFJ competent",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02  SFJ deeper",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move"),

        # Systematic scan of anterior thigh (no clips expected)
        Step("P03  Upper Thigh — antegrade flow only, SFJ competent",
             "probe_move", _pm(sid, "Upper Thigh", 0.12, "medial"),
             expected_action="move"),

        Step("P04  Upper Thigh mid",
             "probe_move", _pm(sid, "Upper Thigh", 0.17, "medial"),
             expected_action="move"),

        Step("P05  Dodd proximal",
             "probe_move", _pm(sid, "Dodd", 0.24, "medial"),
             expected_action="move"),

        Step("P06  Dodd mid — retrograde flow in trunk detected here",
             "probe_move", _pm(sid, "Dodd", 0.29, "medial"),
             expected_action="move"),

        Step("P07  Dodd distal",
             "probe_move", _pm(sid, "Dodd", 0.33, "medial"),
             expected_action="move"),

        Step("P08  Hunterian proximal — trunk fills retrogradely (from below)",
             "probe_move", _pm(sid, "Hunterian", 0.37, "medial"),
             expected_action="move"),

        Step("P09  Hunterian mid",
             "probe_move", _pm(sid, "Hunterian", 0.42, "medial"),
             expected_action="move"),

        Step("P10  Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.46, "medial"),
             expected_action="move"),

        # SPJ: competent
        Step("P11  SPJ — Paranà: only diastolic inflow, no SPJ entry",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P12  SPJ mid",
             "probe_move", _pm(sid, "SPJ", 0.53, "posterior"),
             expected_action="move"),

        # Calf: systematic search for entry perforator
        Step("P13  SSV upper — posterior view, no SPJ entry",
             "probe_move", _pm(sid, "SSV", 0.58, "posterior"),
             expected_action="move"),

        Step("P14  Calf medial upper — scanning for pathological perforator",
             "probe_move", _pm(sid, "Calf", 0.62, "medial"),
             expected_action="move"),

        Step("P15  Calf medial mid",
             "probe_move", _pm(sid, "Calf", 0.67, "medial"),
             expected_action="move"),

        Step("P16  Calf medial lower",
             "probe_move", _pm(sid, "Calf", 0.72, "medial"),
             expected_action="move"),

        Step("P17  Calf medial — Valsalva: outward flow at this perforator",
             "probe_move", _pm(sid, "Calf", 0.77, "medial"),
             expected_action="move"),

        Step("P18  Ankle medial",
             "probe_move", _pm(sid, "Ankle", 0.86, "medial"),
             expected_action="move"),

        Step("P19  Ankle distal",
             "probe_move", _pm(sid, "Ankle", 0.92, "medial"),
             expected_action="move"),

        # Return sweep — pinpointing EP N1->N3
        Step("P20  Return: lower calf",
             "probe_move", _pm(sid, "Calf", 0.85, "medial"),
             expected_action="move"),

        Step("P21  Return: calf lower",
             "probe_move", _pm(sid, "Calf", 0.79, "medial"),
             expected_action="move"),

        # EP N1->N3: perforator exits deep directly into N3 tributary (not SFJ/SPJ).
        Step("CM-1  EP N1->N3 at calf (posY=0.73) — perforator exits deep into N3 tributary (Type IV entry)",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.73, "Calf", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Trace tributary proximally to where it fills trunk retrogradely
        Step("P22  Return: calf mid",
             "probe_move", _pm(sid, "Calf", 0.72, "medial"),
             expected_action="move"),

        Step("P23  Return: calf upper — N3 joins N2 trunk here",
             "probe_move", _pm(sid, "Calf", 0.65, "medial"),
             expected_action="move"),

        Step("P24  Return: SPJ zone",
             "probe_move", _pm(sid, "SPJ", 0.52, "posterior"),
             expected_action="move"),

        Step("P25  Return: Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.44, "medial"),
             expected_action="move"),

        Step("P26  Return: Hunterian mid — trunk filling retrograde from N3 junction",
             "probe_move", _pm(sid, "Hunterian", 0.38, "medial"),
             expected_action="move"),

        Step("P27  Return: Dodd distal",
             "probe_move", _pm(sid, "Dodd", 0.31, "medial"),
             expected_action="move"),

        Step("P28  Return: Dodd mid",
             "probe_move", _pm(sid, "Dodd", 0.25, "medial"),
             expected_action="move"),

        # RP N2->N1: trunk re-enters deep via proximal Hunterian perforator.
        # R1->N3->N2->N1 path complete. Rule 4 fires immediately.
        Step("CM-2  RP N2->N1 at Dodd (posY=0.24) — trunk re-enters deep — COMPLETE (Type IV)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.24, "Dodd", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P29  Verify complete — Dodd",
             "probe_move", _pm(sid, "Dodd", 0.25, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P30  Verify complete — calf",
             "probe_move", _pm(sid, "Calf", 0.65, "medial"),
             expected_action="complete"),

        Step("P31  Verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 6 — mendoza_type5  (Type V: R1→R3→R2→R3→R1)  37 steps
# ---------------------------------------------------------------------------
# Source perforator exits deep into tributary: EP N1->N3.
# Tributary drains into saphenous trunk: RP N3->N2.
# Trunk escapes into a SECOND tributary: EP N2->N3.
# Second tributary re-enters deep: RP N3->N1.
# No RP N2->N1.  Both junctions competent.
#
# Rule 5b: ep_n1_n3 + rp_n3_n2 + ep_n2_n3 + rp_n3_n1
# ---------------------------------------------------------------------------

def _type5_right(sid="s6_type5_r") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # Both junctions competent
        Step("P01  SFJ — competent",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02  SFJ deeper",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move"),

        Step("P03  Upper Thigh — antegrade trunk flow",
             "probe_move", _pm(sid, "Upper Thigh", 0.13, "medial"),
             expected_action="move"),

        Step("P04  Upper Thigh mid",
             "probe_move", _pm(sid, "Upper Thigh", 0.18, "medial"),
             expected_action="move"),

        Step("P05  Dodd proximal",
             "probe_move", _pm(sid, "Dodd", 0.23, "medial"),
             expected_action="move"),

        Step("P06  Dodd mid — retrograde fill detected in trunk",
             "probe_move", _pm(sid, "Dodd", 0.28, "medial"),
             expected_action="move"),

        Step("P07  Dodd distal",
             "probe_move", _pm(sid, "Dodd", 0.33, "medial"),
             expected_action="move"),

        Step("P08  Hunterian proximal",
             "probe_move", _pm(sid, "Hunterian", 0.37, "medial"),
             expected_action="move"),

        Step("P09  Hunterian mid",
             "probe_move", _pm(sid, "Hunterian", 0.42, "medial"),
             expected_action="move"),

        Step("P10  Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.46, "medial"),
             expected_action="move"),

        Step("P11  SPJ — competent",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P12  SPJ mid",
             "probe_move", _pm(sid, "SPJ", 0.54, "posterior"),
             expected_action="move"),

        # Calf: searching for source perforator
        Step("P13  SSV upper",
             "probe_move", _pm(sid, "SSV", 0.58, "posterior"),
             expected_action="move"),

        Step("P14  Calf medial upper",
             "probe_move", _pm(sid, "Calf", 0.62, "medial"),
             expected_action="move"),

        Step("P15  Calf medial — Paranà: outward flow at Boyd perforator",
             "probe_move", _pm(sid, "Calf", 0.67, "medial"),
             expected_action="move"),

        # EP N1->N3: source perforator exits deep into N3 tributary.
        Step("CM-1  EP N1->N3 at calf (posY=0.65) — source perforator into N3 tributary",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.65, "Calf", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Trace N3 tributary proximally to where it connects to trunk
        Step("P16  Calf upper — tracing N3 tributary proximally",
             "probe_move", _pm(sid, "Calf", 0.60, "medial"),
             expected_action="move"),

        Step("P17  SPJ zone — following N3 into popliteal region",
             "probe_move", _pm(sid, "SPJ", 0.53, "posterior"),
             expected_action="move"),

        Step("P18  Hunterian distal — N3 approaches trunk here",
             "probe_move", _pm(sid, "Hunterian", 0.45, "medial"),
             expected_action="move"),

        Step("P19  Hunterian mid — N3 joins N2 trunk at junction",
             "probe_move", _pm(sid, "Hunterian", 0.39, "medial"),
             expected_action="move"),

        Step("P20  Hunterian proximal",
             "probe_move", _pm(sid, "Hunterian", 0.34, "medial"),
             expected_action="move"),

        Step("P21  Dodd distal — N3 tributary-to-trunk junction point",
             "probe_move", _pm(sid, "Dodd", 0.31, "medial"),
             expected_action="move"),

        # RP N3->N2: N3 tributary drains INTO N2 trunk.
        # ep_n1_n3 + rp_n3_n2 now present. rp_n3_n2 present -> Rule 5 (Type VI) blocked.
        Step("CM-2  RP N3->N2 at Dodd (posY=0.29) — N3 drains into N2 trunk",
             "clip_mark", _cm(sid, "RP", "N3", "N2", 0.29, "Dodd", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Follow trunk distally from N3 junction to find second escape (EP N2->N3)
        Step("P22  Dodd mid — trunk fills distally from N3 junction",
             "probe_move", _pm(sid, "Dodd", 0.28, "medial"),
             expected_action="move"),

        Step("P23  Hunterian — trunk sweeping distally",
             "probe_move", _pm(sid, "Hunterian", 0.38, "medial"),
             expected_action="move"),

        Step("P24  Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.44, "medial"),
             expected_action="move"),

        Step("P25  SPJ zone — trunk continues into popliteal",
             "probe_move", _pm(sid, "SPJ", 0.50, "posterior"),
             expected_action="move"),

        Step("P26  Calf upper medial — trunk escapes here into second N3",
             "probe_move", _pm(sid, "Calf", 0.60, "medial"),
             expected_action="move"),

        Step("P27  Calf medial mid",
             "probe_move", _pm(sid, "Calf", 0.65, "medial"),
             expected_action="move"),

        # EP N2->N3: trunk escapes into SECOND tributary.
        Step("CM-3  EP N2->N3 at calf (posY=0.68) — trunk escapes into second N3 tributary",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.68, "Calf", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Follow second N3 tributary to its re-entry perforator
        Step("P28  Calf lower — tracing second N3 distally",
             "probe_move", _pm(sid, "Calf", 0.74, "medial"),
             expected_action="move"),

        Step("P29  Calf lower — Paranà: diastolic inward flow at paratibial perforator",
             "probe_move", _pm(sid, "Calf", 0.80, "medial"),
             expected_action="move"),

        Step("P30  Ankle approach",
             "probe_move", _pm(sid, "Ankle", 0.89, "medial"),
             expected_action="move"),

        Step("P31  Ankle mid",
             "probe_move", _pm(sid, "Ankle", 0.93, "medial"),
             expected_action="move"),

        Step("P32  Return: ankle to lower calf",
             "probe_move", _pm(sid, "Calf", 0.87, "medial"),
             expected_action="move"),

        # RP N3->N1: second N3 tributary re-enters deep.
        # Rule 5b fires: ep_n1_n3 + rp_n3_n2 + ep_n2_n3 + rp_n3_n1. NO rp_n2_n1.
        Step("CM-4  RP N3->N1 at calf (posY=0.83) — second N3 re-enters deep — COMPLETE (Type V)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.83, "Calf", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P33  Verify complete — calf",
             "probe_move", _pm(sid, "Calf", 0.67, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P34  Verify complete — Dodd",
             "probe_move", _pm(sid, "Dodd", 0.29, "medial"),
             expected_action="complete"),

        Step("P35  Verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 7 — mendoza_type6  (Type VI: R1→R3→R1)  32 steps
# ---------------------------------------------------------------------------
# Perforator exits deep into N3 tributary (EP N1->N3).
# Blood travels through N3 ONLY — saphenous trunk (N2) NOT involved.
# Second perforator re-enters deep (RP N3->N1).  No RP N2->N1; no RP N3->N2.
#
# Rule 5: ep_n1_n3 + rp_n3_n1 + NOT rp_n2_n1 + NOT rp_n3_n2
# ---------------------------------------------------------------------------

def _type6_right(sid="s7_type6_r") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # Both junctions competent
        Step("P01  SFJ — competent",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02  SFJ deeper",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move"),

        # Full anterior thigh scan (trunk antegrade throughout)
        Step("P03  Upper Thigh",
             "probe_move", _pm(sid, "Upper Thigh", 0.12, "medial"),
             expected_action="move"),

        Step("P04  Upper Thigh mid",
             "probe_move", _pm(sid, "Upper Thigh", 0.18, "medial"),
             expected_action="move"),

        Step("P05  Dodd proximal",
             "probe_move", _pm(sid, "Dodd", 0.23, "medial"),
             expected_action="move"),

        Step("P06  Dodd mid — GSV trunk: antegrade only",
             "probe_move", _pm(sid, "Dodd", 0.28, "medial"),
             expected_action="move"),

        Step("P07  Dodd distal",
             "probe_move", _pm(sid, "Dodd", 0.33, "medial"),
             expected_action="move"),

        Step("P08  Hunterian proximal",
             "probe_move", _pm(sid, "Hunterian", 0.36, "medial"),
             expected_action="move"),

        Step("P09  Hunterian mid",
             "probe_move", _pm(sid, "Hunterian", 0.41, "medial"),
             expected_action="move"),

        Step("P10  Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.46, "medial"),
             expected_action="move"),

        Step("P11  SPJ — competent",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P12  SPJ mid",
             "probe_move", _pm(sid, "SPJ", 0.53, "posterior"),
             expected_action="move"),

        # Calf lateral — searching for source perforator (lateral territory)
        Step("P13  SSV upper",
             "probe_move", _pm(sid, "SSV", 0.58, "posterior"),
             expected_action="move"),

        Step("P14  Calf lateral upper",
             "probe_move", _pm(sid, "Calf", 0.63, "lateral"),
             expected_action="move"),

        Step("P15  Calf lateral mid",
             "probe_move", _pm(sid, "Calf", 0.68, "lateral"),
             expected_action="move"),

        Step("P16  Calf lateral lower",
             "probe_move", _pm(sid, "Calf", 0.73, "lateral"),
             expected_action="move"),

        Step("P17  Calf lateral — Valsalva: outward flow ≥500 ms, diameter 4 mm",
             "probe_move", _pm(sid, "Calf", 0.78, "lateral"),
             expected_action="move"),

        Step("P18  Ankle lateral",
             "probe_move", _pm(sid, "Ankle", 0.86, "lateral"),
             expected_action="move"),

        Step("P19  Ankle lateral distal",
             "probe_move", _pm(sid, "Ankle", 0.92, "lateral"),
             expected_action="move"),

        # Return sweep to pinpoint source perforator
        Step("P20  Return: lower calf lateral",
             "probe_move", _pm(sid, "Calf", 0.87, "lateral"),
             expected_action="move"),

        Step("P21  Return: calf lateral",
             "probe_move", _pm(sid, "Calf", 0.81, "lateral"),
             expected_action="move"),

        # EP N1->N3: source perforator exits deep into N3 lateral tributary.
        Step("CM-1  EP N1->N3 at ankle (posY=0.89) — perforator into N3 lateral tributary",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.89, "Ankle", "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        # Trace N3 proximally — confirm NO RP N3->N2 (trunk not involved)
        Step("P22  Return: lower calf lateral — tracing N3 proximally",
             "probe_move", _pm(sid, "Calf", 0.82, "lateral"),
             expected_action="move"),

        Step("P23  Calf lateral mid",
             "probe_move", _pm(sid, "Calf", 0.76, "lateral"),
             expected_action="move"),

        Step("P24  Calf lateral — checking GSV trunk: NO retrograde flow",
             "probe_move", _pm(sid, "Calf", 0.70, "medial"),
             expected_action="move"),

        Step("P25  Calf lateral — trunk excluded: pure N3 path confirmed",
             "probe_move", _pm(sid, "Calf", 0.70, "lateral"),
             expected_action="move"),

        Step("P26  Calf lateral upper — finding re-entry perforator",
             "probe_move", _pm(sid, "Calf", 0.64, "lateral"),
             expected_action="move"),

        Step("P27  Calf lateral — diastolic inward flow: re-entry perforator",
             "probe_move", _pm(sid, "Calf", 0.68, "lateral"),
             expected_action="move"),

        # RP N3->N1: N3 re-enters deep via second perforator.
        # NOT rp_n2_n1; NOT rp_n3_n2. Rule 5 fires.
        Step("CM-2  RP N3->N1 at calf lateral (posY=0.71) — N3 re-enters deep — COMPLETE (Type VI)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.71, "Calf", "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P28  Verify complete — calf lateral",
             "probe_move", _pm(sid, "Calf", 0.70, "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P29  Verify complete — SPJ",
             "probe_move", _pm(sid, "SPJ", 0.50, "posterior"),
             expected_action="complete"),

        Step("P30  Verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 8 — spj_type2a  (Type I via SPJ, classic "Type 2A")  30 steps
# ---------------------------------------------------------------------------
# SFJ competent.  SPJ terminal valve INCOMPETENT — deep blood enters SSV trunk
# (EP N1->N2 at popliteal).  SSV refluxes distally.  Re-enters deep via
# calf perforator from trunk (RP N2->N1).  No N3 escape.
#
# Rule 6: ep_n1_n2 + rp_n2_n1 + NOT ep_n2_n3 + old_max >= 0.48.
# SPJ mark posY=0.50 sets old_max to 0.50 before RP N2->N1.
# ---------------------------------------------------------------------------

def _spj_type2a(sid="s8_spj_2a") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # SFJ: competent
        Step("P01  SFJ — Valsalva + Paranà: SFJ competent",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02  SFJ deeper",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move"),

        # GSV thigh: antegrade throughout (no SFJ entry)
        Step("P03  Upper Thigh — antegrade GSV",
             "probe_move", _pm(sid, "Upper Thigh", 0.12, "medial"),
             expected_action="move"),

        Step("P04  Upper Thigh mid",
             "probe_move", _pm(sid, "Upper Thigh", 0.18, "medial"),
             expected_action="move"),

        Step("P05  Dodd proximal",
             "probe_move", _pm(sid, "Dodd", 0.24, "medial"),
             expected_action="move"),

        Step("P06  Dodd mid",
             "probe_move", _pm(sid, "Dodd", 0.29, "medial"),
             expected_action="move"),

        Step("P07  Hunterian",
             "probe_move", _pm(sid, "Hunterian", 0.36, "medial"),
             expected_action="move"),

        Step("P08  Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.43, "medial"),
             expected_action="move"),

        # Now posterior surface — checking for Giacomini vein on posterior thigh
        Step("P09  Giacomini — posterior thigh: no significant Giacomini vein",
             "probe_move", _pm(sid, "Giacomini", 0.30, "posterior"),
             expected_action="move"),

        # SPJ assessment: BOTH Paranà AND compression/relaxation positive
        Step("P10  SPJ approach — lateral decubitus position",
             "probe_move", _pm(sid, "SPJ", 0.46, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior", "junction"]),

        Step("P11  SPJ — Paranà positive AND CR positive: SPJ incompetent",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior"),
             expected_action="move"),

        Step("P12  SPJ mid — Valsalva: deep blood enters SSV trunk",
             "probe_move", _pm(sid, "SPJ", 0.52, "posterior"),
             expected_action="move"),

        # EP N1->N2: SPJ incompetent — deep blood enters SSV trunk at SPJ
        Step("CM-1  EP N1->N2 at SPJ (posY=0.50) — SPJ incompetent confirmed",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.50, "SPJ", "posterior"),
             expected_action="move",
             guidance_must_contain=["SSV", "calf", "posterior", "trunk"]),

        # SSV sweep distally — looking for re-entry (no N3 escape)
        Step("P13  SSV upper — tracing SSV trunk distally",
             "probe_move", _pm(sid, "SSV", 0.58, "posterior"),
             expected_action="move"),

        Step("P14  SSV mid",
             "probe_move", _pm(sid, "SSV", 0.63, "posterior"),
             expected_action="move"),

        Step("P15  SSV — retrograde trunk flow confirmed on Paranà",
             "probe_move", _pm(sid, "SSV", 0.68, "posterior"),
             expected_action="move"),

        Step("P16  SSV lower",
             "probe_move", _pm(sid, "SSV", 0.73, "posterior"),
             expected_action="move"),

        Step("P17  SSV lower — no N3 tributary escape",
             "probe_move", _pm(sid, "SSV", 0.78, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P18  SSV lower",
             "probe_move", _pm(sid, "SSV", 0.83, "posterior"),
             expected_action="move"),

        Step("P19  Ankle posterior",
             "probe_move", _pm(sid, "Ankle", 0.90, "posterior"),
             expected_action="move"),

        Step("P20  Ankle distal posterior",
             "probe_move", _pm(sid, "Ankle", 0.94, "posterior"),
             expected_action="move"),

        # Return sweep — identify perforator re-entry from SSV trunk
        Step("P21  Return: lower calf posterior",
             "probe_move", _pm(sid, "SSV", 0.87, "posterior"),
             expected_action="move"),

        Step("P22  Return: calf lower",
             "probe_move", _pm(sid, "SSV", 0.81, "posterior"),
             expected_action="move"),

        Step("P23  Return: calf mid — biphasic paratibial perforator",
             "probe_move", _pm(sid, "SSV", 0.75, "posterior"),
             expected_action="move"),

        Step("P24  Return: calf mid — diastolic inward flow confirmed",
             "probe_move", _pm(sid, "SSV", 0.70, "posterior"),
             expected_action="move"),

        # RP N2->N1: SSV trunk blood re-enters deep.
        # old_max_visited=0.94 >= 0.48; ep_n1_n2=T; ep_n2_n3=F -> Rule 6 fires.
        Step("CM-2  RP N2->N1 SSV (posY=0.67) — trunk re-enters deep — COMPLETE (SPJ Type 2A)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.67, "SSV", "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P25  Verify complete — SSV calf",
             "probe_move", _pm(sid, "SSV", 0.65, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P26  Verify complete — SPJ",
             "probe_move", _pm(sid, "SPJ", 0.50, "posterior"),
             expected_action="complete"),

        Step("P27  Verify complete — Hunterian",
             "probe_move", _pm(sid, "Hunterian", 0.40, "medial"),
             expected_action="complete"),

        Step("P28  Verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 9 — left_leg_type1  (Type I on left leg)  28 steps
# ---------------------------------------------------------------------------
# Identical CHIVA circuit to mendoza_type1 but on the LEFT leg.
# Tests that the guidance engine handles leg="left" identically.
# ---------------------------------------------------------------------------

def _left_leg_type1(sid="s9_left_t1") -> list[Step]:
    L = "left"
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01  LEFT SFJ — Mickey Mouse sign, femoral side",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial", L),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02  LEFT SFJ mid",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial", L),
             expected_action="move"),

        Step("CM-1  EP N1->N2 at LEFT SFJ (posY=0.06) — incompetent",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.06, "SFJ", "anterior-medial", leg=L),
             expected_action="move",
             guidance_must_contain=["distal", "trunk", "GSV", "thigh"]),

        Step("P03  LEFT Upper Thigh",
             "probe_move", _pm(sid, "Upper Thigh", 0.11, "medial", L),
             expected_action="move"),

        Step("P04  LEFT Upper Thigh mid",
             "probe_move", _pm(sid, "Upper Thigh", 0.16, "medial", L),
             expected_action="move"),

        Step("P05  LEFT Upper Thigh distal",
             "probe_move", _pm(sid, "Upper Thigh", 0.20, "medial", L),
             expected_action="move"),

        Step("P06  LEFT Dodd proximal",
             "probe_move", _pm(sid, "Dodd", 0.24, "medial", L),
             expected_action="move",
             guidance_must_contain=["Hunterian", "Dodd", "perforator", "escape", "tributary"]),

        Step("P07  LEFT Dodd mid — no escape tributary",
             "probe_move", _pm(sid, "Dodd", 0.29, "medial", L),
             expected_action="move",
             forbidden_action="complete"),

        Step("P08  LEFT Dodd distal",
             "probe_move", _pm(sid, "Dodd", 0.33, "medial", L),
             expected_action="move"),

        Step("P09  LEFT Hunterian proximal",
             "probe_move", _pm(sid, "Hunterian", 0.37, "medial", L),
             expected_action="move"),

        Step("P10  LEFT Hunterian mid",
             "probe_move", _pm(sid, "Hunterian", 0.42, "medial", L),
             expected_action="move",
             forbidden_action="complete"),

        Step("P11  LEFT Hunterian distal",
             "probe_move", _pm(sid, "Hunterian", 0.46, "medial", L),
             expected_action="move"),

        # max_visited now >= 0.46; SPJ will push it past 0.48 gate
        Step("P12  LEFT SPJ — lateral decubitus; SPJ competent",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior", L),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior", "junction"]),

        Step("P13  LEFT SPJ mid",
             "probe_move", _pm(sid, "SPJ", 0.53, "posterior", L),
             expected_action="move"),

        Step("P14  LEFT Calf medial upper",
             "probe_move", _pm(sid, "Calf", 0.60, "medial", L),
             expected_action="move"),

        Step("P15  LEFT Calf medial mid — retrograde trunk flow",
             "probe_move", _pm(sid, "Calf", 0.65, "medial", L),
             expected_action="move"),

        Step("P16  LEFT Calf medial lower",
             "probe_move", _pm(sid, "Calf", 0.71, "medial", L),
             expected_action="move"),

        Step("P17  LEFT Calf medial — no N3 escape at calf",
             "probe_move", _pm(sid, "Calf", 0.77, "medial", L),
             expected_action="move",
             forbidden_action="complete"),

        Step("P18  LEFT Calf lower",
             "probe_move", _pm(sid, "Calf", 0.83, "medial", L),
             expected_action="move"),

        Step("P19  LEFT Ankle medial",
             "probe_move", _pm(sid, "Ankle", 0.90, "medial", L),
             expected_action="move"),

        Step("P20  LEFT Ankle distal",
             "probe_move", _pm(sid, "Ankle", 0.94, "medial", L),
             expected_action="move"),

        # Return sweep
        Step("P21  Return LEFT: lower calf",
             "probe_move", _pm(sid, "Calf", 0.86, "medial", L),
             expected_action="move"),

        Step("P22  Return LEFT: calf mid",
             "probe_move", _pm(sid, "Calf", 0.79, "medial", L),
             expected_action="move"),

        Step("P23  Return LEFT: calf upper — scanning perforators",
             "probe_move", _pm(sid, "Calf", 0.73, "medial", L),
             expected_action="move"),

        Step("P24  Return LEFT: calf — diastolic inward perforator flow",
             "probe_move", _pm(sid, "Calf", 0.67, "medial", L),
             expected_action="move"),

        # RP N2->N1: trunk re-enters deep. old_max=0.94>=0.48. Rule 6 fires.
        Step("CM-2  RP N2->N1 LEFT calf (posY=0.69) — trunk re-enters deep — COMPLETE (Left Type I)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.69, "Calf", "medial", leg=L),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P25  Verify LEFT complete — calf",
             "probe_move", _pm(sid, "Calf", 0.67, "medial", L),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P26  Verify LEFT complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial", L),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

_SCENARIOS: dict[str, callable] = {
    "type1":         _type1_right,
    "type2":         _type2_right,
    "type3_noreflux": _type3_no_reflux,
    "type3_reflux":  _type3_reflux,
    "type4":         _type4_right,
    "type5":         _type5_right,
    "type6":         _type6_right,
    "spj_2a":        _spj_type2a,
    "left_type1":    _left_leg_type1,
}

_SCENARIO_DESCRIPTIONS: dict[str, str] = {
    "type1": (
        "Type I (R1→R2→R1): SFJ incompetent (EP N1->N2); GSV trunk reflux; "
        "no tributary escape; calf perforator re-entry (RP N2->N1). "
        "Rule 6: old_max ≥ 0.48.  35 steps."
    ),
    "type2": (
        "Type II (R2→R3→R1): Both junctions competent; trunk is source; "
        "Dodd escape to N3 (EP N2->N3); N3 re-enters deep (RP N3->N1). "
        "Rule 7: no EP N1->N2, no RP N2->N1.  33 steps."
    ),
    "type3_noreflux": (
        "Type III / No Reflux (R1→R2→R3→R1): SFJ entry; EP N2->N3 before RP N2->N1; "
        "RP N3->N1 + RP N2->N1 => maneuver; compress tributary => 'No Reflux' => complete. "
        "Rules 1 then 2.  37 steps."
    ),
    "type3_reflux": (
        "Type I+2 / Reflux (R1→R2→R3→R1): Same path as Type III but compression shows "
        "'Reflux' => trunk has independent reflux => Type I+2 complete. "
        "Rules 1 then 3.  30 steps."
    ),
    "type4": (
        "Type IV (R1→R3→R2→R1): Both junctions competent; perforator exits direct into "
        "N3 (EP N1->N3); N3 fills trunk; trunk re-enters deep (RP N2->N1). "
        "Rule 4 (no max_visited gate).  33 steps."
    ),
    "type5": (
        "Type V (R1→R3→R2→R3→R1): EP N1->N3 source; RP N3->N2 trunk fill; "
        "EP N2->N3 second escape; RP N3->N1 second re-entry. "
        "Rule 5b.  37 steps."
    ),
    "type6": (
        "Type VI (R1→R3→R1): Both junctions competent; purely N3 lateral path; "
        "EP N1->N3 source perforator; RP N3->N1 re-entry; trunk NOT involved. "
        "Rule 5.  32 steps."
    ),
    "spj_2a": (
        "SPJ Type 2A: SFJ competent; SPJ incompetent (EP N1->N2 at posY=0.50); "
        "SSV trunk reflux; calf perforator re-entry (RP N2->N1). "
        "Rule 6 via SPJ.  30 steps."
    ),
    "left_type1": (
        "Left Leg Type I: Identical circuit to Type I but leg='left'. "
        "Tests left-leg handling.  Rule 6.  28 steps."
    ),
}


# ---------------------------------------------------------------------------
# Runner utilities
# ---------------------------------------------------------------------------

def _describe_movement(event: str, data: dict) -> str:
    if event == "stream_start":
        return f"Start session  (id: {data.get('session_id', '?')})"
    if event == "probe_move":
        return (f"Move -> {data.get('region', '?')}\n"
                f"posY={data.get('pos_y_ratio', 0.0):.2f}  "
                f"{data.get('surface', '?')}  {data.get('leg', 'right')} leg")
    if event == "clip_mark":
        base = (f"Mark: {data.get('flow', '?')} "
                f"{data.get('from_type', '?')}->{data.get('to_type', '?')}  "
                f"posY={data.get('pos_y_ratio', 0.0):.2f}  "
                f"{data.get('region', '?')}  {data.get('leg', 'right')} leg")
        if data.get("elimination_test"):
            base += f"\nElim: {data['elimination_test']}"
        return base
    return event


def _describe_expectation(step: Step) -> str:
    if step.event == "stream_start":
        return "session_ready event"
    parts = []
    if step.expected_action:
        parts.append(f'Action = "{step.expected_action}"')
    if step.forbidden_action:
        parts.append(f'Must NOT = "{step.forbidden_action}"')
    if step.guidance_must_contain:
        parts.append("Guidance contains any of: " + ", ".join(step.guidance_must_contain))
    return "\n".join(parts) if parts else "observe only"


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

CYAN   = "\033[96m"
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

_SKIP_EVENTS = {"stream_start"}


class ScenarioRunner:
    def __init__(self, api_base: str):
        self.api_base = api_base.rstrip("/")
        self._sio = socketio.SimpleClient()

    def connect(self) -> None:
        print(f"Connecting to {self.api_base} ...")
        self._sio.connect(self.api_base, wait_timeout=10)
        print("Connected.\n")

    def disconnect(self) -> None:
        self._sio.disconnect()

    def _emit_and_wait(self, event: str, data: dict, timeout: float = 60.0):
        self._sio.emit(event, data)
        if event in _SKIP_EVENTS:
            deadline = time.time() + timeout
            while time.time() < deadline:
                try:
                    ev = self._sio.receive(timeout=2)
                    if ev and ev[0] == "session_ready":
                        return ev[1]
                except Exception:
                    pass
            return None
        deadline = time.time() + timeout
        while time.time() < deadline:
            try:
                ev = self._sio.receive(timeout=2)
                if ev and ev[0] == "guidance_update":
                    return ev[1]
            except Exception:
                pass
        return None

    def run(self, steps: list[Step]) -> list[dict]:
        results = []
        for step in steps:
            print(f"\n{CYAN}--- {step.label} ---{RESET}")

            movement    = _describe_movement(step.event, step.data)
            expectation = _describe_expectation(step)
            response    = self._emit_and_wait(step.event, step.data)

            if step.event in _SKIP_EVENTS:
                status = "SKIP" if response else "TIMEOUT"
                if response:
                    print(f"  session_ready: {response}")
                else:
                    print(f"  {YELLOW}No session_ready{RESET}")
                results.append({"step": step.label, "status": status,
                                "movement": movement, "expectation": expectation,
                                "llm_action": "--", "llm_guidance": "--",
                                "failures": []})
                continue

            if (step.expected_action is None
                    and not step.guidance_must_contain
                    and step.forbidden_action is None):
                if response:
                    g = response.get("guidance", "")
                    a = response.get("action", "?")
                    print(f"  guidance: {g!r}  action: {a}")
                    results.append({"step": step.label, "status": "SKIP",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": a, "llm_guidance": g, "failures": []})
                else:
                    print(f"  {YELLOW}timeout{RESET}")
                    results.append({"step": step.label, "status": "TIMEOUT",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                continue

            if response is None:
                print(f"  {YELLOW}No response (timeout){RESET}")
                results.append({"step": step.label, "status": "TIMEOUT",
                                "movement": movement, "expectation": expectation,
                                "llm_action": "--", "llm_guidance": "--",
                                "failures": []})
                continue

            llm_guidance = response.get("guidance", "") or ""
            llm_action   = response.get("action", "move") or "move"
            print(f"  guidance: {llm_guidance!r}")
            print(f"  action  : {llm_action}")

            passed   = True
            failures = []

            if step.expected_action and llm_action != step.expected_action:
                passed = False
                failures.append(
                    f'Action: got "{llm_action}", expected "{step.expected_action}"')

            if step.guidance_must_contain:
                if not any(kw.lower() in llm_guidance.lower()
                           for kw in step.guidance_must_contain):
                    passed = False
                    failures.append(
                        "Guidance missing any of: "
                        + ", ".join(step.guidance_must_contain))

            if step.forbidden_action and llm_action == step.forbidden_action:
                passed = False
                failures.append(f'Action "{llm_action}" must NOT fire here')

            status = "PASS" if passed else "FAIL"
            if passed:
                print(f"  {GREEN}{BOLD}PASS{RESET}")
            else:
                for f in failures:
                    print(f"  {RED}FAIL: {f}{RESET}")

            results.append({"step": step.label, "status": status,
                            "movement": movement, "expectation": expectation,
                            "llm_action": llm_action, "llm_guidance": llm_guidance,
                            "failures": failures})
            time.sleep(0.4)

        return results

    def print_summary(self, label: str, results: list[dict]) -> bool:
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        passed  = [r for r in checked if r["status"] == "PASS"]
        failed  = [r for r in checked if r["status"] == "FAIL"]
        timeout = [r for r in results  if r["status"] == "TIMEOUT"]
        print(f"\n{'─'*65}")
        print(f"{BOLD}{label}{RESET}  "
              f"{GREEN}{len(passed)} passed{RESET}  "
              f"{RED}{len(failed)} failed{RESET}  "
              f"{YELLOW}{len(timeout)} timeout{RESET}  "
              f"/ {len(checked)} checked")
        for r in failed:
            print(f"  {RED}x {r['step']}{RESET}")
            for f in r.get("failures", []):
                print(f"      {f}")
        return len(failed) == 0 and len(timeout) == 0


# ---------------------------------------------------------------------------
# Word report
# ---------------------------------------------------------------------------

_STATUS_COLORS = {"PASS": "00B050", "FAIL": "FF0000",
                  "TIMEOUT": "FF8C00", "SKIP": "808080"}


def _colored_cell(cell, text: str, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    from docx.shared import RGBColor
    cell.text = text
    run = (cell.paragraphs[0].runs[0]
           if cell.paragraphs[0].runs
           else cell.paragraphs[0].add_run(text))
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, pt: int = 9, bold=False,
                   italic=False, color_rgb=None) -> None:
    cell.text = ""
    for i, line in enumerate(text.split("\n")):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run  = para.add_run(line)
        run.font.size   = __import__("docx").shared.Pt(pt)
        run.font.bold   = bold
        run.font.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb


def write_word_report(all_results: dict[str, list[dict]],
                      api_base: str, out_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin  = sec.bottom_margin = Inches(0.7)

    t = doc.add_heading(
        "CHIVA Streaming Guidance — Consolidated Scenario Test Report", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = m.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}    |    "
                   f"API: {api_base}    |    Scenarios: {len(all_results)}")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    # Summary table
    doc.add_heading("Overall Summary", level=1)
    stbl = doc.add_table(rows=1, cols=5)
    stbl.style = "Table Grid"
    for i, h in enumerate(["Scenario", "Description", "Passed", "Failed", "Timeouts"]):
        c = stbl.rows[0].cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True

    overall_pass = True
    for name, results in all_results.items():
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")
        row = stbl.add_row().cells
        row[0].text = name
        row[1].text = _SCENARIO_DESCRIPTIONS.get(name, "")
        row[2].text = str(p)
        row[3].text = str(f)
        row[4].text = str(t)
        if f or t:
            overall_pass = False

    doc.add_paragraph()
    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run("OVERALL: " + ("ALL PASS" if overall_pass else "FAILURES PRESENT"))
    vr.font.bold = True
    vr.font.size = Pt(13)
    vr.font.color.rgb = (RGBColor(0x00, 0xB0, 0x50) if overall_pass
                         else RGBColor(0xFF, 0x00, 0x00))

    # Per-scenario detail pages
    for name, results in all_results.items():
        doc.add_page_break()
        doc.add_heading(f"Scenario: {name.upper()}", level=1)
        dp = doc.add_paragraph(_SCENARIO_DESCRIPTIONS.get(name, ""))
        if dp.runs:
            dp.runs[0].font.italic = True

        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")

        sl = doc.add_paragraph()
        sl.add_run(f"Checked: {len(checked)}   |   ")
        pr = sl.add_run(f"Passed: {p}")
        pr.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
        sl.add_run("   |   ")
        fr = sl.add_run(f"Failed: {f}")
        fr.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        sl.add_run("   |   ")
        tr = sl.add_run(f"Timeouts: {t}")
        tr.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        doc.add_paragraph()

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        for i, h in enumerate(["#", "Surgeon Action", "Expected",
                                "LLM Guidance", "Action", "Status"]):
            c = tbl.rows[0].cells[i]
            c.text = h
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.bold = True
        for i, w in enumerate([0.25, 1.80, 2.10, 2.10, 0.80, 0.65]):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)

        for idx, r in enumerate(results, 1):
            status = r.get("status", "SKIP")
            row    = tbl.add_row().cells
            row[0].text = str(idx)
            _set_cell_text(row[1], r.get("movement", "--"), pt=8)
            _set_cell_text(row[2], r.get("expectation", "--"), pt=8, italic=True)
            _set_cell_text(row[3], r.get("llm_guidance", "--") or "--", pt=8)
            _set_cell_text(row[4], r.get("llm_action", "--") or "--", pt=9, bold=True)
            _colored_cell(row[5], status, _STATUS_COLORS.get(status, "808080"))
            if status == "FAIL" and r.get("failures"):
                fr_row = tbl.add_row().cells
                merged = fr_row[0].merge(fr_row[5])
                merged.text = ""
                for fn in r["failures"]:
                    pp = merged.add_paragraph(f"  x  {fn}")
                    if pp.runs:
                        pp.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        pp.runs[0].font.size      = Pt(8)
                        pp.runs[0].font.italic    = True

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"\nReport saved: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="CHIVA streaming guidance — consolidated scenario runner")
    parser.add_argument("scenario", nargs="?", default="all",
                        choices=list(_SCENARIOS.keys()) + ["all"],
                        help="Scenario ID or 'all' (default: all)")
    parser.add_argument("--api", default="http://localhost:7861")
    parser.add_argument("--all", action="store_true", dest="run_all")
    args = parser.parse_args()

    to_run = list(_SCENARIOS.items()) if (args.run_all or args.scenario == "all") \
             else [(args.scenario, _SCENARIOS[args.scenario])]

    runner     = ScenarioRunner(args.api)
    all_ok     = True
    all_results: dict[str, list[dict]] = {}

    try:
        runner.connect()
        for name, builder in to_run:
            print(f"\n{'='*65}")
            print(f"{BOLD}SCENARIO: {name.upper()}{RESET}")
            desc = _SCENARIO_DESCRIPTIONS.get(name, "")
            if desc:
                print(f"  {desc}")
            print(f"{'='*65}")
            steps   = builder()
            results = runner.run(steps)
            ok      = runner.print_summary(name, results)
            all_ok  = all_ok and ok
            all_results[name] = results
    except KeyboardInterrupt:
        print("\nInterrupted.")
    finally:
        runner.disconnect()

    if len(to_run) > 1:
        print(f"\n{'='*65}")
        print(f"{BOLD}OVERALL: {'ALL PASS' if all_ok else 'FAILURES PRESENT'}{RESET}")

    if all_results:
        ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
        here     = os.path.dirname(os.path.abspath(__file__))
        out_path = os.path.join(here, "results", f"scenarios_{ts}.docx")
        try:
            write_word_report(all_results, args.api, out_path)
        except Exception as exc:
            print(f"{YELLOW}Warning: could not write Word report — {exc}{RESET}")

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
