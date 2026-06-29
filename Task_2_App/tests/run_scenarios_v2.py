"""
CHIVA Streaming Guidance — Extended Scenario Set (v2)

5 new clinically-grounded scenarios covering:
  1. hunterian_perf_entry   — Hunterian perforator as sole EP N1→N2 (SFJ Valsalva-negative).
  2. aasv_valsalva_negative — Paranà positive + Valsalva negative at SFJ → SFJ competent;
                              AASV classified as N3 (not N2); EP N1→N3 source perf + RP N3→N1.
  3. ssv_type2_autonomous   — SSV pre-terminal autonomous reflux (SPJ one-maneuver ≠ incompetent);
                              EP N2→N3 from SSV trunk + RP N3→N1 → Type II (Rule 7).
  4. spj_type3_noreflux     — SPJ entry (EP N1→N2 at SPJ); SSV escape (EP N2→N3) marked before
                              RP N2→N1; maneuver fires; compress → No Reflux → Type III (Rules 1→2).
  5. rule6_max_visited_gate — RP N2→N1 marked at upper thigh (posY 0.15) before probe ever visits
                              posY ≥ 0.48; confirms Rule 6 gate holds; complete fires only after
                              SPJ zone is covered.

Literature basis:
  - Franceschi & Zamboni 2009, "Principles of Venous Hemodynamics", ch. 9 (pp. 111–126).
  - Gianesini et al. 2014, Phlebology (CHIVA strategy), pp. 11–15.
  - Delfrate 2023, JTAVR 8(2), pp. 21–26.
  - Adler et al. 2022, RadioGraphics 42:2184–2200, pp. 2190–2193.

Region / posY mapping (must match stream.html mapReg()):
  SFJ          posY 0.00–0.07   surface "anterior-medial"
  Upper Thigh  posY 0.08–0.20   surface "medial"
  Dodd         posY 0.21–0.33   surface "medial"
  Hunterian    posY 0.34–0.47   surface "medial"
  Giacomini    posY 0.08–0.47   surface "posterior"
  SPJ          posY 0.48–0.57   surface "posterior"
  Calf         posY 0.58–0.88   surface "medial" / "lateral"
  SSV          posY 0.58–0.88   surface "posterior"
  Ankle        posY 0.89–1.00   surface "medial" / "lateral" / "posterior"

Engine rules relevant to these scenarios:
  Rule 1  maneuver : ep_n2_n3 (no elimTest) + rp_n3_n1 + rp_n2_n1
  Rule 2  complete : elimTest on EP N2→N3 == "No Reflux"                  (Type III)
  Rule 5  complete : ep_n1_n3 + rp_n3_n1 + NOT rp_n2_n1 + NOT rp_n3_n2   (Type VI)
  Rule 6  complete : ep_n1_n2 + rp_n2_n1 + NOT ep_n2_n3 + old_max ≥ 0.48 (Type I / 2A)
  Rule 7  complete : ep_n2_n3 + rp_n3_n1 + NOT ep_n1_n2 + NOT ep_n1_n3 + NOT rp_n2_n1  (Type II)

Usage:
    python tests/run_scenarios_v2.py [--api http://localhost:7861] [<scenario>|all]

Output: tests/results/scenarios_v2_<timestamp>.docx
"""
from __future__ import annotations

import argparse
import io
import os
import sys
import time

if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

try:
    import socketio
except ImportError:
    print("ERROR: python-socketio not installed.  Run: pip install python-socketio[client]")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Step definition  (identical to run_scenarios.py)
# ---------------------------------------------------------------------------

@dataclass
class Step:
    label: str
    event: str
    data: dict
    expected_action: Optional[str] = None
    guidance_must_contain: list[str] = field(default_factory=list)
    forbidden_action: Optional[str] = None


def _pm(sid, region, pos_y, surface, leg="right") -> dict:
    return {"session_id": sid, "region": region, "pos_y_ratio": pos_y,
            "surface": surface, "leg": leg}

def _cm(sid, flow, ft, tt, pos_y, region, surface, elim="", leg="right") -> dict:
    return {"session_id": sid, "flow": flow, "from_type": ft, "to_type": tt,
            "pos_y_ratio": pos_y, "leg": leg, "region": region,
            "surface": surface, "elimination_test": elim}


# ---------------------------------------------------------------------------
# Scenario 10 — hunterian_perf_entry  (Type I via Hunterian perforator EP)
# ---------------------------------------------------------------------------
# Source: Gianesini et al. 2014, p.11:
#   "N1-N2 compartment jumps arising not from an incompetent sapheno-femoral
#    junction but rather by an incompetent Hunterian perforating vein."
#
# Clinical picture:
#   SFJ assessed — Paranà positive (pre-terminal reflux present) but Valsalva
#   NEGATIVE.  Per Delfrate 2023, p.24: "N1→N2 reflux at the GSV end assessed
#   by Paraná/squeezing is not specific of the incompetence of the terminal
#   valve and must be every time confirmed by positive Valsalva."  Therefore
#   SFJ TERMINAL VALVE IS COMPETENT — NO EP N1→N2 marked at SFJ.
#
#   GSV thigh shows retrograde flow whose source is not the SFJ.  Systematic
#   scan reaches Hunterian zone (posY 0.34–0.47): outward perforator flow
#   ≥500 ms, diameter ≥3.5 mm (AVF 2023 criteria) → EP N1→N2 at Hunterian.
#   SPJ competent.  Probe sweeps calf to ankle (max_visited ≥ 0.48).
#   RP N2→N1 at calf → Rule 6 complete (Type I equivalent).
#
# Rule 6: ep_n1_n2 + rp_n2_n1 + NOT ep_n2_n3 + old_max ≥ 0.48.
# ---------------------------------------------------------------------------

def _hunterian_perf_entry(sid="s10_hunt_ep") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # --- SFJ: Paranà positive but Valsalva NEGATIVE → COMPETENT ---
        # Per Delfrate 2023 p.24: concordance of BOTH maneuvers required.
        Step("P01  SFJ — transverse B-mode: Mickey Mouse sign confirmed",
             "probe_move", _pm(sid, "SFJ", 0.04, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction", "saphenofemoral"]),

        Step("P02  SFJ — Doppler on femoral side of terminal valve",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move"),

        # Valsalva: NEGATIVE (no reflux through terminal valve)
        # Paranà: POSITIVE (pre-terminal reflux detected below competent terminal valve)
        # Conclusion: terminal valve COMPETENT — do NOT mark EP N1→N2 here.
        Step("P03  SFJ — Valsalva negative; Paranà positive pre-terminal only → SFJ COMPETENT",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move",
             forbidden_action="complete"),

        # --- Upper thigh: GSV shows retrograde flow despite competent SFJ ---
        # Source must be a perforator further distal (Hunterian zone).
        Step("P04  Upper Thigh — saphenous eye confirmed; retrograde trunk flow detected",
             "probe_move", _pm(sid, "Upper Thigh", 0.10, "medial"),
             expected_action="move",
             guidance_must_contain=["thigh", "trunk", "reflux", "perforator"]),

        Step("P05  Upper Thigh mid — SFJ competent yet retrograde flow present → source distal",
             "probe_move", _pm(sid, "Upper Thigh", 0.15, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P06  Upper Thigh distal — no escape into N3 tributary here",
             "probe_move", _pm(sid, "Upper Thigh", 0.19, "medial"),
             expected_action="move"),

        # --- Dodd zone: scanning medial mid-thigh ---
        Step("P07  Dodd proximal — continuing distal search for source perforator",
             "probe_move", _pm(sid, "Dodd", 0.23, "medial"),
             expected_action="move",
             guidance_must_contain=["Hunterian", "Dodd", "perforator", "escape", "tributary"]),

        Step("P08  Dodd mid — trunk retrograde; no N3 escape at Dodd level",
             "probe_move", _pm(sid, "Dodd", 0.28, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P09  Dodd distal",
             "probe_move", _pm(sid, "Dodd", 0.32, "medial"),
             expected_action="move"),

        # --- Hunterian zone: source perforator identified ---
        # Gianesini 2014 p.11: incompetent Hunterian perforating vein as N1→N2 entry.
        # Delfrate 2023 p.25 (AVF 2023 criteria): outward flow ≥500 ms, diameter ≥3.5 mm.
        # Three maneuvers applied: squeezing, Paranà, Valsalva — all confirm outward N1→N2.
        Step("P10  Hunterian proximal — approaching Hunterian perforator zone",
             "probe_move", _pm(sid, "Hunterian", 0.36, "medial"),
             expected_action="move"),

        Step("P11  Hunterian mid — Paranà: outward perforator flow >500 ms; diam 4.2 mm",
             "probe_move", _pm(sid, "Hunterian", 0.39, "medial"),
             expected_action="move"),

        Step("P12  Hunterian mid-lower — Valsalva confirms N1→N2 outward flow at perforator",
             "probe_move", _pm(sid, "Hunterian", 0.43, "medial"),
             expected_action="move"),

        # EP N1→N2 at Hunterian: deep blood exits into GSV trunk via Hunterian perforator.
        # No EP N1→N2 at SFJ. Rule 6 requires old_max ≥ 0.48 still to be met.
        Step("CM-1  EP N1→N2 at Hunterian (posY=0.41) — Hunterian perf exits deep into GSV trunk",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.41, "Hunterian", "medial"),
             expected_action="move",
             guidance_must_contain=["distal", "calf", "trunk", "reflux"]),

        Step("P13  Hunterian distal — tracing trunk distally past entry point",
             "probe_move", _pm(sid, "Hunterian", 0.46, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # --- SPJ: must assess to confirm no second EP N1→N2 via SSV ---
        # Adler et al. 2022 p.2190: augmentation required at SPJ when SFJ competent.
        # Gianesini 2014 p.14: both Paranà AND CR required for SPJ incompetence.
        Step("P14  SPJ — lateral decubitus; Paranà + CR: SPJ COMPETENT (one maneuver only positive)",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior", "junction"]),

        Step("P15  SPJ mid — confirming SPJ competence; no N1→N2 entry here",
             "probe_move", _pm(sid, "SPJ", 0.53, "posterior"),
             expected_action="move"),

        # old_max now 0.53 ≥ 0.48 — gate condition is met once RP N2→N1 is marked.

        # --- Calf: GSV trunk continues retrograde from Hunterian perforator source ---
        Step("P16  SSV upper — checking posterior calf; no Giacomini involvement",
             "probe_move", _pm(sid, "SSV", 0.59, "posterior"),
             expected_action="move"),

        Step("P17  Calf medial upper — GSV trunk in saphenous eye; retrograde throughout",
             "probe_move", _pm(sid, "Calf", 0.62, "medial"),
             expected_action="move"),

        Step("P18  Calf medial mid — Paranà: sustained retrograde trunk flow >500 ms",
             "probe_move", _pm(sid, "Calf", 0.67, "medial"),
             expected_action="move"),

        Step("P19  Calf medial lower — no EP N2→N3 escape at any calf level",
             "probe_move", _pm(sid, "Calf", 0.72, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P20  Calf lower",
             "probe_move", _pm(sid, "Calf", 0.78, "medial"),
             expected_action="move"),

        # --- Ankle: completing distal coverage ---
        Step("P21  Ankle medial — GSV at medial malleolus; trunk continuity confirmed",
             "probe_move", _pm(sid, "Ankle", 0.90, "medial"),
             expected_action="move"),

        Step("P22  Ankle distal",
             "probe_move", _pm(sid, "Ankle", 0.95, "medial"),
             expected_action="move"),

        # --- Return sweep: identify calf re-entry perforator ---
        # Gianesini 2014 p.15: inward diastolic flow during Paranà release = RP.
        Step("P23  Return: ankle to lower calf",
             "probe_move", _pm(sid, "Calf", 0.87, "medial"),
             expected_action="move"),

        Step("P24  Return: calf lower",
             "probe_move", _pm(sid, "Calf", 0.80, "medial"),
             expected_action="move"),

        Step("P25  Return: calf mid — biphasic perforator: systolic outward + diastolic inward",
             "probe_move", _pm(sid, "Calf", 0.73, "medial"),
             expected_action="move"),

        # RP N2→N1: GSV trunk re-enters deep via Boyd/paratibial perforator.
        # old_max = 0.95 ≥ 0.48; ep_n1_n2=T (Hunterian); ep_n2_n3=F → Rule 6 fires.
        Step("CM-2  RP N2→N1 calf (posY=0.70) — trunk re-enters deep — COMPLETE (Hunterian EP Type I)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.70, "Calf", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P26  Verify complete — SPJ",
             "probe_move", _pm(sid, "SPJ", 0.51, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P27  Verify complete — Hunterian (EP site)",
             "probe_move", _pm(sid, "Hunterian", 0.41, "medial"),
             expected_action="complete"),

        Step("P28  Verify complete — SFJ (competent — no clip)",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 11 — aasv_valsalva_negative  (AASV pitfall → Type VI via N1→N3)
# ---------------------------------------------------------------------------
# Source: Delfrate 2023, JTAVR 8(2), p.24:
#   "The N1→N2 reflux at the GSV end assessed by Paraná/squeezing is not
#    specific of the incompetence of the terminal valve and must be every time
#    confirmed by positive Valsalva… If Valsalva is negative, the Paraná
#    reflux is just due to the incompetence of the pre-terminal valve below
#    a competent terminal valve."
#
# PROTOCOL_REFERENCES.md step 3:
#   "AASV (anterior accessory saphenous vein) lies anterior to GSV in upper
#    thigh — assess separately; classified N3, not N2 (common duplex pitfall)."
#
# Clinical picture:
#   SFJ — Paranà positive but Valsalva NEGATIVE → SFJ terminal valve COMPETENT.
#   NO EP N1→N2 placed at SFJ.  Surgeon identifies AASV (anterior to GSV,
#   classified N3) showing reflux fed by an upper-thigh perforator.
#   EP N1→N3: perforator exits deep directly into AASV (N3) at upper thigh.
#   AASV (N3) refluxes distally into calf.
#   RP N3→N1: calf perforator re-entry into deep system.
#   Rule 5 fires: ep_n1_n3 + rp_n3_n1 + NOT rp_n2_n1 + NOT rp_n3_n2 (Type VI).
# ---------------------------------------------------------------------------

def _aasv_valsalva_negative(sid="s11_aasv_vn") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # --- SFJ: Paranà positive, Valsalva NEGATIVE → pre-terminal only → COMPETENT ---
        Step("P01  SFJ — transverse B-mode; Mickey Mouse sign; Doppler on femoral side of valve",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "Valsalva"]),

        Step("P02  SFJ — Valsalva: NO reflux through terminal valve",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move"),

        # Paranà positive → pre-terminal valve reflux only.  Delfrate 2023 p.24:
        # check arch tributaries (AASV, epigastric, pudendal) when Valsalva negative.
        Step("P03  SFJ — Paranà positive but Valsalva negative → TERMINAL VALVE COMPETENT",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move",
             forbidden_action="complete"),

        # --- GSV trunk: antegrade flow (SFJ competent, no N1→N2 from groin) ---
        Step("P04  Upper Thigh medial — GSV in saphenous eye; antegrade flow",
             "probe_move", _pm(sid, "Upper Thigh", 0.10, "medial"),
             expected_action="move"),

        # --- AASV check: anterior to GSV in upper thigh ---
        # AASV lies anterior to GSV, inside saphenous fascia but in a separate compartment.
        # Classified as N3 (superficial tributary), NOT N2.  A common pitfall (Adler 2022 p.2191).
        Step("P05  Upper Thigh anterior — rotating probe: AASV visible anterior to GSV",
             "probe_move", _pm(sid, "Upper Thigh", 0.12, "medial"),
             expected_action="move",
             guidance_must_contain=["thigh", "anterior", "tributary", "AASV"]),

        Step("P06  Upper Thigh — AASV: reflux on Paranà; diameter 4.0 mm (N3 vessel)",
             "probe_move", _pm(sid, "Upper Thigh", 0.14, "medial"),
             expected_action="move"),

        # Source of AASV reflux: perforator in upper thigh exits deep into AASV (N3).
        # Outward flow >500 ms on all 3 maneuvers (AVF 2023 criteria).
        # EP N1→N3: deep blood exits into AASV (N3), NOT into GSV trunk (N2).
        Step("CM-1  EP N1→N3 at Upper Thigh (posY=0.13) — perf exits deep into AASV (N3, NOT N2)",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.13, "Upper Thigh", "medial"),
             expected_action="move",
             forbidden_action="complete",
             guidance_must_contain=["distal", "tributary", "calf", "reflux"]),

        # --- GSV trunk: still antegrade — AASV is a separate N3 vessel ---
        Step("P07  Upper Thigh medial — GSV trunk: antegrade (confirms trunk not involved)",
             "probe_move", _pm(sid, "Upper Thigh", 0.17, "medial"),
             expected_action="move"),

        Step("P08  Dodd proximal — AASV (N3) continues distally; GSV trunk clean",
             "probe_move", _pm(sid, "Dodd", 0.24, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P09  Dodd mid — checking trunk for N2→N3 escape: none present",
             "probe_move", _pm(sid, "Dodd", 0.29, "medial"),
             expected_action="move"),

        Step("P10  Hunterian — no RP N2→N1 in trunk; no trunk involvement",
             "probe_move", _pm(sid, "Hunterian", 0.38, "medial"),
             expected_action="move"),

        Step("P11  Hunterian distal — AASV/N3 swings posterior toward popliteal region",
             "probe_move", _pm(sid, "Hunterian", 0.45, "medial"),
             expected_action="move"),

        # --- SPJ: both maneuvers negative — no SSV circuit ---
        Step("P12  SPJ — Paranà + CR: SPJ terminal valve COMPETENT",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior", "junction"]),

        Step("P13  SPJ mid — Giacomini vein check: no Giacomini involvement",
             "probe_move", _pm(sid, "SPJ", 0.54, "posterior"),
             expected_action="move"),

        # --- Calf: AASV/N3 now in medial calf territory ---
        Step("P14  Calf medial upper — tracing N3 tributary (ex-AASV) distally",
             "probe_move", _pm(sid, "Calf", 0.60, "medial"),
             expected_action="move"),

        Step("P15  Calf medial — N3 continues; no N3→N2 connection (not a RP N3→N2)",
             "probe_move", _pm(sid, "Calf", 0.65, "medial"),
             expected_action="move"),

        Step("P16  Calf medial mid — N3 diameter 3.8 mm; retrograde flow on Paranà",
             "probe_move", _pm(sid, "Calf", 0.71, "medial"),
             expected_action="move"),

        Step("P17  Calf lower — N3 continues toward ankle",
             "probe_move", _pm(sid, "Calf", 0.77, "medial"),
             expected_action="move"),

        Step("P18  Ankle medial",
             "probe_move", _pm(sid, "Ankle", 0.90, "medial"),
             expected_action="move"),

        Step("P19  Ankle distal",
             "probe_move", _pm(sid, "Ankle", 0.94, "medial"),
             expected_action="move"),

        # --- Return sweep: find RP N3→N1 ---
        # Gianesini 2014 p.15: inward diastolic flow during Paranà release = re-entry RP.
        Step("P20  Return: lower calf",
             "probe_move", _pm(sid, "Calf", 0.87, "medial"),
             expected_action="move"),

        Step("P21  Return: calf lower mid — diastolic inward perforator flow detected",
             "probe_move", _pm(sid, "Calf", 0.80, "medial"),
             expected_action="move"),

        Step("P22  Return: calf mid — confirming RP site: inward flow dominant on release",
             "probe_move", _pm(sid, "Calf", 0.73, "medial"),
             expected_action="move"),

        # RP N3→N1: AASV-derived N3 tributary re-enters deep.
        # ep_n1_n3=T, rp_n3_n1=T, rp_n2_n1=F, rp_n3_n2=F → Rule 5 fires (Type VI).
        Step("CM-2  RP N3→N1 at Calf (posY=0.72) — N3 re-enters deep — COMPLETE (Type VI / AASV path)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.72, "Calf", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P23  Verify complete — Hunterian (EP site)",
             "probe_move", _pm(sid, "Hunterian", 0.38, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P24  Verify complete — SFJ (competent — Valsalva negative)",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 12 — ssv_type2_autonomous  (Type II: SSV autonomous pre-terminal reflux)
# ---------------------------------------------------------------------------
# Source: Gianesini et al. 2014, p.14:
#   "The sapheno-popliteal junction must be assessed both by active (Paranà) and
#    passive (CR) manoeuvres, whose positivity must be contemporaneously present
#    to diagnose a true junctional incompetence."
#
# Delfrate 2023, p.26 (Shunt Type 2 / Type II):
#   SSV trunk with autonomous reflux due to pre-terminal valve incompetence,
#   while the SPJ terminal valve remains competent (only Paranà positive, CR negative).
#
# Clinical picture:
#   SFJ: competent (both maneuvers negative).
#   GSV thigh: antegrade.
#   SPJ: Paranà positive (reflux detected) BUT compression/relaxation (CR) NEGATIVE
#     → by Gianesini 2014 both-must-be-positive rule, SPJ TERMINAL VALVE IS COMPETENT.
#     However, below the competent terminal valve the SSV trunk itself has pre-terminal
#     autonomous reflux (intrinsic wall incompetence, not junctional).
#   SSV trunk (acting as autonomous N2 source) escapes into posterior calf N3 tributary
#     (EP N2→N3 at SSV mid-calf).
#   No EP N1→N2 at any junction.  No RP N2→N1 in SSV trunk.
#   RP N3→N1: N3 tributary re-enters deep via posterior ankle perforator.
#   Rule 7: ep_n2_n3 + rp_n3_n1 + NOT ep_n1_n2 + NOT ep_n1_n3 + NOT rp_n2_n1 → Type II.
# ---------------------------------------------------------------------------

def _ssv_type2_autonomous(sid="s12_ssv_t2") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # --- SFJ: both maneuvers negative → COMPETENT ---
        Step("P01  SFJ — Valsalva AND Paranà both negative → SFJ COMPETENT",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02  SFJ deeper",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move"),

        # --- GSV thigh: antegrade throughout ---
        Step("P03  Upper Thigh — GSV in saphenous eye; Paranà: antegrade only",
             "probe_move", _pm(sid, "Upper Thigh", 0.12, "medial"),
             expected_action="move"),

        Step("P04  Dodd — GSV trunk: antegrade; no N3 escape",
             "probe_move", _pm(sid, "Dodd", 0.27, "medial"),
             expected_action="move"),

        Step("P05  Hunterian — GSV antegrade; no perforator outward flow",
             "probe_move", _pm(sid, "Hunterian", 0.40, "medial"),
             expected_action="move"),

        # --- SPJ: Paranà positive BUT CR (compression/relaxation) NEGATIVE ---
        # Gianesini 2014 p.14: BOTH Paranà AND CR must be positive for true SPJ incompetence.
        # Only Paranà positive here → terminal valve COMPETENT.
        # Pre-terminal SSV reflux is present but not junctional.
        Step("P06  SPJ approach — lateral decubitus (left lat for right leg)",
             "probe_move", _pm(sid, "SPJ", 0.48, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior", "junction"]),

        Step("P07  SPJ — Paranà positive; CR (compression/relaxation) NEGATIVE → TERMINAL VALVE COMPETENT",
             "probe_move", _pm(sid, "SPJ", 0.50, "posterior"),
             expected_action="move"),

        Step("P08  SPJ mid — pre-terminal SSV reflux noted below competent terminal valve",
             "probe_move", _pm(sid, "SPJ", 0.54, "posterior"),
             expected_action="move"),

        # --- SSV trunk: autonomous pre-terminal reflux (below competent terminal valve) ---
        # This is the N2 reflux source (SSV = N2 in context of autonomous SSV reflux).
        # No EP N1→N2 has been confirmed (terminal valve competent).
        Step("P09  SSV upper — SSV trunk shows retrograde flow (autonomous, not junctional)",
             "probe_move", _pm(sid, "SSV", 0.59, "posterior"),
             expected_action="move"),

        Step("P10  SSV mid — retrograde SSV trunk continues; searching for N3 escape",
             "probe_move", _pm(sid, "SSV", 0.64, "posterior"),
             expected_action="move"),

        Step("P11  SSV mid — colour Doppler: tributary filling from SSV trunk detected",
             "probe_move", _pm(sid, "SSV", 0.68, "posterior"),
             expected_action="move"),

        # EP N2→N3: SSV trunk (autonomous N2) escapes into posterior calf N3 tributary.
        # No EP N1→N2 present. Rule 7 will need rp_n3_n1 and no rp_n2_n1.
        Step("CM-1  EP N2→N3 at SSV (posY=0.66) — SSV trunk escapes into posterior N3 tributary",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.66, "SSV", "posterior"),
             expected_action="move",
             forbidden_action="complete",
             guidance_must_contain=["calf", "posterior", "tributary", "distal"]),

        # --- Following N3 distally ---
        Step("P12  SSV lower — N3 tributary continues on posterior calf",
             "probe_move", _pm(sid, "SSV", 0.74, "posterior"),
             expected_action="move"),

        Step("P13  SSV lower — N3 swings toward lateral ankle",
             "probe_move", _pm(sid, "SSV", 0.80, "posterior"),
             expected_action="move"),

        Step("P14  Ankle posterior",
             "probe_move", _pm(sid, "Ankle", 0.90, "posterior"),
             expected_action="move"),

        Step("P15  Ankle lateral — N3 at lateral malleolus",
             "probe_move", _pm(sid, "Ankle", 0.93, "lateral"),
             expected_action="move"),

        # --- Return sweep: finding RP N3→N1 ---
        Step("P16  Return: lower calf lateral",
             "probe_move", _pm(sid, "Calf", 0.87, "lateral"),
             expected_action="move"),

        Step("P17  Return: calf lateral mid — biphasic perforator: inward diastolic dominant",
             "probe_move", _pm(sid, "Calf", 0.80, "lateral"),
             expected_action="move"),

        Step("P18  Return: calf lateral — diastolic inward flow confirmed at perforator",
             "probe_move", _pm(sid, "Calf", 0.74, "lateral"),
             expected_action="move"),

        # RP N3→N1: N3 posterior tributary re-enters deep.
        # ep_n2_n3=T, rp_n3_n1=T, ep_n1_n2=F, ep_n1_n3=F, rp_n2_n1=F → Rule 7 fires (Type II).
        Step("CM-2  RP N3→N1 at Ankle (posY=0.91) — N3 re-enters deep — COMPLETE (SSV Type II)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.91, "Ankle", "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P19  Verify complete — SSV mid",
             "probe_move", _pm(sid, "SSV", 0.65, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P20  Verify complete — SPJ (one-maneuver: competent)",
             "probe_move", _pm(sid, "SPJ", 0.51, "posterior"),
             expected_action="complete"),

        Step("P21  Verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 13 — spj_type3_noreflux  (Type III via SPJ entry)
# ---------------------------------------------------------------------------
# Source: Franceschi & Zamboni 2009, p.114 (reflux elimination test):
#   "Reflux is reduced but still active at the level of the GSV main trunk
#    (B1-B2), despite finger compression of reflux points 2 and 3."
# Gianesini 2014, p.14 (SPJ protocol):
#   "In the same way, the sapheno-popliteal junction must be assessed both by
#    active (Paranà) and passive (CR) manoeuvres."
#
# Clinical picture:
#   SFJ: competent.  SPJ: BOTH Paranà AND CR positive → EP N1→N2 at SPJ.
#   SSV trunk (N2) refluxes distally.
#   EP N2→N3: SSV trunk escapes into N3 posterior calf tributary (MARKED BEFORE RP N2→N1).
#   RP N3→N1: N3 re-enters deep via posterior ankle perforator.
#   RP N2→N1: SSV trunk conducted reflux (marked AFTER EP N2→N3).
#   → Rule 1 MANEUVER fires (ep_n2_n3_no_elim + rp_n3_n1 + rp_n2_n1).
#   Surgeon compresses N3 tributary at SSV escape point → SSV trunk reflux disappears.
#   Elimination test → "No Reflux" → SSV trunk was conduit only → Type III.
#   Rule 2 COMPLETE.
#
# ORDERING CRITICAL: EP N2→N3 must be marked BEFORE RP N2→N1 to prevent
#   Rule 6 (Type I) from firing on the RP N2→N1 mark.
# ---------------------------------------------------------------------------

def _spj_type3_noreflux(sid="s13_spj_t3") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # --- SFJ: competent ---
        Step("P01  SFJ — Valsalva + Paranà both negative → SFJ COMPETENT",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02  SFJ deeper",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move"),

        # --- GSV thigh: antegrade (SFJ competent, no N1→N2 at groin) ---
        Step("P03  Upper Thigh — GSV antegrade; proceeding to SPJ",
             "probe_move", _pm(sid, "Upper Thigh", 0.13, "medial"),
             expected_action="move"),

        Step("P04  Hunterian — GSV antegrade; no escape into N3",
             "probe_move", _pm(sid, "Hunterian", 0.41, "medial"),
             expected_action="move"),

        # --- SPJ: BOTH Paranà AND CR positive → INCOMPETENT ---
        # Gianesini 2014 p.14: concordance of both maneuvers = true junctional incompetence.
        # Adler et al. 2022 p.2190: lateral decubitus; left lat for right leg.
        Step("P05  SPJ approach — lateral decubitus; popliteal fossa visualised",
             "probe_move", _pm(sid, "SPJ", 0.48, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior", "junction"]),

        Step("P06  SPJ — Paranà POSITIVE AND CR (compression/relaxation) POSITIVE → SPJ INCOMPETENT",
             "probe_move", _pm(sid, "SPJ", 0.51, "posterior"),
             expected_action="move"),

        # EP N1→N2 at SPJ: deep (popliteal vein) blood enters SSV trunk.
        Step("CM-1  EP N1→N2 at SPJ (posY=0.51) — SPJ incompetent; popliteal → SSV trunk",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.51, "SPJ", "posterior"),
             expected_action="move",
             guidance_must_contain=["SSV", "calf", "posterior", "trunk"]),

        # --- SSV trunk: tracing distally for N3 escape ---
        Step("P07  SSV upper — SSV trunk retrograde on Paranà; searching for escape",
             "probe_move", _pm(sid, "SSV", 0.58, "posterior"),
             expected_action="move"),

        Step("P08  SSV mid — retrograde sustained >500 ms; tributaries assessed",
             "probe_move", _pm(sid, "SSV", 0.63, "posterior"),
             expected_action="move"),

        Step("P09  SSV mid — colour Doppler: posterior tributary filling from SSV trunk",
             "probe_move", _pm(sid, "SSV", 0.67, "posterior"),
             expected_action="move"),

        # EP N2→N3: SSV trunk escapes into N3 posterior calf tributary.
        # CRITICAL: marked BEFORE RP N2→N1 to prevent Rule 6 triggering.
        Step("CM-2  EP N2→N3 at SSV (posY=0.65) — SSV trunk escapes into N3 tributary (no elimTest yet)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.65, "SSV", "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        # --- N3 tributary: trace distally to re-entry ---
        Step("P10  SSV lower — N3 continues on posterior calf",
             "probe_move", _pm(sid, "SSV", 0.74, "posterior"),
             expected_action="move"),

        Step("P11  Calf posterior — N3 at posterior lower calf",
             "probe_move", _pm(sid, "Calf", 0.80, "posterior"),
             expected_action="move"),

        Step("P12  Ankle posterior — N3 approaching ankle re-entry",
             "probe_move", _pm(sid, "Ankle", 0.90, "posterior"),
             expected_action="move"),

        Step("P13  Ankle lateral — N3 lateral to Achilles tendon",
             "probe_move", _pm(sid, "Ankle", 0.93, "lateral"),
             expected_action="move"),

        Step("P14  Return: calf posterior — scanning for inward diastolic perforator flow",
             "probe_move", _pm(sid, "Calf", 0.84, "posterior"),
             expected_action="move"),

        # RP N3→N1: N3 tributary re-enters deep via posterior calf perforator.
        # ep_n2_n3_no_elim=T, rp_n3_n1=T, rp_n2_n1=F → Rule 1 NOT yet (need rp_n2_n1).
        Step("CM-3  RP N3→N1 at SSV calf (posY=0.78) — N3 re-enters deep (maneuver not yet: need RP N2→N1)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.78, "SSV", "posterior"),
             expected_action="move",
             forbidden_action="maneuver"),

        # --- Return sweep toward SPJ: SSV trunk conducted reflux ---
        Step("P15  Return: calf posterior — SSV trunk still retrograde on relaxation",
             "probe_move", _pm(sid, "SSV", 0.71, "posterior"),
             expected_action="move"),

        Step("P16  Return: SSV mid — trunk shows conducted retrograde flow (N1→N2→N2→N1 path)",
             "probe_move", _pm(sid, "SSV", 0.65, "posterior"),
             expected_action="move"),

        Step("P17  Return: SPJ zone — SSV trunk proximal to escape point",
             "probe_move", _pm(sid, "SPJ", 0.54, "posterior"),
             expected_action="move"),

        # RP N2→N1: conducted SSV trunk reflux (potentially only conducted, not independent).
        # ep_n2_n3_no_elim + rp_n3_n1 + rp_n2_n1 → Rule 1 MANEUVER fires.
        Step("CM-4  RP N2→N1 at SSV (posY=0.61) — conducted SSV trunk reflux — MANEUVER fires",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.61, "SSV", "posterior"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary", "Doppler", "record"]),

        # --- Maneuver state: compress N3 tributary at escape point ---
        # Gianesini 2014, p.13: digit compression of N3 escape point abolishes
        # trunk reflux if RP is on N3 (Type III). Persistent trunk reflux = Type I+2.
        Step("P18  MANEUVER — return to SSV escape point (posY=0.65) for compression test",
             "probe_move", _pm(sid, "SSV", 0.65, "posterior"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary"]),

        Step("P19  MANEUVER — SSV trunk checked: reflux DISAPPEARS on N3 compression",
             "probe_move", _pm(sid, "SSV", 0.61, "posterior"),
             expected_action="maneuver"),

        Step("P20  Return to SSV escape for elimination clip",
             "probe_move", _pm(sid, "SSV", 0.65, "posterior"),
             expected_action="maneuver"),

        # Elimination test: compress N3 tributary at SSV escape point.
        # SSV trunk reflux disappears → "No Reflux" → trunk was conduit only → Type III.
        # Rule 2 fires: complete.
        Step("CM-5  EP N2→N3 elim='No Reflux' at SSV (posY=0.65) — trunk was conduit — COMPLETE (SPJ Type III)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.65, "SSV", "posterior", elim="No Reflux"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P21  Verify complete — SSV mid",
             "probe_move", _pm(sid, "SSV", 0.63, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P22  Verify complete — SPJ",
             "probe_move", _pm(sid, "SPJ", 0.51, "posterior"),
             expected_action="complete"),

        Step("P23  Verify complete — SFJ (competent)",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 14 — rule6_max_visited_gate  (Rule 6 max_visited gate: no premature complete)
# ---------------------------------------------------------------------------
# Source: run_scenarios.py header:
#   "Complete-rule gate: old_max_visited_pos_y >= 0.48 (probe must reach SPJ zone
#    BEFORE RP N2->N1)"
# Franceschi & Zamboni 2009, p.115: "No reflux, no re-entry and viceversa" —
#   the surgeon must confirm the full circuit including the popliteal/SPJ zone
#   before declaring a Type I classification.
#
# Clinical picture:
#   Surgeon correctly identifies EP N1→N2 at SFJ and then marks RP N2→N1
#   EARLY — at upper thigh (posY=0.15) — without having covered the SPJ zone.
#   old_max_visited = 0.15 < 0.48 → Rule 6 gate NOT met → action="move".
#   Surgeon then continues distal (Dodd → Hunterian → SPJ).
#   Once the probe covers SPJ (max_visited ≥ 0.48), the surgeon marks a second
#   RP N2→N1 at calf level.  Now old_max ≥ 0.48 → Rule 6 fires → complete.
#
# Important engine note: Rule 6 fires on CLIP_MARK events when old_max ≥ 0.48
#   at the moment of that clip.  Probe_moves alone do not retroactively trigger
#   complete for clips that were already marked before the gate was met.
#   Therefore the test needs a second RP N2→N1 mark after SPJ coverage.
#
# Key assertions:
#   CM-2: forbidden_action="complete" — early RP at posY=0.15, old_max<0.48.
#   P05–P08: forbidden_action="complete" throughout Dodd/Hunterian scan.
#   CM-3: expected_action="complete" — RP at posY=0.70, old_max≥0.48.
# ---------------------------------------------------------------------------

def _rule6_max_visited_gate(sid="s14_gate") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # --- SFJ: incompetent (both maneuvers positive) ---
        Step("P01  SFJ — Valsalva positive AND Paranà positive → SFJ INCOMPETENT",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02  SFJ mid — confirming both-maneuver positivity",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move"),

        # EP N1→N2 at SFJ: both maneuvers positive.
        Step("CM-1  EP N1→N2 at SFJ (posY=0.06) — SFJ confirmed incompetent",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.06, "SFJ", "anterior-medial"),
             expected_action="move"),

        # --- Upper Thigh: retrograde trunk flow ---
        Step("P03  Upper Thigh — GSV retrograde trunk; Paranà confirms reflux",
             "probe_move", _pm(sid, "Upper Thigh", 0.12, "medial"),
             expected_action="move"),

        Step("P04  Upper Thigh mid — retrograde sustained >500 ms",
             "probe_move", _pm(sid, "Upper Thigh", 0.16, "medial"),
             expected_action="move"),

        # RP N2→N1 marked EARLY — posY=0.15, old_max=0.16 < 0.48.
        # Rule 6 gate NOT met → action MUST be "move".  KEY gate assertion.
        Step("CM-2  RP N2→N1 at Upper Thigh (posY=0.15) — EARLY; old_max=0.16<0.48 → gate blocks complete",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.15, "Upper Thigh", "medial"),
             expected_action="move",
             forbidden_action="complete"),  # PRIMARY GATE ASSERTION

        # --- Dodd: max_visited < 0.48 throughout, complete must NOT fire ---
        Step("P05  Dodd proximal — max_visited=0.24 < 0.48; gate holds",
             "probe_move", _pm(sid, "Dodd", 0.24, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P06  Dodd mid — scanning trunk; no N3 escape",
             "probe_move", _pm(sid, "Dodd", 0.30, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # --- Hunterian: max_visited still < 0.48 ---
        Step("P07  Hunterian proximal — max_visited=0.36 < 0.48; complete must not fire",
             "probe_move", _pm(sid, "Hunterian", 0.36, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P08  Hunterian distal — max_visited=0.46 < 0.48; gate still not met",
             "probe_move", _pm(sid, "Hunterian", 0.46, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # --- SPJ zone: probe crosses 0.48; now old_max will be ≥ 0.48 ---
        Step("P09  SPJ approach (posY=0.49) — max_visited now 0.49 ≥ 0.48",
             "probe_move", _pm(sid, "SPJ", 0.49, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior", "junction"]),

        Step("P10  SPJ mid — confirming SPJ competent; max_visited=0.53",
             "probe_move", _pm(sid, "SPJ", 0.53, "posterior"),
             expected_action="move"),

        # --- Calf: move to calf to mark second RP N2→N1 ---
        # At this point old_max is ≥ 0.53 from SPJ probe steps.
        Step("P11  Calf medial — trunk retrograde; paratibial perforator visible",
             "probe_move", _pm(sid, "Calf", 0.65, "medial"),
             expected_action="move"),

        Step("P12  Calf medial — diastolic inward flow confirmed at perforator",
             "probe_move", _pm(sid, "Calf", 0.70, "medial"),
             expected_action="move"),

        # Second RP N2→N1: old_max ≥ 0.53 at this point (SPJ was covered).
        # ep_n1_n2=T, rp_n2_n1=T, ep_n2_n3=F, old_max ≥ 0.48 → Rule 6 fires.
        Step("CM-3  RP N2→N1 at Calf (posY=0.70) — old_max≥0.53; gate met → COMPLETE (Rule 6)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.70, "Calf", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P13  Verify complete — SPJ",
             "probe_move", _pm(sid, "SPJ", 0.51, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P14  Verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

_SCENARIOS: dict[str, callable] = {
    "hunterian_perf_entry":   _hunterian_perf_entry,
    "aasv_valsalva_negative": _aasv_valsalva_negative,
    "ssv_type2_autonomous":   _ssv_type2_autonomous,
    "spj_type3_noreflux":     _spj_type3_noreflux,
    "rule6_max_visited_gate": _rule6_max_visited_gate,
}

_SCENARIO_DESCRIPTIONS: dict[str, str] = {
    "hunterian_perf_entry": (
        "Type I via Hunterian perforator EP: SFJ Valsalva-negative (terminal valve competent); "
        "EP N1→N2 at Hunterian perf (posY=0.41); GSV trunk reflux; SPJ competent; "
        "calf RP N2→N1 → Rule 6.  Source: Gianesini 2014 p.11, Delfrate 2023 p.24.  28 steps."
    ),
    "aasv_valsalva_negative": (
        "AASV pitfall → Type VI: SFJ Paranà-only positive (Valsalva negative) → SFJ COMPETENT; "
        "AASV classified N3 (not N2); EP N1→N3 upper-thigh perf → AASV; RP N3→N1 calf → Rule 5. "
        "Source: Delfrate 2023 p.24, PROTOCOL_REFERENCES step 3.  24 steps."
    ),
    "ssv_type2_autonomous": (
        "Type II via SSV autonomous pre-terminal reflux: SPJ one-maneuver-only positive → terminal "
        "valve COMPETENT; SSV below competent valve has autonomous reflux; EP N2→N3 SSV escape; "
        "RP N3→N1 ankle lateral → Rule 7.  Source: Gianesini 2014 p.14.  21 steps."
    ),
    "spj_type3_noreflux": (
        "Type III via SPJ entry: SPJ both-maneuvers positive → EP N1→N2 (SPJ); EP N2→N3 SSV escape "
        "(marked before RP N2→N1); RP N3→N1; RP N2→N1 → Rule 1 maneuver; compress N3 → No Reflux "
        "→ Rule 2 complete.  Source: Franceschi & Zamboni 2009 p.114, Gianesini 2014 p.13.  23 steps."
    ),
    "rule6_max_visited_gate": (
        "Rule 6 max_visited gate regression: EP N1→N2 SFJ + RP N2→N1 at posY=0.15 (old_max<0.48) "
        "→ forbidden complete (gate holds).  Second RP N2→N1 at calf after SPJ coverage "
        "(old_max≥0.48) → complete fires.  Source: run_scenarios.py gate spec.  15 steps."
    ),
}


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

CYAN   = "\033[96m"
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

_SKIP_EVENTS = {"stream_start"}


def _describe_movement(event: str, data: dict) -> str:
    if event == "stream_start":
        return f"Start session  (id: {data.get('session_id', '?')})"
    if event == "probe_move":
        return (f"Move → {data.get('region', '?')}\n"
                f"posY={data.get('pos_y_ratio', 0.0):.2f}  "
                f"{data.get('surface', '?')}  {data.get('leg', 'right')} leg")
    if event == "clip_mark":
        base = (f"Mark: {data.get('flow', '?')} "
                f"{data.get('from_type', '?')}→{data.get('to_type', '?')}  "
                f"posY={data.get('pos_y_ratio', 0.0):.2f}  "
                f"{data.get('region', '?')}  {data.get('leg', 'right')} leg")
        if data.get("elimination_test"):
            base += f"\nElim: {data['elimination_test']}"
        return base
    return event


def _describe_expectation(step: Step) -> str:
    if step.event == "stream_start":
        return "session_ready event"
    parts = []
    if step.expected_action:
        parts.append(f'Action = "{step.expected_action}"')
    if step.forbidden_action:
        parts.append(f'Must NOT = "{step.forbidden_action}"')
    if step.guidance_must_contain:
        parts.append("Guidance contains any of: " + ", ".join(step.guidance_must_contain))
    return "\n".join(parts) if parts else "observe only"


class ScenarioRunner:
    def __init__(self, api_base: str):
        self.api_base = api_base.rstrip("/")
        self._sio = socketio.SimpleClient()

    def connect(self) -> None:
        print(f"Connecting to {self.api_base} ...")
        self._sio.connect(self.api_base, wait_timeout=10)
        print("Connected.\n")

    def disconnect(self) -> None:
        self._sio.disconnect()

    def _emit_and_wait(self, event: str, data: dict, timeout: float = 60.0):
        self._sio.emit(event, data)
        if event in _SKIP_EVENTS:
            deadline = time.time() + timeout
            while time.time() < deadline:
                try:
                    ev = self._sio.receive(timeout=2)
                    if ev and ev[0] == "session_ready":
                        return ev[1]
                except Exception:
                    pass
            return None
        deadline = time.time() + timeout
        while time.time() < deadline:
            try:
                ev = self._sio.receive(timeout=2)
                if ev and ev[0] == "guidance_update":
                    return ev[1]
            except Exception:
                pass
        return None

    def run(self, steps: list[Step]) -> list[dict]:
        results = []
        for step in steps:
            print(f"\n{CYAN}--- {step.label} ---{RESET}")

            movement    = _describe_movement(step.event, step.data)
            expectation = _describe_expectation(step)
            response    = self._emit_and_wait(step.event, step.data)

            if step.event in _SKIP_EVENTS:
                status = "SKIP" if response else "TIMEOUT"
                if response:
                    print(f"  session_ready: {response}")
                else:
                    print(f"  {YELLOW}No session_ready{RESET}")
                results.append({"step": step.label, "status": status,
                                "movement": movement, "expectation": expectation,
                                "llm_action": "--", "llm_guidance": "--",
                                "failures": []})
                continue

            if (step.expected_action is None
                    and not step.guidance_must_contain
                    and step.forbidden_action is None):
                if response:
                    g = response.get("guidance", "")
                    a = response.get("action", "?")
                    print(f"  guidance: {g!r}  action: {a}")
                    results.append({"step": step.label, "status": "SKIP",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": a, "llm_guidance": g, "failures": []})
                else:
                    print(f"  {YELLOW}timeout{RESET}")
                    results.append({"step": step.label, "status": "TIMEOUT",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                continue

            if response is None:
                print(f"  {YELLOW}No response (timeout){RESET}")
                results.append({"step": step.label, "status": "TIMEOUT",
                                "movement": movement, "expectation": expectation,
                                "llm_action": "--", "llm_guidance": "--",
                                "failures": []})
                continue

            llm_guidance = response.get("guidance", "") or ""
            llm_action   = response.get("action", "move") or "move"
            print(f"  guidance: {llm_guidance!r}")
            print(f"  action  : {llm_action}")

            passed   = True
            failures = []

            if step.expected_action and llm_action != step.expected_action:
                passed = False
                failures.append(
                    f'Action: got "{llm_action}", expected "{step.expected_action}"')

            if step.guidance_must_contain:
                if not any(kw.lower() in llm_guidance.lower()
                           for kw in step.guidance_must_contain):
                    passed = False
                    failures.append(
                        "Guidance missing any of: "
                        + ", ".join(step.guidance_must_contain))

            if step.forbidden_action and llm_action == step.forbidden_action:
                passed = False
                failures.append(f'Action "{llm_action}" must NOT fire here')

            status = "PASS" if passed else "FAIL"
            if passed:
                print(f"  {GREEN}{BOLD}PASS{RESET}")
            else:
                for f in failures:
                    print(f"  {RED}FAIL: {f}{RESET}")

            results.append({"step": step.label, "status": status,
                            "movement": movement, "expectation": expectation,
                            "llm_action": llm_action, "llm_guidance": llm_guidance,
                            "failures": failures})
            time.sleep(0.4)

        return results

    def print_summary(self, label: str, results: list[dict]) -> bool:
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        passed  = [r for r in checked if r["status"] == "PASS"]
        failed  = [r for r in checked if r["status"] == "FAIL"]
        timeout = [r for r in results  if r["status"] == "TIMEOUT"]
        print(f"\n{'─'*65}")
        print(f"{BOLD}{label}{RESET}  "
              f"{GREEN}{len(passed)} passed{RESET}  "
              f"{RED}{len(failed)} failed{RESET}  "
              f"{YELLOW}{len(timeout)} timeout{RESET}  "
              f"/ {len(checked)} checked")
        for r in failed:
            print(f"  {RED}x {r['step']}{RESET}")
            for f in r.get("failures", []):
                print(f"      {f}")
        return len(failed) == 0 and len(timeout) == 0


# ---------------------------------------------------------------------------
# Word report
# ---------------------------------------------------------------------------

_STATUS_COLORS = {"PASS": "00B050", "FAIL": "FF0000",
                  "TIMEOUT": "FF8C00", "SKIP": "808080"}


def _colored_cell(cell, text: str, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    from docx.shared import RGBColor
    cell.text = text
    run = (cell.paragraphs[0].runs[0]
           if cell.paragraphs[0].runs
           else cell.paragraphs[0].add_run(text))
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, pt: int = 9, bold=False,
                   italic=False, color_rgb=None) -> None:
    cell.text = ""
    for i, line in enumerate(text.split("\n")):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run  = para.add_run(line)
        run.font.size   = __import__("docx").shared.Pt(pt)
        run.font.bold   = bold
        run.font.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb


def write_word_report(all_results: dict[str, list[dict]],
                      api_base: str, out_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin  = sec.bottom_margin = Inches(0.7)

    t = doc.add_heading(
        "CHIVA Streaming Guidance — Extended Scenario Test Report (v2)", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = m.add_run(
        f"Date: {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}    |    "
        f"API: {api_base}    |    Scenarios: {len(all_results)}\n"
        "Literature: Franceschi & Zamboni 2009 · Gianesini 2014 · Delfrate 2023 · Adler 2022")
    mr.font.size = Pt(9)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    doc.add_heading("Overall Summary", level=1)
    stbl = doc.add_table(rows=1, cols=5)
    stbl.style = "Table Grid"
    for i, h in enumerate(["Scenario", "Description", "Passed", "Failed", "Timeouts"]):
        c = stbl.rows[0].cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True

    overall_pass = True
    for name, results in all_results.items():
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")
        row = stbl.add_row().cells
        row[0].text = name
        row[1].text = _SCENARIO_DESCRIPTIONS.get(name, "")
        row[2].text = str(p)
        row[3].text = str(f)
        row[4].text = str(t)
        if f or t:
            overall_pass = False

    doc.add_paragraph()
    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run("OVERALL: " + ("ALL PASS" if overall_pass else "FAILURES PRESENT"))
    vr.font.bold = True
    vr.font.size = Pt(13)
    vr.font.color.rgb = (RGBColor(0x00, 0xB0, 0x50) if overall_pass
                         else RGBColor(0xFF, 0x00, 0x00))

    for name, results in all_results.items():
        doc.add_page_break()
        doc.add_heading(f"Scenario: {name.upper()}", level=1)
        dp = doc.add_paragraph(_SCENARIO_DESCRIPTIONS.get(name, ""))
        if dp.runs:
            dp.runs[0].font.italic = True

        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")

        sl = doc.add_paragraph()
        sl.add_run(f"Checked: {len(checked)}   |   ")
        pr = sl.add_run(f"Passed: {p}")
        pr.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
        sl.add_run("   |   ")
        fr = sl.add_run(f"Failed: {f}")
        fr.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        sl.add_run("   |   ")
        tr = sl.add_run(f"Timeouts: {t}")
        tr.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        doc.add_paragraph()

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        for i, h in enumerate(["#", "Surgeon Action", "Expected",
                                "LLM Guidance", "Action", "Status"]):
            c = tbl.rows[0].cells[i]
            c.text = h
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.bold = True
        for i, w in enumerate([0.25, 1.80, 2.10, 2.10, 0.80, 0.65]):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)

        for idx, r in enumerate(results, 1):
            status = r.get("status", "SKIP")
            row    = tbl.add_row().cells
            row[0].text = str(idx)
            _set_cell_text(row[1], r.get("movement", "--"), pt=8)
            _set_cell_text(row[2], r.get("expectation", "--"), pt=8, italic=True)
            _set_cell_text(row[3], r.get("llm_guidance", "--") or "--", pt=8)
            _set_cell_text(row[4], r.get("llm_action", "--") or "--", pt=9, bold=True)
            _colored_cell(row[5], status, _STATUS_COLORS.get(status, "808080"))
            if status == "FAIL" and r.get("failures"):
                fr_row = tbl.add_row().cells
                merged = fr_row[0].merge(fr_row[5])
                merged.text = ""
                for fn in r["failures"]:
                    pp = merged.add_paragraph(f"  x  {fn}")
                    if pp.runs:
                        pp.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        pp.runs[0].font.size      = Pt(8)
                        pp.runs[0].font.italic    = True

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"\nReport saved: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="CHIVA streaming guidance — extended scenario runner (v2)")
    parser.add_argument("scenario", nargs="?", default="all",
                        choices=list(_SCENARIOS.keys()) + ["all"],
                        help="Scenario ID or 'all' (default: all)")
    parser.add_argument("--api", default="http://localhost:7861")
    parser.add_argument("--all", action="store_true", dest="run_all")
    args = parser.parse_args()

    to_run = list(_SCENARIOS.items()) if (args.run_all or args.scenario == "all") \
             else [(args.scenario, _SCENARIOS[args.scenario])]

    runner     = ScenarioRunner(args.api)
    all_ok     = True
    all_results: dict[str, list[dict]] = {}

    try:
        runner.connect()
        for name, builder in to_run:
            print(f"\n{'='*65}")
            print(f"{BOLD}SCENARIO: {name.upper()}{RESET}")
            desc = _SCENARIO_DESCRIPTIONS.get(name, "")
            if desc:
                print(f"  {desc}")
            print(f"{'='*65}")
            steps   = builder()
            results = runner.run(steps)
            ok      = runner.print_summary(name, results)
            all_ok  = all_ok and ok
            all_results[name] = results
    except KeyboardInterrupt:
        print("\nInterrupted.")
    finally:
        runner.disconnect()

    if len(to_run) > 1:
        print(f"\n{'='*65}")
        print(f"{BOLD}OVERALL: {'ALL PASS' if all_ok else 'FAILURES PRESENT'}{RESET}")

    if all_results:
        ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
        here     = os.path.dirname(os.path.abspath(__file__))
        out_path = os.path.join(here, "results", f"scenarios_v2_{ts}.docx")
        try:
            write_word_report(all_results, args.api, out_path)
        except Exception as exc:
            print(f"{YELLOW}Warning: could not write Word report — {exc}{RESET}")

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
