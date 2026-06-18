"""
adversarial_testing — scenarios where the surgeon repeatedly ignores LLM guidance.

Design intent
-------------
At each step labelled "IGNORE #N" the surgeon moves the probe to a region or
posY that is clinically unexpected given the guidance just received.  We then
assess:
  1. Does the action remain "move" (no premature complete/maneuver)?
  2. Does the LLM still produce valid guidance at the actual probe location?
  3. Does the circuit still complete correctly once the surgeon assembles all
     required clips — even after a chaotic scan path?

Scenarios
---------
sfj_bypass_calf_first
    Surgeon opens session and immediately scans calf/ankle (bypassing SFJ).
    Ignores 3 suggestions to start at SFJ.  Eventually marks EP N1→N2 at SFJ.
    Ignores distal-guidance after that, hops randomly.  RP N2→N1 at calf
    → Rule 6 complete (max_visited reached ankle during early wander).

retrograde_multiple_jumps
    After EP N1→N2 at SFJ, surgeon scans retrograde (back toward groin) three
    times, makes two large jumps to ankle and back.  Never follows "trace
    distally" instruction.  RP N2→N1 at calf → Rule 6 complete.

wrong_region_type4
    EP N1→N3 found at ankle.  Surgeon ignores all "trace N3 / find trunk
    reflux" guidance and hops to SFJ, upper thigh, lateral calf, popliteal six
    times before finally finding RP N2→N1 at Hunterian → Rule 4 complete
    (immediate, no max_visited gate).

systematic_ignore_type6
    EP N1→N3 found at lateral calf.  Surgeon ignores N3-tracing guidance,
    zigzags through SFJ, upper thigh, wrong surfaces, wrong regions — six
    ignores.  Eventually traces lateral to ankle and confirms RP N3→N1
    → Rule 5 complete (ep_n1_n3 + rp_n3_n1, no rp_n2_n1).

Usage
-----
    python tests/run_stream_adversarial.py [--api http://localhost:7861] [--all]
    python tests/run_stream_adversarial.py sfj_bypass_calf_first
"""
from __future__ import annotations

import argparse
import io
import os
import sys
import time

if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

try:
    import socketio
except ImportError:
    print("ERROR: python-socketio not installed.  Run: pip install python-socketio[client]")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Step definition
# ---------------------------------------------------------------------------

@dataclass
class Step:
    label: str
    event: str
    data: dict
    expected_action: Optional[str] = None
    guidance_must_contain: list[str] = field(default_factory=list)
    forbidden_action: Optional[str] = None


def _pm(sid, region, pos_y, surface, leg="right"):
    return {"session_id": sid, "region": region, "pos_y_ratio": pos_y,
            "surface": surface, "leg": leg}

def _cm(sid, flow, ft, tt, pos_y, region, surface, elim="", leg="right"):
    return {"session_id": sid, "flow": flow, "from_type": ft, "to_type": tt,
            "pos_y_ratio": pos_y, "leg": leg, "region": region,
            "surface": surface, "elimination_test": elim}


# ---------------------------------------------------------------------------
# Scenario 1 — sfj_bypass_calf_first  (22 checked steps)
# ---------------------------------------------------------------------------
# Surgeon starts scanning at calf/ankle (completely bypassing SFJ).  LLM
# guidance presumably says "assess SFJ at groin" — surgeon ignores it 3 times.
# After eventually reaching SFJ and marking EP N1→N2, surgeon ignores "trace
# distally" and jumps randomly.  Circuit completes via Rule 6 once RP N2→N1
# is marked (max_visited reached 0.91 during the early ankle wander).

def _sfj_bypass_calf_first(sid="adv_sfj_bypass") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # IGNORE #1 — should start at SFJ, goes to calf instead
        Step("P01 [IGNORE #1] calf medial — surgeon skips SFJ entirely",
             "probe_move", _pm(sid, "CALF", 0.63, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #2 — guidance likely says "go to SFJ", surgeon goes to ankle
        Step("P02 [IGNORE #2] ankle medial — deeper into wrong zone",
             "probe_move", _pm(sid, "ANKLE", 0.87, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P03 ankle lateral — continuing wrong direction",
             "probe_move", _pm(sid, "ANKLE", 0.91, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P04 calf lateral return",
             "probe_move", _pm(sid, "CALF", 0.80, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #3 — still not going to SFJ, goes to popliteal
        Step("P05 [IGNORE #3] popliteal — ignores SFJ guidance again",
             "probe_move", _pm(sid, "POPLITEAL", 0.48, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P06 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.52, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P07 upper thigh — finally heading toward SFJ",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.15, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P08 SFJ approach",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             forbidden_action="complete",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P09 SFJ scan",
             "probe_move", _pm(sid, "SFJ", 0.09, "anterior-medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("CM-1  EP N1→N2 at SFJ — entry finally confirmed",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.07, "SFJ", "anterior-medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #4 — guidance says "trace distally toward Hunterian",
        #             surgeon jumps to ankle instead
        Step("P10 [IGNORE #4] ankle medial — ignores distal-trace guidance",
             "probe_move", _pm(sid, "ANKLE", 0.87, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P11 calf medial upper",
             "probe_move", _pm(sid, "CALF", 0.79, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P12 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.73, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #5 — guidance likely says "check popliteal SPJ",
        #             surgeon jumps to Hunterian instead
        Step("P13 [IGNORE #5] Hunterian — skips popliteal assessment",
             "probe_move", _pm(sid, "HUNTERIAN", 0.34, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P14 Hunterian proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.27, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #6 — goes proximal instead of continuing distally
        Step("P15 [IGNORE #6] upper thigh — proximal instead of distal",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.19, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P16 Hunterian — returning to mid-thigh",
             "probe_move", _pm(sid, "HUNTERIAN", 0.29, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #7 — ignores instruction to continue toward popliteal,
        #             jumps to calf
        Step("P17 [IGNORE #7] calf medial — big jump, ignores guidance",
             "probe_move", _pm(sid, "CALF", 0.62, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P18 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.68, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # max_visited = 0.91 (reached at P03), ep_n1_n2=T, rp_n2_n1=T,
        # not ep_n2_n3 → Rule 6 fires COMPLETE
        Step("CM-2  RP N2→N1 calf — COMPLETE (max_visited=0.91 from early wander)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.68, "CALF", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P19 verify complete — popliteal",
             "probe_move", _pm(sid, "POPLITEAL", 0.49, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P20 verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 2 — retrograde_multiple_jumps  (22 checked steps)
# ---------------------------------------------------------------------------
# After EP N1→N2 at SFJ, surgeon repeatedly scans RETROGRADE (back toward
# groin, posY decreasing) instead of tracing distally as instructed.  Between
# retrograde passes the surgeon also makes two large jumps to ankle and back.
# Only after 6 ignores does the surgeon systematically cover the calf, letting
# Rule 6 fire when RP N2→N1 is confirmed.

def _retrograde_multiple_jumps(sid="adv_retro_jumps") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01 SFJ approach",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 SFJ scan",
             "probe_move", _pm(sid, "SFJ", 0.09, "anterior-medial"),
             expected_action="move"),

        Step("CM-1  EP N1→N2 at SFJ",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.07, "SFJ", "anterior-medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #1 — guidance says "scan distally toward thigh",
        #             surgeon goes retrograde toward groin (posY decreasing)
        Step("P03 [IGNORE #1] SFJ 0.04 — retrograde, back toward groin",
             "probe_move", _pm(sid, "SFJ", 0.04, "anterior-medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P04 upper thigh — starting to correct",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.17, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #2 — retrograde again
        Step("P05 [IGNORE #2] SFJ 0.08 — retrograde again",
             "probe_move", _pm(sid, "SFJ", 0.08, "anterior-medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #3 — huge jump to ankle instead of systematic thigh scan
        Step("P06 [IGNORE #3] ankle medial — giant jump, skips thigh entirely",
             "probe_move", _pm(sid, "ANKLE", 0.85, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P07 ankle distal",
             "probe_move", _pm(sid, "ANKLE", 0.90, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #4 — jumps all the way back to upper thigh from ankle
        Step("P08 [IGNORE #4] upper thigh — jumps back from ankle",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.14, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P09 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.25, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #5 — retrograde from Hunterian back to SFJ
        Step("P10 [IGNORE #5] SFJ 0.05 — retrograde from Hunterian",
             "probe_move", _pm(sid, "SFJ", 0.05, "anterior-medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #6 — huge jump to calf instead of systematic scan
        Step("P11 [IGNORE #6] calf medial — ignores Hunterian guidance",
             "probe_move", _pm(sid, "CALF", 0.63, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P12 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.69, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P13 popliteal — eventually covering gap",
             "probe_move", _pm(sid, "POPLITEAL", 0.48, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P14 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.53, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P15 calf posterior",
             "probe_move", _pm(sid, "CALF", 0.60, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P16 calf posterior mid",
             "probe_move", _pm(sid, "CALF", 0.67, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P17 ankle medial",
             "probe_move", _pm(sid, "ANKLE", 0.86, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P18 calf medial return",
             "probe_move", _pm(sid, "CALF", 0.74, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # max_visited = 0.90 (from P07) ≥ 0.44, ep_n1_n2=T, rp_n2_n1=T,
        # not ep_n2_n3 → Rule 6 COMPLETE
        Step("CM-2  RP N2→N1 calf — COMPLETE (max_visited=0.90 from ankle wander)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.72, "CALF", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P19 verify complete — Hunterian",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P20 verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 3 — wrong_region_type4  (24 checked steps)
# ---------------------------------------------------------------------------
# EP N1→N3 found at ankle.  Rule 4 needs RP N2→N1 in the GSV trunk.  Surgeon
# ignores all guidance pointing to trunk/thigh and instead visits SFJ, lateral
# calf, and popliteal posterior six times before eventually finding RP N2→N1
# at Hunterian.  Rule 4 fires immediately at that clip_mark.

def _wrong_region_type4(sid="adv_wrong_t4") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01 SFJ — competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 popliteal — SPJ competent check",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P03 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.53, "posterior"),
             expected_action="move"),

        Step("P04 calf posterior upper",
             "probe_move", _pm(sid, "CALF", 0.60, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P05 calf medial",
             "probe_move", _pm(sid, "CALF", 0.66, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P06 ankle medial",
             "probe_move", _pm(sid, "ANKLE", 0.78, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P07 ankle distal medial",
             "probe_move", _pm(sid, "ANKLE", 0.86, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P08 ankle lateral",
             "probe_move", _pm(sid, "ANKLE", 0.91, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        Step("CM-1  EP N1→N3 ankle lateral — perforator entry",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.88, "ANKLE", "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #1 — guidance should say "trace N3 tributary / find trunk
        #             reflux in thigh".  Surgeon goes back to ankle instead.
        Step("P09 [IGNORE #1] ankle lateral — returns to ankle, ignores N3 trace",
             "probe_move", _pm(sid, "ANKLE", 0.83, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #2 — huge jump to SFJ
        Step("P10 [IGNORE #2] SFJ — jumps to groin, completely wrong zone",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P11 upper thigh",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.14, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #3 — goes to popliteal instead of staying in thigh for RP N2→N1
        Step("P12 [IGNORE #3] popliteal posterior — wrong zone for trunk reflux",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P13 calf posterior",
             "probe_move", _pm(sid, "CALF", 0.57, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #4 — jumps back to ankle instead of continuing to trunk
        Step("P14 [IGNORE #4] ankle medial — back to ankle again",
             "probe_move", _pm(sid, "ANKLE", 0.86, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P15 calf lateral",
             "probe_move", _pm(sid, "CALF", 0.78, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #5 — lateral calf instead of medial trunk
        Step("P16 [IGNORE #5] calf lateral distal — wrong surface for trunk",
             "probe_move", _pm(sid, "CALF", 0.69, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P17 Hunterian distal — first time near trunk zone",
             "probe_move", _pm(sid, "HUNTERIAN", 0.36, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P18 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.29, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #6 — goes proximal past Hunterian to upper thigh instead of staying
        Step("P19 [IGNORE #6] upper thigh — overshoots Hunterian zone",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.17, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P20 Hunterian proximal — returning to correct zone",
             "probe_move", _pm(sid, "HUNTERIAN", 0.25, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # ep_n1_n3=T, rp_n2_n1=T → Rule 4 COMPLETE immediately (no max_visited gate)
        Step("CM-2  RP N2→N1 Hunterian — COMPLETE (Type 4, Rule 4 immediate)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.27, "HUNTERIAN", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P21 verify complete — calf",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P22 verify complete — ankle",
             "probe_move", _pm(sid, "ANKLE", 0.88, "lateral"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 4 — systematic_ignore_type6  (23 checked steps)
# ---------------------------------------------------------------------------
# EP N1→N3 at lateral calf.  Type 6 needs RP N3→N1 (tributary re-entry).
# Surgeon ignores "trace N3 distally" six times — visits SFJ, upper thigh,
# wrong surfaces.  Eventually traces lateral calf to ankle and confirms RP
# N3→N1 → Rule 5 COMPLETE (ep_n1_n3 + rp_n3_n1, no rp_n2_n1, no rp_n3_n2).

def _systematic_ignore_type6(sid="adv_ignore_t6") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01 SFJ — competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 popliteal — SPJ competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P03 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.52, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P04 calf lateral upper — searching perforators",
             "probe_move", _pm(sid, "CALF", 0.59, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P05 calf lateral mid",
             "probe_move", _pm(sid, "CALF", 0.64, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        Step("CM-1  EP N1→N3 lateral calf — perforator found",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.67, "CALF", "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #1 — guidance should say "trace N3 distally / toward ankle",
        #             surgeon jumps to SFJ instead
        Step("P06 [IGNORE #1] SFJ — jumps to groin after lateral calf finding",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #2 — stays in wrong zone
        Step("P07 [IGNORE #2] upper thigh medial — still wrong zone",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.14, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P08 Hunterian proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.23, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #3 — retrograde from Hunterian back to SFJ
        Step("P09 [IGNORE #3] SFJ 0.07 — retrograde from Hunterian",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P10 upper thigh",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.16, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #4 — jumps to wrong calf surface (medial instead of lateral)
        Step("P11 [IGNORE #4] calf medial — wrong surface, ignores lateral trace",
             "probe_move", _pm(sid, "CALF", 0.71, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P12 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.77, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #5 — goes to popliteal posterior instead of lateral calf
        Step("P13 [IGNORE #5] popliteal posterior — wrong zone/surface",
             "probe_move", _pm(sid, "POPLITEAL", 0.48, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P14 calf posterior",
             "probe_move", _pm(sid, "CALF", 0.55, "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        # IGNORE #6 — jumps to upper thigh again instead of lateral calf distal
        Step("P15 [IGNORE #6] upper thigh — another wrong-direction jump",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.18, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P16 Hunterian medial",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Surgeon finally returns to lateral calf to trace N3 distally
        Step("P17 calf lateral — back on correct surface",
             "probe_move", _pm(sid, "CALF", 0.61, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P18 calf lateral mid",
             "probe_move", _pm(sid, "CALF", 0.67, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P19 calf lateral lower",
             "probe_move", _pm(sid, "CALF", 0.74, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P20 ankle lateral approach",
             "probe_move", _pm(sid, "ANKLE", 0.85, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P21 ankle lateral distal",
             "probe_move", _pm(sid, "ANKLE", 0.91, "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        # ep_n1_n3=T, rp_n3_n1=T, rp_n2_n1=F, rp_n3_n2=F → Rule 5 COMPLETE
        Step("CM-2  RP N3→N1 ankle lateral — COMPLETE (Type 6, Rule 5)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.89, "ANKLE", "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P22 verify complete — popliteal",
             "probe_move", _pm(sid, "POPLITEAL", 0.49, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),
    ]


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

_SCENARIOS: dict[str, callable] = {
    "sfj_bypass_calf_first":      _sfj_bypass_calf_first,
    "retrograde_multiple_jumps":  _retrograde_multiple_jumps,
    "wrong_region_type4":         _wrong_region_type4,
    "systematic_ignore_type6":    _systematic_ignore_type6,
}

_SCENARIO_DESCRIPTIONS = {
    "sfj_bypass_calf_first": (
        "Surgeon skips SFJ and starts at calf/ankle (7 ignores). Eventually marks "
        "EP N1→N2 at SFJ. Chaotic path. RP N2→N1 at calf → Rule 6 complete "
        "(max_visited=0.91 from early ankle wander)."
    ),
    "retrograde_multiple_jumps": (
        "After SFJ entry, surgeon scans retrograde 3x and makes 2 ankle jumps "
        "(6 ignores). RP N2→N1 at calf → Rule 6 complete (max_visited=0.90)."
    ),
    "wrong_region_type4": (
        "EP N1→N3 at ankle. Surgeon ignores trunk-reflux guidance 6x, visits "
        "SFJ/popliteal/lateral calf. RP N2→N1 at Hunterian → Rule 4 immediate complete."
    ),
    "systematic_ignore_type6": (
        "EP N1→N3 at lateral calf. Surgeon ignores N3-trace guidance 6x, zigzags "
        "through SFJ/thigh/medial zones. Eventually lateral calf→ankle trace. "
        "RP N3→N1 → Rule 5 complete."
    ),
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _describe_movement(event: str, data: dict) -> str:
    if event == "stream_start":
        return f"Start session  (id: {data.get('session_id', '?')})"
    if event == "probe_move":
        ignore_tag = " *** ADVERSARIAL ***" if "IGNORE" in str(data) else ""
        return (f"Move probe → {data.get('region', '?')}\n"
                f"posY={data.get('pos_y_ratio', 0.0):.2f}  "
                f"{data.get('surface', '?')}  {data.get('leg', 'right')} leg")
    if event == "clip_mark":
        base = (f"Mark: {data.get('flow', '?')} "
                f"{data.get('from_type', '?')}→{data.get('to_type', '?')}  "
                f"posY={data.get('pos_y_ratio', 0.0):.2f}  {data.get('region', '?')}  "
                f"{data.get('leg', 'right')} leg")
        if data.get("elimination_test"):
            base += f"\nElim: {data['elimination_test']}"
        return base
    return event


def _describe_expectation(step: Step) -> str:
    if step.event == "stream_start":
        return "session_ready event"
    parts = []
    if step.expected_action:
        parts.append(f'Action = "{step.expected_action}"')
    if step.forbidden_action:
        parts.append(f'Must NOT = "{step.forbidden_action}"')
    if step.guidance_must_contain:
        parts.append("Guidance contains any of: " + ", ".join(step.guidance_must_contain))
    return "\n".join(parts) if parts else "observe only"


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

CYAN   = "\033[96m"
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

_SKIP_EVENTS = {"stream_start"}


class ScenarioRunner:
    def __init__(self, api_base: str):
        self.api_base = api_base.rstrip("/")
        self._sio = socketio.SimpleClient()

    def connect(self) -> None:
        print(f"Connecting to {self.api_base} ...")
        self._sio.connect(self.api_base, wait_timeout=10)
        print("Connected.\n")

    def disconnect(self) -> None:
        self._sio.disconnect()

    def _emit_and_wait(self, event: str, data: dict, timeout: float = 60.0):
        self._sio.emit(event, data)
        if event in _SKIP_EVENTS:
            deadline = time.time() + timeout
            while time.time() < deadline:
                try:
                    ev = self._sio.receive(timeout=2)
                    if ev and ev[0] == "session_ready":
                        return ev[1]
                except Exception:
                    pass
            return None
        deadline = time.time() + timeout
        while time.time() < deadline:
            try:
                ev = self._sio.receive(timeout=2)
                if ev and ev[0] == "guidance_update":
                    return ev[1]
            except Exception:
                pass
        return None

    def run(self, steps: list[Step]) -> list[dict]:
        results = []
        for step in steps:
            is_adversarial = "IGNORE" in step.label
            label_color = YELLOW if is_adversarial else CYAN
            print(f"\n{label_color}--- {step.label} ---{RESET}")
            print(f"  event: {step.event}")

            movement    = _describe_movement(step.event, step.data)
            expectation = _describe_expectation(step)
            response    = self._emit_and_wait(step.event, step.data)

            if step.event in _SKIP_EVENTS:
                if response:
                    print(f"  session_ready: {response}")
                    results.append({"step": step.label, "status": "SKIP",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                else:
                    print(f"  {YELLOW}No session_ready{RESET}")
                    results.append({"step": step.label, "status": "TIMEOUT",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                continue

            if (step.expected_action is None
                    and not step.guidance_must_contain
                    and step.forbidden_action is None):
                if response:
                    g = response.get("guidance", "")
                    a = response.get("action", "?")
                    print(f"  guidance: {g!r}  action: {a}")
                    results.append({"step": step.label, "status": "SKIP",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": a, "llm_guidance": g, "failures": []})
                else:
                    print(f"  {YELLOW}timeout{RESET}")
                    results.append({"step": step.label, "status": "TIMEOUT",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                continue

            if response is None:
                print(f"  {YELLOW}No response (timeout){RESET}")
                results.append({"step": step.label, "status": "TIMEOUT",
                                "movement": movement, "expectation": expectation,
                                "llm_action": "--", "llm_guidance": "--",
                                "failures": []})
                continue

            llm_guidance = response.get("guidance", "")
            llm_action   = response.get("action", "move")
            print(f"  guidance: {llm_guidance!r}")
            print(f"  action  : {llm_action}")

            passed   = True
            failures = []

            if step.expected_action and llm_action != step.expected_action:
                passed = False
                failures.append(
                    f'Action: got "{llm_action}", expected "{step.expected_action}"')

            if step.guidance_must_contain:
                if not any(kw.lower() in llm_guidance.lower()
                           for kw in step.guidance_must_contain):
                    passed = False
                    failures.append(
                        "Guidance missing any of: "
                        + ", ".join(step.guidance_must_contain))

            if step.forbidden_action and llm_action == step.forbidden_action:
                passed = False
                failures.append(
                    f'Action "{llm_action}" must NOT fire here')

            if passed:
                print(f"  {GREEN}{BOLD}PASS{RESET}")
            else:
                for f in failures:
                    print(f"  {RED}FAIL: {f}{RESET}")

            results.append({"step": step.label,
                            "status": "PASS" if passed else "FAIL",
                            "movement": movement, "expectation": expectation,
                            "llm_action": llm_action, "llm_guidance": llm_guidance,
                            "failures": failures})
            time.sleep(0.4)

        return results

    def print_summary(self, label: str, results: list[dict]) -> bool:
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        passed  = [r for r in checked if r["status"] == "PASS"]
        failed  = [r for r in checked if r["status"] == "FAIL"]
        timeout = [r for r in results if r["status"] == "TIMEOUT"]

        print(f"\n{'─'*60}")
        print(f"{BOLD}{label}{RESET}  "
              f"{GREEN}{len(passed)} passed{RESET}  "
              f"{RED}{len(failed)} failed{RESET}  "
              f"{YELLOW}{len(timeout)} timeout{RESET}  "
              f"/ {len(checked)} checked")
        for r in failed:
            print(f"  {RED}x {r['step']}{RESET}")
            for f in r.get("failures", []):
                print(f"      {f}")
        return len(failed) == 0 and len(timeout) == 0


# ---------------------------------------------------------------------------
# Word report
# ---------------------------------------------------------------------------

_STATUS_COLORS = {"PASS": "00B050", "FAIL": "FF0000",
                  "TIMEOUT": "FF8C00", "SKIP": "808080"}


def _colored_cell(cell, text: str, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    from docx.shared import RGBColor
    cell.text = text
    run = (cell.paragraphs[0].runs[0]
           if cell.paragraphs[0].runs
           else cell.paragraphs[0].add_run(text))
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, pt: int = 9, bold=False,
                   italic=False, color_rgb=None) -> None:
    cell.text = ""
    for i, line in enumerate(text.split("\n")):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run  = para.add_run(line)
        run.font.size   = __import__("docx").shared.Pt(pt)
        run.font.bold   = bold
        run.font.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb


def write_word_report(all_results: dict[str, list[dict]],
                      api_base: str, out_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin  = sec.bottom_margin = Inches(0.7)

    t = doc.add_heading("CHIVA Streaming — Adversarial Testing Report", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = m.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}    |    "
                   f"API: {api_base}    |    Scenarios: {len(all_results)}")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    note = doc.add_paragraph()
    note.add_run(
        "NOTE: Steps labelled [IGNORE #N] are adversarial — the surgeon deliberately "
        "moves to an unexpected location.  The key assertion is that the action remains "
        '"move" (no premature complete/maneuver) throughout these steps.'
    ).font.italic = True
    doc.add_paragraph()

    doc.add_heading("Overall Summary", level=1)
    stbl = doc.add_table(rows=1, cols=5)
    stbl.style = "Table Grid"
    for i, h in enumerate(["Scenario", "Description", "Passed", "Failed", "Timeouts"]):
        c = stbl.rows[0].cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True

    overall_pass = True
    for name, results in all_results.items():
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")
        row = stbl.add_row().cells
        row[0].text = name
        row[1].text = _SCENARIO_DESCRIPTIONS.get(name, "")
        row[2].text = str(p); row[3].text = str(f); row[4].text = str(t)
        if f or t:
            overall_pass = False

    doc.add_paragraph()
    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run("OVERALL: " + ("ALL PASS" if overall_pass else "FAILURES PRESENT"))
    vr.font.bold = True; vr.font.size = Pt(13)
    vr.font.color.rgb = (RGBColor(0x00, 0xB0, 0x50) if overall_pass
                         else RGBColor(0xFF, 0x00, 0x00))

    for name, results in all_results.items():
        doc.add_page_break()
        doc.add_heading(f"Scenario: {name.upper()}", level=1)
        dp = doc.add_paragraph(_SCENARIO_DESCRIPTIONS.get(name, ""))
        if dp.runs:
            dp.runs[0].font.italic = True

        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")

        sl = doc.add_paragraph()
        sl.add_run(f"Checked: {len(checked)}   |   ")
        pr = sl.add_run(f"Passed: {p}"); pr.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
        sl.add_run("   |   ")
        fr = sl.add_run(f"Failed: {f}"); fr.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        sl.add_run("   |   ")
        tr = sl.add_run(f"Timeouts: {t}"); tr.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        doc.add_paragraph()

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        for i, h in enumerate(["#", "Surgeon Action", "Expected",
                                "LLM Guidance", "Action", "Status"]):
            c = tbl.rows[0].cells[i]
            c.text = h
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.bold = True
        for i, w in enumerate([0.25, 1.80, 2.10, 2.10, 0.80, 0.65]):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)

        for idx, r in enumerate(results, 1):
            status   = r.get("status", "SKIP")
            row      = tbl.add_row().cells
            step_lbl = r.get("step", "")
            is_adv   = "IGNORE" in step_lbl
            row[0].text = str(idx)
            _set_cell_text(row[1], r.get("movement", "--"), pt=8,
                           color_rgb=(RGBColor(0xB8, 0x86, 0x0B) if is_adv else None))
            _set_cell_text(row[2], r.get("expectation", "--"), pt=8, italic=True)
            _set_cell_text(row[3], r.get("llm_guidance", "--") or "--", pt=8)
            _set_cell_text(row[4], r.get("llm_action", "--") or "--", pt=9, bold=True)
            _colored_cell(row[5], status, _STATUS_COLORS.get(status, "808080"))
            if status == "FAIL" and r.get("failures"):
                fr_row = tbl.add_row().cells
                merged = fr_row[0].merge(fr_row[5])
                merged.text = ""
                for fn in r["failures"]:
                    pp = merged.add_paragraph(f"  x  {fn}")
                    if pp.runs:
                        pp.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        pp.runs[0].font.size      = Pt(8)
                        pp.runs[0].font.italic    = True

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"\nReport saved: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="CHIVA adversarial streaming test runner")
    parser.add_argument("scenario", nargs="?", default="all",
                        choices=list(_SCENARIOS.keys()) + ["all"],
                        help="Scenario to run (default: all)")
    parser.add_argument("--api",  default="http://localhost:7861")
    parser.add_argument("--all",  action="store_true", dest="run_all")
    args = parser.parse_args()

    to_run = list(_SCENARIOS.items()) if (args.run_all or args.scenario == "all") \
             else [(args.scenario, _SCENARIOS[args.scenario])]

    runner    = ScenarioRunner(args.api)
    all_ok    = True
    all_results: dict[str, list[dict]] = {}

    try:
        runner.connect()
        for name, builder in to_run:
            print(f"\n{'='*60}")
            print(f"{BOLD}ADVERSARIAL SCENARIO: {name.upper()}{RESET}")
            print(f"{'='*60}")
            steps   = builder()
            results = runner.run(steps)
            ok      = runner.print_summary(name, results)
            all_ok  = all_ok and ok
            all_results[name] = results
    except KeyboardInterrupt:
        print("\nInterrupted.")
    finally:
        runner.disconnect()

    if len(to_run) > 1:
        print(f"\n{'='*60}")
        print(f"{BOLD}OVERALL: {'ALL PASS' if all_ok else 'FAILURES PRESENT'}{RESET}")

    if all_results:
        ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
        here     = os.path.dirname(os.path.abspath(__file__))
        out_path = os.path.join(here, "results",
                                f"adversarial_log_{ts}.docx")
        try:
            write_word_report(all_results, args.api, out_path)
        except Exception as exc:
            print(f"{YELLOW}Warning: could not write Word report — {exc}{RESET}")

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
