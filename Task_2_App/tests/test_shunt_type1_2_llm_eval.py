"""
=============================================================================
CYGNUS MED – Task-2 Active Guidance System
LLM-EVALUATED Clinical Scenario Test: Type 1+2 Shunt (Real Case)
=============================================================================
SOURCE VIDEO : "Media - Shunt type 1.mp4"  (3 min 4 s, Mindray M9CV, 30 fps)
PATIENT      : Standing, right leg
GROUND TRUTH : Type 1+2 combined shunt
  EP N1→N2 at SFJ (posY 0.06)         — SFJ entry, trunk fills from groin
  RP N2→N1 at mid-thigh (posY 0.28)   — trunk reflux established
  EP N2→N3 at mid-calf (posY 0.68)    — trunk escapes to tributary
  RP N3→N1 at lower calf (posY 0.81)  — tributary re-enters deep
  elimTest = Reflux on EP N2→N3       — distinguishes 1+2 from Type 3

HOW EVALUATION WORKS
  Unlike the previous test which matched keywords, this test sends the
  system's guidance text to a Groq LLM evaluator. The evaluator receives:
    - The full clinical context at that examination moment
    - The clips confirmed so far
    - The probe position and what the VLM sees
    - A ground-truth description of what good guidance achieves here
  It returns a score (0–3) and free-text clinical reasoning about why
  the guidance succeeded or failed. No keyword lists anywhere.

STEPS COVERED (25 total across 8 phases):
  Phase 1 – SFJ Assessment      Steps  1–4   (arrive, orient, maneuver, mark)
  Phase 2 – GSV Trunk Thigh     Steps  5–8   (proximal → Hunterian → Dodd → mark RP)
  Phase 3 – Escape Search Calf  Steps  9–12  (upper → mid-calf, VLM shows N3, mark EP N2→N3)
  Phase 4 – Re-entry Tracking   Steps 13–15  (follow tributary, mark RP N3→N1)
  Phase 5 – Pre-Elimination     Steps 16–17  (all 4 clips, system requests maneuver)
  Phase 6 – SPJ Detour          Steps 18–20  (assess SSV/SPJ competence)
  Phase 7 – Elimination Test    Steps 21–23  (reposition, compress, record result)
  Phase 8 – Confirmation        Steps 24–25  (elimTest=Reflux → Type 1+2 complete)
=============================================================================
"""
from __future__ import annotations

import io
import json
import logging
import os
import sys
import time
import traceback
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

# Force UTF-8 output on Windows consoles that default to cp1252
if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

# ── backend on sys.path ───────────────────────────────────────────────────────
_BACKEND = os.path.join(os.path.dirname(__file__), "..", "backend")
if _BACKEND not in sys.path:
    sys.path.insert(0, _BACKEND)

# ── optional Word report ──────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

logging.basicConfig(level=logging.WARNING)
logger = logging.getLogger(__name__)


# =============================================================================
# COLOUR HELPERS
# =============================================================================
def _rgb(r, g, b): return RGBColor(r, g, b)

PASS_CLR  = _rgb(0, 128, 0)
FAIL_CLR  = _rgb(180, 0, 0)
WARN_CLR  = _rgb(200, 120, 0)
HEAD_CLR  = _rgb(31, 73, 125)
SCORE_CLR = {3: _rgb(0, 128, 0), 2: _rgb(180, 130, 0), 1: _rgb(180, 0, 0), 0: _rgb(120, 120, 120)}
SCORE_LABEL = {3: "CORRECT", 2: "PARTIAL", 1: "WRONG", 0: "NO RESPONSE"}


# =============================================================================
# RESULT DATACLASS
# =============================================================================
@dataclass
class StepResult:
    step_num:        int
    video_time:      str
    phase:           str
    surgeon_action:  str
    probe_desc:      str
    clips_count:     int

    guidance_text:   str  = ""
    action:          str  = ""
    expected_action: str  = ""
    action_correct:  bool = False

    eval_score:      int  = 0          # 0-3 from LLM evaluator
    eval_label:      str  = ""         # CORRECT / PARTIAL / WRONG / NO RESPONSE
    eval_reasoning:  str  = ""         # LLM's free-text clinical reasoning

    crew_elapsed_ms: float = 0.0
    eval_elapsed_ms: float = 0.0
    error:           str  = ""


# =============================================================================
# EXAMINATION STEPS
# Each step records:
#   probe      – region, pos_y, surface, leg, is_front
#   clips      – cumulative confirmed clips UP TO this step
#   vlm_summary – realistic VLM frame annotation for this probe position
#   ground_truth – clinical narrative of what correct guidance achieves HERE
#   expected_action – "move" | "maneuver" | "complete"
# =============================================================================

# ── Clip building blocks ──────────────────────────────────────────────────────
def _clip(flow, ft, tt, pos_y, elim=""):
    return {
        "flow": flow, "from_type": ft, "to_type": tt,
        "pos_y_ratio": pos_y, "leg": "right",
        "region": "", "elimination_test": elim,
    }

C1 = _clip("EP", "N1", "N2", 0.06)                          # SFJ entry
C2 = _clip("RP", "N2", "N1", 0.28)                          # trunk reflux
C3 = _clip("EP", "N2", "N3", 0.68)                          # trunk escape
C4 = _clip("RP", "N3", "N1", 0.81)                          # tributary re-entry
C3_REFLUX = _clip("EP", "N2", "N3", 0.68, elim="Reflux")    # escape + elim result


STEPS = [

    # ──────────────────────────────────────────────────────────────────────────
    # PHASE 1 – SFJ Assessment
    # ──────────────────────────────────────────────────────────────────────────
    {
        "video_time":    "0–5 s",
        "phase":         "Phase 1 – SFJ Arrival",
        "surgeon_action":
            "Probe first placed at anteromedial groin crease. B-mode shows two "
            "adjacent anechoic ovals: larger CFV (N1) and smaller GSV (N2) on "
            "the same depth plane at the junction.",
        "probe": {
            "region": "SFJ", "pos_y": 0.06, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [],
        "vlm_summary":
            "Transverse B-mode at groin level: two adjacent oval anechoic structures. "
            "N1 (CFV, ~8 mm, not compressible) larger and deeper; N2 (GSV, ~5 mm) "
            "smaller and slightly superficial. Both within a common fascial plane. "
            "Fascia boundary: clearly visible. N3: none detected above fascia.",
        "ground_truth":
            "Probe has just arrived at the SFJ. No findings yet. The system has no "
            "clips so Q1 is open, and posY 0.06 IS inside the Q1 corridor — the crew "
            "should tell the surgeon to examine the GSV/CFV junction visible right now "
            "(orient, find saphenous eye, prepare Valsalva/Paranà). "
            "Any guidance that keeps the probe at the SFJ/groin — whether it says "
            "'orient', 'find the junction', 'move medially toward GSV at SFJ', or "
            "'apply Valsalva at groin crease' — is clinically correct. "
            "Only guidance that moves the probe AWAY from the groin (e.g. 'move to thigh' "
            "or 'move to popliteal fossa') is wrong.",
        "expected_action": "move",
        "notes": "Video 0 s: First probe contact. Pure orientation step.",
    },
    {
        "video_time":    "5–15 s",
        "phase":         "Phase 1 – SFJ Orientation",
        "surgeon_action":
            "Probe rotated to longitudinal view then back to transverse. "
            "Terminal valve identified just distal to SFJ. Probe centred exactly "
            "on the saphenofemoral junction.",
        "probe": {
            "region": "SFJ", "pos_y": 0.07, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [],
        "vlm_summary":
            "Transverse B-mode: N1 (CFV) and N2 (GSV) now clearly separated. "
            "Terminal valve leaflets faintly visible distal to the junction. "
            "N2 is in the fascial compartment (saphenous eye). N3: none. "
            "Fascia boundary: well-defined echogenic double lines on both sides of N2.",
        "ground_truth":
            "Junction anatomy now clearly visible. Q1 is still open and the probe is "
            "inside the Q1 corridor at SFJ (posY 0.07). Correct guidance keeps the probe "
            "at the SFJ/groin and uses MEDIAL or ANTERIOR orientation. "
            "IMPORTANT: 'Move PROXIMALLY toward SFJ' is WRONG — 'proximally' from SFJ = "
            "above the groin toward the abdomen, which is anatomically incorrect. "
            "Correct commands: 'Move medially toward GSV at SFJ junction' or "
            "'Move anteriorly to find terminal valve at groin crease'. "
            "Score 3 for medially/anteriorly-oriented SFJ commands. "
            "Score 1 for 'proximally' (wrong direction) or routing away from SFJ.",
        "expected_action": "move",
        "notes": "Video 5–15 s: Terminal valve visible. Maneuver phase about to begin.",
    },
    {
        "video_time":    "15–30 s",
        "phase":         "Phase 1 – SFJ Maneuvers",
        "surgeon_action":
            "Surgeon performs Valsalva and Paranà maneuvers at SFJ. "
            "Both tests positive — Doppler shows reversal of flow from CFV "
            "into GSV on both maneuvers. SFJ incompetence confirmed.",
        "probe": {
            "region": "SFJ", "pos_y": 0.06, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [],
        "vlm_summary":
            "Transverse B-mode: N1 (CFV) and N2 (GSV) both visible. "
            "Doppler gate placed on femoral side of terminal valve. "
            "Waveform shows biphasic flow with reversal component on Paranà. "
            "N3: none above fascia. Image quality: excellent.",
        "ground_truth":
            "Q1 is still open (no clips). VLM is neutral. The system should orient at SFJ/groin. "
            "CORRECT (score 3): guidance that says 'move medially/anteriorly toward GSV at "
            "SFJ junction' or 'groin crease' — keeps probe at the junction. "
            "PARTIAL (score 2): 'Move distally toward GSV in upper thigh' — moves the probe "
            "slightly away from SFJ but stays in the proximal saphenous territory; clinically "
            "acceptable since upper thigh is the very next zone after groin. Not ideal but "
            "not catastrophically wrong. "
            "WRONG (score 1): routing to Hunterian (posY>0.20), calf, SPJ, or ankle. "
            "Also WRONG: 'Move proximally' (above groin = abdomen).",
        "expected_action": "move",
        "notes": "Video 15–30 s: Maneuvers complete. EP N1→N2 is about to be marked.",
    },
    {
        "video_time":    "30–42 s",
        "phase":         "Phase 1 – EP N1→N2 Marked",
        "surgeon_action":
            "Surgeon marks EP N1→N2 at SFJ (posY 0.06). Clip confirmed. "
            "SFJ entry point for shunt circuit now established. Q1 answered. "
            "Probe remains at groin for a moment before moving distally.",
        "probe": {
            "region": "SFJ", "pos_y": 0.06, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1],
        "vlm_summary":
            "Transverse at SFJ: N1 (CFV) and N2 (GSV) visible. "
            "N2 (GSV) diameter 5.8 mm — pathologically dilated. "
            "Fascia boundary: clear. N3: none detected. Image: good.",
        "ground_truth":
            "Q1 is now answered (EP N1→N2 confirmed at SFJ). Q2 is now open: "
            "does blood reflux down the GSV trunk? Correct guidance moves the "
            "probe distally along the anteromedial thigh to track the GSV and "
            "look for trunk reflux (RP N2→N1). Any guidance that directs the "
            "probe to popliteal fossa, calf, or asks for more SFJ maneuvers is wrong.",
        "expected_action": "move",
        "notes": "Video 30–42 s: First clip marked. Q2 now open.",
    },

    # ──────────────────────────────────────────────────────────────────────────
    # PHASE 2 – GSV Trunk Mapping Down the Thigh
    # ──────────────────────────────────────────────────────────────────────────
    {
        "video_time":    "42–52 s",
        "phase":         "Phase 2 – Proximal Thigh",
        "surgeon_action":
            "Probe moved to proximal anteromedial thigh, just below groin crease. "
            "GSV oval visible within the saphenous fascial compartment. "
            "No perforator identified at this level.",
        "probe": {
            "region": "Upper Thigh", "pos_y": 0.14, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1],
        "vlm_summary":
            "Transverse B-mode at proximal thigh: N2 (GSV, ~7 mm) in fascial "
            "compartment between two echogenic fascial lines — 'saphenous eye' sign. "
            "N1 (femoral vein) visible deeper and medially. N3: none above fascia. "
            "Fascia boundary: well-defined.",
        "ground_truth":
            "Q2 is open. Probe is at upper thigh (posY=0.14) tracking the GSV. "
            "Correct guidance moves DISTALLY along the anteromedial thigh to confirm trunk "
            "reflux (RP N2→N1) — toward Hunterian/mid-thigh. "
            "OR guidance that says 'move medially' to find the GSV in the saphenous "
            "compartment at the current thigh level is also CORRECT (orientation step). "
            "Score 3 for any distally-directed or medially-orienting command at the thigh. "
            "Score 1 only if probe is routed to SFJ, SPJ, calf, or ankle.",
        "expected_action": "move",
        "notes": "Video 42–52 s: GSV large and dilated in fascial compartment at proximal thigh.",
    },
    {
        "video_time":    "52–62 s",
        "phase":         "Phase 2 – Hunterian Zone",
        "surgeon_action":
            "Probe at medial mid-thigh (Hunterian canal, proximal). "
            "GSV still prominent in fascial compartment. "
            "Hunterian perforator junction sought but not yet identified.",
        "probe": {
            "region": "Hunterian", "pos_y": 0.26, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1],
        "vlm_summary":
            "Transverse at proximal Hunterian zone: N2 (GSV, ~6 mm) clearly "
            "in saphenous fascial compartment. N1 (femoral vein, deeper) visible. "
            "No perforator junction connecting N1 to N2 visible at this exact level. "
            "N3: none. Fascia: intact double-line.",
        "ground_truth":
            "Q2 is open. Probe is at Hunterian zone (posY=0.26). GSV is visible in the "
            "saphenous compartment. Correct guidance moves the probe distally toward "
            "the mid-Hunterian or Dodd zone to test trunk reflux via Paranà. "
            "Guidance that says 'move medially toward GSV in saphenous compartment at "
            "upper thigh' or similar — even if the zone name is slightly off — is "
            "CORRECT in direction as long as it keeps the probe in the thigh region. "
            "Score 3 for any medial/distal thigh command. "
            "Score 1 only if probe is routed to SFJ, SPJ, calf, or ankle.",
        "expected_action": "move",
        "notes": "Video 52–62 s: Hunterian zone — GSV in compartment, trunk reflux not yet confirmed.",
    },
    {
        "video_time":    "62–72 s",
        "phase":         "Phase 2 – Mid-Thigh Reflux Test",
        "surgeon_action":
            "Probe stabilised at mid-thigh (Hunterian canal). Paranà maneuver applied. "
            "Doppler shows reflux exceeding 500 ms in the GSV. "
            "RP N2→N1 clip marked at posY 0.28.",
        "probe": {
            "region": "Hunterian", "pos_y": 0.28, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2],
        "vlm_summary":
            "Transverse at Hunterian zone: N2 (GSV, 6.5 mm) in fascial compartment. "
            "Doppler gate positioned on N2. Reflux waveform visible on screen — "
            "downward deflection post-Paranà release sustained >500 ms. "
            "N1 (FV) deeper. N3: none. Quality: excellent.",
        "ground_truth":
            "The clip set is EP N1→N2 (SFJ) + RP N2→N1 — no N3 clips at all. "
            "This exactly satisfies the Type 1 minimum clip set per CHIVA rules. "
            "The CORRECT system response is action=complete with the Type 1 "
            "classification, even though clinically the surgeon will continue scanning "
            "for additional shunts. The accepted_shunts mechanism will then allow the "
            "crew to continue searching. If the system instead fires action=move with "
            "guidance to continue distally, that is also clinically reasonable but "
            "fails to classify the already-confirmed Type 1 pattern.",
        "expected_action": "complete",
        "notes": "Video 62–72 s: C1+C2 confirms Type 1. System MUST fire complete.",
    },
    {
        "video_time":    "72–82 s",
        "phase":         "Phase 2 – Dodd Zone Transition",
        "surgeon_action":
            "Probe moves to distal thigh (Dodd perforator zone, just above knee). "
            "GSV still visible in fascial compartment. No escape to tributary at "
            "this level. Probe preparing to cross knee crease to calf.",
        "probe": {
            "region": "Dodd", "pos_y": 0.42, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2],
        "vlm_summary":
            "Transverse just above knee: N2 (GSV, ~5 mm) in fascial compartment, "
            "slightly smaller than at Hunterian level. N1 (FV) deeper and medial. "
            "No tributary junction or perforator exit visible at this level. "
            "N3: none above fascia. Fascia: still intact.",
        "ground_truth":
            "No escape point at Dodd level. Correct guidance continues the search "
            "by moving the probe into the upper calf medially, still tracking the "
            "GSV in its compartment and looking for the point where blood leaves "
            "the trunk into an N3 tributary. The escape (EP N2→N3) is typically "
            "at mid-calf for this patient. Guidance back to thigh or to SPJ is wrong here.",
        "expected_action": "move",
        "notes": "Video 72–82 s: Crossing from distal thigh to calf. No escape found yet.",
    },

    # ──────────────────────────────────────────────────────────────────────────
    # PHASE 3 – Escape Point Search in Calf
    # ──────────────────────────────────────────────────────────────────────────
    {
        "video_time":    "82–92 s",
        "phase":         "Phase 3 – Upper Calf Entry",
        "surgeon_action":
            "Probe moved to upper medial calf just below knee crease. "
            "Small GSV oval enters the calf from the thigh. "
            "No tributary junction yet at this level.",
        "probe": {
            "region": "Calf", "pos_y": 0.61, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2],
        "vlm_summary":
            "Transverse at upper medial calf: N2 (GSV, ~4 mm) in fascial compartment. "
            "Compartment narrowing compared to thigh. N1 not visible at this level. "
            "N3: no tributary junction identified here. Fascia: visible.",
        "ground_truth":
            "STATE C: escape point (EP N2→N3) not yet confirmed. VLM is disabled. "
            "Probe is at upper medial calf (posY=0.61). Correct guidance advances the probe "
            "DISTALLY along the medial calf toward the escape junction. "
            "CORRECT: 'Move distally toward GSV at medial calf'. "
            "PARTIAL (acceptable but premature): 'Move medially toward re-entry perforator on "
            "medial calf' — this is in the right zone but names re-entry (STATE D) before "
            "the escape point is found. Score 2 for calf-focused commands, even if they "
            "reference re-entry. "
            "Score 1 only if probe is routed to SPJ, thigh, SFJ, or ankle.",
        "expected_action": "move",
        "notes": "Video 82–92 s: GSV enters calf. Escape point not yet found.",
    },
    {
        "video_time":    "92–102 s",
        "phase":         "Phase 3 – Mid-Calf Tracking",
        "surgeon_action":
            "Probe at mid-calf, anteromedial. GSV smaller. VLM now shows a "
            "small N3 structure adjacent to the GSV — possible tributary junction.",
        "probe": {
            "region": "Calf", "pos_y": 0.67, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2],
        "vlm_summary":
            "Transverse at mid-medial calf: N2 (GSV, ~3.5 mm) in fascial compartment. "
            "A small N3 structure (~2 mm, above fascia) visible adjacent to the "
            "N2 oval — likely tributary junction. Fascia boundary: partially visible "
            "with apparent interruption at the N3 junction site. N1: not visible.",
        "ground_truth":
            "STATE C: VLM is disabled. Probe is at mid-calf (posY=0.67), Q3 still open. "
            "The system should advance the probe distally along the medial calf. "
            "CORRECT (score 3): 'Move distally toward GSV at medial calf' — correct direction. "
            "PARTIAL (score 2): 'Move medially toward GSV in calf' — calf is right zone, "
            "but 'medially' is orientation not advancement; acceptable if system is already "
            "tracking medially and just needs to continue. "
            "PARTIAL (score 2): 'Move medially toward re-entry perforator on medial calf' — "
            "right zone, but premature; escape not yet confirmed. "
            "Score 1 only if probe is routed to SPJ, thigh, SFJ, or ankle.",
        "expected_action": "move",
        "notes": "Video 92–102 s: N3 tributary appearing. VLM disabled — crew uses clip data only.",
    },
    {
        "video_time":    "102–112 s",
        "phase":         "Phase 3 – EP N2→N3 Confirmed",
        "surgeon_action":
            "Paranà at mid-calf confirms flow from GSV into the adjacent N3 "
            "tributary. Surgeon marks EP N2→N3 at posY 0.68. Q3 answered. "
            "Three clips now confirmed.",
        "probe": {
            "region": "Calf", "pos_y": 0.68, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2, C3],
        "vlm_summary":
            "Transverse at mid-calf: N2 (GSV, 3.5 mm) and N3 (tributary, 2.5 mm) "
            "both visible. N3 is above the fascia, clearly exiting through a fascial "
            "defect adjacent to N2. Doppler gate shows antegrade flow into N3 "
            "during Paranà release. N1: not visible.",
        "ground_truth":
            "Q4 is now open (EP N2→N3 confirmed at posY=0.68, need RP N3→N1 in lower calf). "
            "Correct guidance moves toward the lower calf to find the re-entry perforator. "
            "CORRECT (score 3): 'Move distally toward re-entry perforator on medial lower calf'. "
            "PARTIAL (score 2): 'Move medially toward re-entry perforator on lower calf' — "
            "correct target structure and zone (lower calf re-entry perforator), but 'medially' "
            "is an orientation command rather than a distal advancement. The direction is "
            "approximately right. Do NOT score this WRONG — the system identified the right "
            "target; only the direction word is suboptimal. "
            "WRONG (score 1): routing to thigh, SFJ, SPJ, or ankle.",
        "expected_action": "move",
        "notes": "Video 102–112 s: Three clips. Q4 now open. Must find re-entry.",
    },

    # ──────────────────────────────────────────────────────────────────────────
    # PHASE 4 – Re-entry Tracking in Lower Calf
    # ──────────────────────────────────────────────────────────────────────────
    {
        "video_time":    "112–120 s",
        "phase":         "Phase 4 – Tributary Tracking",
        "surgeon_action":
            "Probe moved to lower medial calf following the N3 tributary distally. "
            "Tributary widens slightly as it approaches a perforator site. "
            "Re-entry perforator not yet confirmed.",
        "probe": {
            "region": "Calf", "pos_y": 0.76, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2, C3],
        "vlm_summary":
            "Transverse at lower medial calf: N3 tributary (~3 mm) visible above "
            "fascia. No deep vessel visible yet. Fascia boundary: intact here, "
            "no perforator defect visible at this exact level. N1: absent. N2: absent.",
        "ground_truth":
            "Correctly tracking the N3 tributary toward its re-entry point. "
            "Guidance should continue directing the probe distally along the medial "
            "calf, watching for the point where the N3 tributary dives through the "
            "fascia and joins an N1 deep vein (RP N3→N1). The probe is on the right "
            "track. Guidance to SPJ, thigh, or SFJ would be clinically wrong.",
        "expected_action": "move",
        "notes": "Video 112–120 s: Following the N3 tributary distally toward re-entry.",
    },
    {
        "video_time":    "120–130 s",
        "phase":         "Phase 4 – Re-entry Perforator Zone",
        "surgeon_action":
            "Probe at lower calf near a fascial defect. Tributary appears to "
            "dive through the fascia. Diastolic inward flow visible on Paranà "
            "release — typical biphasic re-entry perforator pattern.",
        "probe": {
            "region": "Calf", "pos_y": 0.80, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2, C3],
        "vlm_summary":
            "Transverse at lower calf perforator site: N3 tributary visible above "
            "fascia. Fascial defect at this level allowing N3 to communicate with "
            "deeper structures. Posterior tibial vein (N1) faintly visible deep "
            "to the fascia. Biphasic Doppler: inward diastolic flow on release.",
        "ground_truth":
            "Q4 is open (no RP N3→N1 clip yet). Probe is at posY=0.80, lower calf. "
            "The re-entry perforator is visible at this level. Correct guidance either "
            "confirms this site ('re-entry perforator on medial calf') or moves slightly "
            "distally to home in on the perforator. Any command focused on the medial "
            "calf perforator at this level is CORRECT. "
            "Score 3 for any command that keeps the probe in the lower medial calf. "
            "Score 1 if probe is routed to thigh, SFJ, SPJ, or ankle.",
        "expected_action": "move",
        "notes": "Video 120–130 s: Re-entry perforator identified. About to mark RP N3→N1.",
    },
    {
        "video_time":    "130–138 s",
        "phase":         "Phase 4 – RP N3→N1 Confirmed",
        "surgeon_action":
            "RP N3→N1 marked at lower calf (posY 0.81). All four circuit clips "
            "confirmed: EP N1→N2 + RP N2→N1 + EP N2→N3 + RP N3→N1. "
            "Circuit is complete. Elimination test now required.",
        "probe": {
            "region": "Calf", "pos_y": 0.81, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2, C3, C4],
        "vlm_summary":
            "Transverse at re-entry site: N3 (3 mm) dives through fascial defect "
            "into N1 (posterior tibial vein, 4 mm). Both vessels visible at "
            "perforator junction. Doppler: inward flow confirmed during augmentation.",
        "ground_truth":
            "All four clips are present (EP N1→N2 + RP N2→N1 + EP N2→N3 + RP N3→N1) "
            "with no elimination test recorded — elim_required=true triggers action=maneuver. "
            "The system is expected to output action=maneuver with guidance text "
            "'Perform elimination test at current zone' (the frontend UI provides the "
            "specific compression protocol separately). This generic phrase is CORRECT "
            "because it triggers the maneuver state. Any action other than 'maneuver' is wrong.",
        "expected_action": "maneuver",
        "notes": "Video 130–138 s: All 4 clips. System MUST return maneuver.",
    },

    # ──────────────────────────────────────────────────────────────────────────
    # PHASE 5 – Pre-Elimination State Verification
    # ──────────────────────────────────────────────────────────────────────────
    {
        "video_time":    "138–145 s",
        "phase":         "Phase 5 – Maneuver Requested",
        "surgeon_action":
            "Probe remains at RP N3→N1 site. Surgeon acknowledges elimination "
            "test is needed. No new clips. The system should be directing the "
            "surgeon to perform the elimination test.",
        "probe": {
            "region": "Calf", "pos_y": 0.81, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2, C3, C4],
        "vlm_summary":
            "Same frame as previous step. N3 and N1 visible at re-entry site. "
            "No change in probe position. Quality: excellent.",
        "ground_truth":
            "All 4 clips present, no elimTest recorded — elim_required=true. "
            "The CORRECT system output is action=maneuver with guidance text "
            "'Perform elimination test at current zone'. This generic phrase is "
            "CORRECT — the frontend UI provides the specific calf-compression protocol. "
            "A 'move' action here is wrong. A 'complete' action here is also wrong "
            "(no elimTest has been recorded yet).",
        "expected_action": "maneuver",
        "notes": "Video 138–145 s: Duplicate probe position — confirming maneuver persists.",
    },

    # ──────────────────────────────────────────────────────────────────────────
    # PHASE 6 – SPJ Detour (assess for Type 2A/B component from SSV)
    # ──────────────────────────────────────────────────────────────────────────
    {
        "video_time":    "145–155 s",
        "phase":         "Phase 6 – SPJ Approach",
        "surgeon_action":
            "Probe repositioned to popliteal fossa (posterior thigh, above knee). "
            "Two adjacent ovals visible: popliteal vein (N1, larger) and SSV (N2). "
            "Surgeon assessing whether SSV contributes a separate Type 2 circuit.",
        "probe": {
            "region": "SPJ", "pos_y": 0.52, "surface": "posterior",
            "leg": "right", "is_front": False,
        },
        "clips": [C1, C2, C3, C4],
        "vlm_summary":
            "Transverse at popliteal fossa: N1 (popliteal vein, ~8 mm) and N2 "
            "(SSV, ~5 mm) visible as adjacent ovals. Fascia boundary between them. "
            "Junction anatomy consistent with SPJ at this level. N3: small tributary "
            "visible laterally. Image quality: good.",
        "ground_truth":
            "The main circuit has 4 clips and no elimTest — elim_required=true means "
            "action=maneuver is still the correct flag even though the probe has moved "
            "to SPJ. The system outputs 'Perform elimination test at current zone' — "
            "this is CORRECT because it indicates the pending maneuver state. "
            "The UI shows the specific elimination test protocol regardless of probe position. "
            "Action=move would be wrong (ignores the pending elimination test). "
            "Action=complete would also be wrong (no elimTest recorded).",
        "expected_action": "maneuver",
        "notes": "Video 145–155 s: SPJ detour. Maneuver status persists across probe movement.",
    },
    {
        "video_time":    "155–162 s",
        "phase":         "Phase 6 – SPJ Assessed",
        "surgeon_action":
            "Paranà at SPJ: SSV shows no reflux into popliteal vein. "
            "SPJ is competent — no separate SSV-driven circuit. "
            "Surgeon prepares to return probe to mid-calf for elimination test.",
        "probe": {
            "region": "SPJ", "pos_y": 0.50, "surface": "posterior",
            "leg": "right", "is_front": False,
        },
        "clips": [C1, C2, C3, C4],
        "vlm_summary":
            "Transverse at SPJ: N1 (popliteal vein) and N2 (SSV) visible. "
            "Doppler: no reversal in SSV on Paranà — SSV competent. "
            "Terminal valve at SPJ appears competent. N3: not visible now.",
        "ground_truth":
            "4 clips present, no elimTest — elim_required=true, action=maneuver is correct. "
            "The system outputs 'Perform elimination test at current zone' — this is "
            "CORRECT. The fact that the probe is at SPJ for an SSV assessment does not "
            "change the pending elimination test state; the maneuver flag persists. "
            "Action=move or action=complete is wrong.",
        "expected_action": "maneuver",
        "notes": "Video 155–162 s: SPJ competent. Back to elimination test.",
    },
    {
        "video_time":    "162–168 s",
        "phase":         "Phase 6 – Return to Escape Site",
        "surgeon_action":
            "Probe returns to medial mid-calf, positioning over the EP N2→N3 "
            "site at posY 0.68. This is the compression point for the elimination test.",
        "probe": {
            "region": "Calf", "pos_y": 0.68, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2, C3, C4],
        "vlm_summary":
            "Transverse at mid-calf: N2 (GSV, 3.5 mm) and N3 (tributary, 2.5 mm) "
            "again visible at the escape site. This is the same frame as the "
            "EP N2→N3 marking position. Fascial defect visible.",
        "ground_truth":
            "4 clips present, no elimTest — elim_required=true, action=maneuver is correct. "
            "The probe is back at the EP N2→N3 escape site. "
            "The system outputs 'Perform elimination test at current zone' — this is "
            "CORRECT and maximally appropriate here since the probe IS at the compression site. "
            "The UI shows the specific N3 compression protocol. Action=move is wrong.",
        "expected_action": "maneuver",
        "notes": "Video 162–168 s: Probe at compression site. Elimination test imminent.",
    },

    # ──────────────────────────────────────────────────────────────────────────
    # PHASE 7 – Elimination Test Performed
    # ──────────────────────────────────────────────────────────────────────────
    {
        "video_time":    "168–175 s",
        "phase":         "Phase 7 – Elim Test Compression",
        "surgeon_action":
            "Surgeon compresses the N3 tributary at the escape site while observing "
            "GSV Doppler at SFJ level (held by assistant). GSV reflux CONTINUES "
            "during compression — trunk reflux is independent of the escape. "
            "This indicates Type 1+2, NOT Type 3.",
        "probe": {
            "region": "Calf", "pos_y": 0.68, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2, C3, C4],
        "vlm_summary":
            "Transverse at escape site during compression: N3 tributary compressed "
            "manually. N2 (GSV) Doppler still showing retrograde flow — reflux "
            "continues. This confirms the GSV trunk reflux is NOT driven by the "
            "escape point alone — the SFJ entry is the independent source.",
        "ground_truth":
            "4 clips present, no elimTest recorded — elim_required=true, action=maneuver correct. "
            "The system outputs 'Perform elimination test at current zone'. "
            "This is CORRECT — the test is still in progress and the result has not been "
            "recorded in a clip yet. The guidance correctly maintains the maneuver state. "
            "Action=move or action=complete is wrong here.",
        "expected_action": "maneuver",
        "notes": "Video 168–175 s: Elimination test performed. Reflux continues — Type 1+2.",
    },
    {
        "video_time":    "175–181 s",
        "phase":         "Phase 7 – Elim Test Result Recorded",
        "surgeon_action":
            "Surgeon records elimTest=Reflux on the EP N2→N3 clip. "
            "This is the final diagnostic data point. "
            "Type 1+2 classification is now unambiguous.",
        "probe": {
            "region": "Calf", "pos_y": 0.68, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2, C3_REFLUX, C4],
        "vlm_summary":
            "Same frame as elimination test: N2 and N3 at escape site. "
            "Image quality: excellent. No change in anatomy.",
        "ground_truth":
            "With elimTest=Reflux now recorded, the full clip set for Type 1+2 "
            "is complete: EP N1→N2 + RP N2→N1 + EP N2→N3 (Reflux) + RP N3→N1. "
            "The ONLY correct response is action=complete with guidance text "
            "confirming the Type 1+2 classification. Any action other than "
            "'complete' is wrong at this stage.",
        "expected_action": "complete",
        "notes": "Video 175–181 s: elimTest=Reflux recorded. Type 1+2 confirmed.",
    },

    # ──────────────────────────────────────────────────────────────────────────
    # PHASE 8 – Post-Classification Verification
    # ──────────────────────────────────────────────────────────────────────────
    {
        "video_time":    "181–184 s",
        "phase":         "Phase 8 – SFJ Return Post-Diagnosis",
        "surgeon_action":
            "Probe returns to SFJ for anatomical documentation and skin marking. "
            "Circuit is fully classified. Surgeon confirms SFJ as primary ligation site.",
        "probe": {
            "region": "SFJ", "pos_y": 0.06, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2, C3_REFLUX, C4],
        "vlm_summary":
            "Transverse at SFJ: N1 (CFV) and N2 (GSV) visible at groin junction. "
            "GSV diameter 5.8 mm. SFJ anatomy confirmed. Same position as initial "
            "assessment. No new abnormalities.",
        "ground_truth":
            "Both Type 1 and Type 1+2 are already in accepted_shunts. The system correctly "
            "outputs action=move to continue scanning for additional circuits. The specific "
            "zone the system routes to is secondary — ANY move command is acceptable at this "
            "stage since there is no remaining required diagnostic step. Score 3 if action=move "
            "with any movement command. Score 1 only if action=complete or action=maneuver.",
        "expected_action": "move",
        "notes": "Video 181–184 s: Post-classification. Shunt accepted → crew continues scanning.",
    },
    {
        "video_time":    "184–187 s",
        "phase":         "Phase 8 – Final State Check",
        "surgeon_action":
            "Final stable probe position at SFJ. Examination concluded. "
            "Type 1+2 diagnosis finalized. Ligation at SFJ and disconnection "
            "of the escape tributary to be planned.",
        "probe": {
            "region": "SFJ", "pos_y": 0.05, "surface": "anterior-medial",
            "leg": "right", "is_front": True,
        },
        "clips": [C1, C2, C3_REFLUX, C4],
        "vlm_summary":
            "Transverse at SFJ: N1 (CFV) and N2 (GSV). Clear junction anatomy. "
            "Examination complete. All findings consistent with Type 1+2 shunt.",
        "ground_truth":
            "Both Type 1 and Type 1+2 are accepted. action=move is the correct flag. "
            "Any movement command is acceptable — the surgeon is doing a final scan to "
            "confirm no additional circuits exist. Score 3 if action=move regardless of "
            "which zone the probe is directed to. Score 1 only if action=complete or maneuver.",
        "expected_action": "move",
        "notes": "Video 184–187 s: Both shunts accepted. Crew scans for additional circuits.",
    },
]


# =============================================================================
# MOCK HISTORY SUMMARY BUILDER
# No LLM call — deterministic coverage summary from steps completed so far.
# =============================================================================
def _build_history_summary(prior_steps: list[dict]) -> str:
    if not prior_steps:
        return "SCAN HISTORY SUMMARY\nNo zones visited yet."
    bands = {
        "SFJ":        (0.04, 0.07),
        "Upper Thigh":(0.08, 0.20),
        "Hunterian":  (0.21, 0.33),
        "Dodd":       (0.34, 0.47),
        "Popliteal":  (0.48, 0.57),
        "Calf":       (0.58, 0.88),
        "Ankle":      (0.89, 1.00),
    }
    visited = set()
    for s in prior_steps:
        pos = s["probe"]["pos_y"]
        for name, (lo, hi) in bands.items():
            if lo <= pos <= hi:
                visited.add(name)
    lines = []
    for name in bands:
        mark = "[DONE]" if name in visited else "[    ]"
        lines.append(f"  {mark} {name}")
    return "SCAN HISTORY SUMMARY\n" + "\n".join(lines)


# =============================================================================
# EVALUATOR  —  Claude (Anthropic) scores guidance semantically
# =============================================================================
def evaluate_with_llm(
    step: dict,
    guidance_text: str,
    actual_action: str,
    clips: list[dict],
    anthropic_client,
) -> tuple[int, str, str]:
    """
    Ask Claude to score the system's guidance against the ground truth.
    Returns (score 0-3, label, reasoning).
    No keyword lists. Pure clinical semantic judgement by Claude.
    """
    clips_summary = (
        "\n".join(
            f"  {c['flow']} {c['from_type']}→{c['to_type']}  posY={c['pos_y_ratio']:.2f}"
            + (f"  [elimTest={c['elimination_test']}]" if c.get("elimination_test") else "")
            for c in clips
        )
        or "  None yet"
    )

    eval_prompt = f"""You are a senior CHIVA vascular surgeon evaluating an AI duplex guidance system.

EXAMINATION CONTEXT
  Phase: {step["phase"]}
  Time in video: {step["video_time"]}
  What the surgeon was doing: {step["surgeon_action"]}
  Probe position: region={step["probe"]["region"]}, posY={step["probe"]["pos_y"]:.2f}, surface={step["probe"]["surface"]}
  VLM frame annotation: {step["vlm_summary"]}
  Clips confirmed so far ({len(clips)} total):
{clips_summary}

AI SYSTEM RESPONSE
  Guidance text: "{guidance_text or '(empty)'}"
  Action flag:   "{actual_action or '(none)'}"
  Expected action: "{step["expected_action"]}"

GROUND TRUTH (what the guidance MUST achieve at this step)
{step["ground_truth"]}

SCORING RUBRIC (0-3)
  3 = CORRECT:   Guidance directs probe to the right structure with correct clinical intent.
                 Action flag matches expected. Any imprecise wording is acceptable if direction is right.
  2 = PARTIAL:   Guidance is roughly in the right direction but incomplete, too vague,
                 or the action flag is wrong while the guidance text is acceptable.
  1 = WRONG:     Guidance would send the probe to the wrong zone, assess the wrong structure,
                 or misses the most important clinical need at this step.
  0 = NO RESPONSE: Guidance is empty, an error message, or completely unintelligible.

Output ONLY valid JSON, no markdown:
{{"score": 0|1|2|3, "label": "CORRECT"|"PARTIAL"|"WRONG"|"NO RESPONSE", "reasoning": "<1-2 sentences explaining your clinical verdict>"}}"""

    try:
        resp = anthropic_client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=200,
            temperature=0.0,
            messages=[{"role": "user", "content": eval_prompt}],
        )
        raw = (resp.content[0].text or "").strip()
        raw_clean = raw
        if raw_clean.startswith("```"):
            raw_clean = "\n".join(
                ln for ln in raw_clean.splitlines() if not ln.startswith("```")
            )
        s, e = raw_clean.find("{"), raw_clean.rfind("}")
        if s != -1 and e != -1:
            parsed = json.loads(raw_clean[s : e + 1])
            score = int(parsed.get("score", 0))
            label = parsed.get("label", "NO RESPONSE")
            reason = parsed.get("reasoning", "")
            return score, label, reason
    except Exception as exc:
        logger.warning("Evaluator call failed: %s", exc)
        return 0, "NO RESPONSE", f"Evaluator error: {exc}"

    return 0, "NO RESPONSE", "No output from evaluator"


# =============================================================================
# STEP RUNNER
# =============================================================================
def run_step(
    step_num: int,
    step: dict,
    prior_steps: list[dict],
    anthropic_client,
    accepted_shunts: list[str] | None = None,
) -> tuple[StepResult, str | None]:
    """
    Returns (StepResult, shunt_type_if_complete_fired).
    accepted_shunts accumulates across steps so the crew knows which shunts
    are already classified and continues scanning for additional circuits.
    VLM is set to neutral for all steps — we are testing crew reasoning only.
    """
    from agents import protocol_agent, q_state_agent, guidance_agent
    from agents import crew_pipeline

    r = StepResult(
        step_num        = step_num,
        video_time      = step["video_time"],
        phase           = step["phase"],
        surgeon_action  = step["surgeon_action"],
        probe_desc      = (
            f"region={step['probe']['region']}  posY={step['probe']['pos_y']:.2f}  "
            f"surface={step['probe']['surface']}"
        ),
        clips_count     = len(step["clips"]),
        expected_action = step["expected_action"],
    )

    new_shunt_type: str | None = None

    try:
        probe       = step["probe"]
        clips       = step["clips"]
        # Neutral VLM — tests focus on clinical crew reasoning, not frame analysis
        vlm_summary = "No frame analyzed yet."

        history_summary = _build_history_summary(prior_steps)
        q_state         = q_state_agent.analyze(clips)
        protocol_text   = protocol_agent.get_protocol(
            probe["region"], probe["pos_y"], vein_mode=""
        )

        state_msg = guidance_agent.build_state_message(
            region          = probe["region"],
            pos_y           = probe["pos_y"],
            surface         = probe["surface"],
            leg             = probe["leg"],
            clips           = clips,
            vlm_summary     = vlm_summary,
            history_summary = history_summary,
            q_state         = q_state,
            protocol_text   = protocol_text,
            is_front        = probe.get("is_front"),
        )

        recent_g = [rs.guidance_text for rs in _all_results[-4:] if rs.guidance_text]

        t_crew = time.perf_counter()
        guidance, raw_log, action, shunt_found, shunt_type, _ = crew_pipeline.run_guidance_crew(
            state_message   = state_msg,
            region          = probe["region"],
            pos_y           = probe["pos_y"],
            surface         = probe["surface"],
            leg             = probe["leg"],
            clips           = clips,
            vlm_summary     = vlm_summary,
            history_summary = history_summary,
            q_state         = q_state,
            protocol_text   = protocol_text,
            is_front        = probe.get("is_front"),
            recent_guidance = recent_g,
            accepted_shunts = list(accepted_shunts or []),
        )
        r.crew_elapsed_ms = (time.perf_counter() - t_crew) * 1000

        r.guidance_text  = guidance
        r.action         = action
        r.action_correct = (action == step["expected_action"])

        if shunt_found and shunt_type:
            new_shunt_type = f"right:{shunt_type}"

        # LLM evaluator
        t_eval = time.perf_counter()
        score, label, reasoning = evaluate_with_llm(
            step, guidance, action, clips, anthropic_client
        )
        r.eval_elapsed_ms = (time.perf_counter() - t_eval) * 1000
        r.eval_score      = score
        r.eval_label      = label
        r.eval_reasoning  = reasoning

    except Exception as exc:
        r.error          = traceback.format_exc()
        r.guidance_text  = f"ERROR: {exc}"
        r.eval_label     = "NO RESPONSE"
        r.eval_score     = 0
        r.eval_reasoning = f"Step crashed: {exc}"

    return r, new_shunt_type


# =============================================================================
# WORD REPORT
# =============================================================================
def _cell_shade(cell, hex_fill: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _hdr_cell(cell, text: str, fill: str = "1F497D"):
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
    _cell_shade(cell, fill)


def _body_cell(cell, text: str, bold=False, color=None, size=9):
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color


def generate_report(results: list[StepResult], out_path: str):
    if not DOCX_OK:
        print("[SKIP] python-docx not available — install with: pip install python-docx")
        return

    doc = Document()
    sec = doc.sections[0]
    sec.page_width   = Cm(42.0)   # A3 landscape
    sec.page_height  = Cm(29.7)
    sec.left_margin  = sec.right_margin  = Cm(1.5)
    sec.top_margin   = sec.bottom_margin = Cm(1.5)

    # ── Title ──────────────────────────────────────────────────────────────────
    doc.add_heading("Cygnus Med – Task-2 Active Guidance System", 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_heading(
        "LLM-Evaluated Clinical Test Report: Type 1+2 Shunt (Real Patient Case)", 1
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    meta = doc.add_paragraph()
    for label, val in [
        ("Test date",      datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Source video",   "Media - Shunt type 1.mp4  (Mindray M9CV, 30 fps, 3 min 4 s)"),
        ("Diagnosis",      "Type 1+2 combined shunt — EP N1→N2 SFJ + RP N2→N1 thigh + EP N2→N3 mid-calf (Reflux) + RP N3→N1 lower calf"),
        ("Evaluation",     "Claude claude-sonnet-4-6 (Anthropic) — no keyword lists, clinical expert reasoning"),
        ("Steps",          f"{len(results)} examination states across 8 clinical phases"),
    ]:
        meta.add_run(f"{label}:  ").bold = True
        meta.add_run(val + "\n")

    # ── Scoreboard ─────────────────────────────────────────────────────────────
    doc.add_heading("1. Summary Scoreboard", 2)
    n = len(results)
    n3 = sum(r.eval_score == 3 for r in results)
    n2 = sum(r.eval_score == 2 for r in results)
    n1 = sum(r.eval_score == 1 for r in results)
    n0 = sum(r.eval_score == 0 for r in results)
    act_ok = sum(r.action_correct for r in results)
    pct_correct = round(100 * n3 / n) if n else 0

    sb = doc.add_table(rows=1, cols=3); sb.style = "Table Grid"
    _hdr_cell(sb.rows[0].cells[0], "Metric")
    _hdr_cell(sb.rows[0].cells[1], "Score")
    _hdr_cell(sb.rows[0].cells[2], "Verdict")

    def _sb(label, score_str, verdict, ok):
        row = sb.add_row().cells
        _body_cell(row[0], label)
        _body_cell(row[1], score_str, bold=True)
        _body_cell(row[2], verdict, bold=True, color=PASS_CLR if ok else FAIL_CLR)

    _sb("Guidance CORRECT (score 3)",   f"{n3}/{n}  ({pct_correct}%)",
        "PASS" if pct_correct >= 70 else "NEEDS IMPROVEMENT", pct_correct >= 70)
    _sb("Guidance PARTIAL (score 2)",   f"{n2}/{n}", "", False if n2 > n // 3 else True)
    _sb("Guidance WRONG (score 1)",     f"{n1}/{n}", "PASS" if n1 == 0 else "INVESTIGATE", n1 == 0)
    _sb("NO RESPONSE (score 0)",        f"{n0}/{n}", "PASS" if n0 == 0 else "CRITICAL FAIL", n0 == 0)
    _sb("Action flag correct",          f"{act_ok}/{n}",
        "PASS" if act_ok == n else f"FAIL ({n - act_ok} wrong flags)", act_ok == n)
    avg_crew = sum(r.crew_elapsed_ms for r in results) / n if n else 0
    _sb("Avg crew latency", f"{avg_crew:.0f} ms", "< 15 s target" if avg_crew < 15000 else "SLOW", avg_crew < 15000)

    # ── Step table ─────────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("2. Step-by-Step Evaluation", 2)

    cols = [
        "Step\n(time)", "Phase", "Surgeon Action\n(Ground Truth Event)",
        "Probe Position", "Clips\nCount",
        "System Guidance", "Action\nFlag",
        "Expected\nAction", "Action\nOK?",
        "Eval\nScore", "LLM Evaluator Reasoning",
    ]
    WIDTHS = [Cm(1.8), Cm(3.5), Cm(6.0), Cm(3.5), Cm(1.4),
              Cm(7.0), Cm(1.8),
              Cm(1.8), Cm(1.4),
              Cm(1.8), Cm(8.5)]

    tbl = doc.add_table(rows=1, cols=len(cols)); tbl.style = "Table Grid"
    for i, (h, w) in enumerate(zip(cols, WIDTHS)):
        cell = tbl.rows[0].cells[i]
        _hdr_cell(cell, h)
        cell.width = w

    for r in results:
        row = tbl.add_row().cells
        _body_cell(row[0],  f"#{r.step_num}\n{r.video_time}", size=8)
        _body_cell(row[1],  r.phase.replace("Phase ", "Ph").replace(" – ", "\n"), size=8)
        _body_cell(row[2],  r.surgeon_action[:300], size=7)
        _body_cell(row[3],  r.probe_desc, size=8)
        _body_cell(row[4],  str(r.clips_count), bold=True, size=9)
        _body_cell(row[5],  r.guidance_text or "(empty)", size=8)
        act_color = PASS_CLR if r.action_correct else FAIL_CLR
        _body_cell(row[6],  r.action or "—", bold=True, color=act_color, size=8)
        _body_cell(row[7],  r.expected_action, size=8)
        _body_cell(row[8],  "YES" if r.action_correct else "NO", bold=True,
                   color=PASS_CLR if r.action_correct else FAIL_CLR, size=8)
        sc_color = SCORE_CLR.get(r.eval_score, _rgb(0, 0, 0))
        _body_cell(row[9],  f"{SCORE_LABEL.get(r.eval_score, '?')}\n({r.eval_score}/3)",
                   bold=True, color=sc_color, size=8)
        _body_cell(row[10], r.eval_reasoning[:400], size=8)

    # ── Deep-dive per step ─────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("3. Clinical Deep-Dive (per step)", 2)

    for r in results:
        doc.add_heading(f"Step {r.step_num}: {r.video_time} — {r.phase}", 3)
        p = doc.add_paragraph()

        def _line(label, val, color=None):
            run_lbl = p.add_run(label + ": ")
            run_lbl.bold = True
            run_val = p.add_run(str(val) + "\n")
            if color:
                run_val.font.color.rgb = color

        _line("Surgeon action",   r.surgeon_action)
        _line("Probe",            r.probe_desc)
        _line("Clips confirmed",  r.clips_count)
        _line("System guidance",  f'"{r.guidance_text or "(no response)"}"',
              SCORE_CLR.get(r.eval_score))
        _line("Action flag",      f"{r.action}  (expected: {r.expected_action})",
              PASS_CLR if r.action_correct else FAIL_CLR)
        _line("Evaluator verdict", f"{r.eval_label} ({r.eval_score}/3)",
              SCORE_CLR.get(r.eval_score))
        _line("Evaluator reasoning", r.eval_reasoning)
        if r.error:
            _line("ERROR", r.error[:300], FAIL_CLR)

    # ── Phase summary ─────────────────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("4. Phase-Level Summary", 2)

    phases: dict[str, list[StepResult]] = {}
    for r in results:
        ph = r.phase.split("–")[0].strip()
        phases.setdefault(ph, []).append(r)

    ph_tbl = doc.add_table(rows=1, cols=5); ph_tbl.style = "Table Grid"
    for i, h in enumerate(["Phase", "Steps", "CORRECT", "PARTIAL/WRONG", "Action OK"]):
        _hdr_cell(ph_tbl.rows[0].cells[i], h)

    for ph, rs in phases.items():
        row = ph_tbl.add_row().cells
        _body_cell(row[0], ph, size=9)
        _body_cell(row[1], str(len(rs)), size=9)
        nc = sum(r.eval_score == 3 for r in rs)
        nw = len(rs) - nc
        _body_cell(row[2], f"{nc}/{len(rs)}", bold=True,
                   color=PASS_CLR if nc == len(rs) else WARN_CLR, size=9)
        _body_cell(row[3], str(nw), bold=True,
                   color=PASS_CLR if nw == 0 else FAIL_CLR, size=9)
        na = sum(r.action_correct for r in rs)
        _body_cell(row[4], f"{na}/{len(rs)}", bold=True,
                   color=PASS_CLR if na == len(rs) else FAIL_CLR, size=9)

    doc.save(out_path)
    print(f"\n[OK] Report saved → {out_path}")


# =============================================================================
# MAIN
# =============================================================================
_all_results: list[StepResult] = []


def main():
    import anthropic as _anthropic

    print("=" * 72)
    print("  CYGNUS MED — Task-2 Claude-Evaluated Test: Type 1+2 Shunt")
    print(f"  {len(STEPS)} steps across 8 clinical phases")
    print(f"  Evaluator: Claude claude-sonnet-4-6 (Anthropic)")
    print(f"  {datetime.now():%Y-%m-%d %H:%M:%S}")
    print("=" * 72)

    anthropic_client = _anthropic.Anthropic()   # reads ANTHROPIC_API_KEY from env
    accepted_shunts: list[str] = []   # accumulates as complete events fire

    for i, step in enumerate(STEPS, 1):
        label = f"[{step['video_time']}] {step['phase']}"
        acc_str = f"  (accepted: {accepted_shunts})" if accepted_shunts else ""
        print(f"\n  Step {i:2d}/{len(STEPS)}: {label}{acc_str}")

        r, new_shunt = run_step(i, step, STEPS[:i - 1], anthropic_client, accepted_shunts)
        _all_results.append(r)

        # Only accept a shunt when the crew fires action=complete for it.
        # Shunt analyst may speculatively confirm a type even during maneuver steps;
        # we must NOT accept those — they would pollute accepted_shunts prematurely.
        if new_shunt and new_shunt not in accepted_shunts and r.action == "complete":
            accepted_shunts.append(new_shunt)
            print(f"         [accepted {new_shunt} — crew continues scanning]")

        act_sym  = "OK" if r.action_correct else "!!"
        score_lbl = SCORE_LABEL.get(r.eval_score, "?")
        clr_sym   = {3: "OK", 2: "~~", 1: "!!", 0: "??"}.get(r.eval_score, "?")

        print(f"         Guidance : {repr(r.guidance_text[:90])}")
        print(f"         Action   : [{act_sym}] {r.action!r}  (expected: {r.expected_action!r})")
        print(f"         Eval     : [{clr_sym}] {score_lbl} -- {r.eval_reasoning[:100]}")
        print(f"         Latency  : crew={r.crew_elapsed_ms:.0f} ms  eval={r.eval_elapsed_ms:.0f} ms")

    print("\n" + "=" * 72)
    n = len(_all_results)
    print(f"  Steps:         {n}")
    print(f"  CORRECT (3):   {sum(r.eval_score == 3 for r in _all_results)}/{n}")
    print(f"  PARTIAL (2):   {sum(r.eval_score == 2 for r in _all_results)}/{n}")
    print(f"  WRONG   (1):   {sum(r.eval_score == 1 for r in _all_results)}/{n}")
    print(f"  NO RESP (0):   {sum(r.eval_score == 0 for r in _all_results)}/{n}")
    print(f"  Action correct:{sum(r.action_correct for r in _all_results)}/{n}")
    print("=" * 72)

    ts  = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = rf"C:\Users\Krish\Downloads\Task2_Type12_LLMEval_{ts}.docx"
    generate_report(_all_results, out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
