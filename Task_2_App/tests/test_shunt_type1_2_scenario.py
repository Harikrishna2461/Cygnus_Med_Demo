"""
=============================================================================
CYGNUS MED – Task-2 Active Guidance System
CLINICAL SCENARIO TEST: Type 1 + 2 Shunt Examination
=============================================================================
SOURCE VIDEO : "Media - Shunt type 1.mp4"  (3 min 4 s, Mindray M9CV, 30 fps)
PATIENT      : Standing, right leg
DIAGNOSIS    : Type 1 + Type 2 combined shunt (CHIVA classification)

WHAT THIS TESTS:
  At each step of the real surgeon's examination, we send the probe position
  to the Task-2 system and evaluate:
    - Was the anatomical region correctly identified?
    - What guidance text did the system generate?
    - Does that guidance correctly direct the sonographer to the NEXT step?
    - Does the system correctly reject wrong-region positions?
    - Does the VLM prompt contain the right clinical context?

EXAMINATION TIMELINE (from frame-by-frame video analysis):
  0–40 s    SFJ (3 steps) — Probe arrives groin crease; CFV + GSV identified at junction
  40–82 s   Thigh (4 steps) — GSV tracked proximal→Hunterian→Dodd zone, N2 in compartment
  82–130 s  Calf (4 steps)  — GSV mapped upper calf→mid→lower→ankle, N2 continuous
  130–165 s SPJ (3 steps)  — Popliteal fossa; SSV + popliteal N1/N2 junction confirmed
  165–183 s Elim (2 steps) — Probe returns to SFJ; elimination test positioning
=============================================================================
"""

from __future__ import annotations

import json
import sys
import time
import traceback
import urllib.request
import urllib.error
import urllib.parse
from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_OK = True
except ImportError:
    DOCX_OK = False

BASE_URL = "http://127.0.0.1:7861"
SESSION  = "shunt_type12_clinical_test"
TIMEOUT  = 30


# ═══════════════════════════════════════════════════════════════════════════
# Data model
# ═══════════════════════════════════════════════════════════════════════════

@dataclass
class StepResult:
    # Identity
    video_time:      str        # e.g. "0–40 s"
    clinical_phase:  str        # e.g. "SFJ Assessment"
    clinical_action: str        # what the surgeon did in the video

    # Probe input
    segment_id:      int
    segment_dist:    float
    is_front:        bool
    position_label:  str        # human-readable probe position description

    # System outputs
    region_detected: str        = ""
    region_expected: str        = ""
    region_correct:  bool       = False

    guidance_text:   str        = ""   # actual LLM output
    guidance_ok:     bool       = False   # did LLM respond (non-error)?
    expected_next_step: str     = ""   # what should the guidance say?
    guidance_match:  str        = ""   # CORRECT / PARTIAL / WRONG / NO_RESPONSE
    guidance_score:  int        = 0    # 0-3: 0=no resp, 1=wrong, 2=partial, 3=correct

    wrong_region_triggered: Optional[bool] = None   # if wrong-region test was run
    wrong_region_reason:    str             = ""

    vein_frame_ok:   bool       = False
    vein_type:       str        = ""

    llm_prompt_ok:   bool       = False   # prompt contains required clinical keywords
    missing_kws:     list       = field(default_factory=list)

    elapsed_ms:      float      = 0.0
    notes:           str        = ""
    error:           str        = ""


# ═══════════════════════════════════════════════════════════════════════════
# API helpers
# ═══════════════════════════════════════════════════════════════════════════

def _post(endpoint: str, payload: dict) -> tuple[dict, float]:
    url  = BASE_URL + endpoint
    data = json.dumps(payload).encode()
    req  = urllib.request.Request(url, data=data,
                                  headers={"Content-Type": "application/json"},
                                  method="POST")
    t0 = time.perf_counter()
    with urllib.request.urlopen(req, timeout=TIMEOUT) as resp:
        body = json.loads(resp.read())
    return body, (time.perf_counter() - t0) * 1000


def _get(endpoint: str, params: dict | None = None) -> tuple[dict, float, int]:
    url = BASE_URL + endpoint
    if params:
        url += "?" + urllib.parse.urlencode(params)
    t0 = time.perf_counter()
    try:
        with urllib.request.urlopen(url, timeout=TIMEOUT) as resp:
            status = resp.status
            ctype  = resp.headers.get("Content-Type", "")
            body   = json.loads(resp.read()) if "json" in ctype else {
                "_headers": dict(resp.headers), "_status": status
            }
        return body, (time.perf_counter() - t0) * 1000, status
    except urllib.error.HTTPError as e:
        return {"error": str(e)}, (time.perf_counter() - t0) * 1000, e.code


def _reset():
    try:
        _post("/api/localize/reset", {"session_id": SESSION})
    except Exception:
        pass


def _localize_batch(readings: list[dict]) -> dict:
    body, _ = _post("/api/localize/batch",
                    {"session_id": SESSION, "readings": readings})
    return body


def _guidance_call(loc: dict) -> dict:
    body, _ = _post("/api/guidance", {"session_id": SESSION, "location": loc})
    return body


def _wrong_region(current: str, expected: str) -> dict:
    body, _ = _post("/api/wrong-region",
                    {"current_region": current, "expected_region": expected})
    return body


def _stable_loc(segment_id: int, segment_dist: float, is_front: bool,
                n_readings: int = 10) -> dict:
    """Feed n_readings to the sliding window and return the stable location."""
    readings = [
        {"segment_id": segment_id,
         "segment_dist": segment_dist + i * 0.001,
         "is_front": is_front}
        for i in range(n_readings)
    ]
    return _localize_batch(readings)


def _score_guidance(text: str, expected_step: str,
                    must_contain: list[str],
                    must_not_contain: list[str] | None = None) -> tuple[str, int]:
    """
    Score guidance text against expected next step.
    Returns (label, score): score 0–3.
      3 = CORRECT   — all required keywords present, no forbidden words
      2 = PARTIAL   — ≥50% required keywords present
      1 = WRONG     — guidance present but off-topic
      0 = NO_RESPONSE — empty or LLM error
    """
    if not text or "llm error" in text.lower() or "non-json" in text.lower():
        return "NO_RESPONSE", 0
    tl = text.lower()
    found = [kw for kw in must_contain if kw.lower() in tl]
    ratio = len(found) / len(must_contain) if must_contain else 1.0
    forbidden_hit = any(w.lower() in tl for w in (must_not_contain or []))
    if ratio >= 1.0 and not forbidden_hit:
        return "CORRECT", 3
    elif ratio >= 0.5:
        return "PARTIAL", 2
    else:
        return "WRONG", 1


# ═══════════════════════════════════════════════════════════════════════════
# EXAMINATION STEPS (ground-truth from video analysis)
# Each step records:
#   - what happened in the video
#   - the probe position to simulate
#   - what the system SHOULD guide the sonographer to do next
#   - keywords that must appear in the guidance (and optionally must NOT appear)
# ═══════════════════════════════════════════════════════════════════════════

STEPS = [
    # ── PHASE 1: SFJ Initial ──────────────────────────────────────────────
    {
        "video_time":    "0–5 s",
        "clinical_phase":"Phase 1 – SFJ Arrival",
        "clinical_action":
            "Surgeon places probe at groin crease, anteromedial surface. "
            "First skin contact. Two adjacent anechoic ovals visible on B-mode: "
            "CFV (larger) and GSV (smaller).",
        "segment_id": 0, "segment_dist": 0.05, "is_front": True,
        "position_label": "Groin crease, anteromedial (dist=0.05) — SFJ arrival",
        "expected_region": "SFJ",
        "expected_next_step":
            "Identify N1 (CFV) and N2 (GSV) at groin; confirm the two adjacent "
            "ovals represent the saphenofemoral junction before moving distally.",
        "guidance_must_contain": ["sfj", "n2"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["sfj", "saphenofemoral", "n2"],
        "wrong_region_test": {"current": "SPJ", "expected": "SFJ", "should_flag": True},
        "vein_frame_region": "SFJ",
        "notes":
            "Video 0 s: Probe first placed at groin crease. CFV + GSV visible "
            "as adjacent oval structures on B-mode."
    },
    {
        "video_time":    "5–15 s",
        "clinical_phase":"Phase 1 – SFJ Orientation",
        "clinical_action":
            "Probe adjusted medially to centre on SFJ junction. "
            "GSV and CFV clearly separated as distinct structures. "
            "Junction anatomy confirmed.",
        "segment_id": 0, "segment_dist": 0.07, "is_front": True,
        "position_label": "Groin crease, medial adjustment (dist=0.07)",
        "expected_region": "SFJ",
        "expected_next_step":
            "Confirm N1 (CFV) and N2 (GSV) are both visible at same depth; "
            "ensure junction anatomy is clear before proceeding distally.",
        "guidance_must_contain": ["sfj", "n1", "n2"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["sfj", "saphenofemoral", "n1"],
        "wrong_region_test": None,
        "vein_frame_region": "SFJ",
        "notes":
            "Video 5–15 s: Probe repositioned medially. CFV larger, GSV smaller. "
            "SFJ anatomy confirmed."
    },
    {
        "video_time":    "15–40 s",
        "clinical_phase":"Phase 1 – SFJ Confirmed",
        "clinical_action":
            "Probe stable at SFJ. GSV junction with CFV confirmed. "
            "SFJ incompetence identified (Type 1 escape point). "
            "Surgeon completes SFJ assessment.",
        "segment_id": 0, "segment_dist": 0.06, "is_front": True,
        "position_label": "Groin crease, stable SFJ position (dist=0.06)",
        "expected_region": "SFJ",
        "expected_next_step":
            "Move probe distally along anteromedial thigh to track GSV from SFJ "
            "through proximal thigh and map the reflux pathway.",
        "guidance_must_contain": ["thigh", "gsv"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["sfj", "saphenofemoral", "n2"],
        "wrong_region_test": None,
        "vein_frame_region": "SFJ",
        "notes":
            "Video 15–40 s: SFJ assessment complete. Type 1 escape point identified. "
            "Next step is to track GSV distally through thigh."
    },

    # ── PHASE 2: Proximal Thigh ───────────────────────────────────────────
    {
        "video_time":    "40–48 s",
        "clinical_phase":"Phase 2 – Proximal Thigh",
        "clinical_action":
            "Probe moved to proximal anteromedial thigh, just below groin crease. "
            "B-mode: large anechoic GSV oval visible within the fascial compartment. "
            "N2 clearly identified between fascial layers.",
        "segment_id": 0, "segment_dist": 0.18, "is_front": True,
        "position_label": "Proximal thigh, anteromedial (dist=0.18)",
        "expected_region": "GSV-THI",
        "expected_next_step":
            "Continue scanning distally along thigh; track GSV in fascial compartment "
            "toward Hunterian canal region at mid-thigh.",
        "guidance_must_contain": ["gsv", "thigh"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["gsv", "thigh", "n2"],
        "wrong_region_test": {"current": "SFJ", "expected": "GSV-THI", "should_flag": True},
        "vein_frame_region": "GSV-THI",
        "notes":
            "Video 40–45 s: Probe swept to proximal thigh. GSV large anechoic oval "
            "in fascial compartment clearly visible."
    },
    {
        "video_time":    "48–56 s",
        "clinical_phase":"Phase 2 – Upper Thigh",
        "clinical_action":
            "Probe at upper thigh (Hunterian proximal zone). "
            "GSV remains visible as anechoic oval between fascial layers. "
            "Fascial compartment boundaries clearly identified.",
        "segment_id": 0, "segment_dist": 0.28, "is_front": True,
        "position_label": "Upper thigh / Hunterian proximal (dist=0.28)",
        "expected_region": "GSV-THI",
        "expected_next_step":
            "Continue distally through Hunterian canal; confirm GSV continuity "
            "in fascial compartment at mid-thigh level.",
        "guidance_must_contain": ["gsv", "thigh"],
        "guidance_must_not_contain": ["sfj"],
        "prompt_must_contain": ["gsv", "thigh", "n2", "fascial"],
        "wrong_region_test": None,
        "vein_frame_region": "GSV-THI",
        "notes":
            "Video 45–50 s: GSV clearly within the saphenous fascial compartment "
            "at Hunterian proximal level."
    },
    {
        "video_time":    "56–65 s",
        "clinical_phase":"Phase 2 – Hunterian Canal",
        "clinical_action":
            "Probe at mid-thigh (Hunterian canal). GSV trunk prominent at this level. "
            "Single large oval between fascial layers with clear compartment boundaries.",
        "segment_id": 0, "segment_dist": 0.40, "is_front": True,
        "position_label": "Mid-thigh / Hunterian canal (dist=0.40)",
        "expected_region": "GSV-THI",
        "expected_next_step":
            "Continue distally to map GSV extent; move probe toward "
            "distal thigh / Dodd perforator zone approaching the knee.",
        "guidance_must_contain": ["gsv", "thigh"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["gsv", "thigh", "n2"],
        "wrong_region_test": {"current": "SFJ", "expected": "GSV-THI", "should_flag": True},
        "vein_frame_region": "GSV-THI",
        "notes":
            "Video 55–65 s: Mid-thigh Hunterian canal level. GSV trunk between "
            "fascial layers, large calibre oval."
    },

    # ── PHASE 3: Distal Thigh ─────────────────────────────────────────────
    {
        "video_time":    "65–73 s",
        "clinical_phase":"Phase 3 – Mid-Distal Thigh",
        "clinical_action":
            "Probe at mid-distal thigh. Single GSV oval still in fascial compartment. "
            "GSV remains well-defined within saphenous fascia.",
        "segment_id": 0, "segment_dist": 0.52, "is_front": True,
        "position_label": "Mid-distal thigh (dist=0.52)",
        "expected_region": "GSV-THI",
        "expected_next_step":
            "Move probe to distal thigh (Dodd zone); then transition to calf "
            "to continue tracking GSV distally.",
        "guidance_must_contain": ["gsv", "calf"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["gsv", "thigh", "n2"],
        "wrong_region_test": None,
        "vein_frame_region": "GSV-THI",
        "notes":
            "Video 65–70 s: Probe at mid-distal thigh. GSV consistent in "
            "fascial compartment, reflux extent being mapped."
    },
    {
        "video_time":    "73–82 s",
        "clinical_phase":"Phase 3 – Dodd Zone",
        "clinical_action":
            "Probe at distal thigh, Dodd perforator zone (above knee medially). "
            "GSV still within fascial compartment. Reflux column tracked to distal thigh. "
            "GSV visible just above the knee.",
        "segment_id": 0, "segment_dist": 0.65, "is_front": True,
        "position_label": "Distal thigh / Dodd zone, anteromedial (dist=0.65)",
        "expected_region": "GSV-THI",
        "expected_next_step":
            "Transition probe from thigh to calf; begin mapping GSV from upper calf "
            "continuing the reflux column distally.",
        "guidance_must_contain": ["gsv", "calf"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["gsv", "thigh", "n2"],
        "wrong_region_test": {"current": "SFJ", "expected": "GSV-THI", "should_flag": True},
        "vein_frame_region": "GSV-THI",
        "notes":
            "Video 70–80 s: Dodd level, distal thigh. Fascia clearly visible, "
            "GSV in compartment. Transition to calf is next step."
    },

    # ── PHASE 4: Calf Mapping ─────────────────────────────────────────────
    {
        "video_time":    "82–92 s",
        "clinical_phase":"Phase 4 – Upper Calf",
        "clinical_action":
            "Probe moved to upper calf, just below knee crease, anteromedial. "
            "Small GSV oval visible in fascial compartment. "
            "GSV enters calf from thigh; N2 confirmed at upper calf.",
        "segment_id": 1, "segment_dist": 0.10, "is_front": True,
        "position_label": "Upper calf, anteromedial (seg=1, dist=0.10)",
        "expected_region": "GSV-CAL",
        "expected_next_step":
            "Scan distally along medial calf, tracking GSV in fascial compartment "
            "toward mid-calf level.",
        "guidance_must_contain": ["gsv", "calf"],
        "guidance_must_not_contain": ["sfj"],
        "prompt_must_contain": ["gsv", "calf", "n2"],
        "wrong_region_test": {"current": "SPJ", "expected": "GSV-CAL", "should_flag": True},
        "vein_frame_region": "GSV-CAL",
        "notes":
            "Video 80–90 s: Probe at upper calf. Small GSV oval entering calf "
            "within fascial compartment."
    },
    {
        "video_time":    "92–102 s",
        "clinical_phase":"Phase 4 – Mid-Calf",
        "clinical_action":
            "Probe at mid-calf, anteromedial. GSV visible in fascial compartment. "
            "Consistent oval appearance maintained through mid-calf level.",
        "segment_id": 1, "segment_dist": 0.35, "is_front": True,
        "position_label": "Mid-calf, anteromedial (seg=1, dist=0.35)",
        "expected_region": "GSV-CAL",
        "expected_next_step":
            "Continue scanning distally along medial calf to lower calf; "
            "map the extent of the GSV compartment toward ankle.",
        "guidance_must_contain": ["gsv", "calf"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["gsv", "calf", "n2"],
        "wrong_region_test": None,
        "vein_frame_region": "GSV-CAL",
        "notes":
            "Video 90–100 s: Probe at mid-calf. GSV oval in fascial compartment "
            "consistent with N2 appearance."
    },
    {
        "video_time":    "102–112 s",
        "clinical_phase":"Phase 4 – Lower Calf",
        "clinical_action":
            "Probe at lower calf. GSV calibre decreasing toward ankle. "
            "Fascial compartment still visible but narrowing. "
            "GSV N2 tracked to lower leg.",
        "segment_id": 1, "segment_dist": 0.60, "is_front": True,
        "position_label": "Lower calf, anteromedial (seg=1, dist=0.60)",
        "expected_region": "GSV-CAL",
        "expected_next_step":
            "Continue to ankle level to confirm distal extent of GSV; "
            "then reposition probe to popliteal fossa to assess SPJ.",
        "guidance_must_contain": ["gsv", "calf"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["gsv", "calf", "n2"],
        "wrong_region_test": None,
        "vein_frame_region": "GSV-CAL",
        "notes":
            "Video 100–110 s: Lower calf. GSV getting smaller in calibre as "
            "probe approaches ankle."
    },
    {
        "video_time":    "112–130 s",
        "clinical_phase":"Phase 4 – Ankle Level",
        "clinical_action":
            "Probe at ankle, medial malleolus. Small GSV oval visible at ankle level. "
            "Distal extent of GSV mapping confirmed. "
            "Fascial compartment at its narrowest here.",
        "segment_id": 1, "segment_dist": 0.82, "is_front": True,
        "position_label": "Ankle level, medial (seg=1, dist=0.82)",
        "expected_region": "GSV-CAL",
        "expected_next_step":
            "GSV reflux column mapped to ankle. Reposition probe to popliteal "
            "fossa (posterior knee) to assess SPJ for Type 2 shunt.",
        "guidance_must_contain": ["spj", "gsv"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["gsv", "calf", "n2"],
        "wrong_region_test": {"current": "SFJ", "expected": "GSV-CAL", "should_flag": True},
        "vein_frame_region": "GSV-CAL",
        "notes":
            "Video 110–130 s: Probe at ankle. GSV very small at malleolus level. "
            "Distal extent of Type 1 shunt pathway confirmed."
    },

    # ── PHASE 5: SPJ Assessment ───────────────────────────────────────────
    {
        "video_time":    "130–142 s",
        "clinical_phase":"Phase 5 – SPJ Approach",
        "clinical_action":
            "Probe repositioned to popliteal fossa, posterior knee. "
            "B-mode: two large adjacent anechoic ovals — popliteal vein (N1, larger) "
            "and SSV (N2, smaller) — visible at junction.",
        "segment_id": 0, "segment_dist": 0.87, "is_front": False,
        "position_label": "Popliteal fossa, posterior thigh (dist=0.87, back)",
        "expected_region": "SPJ",
        "expected_next_step":
            "Identify SSV junction with popliteal vein; confirm N1 (popliteal) "
            "and N2 (SSV) both visible at same depth at SPJ.",
        "guidance_must_contain": ["spj", "n2"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["spj", "saphenopopliteal", "n2"],
        "wrong_region_test": {"current": "GSV-CAL", "expected": "SPJ", "should_flag": True},
        "vein_frame_region": "SPJ",
        "notes":
            "Video 130 s: Probe placed at popliteal fossa. Two large adjacent "
            "ovals visible — popliteal vein and SSV at junction."
    },
    {
        "video_time":    "142–155 s",
        "clinical_phase":"Phase 5 – SPJ Confirmed",
        "clinical_action":
            "SPJ junction confirmed. SSV and popliteal vein identified as adjacent "
            "structures. SSV reflux into popliteal vein visible. "
            "Type 2 escape point confirmed.",
        "segment_id": 1, "segment_dist": 0.03, "is_front": False,
        "position_label": "Posterior upper calf / SPJ from below (seg=1, dist=0.03)",
        "expected_region": "SPJ",
        "expected_next_step":
            "SPJ incompetence confirmed. Return probe to groin (SFJ) for "
            "elimination test — probe back to anteromedial groin crease.",
        "guidance_must_contain": ["sfj", "groin"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["spj", "saphenopopliteal"],
        "wrong_region_test": {"current": "GSV-THI", "expected": "SPJ", "should_flag": True},
        "vein_frame_region": "SPJ",
        "notes":
            "Video 140–155 s: SPJ junction confirmed. SSV + popliteal vein + "
            "tributary branches all visible. Type 2 shunt escape point identified."
    },
    {
        "video_time":    "155–165 s",
        "clinical_phase":"Phase 5 – SPJ Complete",
        "clinical_action":
            "Full SPJ anatomy mapped with multiple vessels visible at junction — "
            "popliteal vein, SSV, and tributary branches. "
            "SPJ assessment complete.",
        "segment_id": 1, "segment_dist": 0.05, "is_front": False,
        "position_label": "SPJ posterior, probe stable (seg=1, dist=0.05, back)",
        "expected_region": "SPJ",
        "expected_next_step":
            "Return to SFJ at groin for elimination test; confirm SFJ as the "
            "sole escape point for the Type 1 shunt.",
        "guidance_must_contain": ["sfj", "groin"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["spj", "saphenopopliteal"],
        "wrong_region_test": None,
        "vein_frame_region": "SPJ",
        "notes":
            "Video 155–165 s: SPJ fully mapped. Type 2 escape point confirmed. "
            "Probe now needs to return to SFJ."
    },

    # ── PHASE 6: Elimination Test ─────────────────────────────────────────
    {
        "video_time":    "165–175 s",
        "clinical_phase":"Phase 6 – Return to SFJ",
        "clinical_action":
            "Probe returns to groin crease (SFJ), anteromedial surface. "
            "Same anatomical position as initial assessment. "
            "N1 (CFV) and N2 (GSV) again visible as adjacent ovals.",
        "segment_id": 0, "segment_dist": 0.05, "is_front": True,
        "position_label": "Groin crease, return for elimination test (dist=0.05)",
        "expected_region": "SFJ",
        "expected_next_step":
            "Confirm SFJ anatomy with N2 (GSV) and N1 (CFV) visible at junction; "
            "ensure same imaging position as initial SFJ assessment.",
        "guidance_must_contain": ["sfj", "n2"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["sfj", "saphenofemoral", "n1"],
        "wrong_region_test": {"current": "SPJ", "expected": "SFJ", "should_flag": True},
        "vein_frame_region": "SFJ",
        "notes":
            "Video 165 s: Probe physically back at groin crease. Same position "
            "as examination start. N1 and N2 again visible."
    },
    {
        "video_time":    "175–183 s",
        "clinical_phase":"Phase 6 – Elimination Test",
        "clinical_action":
            "Probe stable at SFJ during elimination test. "
            "Bidirectional flow pattern at SFJ confirmed. "
            "SFJ is sole escape point for Type 1 shunt. "
            "Final diagnosis: Type 1 + Type 2 combined shunt.",
        "segment_id": 0, "segment_dist": 0.07, "is_front": True,
        "position_label": "SFJ, elimination test — probe stable (dist=0.07)",
        "expected_region": "SFJ",
        "expected_next_step":
            "Document findings: SFJ incompetent (Type 1 shunt), SPJ incompetent (Type 2 shunt). "
            "Mark anatomy on patient skin. Plan CHIVA procedure accordingly.",
        "guidance_must_contain": ["sfj"],
        "guidance_must_not_contain": [],
        "prompt_must_contain": ["sfj", "saphenofemoral"],
        "wrong_region_test": None,
        "vein_frame_region": "SFJ",
        "notes":
            "Video 175–183 s: Elimination test complete at SFJ. Positive result. "
            "Type 1 + 2 combined shunt diagnosis finalised."
    },
]


# ═══════════════════════════════════════════════════════════════════════════
# Run all steps
# ═══════════════════════════════════════════════════════════════════════════

all_results: list[StepResult] = []


def run_step(step: dict) -> StepResult:
    r = StepResult(
        video_time      = step["video_time"],
        clinical_phase  = step["clinical_phase"],
        clinical_action = step["clinical_action"],
        segment_id      = step["segment_id"],
        segment_dist    = step["segment_dist"],
        is_front        = step["is_front"],
        position_label  = step["position_label"],
        region_expected = step["expected_region"],
        expected_next_step = step["expected_next_step"],
    )

    t0 = time.perf_counter()
    try:
        _reset()

        # ── 1. Stable localisation ────────────────────────────────────────
        loc_body = _stable_loc(step["segment_id"], step["segment_dist"],
                               step["is_front"], n_readings=12)
        loc      = loc_body.get("location", {})
        r.region_detected = loc.get("region", "UNKNOWN")
        r.region_correct  = (r.region_detected == r.region_expected)

        # ── 2. Guidance call ──────────────────────────────────────────────
        g_body        = _guidance_call(loc)
        r.guidance_text = g_body.get("guidance", "")
        r.guidance_ok   = (
            bool(r.guidance_text)
            and "llm error" not in r.guidance_text.lower()
            and "non-json" not in r.guidance_text.lower()
        )

        # ── 3. Score guidance against expected next step ──────────────────
        label, score = _score_guidance(
            r.guidance_text,
            step["expected_next_step"],
            step["guidance_must_contain"],
            step.get("guidance_must_not_contain"),
        )
        r.guidance_match = label
        r.guidance_score = score

        # ── 4. LLM prompt quality ─────────────────────────────────────────
        prompt = g_body.get("debug", {}).get("llm_prompt", "").lower()
        kws    = step.get("prompt_must_contain", [])
        r.missing_kws   = [k for k in kws if k.lower() not in prompt]
        r.llm_prompt_ok = len(r.missing_kws) == 0

        # ── 5. Wrong-region test ──────────────────────────────────────────
        wr_test = step.get("wrong_region_test")
        if wr_test:
            wr = _wrong_region(wr_test["current"], wr_test["expected"])
            r.wrong_region_triggered = wr.get("is_wrong", False)
            r.wrong_region_reason    = wr.get("reason", "") or wr.get("suggestion", "")

        # ── 6. Vein frame endpoint ────────────────────────────────────────
        vf_region = step.get("vein_frame_region", r.region_expected)
        vf_body, _, vf_status = _get("/api/vein-frame",
                                      {"region": vf_region,
                                       "pos_y": f"{step['segment_dist']:.3f}"})
        r.vein_frame_ok = vf_status in (200, 206)
        r.vein_type     = (vf_body.get("_headers", {}).get("X-Vein-Label", "")
                           or vf_body.get("X-Vein-Label", ""))

        r.notes = step.get("notes", "")

    except Exception as exc:
        r.error = traceback.format_exc()
        r.guidance_text = f"ERROR: {exc}"
        r.guidance_match = "ERROR"

    r.elapsed_ms = (time.perf_counter() - t0) * 1000
    return r


def _rgb(r, g, b): return RGBColor(r, g, b)
PASS_CLR  = _rgb(0, 128, 0)
FAIL_CLR  = _rgb(180, 0, 0)
WARN_CLR  = _rgb(200, 120, 0)
HEAD_CLR  = _rgb(31, 73, 125)
SCORE_CLR = {3: _rgb(0,128,0), 2: _rgb(180,130,0), 1: _rgb(180,0,0), 0: _rgb(120,120,120)}

SCORE_LABEL = {3: "CORRECT", 2: "PARTIAL", 1: "WRONG", 0: "NO RESPONSE"}

def _cell_shade(cell, hex_fill: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_fill)
    shd.set(qn("w:val"),  "clear")
    tcPr.append(shd)


def _hdr_cell(cell, text: str, hex_fill: str = "1F497D"):
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(255, 255, 255)
    _cell_shade(cell, hex_fill)


def _body_cell(cell, text: str, bold=False, color=None, size=9):
    cell.text = text
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = bold
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = color


# ═══════════════════════════════════════════════════════════════════════════
# Word report
# ═══════════════════════════════════════════════════════════════════════════

def generate_report(results: list[StepResult], out_path: str):
    if not DOCX_OK:
        print("[SKIP] python-docx not available"); return

    doc = Document()
    sec = doc.sections[0]
    sec.page_width  = Cm(29.7)   # A4 landscape
    sec.page_height = Cm(21.0)
    sec.left_margin = sec.right_margin = Cm(1.5)
    sec.top_margin  = sec.bottom_margin = Cm(1.5)

    # ── Title ──────────────────────────────────────────────────────────────
    t = doc.add_heading("Cygnus Med – Task-2 Active Guidance System", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = doc.add_heading("Clinical Scenario Test Report: Type 1 + 2 Shunt", 1)
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    meta = doc.add_paragraph()
    for label, val in [
        ("Test date",     datetime.now().strftime("%Y-%m-%d %H:%M")),
        ("Source video",  "Media - Shunt type 1.mp4  (3 min 4 s, Mindray M9CV, 8 MHz linear)"),
        ("Patient",       "Standing, right leg"),
        ("Diagnosis",     "Type 1 + 2 combined shunt (CHIVA classification)"),
        ("Backend",       BASE_URL),
    ]:
        meta.add_run(f"{label}:  ").bold = True
        meta.add_run(val + "\n")

    # ── Examination narrative ──────────────────────────────────────────────
    doc.add_heading("1. Examination Narrative (from video)", 2)
    narrative = [
        ("0–40 s   SFJ (3 positions)",
         "Probe placed at groin crease (anteromedial). CFV (N1, larger oval) and GSV (N2, smaller) "
         "identified as adjacent structures on B-mode. Probe adjusted medially to centre on junction. "
         "SFJ incompetence confirmed — TYPE 1 ESCAPE POINT identified."),
        ("40–82 s  GSV Thigh (4 positions)",
         "Probe swept distally along anteromedial thigh. GSV oval tracked in fascial compartment "
         "from proximal thigh → Hunterian canal (mid-thigh) → Dodd zone (distal thigh). "
         "N2 consistently visible between fascial layers at each level. TYPE 1 SHUNT PATH mapped."),
        ("82–130 s Calf (4 positions)",
         "Probe moved to calf, anteromedial. Small GSV oval in fascial compartment tracked from "
         "upper calf → mid-calf → lower calf → ankle. GSV calibre decreases distally. "
         "N2 continuously visible within saphenous compartment to ankle level."),
        ("130–165 s SPJ (3 positions)",
         "Probe repositioned to popliteal fossa (posterior knee). Two large adjacent ovals visible: "
         "popliteal vein (N1, larger) + SSV (N2, smaller). SSV junction with popliteal vein "
         "confirmed. Full SPJ anatomy mapped with tributary branches. TYPE 2 ESCAPE POINT confirmed."),
        ("165–183 s Elimination Test (2 positions)",
         "Probe returns to groin crease (SFJ). Same anatomical position as initial assessment. "
         "N1 (CFV) and N2 (GSV) re-identified at junction. Elimination test performed at SFJ — "
         "POSITIVE. SFJ confirmed as sole escape point for Type 1 shunt. "
         "FINAL DIAGNOSIS: Type 1 + Type 2 combined shunt."),
    ]
    for phase, text in narrative:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(phase + ": ").bold = True
        p.add_run(text)

    # ── Summary scoreboard ─────────────────────────────────────────────────
    doc.add_heading("2. Summary Scoreboard", 2)

    n_steps  = len(results)
    n_reg_ok = sum(1 for r in results if r.region_correct)
    n_g_ok   = sum(1 for r in results if r.guidance_ok)
    n_g3     = sum(1 for r in results if r.guidance_score == 3)
    n_g2     = sum(1 for r in results if r.guidance_score == 2)
    n_g1     = sum(1 for r in results if r.guidance_score == 1)
    n_g0     = sum(1 for r in results if r.guidance_score == 0)
    n_pq_ok  = sum(1 for r in results if r.llm_prompt_ok)
    n_vf_ok  = sum(1 for r in results if r.vein_frame_ok)
    n_wr_ok  = sum(1 for r in results
                   if r.wrong_region_triggered is not None and r.wrong_region_triggered)
    n_wr_tot = sum(1 for r in results if r.wrong_region_triggered is not None)

    sb = doc.add_table(rows=1, cols=3)
    sb.style = "Table Grid"
    _hdr_cell(sb.rows[0].cells[0], "Metric",           "1F497D")
    _hdr_cell(sb.rows[0].cells[1], "Score",            "1F497D")
    _hdr_cell(sb.rows[0].cells[2], "Assessment",       "1F497D")

    def _sb_row(label, score_str, verdict_str, verdict_ok: bool):
        row = sb.add_row().cells
        _body_cell(row[0], label)
        _body_cell(row[1], score_str, bold=True)
        _body_cell(row[2], verdict_str, bold=True,
                   color=PASS_CLR if verdict_ok else FAIL_CLR)

    _sb_row("Region localisation accuracy",
            f"{n_reg_ok}/{n_steps}",
            "PASS" if n_reg_ok == n_steps else f"FAIL ({n_steps-n_reg_ok} wrong)",
            n_reg_ok == n_steps)
    _sb_row("LLM guidance generated (non-error)",
            f"{n_g_ok}/{n_steps}",
            "PASS" if n_g_ok == n_steps else f"FAIL ({n_steps-n_g_ok} empty/error)",
            n_g_ok == n_steps)
    _sb_row("Guidance clinically CORRECT",
            f"{n_g3}/{n_steps}",
            f"CORRECT={n_g3}  PARTIAL={n_g2}  WRONG={n_g1}  NO_RESP={n_g0}",
            n_g3 >= n_steps // 2)
    _sb_row("LLM prompt contains clinical context",
            f"{n_pq_ok}/{n_steps}",
            "PASS" if n_pq_ok == n_steps else f"FAIL ({n_steps-n_pq_ok} missing keywords)",
            n_pq_ok == n_steps)
    _sb_row("Vein frame endpoint returns image",
            f"{n_vf_ok}/{n_steps}",
            "PASS" if n_vf_ok == n_steps else f"FAIL ({n_steps-n_vf_ok} errors)",
            n_vf_ok == n_steps)
    _sb_row("Wrong-region detection",
            f"{n_wr_ok}/{n_wr_tot}" if n_wr_tot else "N/A",
            "PASS" if (n_wr_tot == 0 or n_wr_ok == n_wr_tot) else "FAIL",
            n_wr_tot == 0 or n_wr_ok == n_wr_tot)

    # ── Main clinical step table ───────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("3. Step-by-Step Clinical Assessment", 2)

    # Column headers
    cols = [
        "Time\n(video)", "Phase", "What Surgeon Did\n(Video Finding)",
        "Probe Position\nSent to System", "Region\nDetected",
        "Guidance Text\nGenerated by System",
        "Expected Next Step\n(Ground Truth)",
        "Guidance\nScore", "Prompt\nContext OK?",
        "Vein\nFrame", "Wrong-Region\nDetected?", "Notes"
    ]
    WIDTHS = [Cm(1.8), Cm(2.8), Cm(5.5), Cm(3.0), Cm(1.6),
              Cm(5.5), Cm(4.5), Cm(1.6), Cm(1.4),
              Cm(1.4), Cm(2.2), Cm(3.5)]

    tbl = doc.add_table(rows=1, cols=len(cols))
    tbl.style = "Table Grid"
    for i, (hdr, w) in enumerate(zip(cols, WIDTHS)):
        cell = tbl.rows[0].cells[i]
        _hdr_cell(cell, hdr)
        cell.width = w

    for r in results:
        row = tbl.add_row().cells

        # Time
        _body_cell(row[0], r.video_time, size=8)

        # Phase (short)
        phase_short = r.clinical_phase.replace("Phase ", "Ph").replace(" – ", "\n")
        _body_cell(row[1], phase_short, size=8)

        # Clinical action (truncated)
        action = r.clinical_action[:280]
        _body_cell(row[2], action, size=8)

        # Probe position
        _body_cell(row[3], r.position_label, size=8)

        # Region detected
        reg_ok = r.region_correct
        _body_cell(row[4],
                   f"{r.region_detected}\n({'✓' if reg_ok else '✗'} exp: {r.region_expected})",
                   bold=True,
                   color=PASS_CLR if reg_ok else FAIL_CLR,
                   size=8)

        # Guidance text
        gt = r.guidance_text or "(empty)"
        _body_cell(row[5], gt, size=8)

        # Expected next step
        _body_cell(row[6], r.expected_next_step, size=8)

        # Guidance score
        score_label = SCORE_LABEL.get(r.guidance_score, "?")
        _body_cell(row[7], score_label, bold=True,
                   color=SCORE_CLR.get(r.guidance_score, _rgb(0,0,0)), size=8)

        # Prompt context OK
        pq = r.llm_prompt_ok
        pq_text = "YES" if pq else f"NO\nMissing: {r.missing_kws}"
        _body_cell(row[8], pq_text, bold=True,
                   color=PASS_CLR if pq else FAIL_CLR, size=8)

        # Vein frame
        vf_text = f"YES\n{r.vein_type[:30]}" if r.vein_frame_ok else "FAIL"
        _body_cell(row[9], vf_text, bold=True,
                   color=PASS_CLR if r.vein_frame_ok else FAIL_CLR, size=8)

        # Wrong-region
        if r.wrong_region_triggered is None:
            wr_text, wr_ok = "N/A", True
        else:
            wr_text = f"{'YES ✓' if r.wrong_region_triggered else 'NOT DETECTED ✗'}\n{r.wrong_region_reason[:60]}"
            wr_ok   = r.wrong_region_triggered
        _body_cell(row[10], wr_text, bold=True,
                   color=PASS_CLR if wr_ok else FAIL_CLR, size=8)

        # Notes
        _body_cell(row[11], r.notes[:200], size=7)

    # ── Guidance text deep-dive ────────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("4. Guidance Quality Deep-Dive", 2)
    doc.add_paragraph(
        "For each examination step, the full guidance text is shown alongside "
        "the expected clinical next action and a pass/fail verdict."
    )

    for i, r in enumerate(results):
        doc.add_heading(f"Step {i+1}: {r.video_time} — {r.clinical_phase}", 3)
        p = doc.add_paragraph()
        p.add_run("Surgeon action: ").bold = True
        p.add_run(r.clinical_action + "\n")
        p.add_run("Probe position: ").bold = True
        p.add_run(f"{r.position_label}\n")
        p.add_run("Region detected: ").bold = True
        det_run = p.add_run(f"{r.region_detected}  ({'CORRECT' if r.region_correct else 'WRONG — expected ' + r.region_expected})\n")
        det_run.bold = True
        det_run.font.color.rgb = PASS_CLR if r.region_correct else FAIL_CLR
        p.add_run("System guidance: ").bold = True
        g_run = p.add_run(f'"{r.guidance_text or "(no response)"}"\n')
        g_run.font.color.rgb = SCORE_CLR.get(r.guidance_score, _rgb(0,0,0))
        p.add_run("Expected next step: ").bold = True
        p.add_run(r.expected_next_step + "\n")
        p.add_run("Guidance score: ").bold = True
        sc_run = p.add_run(f"{SCORE_LABEL.get(r.guidance_score,'?')} ({r.guidance_score}/3)\n")
        sc_run.bold = True
        sc_run.font.color.rgb = SCORE_CLR.get(r.guidance_score, _rgb(0,0,0))
        if r.missing_kws:
            p.add_run("Prompt missing keywords: ").bold = True
            p.add_run(str(r.missing_kws) + "\n")
        if r.wrong_region_triggered is not None:
            p.add_run("Wrong-region detection: ").bold = True
            wr_run = p.add_run(
                f"{'CORRECT — flagged' if r.wrong_region_triggered else 'FAILED — not flagged'}\n"
                f"Reason: {r.wrong_region_reason[:120]}\n"
            )
            wr_run.font.color.rgb = PASS_CLR if r.wrong_region_triggered else FAIL_CLR
        p.add_run("Video notes: ").bold = True
        p.add_run(r.notes)

    # ── Findings & Recommendations ─────────────────────────────────────────
    doc.add_page_break()
    doc.add_heading("5. Key Findings & Recommendations", 2)

    findings = []

    if n_reg_ok < n_steps:
        wrong_steps = [r for r in results if not r.region_correct]
        findings.append((
            "FAIL",
            "Region localisation errors",
            "System misclassified probe position at: " +
            ", ".join(f"{r.video_time} ({r.region_detected} ≠ {r.region_expected})"
                      for r in wrong_steps)
        ))
    else:
        findings.append(("PASS", "Region localisation", "All positions correctly classified."))

    if n_g_ok < n_steps:
        findings.append((
            "FAIL",
            "LLM guidance generation failure",
            f"Guidance LLM returned empty/error for {n_steps - n_g_ok}/{n_steps} steps. "
            "Root cause identified: generate_guidance() was calling GROQ_TEXT_MODEL "
            "('openai/gpt-oss-120b') which does not exist on Groq. "
            "Fix: change import to GROQ_FAST_MODEL ('llama-3.1-8b-instant')."
        ))
    else:
        findings.append(("PASS", "LLM guidance generation", f"Responses received for all {n_steps} steps."))

    if n_g3 < n_steps * 0.7:
        findings.append((
            "WARN",
            "Guidance clinical correctness below 70%",
            f"Only {n_g3}/{n_steps} steps received CORRECT guidance. "
            f"PARTIAL: {n_g2}, WRONG: {n_g1}, NO_RESPONSE: {n_g0}. "
            "The 12-word guidance constraint may be too tight to include all required clinical terms."
        ))
    else:
        findings.append(("PASS", "Guidance clinical relevance", f"{n_g3}/{n_steps} steps CORRECT."))

    if n_pq_ok < n_steps:
        miss = [r for r in results if not r.llm_prompt_ok]
        findings.append((
            "WARN",
            "LLM prompt context gaps",
            "Missing keywords in prompt for: " +
            ", ".join(f"{r.video_time}: {r.missing_kws}" for r in miss)
        ))
    else:
        findings.append(("PASS", "LLM prompt clinical context", "All required terms present in all prompts."))

    findings.append((
        "INFO",
        "SPJ prompt does not mention SSV explicitly",
        "The SPJ prompt says 'Saphenopopliteal Junction' but does not include 'SSV' or "
        "'small saphenous vein'. For Type 2 shunt guidance, the prompt should explicitly "
        "mention SSV and its connection to the popliteal vein to guide the LLM correctly."
    ))

    findings.append((
        "INFO",
        "Elimination test context not injected",
        "When the probe returns to SFJ at the end of the examination (elimination test), "
        "the system treats it identically to the initial SFJ scan. It does not know the "
        "surgeon is performing the elimination test. A phase-aware context field or "
        "confirmation of prior SPJ findings would help the guidance explicitly suggest "
        "calf compression and bidirectional PW Doppler."
    ))

    finding_colors = {"PASS": "PASS", "FAIL": "FAIL", "WARN": "WARN", "INFO": "INFO"}
    for ftype, title, desc in findings:
        p = doc.add_paragraph(style="List Bullet")
        color = {"PASS": PASS_CLR, "FAIL": FAIL_CLR,
                 "WARN": WARN_CLR, "INFO": HEAD_CLR}.get(ftype, _rgb(0,0,0))
        run = p.add_run(f"[{ftype}] {title}: ")
        run.bold = True
        run.font.color.rgb = color
        p.add_run(desc)

    doc.save(out_path)
    print(f"\n[OK] Report saved → {out_path}")


# ═══════════════════════════════════════════════════════════════════════════
# Main
# ═══════════════════════════════════════════════════════════════════════════

def main():
    print("=" * 72)
    print("  CYGNUS MED — Task-2 Clinical Scenario Test: Type 1 + 2 Shunt")
    print(f"  Backend : {BASE_URL}")
    print(f"  Time    : {datetime.now():%Y-%m-%d %H:%M:%S}")
    print("=" * 72)

    for i, step in enumerate(STEPS):
        label = f"[{step['video_time']}] {step['clinical_phase']}"
        print(f"\n  Step {i+1:2d}/{len(STEPS)}: {label}")
        r = run_step(step)
        all_results.append(r)

        reg_sym = "✓" if r.region_correct  else "✗"
        g_sym   = "✓" if r.guidance_ok     else "✗"
        pq_sym  = "✓" if r.llm_prompt_ok   else "✗"
        vf_sym  = "✓" if r.vein_frame_ok   else "✗"
        score_s = SCORE_LABEL.get(r.guidance_score, "?")

        print(f"         Region  : {reg_sym} {r.region_detected} (exp: {r.region_expected})")
        print(f"         Guidance: {g_sym} {repr(r.guidance_text[:80])}")
        print(f"         Score   : {score_s}")
        print(f"         Prompt  : {pq_sym}  VeinFrame: {vf_sym}  ({r.elapsed_ms:.0f} ms)")

    print("\n" + "=" * 72)
    n = len(all_results)
    print(f"  Region correct    : {sum(r.region_correct for r in all_results)}/{n}")
    print(f"  Guidance OK       : {sum(r.guidance_ok for r in all_results)}/{n}")
    print(f"  Guidance CORRECT  : {sum(r.guidance_score==3 for r in all_results)}/{n}")
    print(f"  Guidance PARTIAL  : {sum(r.guidance_score==2 for r in all_results)}/{n}")
    print(f"  Prompt context OK : {sum(r.llm_prompt_ok for r in all_results)}/{n}")
    print("=" * 72)

    ts  = datetime.now().strftime("%Y%m%d_%H%M%S")
    out = rf"C:\Users\Krish\Downloads\Task2_Shunt_Type12_Test_{ts}.docx"
    generate_report(all_results, out)
    return 0


if __name__ == "__main__":
    sys.exit(main())
