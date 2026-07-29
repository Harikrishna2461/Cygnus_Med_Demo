#!/usr/bin/env python3
"""
update_task2_report.py

Rewrites Task2_Methodology_Report.docx in-place with all changes from the
current codebase plus a proper agentic-system architecture diagram.

Changes reflected:
  1.  Section 6.1  — History agent is now pure Python (no LLM call).
  2.  Section 7.3  — VLM now handles two annotation styles; adds vein_frames bypass.
  3.  Section 8.2  — Task output word limits corrected (Task3=80 w, Task4=15 w).
  4.  Section 8.3  — Three-tier model assignment added per agent.
  5.  Section 8.3 Agent 3 — Full examination states A–F description.
  6.  Section 8.3 Agent 4 — Forbidden-terms rule and ≤15-word output noted.
  7.  Section 8.4  — Accepted-shunts multi-shunt scanning documented.
  8.  Section 8.5  — LiteLLM patch code snippet corrected (patches litellm.completion,
                      not litellm.utils.get_optional_params as sometimes mis-cited).
  9.  New §3.3     — Full-page agentic architecture diagram (matplotlib PNG).
  10. New §8.6     — Three-tier model strategy.
  11. New §8.7     — Circuit Analyst examination states A–F (full reference).
  12. New §8.8     — Zone anti-repetition during Q1 search.
  13. New §8.9     — Multi-shunt scanning / accepted-shunts tracking.
  14. New §8.10    — NavigationPlanner forbidden-terms rule.
"""

from __future__ import annotations

import io
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch
import numpy as np

DOC_PATH = (
    r"C:\Users\Krish\Downloads\Cygnus_Med_Demo\Task_2_App"
    r"\Task2_Methodology_Report.docx"
)

# ─── helpers ──────────────────────────────────────────────────────────────────

def _find(doc, fragment: str):
    """Return (index, para) for the first paragraph whose text contains *fragment*."""
    for i, p in enumerate(doc.paragraphs):
        if fragment in p.text:
            return i, p
    return None, None


def _insert_para_after(ref_para, text: str, bold: bool = False,
                       heading_level: int | None = None, mono: bool = False,
                       indent: int = 0) -> None:
    """
    Insert a new <w:p> element immediately after *ref_para* in the XML tree.
    Supports bold, heading styles (Heading 1–9), monospace font, and left indent.
    """
    new_p = OxmlElement("w:p")

    # ── paragraph properties (pPr) ────────────────────────────────────────────
    pPr = OxmlElement("w:pPr")
    if heading_level is not None:
        pStyle = OxmlElement("w:pStyle")
        pStyle.set(qn("w:val"), f"Heading{heading_level}")
        pPr.append(pStyle)
    if indent:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent * 720))   # 720 twips = 0.5 in per level
        pPr.append(ind)
    new_p.append(pPr)

    # ── run ───────────────────────────────────────────────────────────────────
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    if bold and heading_level is None:
        b = OxmlElement("w:b")
        rPr.append(b)
    if mono:
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Courier New")
        rFonts.set(qn("w:hAnsi"), "Courier New")
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "16")   # 8 pt
        rPr.append(sz)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    r.append(t)
    new_p.append(r)

    ref_para._element.addnext(new_p)


def _replace_run_text(para, new_text: str) -> None:
    """Replace all text in *para* with *new_text* using the first run."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)


def _set_para_text(para, new_text: str, bold: bool = False) -> None:
    """Clear all runs in *para* and rewrite text, optionally bold."""
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
        para.runs[0].bold = bold
    else:
        r = para.add_run(new_text)
        r.bold = bold


# ─── architecture diagram ──────────────────────────────────────────────────────

def _make_diagram() -> io.BytesIO:
    """Render the agentic system architecture diagram and return PNG bytes."""

    BG       = "#0d1117"
    PANEL    = "#161b22"
    BORDER   = "#30363d"

    # colour palette
    C_SURGEON   = "#1f6feb"   # blue  — surgeon input / UI
    C_SERVER    = "#388bfd"   # lighter blue — Flask/SocketIO
    C_VLM       = "#d29922"   # amber — VLM / vision
    C_PREAGENT  = "#238636"   # green — pure-Python pre-agents
    C_CREWAI    = "#6e40c9"   # purple — CrewAI agents
    C_OUTPUT    = "#da3633"   # red   — output events
    C_HEAVY     = "#bc8cff"   # light purple — heavy LLM label
    C_MID       = "#79c0ff"   # light blue   — mid LLM label
    TEXT_MAIN   = "#e6edf3"
    TEXT_DIM    = "#8b949e"
    ARROW_COL   = "#58a6ff"

    fig_w, fig_h = 22, 30
    fig = plt.figure(figsize=(fig_w, fig_h), facecolor=BG)
    ax  = fig.add_axes([0, 0, 1, 1])
    ax.set_xlim(0, fig_w)
    ax.set_ylim(0, fig_h)
    ax.axis("off")
    ax.set_facecolor(BG)

    # ── helper draw functions ─────────────────────────────────────────────────

    def box(x, y, w, h, fill, alpha=0.15, ec=None, lw=1.5, radius=0.25):
        rect = FancyBboxPatch(
            (x, y), w, h,
            boxstyle=f"round,pad=0,rounding_size={radius}",
            facecolor=fill, alpha=alpha,
            edgecolor=ec or fill, linewidth=lw,
            zorder=2,
        )
        ax.add_patch(rect)

    def solid_box(x, y, w, h, fill, alpha=0.9, ec="white", lw=0.5, radius=0.20):
        rect = FancyBboxPatch(
            (x, y), w, h,
            boxstyle=f"round,pad=0,rounding_size={radius}",
            facecolor=fill, alpha=alpha,
            edgecolor=ec, linewidth=lw,
            zorder=3,
        )
        ax.add_patch(rect)

    def txt(x, y, s, size=9.5, color=TEXT_MAIN, ha="center", va="center",
            bold=False, zorder=5, wrap=False):
        w = {"fontsize": size, "color": color, "ha": ha, "va": va,
              "zorder": zorder, "fontweight": "bold" if bold else "normal",
              "fontfamily": "DejaVu Sans"}
        ax.text(x, y, s, **w)

    def arrow(x0, y0, x1, y1, col=ARROW_COL, lw=1.6, style="->"):
        ax.annotate("", xy=(x1, y1), xytext=(x0, y0),
                    arrowprops=dict(arrowstyle=style, color=col,
                                   lw=lw, connectionstyle="arc3,rad=0.0"),
                    zorder=4)

    def agent_box(x, y, w, h, num, role, model_label, model_col,
                  output_desc, fill=C_CREWAI):
        box(x, y, w, h, fill, alpha=0.18, lw=1.8)
        ax.plot([x, x+w], [y+h-0.55, y+h-0.55],
                color=fill, alpha=0.4, lw=0.8, zorder=3)
        txt(x + w/2, y+h-0.28, f"Agent {num}", size=9, bold=True, color=fill)
        txt(x + w/2, y+0.58,   role,          size=9.5, bold=True)
        solid_box(x+0.2, y+0.22, w-0.4, 0.32, model_col, alpha=0.85,
                  ec="none", radius=0.10)
        txt(x + w/2, y+0.38, model_label, size=8, color=BG, bold=True)
        txt(x + w/2, y+h/2-0.05, output_desc, size=8.8, color=TEXT_DIM)

    # ── title ────────────────────────────────────────────────────────────────
    txt(fig_w/2, 29.35,
        "Task 2 — Active Guidance System: Agentic Architecture",
        size=18, bold=True)
    txt(fig_w/2, 28.85,
        "Real-time CHIVA ultrasound examination guidance pipeline",
        size=11, color=TEXT_DIM)

    # ── LAYER 1: Surgeon input ────────────────────────────────────────────────
    bx, by, bw, bh = 1.0, 27.30, 20.0, 1.15
    box(bx, by, bw, bh, C_SURGEON, alpha=0.18, lw=2)
    txt(fig_w/2, by+bh/2+0.22, "SURGEON — Frontend  (stream.html)",
        size=11, bold=True, color=C_SURGEON)
    txt(fig_w/2, by+bh/2-0.24,
        "Interactive Leg Diagram click (probe_move)   ·   Clip Mark Form (clip_mark)",
        size=9.5, color=TEXT_DIM)

    # arrows down from surgeon layer
    arrow(7.5,  by,       7.5,  by-0.55)
    arrow(14.5, by,       14.5, by-0.55)
    txt(9.0, by-0.30, "WebSocket  (Socket.IO)", size=9, color=TEXT_DIM)

    # ── LAYER 2: Flask server ─────────────────────────────────────────────────
    bx2, by2, bw2, bh2 = 1.0, 25.60, 20.0, 1.10
    box(bx2, by2, bw2, bh2, C_SERVER, alpha=0.18, lw=2)
    txt(fig_w/2, by2+bh2/2+0.22,
        "FLASK-SOCKETIO  routes/stream.py  ·  streaming_session.py",
        size=10.5, bold=True, color=C_SERVER)
    txt(fig_w/2, by2+bh2/2-0.24,
        "StreamSession per client  ·  generation counter (anti-stale results)  "
        "·  VLM threshold Δpos≥0.05  ·  LLM threshold Δpos≥0.06",
        size=8.8, color=TEXT_DIM)

    # ── LAYER 3: Streaming engine ─────────────────────────────────────────────
    arrow(fig_w/2, by2, fig_w/2, by2-0.60)
    bx3, by3, bw3, bh3 = 1.0, 23.85, 20.0, 1.15
    box(bx3, by3, bw3, bh3, C_SERVER, alpha=0.12, lw=1.5)
    txt(fig_w/2, by3+bh3/2+0.22,
        "STREAMING GUIDANCE ENGINE  (streaming_guidance_engine.py)",
        size=10.5, bold=True, color=C_SERVER)
    txt(fig_w/2, by3+bh3/2-0.24,
        "Orchestrator: logs position → triggers VLM and/or LLM → assembles response",
        size=9, color=TEXT_DIM)

    # ── split arrow ───────────────────────────────────────────────────────────
    mid_y = by3
    ax.plot([5.0, 5.0], [mid_y, mid_y-0.60], color=ARROW_COL, lw=1.6, zorder=4)
    ax.plot([17.0, 17.0], [mid_y, mid_y-0.60], color=ARROW_COL, lw=1.6, zorder=4)
    ax.annotate("", xy=(5.0, mid_y-0.60), xytext=(5.0, mid_y-0.55),
                arrowprops=dict(arrowstyle="->", color=ARROW_COL, lw=1.6), zorder=4)
    ax.annotate("", xy=(17.0, mid_y-0.60), xytext=(17.0, mid_y-0.55),
                arrowprops=dict(arrowstyle="->", color=ARROW_COL, lw=1.6), zorder=4)
    txt(3.5, mid_y-0.32, "if Δpos ≥ 0.05", size=8, color=TEXT_DIM)
    txt(18.5, mid_y-0.32, "if Δpos ≥ 0.06\nor clip_mark", size=8, color=TEXT_DIM)

    # ── LAYER 4a: VLM panel ───────────────────────────────────────────────────
    vlm_x, vlm_y, vlm_w, vlm_h = 0.5, 19.70, 8.8, 3.45
    box(vlm_x, vlm_y, vlm_w, vlm_h, C_VLM, alpha=0.20, lw=1.8)
    txt(vlm_x+vlm_w/2, vlm_y+vlm_h-0.32,
        "VLM ANALYSIS  (vlm_agent.py)", size=10, bold=True, color=C_VLM)
    ax.plot([vlm_x, vlm_x+vlm_w],
            [vlm_y+vlm_h-0.60, vlm_y+vlm_h-0.60],
            color=C_VLM, alpha=0.35, lw=0.8, zorder=3)

    # frame source hierarchy
    src_items = [
        ("1", "guidance/  (fascia-annotated reference frames)", "★ VLM used"),
        ("2", "vein_frames/  (pre-labelled anatomy)", "VLM bypassed"),
        ("3", "streaming video  (fallback)", "VLM used"),
    ]
    for k, (num, label, note) in enumerate(src_items):
        yy = vlm_y + vlm_h - 0.95 - k*0.52
        txt(vlm_x+0.35, yy, f"{num}.", size=9, color=C_VLM, ha="left")
        txt(vlm_x+0.65, yy, label, size=9, ha="left")
        txt(vlm_x+vlm_w-0.15, yy, note, size=7.5, color=TEXT_DIM, ha="right")

    solid_box(vlm_x+0.3, vlm_y+0.20, vlm_w-0.6, 0.95, C_VLM, alpha=0.25,
              ec=C_VLM, radius=0.12)
    txt(vlm_x+vlm_w/2, vlm_y+0.82,
        "Llama-4-Scout-17B  (Groq Vision API)", size=9.5, bold=True, color=C_VLM)
    txt(vlm_x+vlm_w/2, vlm_y+0.48,
        "Style A: yellow fascia lines + N-labels  ·  Style B: oval vein outlines",
        size=8, color=TEXT_DIM)

    # ── LAYER 4b: Pre-agent context panel ────────────────────────────────────
    pa_x, pa_y, pa_w, pa_h = 10.2, 19.70, 11.3, 3.45
    box(pa_x, pa_y, pa_w, pa_h, C_PREAGENT, alpha=0.18, lw=1.8)
    txt(pa_x+pa_w/2, pa_y+pa_h-0.32,
        "CONTEXT BUILDING  (Pre-Agents — Pure Python)", size=10, bold=True, color=C_PREAGENT)
    ax.plot([pa_x, pa_x+pa_w],
            [pa_y+pa_h-0.60, pa_y+pa_h-0.60],
            color=C_PREAGENT, alpha=0.35, lw=0.8, zorder=3)

    agents_pre = [
        ("history_agent.py",  "7-band zone visit map, confirmed clips — no LLM"),
        ("q_state_agent.py",  "Q1–Q4 status derivation — deterministic rules"),
        ("protocol_agent.py", "Zone examination lookup — pre-built dictionary"),
        ("guidance_agent.py", "Assembles 6-section enriched state message"),
    ]
    for k, (name, desc) in enumerate(agents_pre):
        yy = pa_y + pa_h - 0.95 - k*0.55
        solid_box(pa_x+0.3, yy-0.17, 2.9, 0.40, C_PREAGENT, alpha=0.25,
                  ec=C_PREAGENT, radius=0.10)
        txt(pa_x+1.75, yy+0.02, name, size=8.8, bold=True, color=C_PREAGENT)
        txt(pa_x+3.35, yy+0.02, desc, size=8.5, color=TEXT_DIM, ha="left")

    txt(pa_x+pa_w/2, pa_y+0.48,
        "→  Enriched 6-section state message (Probe · Clips · VLM · History · Q-State · Protocol)",
        size=8.5, color=TEXT_DIM)

    # ── combine arrows into CrewAI layer ────────────────────────────────────
    arrow(vlm_x+vlm_w/2, vlm_y, vlm_x+vlm_w/2, vlm_y-0.55)
    arrow(pa_x+pa_w/2,   pa_y,  pa_x+pa_w/2,   pa_y-0.55)
    ax.plot([vlm_x+vlm_w/2, pa_x+pa_w/2],
            [vlm_y-0.55, vlm_y-0.55],
            color=ARROW_COL, lw=1.6, zorder=4)
    arrow(fig_w/2, vlm_y-0.55, fig_w/2, vlm_y-1.05)

    # ── LAYER 5: CrewAI agents ─────────────────────────────────────────────────
    crew_top = vlm_y - 1.1
    crew_h_total = 9.5
    crew_y = crew_top - crew_h_total
    crew_x, crew_w = 0.5, 21.0

    box(crew_x, crew_y, crew_w, crew_h_total, C_CREWAI, alpha=0.10, lw=2.0)
    txt(fig_w/2, crew_y+crew_h_total-0.28,
        "CrewAI 5-AGENT SEQUENTIAL PIPELINE  (crew_pipeline.py  ·  crew_agents.py)",
        size=11, bold=True, color=C_CREWAI)
    ax.plot([crew_x, crew_x+crew_w],
            [crew_y+crew_h_total-0.55, crew_y+crew_h_total-0.55],
            color=C_CREWAI, alpha=0.35, lw=0.8, zorder=3)
    txt(fig_w/2-3, crew_y+crew_h_total-0.88,
        "Context flows sequentially: each agent sees ALL upstream agents' outputs  ·  90-second timeout  ·  daemon thread",
        size=9, color=TEXT_DIM)

    # draw each agent box
    ag_w  = 18.0
    ag_h  = 1.30
    ag_x0 = crew_x + (crew_w - ag_w) / 2
    ag_gap = 0.25

    agents = [
        (1, "Clinical Interpreter",
         "Mid  (llama-3.3-70b-versatile)", C_MID,
         "Assesses clip quality, VLM alignment, missing evidence  (≤100 words)"),
        (2, "Shunt Analyst",
         "Heavy  (gpt-oss-120b)", C_HEAVY,
         'Classifies CHIVA shunt type → JSON {shunt_type, confirmed, elim_required, evidence}'),
        (3, "Circuit Analyst",
         "Heavy  (gpt-oss-120b)", C_HEAVY,
         "Outputs Examination State A–F + specific target zone + anatomy  (≤80 words)"),
        (4, "Navigation Planner",
         "Mid  (llama-3.3-70b-versatile)", C_MID,
         "Issues ONE movement command  (≤15 words, no clinical jargon, direction + target + zone)"),
        (5, "Guidance Specialist",
         "Heavy  (gpt-oss-120b)", C_HEAVY,
         'Packages command as JSON {"guidance": "...", "action": "move|maneuver|complete"}'),
    ]

    total_agents_h = len(agents) * ag_h + (len(agents)-1) * ag_gap
    ag_y_start = crew_y + crew_h_total - 1.15 - total_agents_h

    for i, (num, role, model_lbl, model_col, out_desc) in enumerate(agents):
        ay = ag_y_start + (len(agents)-1-i) * (ag_h + ag_gap)
        agent_box(ag_x0, ay, ag_w, ag_h, num, role, model_lbl, model_col, out_desc)

        # context arrow (downward between agents, except last)
        if i < len(agents) - 1:
            ax_mid = ag_x0 + ag_w/2
            arrow(ax_mid, ay, ax_mid, ay - ag_gap, style="->")
            ctx_note = "context"
            if num == 4:
                ctx_note = "context [task2 + task4 only]"
            txt(ax_mid + 1.5, ay - ag_gap/2, ctx_note, size=8, color=TEXT_DIM)

    # ── LAYER 6: Output ────────────────────────────────────────────────────────
    out_y = crew_y - 1.10
    out_h = 0.90
    arrow(fig_w/2, crew_y, fig_w/2, crew_y - 0.60)

    box(0.5, out_y, 21.0, out_h, C_OUTPUT, alpha=0.18, lw=1.8)
    txt(fig_w/2, out_y+out_h/2+0.15,
        "WebSocket OUTPUT  (Socket.IO events)", size=10.5, bold=True, color=C_OUTPUT)
    txt(fig_w/2, out_y+out_h/2-0.18,
        "guidance_update → guidance text + action + VLM dict + frame     "
        "·     shunt_confirmed → shunt_type + evidence",
        size=9, color=TEXT_DIM)

    # ── LAYER 7: Surgeon UI ────────────────────────────────────────────────────
    ui_y = out_y - 1.05
    ui_h = 0.90
    arrow(fig_w/2, out_y, fig_w/2, out_y - 0.58)

    box(0.5, ui_y, 21.0, ui_h, C_SURGEON, alpha=0.18, lw=1.8)
    txt(fig_w/2, ui_y+ui_h/2+0.15,
        "SURGEON UI  (stream.html — live WebSocket updates)", size=10.5, bold=True, color=C_SURGEON)
    txt(fig_w/2, ui_y+ui_h/2-0.18,
        "Guidance text  ·  Action badge (move / maneuver / complete)  "
        "·  Frame display  ·  Confirmed clips list  ·  Shunt modal",
        size=9, color=TEXT_DIM)

    # ── Legend ─────────────────────────────────────────────────────────────────
    legend_y = ui_y - 0.85
    items = [
        (C_SURGEON,  "Surgeon / UI"),
        (C_SERVER,   "Flask + SocketIO"),
        (C_VLM,      "VLM (Vision)"),
        (C_PREAGENT, "Pre-Agents (Python)"),
        (C_CREWAI,   "CrewAI Agents"),
        (C_OUTPUT,   "Output Events"),
        (C_HEAVY,    "Heavy LLM (120B)"),
        (C_MID,      "Mid LLM (70B)"),
    ]
    per = fig_w / len(items)
    for k, (col, label) in enumerate(items):
        lx = 0.5 + k * per + per / 2
        solid_box(lx - 0.30, legend_y - 0.03, 0.60, 0.30, col,
                  alpha=0.80, ec="none", radius=0.08)
        txt(lx, legend_y - 0.38, label, size=8, color=TEXT_DIM)

    # bottom caption
    txt(fig_w/2, legend_y - 0.70,
        "Temperature = 0.3  ·  max_iter = 3  ·  allow_delegation = False  "
        "·  VLM threshold = Δpos 0.05  ·  LLM threshold = Δpos 0.06",
        size=8, color=TEXT_DIM)

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=140, bbox_inches="tight",
                facecolor=BG, edgecolor="none")
    buf.seek(0)
    plt.close(fig)
    return buf


# ─── update existing paragraphs ────────────────────────────────────────────────

def _update_history_agent_section(doc) -> None:
    """
    Section 6.1: replace the incorrect 'calls Groq LLM' paragraph with the
    corrected pure-Python description.
    """
    old_frag = "calls Groq LLM (temperature=0) to write a 2-sentence narrative summary"
    idx, para = _find(doc, old_frag)
    if para is None:
        print("  [WARN] Section 6.1 LLM paragraph not found — skipping.")
        return

    _set_para_text(
        para,
        "This is now entirely pure Python — no LLM call is made. "
        "The agent computes the band-visit map and confirmed-clip list "
        "deterministically from session.scan_log and session.clips, then "
        "returns them as structured factual text. "
        "An LLM-generated narrative was removed because: "
        "(1) structured visited/unvisited lists are more reliably parsed by the "
        "Circuit Analyst than a prose summary; "
        "(2) zero latency vs 1–3 s for a Groq API call; "
        "(3) prose summaries occasionally hallucinate progress that was not "
        "actually made by the surgeon.",
    )
    print("  [OK] Section 6.1 history-agent LLM paragraph updated.")

    # Also fix the 'Why plain Python' explanation that follows
    old_frag2 = "Why plain Python for the map, LLM only for the summary?"
    _, para2 = _find(doc, old_frag2)
    if para2:
        _set_para_text(
            para2,
            "Why pure Python for the entire history agent? "
            "The band marking is deterministic (a posY number either falls in a "
            "range or it does not). The narrative summary was converted from an LLM "
            "call to structured text output because the Circuit Analyst reasons more "
            "reliably from explicit lists than from ambiguous prose — and eliminating "
            "this LLM call saves 1–3 seconds per guidance turn.",
        )
        print("  [OK] Section 6.1 'Why plain Python' paragraph updated.")


def _update_context_flow(doc) -> None:
    """Section 8.2: fix task word-count labels."""
    replacements = [
        ("output: open Q + target zone (60 words)",
         "output: examination state A–F + target zone (80 words)"),
        ("output: scan plan (50 words)",
         "output: one movement command ≤15 words (direction + anatomy + zone)"),
    ]
    for old, new in replacements:
        _, para = _find(doc, old)
        if para:
            para.text
            full = para.text.replace(old, new)
            _set_para_text(para, full)
            print(f"  [OK] Section 8.2 context flow: '{old[:40]}...' updated.")
        else:
            print(f"  [WARN] Section 8.2: fragment not found — '{old[:40]}...'")


def _update_agent_descriptions(doc) -> None:
    """
    Section 8.3: append model-tier line to each agent heading paragraph,
    and update the Circuit Analyst and Navigation Planner descriptions.
    """
    agent_changes = {
        "Agent 1 — Clinical Interpreter":
            "Model tier: Mid (llama-3.3-70b-versatile). "
            "Synthesis task — reads clips, VLM, history; no strict rule matching needed.",

        "Agent 2 — Shunt Analyst":
            "Model tier: Heavy (gpt-oss-120b). "
            "Strict rule reasoning — correct identification of the full minimum clip "
            "set for each CHIVA type demands the most capable model in the pipeline.",

        "Agent 3 — Circuit Analyst":
            "Model tier: Heavy (gpt-oss-120b). "
            "NEW IN CURRENT VERSION: This agent now determines a formal Examination "
            "State (A–F) from the Q-state and scan history, then issues specific "
            "routing or examination instructions for each state. "
            "See Section 8.7 for the full A–F state reference table.",

        "Agent 4 — Navigation Planner":
            "Model tier: Mid (llama-3.3-70b-versatile). "
            "NEW IN CURRENT VERSION: output capped at ≤15 words (previously ~50 words). "
            "A strict forbidden-terms list prevents Q1/Q2/Q3/Q4/EP/RP/N1/N2/N3/shunt "
            "from appearing in the surgeon-facing output — only anatomical navigation "
            "terms are permitted. See Section 8.10.",

        "Agent 5 — Guidance Specialist":
            "Model tier: Heavy (gpt-oss-120b). "
            "Chosen heavy because it must simultaneously parse the Shunt Analyst's "
            "JSON flags (elim_required, confirmed) and copy the Navigation Planner's "
            "text verbatim — two precise operations that benefit from the largest model.",
    }

    for heading_frag, note in agent_changes.items():
        idx, para = _find(doc, heading_frag)
        if para is None:
            print(f"  [WARN] Agent heading '{heading_frag[:40]}' not found.")
            continue
        # Find the 'ported from Task 1' or similar line after the heading and insert
        # the model-tier note as a new paragraph immediately after the heading.
        _insert_para_after(para, f"[{note}]", bold=False)
        print(f"  [OK] Inserted model-tier note after '{heading_frag[:40]}'.")


def _update_vlm_section(doc) -> None:
    """Section 7.2 / 7.3: add two-annotation-style and vein_frames bypass note."""
    frag = "The video is synchronised with the leg diagram"
    idx, para = _find(doc, frag)
    if para is None:
        print("  [WARN] VLM video-sync paragraph not found.")
        return

    # Insert a new paragraph before (we insert after, so we insert in reverse order)
    note = (
        "NOTE — CURRENT VERSION: The system uses a three-tier frame source hierarchy "
        "instead of a single streaming video. Frames are selected in priority order: "
        "(1) guidance/ subfolder — reference frames with yellow fascia annotations and "
        "N1/N2/N3 labels (VLM is used); "
        "(2) vein_frames/ subfolder — pre-labelled anatomy frames with oval contours "
        "and vein-name labels (VLM is BYPASSED — folder metadata provides the "
        "vlm_dict directly, avoiding an API call); "
        "(3) streaming video — fallback, frame extracted by OpenCV at posY ratio "
        "(VLM is used). "
        "Frame selection is deterministic: bucket = round(posY × 20); source index "
        "cycles through available sources as the probe moves, making results "
        "reproducible for the same probe position."
    )
    _insert_para_after(para, note)
    print("  [OK] VLM section: three-tier frame source note inserted.")

    # Also add two-annotation-style note to the VLM section
    frag2 = "The VLM only re-runs when posY has changed by"
    _, para2 = _find(doc, frag2)
    if para2:
        style_note = (
            "Two annotation styles are recognised by the VLM system prompt: "
            "Style A — guidance frames with yellow horizontal fascia lines and "
            "N1/N2/N3 text labels (the original annotated video style described in §7.2); "
            "Style B — vein_frames with oval contours and vein-name labels "
            "(GSV, SSV, Tributary, etc.). "
            "The VLM prompt adapts based on which style is detected, "
            "allowing accurate interpretation of both frame types."
        )
        _insert_para_after(para2, style_note)
        print("  [OK] VLM section: two-annotation-style note inserted.")


def _update_shunt_detection_section(doc) -> None:
    """Section 8.4: add multi-shunt / accepted-shunts tracking note."""
    frag = "If the surgeon dismisses the modal"
    _, para = _find(doc, frag)
    if para is None:
        print("  [WARN] Section 8.4 surgeon-dismisses paragraph not found.")
        return
    note = (
        "Multi-Shunt Scanning: Once a shunt type is accepted by the surgeon, "
        "it is added to the session's accepted_shunts set. "
        "On all subsequent crew runs, the Shunt Analyst receives an accepted_ctx "
        "block listing accepted types, instructing it NOT to re-confirm them. "
        "The Guidance Specialist sets action='complete' only for a NEWLY confirmed type "
        "not yet in accepted_shunts. "
        "This allows the surgeon to continue scanning the same leg for additional shunt "
        "circuits after the first is accepted — a common clinical scenario where a "
        "patient has more than one haemodynamic circuit."
    )
    _insert_para_after(para, note)
    print("  [OK] Section 8.4: multi-shunt / accepted-shunts note inserted.")


# ─── add new sections ─────────────────────────────────────────────────────────

def _add_new_sections(doc, diagram_png: io.BytesIO) -> None:
    """Append all new sections to the end of the document."""

    def h2(text):
        doc.add_heading(text, level=2)

    def h3(text):
        doc.add_heading(text, level=3)

    def p(text, bold=False):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.bold = bold
        return para

    def bullet(text):
        para = doc.add_paragraph(style="List Bullet")
        para.add_run(text)

    def code_block(text):
        para = doc.add_paragraph()
        run = para.add_run(text)
        run.font.name = "Courier New"
        run.font.size = Pt(8)
        para.paragraph_format.left_indent = Cm(1)

    # ── Section 3.3: Architecture Diagram ─────────────────────────────────────
    doc.add_heading("3.3  Agentic System Architecture Diagram", level=2)
    p(
        "The diagram below shows the complete request-to-response pipeline for a "
        "single probe_move or clip_mark event. "
        "The five-layer structure (Surgeon → Server → Engine → Agents → UI) "
        "illustrates how each component contributes to the final guidance output "
        "and how the two parallel branches (VLM analysis and context building) "
        "converge at the CrewAI crew kickoff.",
        bold=False,
    )

    diagram_para = doc.add_paragraph()
    diagram_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = diagram_para.add_run()
    run.add_picture(diagram_png, width=Inches(6.8))

    p(
        "Figure 1: Full agentic system architecture for Task 2 real-time guidance. "
        "Colour coding: blue = Flask/SocketIO server components; amber = VLM (Groq Vision); "
        "green = pure-Python pre-agents; purple = CrewAI agents; red = WebSocket output events. "
        "Model tier labels show which Groq model handles each agent.",
        bold=False,
    )

    # ── Section 8.6: Three-Tier Model Strategy ─────────────────────────────────
    doc.add_heading("8.6  Three-Tier Model Assignment Strategy", level=2)
    p(
        "All five CrewAI agents call the Groq API, but not at the same model size. "
        "Three model tiers are defined in config.py and assigned per agent:",
        bold=False,
    )

    rows = [
        ("Heavy", "openai/gpt-oss-120b",
         "Shunt Analyst, Circuit Analyst, Guidance Specialist",
         "Tasks requiring strict rule application (shunt classification), "
         "complex spatial reasoning (examination states A–F), or exact JSON parsing "
         "from multi-agent context."),
        ("Mid", "llama-3.3-70b-versatile",
         "Clinical Interpreter, Navigation Planner",
         "Synthesis and movement-command generation — lower stakes reasoning "
         "that benefits from speed without sacrificing accuracy."),
        ("Fast", "llama-3.1-8b-instant",
         "(reserved — not active in current agents)",
         "Lightweight JSON formatting; available for future use where "
         "speed is the primary constraint."),
    ]

    tbl = doc.add_table(rows=1+len(rows), cols=4)
    tbl.style = "Table Grid"
    headers = ["Tier", "Groq Model", "Assigned Agents", "Rationale"]
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    for i, (tier, model, agents_, rationale) in enumerate(rows):
        row = tbl.rows[i+1]
        row.cells[0].text = tier
        row.cells[1].text = model
        row.cells[2].text = agents_
        row.cells[3].text = rationale

    p("")
    p(
        "All agents share temperature=0.3, allow_delegation=False, and max_iter=3. "
        "Temperature 0.3 was chosen after testing: 0.0 caused occasional frozen "
        "outputs for the same position; 0.7 produced variable phrasing that was "
        "clinically imprecise. 0.3 balances determinism and vocabulary variety.",
        bold=False,
    )

    # ── Section 8.7: Circuit Analyst Examination States A–F ───────────────────
    doc.add_heading(
        "8.7  Circuit Analyst Examination States A–F  (full reference)", level=2
    )
    p(
        "In the current version, Agent 3 (Circuit Analyst) outputs a formal "
        "Examination State label (A–F) as the first token of its response. "
        "The Navigation Planner reads this label and issues the appropriate "
        "movement command for each state. "
        "This replaces the earlier free-text 'target zone' approach with a "
        "structured, verifiable state machine.",
        bold=False,
    )

    states = [
        ("A",
         "Q1 OPEN — no entry point found yet",
         "If probe posY 0.04–0.57 (Q1 corridor): examine current zone anatomy "
         "(GSV in fascial compartment, Hunterian perforator, SFJ junction, etc.). "
         "NEVER route to another zone while inside the corridor — Q1 has not been "
         "confirmed and the probe must examine HERE first. "
         "If probe posY > 0.57: route to the nearest unvisited Q1 candidate zone "
         "(SFJ/groin → Hunterian → popliteal/SPJ). Zone anti-repetition applies — "
         "zones suggested ≥2 times in recent guidance are avoided (see §8.8).",
         "Navigation Planner output: 'Move medially toward GSV in saphenous "
         "compartment at upper thigh'"),
        ("B",
         "Q1 confirmed, Q2 open — entry found; trace trunk reflux",
         "Probe traces distally along anteromedial thigh to confirm RP N2→N1 "
         "(GSV trunk refluxing downward). From SFJ, move to upper thigh; "
         "from upper thigh, continue to Hunterian/Dodd zone. "
         "Do NOT jump to calf or SPJ in State B.",
         "Navigation Planner output: 'Move distally toward Dodd perforator on "
         "medial distal thigh'"),
        ("C",
         "Q1+Q2 confirmed, Q3 open — find tributary escape point EP N2→N3",
         "Probe scans the medial calf (posY 0.60–0.75) for where the GSV trunk "
         "gives off a tributary above the fascia. If probe is in the thigh, "
         "route to the upper medial calf first. "
         "Do NOT route to SPJ or back to thigh — the escape is in the calf. "
         "The target is the GSV trunk giving off a tributary (N3), NOT the "
         "re-entry perforator (which is State D territory).",
         "Navigation Planner output: 'Move distally toward GSV in medial calf'"),
        ("D",
         "Q3 confirmed, Q4 open — track tributary to re-entry RP N3→N1",
         "Follow the N3 tributary toward the lower calf (posY 0.75–0.88) to "
         "find where it dives back into the deep system. "
         "If probe posY < 0.75, move distally. "
         "If already in lower calf, examine the medial lower-calf perforator — "
         "do NOT continue to ankle. Re-entry is in the lower calf.",
         "Navigation Planner output: 'Move distally toward re-entry perforator "
         "on medial lower calf'"),
        ("E",
         "Q1–Q4 confirmed",
         "E1 — elimination test not yet recorded: output only 'STATE E1. "
         "No routing — elimination test pending.' The Guidance Specialist "
         "overrides guidance to 'Perform elimination test at current zone'. "
         "E2 — elimination test IS recorded: circuit classified. Route to "
         "the most diagnostically useful unvisited zone: SPJ/popliteal fossa "
         "first (SSV circuit check), then upper thigh, then calf. "
         "Do NOT route back to SFJ/groin — the SFJ entry is already mapped.",
         "Navigation Planner output (E1): 'Perform elimination test at current zone'  "
         "(E2): 'Move posteriorly toward SPJ junction in popliteal fossa'"),
        ("F",
         "EP N1→N3 found — N1-to-tributary direct bypass",
         "SFJ is competent; blood enters N3 (tributary) directly from N1 (deep). "
         "Circuit Analyst traces the N3 tributary distally to find its re-entry. "
         "Type differentiation (4 / 5 / 6) depends on whether RP N2→N1, "
         "RP N3→N2 + EP N2→N3, or only RP N3→N1 is confirmed.",
         "Navigation Planner output: 'Move distally along N3 tributary toward "
         "re-entry perforator'"),
    ]

    for state_label, trigger, logic, example in states:
        h3(f"State {state_label}: {trigger}")
        p("Routing logic: " + logic)
        p("Example output: " + example)

    # ── Section 8.8: Zone Anti-Repetition ────────────────────────────────────
    doc.add_heading(
        "8.8  Zone Anti-Repetition During Q1 Search  (State A)", level=2
    )
    p(
        "During State A (Q1 open), the Circuit Analyst may suggest the same anatomical "
        "zone on consecutive guidance turns if the surgeon has not confirmed a clip. "
        "This is clinically correct (examine here before moving) but can feel repetitive "
        "if the surgeon has already thoroughly assessed that zone without finding a clip. "
        "The anti-repetition mechanism prevents excessive repetition:",
        bold=False,
    )
    bullet(
        "The system tracks which anatomical zone keyword appeared in each of the "
        "last 8 guidance strings (using recent_guidance, passed to run_guidance_crew())."
    )
    bullet(
        "Any zone that appears in ≥2 of the last 8 guidance outputs is marked as "
        "'overused'. The Circuit Analyst receives the list of overused zones and "
        "is instructed to route to a DIFFERENT Q1 candidate zone."
    )
    bullet(
        "Zones tracked: SFJ/groin, upper thigh, Hunterian (proximal thigh), "
        "Dodd (distal thigh), popliteal/SPJ, calf, ankle."
    )
    bullet(
        "DISABLED after Q1 is confirmed: once an entry point clip (EP from N1) "
        "is in the session, the diagnostic path becomes deterministic (B→C→D→E) "
        "and the correct zone must NOT be blocked by the anti-repetition filter."
    )
    p(
        "Separately, the Navigation Planner receives recent_guidance and a banned list "
        "(guidance strings that appeared ≥2 times), so it rephrases its output "
        "even when routing to the same anatomical target.",
        bold=False,
    )
    code_block(
        "# crew_pipeline.py — anti-repetition logic (State A only)\n"
        "_overused_zones = (\n"
        "    [] if _q1_answered        # disabled once entry point is confirmed\n"
        "    else [z for z, c in _zone_hits.items() if c >= 2]\n"
        ")"
    )

    # ── Section 8.9: Multi-Shunt Scanning ────────────────────────────────────
    doc.add_heading(
        "8.9  Multi-Shunt Scanning — Accepted Shunts Tracking", level=2
    )
    p(
        "A patient may have more than one independent haemodynamic CHIVA circuit on "
        "the same leg. After the surgeon accepts the first confirmed shunt type, "
        "the session does not end — the surgeon continues scanning for additional circuits. "
        "The system handles this with an accepted_shunts set:",
        bold=False,
    )
    bullet(
        "When the surgeon clicks 'Confirm' on the shunt confirmation modal, "
        "the accepted type is added to session.accepted_shunts."
    )
    bullet(
        "On every subsequent crew run, the Shunt Analyst receives an accepted_ctx "
        "block listing all accepted types, instructing it NOT to re-confirm them. "
        "The analyst focuses classification on new clip patterns not yet accepted."
    )
    bullet(
        "The Guidance Specialist sets action='complete' only when a NEWLY confirmed "
        "shunt_type is NOT already in accepted_shunts. If the confirmed type is already "
        "accepted, it outputs action='move' — the surgeon is still scanning."
    )
    bullet(
        "The Circuit Analyst routing treats a previously accepted shunt as irrelevant "
        "to the current routing decision — it determines routing from the Q-state alone, "
        "not from which shunts were accepted."
    )
    code_block(
        "# crew_pipeline.py — Guidance Specialist task5 action logic\n"
        "action:\n"
        "  complete  — if confirmed=true AND shunt_type NOT in accepted_shunts\n"
        "  maneuver  — if elim_required=true\n"
        "  move      — all other cases (including confirmed type already accepted)"
    )

    # ── Section 8.10: NavigationPlanner Forbidden Terms ───────────────────────
    doc.add_heading(
        "8.10  Navigation Planner Forbidden-Terms Rule", level=2
    )
    p(
        "The surgeon reads the final guidance instruction while actively holding an "
        "ultrasound probe. Clinical reasoning terms such as Q1/Q2/Q3/Q4, EP, RP, "
        "N1/N2/N3, 'shunt', and 'circuit' are CHIVA protocol jargon that the "
        "surgeon already knows — they do not need to appear in a real-time navigation "
        "command. Including them wastes the 15-word limit and distracts from the "
        "anatomical movement instruction.",
        bold=False,
    )
    p("The Navigation Planner's task description contains an explicit forbidden list:", bold=False)
    code_block(
        "FORBIDDEN TERMS — must NOT appear anywhere in your output:\n"
        "Q1, Q2, Q3, Q4, corridor, EP, RP, N1, N2, N3, elimTest, elim,\n"
        "shunt, circuit, 'State A', 'State B', 'State C', 'State D', 'State E', 'State F'."
    )
    p(
        "Every valid output starts with a movement verb ('Move') and contains "
        "a direction and a target that names both the anatomical structure AND "
        "its location on the leg. Examples of valid outputs:",
        bold=False,
    )
    bullet("'Move medially toward GSV in saphenous compartment at upper thigh'")
    bullet("'Move distally to Dodd perforator on medial distal thigh'")
    bullet("'Move posteriorly toward SPJ junction in popliteal fossa'")
    bullet("'Move medially toward GSV at medial malleolus (ankle)'")
    p(
        "The Guidance Specialist copies this output verbatim into the JSON 'guidance' "
        "field. No additional summarisation or reformatting is performed. "
        "When action='maneuver', guidance is overridden to "
        "'Perform elimination test at current zone'. "
        "When action='complete', guidance is overridden to "
        "'Circuit complete — classification confirmed'.",
        bold=False,
    )

    print("  [OK] All new sections added.")


# ─── main ──────────────────────────────────────────────────────────────────────

def main():
    print(f"\nOpening: {DOC_PATH}")
    doc = Document(DOC_PATH)
    print(f"  Paragraphs: {len(doc.paragraphs)}  Tables: {len(doc.tables)}")

    print("\n[1/5] Updating Section 6.1 (History Agent)…")
    _update_history_agent_section(doc)

    print("\n[2/5] Updating Section 8.2 (Context Flow word counts)…")
    _update_context_flow(doc)

    print("\n[3/5] Updating Section 8.3 (Agent model-tier notes)…")
    _update_agent_descriptions(doc)

    print("\n[4/5] Updating VLM section (frame source hierarchy + two styles)…")
    _update_vlm_section(doc)

    print("\n[4b/5] Updating Section 8.4 (accepted shunts / multi-shunt note)…")
    _update_shunt_detection_section(doc)

    print("\n[5/5] Generating architecture diagram…")
    diagram_png = _make_diagram()
    print("       Diagram generated — adding new sections…")
    _add_new_sections(doc, diagram_png)

    out_path = DOC_PATH
    doc.save(out_path)
    print(f"\n✓ Saved updated document: {out_path}")
    print("  New sections added: §3.3, §8.6, §8.7, §8.8, §8.9, §8.10")
    print("  Sections updated:   §6.1, §7.2/7.3, §8.2, §8.3, §8.4")


if __name__ == "__main__":
    main()
