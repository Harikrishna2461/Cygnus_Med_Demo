"""
test_set_2 — Extended scenarios (20+ checked steps each) for CHIVA streaming guidance.

Scenarios
---------
type1_extended  Full leg sweep before confirming RP N2→N1.  Tests navigation stamina,
                popliteal/calf coverage gate, and complete firing after thorough scan.
type3_full      Thorough Type 3 path: SFJ entry, trunk reflux, Hunterian escape, calf
                re-entry, maneuver, elimination test (No Reflux), complete.
type5_full      Type 5 circuit (4-clip path: EP N1→N3 → RP N3→N2 → EP N2→N3 → RP N3→N1).
                Tests circuits without SFJ/SPJ entry and without elimination test.
type4_extended  Extended Type 4: bilateral SFJ/SPJ check, calf perforator search, trunk
                sweep for RP N2→N1, complete fires at clip_mark.

Usage
-----
    python tests/run_stream_scenario_2.py [--api http://localhost:7861] [--scenario <id>] [--all]
"""
from __future__ import annotations

import argparse
import io
import json
import os
import sys
import time

if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

try:
    import socketio
except ImportError:
    print("ERROR: python-socketio not installed.  Run: pip install python-socketio[client]")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Step definition (same as test_set_1)
# ---------------------------------------------------------------------------

@dataclass
class Step:
    label: str
    event: str
    data: dict
    expected_action: Optional[str] = None
    guidance_must_contain: list[str] = field(default_factory=list)
    forbidden_action: Optional[str] = None


def _pm(sid, region, pos_y, surface, leg="right"):
    return {"session_id": sid, "region": region, "pos_y_ratio": pos_y,
            "surface": surface, "leg": leg}

def _cm(sid, flow, ft, tt, pos_y, region, surface, elim="", leg="right"):
    return {"session_id": sid, "flow": flow, "from_type": ft, "to_type": tt,
            "pos_y_ratio": pos_y, "leg": leg, "region": region,
            "surface": surface, "elimination_test": elim}


# ---------------------------------------------------------------------------
# Scenario 1 — type1_extended  (25 checked steps)
# ---------------------------------------------------------------------------
# Clinical story:
#   Surgeon confirms SFJ entry early, then does a full sweep of the entire leg
#   (Hunterian zone × 3, popliteal/SPJ, calf/SSV, ankle) looking for a
#   tributary escape.  None found.  Confirms RP N2→N1 late — max_visited is
#   already at ankle, so complete fires immediately at the clip_mark.

def _type1_extended_scenario(sid="ts2_type1_ext") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # ── SFJ assessment ────────────────────────────────────────────────
        Step("P01 SFJ approach — no clips",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 SFJ scan deeper",
             "probe_move", _pm(sid, "SFJ", 0.09, "anterior-medial"),
             expected_action="move"),

        Step("CM-1  clip_mark EP N1→N2 SFJ",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.07, "SFJ", "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["distal", "thigh", "GSV", "trunk", "medial"]),

        # ── Upper thigh sweep ─────────────────────────────────────────────
        Step("P03 upper thigh — tracing GSV",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.13, "medial"),
             expected_action="move"),

        Step("P04 upper thigh distal",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.19, "medial"),
             expected_action="move"),

        # ── Hunterian zone sweep (no RP yet — no hint) ───────────────────
        Step("P05 Hunterian proximal — no escape",
             "probe_move", _pm(sid, "HUNTERIAN", 0.24, "medial"),
             expected_action="move"),

        Step("P06 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.29, "medial"),
             expected_action="move"),

        Step("P07 Hunterian distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.34, "medial"),
             expected_action="move"),

        # ── Distal thigh / popliteal ──────────────────────────────────────
        Step("P08 distal thigh",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.40, "medial"),
             expected_action="move"),

        Step("P09 popliteal — SPJ assessment",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior", "junction"]),

        Step("P10 popliteal deeper — SSV assessment",
             "probe_move", _pm(sid, "POPLITEAL", 0.52, "posterior"),
             expected_action="move"),

        # ── Calf / ankle sweep ────────────────────────────────────────────
        Step("P11 calf posterior upper — SSV trace",
             "probe_move", _pm(sid, "CALF", 0.61, "posterior"),
             expected_action="move",
             guidance_must_contain=["calf", "posterior", "distal"]),

        Step("P12 calf posterior mid",
             "probe_move", _pm(sid, "CALF", 0.67, "posterior"),
             expected_action="move"),

        Step("P13 calf medial — GSV calf",
             "probe_move", _pm(sid, "CALF", 0.73, "medial"),
             expected_action="move"),

        Step("P14 ankle — distal endpoint",
             "probe_move", _pm(sid, "ANKLE", 0.87, "medial"),
             expected_action="move"),

        # ── Sweep back proximally ─────────────────────────────────────────
        Step("P15 calf medial return",
             "probe_move", _pm(sid, "CALF", 0.78, "medial"),
             expected_action="move"),

        Step("P16 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="move"),

        Step("P17 popliteal re-check",
             "probe_move", _pm(sid, "POPLITEAL", 0.51, "posterior"),
             expected_action="move"),

        Step("P18 distal thigh return",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.41, "medial"),
             expected_action="move"),

        # ── Hunterian re-sweep (still no RP — no hint) ───────────────────
        Step("P19 Hunterian re-sweep distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.35, "medial"),
             expected_action="move"),

        Step("P20 Hunterian re-sweep mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="move"),

        Step("P21 upper thigh — final check",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.20, "medial"),
             expected_action="move"),

        # ── Confirm RP N2→N1 — max_visited = 0.87 ≥ 0.44 → complete ─────
        Step("CM-2  clip_mark RP N2→N1 — COMPLETE fires (max_visited=0.87)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.20, "UPPER_THIGH", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P22 verify complete — Hunterian",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P23 verify complete — popliteal",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 2 — type3_full  (23 checked steps)
# ---------------------------------------------------------------------------
# Clinical story:
#   SFJ entry confirmed, trunk reflux confirmed at upper thigh.  Hunterian zone
#   swept with hint active.  SPJ checked (competent).  Escape found at Hunterian.
#   Tributary traced down to calf; re-entry (RP N3→N1) confirmed → maneuver.
#   Multiple probe moves under maneuver.  Surgeon records No Reflux → complete.

def _type3_full_scenario(sid="ts2_type3_full") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # ── SFJ ───────────────────────────────────────────────────────────
        Step("P01 SFJ approach",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral"]),

        Step("CM-1  clip_mark EP N1→N2 SFJ",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.06, "SFJ", "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["distal", "thigh", "GSV", "trunk"]),

        # ── Upper thigh — confirm trunk reflux ────────────────────────────
        Step("P02 upper thigh — tracing trunk",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.15, "medial"),
             expected_action="move"),

        Step("CM-2  clip_mark RP N2→N1 upper thigh",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.15, "UPPER_THIGH", "medial"),
             expected_action="move",
             guidance_must_contain=["Hunterian", "mid-thigh", "tributary", "escape"]),

        # ── Hunterian sweep with hint active ─────────────────────────────
        Step("P03 Hunterian proximal — hint active",
             "probe_move", _pm(sid, "HUNTERIAN", 0.22, "medial"),
             expected_action="move",
             guidance_must_contain=["Hunterian", "tributary", "escape", "mid-thigh"]),

        Step("P04 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.27, "medial"),
             expected_action="move",
             guidance_must_contain=["Hunterian", "tributary", "escape", "perforator"]),

        Step("P05 Hunterian distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.32, "medial"),
             expected_action="move",
             guidance_must_contain=["Hunterian", "tributary", "escape"]),

        # ── SPJ check (popliteal, out of Hunterian range) ─────────────────
        Step("P06 popliteal — SPJ competent check",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P07 popliteal deeper — SSV check",
             "probe_move", _pm(sid, "POPLITEAL", 0.51, "posterior"),
             expected_action="move"),

        # ── Return to Hunterian — find escape ────────────────────────────
        Step("P08 Hunterian return — hint fires again",
             "probe_move", _pm(sid, "HUNTERIAN", 0.29, "medial"),
             expected_action="move",
             guidance_must_contain=["Hunterian", "tributary", "escape"]),

        Step("CM-3  clip_mark EP N2→N3 Hunterian — escape found (no elimTest)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.29, "HUNTERIAN", "medial"),
             expected_action="move",
             guidance_must_contain=["tributary", "distal", "follow"],
             forbidden_action="maneuver"),

        # ── Trace tributary distally ──────────────────────────────────────
        Step("P09 distal thigh — following tributary",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.38, "medial"),
             expected_action="move"),

        Step("P10 popliteal medial — tributary near knee",
             "probe_move", _pm(sid, "POPLITEAL", 0.44, "medial"),
             expected_action="move"),

        Step("P11 calf medial upper — in calf",
             "probe_move", _pm(sid, "CALF", 0.58, "medial"),
             expected_action="move"),

        Step("P12 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.64, "medial"),
             expected_action="move"),

        Step("P13 calf medial lower — approaching re-entry",
             "probe_move", _pm(sid, "CALF", 0.70, "medial"),
             expected_action="move"),

        # ── Confirm re-entry → maneuver ───────────────────────────────────
        Step("CM-4  clip_mark RP N3→N1 calf — MANEUVER fires",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.70, "CALF", "medial"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary", "Doppler", "record"]),

        Step("P14 popliteal — still under maneuver",
             "probe_move", _pm(sid, "POPLITEAL", 0.46, "posterior"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary", "record"]),

        Step("P15 distal thigh — approach Hunterian",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.38, "medial"),
             expected_action="maneuver"),

        Step("P16 Hunterian — maneuver zone",
             "probe_move", _pm(sid, "HUNTERIAN", 0.29, "medial"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary", "Doppler"]),

        # ── Elimination test → No Reflux → complete ───────────────────────
        Step("CM-5  clip_mark EP N2→N3 elimTest=No Reflux — COMPLETE",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.29, "HUNTERIAN", "medial",
                               elim="No Reflux"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P17 verify complete — Hunterian",
             "probe_move", _pm(sid, "HUNTERIAN", 0.29, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P18 verify complete — upper thigh",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.17, "medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 3 — type5_full  (21 checked steps)
# ---------------------------------------------------------------------------
# Clinical story:
#   SFJ competent, SPJ competent.  EP N1→N3 found at mid-calf.  Tributary
#   traced back up to trunk → RP N3→N2 at distal thigh.  Trunk followed
#   distally; EP N2→N3 found in calf.  N3 traced; RP N3→N1 at distal calf →
#   complete (Type 5 rule: ep_n1_n3 + rp_n3_n2 + ep_n2_n3 + rp_n3_n1).
#   NOTE: no elimination test needed; no RP N2→N1 so maneuver rule never fires.

def _type5_full_scenario(sid="ts2_type5") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # ── SFJ — competent ───────────────────────────────────────────────
        Step("P01 SFJ — competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral"]),

        # ── SPJ — competent ───────────────────────────────────────────────
        Step("P02 popliteal — SPJ competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        # ── Scan calf for perforators ─────────────────────────────────────
        Step("P03 calf medial — scanning for perforators",
             "probe_move", _pm(sid, "CALF", 0.58, "medial"),
             expected_action="move"),

        Step("P04 calf medial upper",
             "probe_move", _pm(sid, "CALF", 0.63, "medial"),
             expected_action="move"),

        Step("P05 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.68, "medial"),
             expected_action="move"),

        # ── Find EP N1→N3 ─────────────────────────────────────────────────
        Step("CM-1  clip_mark EP N1→N3 mid-calf perforator",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.68, "CALF", "medial"),
             expected_action="move"),

        # ── Trace N3 tributary toward trunk ───────────────────────────────
        Step("P06 calf medial — tracing N3 tributary",
             "probe_move", _pm(sid, "CALF", 0.63, "medial"),
             expected_action="move"),

        Step("P07 popliteal medial — N3 approaching trunk",
             "probe_move", _pm(sid, "POPLITEAL", 0.50, "medial"),
             expected_action="move"),

        Step("P08 distal thigh — N3 near trunk",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.41, "medial"),
             expected_action="move"),

        # ── Find RP N3→N2 (N3 enters trunk) ──────────────────────────────
        Step("CM-2  clip_mark RP N3→N2 distal thigh — N3 joins trunk",
             "clip_mark", _cm(sid, "RP", "N3", "N2", 0.41, "DISTAL_THIGH", "medial"),
             expected_action="move"),

        # ── Follow trunk distally for escape ──────────────────────────────
        Step("P09 popliteal — following trunk distally",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move"),

        Step("P10 calf posterior upper — trunk in calf",
             "probe_move", _pm(sid, "CALF", 0.57, "posterior"),
             expected_action="move"),

        Step("P11 calf posterior mid",
             "probe_move", _pm(sid, "CALF", 0.63, "posterior"),
             expected_action="move"),

        # ── Find EP N2→N3 (trunk escapes to N3) ──────────────────────────
        Step("CM-3  clip_mark EP N2→N3 calf — trunk escapes to tributary",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.64, "CALF", "posterior"),
             expected_action="move",
             forbidden_action="maneuver"),   # no rp_n2_n1 → maneuver rule cannot fire

        # ── Follow N3 toward re-entry ─────────────────────────────────────
        Step("P12 calf posterior lower — following N3",
             "probe_move", _pm(sid, "CALF", 0.70, "posterior"),
             expected_action="move"),

        Step("P13 calf lateral — N3 crosses toward lateral perforator",
             "probe_move", _pm(sid, "CALF", 0.74, "lateral"),
             expected_action="move"),

        Step("P14 calf lateral distal",
             "probe_move", _pm(sid, "CALF", 0.78, "lateral"),
             expected_action="move"),

        Step("P15 calf lateral — approaching re-entry perforator",
             "probe_move", _pm(sid, "CALF", 0.81, "lateral"),
             expected_action="move"),

        # ── Find RP N3→N1 → COMPLETE (Type 5 rule) ───────────────────────
        Step("CM-4  clip_mark RP N3→N1 distal calf — circuit closed → COMPLETE",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.81, "CALF", "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P16 verify complete — calf lateral",
             "probe_move", _pm(sid, "CALF", 0.81, "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P17 verify complete — Hunterian",
             "probe_move", _pm(sid, "HUNTERIAN", 0.30, "medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 4 — type4_extended  (22 checked steps)
# ---------------------------------------------------------------------------
# Clinical story:
#   SFJ competent (double-checked), SPJ competent.  Full calf sweep for
#   perforators; EP N1→N3 found at mid-calf.  Trunk swept from calf to upper
#   thigh for RP N2→N1.  Complete fires at clip_mark RP N2→N1 (Type 4 rule).

def _type4_extended_scenario(sid="ts2_type4_ext") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # ── SFJ — double check ────────────────────────────────────────────
        Step("P01 SFJ approach — competent",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral"]),

        Step("P02 SFJ deeper — confirm competence",
             "probe_move", _pm(sid, "SFJ", 0.09, "anterior-medial"),
             expected_action="move"),

        # ── SPJ — check ───────────────────────────────────────────────────
        Step("P03 popliteal — SPJ assessment",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P04 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.52, "posterior"),
             expected_action="move"),

        # ── Calf sweep for perforators ────────────────────────────────────
        Step("P05 calf medial upper — scanning for perforators",
             "probe_move", _pm(sid, "CALF", 0.58, "medial"),
             expected_action="move"),

        Step("P06 calf medial",
             "probe_move", _pm(sid, "CALF", 0.63, "medial"),
             expected_action="move"),

        Step("P07 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.68, "medial"),
             expected_action="move"),

        Step("P08 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.74, "medial"),
             expected_action="move"),

        Step("P09 ankle approach",
             "probe_move", _pm(sid, "ANKLE", 0.87, "medial"),
             expected_action="move"),

        # ── Find EP N1→N3 ─────────────────────────────────────────────────
        Step("CM-1  clip_mark EP N1→N3 mid-calf perforator",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.68, "CALF", "medial"),
             expected_action="move"),

        # ── Sweep trunk from popliteal to upper thigh for RP N2→N1 ───────
        Step("P10 popliteal — checking trunk",
             "probe_move", _pm(sid, "POPLITEAL", 0.48, "posterior"),
             expected_action="move"),

        Step("P11 distal thigh — trunk sweep",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.40, "medial"),
             expected_action="move"),

        Step("P12 distal thigh proximal",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.37, "medial"),
             expected_action="move"),

        Step("P13 Hunterian distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.33, "medial"),
             expected_action="move"),

        Step("P14 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.27, "medial"),
             expected_action="move"),

        Step("P15 Hunterian proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.22, "medial"),
             expected_action="move"),

        Step("P16 upper thigh",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.17, "medial"),
             expected_action="move"),

        Step("P17 upper thigh proximal",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.13, "medial"),
             expected_action="move"),

        # ── Confirm RP N2→N1 → COMPLETE (Type 4 rule fires immediately) ──
        Step("CM-2  clip_mark RP N2→N1 Hunterian — COMPLETE fires",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.30, "HUNTERIAN", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P18 verify complete — upper thigh",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.17, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P19 verify complete — calf",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="complete"),

        Step("P20 verify complete — popliteal",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

_SCENARIOS: dict[str, callable] = {
    "type1_extended": _type1_extended_scenario,
    "type3_full":     _type3_full_scenario,
    "type5_full":     _type5_full_scenario,
    "type4_extended": _type4_extended_scenario,
}

_SCENARIO_DESCRIPTIONS = {
    "type1_extended": (
        "Type 1 extended — full leg sweep before confirming RP N2→N1; "
        "complete fires at clip_mark when max_visited=0.87."
    ),
    "type3_full": (
        "Type 3 full — SFJ entry, trunk reflux, Hunterian escape, calf re-entry, "
        "maneuver, elimination test (No Reflux), complete."
    ),
    "type5_full": (
        "Type 5 full — EP N1→N3 → RP N3→N2 → EP N2→N3 → RP N3→N1; "
        "no SFJ/SPJ entry, no elimination test, complete at 4th clip."
    ),
    "type4_extended": (
        "Type 4 extended — bilateral SFJ/SPJ check, full calf perforator sweep, "
        "trunk sweep for RP N2→N1; complete at clip_mark (Type 4 rule)."
    ),
}


# ---------------------------------------------------------------------------
# Human-readable helpers
# ---------------------------------------------------------------------------

def _describe_movement(event: str, data: dict) -> str:
    if event == "stream_start":
        return f"Start session  (id: {data.get('session_id', '?')})"
    if event == "probe_move":
        return (f"Move probe to {data.get('region', '?')}\n"
                f"posY = {data.get('pos_y_ratio', 0.0):.2f}  |  "
                f"{data.get('surface', '?')} surface  |  {data.get('leg', 'right')} leg")
    if event == "clip_mark":
        base = (f"Mark clip: {data.get('flow', '?')} "
                f"{data.get('from_type', '?')}→{data.get('to_type', '?')}\n"
                f"posY = {data.get('pos_y_ratio', 0.0):.2f}  |  "
                f"{data.get('region', '?')}  |  {data.get('leg', 'right')} leg")
        if data.get("elimination_test"):
            base += f"\nElimination test: {data['elimination_test']}"
        return base
    return event


def _describe_expectation(step: Step) -> str:
    if step.event == "stream_start":
        return "session_ready event"
    parts = []
    if step.expected_action:
        parts.append(f'Action = "{step.expected_action}"')
    if step.forbidden_action:
        parts.append(f'Must NOT fire action = "{step.forbidden_action}"')
    if step.guidance_must_contain:
        parts.append("Guidance must mention any of:\n" +
                     ",  ".join(step.guidance_must_contain))
    return "\n".join(parts) if parts else "No assertion (observe only)"


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

CYAN   = "\033[96m"
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

_SKIP_EVENTS = {"stream_start"}


class ScenarioRunner:
    def __init__(self, api_base: str):
        self.api_base = api_base.rstrip("/")
        self._sio = socketio.SimpleClient()

    def connect(self) -> None:
        print(f"Connecting to {self.api_base} ...")
        self._sio.connect(self.api_base, wait_timeout=10)
        print("Connected.\n")

    def disconnect(self) -> None:
        self._sio.disconnect()

    def _emit_and_wait(self, event: str, data: dict, timeout: float = 60.0):
        self._sio.emit(event, data)
        if event in _SKIP_EVENTS:
            deadline = time.time() + timeout
            while time.time() < deadline:
                try:
                    ev = self._sio.receive(timeout=2)
                    if ev and ev[0] == "session_ready":
                        return ev[1]
                except Exception:
                    pass
            return None
        deadline = time.time() + timeout
        while time.time() < deadline:
            try:
                ev = self._sio.receive(timeout=2)
                if ev and ev[0] == "guidance_update":
                    return ev[1]
            except Exception:
                pass
        return None

    def run(self, steps: list[Step]) -> list[dict]:
        results = []
        for step in steps:
            print(f"\n{CYAN}--- {step.label} ---{RESET}")
            print(f"  event: {step.event}")

            movement    = _describe_movement(step.event, step.data)
            expectation = _describe_expectation(step)
            response    = self._emit_and_wait(step.event, step.data)

            # stream_start
            if step.event in _SKIP_EVENTS:
                if response:
                    print(f"  session_ready: {response}")
                    results.append({"step": step.label, "status": "SKIP",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                else:
                    print(f"  {YELLOW}No session_ready received{RESET}")
                    results.append({"step": step.label, "status": "TIMEOUT",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                continue

            # observe-only
            if (step.expected_action is None
                    and not step.guidance_must_contain
                    and step.forbidden_action is None):
                if response:
                    g = response.get("guidance", "")
                    a = response.get("action", "?")
                    print(f"  guidance: {g!r}  action: {a}")
                    results.append({"step": step.label, "status": "SKIP",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": a, "llm_guidance": g, "failures": []})
                else:
                    print(f"  {YELLOW}timeout{RESET}")
                    results.append({"step": step.label, "status": "TIMEOUT",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                continue

            # timeout
            if response is None:
                print(f"  {YELLOW}No response received (timeout){RESET}")
                results.append({"step": step.label, "status": "TIMEOUT",
                                "movement": movement, "expectation": expectation,
                                "llm_action": "--", "llm_guidance": "--",
                                "failures": []})
                continue

            llm_guidance = response.get("guidance", "")
            llm_action   = response.get("action", "move")
            print(f"  guidance: {llm_guidance!r}")
            print(f"  action  : {llm_action}")

            passed   = True
            failures = []

            if step.expected_action and llm_action != step.expected_action:
                passed = False
                failures.append(
                    f'Action: got "{llm_action}", expected "{step.expected_action}"')

            if step.guidance_must_contain:
                if not any(kw.lower() in llm_guidance.lower()
                           for kw in step.guidance_must_contain):
                    passed = False
                    failures.append(
                        f"Guidance missing any of: {', '.join(step.guidance_must_contain)}")

            if step.forbidden_action and llm_action == step.forbidden_action:
                passed = False
                failures.append(
                    f'Action "{llm_action}" must NOT fire at this step')

            if passed:
                print(f"  {GREEN}{BOLD}PASS{RESET}")
            else:
                for f in failures:
                    print(f"  {RED}FAIL: {f}{RESET}")

            results.append({"step": step.label,
                            "status": "PASS" if passed else "FAIL",
                            "movement": movement, "expectation": expectation,
                            "llm_action": llm_action, "llm_guidance": llm_guidance,
                            "failures": failures})
            time.sleep(0.4)

        return results

    def print_summary(self, label: str, results: list[dict]) -> bool:
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        passed  = [r for r in checked if r["status"] == "PASS"]
        failed  = [r for r in checked if r["status"] == "FAIL"]
        timeout = [r for r in results if r["status"] == "TIMEOUT"]

        print(f"\n{'─'*60}")
        print(f"{BOLD}{label}{RESET}  "
              f"{GREEN}{len(passed)} passed{RESET}  "
              f"{RED}{len(failed)} failed{RESET}  "
              f"{YELLOW}{len(timeout)} timeout{RESET}  "
              f"/ {len(checked)} checked")

        for r in failed:
            print(f"  {RED}x {r['step']}{RESET}")
            for f in r.get("failures", []):
                print(f"      {f}")

        return len(failed) == 0 and len(timeout) == 0


# ---------------------------------------------------------------------------
# Word report
# ---------------------------------------------------------------------------

_STATUS_COLORS = {
    "PASS":    "00B050",
    "FAIL":    "FF0000",
    "TIMEOUT": "FF8C00",
    "SKIP":    "808080",
}


def _colored_cell(cell, text: str, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    from docx.shared import RGBColor
    cell.text = text
    run = (cell.paragraphs[0].runs[0]
           if cell.paragraphs[0].runs
           else cell.paragraphs[0].add_run(text))
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, pt: int = 9, bold: bool = False,
                   italic: bool = False, color_rgb=None) -> None:
    cell.text = ""
    lines = text.split("\n")
    for i, line in enumerate(lines):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run  = para.add_run(line)
        run.font.size   = __import__("docx").shared.Pt(pt)
        run.font.bold   = bold
        run.font.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb


def write_word_report(all_results: dict[str, list[dict]], api_base: str,
                      out_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    doc = Document()
    section = doc.sections[0]
    section.left_margin   = Inches(0.6)
    section.right_margin  = Inches(0.6)
    section.top_margin    = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    title = doc.add_heading("CHIVA Streaming Guidance — Test Set 2 Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = meta.add_run(
        f"Date: {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}    |    "
        f"API: {api_base}    |    Scenarios: {len(all_results)}")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    doc.add_heading("Overall Summary", level=1)
    stbl = doc.add_table(rows=1, cols=5)
    stbl.style = "Table Grid"
    for i, h in enumerate(["Scenario", "Description", "Passed", "Failed", "Timeouts"]):
        c = stbl.rows[0].cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True

    overall_pass = True
    for name, results in all_results.items():
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")
        row = stbl.add_row().cells
        row[0].text = name
        row[1].text = _SCENARIO_DESCRIPTIONS.get(name, "")
        row[2].text = str(p)
        row[3].text = str(f)
        row[4].text = str(t)
        if f or t:
            overall_pass = False

    doc.add_paragraph()
    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run("OVERALL RESULT: " + ("ALL PASS" if overall_pass else "FAILURES PRESENT"))
    vr.font.bold = True
    vr.font.size = Pt(13)
    vr.font.color.rgb = (RGBColor(0x00, 0xB0, 0x50) if overall_pass
                         else RGBColor(0xFF, 0x00, 0x00))

    for name, results in all_results.items():
        doc.add_page_break()
        doc.add_heading(f"Scenario: {name.upper()}", level=1)
        dp = doc.add_paragraph(_SCENARIO_DESCRIPTIONS.get(name, ""))
        if dp.runs:
            dp.runs[0].font.italic = True

        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")

        sl = doc.add_paragraph()
        sl.add_run(f"Steps checked: {len(checked)}   |   ")
        pr = sl.add_run(f"Passed: {p}")
        pr.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
        sl.add_run("   |   ")
        fr = sl.add_run(f"Failed: {f}")
        fr.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        sl.add_run("   |   ")
        tr = sl.add_run(f"Timeouts: {t}")
        tr.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        doc.add_paragraph()

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        headers = ["#", "What Surgeon Did", "Expected Guidance",
                   "LLM Guidance", "LLM Action", "Status"]
        col_w   = [0.25, 1.85, 2.10, 2.10, 0.85, 0.65]

        for i, h in enumerate(headers):
            c = tbl.rows[0].cells[i]
            c.text = h
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.bold = True

        for i, w in enumerate(col_w):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)

        for idx, r in enumerate(results, 1):
            status = r.get("status", "SKIP")
            row    = tbl.add_row().cells
            row[0].text = str(idx)
            _set_cell_text(row[1], r.get("movement", "--"), pt=8)
            _set_cell_text(row[2], r.get("expectation", "--"), pt=8, italic=True)
            _set_cell_text(row[3], r.get("llm_guidance", "--") or "--", pt=8)
            _set_cell_text(row[4], r.get("llm_action", "--") or "--", pt=9, bold=True)
            _colored_cell(row[5], status, _STATUS_COLORS.get(status, "808080"))

            if status == "FAIL" and r.get("failures"):
                fr_row = tbl.add_row().cells
                merged = fr_row[0].merge(fr_row[5])
                merged.text = ""
                for fn in r["failures"]:
                    pp = merged.add_paragraph(f"  x  {fn}")
                    if pp.runs:
                        pp.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        pp.runs[0].font.size      = Pt(8)
                        pp.runs[0].font.italic    = True

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"\nReport saved: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="CHIVA stream test set 2 runner")
    parser.add_argument("scenario", nargs="?", default="all",
                        choices=list(_SCENARIOS) + ["all"],
                        help="Scenario to run (default: all)")
    parser.add_argument("--api",  default="http://localhost:7861")
    parser.add_argument("--all",  action="store_true", dest="run_all")
    args = parser.parse_args()

    to_run = list(_SCENARIOS.items()) if (args.run_all or args.scenario == "all") \
             else [(args.scenario, _SCENARIOS[args.scenario])]

    runner     = ScenarioRunner(args.api)
    all_ok     = True
    all_results: dict[str, list[dict]] = {}

    try:
        runner.connect()
        for name, builder in to_run:
            print(f"\n{'='*60}")
            print(f"{BOLD}SCENARIO: {name.upper()}{RESET}")
            print(f"{'='*60}")
            steps   = builder()
            results = runner.run(steps)
            ok      = runner.print_summary(name, results)
            all_ok  = all_ok and ok
            all_results[name] = results
    except KeyboardInterrupt:
        print("\nInterrupted.")
    finally:
        runner.disconnect()

    if len(to_run) > 1:
        print(f"\n{'='*60}")
        print(f"{BOLD}OVERALL: {'ALL PASS' if all_ok else 'FAILURES PRESENT'}{RESET}")

    if all_results:
        ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
        here     = os.path.dirname(os.path.abspath(__file__))
        out_path = os.path.join(here, "results", f"test_set2_log_{ts}.docx")
        try:
            write_word_report(all_results, args.api, out_path)
        except Exception as exc:
            print(f"{YELLOW}Warning: could not write Word report — {exc}{RESET}")

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
