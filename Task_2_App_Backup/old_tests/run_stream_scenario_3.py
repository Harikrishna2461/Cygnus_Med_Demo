"""
test_set_3 — Nine new CHIVA streaming scenarios, each 20+ checked steps.

Scenarios
---------
hunterian_entry  Hunterian perforator as primary entry point (SFJ competent).
                 Type 1 via mid-thigh perforator: EP N1→N2 at Hunterian, full
                 distal sweep, RP N2→N1 confirmed at calf → complete.

ssv_extended     Posterior SSV circuit (SPJ primary). SFJ competent. SPJ entry
                 found, SSV traced extensively. RP N2→N1 confirmed after full
                 calf sweep → complete fires (no escape found, max_visited ≥ 0.44).

left_type4       Left leg Type 4. SFJ and SPJ both competent on left. EP N1→N3
                 found at left calf perforator. GSV trunk swept; RP N2→N1
                 confirmed at Hunterian → complete (Type 4 rule fires immediately).

ankle_type6      Pure perforator circuit at ankle level (Type 6). Both junctions
                 competent. EP N1→N3 at ankle/lower-calf lateral perforator.
                 RP N3→N1 confirmed proximally on calf → complete (Type 6 rule).

type3_no_reflux  Type 3 escape circuit (No Reflux elimination). SFJ entry
                 EP N1→N2. Escape EP N2→N3 found at distal thigh before any RP
                 clip (blocks Rule 6). Full sweep to ankle. RP N3→N1 at calf,
                 then RP N2→N1 at Hunterian → maneuver fires. Elimination test
                 "No Reflux" → Type 3 complete.

type1p2_reflux   Type 1+2 escape circuit (Reflux elimination). SPJ/SSV entry
                 EP N1→N2. Escape EP N2→N3 found along SSV before RP N2→N1.
                 Full ankle sweep. RP N3→N1 at ankle, RP N2→N1 at calf →
                 maneuver. Elimination test "Reflux" → Type 1+2 complete.

type5_complex    Type 5 four-clip perforator circuit. Both junctions competent.
                 EP N1→N3 at Hunterian perforator. RP N3→N2 at calf junction.
                 EP N2→N3 at distal calf. RP N3→N1 at ankle → Type 5 complete
                 (Rule 5b: ep_n1_n3 + rp_n3_n2 + ep_n2_n3 + rp_n3_n1).

sfj_entry_long   Type 1/2A standard SFJ entry with extended thigh sweep (20+
                 steps). EP N1→N2 at SFJ early, thorough multi-zone sweep to
                 ankle, then RP N2→N1 confirmed at distal calf → complete once
                 max_visited ≥ 0.44.

calf_medial_t4   Type 4 with distal perforator entry. EP N1→N3 found at ankle
                 (posY 0.88). Long proximal sweep through calf, popliteal,
                 thigh. RP N2→N1 at Hunterian → Type 4 complete fires
                 immediately (Rule 4: ep_n1_n3 + rp_n2_n1, no max_visited gate).

Usage
-----
    python tests/run_stream_scenario_3.py [--api http://localhost:7861] [--scenario <id>] [--all]
"""
from __future__ import annotations

import argparse
import io
import os
import sys
import time

if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

from dataclasses import dataclass, field
from datetime import datetime
from typing import Optional

try:
    import socketio
except ImportError:
    print("ERROR: python-socketio not installed.  Run: pip install python-socketio[client]")
    sys.exit(1)


# ---------------------------------------------------------------------------
# Step definition
# ---------------------------------------------------------------------------

@dataclass
class Step:
    label: str
    event: str
    data: dict
    expected_action: Optional[str] = None
    guidance_must_contain: list[str] = field(default_factory=list)
    forbidden_action: Optional[str] = None


def _pm(sid, region, pos_y, surface, leg="right"):
    return {"session_id": sid, "region": region, "pos_y_ratio": pos_y,
            "surface": surface, "leg": leg}

def _cm(sid, flow, ft, tt, pos_y, region, surface, elim="", leg="right"):
    return {"session_id": sid, "flow": flow, "from_type": ft, "to_type": tt,
            "pos_y_ratio": pos_y, "leg": leg, "region": region,
            "surface": surface, "elimination_test": elim}


# ---------------------------------------------------------------------------
# Scenario 1 — hunterian_entry  (23 checked steps)
# ---------------------------------------------------------------------------
# SFJ is competent (no clip taken). EP N1→N2 found at Hunterian perforator
# (posY 0.26, medial). Surgeon then sweeps the full leg distally — thigh,
# popliteal, calf, ankle — before confirming RP N2→N1 in calf. Because
# max_visited reaches ankle (0.87), Type 1 complete fires immediately at the
# clip_mark for RP N2→N1.

def _hunterian_entry_scenario(sid="ts3_hunt_entry") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # SFJ probed — competent, no clip taken
        Step("P01 SFJ approach — competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 SFJ deeper scan",
             "probe_move", _pm(sid, "SFJ", 0.09, "anterior-medial"),
             expected_action="move"),

        # Upper thigh — no findings
        Step("P03 upper thigh — no entry found",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.14, "medial"),
             expected_action="move"),

        Step("P04 upper thigh distal",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.20, "medial"),
             expected_action="move"),

        # Hunterian zone — find entry
        Step("P05 Hunterian — scanning for entry",
             "probe_move", _pm(sid, "HUNTERIAN", 0.24, "medial"),
             expected_action="move"),

        Step("CM-1  clip_mark EP N1→N2 Hunterian perforator",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.26, "HUNTERIAN", "medial"),
             expected_action="move"),

        # Trace GSV distally from Hunterian
        Step("P06 Hunterian distal — tracing below entry",
             "probe_move", _pm(sid, "HUNTERIAN", 0.31, "medial"),
             expected_action="move"),

        Step("P07 distal thigh — following GSV",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.38, "medial"),
             expected_action="move"),

        # Popliteal — SPJ competent check
        Step("P08 popliteal — SPJ competent check",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P09 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.53, "posterior"),
             expected_action="move"),

        # Calf sweep
        Step("P10 calf posterior — SSV check",
             "probe_move", _pm(sid, "CALF", 0.60, "posterior"),
             expected_action="move"),

        Step("P11 calf medial — GSV calf",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="move"),

        Step("P12 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.72, "medial"),
             expected_action="move"),

        Step("P13 ankle",
             "probe_move", _pm(sid, "ANKLE", 0.87, "medial"),
             expected_action="move"),

        # Sweep back up
        Step("P14 calf medial return",
             "probe_move", _pm(sid, "CALF", 0.78, "medial"),
             expected_action="move"),

        Step("P15 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.68, "medial"),
             expected_action="move"),

        Step("P16 popliteal return",
             "probe_move", _pm(sid, "POPLITEAL", 0.51, "posterior"),
             expected_action="move"),

        Step("P17 distal thigh return",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.40, "medial"),
             expected_action="move"),

        Step("P18 Hunterian sweep",
             "probe_move", _pm(sid, "HUNTERIAN", 0.32, "medial"),
             expected_action="move"),

        Step("P19 mid-thigh",
             "probe_move", _pm(sid, "HUNTERIAN", 0.26, "medial"),
             expected_action="move"),

        # Confirm RP N2→N1 — max_visited=0.87 ≥ 0.44 → complete fires
        Step("CM-2  clip_mark RP N2→N1 calf — COMPLETE (max_visited=0.87)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.71, "CALF", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P20 verify complete — Hunterian",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P21 verify complete — calf",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 2 — ssv_extended  (21 checked steps)
# ---------------------------------------------------------------------------
# Posterior SSV primary circuit. SFJ probed and confirmed competent. Surgeon
# moves to popliteal, finds SPJ entry (EP N1→N2), traces SSV through calf to
# ankle and back. After thorough SSV sweep, confirms RP N2→N1. Because
# max_visited reached ankle during the sweep, Type 1/2A complete fires at the
# RP N2→N1 clip_mark (no escape was found).

def _ssv_extended_scenario(sid="ts3_ssv_ext") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # SFJ — competent
        Step("P01 SFJ — competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        # Popliteal — find SPJ
        Step("P02 popliteal approach — SPJ assessment",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P03 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.51, "posterior"),
             expected_action="move"),

        Step("CM-1  clip_mark EP N1→N2 at SPJ",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.48, "POPLITEAL", "posterior"),
             expected_action="move"),

        # SSV trace distally
        Step("P04 calf posterior — tracing SSV",
             "probe_move", _pm(sid, "CALF", 0.57, "posterior"),
             expected_action="move"),

        Step("P05 calf posterior mid",
             "probe_move", _pm(sid, "CALF", 0.63, "posterior"),
             expected_action="move"),

        Step("P06 calf posterior lower",
             "probe_move", _pm(sid, "CALF", 0.68, "posterior"),
             expected_action="move"),

        Step("P07 calf posterior distal",
             "probe_move", _pm(sid, "CALF", 0.74, "posterior"),
             expected_action="move"),

        Step("P08 ankle posterior",
             "probe_move", _pm(sid, "ANKLE", 0.88, "posterior"),
             expected_action="move"),

        # Sweep back up along SSV
        Step("P09 calf posterior return lower",
             "probe_move", _pm(sid, "CALF", 0.81, "posterior"),
             expected_action="move"),

        Step("P10 calf posterior return mid",
             "probe_move", _pm(sid, "CALF", 0.73, "posterior"),
             expected_action="move"),

        Step("P11 calf posterior return upper",
             "probe_move", _pm(sid, "CALF", 0.65, "posterior"),
             expected_action="move"),

        Step("P12 calf posterior proximal",
             "probe_move", _pm(sid, "CALF", 0.58, "posterior"),
             expected_action="move"),

        Step("P13 popliteal re-check",
             "probe_move", _pm(sid, "POPLITEAL", 0.53, "posterior"),
             expected_action="move"),

        Step("P14 popliteal proximal",
             "probe_move", _pm(sid, "POPLITEAL", 0.48, "posterior"),
             expected_action="move"),

        # Second pass — confirm reflux
        Step("P15 calf posterior second pass",
             "probe_move", _pm(sid, "CALF", 0.57, "posterior"),
             expected_action="move"),

        Step("P16 calf posterior mid second pass",
             "probe_move", _pm(sid, "CALF", 0.63, "posterior"),
             expected_action="move"),

        # Confirm RP N2→N1 — max_visited=0.88 → complete fires
        Step("CM-2  clip_mark RP N2→N1 SSV calf — COMPLETE (max_visited=0.88)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.65, "CALF", "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P17 verify complete — calf",
             "probe_move", _pm(sid, "CALF", 0.62, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P18 verify complete — popliteal",
             "probe_move", _pm(sid, "POPLITEAL", 0.49, "posterior"),
             expected_action="complete"),

        Step("P19 verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 3 — left_type4  (21 checked steps)
# ---------------------------------------------------------------------------
# Left leg, Type 4 circuit. SFJ left competent. SPJ left competent. Calf
# perforator search on left medial surface → EP N1→N3 found at mid-calf.
# GSV trunk swept from popliteal to upper thigh on left. RP N2→N1 confirmed
# at left Hunterian → Type 4 complete fires immediately (ep_n1_n3 + rp_n2_n1).

def _left_type4_scenario(sid="ts3_left_t4") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # Left SFJ — competent
        Step("P01 SFJ left — competent",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial", leg="left"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        # Left popliteal — SPJ competent
        Step("P02 popliteal left — SPJ competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior", leg="left"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P03 popliteal left deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.52, "posterior", leg="left"),
             expected_action="move"),

        # Left calf — searching for perforators
        Step("P04 calf medial left — searching",
             "probe_move", _pm(sid, "CALF", 0.58, "medial", leg="left"),
             expected_action="move"),

        Step("P05 calf medial left mid",
             "probe_move", _pm(sid, "CALF", 0.63, "medial", leg="left"),
             expected_action="move"),

        Step("P06 calf medial left lower",
             "probe_move", _pm(sid, "CALF", 0.69, "medial", leg="left"),
             expected_action="move"),

        Step("P07 calf medial left distal",
             "probe_move", _pm(sid, "CALF", 0.76, "medial", leg="left"),
             expected_action="move"),

        # Find EP N1→N3 on left calf
        Step("CM-1  clip_mark EP N1→N3 left calf perforator",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.69, "CALF", "medial",
                               leg="left"),
             expected_action="move"),

        # Sweep GSV trunk for RP N2→N1
        Step("P08 popliteal left — checking trunk",
             "probe_move", _pm(sid, "POPLITEAL", 0.50, "posterior", leg="left"),
             expected_action="move"),

        Step("P09 distal thigh left",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.41, "medial", leg="left"),
             expected_action="move"),

        Step("P10 Hunterian left distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.35, "medial", leg="left"),
             expected_action="move"),

        Step("P11 Hunterian left mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.29, "medial", leg="left"),
             expected_action="move"),

        Step("P12 Hunterian left proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.23, "medial", leg="left"),
             expected_action="move"),

        Step("P13 upper thigh left",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.17, "medial", leg="left"),
             expected_action="move"),

        Step("P14 upper thigh left proximal",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.13, "medial", leg="left"),
             expected_action="move"),

        # Return sweep
        Step("P15 Hunterian left — return sweep",
             "probe_move", _pm(sid, "HUNTERIAN", 0.22, "medial", leg="left"),
             expected_action="move"),

        Step("P16 Hunterian left mid return",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial", leg="left"),
             expected_action="move"),

        # Confirm RP N2→N1 → Type 4 complete fires immediately
        Step("CM-2  clip_mark RP N2→N1 left Hunterian — COMPLETE (Type 4)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.30, "HUNTERIAN", "medial",
                               leg="left"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P17 verify complete — upper thigh left",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.17, "medial", leg="left"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P18 verify complete — calf left",
             "probe_move", _pm(sid, "CALF", 0.65, "medial", leg="left"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 4 — ankle_type6  (22 checked steps)
# ---------------------------------------------------------------------------
# Pure perforator circuit at ankle/distal-calf level (Type 6). Both SFJ and
# SPJ confirmed competent. Surgeon searches systematically through calf and
# ankle surfaces. EP N1→N3 found at ankle lateral perforator (posY 0.89).
# Tributary traced back up. RP N3→N1 found on lateral calf (posY 0.76) →
# Type 6 complete (ep_n1_n3 + rp_n3_n1, no rp_n2_n1, no rp_n3_n2).

def _ankle_type6_scenario(sid="ts3_ankle_t6") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        # Both junctions competent
        Step("P01 SFJ — competent",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 popliteal — SPJ competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P03 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.52, "posterior"),
             expected_action="move"),

        # Calf sweep — searching for perforators
        Step("P04 calf medial upper",
             "probe_move", _pm(sid, "CALF", 0.59, "medial"),
             expected_action="move"),

        Step("P05 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="move"),

        Step("P06 calf lateral upper",
             "probe_move", _pm(sid, "CALF", 0.62, "lateral"),
             expected_action="move"),

        Step("P07 calf lateral mid",
             "probe_move", _pm(sid, "CALF", 0.69, "lateral"),
             expected_action="move"),

        Step("P08 calf lateral lower",
             "probe_move", _pm(sid, "CALF", 0.75, "lateral"),
             expected_action="move"),

        Step("P09 calf posterior lower",
             "probe_move", _pm(sid, "CALF", 0.79, "posterior"),
             expected_action="move"),

        Step("P10 ankle lateral",
             "probe_move", _pm(sid, "ANKLE", 0.87, "lateral"),
             expected_action="move"),

        Step("P11 ankle lateral distal",
             "probe_move", _pm(sid, "ANKLE", 0.91, "lateral"),
             expected_action="move"),

        # Find EP N1→N3 at ankle perforator
        Step("CM-1  clip_mark EP N1→N3 ankle lateral perforator",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.89, "ANKLE", "lateral"),
             expected_action="move",
             forbidden_action="complete"),

        # Trace N3 tributary proximally toward re-entry
        Step("P12 ankle — tracing N3 tributary",
             "probe_move", _pm(sid, "ANKLE", 0.91, "lateral"),
             expected_action="move"),

        Step("P13 calf lateral distal",
             "probe_move", _pm(sid, "CALF", 0.84, "lateral"),
             expected_action="move"),

        Step("P14 calf lateral lower",
             "probe_move", _pm(sid, "CALF", 0.80, "lateral"),
             expected_action="move"),

        Step("P15 calf lateral mid-lower",
             "probe_move", _pm(sid, "CALF", 0.76, "lateral"),
             expected_action="move"),

        Step("P16 calf lateral mid",
             "probe_move", _pm(sid, "CALF", 0.72, "lateral"),
             expected_action="move"),

        Step("P17 calf lateral upper",
             "probe_move", _pm(sid, "CALF", 0.67, "lateral"),
             expected_action="move"),

        Step("P18 calf posterior mid",
             "probe_move", _pm(sid, "CALF", 0.63, "posterior"),
             expected_action="move"),

        # Find RP N3→N1 — Type 6 complete fires
        Step("CM-2  clip_mark RP N3→N1 calf lateral — COMPLETE (Type 6)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.76, "CALF", "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P19 verify complete — calf lateral",
             "probe_move", _pm(sid, "CALF", 0.76, "lateral"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P20 verify complete — popliteal",
             "probe_move", _pm(sid, "POPLITEAL", 0.48, "posterior"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 5 — type3_no_reflux  (26 checked steps)
# ---------------------------------------------------------------------------
# Type 3 full escape circuit, elimination test = "No Reflux".
# Ordering: EP N1→N2 (SFJ) → EP N2→N3 (distal thigh, no elim) → full distal
# sweep → RP N3→N1 (calf) → RP N2→N1 (Hunterian) → MANEUVER fires.
# Then surgeon clips EP N2→N3 again with elimination_test="No Reflux" → COMPLETE.
# IMPORTANT: EP N2→N3 is marked BEFORE RP N2→N1 to block Rule 6 (Type 1/2A).

def _type3_no_reflux_scenario(sid="ts3_t3_noreflux") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01 SFJ approach",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("CM-1  EP N1→N2 at SFJ — entry confirmed",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.07, "SFJ", "anterior-medial"),
             expected_action="move"),

        Step("P02 upper thigh — tracing GSV distally",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.13, "medial"),
             expected_action="move"),

        Step("P03 upper thigh mid",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.20, "medial"),
             expected_action="move"),

        Step("P04 Hunterian proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.27, "medial"),
             expected_action="move"),

        Step("P05 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.33, "medial"),
             expected_action="move"),

        Step("P06 Hunterian distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.38, "medial"),
             expected_action="move"),

        # EP N2→N3 marked BEFORE RP N2→N1 — blocks Rule 6 for the rest of the scenario
        Step("CM-2  EP N2→N3 escape perforator — distal thigh (no elim test yet)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.38, "HUNTERIAN", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P07 popliteal — SPJ competent check",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P08 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.53, "posterior"),
             expected_action="move"),

        Step("P09 calf posterior — SSV assessment",
             "probe_move", _pm(sid, "CALF", 0.60, "posterior"),
             expected_action="move"),

        Step("P10 calf medial",
             "probe_move", _pm(sid, "CALF", 0.66, "medial"),
             expected_action="move"),

        Step("P11 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.72, "medial"),
             expected_action="move"),

        Step("P12 ankle",
             "probe_move", _pm(sid, "ANKLE", 0.85, "medial"),
             expected_action="move"),

        Step("P13 calf return lower",
             "probe_move", _pm(sid, "CALF", 0.78, "medial"),
             expected_action="move"),

        Step("P14 calf return mid",
             "probe_move", _pm(sid, "CALF", 0.72, "medial"),
             expected_action="move"),

        # RP N3→N1 first — re-entry from N3 tributary. maneuver needs rp_n2_n1 too → not yet.
        Step("CM-3  RP N3→N1 — re-entry from N3 at calf",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.72, "CALF", "medial"),
             expected_action="move",
             forbidden_action="maneuver"),

        Step("P15 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="move"),

        Step("P16 calf medial upper",
             "probe_move", _pm(sid, "CALF", 0.58, "medial"),
             expected_action="move"),

        Step("P17 Hunterian — returning to re-entry zone",
             "probe_move", _pm(sid, "HUNTERIAN", 0.33, "medial"),
             expected_action="move"),

        Step("P18 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="move"),

        # RP N2→N1: now ep_n2_n3_no_elim + rp_n3_n1 + rp_n2_n1 → Rule 1 MANEUVER
        Step("CM-4  RP N2→N1 at Hunterian — MANEUVER fires",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.28, "HUNTERIAN", "medial"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary"]),

        Step("P19 calf — maneuver state persists",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="maneuver"),

        # Elimination test "No Reflux" → Rule 2 fires COMPLETE
        Step("CM-5  EP N2→N3 elimination='No Reflux' — COMPLETE (Type 3)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.38, "HUNTERIAN", "medial",
                               elim="No Reflux"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P20 verify complete — calf",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P21 verify complete — Hunterian",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 6 — type1p2_reflux  (24 checked steps)
# ---------------------------------------------------------------------------
# Type 1+2 escape circuit (SPJ/SSV entry), elimination test = "Reflux".
# EP N1→N2 at SPJ → EP N2→N3 along SSV (blocks Rule 6) → full ankle sweep →
# RP N3→N1 at ankle → RP N2→N1 at calf → MANEUVER.
# Then EP N2→N3 clip with elimination_test="Reflux" → COMPLETE (Rule 3).

def _type1p2_reflux_scenario(sid="ts3_t1p2_reflux") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01 SFJ — competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 popliteal — SPJ assessment",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P03 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.52, "posterior"),
             expected_action="move"),

        Step("CM-1  EP N1→N2 at SPJ — SSV entry",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.48, "POPLITEAL", "posterior"),
             expected_action="move"),

        Step("P04 calf posterior — tracing SSV",
             "probe_move", _pm(sid, "CALF", 0.55, "posterior"),
             expected_action="move"),

        Step("P05 calf posterior mid",
             "probe_move", _pm(sid, "CALF", 0.61, "posterior"),
             expected_action="move"),

        Step("P06 calf posterior lower",
             "probe_move", _pm(sid, "CALF", 0.67, "posterior"),
             expected_action="move"),

        # EP N2→N3 marked BEFORE RP N2→N1 — blocks Rule 6
        Step("CM-2  EP N2→N3 escape along SSV — no elim test",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.67, "CALF", "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P07 calf posterior distal",
             "probe_move", _pm(sid, "CALF", 0.73, "posterior"),
             expected_action="move"),

        Step("P08 calf lower",
             "probe_move", _pm(sid, "CALF", 0.79, "posterior"),
             expected_action="move"),

        Step("P09 ankle posterior",
             "probe_move", _pm(sid, "ANKLE", 0.89, "posterior"),
             expected_action="move"),

        Step("P10 ankle return",
             "probe_move", _pm(sid, "ANKLE", 0.83, "posterior"),
             expected_action="move"),

        Step("P11 calf return lower",
             "probe_move", _pm(sid, "CALF", 0.77, "posterior"),
             expected_action="move"),

        Step("P12 calf return mid",
             "probe_move", _pm(sid, "CALF", 0.71, "posterior"),
             expected_action="move"),

        # RP N3→N1: re-entry from N3 at ankle. maneuver needs rp_n2_n1 too → not yet.
        Step("CM-3  RP N3→N1 at ankle — re-entry from N3",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.87, "ANKLE", "posterior"),
             expected_action="move",
             forbidden_action="maneuver"),

        Step("P13 calf return",
             "probe_move", _pm(sid, "CALF", 0.77, "posterior"),
             expected_action="move"),

        Step("P14 calf mid",
             "probe_move", _pm(sid, "CALF", 0.70, "posterior"),
             expected_action="move"),

        Step("P15 popliteal return",
             "probe_move", _pm(sid, "POPLITEAL", 0.57, "posterior"),
             expected_action="move"),

        Step("P16 popliteal proximal",
             "probe_move", _pm(sid, "POPLITEAL", 0.50, "posterior"),
             expected_action="move"),

        # RP N2→N1: now ep_n2_n3_no_elim + rp_n3_n1 + rp_n2_n1 → MANEUVER
        Step("CM-4  RP N2→N1 SSV calf — MANEUVER fires",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.65, "CALF", "posterior"),
             expected_action="maneuver",
             guidance_must_contain=["compress", "tributary"]),

        Step("P17 calf — maneuver state persists",
             "probe_move", _pm(sid, "CALF", 0.60, "posterior"),
             expected_action="maneuver"),

        # Elimination test "Reflux" → Rule 3 fires COMPLETE (Type 1+2)
        Step("CM-5  EP N2→N3 elimination='Reflux' — COMPLETE (Type 1+2)",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.67, "CALF", "posterior",
                               elim="Reflux"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P18 verify complete — popliteal",
             "probe_move", _pm(sid, "POPLITEAL", 0.49, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P19 verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 7 — type5_complex  (25 checked steps)
# ---------------------------------------------------------------------------
# Type 5 four-clip perforator circuit. Both junctions competent. No SFJ/SPJ entry.
# Clip order: EP N1→N3 (Hunterian perforator) → RP N3→N2 (calf) → EP N2→N3 (calf
# distal) → RP N3→N1 (ankle) → Rule 5b fires complete.
# No rp_n2_n1 anywhere → Rule 4 never fires; rp_n3_n2 present → Rule 5 blocked.

def _type5_complex_scenario(sid="ts3_t5_complex") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01 SFJ — competent, no clip",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 popliteal — SPJ competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P03 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.52, "posterior"),
             expected_action="move"),

        Step("P04 upper thigh — searching for perforators",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.15, "medial"),
             expected_action="move"),

        Step("P05 Hunterian proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.21, "medial"),
             expected_action="move"),

        Step("P06 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.27, "medial"),
             expected_action="move"),

        Step("P07 Hunterian distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.33, "medial"),
             expected_action="move"),

        # EP N1→N3: direct perforator from N1 trunk to N3 tributary (not via SFJ/SPJ)
        Step("CM-1  EP N1→N3 Hunterian perforator",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.30, "HUNTERIAN", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P08 distal thigh — tracing N3 distally",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.40, "medial"),
             expected_action="move"),

        Step("P09 popliteal — N3 into popliteal region",
             "probe_move", _pm(sid, "POPLITEAL", 0.48, "posterior"),
             expected_action="move"),

        Step("P10 calf posterior upper",
             "probe_move", _pm(sid, "CALF", 0.56, "posterior"),
             expected_action="move"),

        # RP N3→N2: reflux from N3 tributary back toward N2 collecting trunk
        Step("CM-2  RP N3→N2 — N3→N2 reflux at calf junction",
             "clip_mark", _cm(sid, "RP", "N3", "N2", 0.60, "CALF", "posterior"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P11 calf posterior mid",
             "probe_move", _pm(sid, "CALF", 0.65, "posterior"),
             expected_action="move"),

        Step("P12 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.71, "medial"),
             expected_action="move"),

        Step("P13 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.77, "medial"),
             expected_action="move"),

        Step("P14 ankle",
             "probe_move", _pm(sid, "ANKLE", 0.87, "medial"),
             expected_action="move"),

        Step("P15 calf return lower",
             "probe_move", _pm(sid, "CALF", 0.80, "medial"),
             expected_action="move"),

        Step("P16 calf return mid",
             "probe_move", _pm(sid, "CALF", 0.73, "medial"),
             expected_action="move"),

        Step("P17 calf medial",
             "probe_move", _pm(sid, "CALF", 0.66, "medial"),
             expected_action="move"),

        # EP N2→N3: second escape point — N2 trunk → N3 tributary
        Step("CM-3  EP N2→N3 — second escape at distal calf",
             "clip_mark", _cm(sid, "EP", "N2", "N3", 0.68, "CALF", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        Step("P18 calf lower — tracing to re-entry",
             "probe_move", _pm(sid, "CALF", 0.75, "medial"),
             expected_action="move"),

        Step("P19 ankle approach",
             "probe_move", _pm(sid, "ANKLE", 0.85, "medial"),
             expected_action="move"),

        Step("P20 ankle distal",
             "probe_move", _pm(sid, "ANKLE", 0.90, "medial"),
             expected_action="move"),

        # RP N3→N1: Rule 5b fires — ep_n1_n3 + rp_n3_n2 + ep_n2_n3 + rp_n3_n1
        Step("CM-4  RP N3→N1 at ankle — COMPLETE (Type 5)",
             "clip_mark", _cm(sid, "RP", "N3", "N1", 0.88, "ANKLE", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P21 verify complete — calf",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P22 verify complete — Hunterian",
             "probe_move", _pm(sid, "HUNTERIAN", 0.30, "medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 8 — sfj_entry_long  (23 checked steps)
# ---------------------------------------------------------------------------
# Type 1/2A with standard SFJ entry point. Long multi-zone sweep: EP N1→N2 at
# SFJ early, thigh/popliteal/calf/ankle sweep, back up the leg. RP N2→N1
# confirmed at distal calf once max_visited=0.87 ≥ 0.44 → Rule 6 COMPLETE.

def _sfj_entry_long_scenario(sid="ts3_sfj_long") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01 SFJ approach",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 SFJ deeper scan",
             "probe_move", _pm(sid, "SFJ", 0.09, "anterior-medial"),
             expected_action="move"),

        Step("CM-1  EP N1→N2 at SFJ",
             "clip_mark", _cm(sid, "EP", "N1", "N2", 0.07, "SFJ", "anterior-medial"),
             expected_action="move"),

        Step("P03 upper thigh — scanning GSV distally",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.13, "medial"),
             expected_action="move"),

        Step("P04 upper thigh mid",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.18, "medial"),
             expected_action="move"),

        Step("P05 upper thigh distal",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.23, "medial"),
             expected_action="move"),

        Step("P06 Hunterian proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.28, "medial"),
             expected_action="move"),

        Step("P07 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.33, "medial"),
             expected_action="move"),

        Step("P08 Hunterian distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.38, "medial"),
             expected_action="move"),

        Step("P09 distal thigh",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.43, "medial"),
             expected_action="move"),

        # max_visited = 0.43 < 0.44 at this step — Rule 6 still blocked
        Step("P10 popliteal — SPJ competent check",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        # max_visited = 0.47 ≥ 0.44 from this step onward — but rp_n2_n1 still absent
        Step("P11 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.53, "posterior"),
             expected_action="move"),

        Step("P12 calf posterior",
             "probe_move", _pm(sid, "CALF", 0.60, "posterior"),
             expected_action="move"),

        Step("P13 calf medial",
             "probe_move", _pm(sid, "CALF", 0.66, "medial"),
             expected_action="move"),

        Step("P14 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.72, "medial"),
             expected_action="move"),

        Step("P15 calf medial distal",
             "probe_move", _pm(sid, "CALF", 0.78, "medial"),
             expected_action="move"),

        Step("P16 ankle",
             "probe_move", _pm(sid, "ANKLE", 0.87, "medial"),
             expected_action="move"),

        Step("P17 calf return lower",
             "probe_move", _pm(sid, "CALF", 0.80, "medial"),
             expected_action="move"),

        Step("P18 calf return mid",
             "probe_move", _pm(sid, "CALF", 0.74, "medial"),
             expected_action="move"),

        Step("P19 calf return upper",
             "probe_move", _pm(sid, "CALF", 0.67, "medial"),
             expected_action="move"),

        # max_visited = 0.87 ≥ 0.44; ep_n1_n2=T, rp_n2_n1=T, not ep_n2_n3 → Rule 6 COMPLETE
        Step("CM-2  RP N2→N1 distal calf — COMPLETE (Type 1/2A, max_visited=0.87)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.68, "CALF", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P20 verify complete — popliteal",
             "probe_move", _pm(sid, "POPLITEAL", 0.50, "posterior"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P21 verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Scenario 9 — calf_medial_t4  (25 checked steps)
# ---------------------------------------------------------------------------
# Type 4 with perforator entry found at ankle level (distal), then long
# proximal sweep through calf → popliteal → thigh. RP N2→N1 at Hunterian.
# Rule 4 fires COMPLETE immediately (ep_n1_n3 + rp_n2_n1, no max_visited gate).

def _calf_medial_t4_scenario(sid="ts3_calf_t4") -> list[Step]:
    return [
        Step("stream_start", "stream_start", {"session_id": sid}),

        Step("P01 SFJ — competent",
             "probe_move", _pm(sid, "SFJ", 0.06, "anterior-medial"),
             expected_action="move",
             guidance_must_contain=["SFJ", "groin", "femoral", "junction"]),

        Step("P02 popliteal — SPJ competent",
             "probe_move", _pm(sid, "POPLITEAL", 0.47, "posterior"),
             expected_action="move",
             guidance_must_contain=["popliteal", "SPJ", "posterior"]),

        Step("P03 popliteal deeper",
             "probe_move", _pm(sid, "POPLITEAL", 0.52, "posterior"),
             expected_action="move"),

        Step("P04 calf posterior — searching distally",
             "probe_move", _pm(sid, "CALF", 0.59, "posterior"),
             expected_action="move"),

        Step("P05 calf medial upper",
             "probe_move", _pm(sid, "CALF", 0.64, "medial"),
             expected_action="move"),

        Step("P06 calf medial mid",
             "probe_move", _pm(sid, "CALF", 0.70, "medial"),
             expected_action="move"),

        Step("P07 calf medial lower",
             "probe_move", _pm(sid, "CALF", 0.76, "medial"),
             expected_action="move"),

        Step("P08 ankle medial",
             "probe_move", _pm(sid, "ANKLE", 0.85, "medial"),
             expected_action="move"),

        Step("P09 ankle lateral",
             "probe_move", _pm(sid, "ANKLE", 0.91, "lateral"),
             expected_action="move"),

        Step("P10 ankle return medial",
             "probe_move", _pm(sid, "ANKLE", 0.85, "medial"),
             expected_action="move"),

        Step("P11 calf lower return",
             "probe_move", _pm(sid, "CALF", 0.78, "medial"),
             expected_action="move"),

        # EP N1→N3 at ankle — perforator entry directly from N1 to N3 tributary
        Step("CM-1  EP N1→N3 ankle perforator — direct N1→N3 entry",
             "clip_mark", _cm(sid, "EP", "N1", "N3", 0.88, "ANKLE", "medial"),
             expected_action="move",
             forbidden_action="complete"),

        # Sweep proximally looking for RP N2→N1 in GSV trunk
        Step("P12 calf lateral — scanning for trunk reflux",
             "probe_move", _pm(sid, "CALF", 0.80, "lateral"),
             expected_action="move"),

        Step("P13 calf posterior",
             "probe_move", _pm(sid, "CALF", 0.74, "posterior"),
             expected_action="move"),

        Step("P14 calf posterior upper",
             "probe_move", _pm(sid, "CALF", 0.67, "posterior"),
             expected_action="move"),

        Step("P15 popliteal",
             "probe_move", _pm(sid, "POPLITEAL", 0.54, "posterior"),
             expected_action="move"),

        Step("P16 popliteal proximal",
             "probe_move", _pm(sid, "POPLITEAL", 0.48, "posterior"),
             expected_action="move"),

        Step("P17 distal thigh",
             "probe_move", _pm(sid, "DISTAL_THIGH", 0.41, "medial"),
             expected_action="move"),

        Step("P18 Hunterian distal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.35, "medial"),
             expected_action="move"),

        Step("P19 Hunterian mid",
             "probe_move", _pm(sid, "HUNTERIAN", 0.29, "medial"),
             expected_action="move"),

        Step("P20 Hunterian proximal",
             "probe_move", _pm(sid, "HUNTERIAN", 0.23, "medial"),
             expected_action="move"),

        Step("P21 upper thigh",
             "probe_move", _pm(sid, "UPPER_THIGH", 0.16, "medial"),
             expected_action="move"),

        # RP N2→N1 at Hunterian: ep_n1_n3 + rp_n2_n1 → Rule 4 COMPLETE immediately
        Step("CM-2  RP N2→N1 Hunterian — COMPLETE (Type 4, immediate)",
             "clip_mark", _cm(sid, "RP", "N2", "N1", 0.28, "HUNTERIAN", "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification", "complete", "mapped"]),

        Step("P22 verify complete — calf",
             "probe_move", _pm(sid, "CALF", 0.65, "medial"),
             expected_action="complete",
             guidance_must_contain=["circuit", "classification"]),

        Step("P23 verify complete — SFJ",
             "probe_move", _pm(sid, "SFJ", 0.07, "anterior-medial"),
             expected_action="complete"),
    ]


# ---------------------------------------------------------------------------
# Registry
# ---------------------------------------------------------------------------

_SCENARIOS: dict[str, callable] = {
    "hunterian_entry":  _hunterian_entry_scenario,
    "ssv_extended":     _ssv_extended_scenario,
    "left_type4":       _left_type4_scenario,
    "ankle_type6":      _ankle_type6_scenario,
    "type3_no_reflux":  _type3_no_reflux_scenario,
    "type1p2_reflux":   _type1p2_reflux_scenario,
    "type5_complex":    _type5_complex_scenario,
    "sfj_entry_long":   _sfj_entry_long_scenario,
    "calf_medial_t4":   _calf_medial_t4_scenario,
}

_SCENARIO_DESCRIPTIONS = {
    "hunterian_entry": (
        "Hunterian perforator entry (SFJ competent) — Type 1 via mid-thigh EP N1→N2; "
        "full distal sweep to ankle; complete fires at RP N2→N1 clip_mark."
    ),
    "ssv_extended": (
        "SSV extended — SPJ entry, full SSV sweep to ankle and back; "
        "complete fires at RP N2→N1 clip_mark (no escape found, max_visited=0.88)."
    ),
    "left_type4": (
        "Left leg Type 4 — SFJ+SPJ left competent; EP N1→N3 left calf; "
        "RP N2→N1 left Hunterian → Type 4 complete (ep_n1_n3 + rp_n2_n1)."
    ),
    "ankle_type6": (
        "Ankle Type 6 — both junctions competent; EP N1→N3 ankle lateral; "
        "RP N3→N1 calf lateral → Type 6 complete (no trunk)."
    ),
    "type3_no_reflux": (
        "Type 3 escape circuit (No Reflux elim) — SFJ entry; EP N2→N3 marked before "
        "RP N2→N1 to block Rule 6; full sweep; RP N3→N1 + RP N2→N1 → maneuver; "
        "elim='No Reflux' → Type 3 complete."
    ),
    "type1p2_reflux": (
        "Type 1+2 escape circuit (Reflux elim) — SPJ/SSV entry; EP N2→N3 along SSV "
        "before RP N2→N1; full ankle sweep; RP N3→N1 + RP N2→N1 → maneuver; "
        "elim='Reflux' → Type 1+2 complete."
    ),
    "type5_complex": (
        "Type 5 four-clip perforator circuit — both junctions competent; EP N1→N3 "
        "Hunterian; RP N3→N2 calf; EP N2→N3 distal calf; RP N3→N1 ankle → "
        "Rule 5b complete (no rp_n2_n1)."
    ),
    "sfj_entry_long": (
        "Type 1/2A standard SFJ entry — EP N1→N2 at SFJ; long multi-zone sweep to "
        "ankle (max_visited=0.87); RP N2→N1 at calf → Rule 6 complete."
    ),
    "calf_medial_t4": (
        "Type 4 distal perforator — EP N1→N3 at ankle (0.88); long proximal sweep; "
        "RP N2→N1 at Hunterian → Rule 4 complete immediately (no max_visited gate)."
    ),
}


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _describe_movement(event: str, data: dict) -> str:
    if event == "stream_start":
        return f"Start session  (id: {data.get('session_id', '?')})"
    if event == "probe_move":
        return (f"Move probe → {data.get('region', '?')}\n"
                f"posY={data.get('pos_y_ratio', 0.0):.2f}  "
                f"{data.get('surface', '?')}  {data.get('leg', 'right')} leg")
    if event == "clip_mark":
        base = (f"Mark: {data.get('flow', '?')} "
                f"{data.get('from_type', '?')}→{data.get('to_type', '?')}  "
                f"posY={data.get('pos_y_ratio', 0.0):.2f}  {data.get('region', '?')}  "
                f"{data.get('leg', 'right')} leg")
        if data.get("elimination_test"):
            base += f"\nElim: {data['elimination_test']}"
        return base
    return event


def _describe_expectation(step: Step) -> str:
    if step.event == "stream_start":
        return "session_ready event"
    parts = []
    if step.expected_action:
        parts.append(f'Action = "{step.expected_action}"')
    if step.forbidden_action:
        parts.append(f'Must NOT = "{step.forbidden_action}"')
    if step.guidance_must_contain:
        parts.append("Guidance contains any of: " + ", ".join(step.guidance_must_contain))
    return "\n".join(parts) if parts else "observe only"


# ---------------------------------------------------------------------------
# Runner
# ---------------------------------------------------------------------------

CYAN   = "\033[96m"
GREEN  = "\033[92m"
RED    = "\033[91m"
YELLOW = "\033[93m"
RESET  = "\033[0m"
BOLD   = "\033[1m"

_SKIP_EVENTS = {"stream_start"}


class ScenarioRunner:
    def __init__(self, api_base: str):
        self.api_base = api_base.rstrip("/")
        self._sio = socketio.SimpleClient()

    def connect(self) -> None:
        print(f"Connecting to {self.api_base} ...")
        self._sio.connect(self.api_base, wait_timeout=10)
        print("Connected.\n")

    def disconnect(self) -> None:
        self._sio.disconnect()

    def _emit_and_wait(self, event: str, data: dict, timeout: float = 60.0):
        self._sio.emit(event, data)
        if event in _SKIP_EVENTS:
            deadline = time.time() + timeout
            while time.time() < deadline:
                try:
                    ev = self._sio.receive(timeout=2)
                    if ev and ev[0] == "session_ready":
                        return ev[1]
                except Exception:
                    pass
            return None
        deadline = time.time() + timeout
        while time.time() < deadline:
            try:
                ev = self._sio.receive(timeout=2)
                if ev and ev[0] == "guidance_update":
                    return ev[1]
            except Exception:
                pass
        return None

    def run(self, steps: list[Step]) -> list[dict]:
        results = []
        for step in steps:
            print(f"\n{CYAN}--- {step.label} ---{RESET}")
            print(f"  event: {step.event}")

            movement    = _describe_movement(step.event, step.data)
            expectation = _describe_expectation(step)
            response    = self._emit_and_wait(step.event, step.data)

            if step.event in _SKIP_EVENTS:
                if response:
                    print(f"  session_ready: {response}")
                    results.append({"step": step.label, "status": "SKIP",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                else:
                    print(f"  {YELLOW}No session_ready{RESET}")
                    results.append({"step": step.label, "status": "TIMEOUT",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                continue

            if (step.expected_action is None
                    and not step.guidance_must_contain
                    and step.forbidden_action is None):
                if response:
                    g = response.get("guidance", "")
                    a = response.get("action", "?")
                    print(f"  guidance: {g!r}  action: {a}")
                    results.append({"step": step.label, "status": "SKIP",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": a, "llm_guidance": g, "failures": []})
                else:
                    print(f"  {YELLOW}timeout{RESET}")
                    results.append({"step": step.label, "status": "TIMEOUT",
                                    "movement": movement, "expectation": expectation,
                                    "llm_action": "--", "llm_guidance": "--",
                                    "failures": []})
                continue

            if response is None:
                print(f"  {YELLOW}No response (timeout){RESET}")
                results.append({"step": step.label, "status": "TIMEOUT",
                                "movement": movement, "expectation": expectation,
                                "llm_action": "--", "llm_guidance": "--",
                                "failures": []})
                continue

            llm_guidance = response.get("guidance", "")
            llm_action   = response.get("action", "move")
            print(f"  guidance: {llm_guidance!r}")
            print(f"  action  : {llm_action}")

            passed   = True
            failures = []

            if step.expected_action and llm_action != step.expected_action:
                passed = False
                failures.append(
                    f'Action: got "{llm_action}", expected "{step.expected_action}"')

            if step.guidance_must_contain:
                if not any(kw.lower() in llm_guidance.lower()
                           for kw in step.guidance_must_contain):
                    passed = False
                    failures.append(
                        "Guidance missing any of: "
                        + ", ".join(step.guidance_must_contain))

            if step.forbidden_action and llm_action == step.forbidden_action:
                passed = False
                failures.append(
                    f'Action "{llm_action}" must NOT fire here')

            if passed:
                print(f"  {GREEN}{BOLD}PASS{RESET}")
            else:
                for f in failures:
                    print(f"  {RED}FAIL: {f}{RESET}")

            results.append({"step": step.label,
                            "status": "PASS" if passed else "FAIL",
                            "movement": movement, "expectation": expectation,
                            "llm_action": llm_action, "llm_guidance": llm_guidance,
                            "failures": failures})
            time.sleep(0.4)

        return results

    def print_summary(self, label: str, results: list[dict]) -> bool:
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        passed  = [r for r in checked if r["status"] == "PASS"]
        failed  = [r for r in checked if r["status"] == "FAIL"]
        timeout = [r for r in results if r["status"] == "TIMEOUT"]

        print(f"\n{'─'*60}")
        print(f"{BOLD}{label}{RESET}  "
              f"{GREEN}{len(passed)} passed{RESET}  "
              f"{RED}{len(failed)} failed{RESET}  "
              f"{YELLOW}{len(timeout)} timeout{RESET}  "
              f"/ {len(checked)} checked")
        for r in failed:
            print(f"  {RED}x {r['step']}{RESET}")
            for f in r.get("failures", []):
                print(f"      {f}")
        return len(failed) == 0 and len(timeout) == 0


# ---------------------------------------------------------------------------
# Word report
# ---------------------------------------------------------------------------

_STATUS_COLORS = {"PASS": "00B050", "FAIL": "FF0000",
                  "TIMEOUT": "FF8C00", "SKIP": "808080"}


def _colored_cell(cell, text: str, hex_color: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml   import OxmlElement
    from docx.shared import RGBColor
    cell.text = text
    run = (cell.paragraphs[0].runs[0]
           if cell.paragraphs[0].runs
           else cell.paragraphs[0].add_run(text))
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.bold = True
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, pt: int = 9, bold=False,
                   italic=False, color_rgb=None) -> None:
    cell.text = ""
    for i, line in enumerate(text.split("\n")):
        para = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        run  = para.add_run(line)
        run.font.size   = __import__("docx").shared.Pt(pt)
        run.font.bold   = bold
        run.font.italic = italic
        if color_rgb:
            run.font.color.rgb = color_rgb


def write_word_report(all_results: dict[str, list[dict]],
                      api_base: str, out_path: str) -> None:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin  = sec.bottom_margin = Inches(0.7)

    t = doc.add_heading("CHIVA Streaming — Test Set 3 Report (9 Scenarios)", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m = doc.add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mr = m.add_run(f"Date: {datetime.now().strftime('%Y-%m-%d  %H:%M:%S')}    |    "
                   f"API: {api_base}    |    Scenarios: {len(all_results)}")
    mr.font.size = Pt(10)
    mr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    doc.add_heading("Overall Summary", level=1)
    stbl = doc.add_table(rows=1, cols=5)
    stbl.style = "Table Grid"
    for i, h in enumerate(["Scenario", "Description", "Passed", "Failed", "Timeouts"]):
        c = stbl.rows[0].cells[i]
        c.text = h
        if c.paragraphs[0].runs:
            c.paragraphs[0].runs[0].font.bold = True

    overall_pass = True
    for name, results in all_results.items():
        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")
        row = stbl.add_row().cells
        row[0].text = name
        row[1].text = _SCENARIO_DESCRIPTIONS.get(name, "")
        row[2].text = str(p); row[3].text = str(f); row[4].text = str(t)
        if f or t:
            overall_pass = False

    doc.add_paragraph()
    vp = doc.add_paragraph()
    vp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    vr = vp.add_run("OVERALL: " + ("ALL PASS" if overall_pass else "FAILURES PRESENT"))
    vr.font.bold = True; vr.font.size = Pt(13)
    vr.font.color.rgb = (RGBColor(0x00, 0xB0, 0x50) if overall_pass
                         else RGBColor(0xFF, 0x00, 0x00))

    for name, results in all_results.items():
        doc.add_page_break()
        doc.add_heading(f"Scenario: {name.upper()}", level=1)
        dp = doc.add_paragraph(_SCENARIO_DESCRIPTIONS.get(name, ""))
        if dp.runs:
            dp.runs[0].font.italic = True

        checked = [r for r in results if r["status"] in ("PASS", "FAIL")]
        p = sum(1 for r in checked if r["status"] == "PASS")
        f = sum(1 for r in checked if r["status"] == "FAIL")
        t = sum(1 for r in results  if r["status"] == "TIMEOUT")

        sl = doc.add_paragraph()
        sl.add_run(f"Checked: {len(checked)}   |   ")
        pr = sl.add_run(f"Passed: {p}"); pr.font.color.rgb = RGBColor(0x00, 0xB0, 0x50)
        sl.add_run("   |   ")
        fr = sl.add_run(f"Failed: {f}"); fr.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
        sl.add_run("   |   ")
        tr = sl.add_run(f"Timeouts: {t}"); tr.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        doc.add_paragraph()

        tbl = doc.add_table(rows=1, cols=6)
        tbl.style = "Table Grid"
        for i, h in enumerate(["#", "Surgeon Action", "Expected",
                                "LLM Guidance", "Action", "Status"]):
            c = tbl.rows[0].cells[i]
            c.text = h
            if c.paragraphs[0].runs:
                c.paragraphs[0].runs[0].font.bold = True
        for i, w in enumerate([0.25, 1.80, 2.10, 2.10, 0.80, 0.65]):
            for cell in tbl.columns[i].cells:
                cell.width = Inches(w)

        for idx, r in enumerate(results, 1):
            status = r.get("status", "SKIP")
            row    = tbl.add_row().cells
            row[0].text = str(idx)
            _set_cell_text(row[1], r.get("movement", "--"), pt=8)
            _set_cell_text(row[2], r.get("expectation", "--"), pt=8, italic=True)
            _set_cell_text(row[3], r.get("llm_guidance", "--") or "--", pt=8)
            _set_cell_text(row[4], r.get("llm_action", "--") or "--", pt=9, bold=True)
            _colored_cell(row[5], status, _STATUS_COLORS.get(status, "808080"))
            if status == "FAIL" and r.get("failures"):
                fr_row = tbl.add_row().cells
                merged = fr_row[0].merge(fr_row[5])
                merged.text = ""
                for fn in r["failures"]:
                    pp = merged.add_paragraph(f"  x  {fn}")
                    if pp.runs:
                        pp.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                        pp.runs[0].font.size      = Pt(8)
                        pp.runs[0].font.italic    = True

    os.makedirs(os.path.dirname(out_path), exist_ok=True)
    doc.save(out_path)
    print(f"\nReport saved: {out_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(description="CHIVA stream test set 3 runner")
    parser.add_argument("scenario", nargs="?", default="all",
                        choices=list(_SCENARIOS.keys()) + ["all"],
                        help="Scenario to run (default: all)")
    parser.add_argument("--api",  default="http://localhost:7861")
    parser.add_argument("--all",  action="store_true", dest="run_all")
    args = parser.parse_args()

    to_run = list(_SCENARIOS.items()) if (args.run_all or args.scenario == "all") \
             else [(args.scenario, _SCENARIOS[args.scenario])]

    runner = ScenarioRunner(args.api)
    all_ok = True
    all_results: dict[str, list[dict]] = {}

    try:
        runner.connect()
        for name, builder in to_run:
            print(f"\n{'='*60}")
            print(f"{BOLD}SCENARIO: {name.upper()}{RESET}")
            print(f"{'='*60}")
            steps   = builder()
            results = runner.run(steps)
            ok      = runner.print_summary(name, results)
            all_ok  = all_ok and ok
            all_results[name] = results
    except KeyboardInterrupt:
        print("\nInterrupted.")
    finally:
        runner.disconnect()

    if len(to_run) > 1:
        print(f"\n{'='*60}")
        print(f"{BOLD}OVERALL: {'ALL PASS' if all_ok else 'FAILURES PRESENT'}{RESET}")

    if all_results:
        ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
        here     = os.path.dirname(os.path.abspath(__file__))
        out_path = os.path.join(here, "results", f"test_set3_log_{ts}.docx")
        try:
            write_word_report(all_results, args.api, out_path)
        except Exception as exc:
            print(f"{YELLOW}Warning: could not write Word report — {exc}{RESET}")

    sys.exit(0 if all_ok else 1)


if __name__ == "__main__":
    main()
