"""
Run baseline and fine-tuned models on Seen/Unseen test frames,
produce side-by-side comparison panels, and insert into the report Word doc.
"""
import sys, os, io, glob as _glob
import numpy as np
import torch
import torch.nn.functional as F
from PIL import Image, ImageDraw, ImageFont
import cv2

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, os.path.join(BASE_DIR, 'stubs'))
sys.path.insert(0, os.path.join(BASE_DIR, 'BiomedParse'))

from detectron2.structures import ImageList
from modeling.BaseModel import BaseModel
from modeling import build_model
from utilities.distributed import init_distributed
from utilities.arguments import load_opt_from_config_files
from utilities.constants import BIOMED_CLASSES

# ── Word imports ──────────────────────────────────────────────────────────────
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# Image selection
# ─────────────────────────────────────────────────────────────────────────────
IMG_ROOT = r'C:\Users\Krish\Pictures\Screenshots\Task_4_Test_Images'

SEEN_FILES = sorted(_glob.glob(os.path.join(IMG_ROOT, 'Seen_Frames', '*.png')))

UNSEEN_ALL = sorted(_glob.glob(os.path.join(IMG_ROOT, 'Unseen_Frames', '*.png')))
# 0-5   : resized 512x512 images (all distinct) — keep all
# 6-20  : diverse probe positions at 1-3 min intervals — keep all
# 21-50 : dense consecutive frames at ~15s intervals — keep every 3rd to avoid near-duplicates
UNSEEN_FILES = (
    UNSEEN_ALL[0:21] +
    UNSEEN_ALL[21::3]
)

print(f"Seen frames   : {len(SEEN_FILES)}")
print(f"Unseen frames : {len(UNSEEN_FILES)}")

# ─────────────────────────────────────────────────────────────────────────────
# Preprocessing — match training setup: letterbox to 1024×1024
# ─────────────────────────────────────────────────────────────────────────────
TRAIN_SIZE = 1024

def preprocess(path):
    """Load RGBA/RGB image, letterbox to TRAIN_SIZE × TRAIN_SIZE, return RGB PIL."""
    img = Image.open(path).convert('RGB')
    w, h = img.size
    scale = TRAIN_SIZE / max(w, h)
    new_w, new_h = int(w * scale), int(h * scale)
    img_r = img.resize((new_w, new_h), Image.BICUBIC)
    canvas = Image.new('RGB', (TRAIN_SIZE, TRAIN_SIZE), (0, 0, 0))
    pad_x = (TRAIN_SIZE - new_w) // 2
    pad_y = (TRAIN_SIZE - new_h) // 2
    canvas.paste(img_r, (pad_x, pad_y))
    return canvas

# ─────────────────────────────────────────────────────────────────────────────
# Model loading
# ─────────────────────────────────────────────────────────────────────────────
CFG = os.path.join(BASE_DIR, 'BiomedParse', 'configs', 'biomed_fascia_finetuning.yaml')

def load_model(weights_path, label):
    print(f"Loading {label} ...")
    opt = load_opt_from_config_files([CFG])
    opt = init_distributed(opt)
    m = BaseModel(opt, build_model(opt)).from_pretrained(weights_path).eval().cuda()
    with torch.no_grad():
        m.model.sem_seg_head.predictor.lang_encoder.get_text_embeddings(
            BIOMED_CLASSES + ['background'], is_eval=True)
    print(f"  {label} loaded.")
    return m

BASELINE_W = os.path.join(BASE_DIR, 'pretrained', 'biomedparse_v1.pt')
FASCIA_W   = sorted(_glob.glob(os.path.join(BASE_DIR, 'BiomedParse', 'output',
                'fascia_finetuning_v2_production', '**', 'model_state_dict.pt'),
                recursive=True), key=os.path.getmtime)[-1]
VEIN_W     = sorted(_glob.glob(os.path.join(BASE_DIR, 'BiomedParse', 'output',
                'fascia_vein_finetuning', '**', 'model_state_dict.pt'),
                recursive=True), key=os.path.getmtime)[-1]

baseline_model = load_model(BASELINE_W, 'Baseline')
fascia_model   = load_model(FASCIA_W,   'Fine-tuned Fascia')
vein_model     = load_model(VEIN_W,     'Fine-tuned Vein')

# ─────────────────────────────────────────────────────────────────────────────
# Inference helpers
# ─────────────────────────────────────────────────────────────────────────────
FASCIA_TEXT = 'fascia layer in PeripheralVascular Ultrasound'
VEIN_TEXT   = ('small oval anechoic dark void vein lumen in cross-section '
               'peripheral vascular ultrasound below fascia')

def _grounding_prob(mdl, img_pil, text, infer_size=512, top_bias=False):
    m = mdl.model; pred = m.sem_seg_head.predictor; W, H = img_pil.size
    arr = np.asarray(img_pil.resize((infer_size, infer_size), Image.BICUBIC)).astype(np.float32)
    img_t = torch.from_numpy(arr.copy()).permute(2, 0, 1).cuda()
    images = ImageList.from_tensors([(img_t - m.pixel_mean) / m.pixel_std], m.size_divisibility)
    gtext = pred.lang_encoder.get_text_token_embeddings([text], name='grounding', token=False, norm=False)
    tok_emb = gtext['token_emb']; tok_mask = gtext['tokens']['attention_mask'].bool()
    q_emb = tok_emb[tok_mask]
    nz_mask = torch.zeros(q_emb[:, None].shape[:-1], dtype=torch.bool, device=q_emb.device)
    extra = {'grounding_tokens': q_emb[:, None],
             'grounding_nonzero_mask': nz_mask.t(),
             'grounding_class': gtext['class_emb']}
    with torch.no_grad():
        feats = m.backbone(images.tensor)
        mf, _, ms = m.sem_seg_head.pixel_decoder.forward_features(feats)
        outputs = pred(ms, mf, extra=extra, task='grounding_eval')
    all_gm = outputs['pred_gmasks'][0]
    probs = torch.sigmoid(all_gm)
    if top_bias:
        Hm = probs.shape[1]
        vw = torch.linspace(2.0, 0.2, Hm, device=probs.device).view(1, Hm, 1)
        weighted = (probs * vw).reshape(101, -1).max(dim=1).values
    else:
        weighted = probs.reshape(101, -1).max(dim=1).values
    best_q = weighted.argmax().item()
    return F.interpolate(
        all_gm[best_q:best_q + 1][None], (H, W), mode='bilinear', align_corners=False
    )[0, 0].sigmoid().detach().cpu().numpy().astype(np.float32)

def run_inference(img_pil, fascia_mdl, vein_mdl):
    """Returns (fascia_prob, vein_prob) float32 arrays same size as img_pil."""
    f_prob = _grounding_prob(fascia_mdl, img_pil, FASCIA_TEXT, top_bias=True)
    v_prob = _grounding_prob(vein_mdl,   img_pil, VEIN_TEXT,   top_bias=False)
    return f_prob, v_prob

# ─────────────────────────────────────────────────────────────────────────────
# Overlay visualisation
# ─────────────────────────────────────────────────────────────────────────────
FASCIA_COL = (0,  230, 230)   # cyan
VEIN_COL   = (0,  210,  0)   # green

def make_mask_overlay(img_rgb_np, fascia_prob, vein_prob,
                      fascia_thresh=0.35, vein_thresh=0.5, alpha=0.55):
    """Blend fascia (cyan) + vein (green) filled masks onto the image."""
    out = img_rgb_np.copy().astype(np.float32)
    h, w = out.shape[:2]
    fp = cv2.resize(fascia_prob, (w, h))
    vp = cv2.resize(vein_prob,   (w, h))
    for mask, col in [(fp > fascia_thresh, FASCIA_COL), (vp > vein_thresh, VEIN_COL)]:
        if mask.any():
            layer = np.zeros_like(out)
            layer[mask] = col
            out = np.where(mask[:, :, None], out * (1 - alpha) + layer * alpha, out)
    return out.clip(0, 255).astype(np.uint8)


def make_outline_annotated(img_rgb_np, fascia_prob, vein_prob,
                           fascia_thresh=0.35, vein_thresh=0.5):
    """Draw boundary contour outlines for fascia (cyan) and vein (green) on the image."""
    out = img_rgb_np.copy()
    h, w = out.shape[:2]
    fp = cv2.resize(fascia_prob, (w, h))
    vp = cv2.resize(vein_prob,   (w, h))

    fascia_mask = (fp > fascia_thresh).astype(np.uint8)
    vein_mask   = (vp > vein_thresh).astype(np.uint8)

    if fascia_mask.any():
        cnts, _ = cv2.findContours(fascia_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        cv2.drawContours(out, cnts, -1, FASCIA_COL, 3)

    if vein_mask.any():
        cnts, _ = cv2.findContours(vein_mask, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        cv2.drawContours(out, cnts, -1, VEIN_COL, 3)

    return out


def add_label_bar(img_np, label, font_size=18):
    """Add a dark label bar at the top of the image."""
    h, w = img_np.shape[:2]
    bar_h = 32
    bar = np.ones((bar_h, w, 3), dtype=np.uint8) * 30
    pil = Image.fromarray(np.vstack([bar, img_np]))
    draw = ImageDraw.Draw(pil)
    try:
        font = ImageFont.truetype("arial.ttf", font_size)
    except Exception:
        font = ImageFont.load_default()
    bbox = draw.textbbox((0, 0), label, font=font)
    tw, th = bbox[2] - bbox[0], bbox[3] - bbox[1]
    draw.text(((w - tw) // 2, (bar_h - th) // 2), label, fill=(255, 255, 255), font=font)
    return np.array(pil)


def make_panel(img_pil, baseline_f, baseline_v, ft_f, ft_v, panel_size=320):
    """4-column panel: Input | Baseline mask | Baseline outline | FT mask | FT outline."""
    img_np = np.array(img_pil)

    b_mask    = make_mask_overlay(img_np, baseline_f, baseline_v)
    b_outline = make_outline_annotated(img_np, baseline_f, baseline_v)
    ft_mask   = make_mask_overlay(img_np, ft_f, ft_v)
    ft_outline = make_outline_annotated(img_np, ft_f, ft_v)

    def resize(a):
        return np.array(Image.fromarray(a).resize((panel_size, panel_size), Image.BICUBIC))

    cols = [
        add_label_bar(resize(img_np),     'Input'),
        add_label_bar(resize(b_mask),     'Baseline — mask'),
        add_label_bar(resize(b_outline),  'Baseline — outline'),
        add_label_bar(resize(ft_mask),    'Fine-tuned — mask'),
        add_label_bar(resize(ft_outline), 'Fine-tuned — outline'),
    ]

    sep = np.ones((cols[0].shape[0], 3, 3), dtype=np.uint8) * 160
    panel = cols[0]
    for c in cols[1:]:
        panel = np.hstack([panel, sep, c])
    return Image.fromarray(panel)

# ─────────────────────────────────────────────────────────────────────────────
# Run all images
# ─────────────────────────────────────────────────────────────────────────────
PANEL_DIR = os.path.join(BASE_DIR, 'report_panels')
os.makedirs(PANEL_DIR, exist_ok=True)

def process_set(file_list, label):
    panels = []
    for i, path in enumerate(file_list):
        fname = os.path.basename(path)
        print(f"  [{label}] {i+1}/{len(file_list)} — {fname}")
        img_pil = preprocess(path)

        # Baseline: use baseline model for both fascia and vein
        with torch.no_grad():
            b_fascia, b_vein = run_inference(img_pil, baseline_model, baseline_model)
            ft_fascia, ft_vein = run_inference(img_pil, fascia_model, vein_model)

        panel = make_panel(img_pil, b_fascia, b_vein, ft_fascia, ft_vein)
        out_path = os.path.join(PANEL_DIR, f'{label}_{i:02d}_{fname}')
        panel.save(out_path)
        panels.append((out_path, fname))
    return panels

print("\n=== Processing Seen Frames ===")
seen_panels = process_set(SEEN_FILES, 'seen')

print("\n=== Processing Unseen Frames ===")
unseen_panels = process_set(UNSEEN_FILES, 'unseen')

# ─────────────────────────────────────────────────────────────────────────────
# Load existing Word doc and append results section
# ─────────────────────────────────────────────────────────────────────────────
DOC_PATH = os.path.join(BASE_DIR, 'FasciaVeinDetection_Report.docx')
doc = Document(DOC_PATH)

def set_cell_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto'); shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

BLUE  = RGBColor(0x1e, 0x40, 0xaf)
BLACK = RGBColor(0x0f, 0x17, 0x2a)
GRAY  = RGBColor(0x47, 0x55, 0x69)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    run.font.color.rgb = BLUE if level == 1 else BLACK
    run.font.bold = True
    p.paragraph_format.space_before = Pt(18 if level == 1 else 10)
    p.paragraph_format.space_after  = Pt(6)

def add_body(text, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(10.5); r.font.italic = italic
    p.paragraph_format.space_after = Pt(6)

def add_panels_to_doc(panels, section_title, section_body):
    add_heading(section_title, level=2)
    add_body(section_body)

    PANEL_W = Inches(9.8)   # 5 columns × 320px — fits landscape page with 0.5" margins

    for panel_path, fname in panels:
        # Filename label
        p = doc.add_paragraph()
        r = p.add_run(os.path.splitext(fname)[0])
        r.font.size = Pt(9); r.font.color.rgb = GRAY; r.font.italic = True
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after  = Pt(2)

        # Insert image
        doc.add_picture(panel_path, width=PANEL_W)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Legend
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p.add_run('■ ')
        r1.font.color.rgb = RGBColor(0, 230, 230); r1.font.size = Pt(10)
        p.add_run('Fascia   ').font.size = Pt(9)
        r2 = p.add_run('■ ')
        r2.font.color.rgb = RGBColor(0, 210, 0); r2.font.size = Pt(10)
        p.add_run('Vein').font.size = Pt(9)
        p.paragraph_format.space_after = Pt(8)

doc.add_page_break()

# Switch to landscape for wide panels
from docx.oxml import OxmlElement as _OE
from docx.oxml.ns import qn as _qn
from docx.shared import Inches as _In
new_section = doc.add_section()
new_section.orientation = 1   # WD_ORIENT.LANDSCAPE
new_section.page_width  = _In(11)
new_section.page_height = _In(8.5)
new_section.left_margin = new_section.right_margin = _In(0.5)
new_section.top_margin  = new_section.bottom_margin = _In(0.5)

# Section 9 heading
add_heading('9. Visual Results — Baseline vs. Fine-Tuned')
add_body(
    'Each panel shows five views of the same ultrasound frame: Input (letterboxed to 1024×1024 '
    'to match training format), Baseline mask overlay, Baseline boundary outline annotation, '
    'Fine-tuned mask overlay, and Fine-tuned boundary outline annotation. '
    'Cyan = fascia; Green = vein.'
)

add_panels_to_doc(
    seen_panels,
    '9.1  Seen Frames (in-distribution)',
    'These frames come from the same ultrasound machine and acquisition protocol used for training. '
    'The fine-tuned model is expected to perform well; the baseline may miss fascia entirely.'
)

add_panels_to_doc(
    unseen_panels,
    '9.2  Unseen Frames (out-of-distribution)',
    'These frames come from different sources (different machine, brightness, or gain settings). '
    'They reveal the domain shift limitation: the fine-tuned model can hallucinate detections '
    'on images outside the training distribution, while the baseline is consistently poor on fascia.'
)

doc.save(DOC_PATH)
print(f'\nWord document updated: {DOC_PATH}')
print(f'Panels saved to: {PANEL_DIR}')
