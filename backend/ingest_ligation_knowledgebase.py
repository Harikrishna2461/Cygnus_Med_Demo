"""
Ligation Knowledgebase Ingestion
=================================
Ingests:
  1. Ligation_Knowledgebase_1.docx  (text document)
  2. Ligation_Knowledgebase_2.pdf   (images + diagrams — vision OCR)

Into a dedicated Qdrant collection: ligation_knowledgebase_db

Run:
    cd backend
    python ingest_ligation_knowledgebase.py
"""

import os
import time
import base64
import logging
import numpy as np
import requests
from pathlib import Path
from groq import Groq
import fitz  # pymupdf
import docx  # python-docx

from qdrant_client import QdrantClient
from qdrant_client.models import Distance, VectorParams, PointStruct

from config import (
    CHUNK_SIZE,
    CHUNK_OVERLAP,
    OLLAMA_BASE_URL,
    OLLAMA_EMBEDDING_MODEL,
    QDRANT_PATH,
    QDRANT_HOST,
    QDRANT_PORT,
    QDRANT_API_KEY,
    QDRANT_LIGATION_COLLECTION,
    EMBEDDING_DIMENSION,
    GROQ_API_KEY,
    LOG_FILE,
    LOG_LEVEL,
)

GROQ_VISION_MODEL = "meta-llama/llama-4-scout-17b-16e-instruct"

logging.basicConfig(
    level=getattr(logging, LOG_LEVEL),
    format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
    handlers=[logging.FileHandler(LOG_FILE), logging.StreamHandler()],
)
logger = logging.getLogger(__name__)

BASE_DIR = Path(__file__).resolve().parent
DOCX_PATH = BASE_DIR / "Ligation_Knowledgebase_1.docx"
PDF_PATH = BASE_DIR / "Ligation_Knowledgebase_2.pdf"

# Shunt types to anchor extraction and metadata tagging
SHUNT_TYPES = [
    "Type 1", "Type 2", "Type 3",
    "Type 4 Pelvic", "Type 4 Perforator",
    "Type 5 Pelvic", "Type 5 Perforator",
    "Type 6",
]


# ── Qdrant ────────────────────────────────────────────────────────────────────

def get_qdrant_client() -> QdrantClient:
    if QDRANT_HOST:
        return QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT, api_key=QDRANT_API_KEY)
    os.makedirs(QDRANT_PATH, exist_ok=True)
    return QdrantClient(path=QDRANT_PATH)


# ── Embedding ─────────────────────────────────────────────────────────────────

def get_embedding(text: str) -> np.ndarray:
    try:
        resp = requests.post(
            f"{OLLAMA_BASE_URL}/api/embed",
            json={"model": OLLAMA_EMBEDDING_MODEL, "input": text},
            timeout=30,
        )
        resp.raise_for_status()
        return np.array(resp.json()["embeddings"][0], dtype=np.float32)
    except Exception as e:
        logger.error(f"Embedding error: {e}")
        return np.zeros(EMBEDDING_DIMENSION, dtype=np.float32)


# ── Chunking ──────────────────────────────────────────────────────────────────

def split_into_chunks(text: str, chunk_size=CHUNK_SIZE, overlap=CHUNK_OVERLAP) -> list[str]:
    words = text.split()
    chunks = []
    for i in range(0, len(words), chunk_size - overlap):
        chunk = " ".join(words[i : i + chunk_size])
        if chunk.strip():
            chunks.append(chunk.strip())
    return chunks


def detect_shunt_type(text: str) -> str:
    """Return the shunt type tag if the chunk mentions a specific type."""
    text_lower = text.lower()
    for st in SHUNT_TYPES:
        if st.lower() in text_lower:
            return st
    return "general"


# ── Loaders ───────────────────────────────────────────────────────────────────

def extract_docx_text(docx_path: Path) -> str:
    """Extract all text from a .docx file preserving paragraph structure."""
    logger.info(f"Extracting DOCX: {docx_path.name}")
    doc = docx.Document(str(docx_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # Also pull text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    paragraphs.append(cell_text)

    full_text = "\n".join(paragraphs)
    logger.info(f"  Extracted {len(full_text):,} characters from DOCX")
    return full_text


def extract_pdf_vision(pdf_path: Path) -> str:
    """
    OCR every page of a PDF using Groq vision — handles image-heavy / diagram pages.
    Prompts the model to extract all ligation-related text per shunt type.
    """
    logger.info(f"Extracting PDF (vision OCR): {pdf_path.name}")
    groq_client = Groq(api_key=GROQ_API_KEY)
    doc = fitz.open(str(pdf_path))
    total = len(doc)
    logger.info(f"  {total} pages — sending to Groq vision model")

    all_text: list[str] = []
    for i, page in enumerate(doc):
        logger.info(f"  OCR page {i + 1}/{total}...")
        mat = fitz.Matrix(2, 2)
        pix = page.get_pixmap(matrix=mat)
        img_bytes = pix.tobytes("png")
        img_b64 = base64.standard_b64encode(img_bytes).decode("utf-8")

        try:
            resp = groq_client.chat.completions.create(
                model=GROQ_VISION_MODEL,
                messages=[
                    {
                        "role": "user",
                        "content": [
                            {
                                "type": "image_url",
                                "image_url": {"url": f"data:image/png;base64,{img_b64}"},
                            },
                            {
                                "type": "text",
                                "text": (
                                    "This page is from a medical ligation knowledgebase about venous shunt types "
                                    "(Type 1, Type 2, Type 3, Type 4 Pelvic, Type 4 Perforator, "
                                    "Type 5 Pelvic, Type 5 Perforator, Type 6). "
                                    "Extract ALL text visible on this page — including text inside diagrams, "
                                    "labels, arrows, tables, captions, and figure descriptions. "
                                    "For each shunt type mentioned, preserve its heading and all associated "
                                    "ligation procedure details, anatomical targets, and surgical steps. "
                                    "Transcribe exactly as written. Output only the transcribed text — no commentary."
                                ),
                            },
                        ],
                    }
                ],
                max_tokens=3000,
            )
            page_text = resp.choices[0].message.content or ""
            all_text.append(f"[Page {i + 1}]\n{page_text}")
        except Exception as e:
            logger.warning(f"  OCR failed for page {i + 1}: {e}")
            all_text.append(f"[Page {i + 1}]\n")

        time.sleep(0.5)  # rate-limit buffer

    doc.close()
    full_text = "\n\n".join(all_text)
    logger.info(f"  Extracted {len(full_text):,} characters via vision OCR")
    return full_text


# ── Ingestion pipeline ────────────────────────────────────────────────────────

def ingest():
    logger.info("=" * 70)
    logger.info(f"LIGATION KNOWLEDGEBASE INGESTION -> {QDRANT_LIGATION_COLLECTION}")
    logger.info("=" * 70)

    # 1. Load sources
    sources: list[tuple[str, str]] = []

    if DOCX_PATH.exists():
        docx_text = extract_docx_text(DOCX_PATH)
        if docx_text.strip():
            sources.append((DOCX_PATH.name, docx_text))
        else:
            logger.warning(f"DOCX returned empty text: {DOCX_PATH}")
    else:
        logger.error(f"DOCX not found: {DOCX_PATH}")

    if PDF_PATH.exists():
        pdf_text = extract_pdf_vision(PDF_PATH)
        if pdf_text.strip():
            sources.append((PDF_PATH.name, pdf_text))
        else:
            logger.warning(f"PDF returned empty text: {PDF_PATH}")
    else:
        logger.error(f"PDF not found: {PDF_PATH}")

    if not sources:
        raise RuntimeError("No source documents found — aborting.")

    # 2. Chunk each source with fine-grained size to keep shunt-type sections intact
    all_chunks: list[dict] = []
    for source_name, text in sources:
        # Use smaller chunks (120 words) so each shunt type section stays together
        chunks = split_into_chunks(text, chunk_size=120, overlap=20)
        logger.info(f"  {source_name}: {len(chunks)} chunks")
        for idx, chunk in enumerate(chunks):
            all_chunks.append({
                "text": chunk,
                "source": source_name,
                "chunk_index": idx,
                "shunt_type": detect_shunt_type(chunk),
            })

    logger.info(f"\nTotal chunks: {len(all_chunks)}")

    # Log shunt-type distribution
    from collections import Counter
    dist = Counter(c["shunt_type"] for c in all_chunks)
    for st, count in sorted(dist.items()):
        logger.info(f"  {st}: {count} chunks")

    # 3. Embed
    logger.info(f"\nEmbedding {len(all_chunks)} chunks with {OLLAMA_EMBEDDING_MODEL}...")
    embeddings: list[np.ndarray] = []
    t0 = time.time()
    for i, item in enumerate(all_chunks):
        if (i + 1) % 10 == 0 or (i + 1) == len(all_chunks):
            logger.info(f"  {i + 1}/{len(all_chunks)} embedded...")
        embeddings.append(get_embedding(item["text"]))
        time.sleep(0.05)
    logger.info(f"  Done in {time.time() - t0:.1f}s")

    # 4. Upsert into Qdrant
    logger.info(f"\nUpserting into Qdrant collection '{QDRANT_LIGATION_COLLECTION}'...")
    client = get_qdrant_client()

    existing = [c.name for c in client.get_collections().collections]
    if QDRANT_LIGATION_COLLECTION in existing:
        client.delete_collection(QDRANT_LIGATION_COLLECTION)
        logger.info(f"  Dropped existing collection '{QDRANT_LIGATION_COLLECTION}'")

    client.create_collection(
        collection_name=QDRANT_LIGATION_COLLECTION,
        vectors_config=VectorParams(size=EMBEDDING_DIMENSION, distance=Distance.COSINE),
    )
    logger.info(f"  Created collection (dim={EMBEDDING_DIMENSION}, Cosine)")

    points = [
        PointStruct(
            id=i,
            vector=embeddings[i].tolist(),
            payload={
                "text": all_chunks[i]["text"],
                "source": all_chunks[i]["source"],
                "chunk_index": all_chunks[i]["chunk_index"],
                "shunt_type": all_chunks[i]["shunt_type"],
            },
        )
        for i in range(len(all_chunks))
    ]

    batch_size = 100
    for start in range(0, len(points), batch_size):
        batch = points[start : start + batch_size]
        client.upsert(collection_name=QDRANT_LIGATION_COLLECTION, points=batch)
        logger.info(f"  Upserted {min(start + batch_size, len(points))}/{len(points)} points")

    info = client.get_collection(QDRANT_LIGATION_COLLECTION)
    logger.info("\n" + "=" * 70)
    logger.info("INGESTION COMPLETE")
    logger.info("=" * 70)
    logger.info(f"  Collection : {QDRANT_LIGATION_COLLECTION}")
    logger.info(f"  Points     : {info.points_count}")
    logger.info(f"  Sources    : {[s[0] for s in sources]}")
    logger.info(f"  Storage    : {QDRANT_PATH}")


if __name__ == "__main__":
    try:
        ingest()
    except Exception as e:
        logger.error(f"Ingestion failed: {e}")
        raise
