"""
CHIVA Rules Ablation Test  (4-way)
===================================
Crosses two binary variables to produce 4 prompt configurations per test case:

  Case A — WITH  CHIVA_RULES  +  WITH  RAG   (production baseline)
  Case B — WITHOUT CHIVA_RULES +  WITH  RAG   (RAG-only)
  Case C — WITH  CHIVA_RULES  +  WITHOUT RAG  (rules-only / no retrieval)
  Case D — WITHOUT CHIVA_RULES +  WITHOUT RAG  (pure LLM knowledge, zero context)

RAG uses an enhanced, pattern-focused query that targets CHIVA classification
vocabulary directly — not a generic NL description — so retrieved chunks are
genuinely relevant to the shunt pattern being tested.

Output:  test_chiva_rules_ablation_<timestamp>.docx

Run:
    cd backend
    python test_chiva_rules_ablation.py
"""

#These insights are directly taken from book from pg 573-600.
Detailed_Insights = """
A shunt is defined as a venous pathway carrying not only the physiologically drainage but also the
pathologically deviated blood. Anatomically and hemodynamically, it starts in a refluxing (or escaping)
point (EP) and terminates in the so-called re-entry point (RP).

Closed shunt: it is a refluxing pattern constituting a vicious re-circulation. In this example, an incompetent sapheno
femoral junction represents the EP, pathologically shunting the blood from the deep venous system (N1) to the saphenous
system (N2). A leg perforating vein constitutes the RP, draining the reflux back into the deep venous system.

In type 1 shunt (Figure 4a), the EP allows a reflux from N1 to N2 compartment, eventually interesting
also N3. The main characteristic of this shunt type is the presence of a RP focused on the saphenous trunk.

Being a close shunt, after the systolic push, the blood in a type 1 shunt will reflux again through the EP
during the diastolic phase.

In case of a co-existing N3 refluxing tributary, the refluxing pattern is named Type 1þN3 shunt.1,2
Type 3 shunt (Figure 4b) is mainly characterized by the absence of an RP along the N2 compartment.
The pathological flow is directed from N1 to N2 and then to N3. The difference with the Type 1 (or with the
Type 1-N3shunt) is just the absence of an efficient RP along the saphenous vein.
The blood re-enters into the N1 compartment by an RP focused on the N3, then re-establishing the CS by
escaping once again from the N1-N2 leaking point.

In case of a saphenous trunk RP absence, an interruption of the N1-N2 leaking point
would lead to a not-draining venous system, because of the energy gradient suppression (see therapeutic strategy chapter).

In type 2 shunt (Figure 5a), the leaking point is only from N2 to N3, with a re-entry perforator focused on
the same incompetent tributary and draining toward N1. In this pattern, no recirculation is created, so estab
lishing an ODS. 

Type 2 shunt: it is characterized by an N2-N3 pathological compartment jump, with a RP that is focused along the same
incompetent tributary, so draining into the N1 network. No recirculation occurs and an ODS is formed. (b) Type 2 shunt can also
create a CS, whenever exhibiting an N2-N3 jump, draining back into the same saphenous trunk. In this case, recirculation occurs.

Pg 603-606 : 
Shunt 1 on the left. Escape Points: terminal valve SFJ (red continuous line), SPJ (red dotted line), N2
direct insufficient perforanting vein (blue line and blue circle). N1-N2-N1. Shunt 1 + 2 on the right: one or more
Open Derivative (red dashed) Shunt(s) are added to shunt 1.

Shunt 2a on the left side. Two distinct types of shunt 2a exist: ODS and CS (the only exception in shunts 2a.
Only hydrostatic pressure increases in the ODS. There are no N2 segments with retrograde flow. Shunt 2b in the midle
and Shunt 2c on the right.

Shunt 3. It's a closed shunt. Escape point N1-N2, with an N3 or N4 interposing itself between the escape point
and the re-entry to bypass an N2 continent segment.

Shunt 1: CS. TransMural Pressure increases due to an increase in Hydrostatic Pressure and the addition of
Shunt Pressure It shows an N1>N2>N1 reflux triggered by both Paraná/Squeezing diastole and Valsalva systole. Three
different escape points (N1-N2) are possible: the terminal valve of the SFJ as well as the SPJ, or a direct perforation
of N2 (Fig.12). RP may be multiple along N1 and the more distal one is called Terminal RP.

Shunt 2: it shows an N2>N3 reflux triggered by Paraná/Squeezing diastole but not by Valsalva systole. Three different types of
shunt 2 exist: 2A, 2B, and 2C (Fig.13). Two distinct types of shunt 2a exist: ODS and CS (the only exception
in shunts 2a). Only hydrostatic pressure increases in the ODS. There are no N2 segments with retrograde
flow. Shunt 2b is an Open Derivative Shunt (ODS) with proximal saphenous incontinence and a re-entry path
exclusively through N3, N4L, N4T. Only hydrostatic pressure increases. Shunt 2c is an open derivative shunt
(ODS ) with proximal saphenous incontinence. Saphenous N2 re-entry point, without any interposition of continent
N2 segment, and in addition through N3. Terminal valve continence, possible preterminal incontinence. Only
hydrostatic pressure increases.

Shunt 1+2: CS or Closed Shunt. It's a shunt type 1 plus a shunt type
2 (Fig.12). It is a Shunt 2 connected to a Shunt 1. It shows an N1->N2->N1 + N2->N3 reflux triggered by both Paraná/
Squeezing diastole and Valsalva systole. Notice that when the N1->N2 EP is disconnected, Valsalva doesn't evoke
anymore a N2->N3 reflux. One or more Shunt 2s may connect to the Shunt 1. In a shunt 1+2 the main RP is on
N2, while in the shunt 2+1 it's indirect.

Shunt 3: CS. N1->N2 reflux is evoked by both
Paraná /squeezing diastole and Valsalva systole but N1 doesn't drain directly because intermediate RP is absent. So it drains into an incompetent N3 then directly
or indirectly into a RP. Different types of shunt 3 may exist, with re-entry flow passing through an N4L/N4T to
reach an N3-N1 re-entry point, or an N2. As sometimes the intermediate RP may be not seen whilst present, a re
entry test is possible. It consists in looking at the N2 flow when blocking N3 during a Paraná/squeezing maneuver or
preferably Valsalva, particularly when the RP is searched at the thigh. TMP increases due to the presence of shunt
pressure and increased hydrostatic pressure.

Shunt 4: CS. Its reflux is fed by N1>N3 EP(Perforator or Pelvic Point) then drains into N1 through a
N2 segment: N1->N3->N2->N1.

Shunt 5: CS. Its reflux is fed by N1>N3 EP (Pelvic or Perforator) then drains into N1 through a
N1->N3->N2->N3->N1.


Shunt 1 Flow is : N1->N2->N1
Shunt 1 + 2 Flow is : N1->N2->N1;N2->N3->N1
Shunt 2A (2A can be both ODS and CS (Closed Shunt)) Flow is : N2->N3->N2;N2->N3->N1
Shunt 2B (always ODS) Flow is : N2->N3->N2;N2->N3->N1
Shunt 2C (always ODS) Flow is : N2->N1;N2->N3->N1
Shunt 3 Flow is : N1->N2->N3->N1 or N1->N2->N3->N2-N1
Shunt 4 Flow is : N1->N3->N2->N1 
Shunt 5 Flow is : N1->N3->N2->N3->N1 or N1->N3->N2->N3->N2->N1
Shunt 6 Flow is : 
"""

import os
import sys
import logging
import requests
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import (
    QDRANT_PATH, QDRANT_HOST, QDRANT_PORT, QDRANT_API_KEY,
    QDRANT_SHUNT_COLLECTION, OLLAMA_BASE_URL, OLLAMA_EMBEDDING_MODEL,
    EMBEDDING_DIMENSION, GROQ_API_KEY, GROQ_MODEL,
)
from shunt_llm_classifier import (
    CHIVA_RULES, build_prompt, _summarise_clips, _repair_and_parse,
)

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

NO_RAG_PLACEHOLDER = "No RAG context provided for this run."

# ── Anatomical region boundaries (normalised image coords, top-left = 0,0) ────
ANATOMICAL_REGIONS = {
    "Right": [
        ("thigh",            0.0931, 0.475,  0.0,    0.5497),
        ("calf",             0.105,  0.2947, 0.5497, 1.0),
        ("popliteal region", 0.2827, 0.4386, 0.5497, 1.0),
    ],
    "Left": [
        ("thigh",            0.4985, 0.909,  0.0,    0.5497),
        ("calf",             0.7081, 0.91,   0.5497, 1.0),
        ("popliteal region", 0.588,  0.714,  0.5497, 1.0),
    ],
}


def get_anatomical_region(pos_x: float | None, pos_y: float, leg: str) -> str:
    for name, x_min, x_max, y_min, y_max in ANATOMICAL_REGIONS.get(leg, ANATOMICAL_REGIONS["Right"]):
        if pos_x is not None and x_min <= pos_x <= x_max and y_min <= pos_y <= y_max:
            return name
    return "thigh" if pos_y <= 0.5497 else "lower leg"


# ── Test cases ────────────────────────────────────────────────────────────────
TEST_CASES = [
    {
        "label": "Type 1 — SFJ incompetent, GSV trunk reflux only",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posXRatio": 0.25, "posYRatio": 0.06},
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posXRatio": 0.20, "posYRatio": 0.70},
        ],
        "expected_type": "Type 1",
    },
    {
        "label": "Type 2A — GSV feeds tributary, SFJ competent",
        "leg": "Left",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posXRatio": 0.70, "posYRatio": 0.30},
            {"flow": "RP", "fromType": "N3", "toType": "N2", "posXRatio": 0.72, "posYRatio": 0.65},
        ],
        "expected_type": "Type 2A",
    },
    {
        "label": "Type 2B — perforator entry, tributary reflux, no GSV reflux",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N2", "posXRatio": 0.20, "posYRatio": 0.65},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.35, "posYRatio": 0.82},
        ],
        "expected_type": "Type 2B",
    },
    {
        "label": "Type 2C — perforator entry + tributary reflux + GSV reflux",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N2", "posXRatio": 0.20, "posYRatio": 0.65},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.35, "posYRatio": 0.82},
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posXRatio": 0.25, "posYRatio": 0.35},
        ],
        "expected_type": "Type 2C",
    },
    {
        "label": "Type 3 — SFJ entry + tributary escape, no GSV reflux",
        "leg": "Left",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posXRatio": 0.70, "posYRatio": 0.05},
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posXRatio": 0.65, "posYRatio": 0.30},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.72, "posYRatio": 0.70},
        ],
        "expected_type": "Type 3",
    },
    {
        "label": "Type 1+2 — SFJ + tributary + GSV reflux, eliminationTest=Reflux",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posXRatio": 0.25, "posYRatio": 0.06},
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posXRatio": 0.20, "posYRatio": 0.25},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.20, "posYRatio": 0.70},
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posXRatio": 0.25, "posYRatio": 0.35,
             "eliminationTest": "Reflux"},
        ],
        "expected_type": "Type 1+2",
    },
    {
        "label": "No Shunt — EP only, no RP anywhere",
        "leg": "Left",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N2", "posXRatio": 0.70, "posYRatio": 0.05},
        ],
        "expected_type": "No shunt detected",
    },
]


# ── Infrastructure ────────────────────────────────────────────────────────────

def get_qdrant_client():
    from qdrant_client import QdrantClient
    if QDRANT_HOST:
        return QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT, api_key=QDRANT_API_KEY)
    if not os.path.exists(QDRANT_PATH):
        raise FileNotFoundError(
            f"Qdrant storage not found at {QDRANT_PATH}. "
            "Run ingest_shunt_classification.py first."
        )
    return QdrantClient(path=QDRANT_PATH)


def get_embedding(text: str) -> list[float]:
    try:
        resp = requests.post(
            f"{OLLAMA_BASE_URL}/api/embed",
            json={"model": OLLAMA_EMBEDDING_MODEL, "input": text},
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"][0]
    except Exception as e:
        logger.error(f"Embedding error: {e}")
        return [0.0] * EMBEDDING_DIMENSION


def retrieve_rag_chunks(client, query: str, k: int = 3) -> list[dict]:
    embedding = get_embedding(query)
    response = client.query_points(
        collection_name=QDRANT_SHUNT_COLLECTION,
        query=embedding,
        limit=k,
        with_payload=True,
    )
    return [
        {
            "text":   hit.payload.get("text", ""),
            "score":  hit.score,
            "source": hit.payload.get("source", "unknown"),
        }
        for hit in response.points
    ]


def call_groq(prompt: str) -> tuple[str, dict]:
    from groq import Groq as GroqClient
    client = GroqClient(api_key=GROQ_API_KEY)
    resp = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=512,
    )
    raw = resp.choices[0].message.content or ""
    usage = {
        "prompt_tokens":     resp.usage.prompt_tokens,
        "completion_tokens": resp.usage.completion_tokens,
        "total_tokens":      resp.usage.total_tokens,
    }
    return raw, usage


# ── Three RAG query versions ──────────────────────────────────────────────────

def _clip_to_nl_sentence(clip: dict, leg: str) -> str:
    flow      = clip.get("flow", "?")
    from_node = clip.get("fromType", "?")
    to_node   = clip.get("toType", "?")
    region    = get_anatomical_region(clip.get("posXRatio"), clip.get("posYRatio", 0.5), leg)
    nodes     = [from_node] if from_node == to_node else [from_node, to_node]
    elim      = clip.get("eliminationTest")
    return (
        f"There is {'Escape Point (EP)' if flow == 'EP' else 'Re-entry Point (RP)'} "
        f"around the {region} and the {'antegrade' if flow == 'EP' else 'retrograde'} "
        f"flow is from {' to '.join(nodes)} ({' -> '.join(nodes)})"
        + (f" (elimination test: {elim})" if elim else "") + "."
    )


def _build_flow_chain(clips: list[dict]) -> str:
    """'EP: N1->N2, N2->N3 | RP: N3->N1, N2->N1[elim=Reflux]'"""
    ep_parts, rp_parts = [], []
    for c in clips:
        fn   = c.get("fromType", "?")
        tn   = c.get("toType", "?")
        elim = c.get("eliminationTest", "")
        part = f"{fn}->{tn}" + (f"[elim={elim}]" if elim else "")
        (ep_parts if c.get("flow") == "EP" else rp_parts).append(part)
    segments = []
    if ep_parts:
        segments.append("EP: " + ", ".join(ep_parts))
    if rp_parts:
        segments.append("RP: " + ", ".join(rp_parts))
    return " | ".join(segments)


def build_query_v1(clips: list[dict], leg: str) -> str:
    """V1 — Pattern-focused keyword query."""
    clip_tokens = []
    for c in clips:
        region = get_anatomical_region(c.get("posXRatio"), c.get("posYRatio", 0.5), leg)
        elim   = c.get("eliminationTest", "")
        token  = f"{c.get('flow')} {c.get('fromType')}->{c.get('toType')} {region}"
        if elim:
            token += f" eliminationTest={elim}"
        clip_tokens.append(token)
    vocab_anchor = (
        "CHIVA shunt classification type N1 N2 N3 escape point EP re-entry RP "
        "saphenofemoral junction SFJ perforator tributary GSV SSV saphenous reflux "
        "retrograde antegrade type 1 type 2A type 2B type 2C type 3 closed open shunt"
    )
    return f"CHIVA shunt type classification. Pattern: {' | '.join(clip_tokens)}. {vocab_anchor}"


def build_query_v2(clips: list[dict], leg: str) -> str:
    """V2 — Natural-language description + flow chain suffix (recommended default)."""
    sentences  = " ".join(_clip_to_nl_sentence(c, leg) for c in clips)
    flow_chain = _build_flow_chain(clips)
    return (
        "Given the following information, use the CHIVA shunt classification method "
        f"to classify the shunt type: {sentences} "
        f"Find the shunt type given that the Flow is {flow_chain}."
    )


def build_query_v3(clips: list[dict], leg: str = "Right") -> str:  # leg unused — position-agnostic
    """V3 — Flow chain only (minimal query)."""
    return f"Find the CHIVA shunt type given that the Flow is {_build_flow_chain(clips)}."


# ── 4 prompt builders ─────────────────────────────────────────────────────────

def prompt_A_with_rules_with_rag(clips: list[dict], rag_text: str, leg: str) -> str:
    """Standard production prompt — CHIVA_RULES block + RAG chunks."""
    return build_prompt(clips, rag_text, leg)


def prompt_B_no_rules_with_rag(clips: list[dict], rag_text: str, leg: str) -> str:
    """RAG-only — no embedded CHIVA_RULES, only the retrieved chunks guide the LLM."""
    clips_str = _summarise_clips(clips)
    return f"""=== MEDICAL KNOWLEDGE BASE (retrieved via RAG) ===
{rag_text}

=== ASSESSMENT: {leg} leg ({len(clips)} clips) ===
{clips_str}

=== TASK ===
Using ONLY the medical knowledge retrieved above, classify the venous shunt type
for the {leg} leg following the CHIVA classification framework described in the
retrieved passages.

Output ONLY the JSON below — no other text, no markdown fences.

{{
    "shunt_type": "<Type 1 / Type 2A / Type 2B / Type 2C / Type 3 / Type 1+2 / No shunt detected / Undetermined>",
    "confidence": <0.0-1.0>,
    "reasoning": ["<step 1>", "<step 2>", "..."],
    "ligation": ["<ligation step 1>", "..."],
    "needs_elim_test": <true/false>,
    "ask_diameter": <true/false>,
    "ask_branching": <true/false>,
    "summary": "<1 sentence clinical summary>"
}}"""


def prompt_C_with_rules_no_rag(clips: list[dict], leg: str) -> str:
    """Rules-only — CHIVA_RULES embedded, no RAG context."""
    return build_prompt(clips, NO_RAG_PLACEHOLDER, leg)


def prompt_D_no_rules_no_rag(clips: list[dict], leg: str) -> str:
    """Bare LLM — no CHIVA_RULES, no RAG. Pure model knowledge only."""
    clips_str = _summarise_clips(clips)
    return f"""=== ASSESSMENT: {leg} leg ({len(clips)} clips) ===
{clips_str}

=== TASK ===
You are a CHIVA venous specialist. Using your internal clinical knowledge, classify
the venous shunt type for the {leg} leg.

Key concepts:
  N1 = deep venous system (femoral / popliteal)
  N2 = great or small saphenous vein trunk (GSV / SSV)
  N3 = tributaries / superficial branches
  EP = antegrade (escape point, physiological direction)
  RP = retrograde (re-entry point, pathological reflux)

Output ONLY the JSON below — no other text, no markdown fences.

{{
    "shunt_type": "<Type 1 / Type 2A / Type 2B / Type 2C / Type 3 / Type 1+2 / No shunt detected / Undetermined>",
    "confidence": <0.0-1.0>,
    "reasoning": ["<step 1>", "<step 2>", "..."],
    "ligation": ["<ligation step 1>", "..."],
    "needs_elim_test": <true/false>,
    "ask_diameter": <true/false>,
    "ask_branching": <true/false>,
    "summary": "<1 sentence clinical summary>"
}}"""


def run_llm(prompt: str) -> dict:
    raw, usage = call_groq(prompt)
    result = _repair_and_parse(raw) or {"shunt_type": "PARSE_ERROR", "confidence": 0.0}
    result["_llm_usage"]  = usage
    result["_prompt_len"] = len(prompt)
    result["_raw"]        = raw
    return result


# ── Document helpers ──────────────────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold: bool = False, color: RGBColor | None = None):
    p   = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def _table(doc: Document, rows: list[tuple[str, str]], header: tuple[str, str] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    if header:
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = header
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    for i, (k, v) in enumerate(rows):
        row = table.rows[i + 1].cells
        row[0].text = str(k)
        row[1].text = str(v)
    doc.add_paragraph()



def _correctness(expected: str, result: dict) -> str:
    shunt_type = result.get("shunt_type", "")
    return "CORRECT" if expected.lower() in shunt_type.lower() else "WRONG"


# ── Core test logic ───────────────────────────────────────────────────────────

def run_ablation(client, doc: Document):
    _heading(doc, "3 Query Versions × 4 Prompt Configs Ablation", level=1)
    _para(doc, (
        "Three RAG query versions are tested, each retrieving different chunks.\n"
        "Each version then runs through 4 prompt configurations:\n"
        "  A = +Rules +RAG   B = -Rules +RAG   C = +Rules -RAG   D = -Rules -RAG\n\n"
        "V1 — pattern-focused keyword query\n"
        "V2 — natural-language description + flow chain (recommended default)\n"
        "V3 — flow chain only (minimal)\n\n"
        "C and D don't use RAG so they are run once per case and shared across versions."
    ))
    doc.add_paragraph()

    # scores[version][config]
    scores = {v: {"A": 0, "B": 0, "C": 0, "D": 0} for v in ("v1", "v2", "v3")}
    summary_rows: list[tuple[str, str]] = []

    for case in TEST_CASES:
        leg      = case["leg"]
        clips    = case["clips"]
        expected = case["expected_type"]
        label    = case["label"]

        logger.info(f"\n  [{label}]")

        # Build 3 queries and retrieve their chunks
        queries = {
            "v1": build_query_v1(clips, leg),
            "v2": build_query_v2(clips, leg),
            "v3": build_query_v3(clips, leg),
        }
        chunks = {v: retrieve_rag_chunks(client, q, k=3) for v, q in queries.items()}
        rag_txt = {
            v: "\n\n---\n\n".join(ch["text"][:600] for ch in chunks[v]) if chunks[v] else NO_RAG_PLACEHOLDER
            for v in ("v1", "v2", "v3")
        }
        avg_sim = {
            v: round(sum(ch["score"] for ch in chunks[v]) / len(chunks[v]), 4) if chunks[v] else 0.0
            for v in ("v1", "v2", "v3")
        }

        for v in ("v1", "v2", "v3"):
            logger.info(f"    {v} query: {queries[v][:100]}  sim={avg_sim[v]}")

        # C and D are independent of query version — run once
        res_c = run_llm(prompt_C_with_rules_no_rag(clips, leg))
        res_d = run_llm(prompt_D_no_rules_no_rag(clips, leg))

        # A and B depend on which RAG chunks were retrieved
        ab = {
            v: {
                "A": run_llm(prompt_A_with_rules_with_rag(clips, rag_txt[v], leg)),
                "B": run_llm(prompt_B_no_rules_with_rag(clips, rag_txt[v], leg)),
            }
            for v in ("v1", "v2", "v3")
        }

        results = {
            v: {"A": ab[v]["A"], "B": ab[v]["B"], "C": res_c, "D": res_d}
            for v in ("v1", "v2", "v3")
        }

        ok = {
            v: {k: _correctness(expected, results[v][k]) for k in "ABCD"}
            for v in ("v1", "v2", "v3")
        }

        for v in ("v1", "v2", "v3"):
            for k in "ABCD":
                if ok[v][k] == "CORRECT":
                    scores[v][k] += 1

        logger.info(
            f"    V1 A={ok['v1']['A']} B={ok['v1']['B']}  "
            f"V2 A={ok['v2']['A']} B={ok['v2']['B']}  "
            f"V3 A={ok['v3']['A']} B={ok['v3']['B']}  "
            f"C={ok['v1']['C']} D={ok['v1']['D']}"
        )

        summary_rows.append((label, (
            f"V1 A={ok['v1']['A']} B={ok['v1']['B']} | "
            f"V2 A={ok['v2']['A']} B={ok['v2']['B']} | "
            f"V3 A={ok['v3']['A']} B={ok['v3']['B']} | "
            f"C={ok['v1']['C']} D={ok['v1']['D']}"
        )))

        # ── Per-case doc section ───────────────────────────────────────────────
        _heading(doc, label, level=2)

        _table(doc, [
            ("Expected shunt type", expected),
            ("V1 query", queries["v1"]),
            ("V2 query", queries["v2"]),
            ("V3 query", queries["v3"]),
            ("V1 avg cosine similarity", str(avg_sim["v1"])),
            ("V2 avg cosine similarity", str(avg_sim["v2"])),
            ("V3 avg cosine similarity", str(avg_sim["v3"])),
        ], header=("Field", "Value"))

        # Show retrieved chunks for each version
        for v, vlabel in [("v1", "V1 (pattern-focused)"),
                           ("v2", "V2 (NL + flow chain)"),
                           ("v3", "V3 (flow chain only)")]:
            _para(doc, f"Chunks retrieved by {vlabel}  (avg sim={avg_sim[v]}):", bold=True)
            for i, ch in enumerate(chunks[v]):
                _para(doc, f"  Chunk {i+1}  score={ch['score']:.4f}  source={ch['source']}")
                doc.add_paragraph(ch["text"][:300] + ("…" if len(ch["text"]) > 300 else ""))
            doc.add_paragraph()

        # Results per version (A and B) — shared C/D shown once
        for v, vlabel in [("v1", "V1  (pattern-focused)"),
                           ("v2", "V2  (NL + flow chain)"),
                           ("v3", "V3  (flow chain only)")]:
            _para(doc, f"── {vlabel} ──", bold=True)
            _table(doc, [
                ("A (+Rules +RAG)  type / correct / conf",
                 f"{results[v]['A'].get('shunt_type','?')} / {ok[v]['A']} / {results[v]['A'].get('confidence','?')}"),
                ("B (-Rules +RAG)  type / correct / conf",
                 f"{results[v]['B'].get('shunt_type','?')} / {ok[v]['B']} / {results[v]['B'].get('confidence','?')}"),
                ("A reasoning",
                 " | ".join(results[v]["A"].get("reasoning", [])[:2]) or "(none)"),
                ("B reasoning",
                 " | ".join(results[v]["B"].get("reasoning", [])[:2]) or "(none)"),
            ], header=("Config", "Result"))

        _para(doc, "── No-RAG baseline (shared across V1/V2/V3) ──", bold=True)
        _table(doc, [
            ("C (+Rules -RAG)  type / correct / conf",
             f"{res_c.get('shunt_type','?')} / {ok['v1']['C']} / {res_c.get('confidence','?')}"),
            ("D (-Rules -RAG)  type / correct / conf",
             f"{res_d.get('shunt_type','?')} / {ok['v1']['D']} / {res_d.get('confidence','?')}"),
            ("C reasoning", " | ".join(res_c.get("reasoning", [])[:2]) or "(none)"),
            ("D reasoning", " | ".join(res_d.get("reasoning", [])[:2]) or "(none)"),
        ], header=("Config", "Result"))
        doc.add_paragraph()

    # ── Summary ────────────────────────────────────────────────────────────────
    doc.add_page_break()
    _heading(doc, "Summary", level=1)

    n = len(TEST_CASES)
    _table(doc, [
        ("", "A (+Rules+RAG)  |  B (-Rules+RAG)  |  C (+Rules-RAG)  |  D (-Rules-RAG)"),
        ("V1 (pattern query)",
         f"A={scores['v1']['A']}/{n}  |  B={scores['v1']['B']}/{n}  |  "
         f"C={scores['v1']['C']}/{n}  |  D={scores['v1']['D']}/{n}"),
        ("V2 (NL + flow chain)",
         f"A={scores['v2']['A']}/{n}  |  B={scores['v2']['B']}/{n}  |  "
         f"C={scores['v2']['C']}/{n}  |  D={scores['v2']['D']}/{n}"),
        ("V3 (flow chain only)",
         f"A={scores['v3']['A']}/{n}  |  B={scores['v3']['B']}/{n}  |  "
         f"C={scores['v3']['C']}/{n}  |  D={scores['v3']['D']}/{n}"),
    ], header=("Query version", "Accuracy per config"))

    doc.add_paragraph()
    _table(doc, summary_rows, header=("Test Case", "V1/V2/V3 × A/B  +  C/D"))

    # ── Interpretation guide ───────────────────────────────────────────────────
    doc.add_page_break()
    _heading(doc, "Interpretation Guide", level=1)
    _table(doc, [
        ("A correct, B wrong (any version)",
         "CHIVA_RULES are load-bearing. RAG alone is not enough — keep the embedded rules."),
        ("B correct, A wrong",
         "Rules are overriding correct RAG signal — conflict in CHIVA_RULES. Review."),
        ("A=B correct, C wrong",
         "RAG is the key contributor. Rules alone miss this case."),
        ("A=C correct, B wrong",
         "Rules alone are sufficient. RAG adds no value here."),
        ("V2 A/B beat V1 A/B",
         "Natural-language query retrieves better chunks than pattern-focused. Keep V2."),
        ("V3 A/B beat V2 A/B",
         "Flow chain alone retrieves the most relevant chunks. Consider V3 for efficiency."),
        ("All 4 correct",
         "Pattern is unambiguous — LLM classifies correctly regardless of context."),
        ("All 4 wrong",
         "Critical gap. Enrich the KB and review CHIVA_RULES for this type."),
    ], header=("Observation", "Interpretation / Action"))


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    ts       = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        f"test_chiva_rules_ablation_{ts}.docx",
    )

    logger.info("=" * 60)
    logger.info("CHIVA RULES ABLATION  (4-way)  — starting")
    logger.info("=" * 60)

    doc = Document()
    doc.core_properties.author = "Cygnus Med RAG Test"

    title = doc.add_heading("CHIVA Rules × RAG Ablation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _table(doc, [
        ("Generated",          datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Vector DB",          f"Qdrant  |  collection={QDRANT_SHUNT_COLLECTION}"),
        ("Embedding model",    OLLAMA_EMBEDDING_MODEL),
        ("LLM",                GROQ_MODEL),
        ("RAG k",              "3 chunks per case"),
        ("RAG query style",    "Pattern-focused: clip node types + CHIVA vocabulary anchor"),
        ("Test cases",         str(len(TEST_CASES))),
        ("CHIVA_RULES length", f"{len(CHIVA_RULES):,} chars"),
        ("Configs tested",     "A=+Rules+RAG  B=-Rules+RAG  C=+Rules-RAG  D=-Rules-RAG"),
    ], header=("Config", "Value"))
    doc.add_page_break()

    try:
        logger.info("Connecting to Qdrant...")
        client = get_qdrant_client()
        info   = client.get_collection(QDRANT_SHUNT_COLLECTION)
        logger.info(f"  {info.points_count} points in '{QDRANT_SHUNT_COLLECTION}'")
        _para(doc,
              f"Collection '{QDRANT_SHUNT_COLLECTION}' — {info.points_count} vectors loaded.",
              bold=True)
        doc.add_paragraph()
    except Exception as e:
        logger.error(f"Qdrant init failed: {e}")
        _para(doc, f"ERROR: {e}", bold=True, color=RGBColor(200, 0, 0))
        doc.save(out_path)
        return

    run_ablation(client, doc)

    doc.save(out_path)
    logger.info(f"\nReport saved -> {out_path}")
    print(f"\nDone.  Report: {out_path}")


if __name__ == "__main__":
    main()
