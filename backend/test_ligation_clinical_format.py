"""
Ligation Planning Test — Clinical Format
==========================================
Outputs ligation recommendations matching the clinical format shown in
Shunt Classification Cheatsheet with:
- Shunt Type Assessment Results
- Reasoning (clinical logic)
- Proposed Ligation (Treatment Plan) with specific steps
- Additional Information Required

Run:
    cd backend
    python test_ligation_clinical_format.py
"""

import os
import sys
import json
import requests
import glob
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from qdrant_client import QdrantClient
from groq import Groq as GroqClient
from config import GROQ_API_KEY, GROQ_MODEL

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import RGBColor, Pt
except ImportError:
    print("ERROR: python-docx not installed")
    sys.exit(1)

QDRANT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "qdrant_storage")
QDRANT_COLLECTION = "ligation_knowledgebase_db"
OLLAMA_URL = "http://localhost:11434"
EMBEDDING_MODEL = "llama3.2:1b"
EMBEDDING_DIM = 2048

DATA_FOLDER = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "json samples"))

# CHIVA Ligation Plans (reference)
LIGATION_PLANS = {
    "Type 1": {
        "reasoning": [
            "EP N1->N2 present (SFJ or Hunterian incompetent)",
            "RP N2->N1 present (GSV reflux back to deep)",
            "No EP N2->N3 and no RP at N3",
            "Circular flow pattern: N1->N2->N1"
        ],
        "ligation_steps": [
            "Primary ligation at SFJ (if y <= 0.098) or Hunterian (if y <= 0.353)",
            "If multiple RP N2->N1 present: ligate below each except the most distal"
        ],
        "additional_info": [
            "RP at N2 diameter: Small / Large",
            "Multiple RP N2->N1 sites: Yes / No"
        ]
    },
    "Type 2A": {
        "reasoning": [
            "EP N2->N3 present (saphenous to tributary)",
            "SFJ competent (no EP N1->N2)",
            "RP N3->N2 or RP N3->N1 (tributary reflux)",
            "No N1->N2 involvement"
        ],
        "ligation_steps": [
            "Ligate highest EP at N2->N3 junction",
            "If multiple tributaries: consider calibre, distance to perforator, drainage pattern"
        ],
        "additional_info": [
            "Multiple tributaries at N3: Yes / No",
            "Branching pattern: Calibre / Distance / Drainage"
        ]
    },
    "Type 2B": {
        "reasoning": [
            "EP N2->N2 present (perforator entry, NOT SFJ)",
            "SFJ competent (intact)",
            "RP N3->N1 (from tributary to deep)",
            "No RP N2->N1 (key differentiator from Type 2C)",
            "Open distal shunt (ODS)"
        ],
        "ligation_steps": [
            "Ligate highest EP N2->N2 (perforator entry point)",
            "Preserve SFJ competence",
            "If multiple perforators: ligate based on flow dynamics"
        ],
        "additional_info": [
            "Perforator location: SFJ-Knee / Hunterian / Calf",
            "Multiple perforators: Yes / No"
        ]
    },
    "Type 3": {
        "reasoning": [
            "EP N1->N2 present (SFJ incompetent)",
            "EP N2->N3 also present (secondary tributary entry)",
            "RP only at N3 (no RP N2->N1)",
            "Requires staged approach with follow-up"
        ],
        "ligation_steps": [
            "Stage 1: Ligate EP N2->N3 tributaries at their junctions",
            "Stage 2: Follow-up at 6-12 months to assess SFJ reflux",
            "If N2 reflux develops: then ligate SFJ"
        ],
        "additional_info": [
            "Number of tributaries at N3: Single / Multiple",
            "Follow-up imaging planned: Yes / No"
        ]
    },
    "Type 1+2": {
        "reasoning": [
            "EP N1->N2 present (SFJ or Hunterian incompetent)",
            "EP N2->N3 also present (tributary entry)",
            "RP N3->N2/N1 AND RP N2->N1 (complex reflux pattern)",
            "Elimination test result needed to disambiguate"
        ],
        "ligation_steps": [
            "If small RP N2->N1: Apply CHIVA 2 (ligate EP N2->N3 first, then SFJ)",
            "If large/multiple RP N2->N1: Ligate SFJ + all tributaries simultaneously",
            "Ligate below each RP N2->N1 except the most distal"
        ],
        "additional_info": [
            "RP N2->N1 diameter: Small / Large",
            "Elimination test result: Reflux / No Reflux",
            "Number of RP N2->N1 sites: Single / Multiple"
        ]
    }
}


def embed(text: str) -> list[float]:
    """Embed using llama3.2:1b."""
    try:
        resp = requests.post(
            f"{OLLAMA_URL}/api/embed",
            json={"model": EMBEDDING_MODEL, "input": text},
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"][0]
    except Exception as e:
        print(f"ERROR embedding: {e}")
        return [0.0] * EMBEDDING_DIM


def retrieve_chunks(client: QdrantClient, query: str, k: int = 3) -> list[dict]:
    """Retrieve chunks from ligation KB."""
    try:
        vec = embed(query)
        hits = client.query_points(
            collection_name=QDRANT_COLLECTION,
            query=vec,
            limit=k,
            with_payload=True,
        )
        return [
            {
                "id": h.id,
                "score": h.score,
                "text": h.payload.get("text", ""),
            }
            for h in hits.points
        ]
    except Exception as e:
        print(f"ERROR: {e}")
        raise


def infer_shunt_type(clips: list[dict]) -> tuple[str, dict]:
    """
    Infer shunt type from clips.
    Returns: (shunt_type, ligation_plan_dict)
    """
    flows = {(c.get('flow'), c.get('fromType'), c.get('toType')) for c in clips if c.get('flow')}

    # Simple rule-based classification
    has_ep_n1_n2 = any(f[0] == 'EP' and f[1] == 'N1' and f[2] == 'N2' for f in flows)
    has_ep_n2_n3 = any(f[0] == 'EP' and f[1] == 'N2' and f[2] == 'N3' for f in flows)
    has_ep_n2_n2 = any(f[0] == 'EP' and f[1] == 'N2' and f[2] == 'N2' for f in flows)
    has_rp_n2_n1 = any(f[0] == 'RP' and f[1] == 'N2' and f[2] == 'N1' for f in flows)
    has_rp_n3 = any(f[0] == 'RP' and f[1] == 'N3' for f in flows)

    if has_ep_n1_n2 and not has_ep_n2_n3 and has_rp_n2_n1 and not has_rp_n3:
        return "Type 1", LIGATION_PLANS["Type 1"]
    elif not has_ep_n1_n2 and has_ep_n2_n3 and not has_ep_n2_n2:
        return "Type 2A", LIGATION_PLANS["Type 2A"]
    elif has_ep_n2_n2 and not has_ep_n1_n2 and has_rp_n3 and not has_rp_n2_n1:
        return "Type 2B", LIGATION_PLANS["Type 2B"]
    elif has_ep_n1_n2 and has_ep_n2_n3 and has_rp_n3 and not has_rp_n2_n1:
        return "Type 3", LIGATION_PLANS["Type 3"]
    elif has_ep_n1_n2 and has_ep_n2_n3 and has_rp_n2_n1 and has_rp_n3:
        return "Type 1+2", LIGATION_PLANS["Type 1+2"]
    else:
        return "Undetermined", {"reasoning": ["Cannot classify from clips"], "ligation_steps": [], "additional_info": []}


def add_clinical_format_to_doc(doc: Document, sample_name: str, clips: list[dict],
                                 shunt_type: str, plan: dict, retrieved_chunks: list[dict]):
    """Add clinical format section to Word document."""

    # Title
    doc.add_heading(sample_name, level=1)

    # Retrieved RAG Chunks
    doc.add_heading("Retrieved Knowledge Base Chunks", level=2)
    for i, chunk in enumerate(retrieved_chunks, 1):
        doc.add_paragraph(f"Chunk {i} (ID: {chunk['id']}, score: {chunk['score']:.4f}):", style="List Bullet")
        doc.add_paragraph(chunk['text'], style="List Bullet 2")

    doc.add_paragraph("")  # Spacing

    # Shunt Type Assessment Results
    doc.add_heading("Shunt Type Assessment Results", level=2)
    result_para = doc.add_paragraph()
    result_para.add_run(f"Shunt Type: ").bold = True
    result_para.add_run(shunt_type)

    # Reasoning
    doc.add_heading("Reasoning", level=2)
    for reason in plan.get("reasoning", []):
        doc.add_paragraph(reason, style="List Bullet")

    # Proposed Ligation (Treatment Plan)
    doc.add_heading("Proposed Ligation (Treatment Plan)", level=2)
    for step in plan.get("ligation_steps", []):
        doc.add_paragraph(step, style="List Bullet")

    # Additional Information Required
    doc.add_heading("Additional Information Required", level=2)
    for info in plan.get("additional_info", []):
        doc.add_paragraph(info, style="List Bullet")

    doc.add_paragraph("")  # Spacing


def main():
    print("=" * 80)
    print("LIGATION PLANNING — CLINICAL FORMAT")
    print("=" * 80)

    client = QdrantClient(path=QDRANT_PATH)
    doc = Document()

    # Title
    title = doc.add_heading("Ligation Planning Report — Clinical Format", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "This report follows the clinical format from the Shunt Classification Cheatsheet, "
        "with structured sections for assessment, reasoning, ligation plan, and additional information required."
    )

    doc.add_paragraph("")  # Spacing

    # Test samples
    json_files = sorted(glob.glob(os.path.join(DATA_FOLDER, "*.json")))[:6]

    for json_file in json_files:
        sample_name = os.path.basename(json_file).replace(".json", "")

        try:
            with open(json_file) as f:
                data = json.load(f)

            clips = data.get("clips", [])
            if not clips:
                continue

            print(f"\n{'=' * 80}")
            print(f"SAMPLE: {sample_name}")
            print(f"{'=' * 80}")

            # Infer shunt type and get plan
            shunt_type, plan = infer_shunt_type(clips)
            print(f"Inferred Type: {shunt_type}")
            print(f"Reasoning steps: {len(plan.get('reasoning', []))}")
            print(f"Ligation steps: {len(plan.get('ligation_steps', []))}")

            # Retrieve supporting chunks
            query = f"Ligation planning for {shunt_type} venous shunt"
            chunks = retrieve_chunks(client, query, k=2)

            # Add to document in clinical format
            add_clinical_format_to_doc(doc, sample_name, clips, shunt_type, plan, chunks)

        except Exception as e:
            print(f"ERROR processing {sample_name}: {e}")

    client.close()

    # Save report
    timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"Ligation_Clinical_Format_{timestamp_str}.docx"
    doc.save(filename)
    print(f"\nReport saved to: {filename}")


if __name__ == "__main__":
    main()
