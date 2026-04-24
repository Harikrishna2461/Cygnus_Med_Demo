"""
Comprehensive Ligation Testing Framework
==========================================
Unified script combining:
- K-divergence analysis (finding optimal k for each shunt type)
- Chunk retrieval logging (detailed chunk IDs, scores, text)
- LLM ligation output logging (sites, reasoning, confidence)
- Quality metrics (comparing against baseline CHIVA plans)

Outputs single Word document with all analysis.

Run:
    cd backend
    python test_ligation_comprehensive.py
"""

import os
import sys
import json
import requests
import glob
from datetime import datetime
from collections import defaultdict

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from qdrant_client import QdrantClient
from groq import Groq as GroqClient
from config import GROQ_API_KEY, GROQ_MODEL

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor
except ImportError:
    print("ERROR: python-docx not installed")
    sys.exit(1)

QDRANT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "qdrant_storage")
QDRANT_COLLECTION = "ligation_knowledgebase_db_v2"  # Type-specific chunks
OLLAMA_URL = "http://localhost:11434"
EMBEDDING_MODEL = "nomic-embed-text"  # Matches v2 indexing
EMBEDDING_DIM = 768

DATA_FOLDER = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "json samples"))

# Baseline CHIVA ligation plans for quality comparison
# Includes EP/RP flow notation as expected in output
BASELINE_PLANS = {
    "Type 1": {
        "primary_site": "SFJ",
        "ligation_targets": [
            "high ligation at sfj",
            "ep n1->n2",
            "ligate rp n2->n1",
            "saphenofemoral junction"
        ],
        "approach": "high ligation at SFJ, ligate reflux points",
        "key_points": ["N1->N2 entry", "N2->N1 reflux", "circular flow"]
    },
    "Type 2A": {
        "primary_site": "N2->N3",
        "ligation_targets": [
            "ligate highest ep at n2->n3",
            "ep n2->n3",
            "tributary junction",
            "n2->n3"
        ],
        "approach": "ligate highest EP at tributary junction",
        "key_points": ["SFJ competent", "N2->N3 entry", "N3 reflux"]
    },
    "Type 2B": {
        "primary_site": "Perforator",
        "ligation_targets": [
            "ligate ep n2->n2",
            "ep n2->n2",
            "perforator",
            "selective perforator"
        ],
        "approach": "selective perforator ligation, preserve SFJ",
        "key_points": ["SFJ competent", "N2->N2 entry", "open distal shunt"]
    },
    "Type 3": {
        "primary_site": "Tributary (staged)",
        "ligation_targets": [
            "ligate tributaries at n2->n3",
            "ep n2->n3",
            "staged approach",
            "follow-up sfj"
        ],
        "approach": "staged: ligate tributaries first, then SFJ at follow-up if needed",
        "key_points": ["dual EP", "SFJ + tributary", "staged", "follow-up"]
    },
    "Type 1+2": {
        "primary_site": "SFJ + Tributary",
        "ligation_targets": [
            "ep n1->n2",
            "ep n2->n3",
            "rp n2->n1",
            "chiva 2",
            "simultaneous"
        ],
        "approach": "depends on RP N2->N1 diameter: CHIVA 2 staged or simultaneous ligation",
        "key_points": ["dual entry", "RP diameter", "CHIVA 2", "complex"]
    }
}

# Semantic query pairs for K-divergence testing
# ~30-40% variance: synonyms, reordering, rephrasing but SAME semantic meaning
QUERY_PAIRS = {
    "Type 1": {
        "A": "Type 1 SFJ incompetence with circular reflux pattern N1->N2->N1. Ligation strategy: high tie at saphenofemoral junction. Handle multiple reflux sites.",
        "B": "SFJ incompetency causes circular flow N1->N2->N1. Strategy: high ligation tie at saphenofemoral junction. Manage multiple reflux zones."
    },
    "Type 2A": {
        "A": "Type 2A: tributary enters from saphenous trunk N2->N3 without SFJ involvement. Ligation: highest EP at tributary junction. Consider branching anatomy.",
        "B": "Tributary shunt: GSV trunk entry N2->N3 without SFJ. Strategy: ligate highest EP at tributary junction. Assess branching patterns."
    },
    "Type 2B": {
        "A": "Type 2B perforator entry N2->N2 into saphenous trunk. SFJ remains competent. Open distal shunt with reflux N3->N1. Selective perforator ligation.",
        "B": "Perforator shunt: entry N2->N2 into saphenous trunk. Competent SFJ. Distal reflux N3->N1. Strategy: selective perforator ligation."
    },
    "Type 3": {
        "A": "Type 3: dual incompetence with EP N1->N2 and EP N2->N3. Treatment: staged approach, ligate tributaries first, then follow-up for SFJ.",
        "B": "Dual entry shunt N1->N2 and N2->N3 incompetence. Staged strategy: first ligate tributaries, then assess SFJ at follow-up."
    },
    "Type 1+2": {
        "A": "Type 1+2 complex dual entry: SFJ incompetent (N1->N2) plus tributary involvement (N2->N3). RP N2->N1 diameter directs strategy: CHIVA 2 or simultaneous ligation.",
        "B": "Complex dual entry shunt: SFJ incompetency (N1->N2) with tributary (N2->N3). Strategy depends on RP N2->N1 diameter: CHIVA 2 staged or simultaneous approach."
    }
}


def embed(text: str) -> list:
    """Embed text using nomic-embed-text."""
    try:
        resp = requests.post(
            f"{OLLAMA_URL}/api/embed",
            json={"model": EMBEDDING_MODEL, "input": text},
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"][0]
    except Exception as e:
        print(f"ERROR embedding: {e}")
        return [0.0] * EMBEDDING_DIM


def retrieve_chunks_at_k(client: QdrantClient, query: str, k_max: int = 10) -> dict:
    """
    Retrieve chunks at multiple k values (1 to k_max).
    Returns dict: {k: [(chunk_id, score, text), ...]}
    """
    vec = embed(query)
    results = {}
    try:
        hits = client.query_points(
            collection_name=QDRANT_COLLECTION,
            query=vec,
            limit=k_max,
            with_payload=True,
        )
        for k in range(1, k_max + 1):
            results[k] = [
                (h.id, h.score, h.payload.get("text", "")[:100])
                for h in hits.points[:k]
            ]
    except Exception as e:
        print(f"ERROR: {e}")
        raise
    return results


def find_divergence_point(qa_results: dict, qb_results: dict) -> int:
    """Find k value where retrieved chunk IDs diverge between two queries."""
    for k in range(1, max(qa_results.keys()) + 1):
        qa_ids = {item[0] for item in qa_results[k]}
        qb_ids = {item[0] for item in qb_results[k]}
        if qa_ids != qb_ids:
            return k
    return 0


def calculate_stability(qa_results: dict, qb_results: dict, k: int) -> float:
    """Calculate stability at specific k: Jaccard similarity of chunk IDs."""
    qa_ids = {item[0] for item in qa_results[k]}
    qb_ids = {item[0] for item in qb_results[k]}
    if not qa_ids and not qb_ids:
        return 1.0
    intersection = len(qa_ids & qb_ids)
    union = len(qa_ids | qb_ids)
    return intersection / union if union > 0 else 0.0


def analyze_clip_patterns(clips: list, shunt_type: str) -> dict:
    """Extract clinical decision parameters from clips without hardcoding outcomes."""
    analysis = {
        "primary_eps": [],
        "primary_rps": [],
        "secondary_decision_points": [],
        "ep_positions": [],
        "rp_positions": []
    }

    # Collect all EPs and RPs with positions
    for clip in clips:
        flow = clip.get('flow', '')
        from_type = clip.get('fromType', '')
        to_type = clip.get('toType', '')
        pos_y = clip.get('posYRatio', 0.0)

        if flow == 'EP':
            analysis["primary_eps"].append(f"{from_type}->{to_type}")
            analysis["ep_positions"].append(pos_y)
        elif flow == 'RP':
            analysis["primary_rps"].append(f"{from_type}->{to_type}")
            analysis["rp_positions"].append(pos_y)

    # Identify secondary decision points based on pattern
    unique_eps = set(analysis["primary_eps"])
    unique_rps = set(analysis["primary_rps"])

    # Type 1+2 specific: check RP N2->N1 characteristics for staged vs simultaneous decision
    if shunt_type == "Type 1+2" and "N2->N1" in unique_rps:
        n2_n1_positions = [p for i, p in enumerate(analysis["rp_positions"])
                          if analysis["primary_rps"][i] == "N2->N1"]
        analysis["secondary_decision_points"].append({
            "point": "RP N2->N1 diameter assessment (inferred from presence/count)",
            "clinical_significance": "Determines CHIVA 2 staged vs simultaneous approach",
            "consideration": "Small/single RP N2->N1 suggests staged approach (ligate tributaries first, reassess), large/multiple suggests simultaneous ligation"
        })

    # Type 2A/2B specific: check for multiple tributaries (EP N2->N3 counts)
    ep_n2_n3_count = analysis["primary_eps"].count("N2->N3")
    if shunt_type in ["Type 2A", "Type 2B"] and ep_n2_n3_count > 1:
        analysis["secondary_decision_points"].append({
            "point": f"Multiple tributaries at N3 detected (count: {ep_n2_n3_count})",
            "clinical_significance": "Need to evaluate which tributaries to ligate",
            "consideration": "Consider anatomical factors: calibre differences, distance to perforator, drainage patterns through larger/smaller vessels"
        })

    # Type 3 specific: dual entries (SFJ + tributaries) = staged approach
    if shunt_type == "Type 3" and "N1->N2" in unique_eps and "N2->N3" in unique_eps:
        analysis["secondary_decision_points"].append({
            "point": "Dual entries detected (SFJ + tributaries)",
            "clinical_significance": "Staged approach indicated by CHIVA guidelines",
            "consideration": "Stage 1: ligate tributaries (EP N2->N3). Stage 2 (at 6-12 month follow-up): assess SFJ reflux and ligate if needed"
        })

    # Any multiple RPs suggest multiple reflux zones to assess
    if len(unique_rps) > 1:
        analysis["secondary_decision_points"].append({
            "point": f"Multiple reflux zones detected ({len(unique_rps)} types)",
            "clinical_significance": "Each reflux zone may require specific management",
            "consideration": "Determine which reflux zones need ligation vs which may resolve after primary ligation"
        })

    return analysis


def generate_ligation_plan(clips: list, shunt_type: str) -> dict:
    """Generate ligation plan using LLM with enhanced CHIVA principles and clinical context."""

    # Extract clinical patterns from clips
    clip_analysis = analyze_clip_patterns(clips, shunt_type)

    clips_str = "\n".join([
        f"  Clip {i+1}: {c.get('flow', '?')} {c.get('fromType', '?')}->"
        f"{c.get('toType', '?')} (y={c.get('posYRatio', 0.0):.3f})"
        for i, c in enumerate(clips)
    ])

    # Build clinical context from analysis
    clinical_context = ""
    if clip_analysis["secondary_decision_points"]:
        clinical_context += "\nKey Decision Points Identified from Clips:\n"
        for dp in clip_analysis["secondary_decision_points"]:
            clinical_context += f"  - {dp['point']}\n    Clinical significance: {dp['clinical_significance']}\n    Consideration: {dp['consideration']}\n"

    # CHIVA ligation principles for each type (guidance, not hardcoded outcomes)
    CHIVA_PRINCIPLES = {
        "Type 1": """
CHIVA Principle for Type 1:
- PRIMARY TARGET: Entry point (EP N1->N2) at SFJ or Hunterian junction
- SECONDARY TARGETS: Reflux zones (RP N2->N1) - ligate below each except most distal
- DECISION FACTOR: posYRatio determines location (SFJ if y<=0.098, Hunterian if y<=0.353)
- GOAL: Interrupt circular flow N1->N2->N1 by high ligation at entry
- APPROACH: High ligation to prevent thrombosis and maintain venous return""",

        "Type 2A": """
CHIVA Principle for Type 2A:
- PRIMARY TARGET: Highest escape point (EP N2->N3 at tributary junction)
- SECONDARY DECISION: If multiple tributaries, assess anatomical factors
  * Calibre: Equal vs Unequal vessel sizes
  * Distance to perforator: Affects which tributary drains better
  * Drainage: Presence of drainage through thinner vessel
- GOAL: Selectively interrupt escape route while preserving GSV function
- APPROACH: Ligate at junction level, preserve trunk if diameter is normal""",

        "Type 2B": """
CHIVA Principle for Type 2B:
- PRIMARY TARGET: Perforator entry point (EP N2->N2)
- SECONDARY DECISION: Location along saphenous trunk (SFJ-Knee vs Hunterian vs Calf)
- DECISION FACTOR: posYRatio helps determine location for surgical approach
- GOAL: Selective perforator ligation while preserving SFJ competence
- CONSTRAINT: Do NOT ligate SFJ - it is competent and must remain patent""",

        "Type 3": """
CHIVA Principle for Type 3:
- PRIMARY TARGET (Stage 1): Tributaries at EP N2->N3
- SECONDARY TARGET (Stage 2): SFJ only if reflux develops at follow-up
- DECISION LOGIC: Conservative staged approach - avoid unnecessary SFJ intervention
- TIMING: Initial ligation of tributaries, reassess at 6-12 months
- GOAL: If tributaries alone resolves N2 reflux, avoid SFJ ligation entirely""",

        "Type 1+2": """
CHIVA Principle for Type 1+2:
- PRIMARY DECISION: RP N2->N1 characteristics (inferred from clip count/presence)
  * SMALL/SINGLE RP N2->N1 → CHIVA 2 STAGED APPROACH:
    Stage 1: Ligate EP N2->N3 (tributaries) first
    Stage 2: At follow-up, reassess N2 reflux and ligate SFJ if needed
  * LARGE/MULTIPLE RP N2->N1 → SIMULTANEOUS APPROACH:
    Ligate both EP N1->N2 (SFJ) and EP N2->N3 (tributaries) together
- SECONDARY TARGETS: Ligate below each RP N2->N1 except most distal
- GOAL: Hemodynamic correction with minimal unnecessary intervention""",
    }

    chiva_guidance = CHIVA_PRINCIPLES.get(shunt_type, "Apply standard CHIVA ligation principles")

    prompt = f"""=== LIGATION PLANNING TASK ===

Shunt Type: {shunt_type}

Clips Analysis:
{clips_str}
{clinical_context}

{chiva_guidance}

Based on the CHIVA principles above and the clip patterns provided, generate a specific ligation plan.
Use the secondary decision points to inform your approach selection.
Output ONLY valid JSON — no markdown, no explanation.

For ligation_sites, include the EP/RP flow and clinical strategy:
- Examples: "Ligate highest EP at N2->N3", "High ligation at SFJ (EP N1->N2)",
  "Stage 1: Ligate EP N2->N3", "Ligate EP N2->N2 (perforator)"

{{
    "ligation_sites": ["<EP/RP with flow and clinical action>", "<site 2>"],
    "primary_approach": "<chosen clinical strategy>",
    "reasoning": ["<reason 1: which clips guided this>", "<reason 2>", "<reason 3>"],
    "confidence": 0.85,
    "chiva_alignment": "high"
}}"""

    try:
        client = GroqClient(api_key=GROQ_API_KEY)
        resp = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=512,
        )
        raw = resp.choices[0].message.content or ""

        # Strip markdown code blocks if present
        if raw.startswith("```"):
            raw = raw.split("```")[1]
            if raw.startswith("json"):
                raw = raw[4:]
        raw = raw.strip()

        return json.loads(raw)
    except json.JSONDecodeError as e:
        return {
            "ligation_sites": [],
            "primary_approach": "PARSE_ERROR",
            "reasoning": ["JSON parsing failed"],
            "confidence": 0.0,
            "chiva_alignment": "unknown"
        }
    except Exception as e:
        return {
            "ligation_sites": [],
            "error": str(e),
            "confidence": 0.0,
            "chiva_alignment": "unknown"
        }


def score_ligation_quality(llm_plan: dict, baseline: dict, shunt_type: str) -> dict:
    """
    Compare LLM-generated ligation against baseline CHIVA plan.
    Smart matching that understands node notation and anatomy mapping.

    Metrics:
    - site_match: Does LLM recommend correct anatomical/node ligation site?
    - approach_alignment: Does approach strategy align with type?
    - confidence_calibration: Is stated confidence appropriate?
    """
    # Node-to-anatomy mapping for intelligent matching
    NODE_ANATOMY = {
        "n1": ["deep", "deep vein", "femoral"],
        "n2": ["saphenous", "gsv", "trunk", "saphenofemoral", "sfj", "sapheno"],
        "n3": ["tributary", "tributaries", "branch"],
        "n2->n3": ["tributary junction", "ep n2->n3"],
        "n1->n2": ["sfj", "saphenofemoral junction", "sapheno"],
        "n2->n2": ["perforator"],
    }

    # Type-specific expected node patterns
    EXPECTED_NODES = {
        "Type 1": ["n1->n2", "n2->n1", "sfj", "n2"],  # Should target N1->N2 junction (SFJ)
        "Type 2A": ["n2->n3", "n3"],  # Should target N2->N3 junction
        "Type 2B": ["n2->n2", "perforator", "n2"],  # Should target perforator entry
        "Type 3": ["n2->n3", "n3", "n1->n2"],  # Should target N2->N3 first, then N1->N2
        "Type 1+2": ["n1->n2", "n2->n3"],  # Should target both
    }

    scores = {
        "site_match": 0.0,
        "approach_alignment": 0.0,
        "confidence_calibration": 0.0,
        "overall_quality": 0.0,
        "details": []
    }

    # 1. Site Match: Check if LLM sites match baseline targets OR expected nodes for type
    llm_sites_raw = llm_plan.get("ligation_sites", [])
    llm_sites = set(str(s).lower() for s in llm_sites_raw)
    baseline_targets = set(s.lower() for s in baseline["ligation_targets"])
    expected_nodes = set(EXPECTED_NODES.get(shunt_type, []))

    site_matches = 0
    if llm_sites:
        # Direct match with baseline targets
        direct_matches = len(llm_sites & baseline_targets)

        # Node/anatomy mapping match
        expanded_llm = set()
        for site in llm_sites:
            expanded_llm.add(site)
            # Add anatomy keywords for this node
            for node, anatomy_list in NODE_ANATOMY.items():
                if node in site or site in node:
                    expanded_llm.update(anatomy_list)

        anatomy_matches = len(expanded_llm & baseline_targets)
        node_matches = len(llm_sites & expected_nodes)

        site_matches = max(direct_matches, anatomy_matches, node_matches)
        max_possible = max(len(baseline_targets), len(expected_nodes), 1)
        scores["site_match"] = site_matches / max_possible

        scores["details"].append(
            f"Site matching: LLM=[{', '.join(llm_sites_raw)}] vs "
            f"Baseline={list(baseline_targets)} | Direct:{direct_matches}, "
            f"Anatomy:{anatomy_matches}, Nodes:{node_matches}"
        )
    else:
        scores["details"].append("No ligation sites provided by LLM")

    # 2. Approach Alignment: Check type-specific strategy keywords
    llm_approach = str(llm_plan.get("primary_approach", "")).lower()

    # Type-specific approach keywords to match
    TYPE_KEYWORDS = {
        "Type 1": ["high", "ligation", "sfj", "saphenofemoral", "tie"],
        "Type 2A": ["tributary", "junction", "ligate", "branch"],
        "Type 2B": ["perforator", "selective", "ligation"],
        "Type 3": ["staged", "tributary", "follow", "sfj"],
        "Type 1+2": ["chiva", "dual", "staged", "diameter", "simultaneous"],
    }

    type_keywords = TYPE_KEYWORDS.get(shunt_type, [])
    if type_keywords:
        keyword_matches = sum(1 for kw in type_keywords if kw in llm_approach)
        scores["approach_alignment"] = keyword_matches / len(type_keywords)
        scores["details"].append(
            f"Approach alignment: {keyword_matches}/{len(type_keywords)} type-specific "
            f"keywords matched in '{llm_approach}'"
        )
    else:
        scores["approach_alignment"] = 0.5
        scores["details"].append("No type-specific approach keywords defined")

    # 3. Confidence Calibration: Is confidence appropriate for actual quality?
    confidence = llm_plan.get("confidence", 0.0)
    if isinstance(confidence, (int, float)):
        actual_quality = (scores["site_match"] + scores["approach_alignment"]) / 2

        # Confidence should roughly match actual quality
        confidence_error = abs(confidence - actual_quality)
        scores["confidence_calibration"] = max(0.0, 1.0 - confidence_error)
        scores["details"].append(
            f"Confidence: stated={confidence:.2f}, actual_quality={actual_quality:.2f}, "
            f"calibration={scores['confidence_calibration']:.2f}"
        )

    # 4. Overall quality score
    scores["overall_quality"] = (
        scores["site_match"] * 0.4 +
        scores["approach_alignment"] * 0.4 +
        scores["confidence_calibration"] * 0.2
    )

    return scores


def add_kdivergence_to_doc(doc: Document, shunt_type: str, qa_results: dict, qb_results: dict):
    """Add K-divergence analysis to document."""
    doc.add_heading(f"{shunt_type} — K-Divergence Analysis", level=2)

    divergence_k = find_divergence_point(qa_results, qb_results)
    doc.add_paragraph(f"Divergence Point: k={divergence_k}", style="List Bullet")

    doc.add_heading("Stability at Each K", level=3)
    for k in range(1, 8):
        stability = calculate_stability(qa_results, qb_results, k)
        qa_ids = [item[0] for item in qa_results[k]]
        qb_ids = [item[0] for item in qb_results[k]]

        status = "STABLE" if stability == 1.0 else "DIVERGING"
        doc.add_paragraph(
            f"k={k}: stability={stability:.2f} ({status}) | "
            f"QA_IDs: {qa_ids}, QB_IDs: {qb_ids}",
            style="List Bullet"
        )


def add_retrieval_to_doc(doc: Document, query_label: str, results: dict, k_focus: int = 3):
    """Add chunk retrieval details to document."""
    doc.add_heading(f"Retrieval: {query_label}", level=3)

    chunks = results.get(k_focus, [])
    for i, (chunk_id, score, text) in enumerate(chunks, 1):
        doc.add_paragraph(
            f"Chunk {i}: ID={chunk_id}, Score={score:.4f}",
            style="List Bullet"
        )
        doc.add_paragraph(f"{text}...", style="List Bullet 2")


def add_llm_output_to_doc(doc: Document, llm_plan: dict, quality_scores: dict):
    """Add LLM output and quality metrics to document."""
    doc.add_heading("LLM Ligation Plan", level=3)

    doc.add_paragraph(
        f"Primary Approach: {llm_plan.get('primary_approach', '?')}",
        style="List Bullet"
    )
    doc.add_paragraph(
        f"Ligation Sites: {llm_plan.get('ligation_sites', [])}",
        style="List Bullet"
    )

    confidence = llm_plan.get("confidence", 0.0)
    if isinstance(confidence, (int, float)):
        doc.add_paragraph(f"Confidence: {confidence:.2f}", style="List Bullet")

    reasoning = llm_plan.get("reasoning", [])
    if reasoning:
        doc.add_paragraph("Reasoning:", style="List Bullet")
        for reason in reasoning:
            doc.add_paragraph(reason, style="List Bullet 2")

    # Quality metrics
    doc.add_heading("Quality Assessment", level=3)
    doc.add_paragraph(
        f"Overall Quality Score: {quality_scores['overall_quality']:.2f}/1.0",
        style="List Bullet"
    )
    doc.add_paragraph(
        f"Site Match: {quality_scores['site_match']:.2f}",
        style="List Bullet"
    )
    doc.add_paragraph(
        f"Approach Alignment: {quality_scores['approach_alignment']:.2f}",
        style="List Bullet"
    )
    doc.add_paragraph(
        f"Confidence Calibration: {quality_scores['confidence_calibration']:.2f}",
        style="List Bullet"
    )

    for detail in quality_scores["details"]:
        doc.add_paragraph(detail, style="List Bullet 2")


def main():
    print("=" * 90)
    print("COMPREHENSIVE LIGATION TESTING FRAMEWORK")
    print("=" * 90)

    client = QdrantClient(path=QDRANT_PATH)
    doc = Document()

    # Title
    title = doc.add_heading("Comprehensive Ligation Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This comprehensive test combines K-divergence analysis, chunk retrieval logging, "
        "LLM output logging, and quality metrics for ligation planning. "
        "Tests are run against type-specific knowledge base (ligation_knowledgebase_db_v2) "
        "using semantic query pairs and actual patient data."
    )

    # Section 1: K-Divergence Analysis
    print("\n" + "=" * 90)
    print("SECTION 1: K-DIVERGENCE ANALYSIS")
    print("=" * 90)

    doc.add_heading("Section 1: K-Divergence Analysis", level=1)
    doc.add_paragraph(
        "Finding optimal k for each shunt type by detecting where retrieved chunks "
        "diverge between semantically similar queries."
    )

    kdivergence_results = {}
    for shunt_type, queries in QUERY_PAIRS.items():
        print(f"\n{shunt_type}:")
        qa_results = retrieve_chunks_at_k(client, queries["A"], k_max=10)
        qb_results = retrieve_chunks_at_k(client, queries["B"], k_max=10)
        divergence_k = find_divergence_point(qa_results, qb_results)
        kdivergence_results[shunt_type] = divergence_k
        print(f"  Divergence at k={divergence_k}")

        doc.add_heading(shunt_type, level=2)
        doc.add_paragraph(f"Query A: {queries['A'][:80]}...", style="List Bullet")
        doc.add_paragraph(f"Query B: {queries['B'][:80]}...", style="List Bullet")
        add_kdivergence_to_doc(doc, shunt_type, qa_results, qb_results)

    # Section 2: Detailed Retrieval & LLM Output with Actual Data
    print("\n" + "=" * 90)
    print("SECTION 2: DETAILED RETRIEVAL & LLM OUTPUT ANALYSIS")
    print("=" * 90)

    doc.add_heading("Section 2: Detailed Retrieval & LLM Output Analysis", level=1)
    doc.add_paragraph(
        "Testing with actual patient data from JSON samples. For each sample, "
        "retrieval chunks and LLM-generated ligation plans are logged with quality assessment."
    )

    json_files = sorted(glob.glob(os.path.join(DATA_FOLDER, "*.json")))[:5]
    quality_summary = defaultdict(list)

    for json_file in json_files:
        sample_name = os.path.basename(json_file).replace(".json", "")
        print(f"\n{sample_name}:")

        try:
            with open(json_file) as f:
                data = json.load(f)

            clips = data.get("clips", [])
            if not clips:
                print("  (No clips, skipping)")
                continue

            # Infer shunt type from clips
            flows = {(c.get('flow'), c.get('fromType'), c.get('toType')) for c in clips if c.get('flow')}
            has_ep_n1_n2 = any(f[0] == 'EP' and f[1] == 'N1' and f[2] == 'N2' for f in flows)
            has_ep_n2_n3 = any(f[0] == 'EP' and f[1] == 'N2' and f[2] == 'N3' for f in flows)
            has_ep_n2_n2 = any(f[0] == 'EP' and f[1] == 'N2' and f[2] == 'N2' for f in flows)
            has_rp_n2_n1 = any(f[0] == 'RP' and f[1] == 'N2' and f[2] == 'N1' for f in flows)
            has_rp_n3 = any(f[0] == 'RP' and f[1] == 'N3' for f in flows)

            if has_ep_n1_n2 and not has_ep_n2_n3 and has_rp_n2_n1 and not has_rp_n3:
                inferred_type = "Type 1"
            elif not has_ep_n1_n2 and has_ep_n2_n3 and not has_ep_n2_n2:
                inferred_type = "Type 2A"
            elif has_ep_n2_n2 and not has_ep_n1_n2 and has_rp_n3 and not has_rp_n2_n1:
                inferred_type = "Type 2B"
            elif has_ep_n1_n2 and has_ep_n2_n3 and has_rp_n3 and not has_rp_n2_n1:
                inferred_type = "Type 3"
            elif has_ep_n1_n2 and has_ep_n2_n3 and has_rp_n2_n1 and has_rp_n3:
                inferred_type = "Type 1+2"
            else:
                inferred_type = "Unknown"

            print(f"  Inferred Type: {inferred_type}")

            # Retrieve chunks
            query = QUERY_PAIRS.get(inferred_type, {}).get("A", f"Ligation planning for {inferred_type}")
            qa_results = retrieve_chunks_at_k(client, query, k_max=3)
            chunks = qa_results[3]

            # Generate LLM plan
            llm_plan = generate_ligation_plan(clips, inferred_type)
            print(f"  LLM sites: {llm_plan.get('ligation_sites', [])}")

            # Score quality
            baseline = BASELINE_PLANS.get(inferred_type, {})
            quality_scores = score_ligation_quality(llm_plan, baseline, inferred_type)
            quality_summary[inferred_type].append(quality_scores["overall_quality"])

            print(f"  Quality score: {quality_scores['overall_quality']:.2f}")

            # Add to document
            doc.add_heading(sample_name, level=2)
            doc.add_paragraph(f"Inferred Type: {inferred_type}", style="List Bullet")
            doc.add_paragraph(f"Clips count: {len(clips)}", style="List Bullet")

            doc.add_heading("Retrieved Chunks (k=3)", level=3)
            for i, (chunk_id, score, text) in enumerate(chunks, 1):
                doc.add_paragraph(
                    f"Chunk {i}: ID={chunk_id}, Score={score:.4f}",
                    style="List Bullet"
                )
                doc.add_paragraph(f"{text}...", style="List Bullet 2")

            add_llm_output_to_doc(doc, llm_plan, quality_scores)
            doc.add_paragraph("")  # Spacing

        except Exception as e:
            print(f"  ERROR: {e}")

    client.close()

    # Section 3: Summary & Metrics
    print("\n" + "=" * 90)
    print("SECTION 3: SUMMARY & QUALITY METRICS")
    print("=" * 90)

    doc.add_heading("Section 3: Summary & Quality Metrics", level=1)

    # K-divergence summary
    doc.add_heading("K-Divergence Summary", level=2)
    for shunt_type, divergence_k in kdivergence_results.items():
        if divergence_k == 0:
            msg = "No divergence (highly stable retrieval)"
        else:
            msg = f"Diverges at k={divergence_k} (stable up to k={divergence_k-1})"
        doc.add_paragraph(f"{shunt_type}: {msg}", style="List Bullet")

    # Quality summary
    doc.add_heading("LLM Output Quality Summary", level=2)
    if quality_summary:
        for shunt_type in sorted(quality_summary.keys()):
            scores = quality_summary[shunt_type]
            if scores:
                avg_score = sum(scores) / len(scores)
                doc.add_paragraph(
                    f"{shunt_type}: avg quality={avg_score:.2f} "
                    f"(n={len(scores)} samples)",
                    style="List Bullet"
                )
    else:
        doc.add_paragraph("No quality samples evaluated", style="List Bullet")

    doc.add_heading("Interpretation", level=2)
    doc.add_paragraph(
        "K-Divergence: Lower values indicate less stable retrieval under query variations. "
        "Divergence at k=1 suggests very low stability; k>=4 is more stable.",
        style="List Bullet"
    )
    doc.add_paragraph(
        "Quality Score: Composite of site match (40%), approach alignment (40%), "
        "and confidence calibration (20%). Scores >= 0.7 indicate strong alignment with CHIVA guidelines.",
        style="List Bullet"
    )

    # Save report
    timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"Ligation_Comprehensive_{timestamp_str}.docx"
    doc.save(filename)
    print(f"\n{'=' * 90}")
    print(f"Report saved to: {filename}")
    print(f"{'=' * 90}")


if __name__ == "__main__":
    main()
