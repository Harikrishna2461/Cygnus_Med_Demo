"""
Ligation K-Divergence Test
===========================
Tests retrieval stability under query variations.
Compares chunk differentiation across different ligation query phrasings.

Run:
    cd backend
    python test_ligation_kdivergence.py
"""

import os
import sys
import requests
import json
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from qdrant_client import QdrantClient
from groq import Groq as GroqClient
from config import GROQ_API_KEY, GROQ_MODEL

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)

QDRANT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "qdrant_storage")
QDRANT_COLLECTION = "ligation_knowledgebase_db"
OLLAMA_URL = "http://localhost:11434"
EMBEDDING_MODEL = "llama3.2:1b"
EMBEDDING_DIM = 2048

# Query variations for K-divergence: two phrasings of SAME ligation challenge
# QA (phrasing A): Formal, structured ligation query
# QB (phrasing B): Alternative phrasing with synonyms and different emphasis
LIGATION_QUERY_PAIRS = {
    "Type 1 Ligation": {
        "QA": (
            "Type 1 venous shunt classification. SFJ incompetent entry (EP N1->N2). "
            "Circular flow N1->N2->N1 with saphenous reflux (RP N2->N1). "
            "Ligation decision: Ligate at SFJ junction or Hunterian level? "
            "Multiple reflux points at N2->N1: how to manage?"
        ),
        "QB": (
            "Type 1 closed shunt. Incompetent saphenofemoral junction with antegrade flow "
            "from deep to saphenous (N1->N2). Retrograde pathway N2->N1. "
            "Surgical approach: Which entry point to ligate (SFJ vs Hunterian)? "
            "Handling multiple RP N2->N1 sites during intervention?"
        ),
    },
    "Type 2A Ligation": {
        "QA": (
            "Type 2A venous shunt. SFJ competent. Antegrade flow N2->N3 (saphenous to tributary). "
            "Retrograde N3->N2. No N1->N2 involvement. "
            "Ligation target: Highest EP N2->N3 junction. "
            "Multiple branching tributaries: calibre and distance considerations?"
        ),
        "QB": (
            "Type 2A saphenous-tributary shunt. Intact SFJ. Flow diverted from GSV trunk to tributary (N2->N3) "
            "with reflux back (N3->N2). No deep system entry. "
            "Surgical strategy: Locate and ligate highest escape point (N2->N3). "
            "Branching anatomy management: how to decide which tributaries to ligate?"
        ),
    },
    "Type 2B Ligation": {
        "QA": (
            "Type 2B venous shunt. Perforator entry (EP N2->N2) at saphenous trunk. "
            "SFJ competent. Reflux N3->N1 present. No RP N2->N1. Open distal shunt. "
            "Ligation: Highest EP N2->N2 (perforator entry). "
            "Perforator identification and selective ligation approach?"
        ),
        "QB": (
            "Type 2B perforator-fed shunt. Perforating vein entry into GSV trunk (EP N2->N2). "
            "Saphenofemoral junction intact. Tributary reflux N3->N1 without saphenous reflux. "
            "Intervention focus: Ligate perforator entry point (highest N2->N2). "
            "Accurate perforator localization and targeted ligation technique?"
        ),
    },
}

TEST_CLIPS = {
    "Type 1 Ligation": [
        {"flow": "EP", "fromType": "N1", "toType": "N2", "posYRatio": 0.06},
        {"flow": "RP", "fromType": "N2", "toType": "N1", "posYRatio": 0.70},
    ],
    "Type 2A Ligation": [
        {"flow": "EP", "fromType": "N2", "toType": "N3", "posYRatio": 0.30},
        {"flow": "RP", "fromType": "N3", "toType": "N2", "posYRatio": 0.65},
    ],
    "Type 2B Ligation": [
        {"flow": "EP", "fromType": "N2", "toType": "N2", "posYRatio": 0.65},
        {"flow": "RP", "fromType": "N3", "toType": "N1", "posYRatio": 0.82},
    ],
}


def embed(text: str) -> list[float]:
    """Embed text using nomic-embed-text."""
    try:
        resp = requests.post(
            f"{OLLAMA_URL}/api/embed",
            json={"model": EMBEDDING_MODEL, "input": text},
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"][0]
    except Exception as e:
        print(f"ERROR embedding: {e}")
        return [0.0] * EMBEDDING_DIM


def retrieve_chunks(client: QdrantClient, query: str, k: int = 3) -> list[tuple]:
    """Retrieve top-k chunks from Qdrant. Returns (score, text, point_id)."""
    try:
        vec = embed(query)
        hits = client.query_points(
            collection_name=QDRANT_COLLECTION,
            query=vec,
            limit=k,
            with_payload=True,
        )
        return [(h.score, h.payload.get("text", "")[:200], h.id) for h in hits.points]
    except ValueError as e:
        if "not found" in str(e):
            print(f"\nERROR: Collection '{QDRANT_COLLECTION}' not found")
            raise
        raise


def calculate_jaccard(qa_ids: list, qb_ids: list) -> float:
    """Calculate Jaccard similarity of chunk IDs between two queries."""
    if not qa_ids or not qb_ids:
        return 0.0
    intersection = len(set(qa_ids) & set(qb_ids))
    union = len(set(qa_ids) | set(qb_ids))
    return intersection / union if union > 0 else 0.0


def main():
    print("=" * 80)
    print("LIGATION K-DIVERGENCE TEST")
    print("=" * 80)

    client = QdrantClient(path=QDRANT_PATH)
    results = {}

    for case_name, query_pair in LIGATION_QUERY_PAIRS.items():
        print(f"\n{'=' * 80}")
        print(f"CASE: {case_name}")
        print(f"{'=' * 80}")

        qa_query = query_pair["QA"]
        qb_query = query_pair["QB"]

        print(f"\nQuery A ({len(qa_query)} chars, formal):")
        print(f"  {qa_query[:100]}...")

        print(f"\nQuery B ({len(qb_query)} chars, alternative):")
        print(f"  {qb_query[:100]}...")

        # Retrieve chunks for both queries
        qa_chunks = retrieve_chunks(client, qa_query, k=3)
        qb_chunks = retrieve_chunks(client, qb_query, k=3)

        qa_ids = [chunk[2] for chunk in qa_chunks]
        qb_ids = [chunk[2] for chunk in qb_chunks]

        print(f"\nQuery A retrieved IDs: {qa_ids}")
        print(f"Query B retrieved IDs: {qb_ids}")

        jaccard = calculate_jaccard(qa_ids, qb_ids)
        print(f"\nJaccard Similarity: {jaccard:.2f}")

        results[case_name] = {
            "query_a": qa_query,
            "query_b": qb_query,
            "qa_chunks": qa_chunks,
            "qb_chunks": qb_chunks,
            "qa_ids": qa_ids,
            "qb_ids": qb_ids,
            "jaccard": jaccard,
        }

    client.close()

    # Generate report
    generate_report(results)
    print(f"\n{'=' * 80}")
    print("K-DIVERGENCE TEST COMPLETE")
    print(f"{'=' * 80}")


def generate_report(results: dict):
    """Generate Word document with K-divergence results."""
    doc = Document()
    timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"Ligation_KDivergence_Report_{timestamp_str}.docx"

    # Title
    title = doc.add_heading("Ligation Query K-Divergence Test Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    timestamp = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This K-divergence test evaluates retrieval stability under query variations for ligation planning. "
        "Two phrasings of the same ligation challenge are compared to measure chunk differentiation. "
        "High Jaccard similarity indicates robust retrieval; low similarity suggests query sensitivity."
    )

    # Test Configuration
    doc.add_heading("Test Configuration", level=1)
    doc.add_paragraph(f"Embedding Model: nomic-embed-text (768-dim)")
    doc.add_paragraph(f"Qdrant Collection: {QDRANT_COLLECTION}")
    doc.add_paragraph(f"Test Cases: {len(results)}")

    # Results for each case
    for case_name, data in results.items():
        doc.add_heading(f"Case: {case_name}", level=1)

        doc.add_heading("Query A (Formal Phrasing)", level=2)
        doc.add_paragraph(data["query_a"], style="List Bullet")

        doc.add_heading("Query B (Alternative Phrasing)", level=2)
        doc.add_paragraph(data["query_b"], style="List Bullet")

        doc.add_heading("Retrieval Results", level=2)
        doc.add_paragraph(f"Query A retrieved {len(data['qa_ids'])} chunks (IDs: {data['qa_ids']})", style="List Bullet")
        doc.add_paragraph(f"Query B retrieved {len(data['qb_ids'])} chunks (IDs: {data['qb_ids']})", style="List Bullet")

        doc.add_heading("Similarity Metric", level=2)
        jaccard = data["jaccard"]
        doc.add_paragraph(f"Jaccard Similarity: {jaccard:.2f}", style="List Bullet")
        if jaccard > 0.5:
            doc.add_paragraph("Interpretation: High stability — similar chunks retrieved for both phrasings", style="List Bullet 2")
        elif jaccard > 0.2:
            doc.add_paragraph("Interpretation: Moderate stability — some overlap in retrieved chunks", style="List Bullet 2")
        else:
            doc.add_paragraph("Interpretation: Low stability — different chunks retrieved for alternative phrasing", style="List Bullet 2")

        doc.add_paragraph("")  # Spacing

    # Summary Statistics
    doc.add_heading("Summary Statistics", level=1)
    avg_jaccard = sum(data["jaccard"] for data in results.values()) / len(results) if results else 0
    doc.add_paragraph(f"Average Jaccard Similarity: {avg_jaccard:.2f}")
    doc.add_paragraph(f"Interpretation: Query variations {'show stable retrieval' if avg_jaccard > 0.5 else 'show variable retrieval'}")

    # Save
    doc.save(filename)
    print(f"\nReport saved to: {filename}")


if __name__ == "__main__":
    main()
