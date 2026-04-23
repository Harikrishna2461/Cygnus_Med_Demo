"""
Ligation K-Divergence Test (Proper Implementation)
===================================================
Finds optimal k by detecting where retrieved chunks diverge between
two semantically similar ligation queries.

K-divergence: the k value at which retrieved chunks STOP matching between queries.
- k=1,2,3: chunks match → good stability
- k=4,5: chunks diverge → threshold exceeded

Run:
    cd backend
    python test_ligation_kdivergence_proper.py
"""

import os
import sys
import requests
import json
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from qdrant_client import QdrantClient
from config import GROQ_API_KEY

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed")
    sys.exit(1)

QDRANT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "qdrant_storage")
QDRANT_COLLECTION = "ligation_knowledgebase_db"
OLLAMA_URL = "http://localhost:11434"
EMBEDDING_MODEL = "llama3.2:1b"
EMBEDDING_DIM = 2048

# Semantic query pairs: same ligation context, different phrasing
QUERY_PAIRS = {
    "Type 1 Ligation (SFJ)": {
        "A": "Type 1 venous shunt. SFJ incompetent with EP N1->N2. RP N2->N1. Where to ligate?",
        "B": "Incompetent saphenofemoral junction shunt. Entry at SFJ, circular flow N1->N2->N1. Ligation target?",
    },
    "Type 2A Ligation (Tributary)": {
        "A": "Type 2A shunt: competent SFJ, tributary entry N2->N3, RP N3->N2. Ligate where?",
        "B": "Saphenous-tributary shunt, GSV feeds tributary via N2->N3, reflux back. Primary ligation site?",
    },
    "Type 2B Ligation (Perforator)": {
        "A": "Type 2B shunt: perforator entry N2->N2, competent SFJ, RP N3->N1. Where is EP?",
        "B": "Perforator-fed shunt via N2->N2 entry, open distal with N3->N1 reflux. Ligation approach?",
    },
}


def embed(text: str) -> list[float]:
    """Embed using llama3.2:1b."""
    try:
        resp = requests.post(
            f"{OLLAMA_URL}/api/embed",
            json={"model": EMBEDDING_MODEL, "input": text},
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"][0]
    except Exception as e:
        print(f"ERROR embedding: {e}")
        return [0.0] * EMBEDDING_DIM


def retrieve_chunks_at_k(client: QdrantClient, query: str, k_max: int = 10) -> dict:
    """
    Retrieve chunks at multiple k values (1 to k_max).
    Returns dict: {k: [(chunk_id, score, text), ...]}
    """
    vec = embed(query)
    results = {}

    try:
        hits = client.query_points(
            collection_name=QDRANT_COLLECTION,
            query=vec,
            limit=k_max,
            with_payload=True,
        )

        for k in range(1, k_max + 1):
            results[k] = [
                (h.id, h.score, h.payload.get("text", "")[:100])
                for h in hits.points[:k]
            ]
    except Exception as e:
        print(f"ERROR: {e}")
        raise

    return results


def find_divergence_point(qa_results: dict, qb_results: dict) -> int:
    """
    Find k value where retrieved chunk IDs start to diverge between two queries.
    Returns: k value where divergence occurs (0 if never diverges within range)
    """
    for k in range(1, max(qa_results.keys()) + 1):
        qa_ids = {item[0] for item in qa_results[k]}
        qb_ids = {item[0] for item in qb_results[k]}

        # If sets don't match exactly, chunks diverged at this k
        if qa_ids != qb_ids:
            return k

    return 0  # No divergence found


def calculate_stability(qa_results: dict, qb_results: dict, k: int) -> float:
    """
    Calculate stability at specific k: ratio of matching chunk IDs.
    1.0 = perfect match, 0.0 = no overlap
    """
    qa_ids = {item[0] for item in qa_results[k]}
    qb_ids = {item[0] for item in qb_results[k]}

    if not qa_ids and not qb_ids:
        return 1.0

    intersection = len(qa_ids & qb_ids)
    union = len(qa_ids | qb_ids)
    return intersection / union if union > 0 else 0.0


def main():
    print("=" * 80)
    print("LIGATION K-DIVERGENCE TEST (Finding Optimal K)")
    print("=" * 80)

    client = QdrantClient(path=QDRANT_PATH)
    doc = Document()

    # Title
    title = doc.add_heading("Ligation K-Divergence Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This test identifies the optimal k value for ligation RAG retrieval by finding where "
        "retrieved chunks diverge between semantically similar queries. K-divergence indicates "
        "the threshold beyond which retrieval becomes unstable."
    )

    results_summary = {}

    for case_name, queries in QUERY_PAIRS.items():
        qa_query = queries["A"]
        qb_query = queries["B"]

        print(f"\n{'=' * 80}")
        print(f"CASE: {case_name}")
        print(f"{'=' * 80}")
        print(f"Query A: {qa_query[:80]}...")
        print(f"Query B: {qb_query[:80]}...")

        # Retrieve at multiple k values
        qa_results = retrieve_chunks_at_k(client, qa_query, k_max=10)
        qb_results = retrieve_chunks_at_k(client, qb_query, k_max=10)

        # Find divergence point
        divergence_k = find_divergence_point(qa_results, qb_results)
        print(f"\nDivergence Point: k={divergence_k}")

        # Add to document
        doc.add_heading(case_name, level=1)

        doc.add_heading("Queries", level=2)
        doc.add_paragraph(f"Query A: {qa_query}", style="List Bullet")
        doc.add_paragraph(f"Query B: {qb_query}", style="List Bullet")

        doc.add_heading("K-Divergence Analysis", level=2)

        # Show stability at each k
        for k in range(1, 8):
            stability = calculate_stability(qa_results, qb_results, k)
            qa_ids = [item[0] for item in qa_results[k]]
            qb_ids = [item[0] for item in qb_results[k]]

            status = "STABLE" if stability == 1.0 else "DIVERGING"
            print(f"  k={k}: stability={stability:.2f} {status}")

            doc.add_paragraph(
                f"k={k}: Stability={stability:.2f} "
                f"(QA IDs: {qa_ids}, QB IDs: {qb_ids})",
                style="List Bullet"
            )

        doc.add_heading("Analysis", level=2)
        if divergence_k == 0:
            msg = "No divergence detected within k=1-10. Queries retrieve consistently similar chunks."
            print(f"Result: {msg}")
            doc.add_paragraph(msg)
            results_summary[case_name] = "No divergence (stable)"
        else:
            msg = f"Chunks diverge at k={divergence_k}. Stable retrieval up to k={divergence_k-1}."
            print(f"Result: {msg}")
            doc.add_paragraph(msg)
            results_summary[case_name] = f"k={divergence_k}"

        doc.add_paragraph("")  # Spacing

    client.close()

    # Summary section
    doc.add_heading("Summary", level=1)
    doc.add_paragraph("K-Divergence Results:")
    for case, result in results_summary.items():
        doc.add_paragraph(f"{case}: {result}", style="List Bullet")

    doc.add_paragraph("\nInterpretation:")
    doc.add_paragraph(
        "Lower k-divergence (e.g., k=4) means chunks become unstable quickly under query variations. "
        "Higher k-divergence (k=8+) or no divergence indicates robust retrieval. "
        "Recommended k for ligation queries: use value just before divergence point.",
        style="List Bullet"
    )

    # Save report
    timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"Ligation_KDivergence_Proper_{timestamp_str}.docx"
    doc.save(filename)
    print(f"\nReport saved to: {filename}")


if __name__ == "__main__":
    main()
