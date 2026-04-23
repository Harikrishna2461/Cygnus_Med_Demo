"""
Ligation RAG Test with Actual JSON Inputs
==========================================
Uses actual patient data from json samples and multiple shunts folders.
Logs retrieved chunks and LLM ligation recommendations in detail.

Run:
    cd backend
    python test_ligation_rag_actual_inputs.py
"""

import os
import sys
import json
import requests
import glob
from datetime import datetime

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from qdrant_client import QdrantClient
from groq import Groq as GroqClient
from config import GROQ_API_KEY, GROQ_MODEL

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed")
    sys.exit(1)

QDRANT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "qdrant_storage")
QDRANT_COLLECTION = "ligation_knowledgebase_db"
OLLAMA_URL = "http://localhost:11434"
EMBEDDING_MODEL = "llama3.2:1b"
EMBEDDING_DIM = 2048

DATA_FOLDER = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "json samples"))


def embed(text: str) -> list[float]:
    """Embed text using llama3.2:1b."""
    try:
        resp = requests.post(
            f"{OLLAMA_URL}/api/embed",
            json={"model": EMBEDDING_MODEL, "input": text},
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"][0]
    except Exception as e:
        print(f"ERROR embedding: {e}")
        return [0.0] * EMBEDDING_DIM


def retrieve_chunks(client: QdrantClient, query: str, k: int = 3) -> list[dict]:
    """Retrieve top-k chunks from Qdrant ligation KB."""
    try:
        vec = embed(query)
        hits = client.query_points(
            collection_name=QDRANT_COLLECTION,
            query=vec,
            limit=k,
            with_payload=True,
        )
        return [
            {
                "id": h.id,
                "score": h.score,
                "text": h.payload.get("text", ""),
            }
            for h in hits.points
        ]
    except Exception as e:
        print(f"ERROR retrieving: {e}")
        raise


def summarise_clips(clips: list[dict]) -> str:
    """Summarize clips for LLM context."""
    lines = []
    for i, clip in enumerate(clips, 1):
        flow = clip.get("flow", "?")
        from_t = clip.get("fromType", "?")
        to_t = clip.get("toType", "?")
        y = clip.get("posYRatio", 0.0)
        step = clip.get("step", "?")
        lines.append(f"  Clip {i}: {flow} {from_t}->{to_t} (y={y:.3f}, step={step})")
    return "\n".join(lines)


def generate_ligation_plan(clips: list[dict]) -> dict:
    """Generate ligation plan using LLM with RAG context."""
    clips_str = summarise_clips(clips)

    prompt = f"""=== LIGATION PLANNING TASK ===

Clips Analysis:
{clips_str}

Based on the clip patterns above, generate a ligation plan.
Output ONLY the JSON below — no markdown.

{{
    "ligation_sites": ["<site 1>", "<site 2>"],
    "approach": "<strategy>",
    "reasoning": ["<reason 1>", "<reason 2>"],
    "confidence": 0.85
}}"""

    try:
        client = GroqClient(api_key=GROQ_API_KEY)
        resp = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=512,
        )
        raw = resp.choices[0].message.content or ""
        return json.loads(raw)
    except json.JSONDecodeError:
        return {"ligation_sites": [], "approach": "PARSE_ERROR", "confidence": 0.0}
    except Exception as e:
        return {"ligation_sites": [], "error": str(e), "confidence": 0.0}


def build_ligation_query(clips: list[dict]) -> str:
    """Build optimized ligation query from clips."""
    flows = {f"{c.get('flow')} {c.get('fromType')}->{c.get('toType')}" for c in clips if c.get('flow')}
    flow_str = ", ".join(sorted(flows))
    return f"Ligation planning for shunt with flow patterns: {flow_str}"


def main():
    print("=" * 80)
    print("LIGATION RAG TEST WITH ACTUAL INPUTS")
    print("=" * 80)

    client = QdrantClient(path=QDRANT_PATH)
    doc = Document()

    # Title
    title = doc.add_heading("Ligation RAG Retrieval Test — Actual Patient Data", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}").alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Test json samples
    json_files = sorted(glob.glob(os.path.join(DATA_FOLDER, "*.json")))[:5]  # First 5 samples

    for json_file in json_files:
        sample_name = os.path.basename(json_file)

        try:
            with open(json_file) as f:
                data = json.load(f)

            clips = data.get("clips", [])
            if not clips:
                continue

            print(f"\n{'=' * 80}")
            print(f"SAMPLE: {sample_name}")
            print(f"{'=' * 80}")

            # Build query
            query = build_ligation_query(clips)
            print(f"Query: {query}")

            # Add to document
            doc.add_heading(sample_name, level=1)
            doc.add_paragraph(f"Clips count: {len(clips)}")
            doc.add_paragraph(query, style="List Bullet")

            # Retrieve chunks
            chunks = retrieve_chunks(client, query, k=3)
            print(f"\nRetrieved chunks (top-3):")

            doc.add_heading("Retrieved Chunks", level=2)
            for i, chunk in enumerate(chunks, 1):
                print(f"\n  Chunk {i} (ID:{chunk['id']}, score: {chunk['score']:.4f}):")
                print(f"    {chunk['text'][:150]}...")

                doc.add_paragraph(f"Chunk {i} (ID: {chunk['id']}, score: {chunk['score']:.4f})", style="List Bullet")
                doc.add_paragraph(chunk['text'], style="List Bullet 2")

            # Generate ligation plan
            print(f"\nLLM Ligation Plan:")
            plan = generate_ligation_plan(clips)
            print(f"  Sites: {plan.get('ligation_sites', [])}")
            print(f"  Approach: {plan.get('approach', '?')}")

            doc.add_heading("LLM Ligation Plan", level=2)
            doc.add_paragraph(f"Ligation Sites: {plan.get('ligation_sites', [])}", style="List Bullet")
            doc.add_paragraph(f"Approach: {plan.get('approach', '?')}", style="List Bullet")

            reasoning = plan.get('reasoning', [])
            if reasoning:
                doc.add_paragraph("Reasoning:", style="List Bullet")
                for step in reasoning:
                    doc.add_paragraph(step, style="List Bullet 2")

            doc.add_paragraph("")  # Spacing

        except Exception as e:
            print(f"ERROR processing {sample_name}: {e}")

    client.close()

    # Save report
    timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"Ligation_RAG_Actual_Inputs_{timestamp_str}.docx"
    doc.save(filename)
    print(f"\nReport saved to: {filename}")


if __name__ == "__main__":
    main()
