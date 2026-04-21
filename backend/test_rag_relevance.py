"""
RAG Chunk Relevance Validation Script — Task 1 Enhancement
===========================================================

Tests:
  1. Chunk inspection      — Print every retrieved chunk with its similarity score
  2. Keyword relevance     — Check how many chunks contain CHIVA-relevant keywords
  3. Ablation (with/without RAG) — Compare LLM classification output with and without RAG context
  4. Top-N sweep           — Repeat retrieval at k=1, 3, 5, 7 and record how results change
  5. Query variant test    — Compare the current generic query vs. a more specific query

All findings are written to:  rag_relevance_report.docx  (same folder as this script)

Run:
    cd backend
    python test_rag_relevance.py
"""

import os
import sys
import time
import json
import logging
import datetime
import numpy as np
import requests

# ── path setup ────────────────────────────────────────────────────────────────
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import (
    QDRANT_PATH, QDRANT_HOST, QDRANT_PORT, QDRANT_API_KEY,
    QDRANT_COLLECTION, OLLAMA_BASE_URL, OLLAMA_EMBEDDING_MODEL,
    EMBEDDING_DIMENSION, MAX_RETRIEVAL_RESULTS, GROQ_API_KEY, GROQ_MODEL,
)

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

# ── CHIVA relevance keywords ───────────────────────────────────────────────────
CHIVA_KEYWORDS = [
    "chiva", "shunt", "saphenous", "gsv", "ssv", "sfj", "reflux", "perforator",
    "tributary", "ligation", "n1", "n2", "n3", "ep ", "rp ", "antegrade",
    "retrograde", "type 1", "type 2", "type 3", "venous", "varicose",
    "hunterian", "saphenofemoral", "popliteal", "femoral",
]

# ── Representative test cases (one per shunt type) ────────────────────────────
TEST_CASES = [
    {
        "label": "Type 1 — SFJ incompetent, GSV reflux only",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posYRatio": 0.06},
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posYRatio": 0.25},
        ],
        "expected_type": "Type 1",
    },
    {
        "label": "Type 2A — GSV to tributary, no SFJ entry",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posYRatio": 0.20},
            {"flow": "RP", "fromType": "N3", "toType": "N2", "posYRatio": 0.47},
        ],
        "expected_type": "Type 2A",
    },
    {
        "label": "Type 2B — perforator entry, tributary reflux, no GSV reflux",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N2", "posYRatio": 0.05},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posYRatio": 0.13},
        ],
        "expected_type": "Type 2B",
    },
    {
        "label": "Type 2C — perforator entry, tributary + GSV reflux",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N2", "posYRatio": 0.05},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posYRatio": 0.13},
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posYRatio": 0.21},
        ],
        "expected_type": "Type 2C",
    },
    {
        "label": "Type 3 — SFJ entry + tributary escape, no GSV reflux",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posYRatio": 0.05},
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posYRatio": 0.13},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posYRatio": 0.21},
        ],
        "expected_type": "Type 3",
    },
    {
        "label": "Type 1+2 — SFJ + tributary + GSV reflux, elimination=Reflux",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posYRatio": 0.06},
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posYRatio": 0.13},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posYRatio": 0.21},
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posYRatio": 0.30,
             "eliminationTest": "Reflux"},
        ],
        "expected_type": "Type 1+2",
    },
]

K_VALUES = [1, 3, 5, 7]


# ── Helpers ───────────────────────────────────────────────────────────────────

def get_qdrant_client():
    from qdrant_client import QdrantClient
    if QDRANT_HOST:
        return QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT, api_key=QDRANT_API_KEY)
    if not os.path.exists(QDRANT_PATH):
        raise FileNotFoundError(f"Qdrant storage not found at {QDRANT_PATH}. Run ingest.py first.")
    return QdrantClient(path=QDRANT_PATH)


def get_embedding(text: str) -> list[float]:
    try:
        resp = requests.post(
            f"{OLLAMA_BASE_URL}/api/embed",
            json={"model": OLLAMA_EMBEDDING_MODEL, "input": text},
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"][0]
    except Exception as e:
        logger.error(f"Embedding error: {e}")
        return [0.0] * EMBEDDING_DIMENSION


def retrieve_with_scores(client, query: str, k: int) -> list[dict]:
    """Return chunks with their cosine similarity scores."""
    embedding = get_embedding(query)
    # qdrant-client >= 1.7 replaced .search() with .query_points()
    response = client.query_points(
        collection_name=QDRANT_COLLECTION,
        query=embedding,
        limit=k,
        with_payload=True,
    )
    return [{"text": hit.payload.get("text", ""), "score": hit.score} for hit in response.points]


def build_rag_query(clips: list[dict]) -> str:
    """Same logic as shunt_llm_classifier._retrieve_rag_context."""
    rp_flows = [f"{c['fromType']}→{c['toType']}" for c in clips if c.get("flow") == "RP"]
    ep_flows = [f"{c['fromType']}→{c['toType']}" for c in clips if c.get("flow") == "EP"]
    return (
        f"CHIVA shunt classification venous reflux "
        f"RP flows: {', '.join(rp_flows) or 'none'}. "
        f"EP flows: {', '.join(ep_flows[:4]) or 'none'}. "
        f"saphenofemoral junction GSV tributary ligation treatment"
    )


def build_specific_query(case: dict) -> str:
    """A more targeted query encoding the expected shunt type explicitly."""
    label = case["label"]
    expected = case["expected_type"]
    rp = [f"{c['fromType']}→{c['toType']}" for c in case["clips"] if c["flow"] == "RP"]
    ep = [f"{c['fromType']}→{c['toType']}" for c in case["clips"] if c["flow"] == "EP"]
    return (
        f"Given the following inputs, find how the shunt can be classified"
        #f"CHIVA {expected} venous shunt classification ligation "
        f"EP: {', '.join(ep) or 'none'} "
        f"RP: {', '.join(rp) or 'none'} "
        #f"saphenous perforator tributary reflux treatment"
    )


def score_relevance(chunks: list[dict]) -> dict:
    """Count how many chunks contain CHIVA-relevant keywords."""
    total = len(chunks)
    relevant = 0
    keyword_hits_per_chunk = []
    for ch in chunks:
        text_lower = ch["text"].lower()
        hits = [kw for kw in CHIVA_KEYWORDS if kw in text_lower]
        keyword_hits_per_chunk.append(hits)
        if hits:
            relevant += 1
    avg_score = sum(c["score"] for c in chunks) / total if total else 0.0
    return {
        "total": total,
        "relevant_count": relevant,
        "relevance_pct": round(100 * relevant / total, 1) if total else 0,
        "avg_similarity": round(avg_score, 4),
        "keyword_hits": keyword_hits_per_chunk,
    }


def call_groq_llm(prompt: str) -> tuple[str, dict]:
    """Call Groq LLM, return (raw_text, usage_dict)."""
    from groq import Groq as GroqClient
    client = GroqClient(api_key=GROQ_API_KEY)
    resp = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=512,
    )
    raw = resp.choices[0].message.content or ""
    usage = {
        "prompt_tokens": resp.usage.prompt_tokens,
        "completion_tokens": resp.usage.completion_tokens,
        "total_tokens": resp.usage.total_tokens,
    }
    return raw, usage


def classify_with_llm(clips: list[dict], rag_context: str, leg_label: str) -> dict:
    """Build the prompt and call the LLM; return parsed result."""
    from shunt_llm_classifier import build_prompt, _repair_and_parse

    def _call(prompt, stream=False, return_usage=False):
        raw, usage = call_groq_llm(prompt)
        if return_usage:
            return raw, usage
        return raw

    prompt = build_prompt(clips, rag_context, leg_label)
    raw, usage = _call(prompt, return_usage=True)
    result = _repair_and_parse(raw) or {"shunt_type": "PARSE_ERROR", "confidence": 0.0}
    result["_llm_usage"] = usage
    result["_prompt_len"] = len(prompt)
    return result


# ── Document helpers ──────────────────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold: bool = False, color: RGBColor | None = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def _table_2col(doc: Document, rows: list[tuple[str, str]], header: tuple[str, str] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    if header:
        hdr = table.rows[0].cells
        hdr[0].text = header[0]
        hdr[1].text = header[1]
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    for i, (k, v) in enumerate(rows):
        row = table.rows[i + 1].cells
        row[0].text = str(k)
        row[1].text = str(v)
    doc.add_paragraph()


# ── Test runners ──────────────────────────────────────────────────────────────

def run_chunk_inspection(client, doc: Document) -> list[dict]:
    """Test 1: retrieve top-3 chunks for every test case, show full text + score."""
    _heading(doc, "Test 1 — Chunk Inspection (top-3 per case)", level=1)
    _para(doc,
          "For each representative shunt pattern the RAG query is built exactly as "
          "shunt_llm_classifier._retrieve_rag_context does. The top-3 chunks are shown "
          "with their cosine similarity score so you can judge relevance by eye.",
          bold=False)
    doc.add_paragraph()

    all_results = []

    for case in TEST_CASES:
        query = build_rag_query(case["clips"])
        chunks = retrieve_with_scores(client, query, k=3)
        rel = score_relevance(chunks)

        _heading(doc, case["label"], level=2)
        _table_2col(doc, [
            ("Expected shunt type", case["expected_type"]),
            ("RAG query sent", query),
            ("Chunks retrieved", str(rel["total"])),
            ("Chunks with CHIVA keywords", f"{rel['relevant_count']} / {rel['total']}  ({rel['relevance_pct']}%)"),
            ("Avg cosine similarity", str(rel["avg_similarity"])),
        ], header=("Field", "Value"))

        for j, ch in enumerate(chunks):
            hits = rel["keyword_hits"][j]
            hit_str = ", ".join(hits) if hits else "NONE — not CHIVA-relevant"
            color = RGBColor(0, 128, 0) if hits else RGBColor(200, 0, 0)
            _para(doc, f"Chunk {j+1}  |  score={ch['score']:.4f}  |  keywords: {hit_str}", bold=True, color=color)
            doc.add_paragraph(ch["text"][:800] + ("…" if len(ch["text"]) > 800 else ""))
            doc.add_paragraph()

        all_results.append({"case": case["label"], "chunks": chunks, "relevance": rel})

    return all_results


def run_ablation_test(client, doc: Document) -> list[dict]:
    """Test 2: LLM classification WITH vs WITHOUT RAG context."""
    _heading(doc, "Test 2 — Ablation: LLM With vs Without RAG", level=1)
    _para(doc,
          "Each test case is run twice through the LLM: once with the top-3 RAG chunks "
          "injected (current behaviour) and once with the RAG section replaced by "
          "'No RAG context available.' — the baseline. "
          "We compare shunt_type, confidence, and token usage.",
          bold=False)
    doc.add_paragraph()

    rows_summary = []
    all_results = []

    for case in TEST_CASES:
        query = build_rag_query(case["clips"])
        chunks = retrieve_with_scores(client, query, k=3)
        rag_text = "\n\n---\n\n".join(ch["text"][:600] for ch in chunks) if chunks else ""

        logger.info(f"  Ablation: {case['label']}")

        result_with = classify_with_llm(case["clips"], rag_text, "Left")
        result_without = classify_with_llm(case["clips"], "No RAG context available.", "Left")

        type_with = result_with.get("shunt_type", "?")
        type_without = result_without.get("shunt_type", "?")
        conf_with = result_with.get("confidence", 0.0)
        conf_without = result_without.get("confidence", 0.0)
        expected = case["expected_type"]

        match_with = "CORRECT" if expected.lower() in type_with.lower() else "WRONG"
        match_without = "CORRECT" if expected.lower() in type_without.lower() else "WRONG"

        rows_summary.append((
            case["label"],
            f"WITH={type_with} [{match_with}]  |  WITHOUT={type_without} [{match_without}]",
        ))

        _heading(doc, case["label"], level=2)
        _table_2col(doc, [
            ("Expected", expected),
            ("With RAG — shunt_type", type_with),
            ("With RAG — confidence", str(conf_with)),
            ("With RAG — correct?", match_with),
            ("With RAG — prompt tokens", str(result_with["_llm_usage"].get("prompt_tokens", "?"))),
            ("Without RAG — shunt_type", type_without),
            ("Without RAG — confidence", str(conf_without)),
            ("Without RAG — correct?", match_without),
            ("Without RAG — prompt tokens", str(result_without["_llm_usage"].get("prompt_tokens", "?"))),
            ("Token saving (no RAG)", str(
                result_with["_llm_usage"].get("prompt_tokens", 0) -
                result_without["_llm_usage"].get("prompt_tokens", 0)
            )),
        ], header=("Metric", "Value"))

        all_results.append({
            "case": case["label"],
            "expected": expected,
            "with_rag": result_with,
            "without_rag": result_without,
        })

    _heading(doc, "Ablation Summary", level=2)
    _table_2col(doc, rows_summary, header=("Test Case", "With vs Without RAG"))

    return all_results


def run_top_n_sweep(client, doc: Document) -> list[dict]:
    """Test 3: Retrieve at k=1,3,5,7 for each case and log how scores + relevance change."""
    _heading(doc, "Test 3 — Top-N Sweep (k = 1, 3, 5, 7)", level=1)
    _para(doc,
          "Shows how similarity scores and keyword-relevance percentage change as more "
          "chunks are retrieved. Helps decide the optimal k.",
          bold=False)
    doc.add_paragraph()

    all_results = []

    for case in TEST_CASES:
        query = build_rag_query(case["clips"])
        _heading(doc, case["label"], level=2)

        sweep_rows = []
        case_sweep = []

        for k in K_VALUES:
            chunks = retrieve_with_scores(client, query, k=k)
            rel = score_relevance(chunks)
            scores = [round(c["score"], 4) for c in chunks]
            sweep_rows.append((
                f"k={k}",
                f"avg_score={rel['avg_similarity']}  relevant={rel['relevant_count']}/{k} "
                f"({rel['relevance_pct']}%)  scores={scores}",
            ))
            case_sweep.append({"k": k, "relevance": rel, "chunks": chunks})

        _table_2col(doc, sweep_rows, header=("k", "Results"))
        all_results.append({"case": case["label"], "sweep": case_sweep})

    return all_results


def run_query_variant_test(client, doc: Document) -> list[dict]:
    """Test 4: Compare generic current query vs. type-specific query."""
    _heading(doc, "Test 4 — Query Variant: Generic vs Type-Specific", level=1)
    _para(doc,
          "The current RAG query only lists flow types (N2→N1 etc.). "
          "This test also runs a more specific query that names the expected shunt type "
          "and clinical terms. Compare scores to see if a better query improves chunk relevance.",
          bold=False)
    doc.add_paragraph()

    all_results = []

    for case in TEST_CASES:
        generic_q = build_rag_query(case["clips"])
        specific_q = build_specific_query(case)

        chunks_generic = retrieve_with_scores(client, generic_q, k=3)
        chunks_specific = retrieve_with_scores(client, specific_q, k=3)

        rel_g = score_relevance(chunks_generic)
        rel_s = score_relevance(chunks_specific)

        _heading(doc, case["label"], level=2)
        _table_2col(doc, [
            ("Generic query", generic_q),
            ("Specific query", specific_q),
            ("Generic — avg score", str(rel_g["avg_similarity"])),
            ("Specific — avg score", str(rel_s["avg_similarity"])),
            ("Generic — relevant chunks", f"{rel_g['relevant_count']}/3 ({rel_g['relevance_pct']}%)"),
            ("Specific — relevant chunks", f"{rel_s['relevant_count']}/3 ({rel_s['relevance_pct']}%)"),
            ("Score delta (specific - generic)", str(round(rel_s["avg_similarity"] - rel_g["avg_similarity"], 4))),
        ], header=("Dimension", "Value"))

        all_results.append({
            "case": case["label"],
            "generic": {"query": generic_q, "relevance": rel_g},
            "specific": {"query": specific_q, "relevance": rel_s},
        })

    return all_results


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    out_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "rag_relevance_report.docx")

    logger.info("=" * 60)
    logger.info("RAG RELEVANCE VALIDATION — starting")
    logger.info("=" * 60)

    # ── init doc ──
    doc = Document()
    doc.core_properties.author = "Cygnus Med RAG Test"

    # Title page
    title = doc.add_heading("RAG Chunk Relevance Validation Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_rows = [
        ("Generated", datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Vector DB", f"Qdrant  |  collection={QDRANT_COLLECTION}"),
        ("Embedding model", OLLAMA_EMBEDDING_MODEL),
        ("LLM (ablation)", GROQ_MODEL),
        ("Test cases", str(len(TEST_CASES))),
        ("k values tested", str(K_VALUES)),
        ("CHIVA keywords checked", str(len(CHIVA_KEYWORDS))),
    ]
    _table_2col(doc, meta_rows, header=("Config", "Value"))
    doc.add_page_break()

    # ── connect to Qdrant ──
    try:
        logger.info("Connecting to Qdrant…")
        client = get_qdrant_client()
        info = client.get_collection(QDRANT_COLLECTION)
        logger.info(f"  {info.points_count} points in collection '{QDRANT_COLLECTION}'")
        _para(doc, f"Qdrant collection '{QDRANT_COLLECTION}' — {info.points_count} vectors loaded.", bold=True)
        doc.add_paragraph()
    except Exception as e:
        logger.error(f"Qdrant init failed: {e}")
        _para(doc, f"ERROR: Could not connect to Qdrant — {e}", bold=True,
              color=RGBColor(200, 0, 0))
        doc.save(out_path)
        return

    # ── Test 1: Chunk Inspection ──
    logger.info("\nTest 1 — Chunk Inspection")
    run_chunk_inspection(client, doc)
    doc.add_page_break()

    # ── Test 2: Ablation ──
    logger.info("\nTest 2 — Ablation (with vs without RAG)")
    logger.info("  NOTE: Calls the Groq LLM — may take ~30s per case")
    try:
        run_ablation_test(client, doc)
    except Exception as e:
        logger.warning(f"  Ablation skipped — LLM not available: {e}")
        _para(doc, f"Ablation test skipped (LLM unavailable): {e}",
              color=RGBColor(180, 90, 0))
    doc.add_page_break()

    # ── Test 3: Top-N Sweep ──
    logger.info("\nTest 3 — Top-N Sweep")
    run_top_n_sweep(client, doc)
    doc.add_page_break()

    # ── Test 4: Query Variant ──
    logger.info("\nTest 4 — Query Variant (generic vs specific)")
    run_query_variant_test(client, doc)

    # ── Recommendations page ──
    doc.add_page_break()
    _heading(doc, "Interpretation Guide & Next Steps", level=1)
    recommendations = [
        ("Similarity score < 0.70",
         "Chunk is semantically distant. Consider enriching the knowledge base "
         "with more CHIVA-specific documents or improving the RAG query."),
        ("0.70 – 0.85",
         "Moderate relevance. Chunk may contain some useful context but also "
         "generic venous anatomy text. Acceptable if keyword hits are present."),
        ("> 0.85",
         "Strong semantic match. These chunks are the most reliable for classification support."),
        ("'NONE' keyword hits",
         "Chunk contains no CHIVA vocabulary — it is likely noise. "
         "If most chunks score this way, the knowledge base needs CHIVA documents ingested."),
        ("Ablation: same result with/without RAG",
         "RAG is not contributing unique information. Either the hardcoded CHIVA_RULES "
         "already cover the pattern, or the retrieved chunks are irrelevant."),
        ("Ablation: result differs with RAG and WITH is WRONG",
         "RAG is introducing noise that overrides the correct rule. "
         "Reduce k or filter chunks by a minimum similarity threshold."),
        ("Top-N sweep: relevance drops sharply at k=5 or k=7",
         "The useful signal is concentrated in the top-2 or top-3 chunks. "
         "Use k=3 as the default; only go higher if the 4th+ chunks still score > 0.75."),
        ("Query variant: specific query scores higher",
         "Incorporate the shunt-type signal into the RAG query. "
         "This means doing a quick rule-based pre-classification first, "
         "then using the result to build a richer RAG query."),
    ]
    _table_2col(doc, recommendations, header=("Observation", "What to do"))

    # ── save ──
    doc.save(out_path)
    logger.info(f"\nReport saved → {out_path}")
    print(f"\nDone.  Report: {out_path}")


if __name__ == "__main__":
    main()