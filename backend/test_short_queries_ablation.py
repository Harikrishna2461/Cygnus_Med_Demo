"""
Ablation Test: Short Queries WITH vs WITHOUT CHIVA Rules
=========================================================
Tests whether CHIVA rules in the LLM prompt improve classification accuracy
when using short, focused queries and type-specific retrieved chunks.

Run:
    cd backend
    python test_short_queries_ablation.py

Output: Ablation_Test_Report.docx
"""

import os
import sys
import requests
import json
from datetime import datetime

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, SCRIPT_DIR)
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from qdrant_client import QdrantClient
from groq import Groq as GroqClient
from config import GROQ_API_KEY, GROQ_MODEL

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx not installed. Install with: pip install python-docx")
    sys.exit(1)

QDRANT_PATH = os.path.join(SCRIPT_DIR, "qdrant_storage")
QDRANT_COLLECTION = "shunt_classification_db_v2"
OLLAMA_URL = "http://localhost:11434"
EMBEDDING_MODEL = "nomic-embed-text"
EMBEDDING_DIM = 768

# CHIVA Classification Rules
CHIVA_RULES = """
CHIVA SHUNT CLASSIFICATION RULES:

Type 1: Closed shunt with:
- EP (Escape Point) at SFJ (Saphenofemoral Junction) incompetent: N1->N2
- RP (Re-entry Point) via saphenous trunk reflux: N2->N1
- Flow: N1->N2->N1 (circular pattern)
- SFJ must be incompetent
- No EP at N2->N3 or N3 involvement

Type 2A: Saphenous-tributary shunt with:
- SFJ competent (intact)
- EP at N2->N3 (saphenous to tributary): Flow from GSV/SSV to tributary
- RP at N3->N2 (tributary back to saphenous)
- Flow: N2->N3->N2 (circular at N2-N3 level)
- No N1->N2 involvement
- Can be ODS (Open Distal Shunt) or Closed Shunt

Type 2B: Saphenous-tributary shunt with:
- SFJ competent (intact)
- EP at N2->N2 (perforator entry into saphenous trunk)
- RP at N3->N1 (from tributary to deep system)
- Flow: N2->N3->N1 (linear drainage to deep)
- No direct RP N2->N1
- No EP N1->N2
- Always ODS (Open Distal Shunt)

Type 3: Tributary-to-tributary shunt with:
- Both EP and RP at tributary level (N3->N3)
- No saphenous (N2) or deep (N1) system involvement

Type 4: Complex perforator shunt with:
- EP at N1->N3 (perforator entry from deep to tributary)
- RP via N2 back to N1
- Flow: N1->N3->N2->N1 (closed path through all levels)
- Closed shunt with perforator involvement

Type 5: Extended perforator with:
- EP at N1->N3 (perforator entry)
- Multiple re-entries creating extended cycle
- Flow: N1->N3->N2->N3->N1
- Closed shunt
"""

SHORT_QUERIES_ORIGINAL = {
    "Type 1": (
        "CHIVA shunt. EP N1->N2 SFJ incompetent. RP N2->N1 saphenous trunk reflux. "
        "Flow N1->N2->N1. No EP N2->N3. No RP at N3."
    ),
    "Type 2A": (
        "CHIVA shunt. EP N2->N3 GSV to tributary. SFJ competent. RP N3->N2. "
        "Flow N2->N3->N2. No EP N1->N2."
    ),
    "Type 2B": (
        "CHIVA shunt. EP N2->N2 perforator entry. SFJ competent. RP N3->N1. "
        "Flow N2->N3->N1. No RP N2->N1. No EP N1->N2."
    ),
}

SHORT_QUERIES = {
    "Type 1": (
        "Flow is: N1->N2->N1. There is no EP at N2->N3. There is no RP at N3."
    ),
    "Type 2A": (
        "Flow is: N2->N3->N2. There is no EP at N1->N2."
    ),
    "Type 2B": (
        "Flow is: N2->N3->N1. There is no RP at N2->N1. There is no EP at N1->N2."
    ),
}

TEST_CLIPS = {
    "Type 1": [
        {"flow": "EP", "fromType": "N1", "toType": "N2", "posXRatio": 0.25, "posYRatio": 0.06},
        {"flow": "RP", "fromType": "N2", "toType": "N1", "posXRatio": 0.20, "posYRatio": 0.70},
    ],
    "Type 2A": [
        {"flow": "EP", "fromType": "N2", "toType": "N3", "posXRatio": 0.70, "posYRatio": 0.30},
        {"flow": "RP", "fromType": "N3", "toType": "N2", "posXRatio": 0.72, "posYRatio": 0.65},
    ],
    "Type 2B": [
        {"flow": "EP", "fromType": "N2", "toType": "N2", "posXRatio": 0.20, "posYRatio": 0.65},
        {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.35, "posYRatio": 0.82},
    ],
}


def embed(text: str) -> list[float]:
    """Embed text using nomic-embed-text."""
    try:
        resp = requests.post(
            f"{OLLAMA_URL}/api/embed",
            json={"model": EMBEDDING_MODEL, "input": text},
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"][0]
    except Exception as e:
        print(f"ERROR embedding: {e}")
        return [0.0] * EMBEDDING_DIM


def retrieve_chunks(client: QdrantClient, query: str, k: int = 3) -> list[dict]:
    """Retrieve top-k chunks from Qdrant."""
    try:
        vec = embed(query)
        hits = client.query_points(
            collection_name=QDRANT_COLLECTION,
            query=vec,
            limit=k,
            with_payload=True,
        )
        return [
            {
                "score": h.score,
                "chunk_type": h.payload.get("chunk_type", "unknown"),
                "shunt_types": h.payload.get("shunt_types", []),
                "text": h.payload.get("text", "")[:300],
            }
            for h in hits.points
        ]
    except ValueError as e:
        if "not found" in str(e):
            print(f"\nERROR: Collection '{QDRANT_COLLECTION}' not found at {QDRANT_PATH}")
            print(f"Available collections: {[c.name for c in client.get_collections().collections]}")
            raise
        raise


def summarise_clips(clips: list[dict]) -> str:
    """Summarize clips into readable format."""
    lines = []
    for i, clip in enumerate(clips, 1):
        flow = clip.get("flow", "?")
        from_t = clip.get("fromType", "?")
        to_t = clip.get("toType", "?")
        lines.append(f"  Clip {i}: {flow} {from_t}->{to_t}")
    return "\n".join(lines)


def classify_with_rules(clips: list[dict], short_query: str, leg: str = "Right") -> dict:
    """Classify using LLM WITH CHIVA rules."""
    from shunt_llm_classifier import _summarise_clips, _repair_and_parse

    clips_str = _summarise_clips(clips)

    json_schema = """{
    "shunt_type": "<Type 1 / Type 2A / Type 2B / Type 2C / Type 3 / Type 4 / Type 5 / Type 1+2 / No shunt detected>",
    "confidence": <0.0-1.0>,
    "reasoning": ["<step 1>", "<step 2>"],
    "summary": "<1 sentence>"
}"""

    prompt = f"""=== CHIVA CLASSIFICATION RULES ===
{CHIVA_RULES}

=== ASSESSMENT: {leg} leg ===
Query Context: {short_query}

Clips:
{clips_str}

=== TASK ===
Using the CHIVA classification rules above, classify the venous shunt type.
Output ONLY the JSON below — no other text.

{json_schema}"""

    try:
        client = GroqClient(api_key=GROQ_API_KEY)
        resp = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=512,
        )
        raw = resp.choices[0].message.content or ""
        result = _repair_and_parse(raw) or {"shunt_type": "PARSE_ERROR", "confidence": 0.0}
        return result
    except Exception as e:
        print(f"ERROR calling LLM: {e}")
        return {"shunt_type": "LLM_ERROR", "confidence": 0.0, "error": str(e)}


def classify_without_rules(clips: list[dict], short_query: str, leg: str = "Right") -> dict:
    """Classify using LLM WITHOUT CHIVA rules."""
    from shunt_llm_classifier import _summarise_clips, _repair_and_parse

    clips_str = _summarise_clips(clips)

    json_schema = """{
    "shunt_type": "<Type 1 / Type 2A / Type 2B / Type 2C / Type 3 / Type 4 / Type 5 / Type 1+2 / No shunt detected>",
    "confidence": <0.0-1.0>,
    "reasoning": ["<step 1>", "<step 2>"],
    "summary": "<1 sentence>"
}"""

    prompt = f"""=== QUERY CONTEXT (NO CHIVA RULES) ===
{short_query}

=== ASSESSMENT: {leg} leg ===
{clips_str}

=== TASK ===
Based ONLY on the query context above (no embedded rules, no knowledge base),
classify the venous shunt type for the {leg} leg.
Output ONLY the JSON below — no other text.

{json_schema}"""

    try:
        client = GroqClient(api_key=GROQ_API_KEY)
        resp = client.chat.completions.create(
            model=GROQ_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=512,
        )
        raw = resp.choices[0].message.content or ""
        result = _repair_and_parse(raw) or {"shunt_type": "PARSE_ERROR", "confidence": 0.0}
        return result
    except Exception as e:
        print(f"ERROR calling LLM: {e}")
        return {"shunt_type": "LLM_ERROR", "confidence": 0.0, "error": str(e)}


def main():
    print("=" * 80)
    print("ABLATION TEST: SHORT QUERIES WITH vs WITHOUT CHIVA RULES")
    print("=" * 80)

    client = QdrantClient(path=QDRANT_PATH)

    # Track results for document
    results = {
        "Type 1": {"query": SHORT_QUERIES["Type 1"], "with_rules": None, "without_rules": None, "chunks": None},
        "Type 2A": {"query": SHORT_QUERIES["Type 2A"], "with_rules": None, "without_rules": None, "chunks": None},
        "Type 2B": {"query": SHORT_QUERIES["Type 2B"], "with_rules": None, "without_rules": None, "chunks": None},
    }

    for case_name, query in SHORT_QUERIES.items():
        print(f"\n{'=' * 80}")
        print(f"CASE: {case_name}")
        print(f"{'=' * 80}")
        print(f"\nQuery ({len(query)} chars):\n  {query}")

        # Retrieve chunks
        chunks = retrieve_chunks(client, query, k=3)
        results[case_name]["chunks"] = chunks

        print(f"\nRetrieved chunks (top-3):")
        for i, ch in enumerate(chunks, 1):
            print(f"\n  Chunk {i}:")
            print(f"    Score: {ch['score']:.4f}")
            print(f"    Type: {ch['chunk_type']}")
            print(f"    Shunt types: {ch['shunt_types']}")
            print(f"    Text: {ch['text']}")

        # Classify WITH CHIVA rules
        clips = TEST_CLIPS[case_name]
        result_with = classify_with_rules(clips, query)
        results[case_name]["with_rules"] = result_with

        print(f"\n[WITH CHIVA RULES] LLM classification:")
        print(f"  Shunt type: {result_with.get('shunt_type', '?')}")
        print(f"  Confidence: {result_with.get('confidence', '?'):.2f}")
        reasoning = result_with.get('reasoning', [])
        if reasoning:
            print(f"  Reasoning (steps): {len(reasoning)}")

        # Classify WITHOUT CHIVA rules
        result_without = classify_without_rules(clips, query)
        results[case_name]["without_rules"] = result_without

        print(f"\n[WITHOUT CHIVA RULES] LLM classification:")
        print(f"  Shunt type: {result_without.get('shunt_type', '?')}")
        print(f"  Confidence: {result_without.get('confidence', '?'):.2f}")
        reasoning = result_without.get('reasoning', [])
        if reasoning:
            print(f"  Reasoning (steps): {len(reasoning)}")

    client.close()

    # Generate Word document
    filename = generate_report(results)
    print(f"\n{'=' * 80}")
    print(f"Report saved to: {filename}")
    print(f"{'=' * 80}")


def generate_report(results: dict):
    """Generate a comprehensive Word document with all results."""
    doc = Document()

    # Use timestamp for filename to avoid conflicts
    timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"Ablation_Test_Report_{timestamp_str}.docx"

    # Title
    title = doc.add_heading("Short Query Ablation Test Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Timestamp
    timestamp = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        "This ablation test evaluates the impact of CHIVA classification rules on LLM accuracy "
        "when using short, focused queries for venous shunt classification. "
        "The test compares two scenarios: with explicit CHIVA rules in the LLM prompt, "
        "and without rules (only query context)."
    )

    # Test Configuration
    doc.add_heading("Test Configuration", level=1)
    doc.add_paragraph(f"Embedding Model: nomic-embed-text (768-dim)")
    doc.add_paragraph(f"LLM: {GROQ_MODEL}")
    doc.add_paragraph(f"Qdrant Collection: {QDRANT_COLLECTION}")
    doc.add_paragraph(f"Test Cases: {len(results)} (Type 1, Type 2A, Type 2B)")

    # Results for each case
    for case_name in ["Type 1", "Type 2A", "Type 2B"]:
        case_data = results[case_name]

        doc.add_heading(f"Case: {case_name}", level=1)

        # Query
        doc.add_heading("Query", level=2)
        doc.add_paragraph(case_data["query"], style="List Bullet")

        # Retrieved Chunks
        doc.add_heading("Retrieved Chunks (Top-3)", level=2)
        for i, chunk in enumerate(case_data["chunks"], 1):
            doc.add_paragraph(f"Chunk {i}: {chunk['chunk_type']} (score: {chunk['score']:.4f})", style="List Bullet")
            doc.add_paragraph(f"Shunt types: {chunk['shunt_types']}", style="List Bullet 2")
            doc.add_paragraph(f"Text: {chunk['text']}", style="List Bullet 2")

        # WITH Rules
        doc.add_heading("WITH CHIVA Rules", level=2)
        result_with = case_data["with_rules"]
        doc.add_paragraph(f"Classification: {result_with.get('shunt_type', '?')}", style="List Bullet")
        doc.add_paragraph(f"Confidence: {result_with.get('confidence', '?'):.2f}", style="List Bullet")

        reasoning = result_with.get('reasoning', [])
        if reasoning:
            doc.add_paragraph("Reasoning:", style="List Bullet")
            for step in reasoning:
                doc.add_paragraph(step, style="List Bullet 2")

        # WITHOUT Rules
        doc.add_heading("WITHOUT CHIVA Rules", level=2)
        result_without = case_data["without_rules"]
        doc.add_paragraph(f"Classification: {result_without.get('shunt_type', '?')}", style="List Bullet")
        doc.add_paragraph(f"Confidence: {result_without.get('confidence', '?'):.2f}", style="List Bullet")

        reasoning = result_without.get('reasoning', [])
        if reasoning:
            doc.add_paragraph("Reasoning:", style="List Bullet")
            for step in reasoning:
                doc.add_paragraph(step, style="List Bullet 2")

        # Comparison
        doc.add_heading("Comparison", level=2)
        match_with = result_with.get('shunt_type') == case_name
        match_without = result_without.get('shunt_type') == case_name

        doc.add_paragraph(f"WITH rules: {'CORRECT' if match_with else 'INCORRECT'}", style="List Bullet")
        doc.add_paragraph(f"WITHOUT rules: {'CORRECT' if match_without else 'INCORRECT'}", style="List Bullet")

        doc.add_paragraph("")  # Spacing

    # Summary Statistics
    doc.add_heading("Summary Statistics", level=1)

    correct_with = sum(
        1 for case, data in results.items()
        if data["with_rules"].get('shunt_type') == case
    )
    correct_without = sum(
        1 for case, data in results.items()
        if data["without_rules"].get('shunt_type') == case
    )

    total = len(results)

    doc.add_paragraph(f"Total test cases: {total}")
    doc.add_paragraph(f"WITH CHIVA rules: {correct_with}/{total} correct ({100*correct_with//total}%)", style="List Bullet")
    doc.add_paragraph(f"WITHOUT CHIVA rules: {correct_without}/{total} correct ({100*correct_without//total}%)", style="List Bullet")
    doc.add_paragraph(f"Improvement from rules: {correct_with - correct_without:+d} cases", style="List Bullet")

    # Conclusions
    doc.add_heading("Conclusions", level=1)
    if correct_with > correct_without:
        doc.add_paragraph(
            f"CHIVA rules significantly improve classification accuracy. "
            f"With rules: {correct_with}/{total} cases correct. "
            f"Without rules: {correct_without}/{total} cases correct. "
            f"Adding explicit rules to the LLM prompt increased accuracy by {correct_with - correct_without} case(s)."
        )
    elif correct_with == correct_without:
        doc.add_paragraph(
            "CHIVA rules show no impact on classification accuracy for these test cases. "
            f"Both with and without rules: {correct_with}/{total} cases correct."
        )
    else:
        doc.add_paragraph(
            "CHIVA rules unexpectedly decreased classification accuracy. This may indicate "
            "an issue with rule formulation or LLM prompt structure."
        )

    # Save
    doc.save(filename)
    return filename


if __name__ == "__main__":
    main()
