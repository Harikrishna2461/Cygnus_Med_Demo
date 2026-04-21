"""
Shunt Classification RAG Relevance Validation
==============================================
Tests the shunt_classification_db Qdrant collection.
Focus: shunt TYPE identification only — no ligation content.

Tests:
  1. Chunk Inspection      — top-3 chunks per case with similarity scores
  2. Ablation              — LLM classification WITH vs WITHOUT RAG
  3. Top-N Sweep           — k=1,3,5,7 relevance comparison
  4. Query Variant         — generic vs type-specific query comparison
  5. CHIVA Rules Ablation  — 3 query versions × 4 prompt configs
  6. k-Divergence          — find k where two similar queries retrieve different chunks

Output:  shunt_classification_rag_report.docx  (same folder as this script)

Run:
    cd backend
    python test_shunt_classification_rag.py
"""

import os
import sys
import time
import logging
import datetime
import requests

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import (
    QDRANT_PATH, QDRANT_HOST, QDRANT_PORT, QDRANT_API_KEY,
    QDRANT_SHUNT_COLLECTION, OLLAMA_BASE_URL, OLLAMA_EMBEDDING_MODEL,
    EMBEDDING_DIMENSION, GROQ_API_KEY, GROQ_MODEL,
)

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

# ── Classification-focused keywords (no ligation) ─────────────────────────────
CLASSIFICATION_KEYWORDS = [
    "shunt", "type 1", "type 2", "type 3", "type 2a", "type 2b", "type 2c",
    "chiva", "n1", "n2", "n3", "ep ", "rp ", "escape point", "re-entry",
    "saphenous", "gsv", "ssv", "sfj", "reflux", "retrograde", "antegrade",
    "perforator", "tributary", "saphenofemoral", "closed shunt", "open shunt",
    "hunterian", "femoral", "popliteal", "varicose", "venous insufficiency",
    "n1-n2", "n2-n3", "n1→n2", "n2→n3",
]

# ── Anatomical region boundaries (normalised image coords, top-left = 0,0) ────
# Each region: (x_min, x_max, y_min, y_max) → human-readable label
ANATOMICAL_REGIONS = {
    "Right": [
        ("thigh",            0.0931, 0.475,  0.0,    0.5497),   # SFJ → Knee
        ("calf",             0.105,  0.2947, 0.5497, 1.0),      # Knee → Ankle (medial)
        ("popliteal region", 0.2827, 0.4386, 0.5497, 1.0),      # SPJ → Ankle (posterior)
    ],
    "Left": [
        ("thigh",            0.4985, 0.909,  0.0,    0.5497),   # SFJ → Knee
        ("calf",             0.7081, 0.91,   0.5497, 1.0),      # Knee → Ankle (medial)
        ("popliteal region", 0.588,  0.714,  0.5497, 1.0),      # SPJ → Ankle (posterior)
    ],
}


def get_anatomical_region(pos_x: float | None, pos_y: float, leg: str) -> str:
    """Return the anatomical region name for a clip based on its image position."""
    regions = ANATOMICAL_REGIONS.get(leg, ANATOMICAL_REGIONS["Right"])
    if pos_x is not None:
        for name, x_min, x_max, y_min, y_max in regions:
            if x_min <= pos_x <= x_max and y_min <= pos_y <= y_max:
                return name
    # Fallback: use Y only to decide upper vs lower leg
    return "thigh" if pos_y <= 0.5497 else "lower leg"


# ── Test cases — one per shunt type ───────────────────────────────────────────
# posXRatio + posYRatio are realistic anatomical positions across the full leg.
# Y < 0.5497 → upper leg (thigh); Y ≥ 0.5497 → lower leg (calf / popliteal).
TEST_CASES = [
    {
        # EP at SFJ (top of thigh) → RP refluxes all the way down into the calf.
        "label": "Type 1 — SFJ incompetent, GSV trunk reflux only",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posXRatio": 0.25, "posYRatio": 0.06},   # thigh (SFJ)
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posXRatio": 0.20, "posYRatio": 0.70},   # calf
        ],
        "expected_type": "Type 1",
    },
    {
        # GSV in mid-thigh escapes into a tributary that re-enters in the calf.
        "label": "Type 2A — GSV feeds tributary, SFJ competent",
        "leg": "Left",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posXRatio": 0.70, "posYRatio": 0.30},   # thigh
            {"flow": "RP", "fromType": "N3", "toType": "N2", "posXRatio": 0.72, "posYRatio": 0.65},   # calf
        ],
        "expected_type": "Type 2A",
    },
    {
        # Perforator entry in the calf; tributary reflux continues to ankle.
        "label": "Type 2B — perforator entry, tributary reflux, no GSV reflux",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N2", "posXRatio": 0.20, "posYRatio": 0.65},   # calf (perforator)
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.35, "posYRatio": 0.82},   # popliteal region
        ],
        "expected_type": "Type 2B",
    },
    {
        # Perforator in calf, tributary reflux in lower leg, secondary GSV reflux in thigh.
        "label": "Type 2C — perforator entry + tributary reflux + GSV reflux",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N2", "posXRatio": 0.20, "posYRatio": 0.65},   # calf (perforator)
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.35, "posYRatio": 0.82},   # popliteal region
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posXRatio": 0.25, "posYRatio": 0.35},   # thigh
        ],
        "expected_type": "Type 2C",
    },
    {
        # SFJ entry in thigh, tributary escape in thigh, re-entry in calf.
        "label": "Type 3 — SFJ entry + tributary escape, no GSV reflux",
        "leg": "Left",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posXRatio": 0.70, "posYRatio": 0.05},   # thigh (SFJ)
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posXRatio": 0.65, "posYRatio": 0.30},   # thigh
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.72, "posYRatio": 0.70},   # calf
        ],
        "expected_type": "Type 3",
    },
    {
        # SFJ in thigh, tributary in thigh, dual reflux: N3 in calf + N2 in thigh (elimination=Reflux).
        "label": "Type 1+2 — SFJ + tributary + GSV reflux, eliminationTest=Reflux",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posXRatio": 0.25, "posYRatio": 0.06},   # thigh (SFJ)
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posXRatio": 0.20, "posYRatio": 0.25},   # thigh
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.20, "posYRatio": 0.70},   # calf
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posXRatio": 0.25, "posYRatio": 0.35,
             "eliminationTest": "Reflux"},                                                               # thigh
        ],
        "expected_type": "Type 1+2",
    },
    {
        # Perforator in thigh with no reflux anywhere → no shunt.
        "label": "No Shunt — EP only, no RP anywhere",
        "leg": "Left",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N2", "posXRatio": 0.70, "posYRatio": 0.05},   # thigh
        ],
        "expected_type": "No shunt detected",
    },
]

K_VALUES  = [1, 3, 5, 7]
K_DIVERGE = [1, 2, 3, 4, 5, 6, 7, 8, 10]   # Test 6 sweep


# ── Helpers ───────────────────────────────────────────────────────────────────

def get_qdrant_client():
    from qdrant_client import QdrantClient
    if QDRANT_HOST:
        return QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT, api_key=QDRANT_API_KEY)
    if not os.path.exists(QDRANT_PATH):
        raise FileNotFoundError(f"Qdrant storage not found at {QDRANT_PATH}. Run ingest_shunt_classification.py first.")
    return QdrantClient(path=QDRANT_PATH)


def get_embedding(text: str) -> list[float]:
    try:
        resp = requests.post(
            f"{OLLAMA_BASE_URL}/api/embed",
            json={"model": OLLAMA_EMBEDDING_MODEL, "input": text},
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"][0]
    except Exception as e:
        logger.error(f"Embedding error: {e}")
        return [0.0] * EMBEDDING_DIMENSION


def retrieve_with_scores(client, query: str, k: int) -> list[dict]:
    embedding = get_embedding(query)
    response = client.query_points(
        collection_name=QDRANT_SHUNT_COLLECTION,
        query=embedding,
        limit=k,
        with_payload=True,
    )
    return [
        {
            "text": hit.payload.get("text", ""),
            "score": hit.score,
            "source": hit.payload.get("source", "unknown"),
        }
        for hit in response.points
    ]


def clips_to_raw_str(clips: list[dict]) -> str:
    """Return the clips exactly as received — JSON array."""
    import json
    return json.dumps(clips)


def _clip_to_nl_sentence(clip: dict, leg: str) -> str:
    """Convert a single clip dict into a natural-language sentence."""
    flow      = clip.get("flow", "?")
    from_node = clip.get("fromType", "?")
    to_node   = clip.get("toType", "?")
    pos_x     = clip.get("posXRatio")
    pos_y     = clip.get("posYRatio", 0.5)
    region    = get_anatomical_region(pos_x, pos_y, leg)

    # Build the flow chain description including intermediate nodes when present
    nodes = [from_node]
    # If fromType == toType it's a local perforator — mention only once
    if to_node != from_node:
        nodes.append(to_node)
    chain_arrow = "->".join(nodes)
    chain_words = " to ".join(nodes)

    elim = clip.get("eliminationTest")
    elim_note = f" (elimination test: {elim})" if elim else ""

    flow_label = "Escape Point (EP)" if flow == "EP" else "Re-entry Point (RP)"
    direction  = "antegrade" if flow == "EP" else "retrograde"

    return (
        f"There is {flow_label} around the {region} and the {direction} flow is "
        f"from {chain_words} ({chain_arrow}){elim_note}."
    )


def _build_flow_chain(clips: list[dict]) -> str:
    """Compact EP/RP flow notation: 'EP: N1->N2, N2->N3 | RP: N3->N1, N2->N1[elim=Reflux]'"""
    ep_parts, rp_parts = [], []
    for c in clips:
        fn   = c.get("fromType", "?")
        tn   = c.get("toType", "?")
        elim = c.get("eliminationTest", "")
        part = f"{fn}->{tn}" + (f"[elim={elim}]" if elim else "")
        (ep_parts if c.get("flow") == "EP" else rp_parts).append(part)
    segments = []
    if ep_parts:
        segments.append("EP: " + ", ".join(ep_parts))
    if rp_parts:
        segments.append("RP: " + ", ".join(rp_parts))
    return " | ".join(segments)


def build_query_v1(clips: list[dict], leg: str) -> str:
    """V1 — Pattern-focused keyword query (good for keyword/BM25 retrieval)."""
    clip_tokens = []
    for c in clips:
        region = get_anatomical_region(c.get("posXRatio"), c.get("posYRatio", 0.5), leg)
        elim   = c.get("eliminationTest", "")
        token  = f"{c.get('flow')} {c.get('fromType')}->{c.get('toType')} {region}"
        if elim:
            token += f" eliminationTest={elim}"
        clip_tokens.append(token)
    vocab_anchor = (
        "CHIVA shunt classification type N1 N2 N3 escape point EP re-entry RP "
        "saphenofemoral junction SFJ perforator tributary GSV SSV saphenous reflux "
        "retrograde antegrade type 1 type 2A type 2B type 2C type 3 closed open shunt"
    )
    return f"CHIVA shunt type classification. Pattern: {' | '.join(clip_tokens)}. {vocab_anchor}"


def build_query_v2(clips: list[dict], leg: str) -> str:
    """V2 — Natural-language description + flow chain suffix (recommended default)."""
    sentences = " ".join(_clip_to_nl_sentence(c, leg) for c in clips)
    flow_chain = _build_flow_chain(clips)
    return (
        "Given the following information, use the CHIVA shunt classification method "
        f"to classify the shunt type: {sentences} "
        f"Find the shunt type given that the Flow is {flow_chain}."
    )


def build_query_v3(clips: list[dict], leg: str = "Right") -> str:  # leg unused — V3 is position-agnostic
    """V3 — Flow chain only (minimal query, tests if node pattern alone is enough)."""
    return f"Find the CHIVA shunt type given that the Flow is {_build_flow_chain(clips)}."


def build_classification_query(clips: list[dict], leg: str = "Right") -> str:
    """Default RAG query used across Tests 1-4 — V2 natural language + flow chain."""
    return build_query_v2(clips, leg)


def score_relevance(chunks: list[dict]) -> dict:
    total = len(chunks)
    relevant = 0
    keyword_hits_per_chunk = []
    for ch in chunks:
        text_lower = ch["text"].lower()
        hits = [kw for kw in CLASSIFICATION_KEYWORDS if kw in text_lower]
        keyword_hits_per_chunk.append(hits)
        if hits:
            relevant += 1
    avg_score = sum(c["score"] for c in chunks) / total if total else 0.0
    return {
        "total": total,
        "relevant_count": relevant,
        "relevance_pct": round(100 * relevant / total, 1) if total else 0,
        "avg_similarity": round(avg_score, 4),
        "keyword_hits": keyword_hits_per_chunk,
    }


def call_groq_llm(prompt: str) -> tuple[str, dict]:
    from groq import Groq as GroqClient
    client = GroqClient(api_key=GROQ_API_KEY)
    resp = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=512,
    )
    raw = resp.choices[0].message.content or ""
    usage = {
        "prompt_tokens": resp.usage.prompt_tokens,
        "completion_tokens": resp.usage.completion_tokens,
        "total_tokens": resp.usage.total_tokens,
    }
    return raw, usage


def classify_with_llm(clips: list[dict], rag_context: str, leg_label: str) -> dict:
    from shunt_llm_classifier import build_prompt, _repair_and_parse
    prompt = build_prompt(clips, rag_context, leg_label)
    raw, usage = call_groq_llm(prompt)
    result = _repair_and_parse(raw) or {"shunt_type": "PARSE_ERROR", "confidence": 0.0}
    result["_llm_usage"] = usage
    result["_prompt_len"] = len(prompt)
    return result


def classify_without_chiva_rules(clips: list[dict], rag_context: str, leg_label: str) -> dict:
    """Call LLM with RAG context but NO embedded CHIVA_RULES — only the retrieved chunks guide it."""
    from shunt_llm_classifier import _summarise_clips, _repair_and_parse
    clips_str = _summarise_clips(clips)
    prompt = f"""=== MEDICAL KNOWLEDGE BASE (retrieved via RAG) ===
{rag_context}

=== ASSESSMENT: {leg_label} ({len(clips)} clips) ===
{clips_str}

=== TASK ===
Using ONLY the medical knowledge retrieved above, classify the venous shunt type for the {leg_label} leg.
Output ONLY the JSON below — no other text, no markdown fences.

{{
    "shunt_type": "<Type 1 / Type 2A / Type 2B / Type 2C / Type 3 / Type 1+2 / No shunt detected / Undetermined>",
    "confidence": <0.0-1.0>,
    "reasoning": ["<step 1>", "<step 2>"],
    "ligation": ["<ligation step 1>"],
    "needs_elim_test": <true/false>,
    "ask_diameter": <true/false>,
    "ask_branching": <true/false>,
    "summary": "<1 sentence clinical summary>"
}}"""
    raw, usage = call_groq_llm(prompt)
    result = _repair_and_parse(raw) or {"shunt_type": "PARSE_ERROR", "confidence": 0.0}
    result["_llm_usage"] = usage
    result["_prompt_len"] = len(prompt)
    return result


# ── Document helpers ──────────────────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold: bool = False, color: RGBColor | None = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def _table_2col(doc: Document, rows: list[tuple[str, str]], header: tuple[str, str] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    if header:
        hdr = table.rows[0].cells
        hdr[0].text = header[0]
        hdr[1].text = header[1]
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    for i, (k, v) in enumerate(rows):
        row = table.rows[i + 1].cells
        row[0].text = str(k)
        row[1].text = str(v)
    doc.add_paragraph()


# ── Test 1: Chunk Inspection ──────────────────────────────────────────────────

def run_chunk_inspection(client, doc: Document):
    _heading(doc, "Test 1 — Chunk Inspection (top-3 per case)", level=1)
    _para(doc,
          "For each shunt pattern the classification-focused RAG query is issued against "
          "shunt_classification_db. Top-3 chunks are shown with cosine similarity scores "
          "and matching classification keywords.")
    doc.add_paragraph()

    for case in TEST_CASES:
        leg = case.get("leg", "Right")
        raw_clip_str = clips_to_raw_str(case["clips"])
        query = build_classification_query(case["clips"], leg)
        chunks = retrieve_with_scores(client, query, k=3)
        rel = score_relevance(chunks)

        _heading(doc, case["label"], level=2)
        _table_2col(doc, [
            ("Expected shunt type",  case["expected_type"]),
            ("Raw clip input",       raw_clip_str),
            ("RAG query",            query),
            ("Chunks retrieved",     str(rel["total"])),
            ("Chunks with classification keywords",
             f"{rel['relevant_count']} / {rel['total']}  ({rel['relevance_pct']}%)"),
            ("Avg cosine similarity", str(rel["avg_similarity"])),
        ], header=("Field", "Value"))

        for j, ch in enumerate(chunks):
            hits = rel["keyword_hits"][j]
            hit_str = ", ".join(hits) if hits else "NONE — not classification-relevant"
            color = RGBColor(0, 128, 0) if hits else RGBColor(200, 0, 0)
            _para(doc,
                  f"Chunk {j+1}  |  score={ch['score']:.4f}  |  source={ch['source']}  |  keywords: {hit_str}",
                  bold=True, color=color)
            doc.add_paragraph(ch["text"][:800] + ("…" if len(ch["text"]) > 800 else ""))
            doc.add_paragraph()


# ── Test 2: Ablation ──────────────────────────────────────────────────────────

def run_ablation_test(client, doc: Document):
    _heading(doc, "Test 2 — Ablation: LLM With vs Without RAG", level=1)
    _para(doc,
          "Each test case is classified twice: once with top-3 RAG chunks from "
          "shunt_classification_db injected, once with no RAG context. "
          "We compare shunt_type, confidence, and token usage.")
    doc.add_paragraph()

    summary_rows = []

    for case in TEST_CASES:
        leg = case.get("leg", "Right")
        query = build_classification_query(case["clips"], leg)
        chunks = retrieve_with_scores(client, query, k=3)
        rag_text = "\n\n---\n\n".join(ch["text"][:600] for ch in chunks) if chunks else ""

        logger.info(f"  Ablation: {case['label']}")
        result_with    = classify_with_llm(case["clips"], rag_text, leg)
        result_without = classify_with_llm(case["clips"], "No RAG context available.", "Left")

        type_with    = result_with.get("shunt_type", "?")
        type_without = result_without.get("shunt_type", "?")
        expected     = case["expected_type"]

        match_with    = "CORRECT" if expected.lower() in type_with.lower() else "WRONG"
        match_without = "CORRECT" if expected.lower() in type_without.lower() else "WRONG"

        summary_rows.append((
            case["label"],
            f"WITH={type_with} [{match_with}]  |  WITHOUT={type_without} [{match_without}]",
        ))

        _heading(doc, case["label"], level=2)
        _table_2col(doc, [
            ("Expected",                      expected),
            ("With RAG — shunt_type",         type_with),
            ("With RAG — confidence",         str(result_with.get("confidence", "?"))),
            ("With RAG — correct?",           match_with),
            ("With RAG — prompt tokens",      str(result_with["_llm_usage"].get("prompt_tokens", "?"))),
            ("Without RAG — shunt_type",      type_without),
            ("Without RAG — confidence",      str(result_without.get("confidence", "?"))),
            ("Without RAG — correct?",        match_without),
            ("Without RAG — prompt tokens",   str(result_without["_llm_usage"].get("prompt_tokens", "?"))),
            ("Token saving (no RAG)",         str(
                result_with["_llm_usage"].get("prompt_tokens", 0) -
                result_without["_llm_usage"].get("prompt_tokens", 0)
            )),
        ], header=("Metric", "Value"))

    _heading(doc, "Ablation Summary", level=2)
    _table_2col(doc, summary_rows, header=("Test Case", "With vs Without RAG"))


# ── Test 3: Top-N Sweep ───────────────────────────────────────────────────────

def run_top_n_sweep(client, doc: Document):
    _heading(doc, "Test 3 — Top-N Sweep (k = 1, 3, 5, 7)", level=1)
    _para(doc, "Shows how similarity scores and keyword-relevance change as k increases.")
    doc.add_paragraph()

    for case in TEST_CASES:
        leg = case.get("leg", "Right")
        query = build_classification_query(case["clips"], leg)
        _heading(doc, case["label"], level=2)
        sweep_rows = []
        for k in K_VALUES:
            chunks = retrieve_with_scores(client, query, k=k)
            rel = score_relevance(chunks)
            scores = [round(c["score"], 4) for c in chunks]
            sweep_rows.append((
                f"k={k}",
                f"avg_score={rel['avg_similarity']}  "
                f"relevant={rel['relevant_count']}/{k} ({rel['relevance_pct']}%)  "
                f"scores={scores}",
            ))
        _table_2col(doc, sweep_rows, header=("k", "Results"))


# ── Test 4: Query Variant ─────────────────────────────────────────────────────

def run_query_variant_test(client, doc: Document):
    _heading(doc, "Test 4 — Query Version Comparison: V1 vs V2 vs V3", level=1)
    _para(doc,
          "Compares three RAG query formulations for retrieval quality (no LLM calls).\n"
          "  V1 — Pattern-focused keyword query\n"
          "  V2 — Natural-language description + flow chain (default)\n"
          "  V3 — Flow chain only (minimal)\n"
          "Higher avg cosine similarity + more keyword hits = better retrieval.")
    doc.add_paragraph()

    for case in TEST_CASES:
        leg = case.get("leg", "Right")
        q1 = build_query_v1(case["clips"], leg)
        q2 = build_query_v2(case["clips"], leg)
        q3 = build_query_v3(case["clips"], leg)

        ch1 = retrieve_with_scores(client, q1, k=3)
        ch2 = retrieve_with_scores(client, q2, k=3)
        ch3 = retrieve_with_scores(client, q3, k=3)

        r1 = score_relevance(ch1)
        r2 = score_relevance(ch2)
        r3 = score_relevance(ch3)

        _heading(doc, case["label"], level=2)
        _table_2col(doc, [
            ("V1 query", q1),
            ("V2 query", q2),
            ("V3 query", q3),
            ("V1 — avg cosine score",    str(r1["avg_similarity"])),
            ("V2 — avg cosine score",    str(r2["avg_similarity"])),
            ("V3 — avg cosine score",    str(r3["avg_similarity"])),
            ("V1 — relevant chunks",     f"{r1['relevant_count']}/3 ({r1['relevance_pct']}%)"),
            ("V2 — relevant chunks",     f"{r2['relevant_count']}/3 ({r2['relevance_pct']}%)"),
            ("V3 — relevant chunks",     f"{r3['relevant_count']}/3 ({r3['relevance_pct']}%)"),
            ("Best query version",
             sorted([("V1", r1["avg_similarity"]), ("V2", r2["avg_similarity"]), ("V3", r3["avg_similarity"])],
                    key=lambda x: x[1], reverse=True)[0][0]),
        ], header=("Dimension", "Value"))


# ── Test 5: CHIVA Rules Ablation (RAG always present) ────────────────────────

_NO_RAG = "No RAG context provided for this run."

_JSON_SCHEMA = """{
    "shunt_type": "<Type 1 / Type 2A / Type 2B / Type 2C / Type 3 / Type 1+2 / No shunt detected / Undetermined>",
    "confidence": <0.0-1.0>,
    "reasoning": ["<step 1>", "<step 2>", "..."],
    "ligation": ["<ligation step 1>", "..."],
    "needs_elim_test": <true/false>,
    "ask_diameter": <true/false>,
    "ask_branching": <true/false>,
    "summary": "<1 sentence clinical summary>"
}"""


def _groq_call(prompt: str) -> dict:
    from groq import Groq as GroqClient
    from shunt_llm_classifier import _repair_and_parse
    resp = GroqClient(api_key=GROQ_API_KEY).chat.completions.create(
        model=GROQ_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=512,
    )
    raw    = resp.choices[0].message.content or ""
    result = _repair_and_parse(raw) or {"shunt_type": "PARSE_ERROR", "confidence": 0.0}
    result["_llm_usage"]  = {"prompt_tokens": resp.usage.prompt_tokens,
                              "completion_tokens": resp.usage.completion_tokens}
    result["_prompt_len"] = len(prompt)
    return result


def _run_4way(clips: list[dict], rag_text: str, leg: str) -> dict:
    """4 configs: A=+Rules+RAG, B=-Rules+RAG, C=+Rules-RAG, D=-Rules-RAG."""
    from shunt_llm_classifier import build_prompt as _bp, _summarise_clips
    clips_str = _summarise_clips(clips)

    def no_rules_prompt(rag):
        return (
            f"=== MEDICAL KNOWLEDGE BASE (retrieved via RAG) ===\n{rag}\n\n"
            f"=== ASSESSMENT: {leg} leg ({len(clips)} clips) ===\n{clips_str}\n\n"
            "=== TASK ===\nUsing ONLY the medical knowledge retrieved above, classify "
            f"the venous shunt type for the {leg} leg following the CHIVA classification "
            f"framework described in the retrieved passages.\n\nOutput ONLY the JSON:\n{_JSON_SCHEMA}"
        )

    def no_context_prompt():
        return (
            f"=== ASSESSMENT: {leg} leg ({len(clips)} clips) ===\n{clips_str}\n\n"
            "=== TASK ===\nYou are a CHIVA venous specialist. Using your internal "
            f"clinical knowledge, classify the venous shunt type for the {leg} leg. "
            "N1=deep vein, N2=saphenous trunk (GSV/SSV), N3=tributary, "
            "EP=antegrade escape, RP=retrograde re-entry/reflux.\n\n"
            f"Output ONLY the JSON:\n{_JSON_SCHEMA}"
        )

    return {
        "A": _groq_call(_bp(clips, rag_text, leg)),
        "B": _groq_call(no_rules_prompt(rag_text)),
        "C": _groq_call(_bp(clips, _NO_RAG, leg)),
        "D": _groq_call(no_context_prompt()),
    }


def _ok(expected: str, result: dict) -> str:
    return "CORRECT" if expected.lower() in result.get("shunt_type", "").lower() else "WRONG"


def run_chiva_rules_ablation(client, doc: Document):
    _heading(doc, "Test 5 — 3 Query Versions × 4 Prompt Configs", level=1)
    _para(doc,
          "For each test case, three RAG queries are used to retrieve different chunks,\n"
          "and each is then passed through 4 prompt configurations:\n"
          "  A = +Rules +RAG   B = -Rules +RAG   C = +Rules -RAG   D = -Rules -RAG\n\n"
          "V1 — pattern-focused keyword query\n"
          "V2 — natural-language description + flow chain (default)\n"
          "V3 — flow chain only (minimal)\n\n"
          "C and D are independent of query version (no RAG) so they are run once per case.")
    doc.add_paragraph()

    # scores[version][config]
    scores = {v: {"A": 0, "B": 0, "C": 0, "D": 0} for v in ("v1", "v2", "v3")}
    summary_rows = []

    for case in TEST_CASES:
        leg      = case.get("leg", "Right")
        clips    = case["clips"]
        expected = case["expected_type"]
        label    = case["label"]

        logger.info(f"\n  Test 5: {label}")

        # Build all 3 queries and retrieve their chunks
        queries = {
            "v1": build_query_v1(clips, leg),
            "v2": build_query_v2(clips, leg),
            "v3": build_query_v3(clips, leg),
        }
        chunks  = {v: retrieve_with_scores(client, q, k=3) for v, q in queries.items()}
        rag_txt = {
            v: "\n\n---\n\n".join(ch["text"][:600] for ch in chunks[v]) if chunks[v] else _NO_RAG
            for v in ("v1", "v2", "v3")
        }
        avg_sim = {v: score_relevance(chunks[v])["avg_similarity"] for v in ("v1", "v2", "v3")}

        # C and D don't use RAG — run once and share across versions
        shared_cd = {
            "C": _groq_call(__import__("shunt_llm_classifier").build_prompt(clips, _NO_RAG, leg)),
            "D": _run_4way(clips, _NO_RAG, leg)["D"],
        }

        # A and B depend on which RAG chunks were retrieved
        ab = {v: _run_4way(clips, rag_txt[v], leg) for v in ("v1", "v2", "v3")}

        # Merge: A+B from per-version, C+D shared
        results = {
            v: {"A": ab[v]["A"], "B": ab[v]["B"], "C": shared_cd["C"], "D": shared_cd["D"]}
            for v in ("v1", "v2", "v3")
        }

        ok = {v: {k: _ok(expected, results[v][k]) for k in "ABCD"} for v in ("v1", "v2", "v3")}
        for v in ("v1", "v2", "v3"):
            for k in "ABCD":
                if ok[v][k] == "CORRECT":
                    scores[v][k] += 1

        logger.info(
            f"    V1: A={ok['v1']['A']} B={ok['v1']['B']} C={ok['v1']['C']} D={ok['v1']['D']}  "
            f"V2: A={ok['v2']['A']} B={ok['v2']['B']}  V3: A={ok['v3']['A']} B={ok['v3']['B']}"
        )

        summary_rows.append((label, (
            f"V1 A={ok['v1']['A']} B={ok['v1']['B']} | "
            f"V2 A={ok['v2']['A']} B={ok['v2']['B']} | "
            f"V3 A={ok['v3']['A']} B={ok['v3']['B']} | "
            f"C={ok['v1']['C']} D={ok['v1']['D']}"
        )))

        # ── Per-case doc section ───────────────────────────────────────────────
        _heading(doc, label, level=2)

        # Show all 3 queries and their retrieval quality
        _table_2col(doc, [
            ("Expected shunt type", expected),
            ("V1 query", queries["v1"]),
            ("V2 query", queries["v2"]),
            ("V3 query", queries["v3"]),
            ("V1 avg cosine similarity", str(avg_sim["v1"])),
            ("V2 avg cosine similarity", str(avg_sim["v2"])),
            ("V3 avg cosine similarity", str(avg_sim["v3"])),
        ], header=("Field", "Value"))

        # Per-version results (A and B only — C/D shown once after)
        for v, vlabel in [("v1", "V1  (pattern-focused)"),
                           ("v2", "V2  (NL + flow chain)"),
                           ("v3", "V3  (flow chain only)")]:
            _para(doc, f"── {vlabel} ──", bold=True)
            _table_2col(doc, [
                ("RAG avg similarity", str(avg_sim[v])),
                ("A (+Rules +RAG)  shunt_type / correct / confidence",
                 f"{results[v]['A'].get('shunt_type','?')} / {ok[v]['A']} / {results[v]['A'].get('confidence','?')}"),
                ("B (-Rules +RAG)  shunt_type / correct / confidence",
                 f"{results[v]['B'].get('shunt_type','?')} / {ok[v]['B']} / {results[v]['B'].get('confidence','?')}"),
                ("A reasoning", " | ".join(results[v]["A"].get("reasoning", [])[:2]) or "(none)"),
                ("B reasoning", " | ".join(results[v]["B"].get("reasoning", [])[:2]) or "(none)"),
            ], header=("Config", "Result"))

        # C and D (shared — independent of query version)
        _para(doc, "── No-RAG configs (shared across V1/V2/V3) ──", bold=True)
        _table_2col(doc, [
            ("C (+Rules -RAG)  shunt_type / correct / confidence",
             f"{shared_cd['C'].get('shunt_type','?')} / {ok['v1']['C']} / {shared_cd['C'].get('confidence','?')}"),
            ("D (-Rules -RAG)  shunt_type / correct / confidence",
             f"{shared_cd['D'].get('shunt_type','?')} / {ok['v1']['D']} / {shared_cd['D'].get('confidence','?')}"),
            ("C reasoning", " | ".join(shared_cd["C"].get("reasoning", [])[:2]) or "(none)"),
            ("D reasoning", " | ".join(shared_cd["D"].get("reasoning", [])[:2]) or "(none)"),
        ], header=("Config", "Result"))
        doc.add_paragraph()

    # ── Summary ────────────────────────────────────────────────────────────────
    n = len(TEST_CASES)
    _heading(doc, "Test 5 Summary", level=2)

    _table_2col(doc, [
        ("", "A (+Rules+RAG)  |  B (-Rules+RAG)  |  C (+Rules-RAG)  |  D (-Rules-RAG)"),
        ("V1 (pattern query)",
         f"A={scores['v1']['A']}/{n}  |  B={scores['v1']['B']}/{n}  |  "
         f"C={scores['v1']['C']}/{n}  |  D={scores['v1']['D']}/{n}"),
        ("V2 (NL + flow chain)",
         f"A={scores['v2']['A']}/{n}  |  B={scores['v2']['B']}/{n}  |  "
         f"C={scores['v2']['C']}/{n}  |  D={scores['v2']['D']}/{n}"),
        ("V3 (flow chain only)",
         f"A={scores['v3']['A']}/{n}  |  B={scores['v3']['B']}/{n}  |  "
         f"C={scores['v3']['C']}/{n}  |  D={scores['v3']['D']}/{n}"),
    ], header=("Query version", "Accuracy per config"))

    _table_2col(doc, summary_rows, header=("Test Case", "V1/V2/V3 × A/B  +  C/D"))


# ── Test 6: k-Divergence ─────────────────────────────────────────────────────

def _retrieve_ids(client, query: str, k: int) -> list:
    """Return Qdrant point IDs for top-k results (no payload needed)."""
    embedding = get_embedding(query)
    resp = client.query_points(
        collection_name=QDRANT_SHUNT_COLLECTION,
        query=embedding,
        limit=k,
        with_payload=False,
    )
    return [hit.id for hit in resp.points]


def _jaccard(a: list, b: list) -> float:
    sa, sb = set(a), set(b)
    if not sa and not sb:
        return 1.0
    return len(sa & sb) / len(sa | sb)


def _build_kdiv_query_a(clips: list[dict], leg: str) -> str:
    """QA — 'Find the shunt type given this input: <NL> Flow: <chain>'"""
    desc  = " ".join(_clip_to_nl_sentence(c, leg) for c in clips)
    flow  = _build_flow_chain(clips)
    return f"Find the CHIVA shunt type given this input: {desc} Flow: {flow}"


def _build_kdiv_query_b(clips: list[dict], leg: str) -> str:
    """QB — 'With this given input, classify the shunt type: <NL> Flow: <chain>'"""
    desc  = " ".join(_clip_to_nl_sentence(c, leg) for c in clips)
    flow  = _build_flow_chain(clips)
    return f"With this given input, classify the CHIVA shunt type: {desc} Flow: {flow}"


def run_k_divergence_test(client, doc: Document):
    _heading(doc, "Test 6 — k-Selection: Chunk Divergence Between Two Semantically Similar Queries", level=1)
    _para(doc,
          "Two semantically similar RAG queries are compared as k increases from 1 to 10.\n"
          "  QA — 'Find the CHIVA shunt type given this input: <NL desc> Flow: <chain>'\n"
          "  QB — 'With this given input, classify the CHIVA shunt type: <NL desc> Flow: <chain>'\n\n"
          "Metric: Jaccard similarity of the Qdrant point-ID sets returned by each query.\n"
          "  Jaccard = |QA ∩ QB| / |QA ∪ QB|  (1.0 = identical chunks, 0.0 = no overlap)\n\n"
          "Divergence point = first k where Jaccard < 1.0 (chunk sets are no longer identical).\n"
          "Ideal k          = largest k at which ALL cases still have Jaccard = 1.0 (safe zone).")
    doc.add_paragraph()

    case_results = []   # (label, divergence_k, last_full_agree_k, per_k_jaccard)

    for case in TEST_CASES:
        leg   = case.get("leg", "Right")
        clips = case["clips"]
        label = case["label"]

        qa = _build_kdiv_query_a(clips, leg)
        qb = _build_kdiv_query_b(clips, leg)

        logger.info(f"\n  Test 6: {label}")

        divergence_k    = None
        last_agree_k    = K_DIVERGE[0]
        sweep_rows      = []
        per_k_jaccard   = {}

        for k in K_DIVERGE:
            ids_a = _retrieve_ids(client, qa, k)
            ids_b = _retrieve_ids(client, qb, k)
            jac   = _jaccard(ids_a, ids_b)
            shared = len(set(ids_a) & set(ids_b))
            per_k_jaccard[k] = jac

            if jac == 1.0:
                last_agree_k = k
            if jac < 1.0 and divergence_k is None:
                divergence_k = k

            bar   = "█" * round(jac * 10) + "░" * (10 - round(jac * 10))
            flag  = "  ← FIRST DIVERGE" if k == divergence_k else ""
            sweep_rows.append((
                f"k={k}",
                f"Jaccard={jac:.3f}  shared={shared}/{k}  [{bar}]{flag}",
            ))
            logger.info(f"    k={k}: Jaccard={jac:.3f} shared={shared}/{k}")

        if divergence_k is None:
            divergence_k = K_DIVERGE[-1] + 1   # never diverged within sweep

        case_results.append((label, divergence_k, last_agree_k, per_k_jaccard))

        _heading(doc, label, level=2)
        _table_2col(doc, [("QA query", qa), ("QB query", qb)], header=("Query", "Text"))
        _table_2col(doc, sweep_rows, header=("k", "Chunk Overlap (Jaccard similarity)"))
        agree_str = f"k={last_agree_k}" if divergence_k <= K_DIVERGE[-1] else "never (≥ k=10)"
        _para(doc,
              f"Divergence point: k={divergence_k}   |   Last full agreement: k={last_agree_k}",
              bold=True)
        doc.add_paragraph()

    # ── Summary ────────────────────────────────────────────────────────────────
    _heading(doc, "Test 6 — Summary & Ideal k Recommendation", level=2)

    from collections import Counter
    summary_rows = []
    all_last_agree = [la for _, _, la, _ in case_results]

    for label, div_k, agree_k, _ in case_results:
        div_str   = f"k={div_k}" if div_k <= K_DIVERGE[-1] else "never diverged"
        agree_str = f"k={agree_k}"
        summary_rows.append((label, f"Diverges at {div_str}  |  Last full agree: {agree_str}"))

    # Ideal k = largest k that is <= every case's last_agree_k
    conservative_k = min(all_last_agree)
    # Modal k — most commonly the last agree point
    modal_k = Counter(all_last_agree).most_common(1)[0][0]

    n = len(TEST_CASES)
    agree_at = {k: sum(1 for _, _, la, _ in case_results if la >= k) for k in K_DIVERGE}
    summary_rows.append(("", ""))
    summary_rows.append((
        "Agreement rate per k (cases still fully identical)",
        "  ".join(f"k={k}:{agree_at[k]}/{n}" for k in K_DIVERGE),
    ))
    summary_rows.append(("Conservative ideal k (all cases agree)", str(conservative_k)))
    summary_rows.append(("Modal ideal k (most common last-agree)", str(modal_k)))
    summary_rows.append((
        "Recommendation",
        f"Use k={conservative_k} as default: at this value, both semantically similar queries "
        f"retrieve the exact same chunks across all {n} test cases. "
        f"Beyond k={conservative_k}, retrieval becomes sensitive to minor query phrasing differences, "
        f"introducing inconsistency without guaranteed relevance gain.",
    ))

    _table_2col(doc, summary_rows, header=("Test Case", "Divergence Analysis"))


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    from datetime import datetime
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        f"shunt_classification_rag_report_{ts}.docx",
    )

    logger.info("=" * 60)
    logger.info("SHUNT CLASSIFICATION RAG VALIDATION — starting")
    logger.info("=" * 60)

    doc = Document()
    doc.core_properties.author = "Cygnus Med RAG Test"

    title = doc.add_heading("Shunt Classification RAG Relevance Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _table_2col(doc, [
        ("Generated",            datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Vector DB",            f"Qdrant  |  collection={QDRANT_SHUNT_COLLECTION}"),
        ("Embedding model",      OLLAMA_EMBEDDING_MODEL),
        ("LLM (ablation)",       GROQ_MODEL),
        ("Test cases",           str(len(TEST_CASES))),
        ("k values tested",      str(K_VALUES)),
        ("Keywords checked",     str(len(CLASSIFICATION_KEYWORDS))),
        ("Focus",                "Shunt type classification only"),
        ("Tests",                "1=Chunk Inspection  2=RAG Ablation  3=Top-N Sweep  4=Query Variant  5=CHIVA Rules Ablation  6=k-Divergence"),
        ("k divergence sweep",   str(K_DIVERGE)),
    ], header=("Config", "Value"))
    doc.add_page_break()

    # Connect
    try:
        logger.info("Connecting to Qdrant…")
        client = get_qdrant_client()
        info = client.get_collection(QDRANT_SHUNT_COLLECTION)
        logger.info(f"  {info.points_count} points in '{QDRANT_SHUNT_COLLECTION}'")
        _para(doc, f"Collection '{QDRANT_SHUNT_COLLECTION}' — {info.points_count} vectors loaded.", bold=True)
        doc.add_paragraph()
    except Exception as e:
        logger.error(f"Qdrant init failed: {e}")
        _para(doc, f"ERROR: {e}", bold=True, color=RGBColor(200, 0, 0))
        doc.save(out_path)
        return

    # Test 1
    logger.info("\nTest 1 — Chunk Inspection")
    run_chunk_inspection(client, doc)
    doc.add_page_break()

    # Test 2
    logger.info("\nTest 2 — Ablation")
    try:
        run_ablation_test(client, doc)
    except Exception as e:
        logger.warning(f"  Ablation skipped: {e}")
        _para(doc, f"Ablation skipped (LLM unavailable): {e}", color=RGBColor(180, 90, 0))
    doc.add_page_break()

    # Test 3
    logger.info("\nTest 3 — Top-N Sweep")
    run_top_n_sweep(client, doc)
    doc.add_page_break()

    # Test 4
    logger.info("\nTest 4 — Query Variant")
    run_query_variant_test(client, doc)
    doc.add_page_break()

    # Test 5
    logger.info("\nTest 5 — CHIVA Rules Ablation (RAG always on)")
    try:
        run_chiva_rules_ablation(client, doc)
    except Exception as e:
        logger.warning(f"  CHIVA rules ablation skipped: {e}")
        _para(doc, f"CHIVA rules ablation skipped (LLM unavailable): {e}", color=RGBColor(180, 90, 0))

    # Test 6
    logger.info("\nTest 6 — k-Divergence")
    run_k_divergence_test(client, doc)
    doc.add_page_break()

    # Interpretation guide
    doc.add_page_break()
    _heading(doc, "Interpretation Guide", level=1)
    _table_2col(doc, [
        ("Score > 0.85",
         "Strong match — chunk is highly relevant to shunt classification."),
        ("Score 0.70–0.85",
         "Moderate match — likely useful; verify keyword hits."),
        ("Score < 0.70",
         "Weak match — chunk may be noise. Enrich the knowledge base or tighten the query."),
        ("'NONE' keyword hits",
         "No classification vocabulary found — chunk is irrelevant. "
         "Check if surgical/procedural pages were ingested instead of classification pages."),
        ("Ablation: same result ± RAG",
         "RAG is not adding classification signal beyond the embedded CHIVA_RULES. "
         "This is acceptable if rules already fully cover the pattern."),
        ("Ablation: WITH RAG is WRONG",
         "RAG is injecting noise that overrides the correct rule. "
         "Reduce k or apply a minimum similarity threshold (e.g., drop chunks < 0.72)."),
        ("Top-N: relevance drops at k=5+",
         "Signal is in top-3. Keep k=3 as default."),
        ("Type-named query scores higher",
         "Consider a two-pass approach: quick rule pre-classification → type-named RAG query."),
        ("Test 5: CORRECT with rules, WRONG without",
         "CHIVA_RULES are load-bearing for that type — RAG alone cannot classify it. "
         "Keep the embedded rules. Consider enriching the KB with more type-specific text."),
        ("Test 5: CORRECT both ways",
         "RAG knowledge base covers this shunt type sufficiently. "
         "The embedded rules are redundant for this pattern — good KB coverage."),
        ("Test 5: WRONG with rules, CORRECT without",
         "Embedded rules are overriding correct RAG signal — rule conflict. "
         "Review CHIVA_RULES for that type; the retrieved chunk may be more accurate."),
    ], header=("Observation", "Action"))

    doc.save(out_path)
    logger.info(f"\nReport saved → {out_path}")
    print(f"\nDone.  Report: {out_path}")


if __name__ == "__main__":
    main()