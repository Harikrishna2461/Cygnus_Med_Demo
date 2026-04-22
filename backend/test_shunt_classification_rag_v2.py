"""
Shunt Classification RAG Relevance Validation — v2
====================================================
Tests the shunt_classification_db_v2 Qdrant collection (advanced chunking).
Focus: shunt TYPE identification only — no ligation content.

Changes vs v1:
  • Collection: shunt_classification_db_v2 (advanced chunking pipeline)
  • Query V4 — explicit N-node flow notation query (matches synthetic + flow-pattern chunks)
  • Query V5 — discriminator-focused query (targets type-differentiating features)
  • Extended test cases: Shunt 4, Shunt 5, No-Shunt variants
  • Chunk metadata displayed: chunk_type, shunt_types, flow_patterns
  • Report filename: shunt_classification_rag_v2_report_<timestamp>.docx

Tests:
  1. Chunk Inspection      — top-3 chunks per case with chunk_type + similarity
  2. Ablation              — LLM classification WITH vs WITHOUT RAG
  3. Top-N Sweep           — k=1,3,5,7 relevance comparison
  4. Query Variant         — V1 vs V2 vs V3 vs V4 vs V5
  5. CHIVA Rules Ablation  — 3 query versions × 4 prompt configs
  6. k-Divergence          — find k where two similar queries retrieve different chunks

Run:
    cd backend
    python test_shunt_classification_rag_v2.py
"""

import os
import sys
import time
import json
import logging
import datetime
import requests

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import (
    QDRANT_PATH, QDRANT_HOST, QDRANT_PORT, QDRANT_API_KEY,
    OLLAMA_BASE_URL, GROQ_API_KEY, GROQ_MODEL,
)

# v2 overrides — must match ingest_shunt_classification_v2.py
QDRANT_SHUNT_COLLECTION = "shunt_classification_db_v2"
OLLAMA_EMBEDDING_MODEL  = "nomic-embed-text"   # retrieval-specialized, 768-dim
EMBEDDING_DIMENSION     = 768

logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")
logger = logging.getLogger(__name__)

# ── Classification-focused keywords ───────────────────────────────────────────
CLASSIFICATION_KEYWORDS = [
    "shunt", "type 1", "type 2", "type 3", "type 2a", "type 2b", "type 2c",
    "type 4", "type 5", "type 6", "chiva", "n1", "n2", "n3", "ep ", "rp ",
    "escape point", "re-entry", "saphenous", "gsv", "ssv", "sfj", "reflux",
    "retrograde", "antegrade", "perforator", "tributary", "saphenofemoral",
    "closed shunt", "open shunt", "hunterian", "femoral", "popliteal",
    "varicose", "venous insufficiency", "n1->n2", "n2->n3", "n3->n1",
    "n1->n3", "n3->n2", "n2->n1", "flow pattern", "escape", "re-entry point",
    "ods", "open deviated", "pelvic perforator", "bone perforator",
]

# ── Anatomical region boundaries ───────────────────────────────────────────────
ANATOMICAL_REGIONS = {
    "Right": [
        ("thigh",            0.0931, 0.475,  0.0,    0.5497),
        ("calf",             0.105,  0.2947, 0.5497, 1.0),
        ("popliteal region", 0.2827, 0.4386, 0.5497, 1.0),
    ],
    "Left": [
        ("thigh",            0.4985, 0.909,  0.0,    0.5497),
        ("calf",             0.7081, 0.91,   0.5497, 1.0),
        ("popliteal region", 0.588,  0.714,  0.5497, 1.0),
    ],
}


def get_anatomical_region(pos_x: float | None, pos_y: float, leg: str) -> str:
    regions = ANATOMICAL_REGIONS.get(leg, ANATOMICAL_REGIONS["Right"])
    if pos_x is not None:
        for name, x_min, x_max, y_min, y_max in regions:
            if x_min <= pos_x <= x_max and y_min <= pos_y <= y_max:
                return name
    return "thigh" if pos_y <= 0.5497 else "lower leg"


# ── Test cases ─────────────────────────────────────────────────────────────────
TEST_CASES = [
    {
        "label": "Type 1 — SFJ incompetent, GSV trunk reflux only",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posXRatio": 0.25, "posYRatio": 0.06},
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posXRatio": 0.20, "posYRatio": 0.70},
        ],
        "expected_type": "Type 1",
    },
    {
        "label": "Type 2A — GSV feeds tributary, SFJ competent",
        "leg": "Left",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posXRatio": 0.70, "posYRatio": 0.30},
            {"flow": "RP", "fromType": "N3", "toType": "N2", "posXRatio": 0.72, "posYRatio": 0.65},
        ],
        "expected_type": "Type 2A",
    },
    {
        "label": "Type 2B — perforator entry, tributary reflux, no GSV reflux",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N2", "posXRatio": 0.20, "posYRatio": 0.65},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.35, "posYRatio": 0.82},
        ],
        "expected_type": "Type 2B",
    },
    {
        "label": "Type 2C — perforator entry + tributary reflux + GSV reflux",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N2", "posXRatio": 0.20, "posYRatio": 0.65},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.35, "posYRatio": 0.82},
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posXRatio": 0.25, "posYRatio": 0.35},
        ],
        "expected_type": "Type 2C",
    },
    {
        "label": "Type 3 — SFJ entry + tributary escape, no GSV reflux",
        "leg": "Left",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posXRatio": 0.70, "posYRatio": 0.05},
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posXRatio": 0.65, "posYRatio": 0.30},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.72, "posYRatio": 0.70},
        ],
        "expected_type": "Type 3",
    },
    {
        "label": "Type 1+2 — SFJ + tributary + GSV reflux, eliminationTest=Reflux",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N2", "posXRatio": 0.25, "posYRatio": 0.06},
            {"flow": "EP", "fromType": "N2", "toType": "N3", "posXRatio": 0.20, "posYRatio": 0.25},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.20, "posYRatio": 0.70},
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posXRatio": 0.25, "posYRatio": 0.35,
             "eliminationTest": "Reflux"},
        ],
        "expected_type": "Type 1+2",
    },
    {
        "label": "No Shunt — EP only, no RP anywhere",
        "leg": "Left",
        "clips": [
            {"flow": "EP", "fromType": "N2", "toType": "N2", "posXRatio": 0.70, "posYRatio": 0.05},
        ],
        "expected_type": "No shunt detected",
    },
    {
        "label": "Type 4 — N1 to N3 via perforator, drains via N2 back to N1",
        "leg": "Right",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N3", "posXRatio": 0.22, "posYRatio": 0.60},
            {"flow": "RP", "fromType": "N2", "toType": "N1", "posXRatio": 0.24, "posYRatio": 0.40},
        ],
        "expected_type": "Type 4",
    },
    {
        "label": "Type 5 — N1 to N3 via perforator, complex double N3 return",
        "leg": "Left",
        "clips": [
            {"flow": "EP", "fromType": "N1", "toType": "N3", "posXRatio": 0.72, "posYRatio": 0.65},
            {"flow": "RP", "fromType": "N3", "toType": "N2", "posXRatio": 0.68, "posYRatio": 0.50},
            {"flow": "RP", "fromType": "N3", "toType": "N1", "posXRatio": 0.70, "posYRatio": 0.75},
        ],
        "expected_type": "Type 5",
    },
]

K_VALUES  = [1, 3, 5, 7]
K_DIVERGE = [1, 2, 3, 4, 5, 6, 7, 8, 10]


# ── Helpers ───────────────────────────────────────────────────────────────────

def get_qdrant_client():
    from qdrant_client import QdrantClient
    if QDRANT_HOST:
        return QdrantClient(host=QDRANT_HOST, port=QDRANT_PORT, api_key=QDRANT_API_KEY)
    if not os.path.exists(QDRANT_PATH):
        raise FileNotFoundError(
            f"Qdrant storage not found at {QDRANT_PATH}. "
            "Run ingest_shunt_classification_v2.py first."
        )
    return QdrantClient(path=QDRANT_PATH)


def get_embedding(text: str) -> list[float]:
    try:
        resp = requests.post(
            f"{OLLAMA_BASE_URL}/api/embed",
            json={"model": OLLAMA_EMBEDDING_MODEL, "input": text},
            timeout=30,
        )
        resp.raise_for_status()
        return resp.json()["embeddings"][0]
    except Exception as e:
        logger.error(f"Embedding error: {e}")
        return [0.0] * EMBEDDING_DIMENSION


def retrieve_with_scores(client, query: str, k: int) -> list[dict]:
    embedding = get_embedding(query)
    response = client.query_points(
        collection_name=QDRANT_SHUNT_COLLECTION,
        query=embedding,
        limit=k,
        with_payload=True,
    )
    return [
        {
            "text":          hit.payload.get("text", ""),
            "score":         hit.score,
            "source":        hit.payload.get("source", "unknown"),
            "chunk_type":    hit.payload.get("chunk_type", "unknown"),
            "shunt_types":   hit.payload.get("shunt_types", []),
            "flow_patterns": hit.payload.get("flow_patterns", []),
        }
        for hit in response.points
    ]


def clips_to_raw_str(clips: list[dict]) -> str:
    return json.dumps(clips)


def _clip_to_nl_sentence(clip: dict, leg: str) -> str:
    flow      = clip.get("flow", "?")
    from_node = clip.get("fromType", "?")
    to_node   = clip.get("toType", "?")
    pos_x     = clip.get("posXRatio")
    pos_y     = clip.get("posYRatio", 0.5)
    region    = get_anatomical_region(pos_x, pos_y, leg)

    nodes = [from_node]
    if to_node != from_node:
        nodes.append(to_node)
    chain_arrow = "->".join(nodes)
    chain_words = " to ".join(nodes)

    elim = clip.get("eliminationTest")
    elim_note = f" (elimination test: {elim})" if elim else ""

    flow_label = "Escape Point (EP)" if flow == "EP" else "Re-entry Point (RP)"
    direction  = "antegrade" if flow == "EP" else "retrograde"

    return (
        f"There is {flow_label} around the {region} and the {direction} flow is "
        f"from {chain_words} ({chain_arrow}){elim_note}."
    )


def _build_flow_chain(clips: list[dict]) -> str:
    ep_parts, rp_parts = [], []
    for c in clips:
        fn   = c.get("fromType", "?")
        tn   = c.get("toType", "?")
        elim = c.get("eliminationTest", "")
        part = f"{fn}->{tn}" + (f"[elim={elim}]" if elim else "")
        (ep_parts if c.get("flow") == "EP" else rp_parts).append(part)
    segments = []
    if ep_parts:
        segments.append("EP: " + ", ".join(ep_parts))
    if rp_parts:
        segments.append("RP: " + ", ".join(rp_parts))
    return " | ".join(segments)


def _infer_sfj_status(clips: list[dict]) -> str:
    for c in clips:
        if c.get("flow") == "EP" and c.get("fromType") == "N1" and c.get("toType") == "N2":
            return "INCOMPETENT (EP N1->N2 exists)"
    return "COMPETENT (no EP N1->N2 clip)"


def _infer_n2_reflux(clips: list[dict]) -> str:
    for c in clips:
        if c.get("flow") == "RP" and c.get("fromType") == "N2" and c.get("toType") == "N1":
            return "YES — RP N2->N1 present (retrograde GSV trunk reflux)"
    return "NO — no RP N2->N1 (saphenous trunk competent)"


# ── Query builders ────────────────────────────────────────────────────────────

def build_query_v1(clips: list[dict], leg: str) -> str:
    """V1 — Pattern-focused keyword query."""
    clip_tokens = []
    for c in clips:
        region = get_anatomical_region(c.get("posXRatio"), c.get("posYRatio", 0.5), leg)
        elim   = c.get("eliminationTest", "")
        token  = f"{c.get('flow')} {c.get('fromType')}->{c.get('toType')} {region}"
        if elim:
            token += f" eliminationTest={elim}"
        clip_tokens.append(token)
    vocab_anchor = (
        "CHIVA shunt classification type N1 N2 N3 escape point EP re-entry RP "
        "saphenofemoral junction SFJ perforator tributary GSV SSV saphenous reflux "
        "retrograde antegrade type 1 type 2A type 2B type 2C type 3 closed open shunt"
    )
    return f"CHIVA shunt type classification. Pattern: {' | '.join(clip_tokens)}. {vocab_anchor}"


def build_query_v2(clips: list[dict], leg: str) -> str:
    """V2 — Natural-language description + flow chain suffix (recommended default)."""
    sentences = " ".join(_clip_to_nl_sentence(c, leg) for c in clips)
    flow_chain = _build_flow_chain(clips)
    return (
        "Given the following information, use the CHIVA shunt classification method "
        f"to classify the shunt type: {sentences} "
        f"Find the shunt type given that the Flow is {flow_chain}."
    )


def build_query_v3(clips: list[dict], leg: str = "Right") -> str:
    """V3 — Flow chain only (minimal)."""
    return f"Find the CHIVA shunt type given that the Flow is {_build_flow_chain(clips)}."


def build_query_v4(clips: list[dict], leg: str) -> str:
    """V4 — Explicit N-node flow notation query.
    Mirrors the exact N1->N2->N3 notation used in the knowledge base synthetic chunks
    and flow-pattern lookup chunks. Best for matching the v2 collection's structured entries.
    """
    ep_arrows = [
        f"{c['fromType']}->{c['toType']}"
        for c in clips if c.get("flow") == "EP"
    ]
    rp_arrows = [
        f"{c['fromType']}->{c['toType']}"
        for c in clips if c.get("flow") == "RP"
    ]
    elim_clips = [c for c in clips if c.get("eliminationTest")]

    sfj_status   = _infer_sfj_status(clips)
    n2_reflux    = _infer_n2_reflux(clips)
    flow_chain   = _build_flow_chain(clips)

    # Build the canonical flow path by combining EP and RP arrows
    all_arrows = ep_arrows + rp_arrows
    flow_path = "->".join(dict.fromkeys(
        node for arrow in all_arrows for node in arrow.split("->")
    ))

    elim_note = ""
    if elim_clips:
        elim_note = " eliminationTest=" + ",".join(
            c.get("eliminationTest", "") for c in elim_clips
        ) + "."

    return (
        f"CHIVA shunt classification. "
        f"Flow path: {flow_path}. "
        f"EP arrows: {', '.join(ep_arrows) if ep_arrows else 'none'}. "
        f"RP arrows: {', '.join(rp_arrows) if rp_arrows else 'none'}. "
        f"SFJ status: {sfj_status}. "
        f"N2 retrograde reflux: {n2_reflux}. "
        f"Full flow chain: {flow_chain}.{elim_note} "
        f"What is the CHIVA shunt type? "
        f"Match the flow pattern against: N1->N2->N1 (Type 1), N1->N2->N3->N1 (Type 3), "
        f"N2->N3->N2 (Type 2), N1->N3->N2->N1 (Type 4), N1->N3->N2->N3->N1 (Type 5)."
    )


def build_query_v5(clips: list[dict], leg: str) -> str:
    """V5 — Discriminator-focused query.
    Explicitly calls out the key features that discriminate between similar types.
    Targets the discriminator chunks in the v2 collection.
    """
    sfj_status = _infer_sfj_status(clips)
    n2_reflux  = _infer_n2_reflux(clips)
    flow_chain = _build_flow_chain(clips)

    has_ep_n1n2 = any(c.get("flow") == "EP" and c.get("fromType") == "N1" and c.get("toType") == "N2"
                      for c in clips)
    has_ep_n2n3 = any(c.get("flow") == "EP" and c.get("fromType") == "N2" and c.get("toType") == "N3"
                      for c in clips)
    has_ep_n2n2 = any(c.get("flow") == "EP" and c.get("fromType") == "N2" and c.get("toType") == "N2"
                      for c in clips)
    has_ep_n1n3 = any(c.get("flow") == "EP" and c.get("fromType") == "N1" and c.get("toType") == "N3"
                      for c in clips)
    has_rp_n2n1 = any(c.get("flow") == "RP" and c.get("fromType") == "N2" and c.get("toType") == "N1"
                      for c in clips)
    has_rp_n3   = any(c.get("flow") == "RP" and c.get("fromType") == "N3"
                      for c in clips)
    has_rp_any  = any(c.get("flow") == "RP" for c in clips)
    elim_clips  = [c for c in clips if c.get("eliminationTest")]
    elim_note   = (
        " eliminationTest=" + ",".join(c.get("eliminationTest", "") for c in elim_clips)
        if elim_clips else " no elimination test"
    )

    # Derive likely candidate types from the pattern
    candidates = []
    if has_ep_n1n2 and has_rp_n2n1 and not has_ep_n2n3 and not has_rp_n3:
        candidates.append("Type 1")
    if has_ep_n1n2 and has_ep_n2n3 and has_rp_n3 and not has_rp_n2n1:
        candidates.append("Type 3")
    if has_ep_n1n2 and has_ep_n2n3 and has_rp_n3 and has_rp_n2n1:
        candidates.append("Type 1+2")
    if has_ep_n2n3 and not has_ep_n1n2 and not has_rp_n2n1:
        candidates.append("Type 2A")
    if has_ep_n2n2 and has_rp_n3 and not has_rp_n2n1 and not has_ep_n1n2:
        candidates.append("Type 2B")
    if has_ep_n2n2 and has_rp_n3 and has_rp_n2n1 and not has_ep_n1n2:
        candidates.append("Type 2C")
    if has_ep_n1n3:
        candidates.append("Type 4")
        candidates.append("Type 5")
    if not has_rp_any:
        candidates.append("No shunt detected")

    candidate_str = " or ".join(candidates) if candidates else "unknown"

    return (
        f"CHIVA shunt type classification — discriminator analysis. "
        f"Flow chain: {flow_chain}.{elim_note}. "
        f"Key features: SFJ={sfj_status}. "
        f"EP N1->N2={'YES' if has_ep_n1n2 else 'NO'} "
        f"EP N2->N3={'YES' if has_ep_n2n3 else 'NO'} "
        f"EP N2->N2={'YES' if has_ep_n2n2 else 'NO'} "
        f"EP N1->N3={'YES' if has_ep_n1n3 else 'NO'} "
        f"RP N2->N1={'YES' if has_rp_n2n1 else 'NO'} "
        f"RP at N3={'YES' if has_rp_n3 else 'NO'}. "
        f"Candidate type: {candidate_str}. "
        f"Discriminate between: Type 1 vs Type 3 (both have EP N1->N2; "
        f"Type 1 has RP on N2 trunk, Type 3 has RP only at N3). "
        f"Type 2B vs Type 2C (both have EP N2->N2 perforator; "
        f"2C additionally has RP N2->N1, 2B does not). "
        f"Type 2C vs Type 1+2 (2C has EP N2->N2 SFJ-competent; "
        f"1+2 has EP N1->N2 SFJ-incompetent). "
        f"Retrieve CHIVA shunt classification rules for {candidate_str}."
    )


def build_classification_query(clips: list[dict], leg: str = "Right") -> str:
    """Default RAG query — V4 for v2 collection (best match for synthetic chunks)."""
    return build_query_v4(clips, leg)


def score_relevance(chunks: list[dict]) -> dict:
    total = len(chunks)
    relevant = 0
    keyword_hits_per_chunk = []
    for ch in chunks:
        text_lower = ch["text"].lower()
        hits = [kw for kw in CLASSIFICATION_KEYWORDS if kw in text_lower]
        keyword_hits_per_chunk.append(hits)
        if hits:
            relevant += 1
    avg_score = sum(c["score"] for c in chunks) / total if total else 0.0
    return {
        "total":               total,
        "relevant_count":      relevant,
        "relevance_pct":       round(100 * relevant / total, 1) if total else 0,
        "avg_similarity":      round(avg_score, 4),
        "keyword_hits":        keyword_hits_per_chunk,
    }


def call_groq_llm(prompt: str) -> tuple[str, dict]:
    from groq import Groq as GroqClient
    client = GroqClient(api_key=GROQ_API_KEY)
    resp = client.chat.completions.create(
        model=GROQ_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=512,
    )
    raw = resp.choices[0].message.content or ""
    usage = {
        "prompt_tokens":      resp.usage.prompt_tokens,
        "completion_tokens":  resp.usage.completion_tokens,
        "total_tokens":       resp.usage.total_tokens,
    }
    return raw, usage


def classify_with_llm(clips: list[dict], rag_context: str, leg_label: str) -> dict:
    from shunt_llm_classifier import build_prompt, _repair_and_parse
    prompt = build_prompt(clips, rag_context, leg_label)
    raw, usage = call_groq_llm(prompt)
    result = _repair_and_parse(raw) or {"shunt_type": "PARSE_ERROR", "confidence": 0.0}
    result["_llm_usage"]  = usage
    result["_prompt_len"] = len(prompt)
    return result


def classify_without_chiva_rules(clips: list[dict], rag_context: str, leg_label: str) -> dict:
    from shunt_llm_classifier import _summarise_clips, _repair_and_parse
    clips_str = _summarise_clips(clips)
    _JSON_SCHEMA = """{
    "shunt_type": "<Type 1 / Type 2A / Type 2B / Type 2C / Type 3 / Type 1+2 / No shunt detected / Undetermined>",
    "confidence": <0.0-1.0>,
    "reasoning": ["<step 1>", "<step 2>"],
    "ligation": ["<ligation step 1>"],
    "needs_elim_test": <true/false>,
    "ask_diameter": <true/false>,
    "ask_branching": <true/false>,
    "summary": "<1 sentence clinical summary>"
}"""
    prompt = f"""=== MEDICAL KNOWLEDGE BASE (retrieved via RAG) ===
{rag_context}

=== ASSESSMENT: {leg_label} ({len(clips)} clips) ===
{clips_str}

=== TASK ===
Using ONLY the medical knowledge retrieved above, classify the venous shunt type for the {leg_label} leg.
Output ONLY the JSON below — no other text, no markdown fences.

{_JSON_SCHEMA}"""
    raw, usage = call_groq_llm(prompt)
    result = _repair_and_parse(raw) or {"shunt_type": "PARSE_ERROR", "confidence": 0.0}
    result["_llm_usage"]  = usage
    result["_prompt_len"] = len(prompt)
    return result


# ── Document helpers ──────────────────────────────────────────────────────────

def _heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def _para(doc: Document, text: str, bold: bool = False, color: RGBColor | None = None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p


def _table_2col(doc: Document, rows: list[tuple[str, str]], header: tuple[str, str] | None = None):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Table Grid"
    if header:
        hdr = table.rows[0].cells
        hdr[0].text = header[0]
        hdr[1].text = header[1]
        for cell in hdr:
            for run in cell.paragraphs[0].runs:
                run.bold = True
    for i, (k, v) in enumerate(rows):
        row = table.rows[i + 1].cells
        row[0].text = str(k)
        row[1].text = str(v)
    doc.add_paragraph()


# ── Test 1: Chunk Inspection ──────────────────────────────────────────────────

def run_chunk_inspection(client, doc: Document):
    _heading(doc, "Test 1 — Chunk Inspection (top-3 per case)", level=1)
    _para(doc,
          "For each shunt pattern the V4 N-node flow notation RAG query is issued against "
          "shunt_classification_db_v2. Top-3 chunks are shown with cosine similarity scores, "
          "chunk_type (docx_paragraph / synthetic_type / flow_pattern / context_window / rules), "
          "and matching classification keywords.")
    doc.add_paragraph()

    for case in TEST_CASES:
        leg = case.get("leg", "Right")
        raw_clip_str = clips_to_raw_str(case["clips"])
        query = build_classification_query(case["clips"], leg)
        chunks = retrieve_with_scores(client, query, k=3)
        rel = score_relevance(chunks)

        _heading(doc, case["label"], level=2)
        _table_2col(doc, [
            ("Expected shunt type",  case["expected_type"]),
            ("Raw clip input",       raw_clip_str),
            ("RAG query (V4)",       query),
            ("Chunks retrieved",     str(rel["total"])),
            ("Chunks with keywords",
             f"{rel['relevant_count']} / {rel['total']}  ({rel['relevance_pct']}%)"),
            ("Avg cosine similarity", str(rel["avg_similarity"])),
        ], header=("Field", "Value"))

        for j, ch in enumerate(chunks):
            hits = rel["keyword_hits"][j]
            hit_str = ", ".join(hits[:8]) if hits else "NONE — not classification-relevant"
            color = RGBColor(0, 128, 0) if hits else RGBColor(200, 0, 0)
            _para(doc,
                  f"Chunk {j+1}  |  score={ch['score']:.4f}  |  type={ch['chunk_type']}  "
                  f"|  source={ch['source']}  |  shunt_types={ch['shunt_types']}  |  keywords: {hit_str}",
                  bold=True, color=color)
            doc.add_paragraph(ch["text"][:800] + ("…" if len(ch["text"]) > 800 else ""))
            doc.add_paragraph()


# ── Test 2: Ablation ──────────────────────────────────────────────────────────

def run_ablation_test(client, doc: Document):
    _heading(doc, "Test 2 — Ablation: LLM With vs Without RAG", level=1)
    _para(doc,
          "Each test case is classified twice: once with top-3 RAG chunks from "
          "shunt_classification_db_v2 injected (V4 query), once with no RAG context. "
          "We compare shunt_type, confidence, and token usage.")
    doc.add_paragraph()

    summary_rows = []

    for case in TEST_CASES:
        leg = case.get("leg", "Right")
        query = build_classification_query(case["clips"], leg)
        chunks = retrieve_with_scores(client, query, k=3)
        rag_text = "\n\n---\n\n".join(ch["text"][:600] for ch in chunks) if chunks else ""

        logger.info(f"  Ablation: {case['label']}")
        result_with    = classify_with_llm(case["clips"], rag_text, leg)
        result_without = classify_with_llm(case["clips"], "No RAG context available.", "Left")

        type_with    = result_with.get("shunt_type", "?")
        type_without = result_without.get("shunt_type", "?")
        expected     = case["expected_type"]

        match_with    = "CORRECT" if expected.lower() in type_with.lower() else "WRONG"
        match_without = "CORRECT" if expected.lower() in type_without.lower() else "WRONG"

        summary_rows.append((
            case["label"],
            f"WITH={type_with} [{match_with}]  |  WITHOUT={type_without} [{match_without}]",
        ))

        _heading(doc, case["label"], level=2)
        _table_2col(doc, [
            ("Expected",                      expected),
            ("With RAG — shunt_type",         type_with),
            ("With RAG — confidence",         str(result_with.get("confidence", "?"))),
            ("With RAG — correct?",           match_with),
            ("With RAG — prompt tokens",      str(result_with["_llm_usage"].get("prompt_tokens", "?"))),
            ("Without RAG — shunt_type",      type_without),
            ("Without RAG — confidence",      str(result_without.get("confidence", "?"))),
            ("Without RAG — correct?",        match_without),
            ("Without RAG — prompt tokens",   str(result_without["_llm_usage"].get("prompt_tokens", "?"))),
            ("Token delta (with vs without)", str(
                result_with["_llm_usage"].get("prompt_tokens", 0) -
                result_without["_llm_usage"].get("prompt_tokens", 0)
            )),
        ], header=("Metric", "Value"))

    _heading(doc, "Ablation Summary", level=2)
    _table_2col(doc, summary_rows, header=("Test Case", "With vs Without RAG"))


# ── Test 3: Top-N Sweep ───────────────────────────────────────────────────────

def run_top_n_sweep(client, doc: Document):
    _heading(doc, "Test 3 — Top-N Sweep (k = 1, 3, 5, 7)", level=1)
    _para(doc, "Shows how similarity scores and keyword-relevance change as k increases.")
    doc.add_paragraph()

    for case in TEST_CASES:
        leg = case.get("leg", "Right")
        query = build_classification_query(case["clips"], leg)
        _heading(doc, case["label"], level=2)
        sweep_rows = []
        for k in K_VALUES:
            chunks = retrieve_with_scores(client, query, k=k)
            rel = score_relevance(chunks)
            scores = [round(c["score"], 4) for c in chunks]
            ctypes = [c["chunk_type"][:12] for c in chunks]
            sweep_rows.append((
                f"k={k}",
                f"avg_score={rel['avg_similarity']}  "
                f"relevant={rel['relevant_count']}/{k} ({rel['relevance_pct']}%)  "
                f"scores={scores}  chunk_types={ctypes}",
            ))
        _table_2col(doc, sweep_rows, header=("k", "Results"))


# ── Test 4: Query Variant ─────────────────────────────────────────────────────

def run_query_variant_test(client, doc: Document):
    _heading(doc, "Test 4 — Query Version Comparison: V1 vs V2 vs V3 vs V4 vs V5", level=1)
    _para(doc,
          "Compares five RAG query formulations for retrieval quality (no LLM calls).\n"
          "  V1 — Pattern-focused keyword query\n"
          "  V2 — Natural-language description + flow chain\n"
          "  V3 — Flow chain only (minimal)\n"
          "  V4 — Explicit N-node flow notation (new, targets v2 synthetic chunks)\n"
          "  V5 — Discriminator-focused query (targets type-differentiating features)\n"
          "Higher avg cosine similarity + more keyword hits = better retrieval.")
    doc.add_paragraph()

    for case in TEST_CASES:
        leg = case.get("leg", "Right")
        q1 = build_query_v1(case["clips"], leg)
        q2 = build_query_v2(case["clips"], leg)
        q3 = build_query_v3(case["clips"], leg)
        q4 = build_query_v4(case["clips"], leg)
        q5 = build_query_v5(case["clips"], leg)

        ch1 = retrieve_with_scores(client, q1, k=3)
        ch2 = retrieve_with_scores(client, q2, k=3)
        ch3 = retrieve_with_scores(client, q3, k=3)
        ch4 = retrieve_with_scores(client, q4, k=3)
        ch5 = retrieve_with_scores(client, q5, k=3)

        r1 = score_relevance(ch1)
        r2 = score_relevance(ch2)
        r3 = score_relevance(ch3)
        r4 = score_relevance(ch4)
        r5 = score_relevance(ch5)

        ranked = sorted(
            [("V1", r1["avg_similarity"]), ("V2", r2["avg_similarity"]),
             ("V3", r3["avg_similarity"]), ("V4", r4["avg_similarity"]),
             ("V5", r5["avg_similarity"])],
            key=lambda x: x[1], reverse=True,
        )

        _heading(doc, case["label"], level=2)
        _table_2col(doc, [
            ("V1 query", q1[:300] + "…"),
            ("V2 query", q2[:300] + "…"),
            ("V3 query", q3),
            ("V4 query", q4[:300] + "…"),
            ("V5 query", q5[:300] + "…"),
            ("V1 — avg cosine score",    str(r1["avg_similarity"])),
            ("V2 — avg cosine score",    str(r2["avg_similarity"])),
            ("V3 — avg cosine score",    str(r3["avg_similarity"])),
            ("V4 — avg cosine score",    str(r4["avg_similarity"])),
            ("V5 — avg cosine score",    str(r5["avg_similarity"])),
            ("V1 — relevant chunks",     f"{r1['relevant_count']}/3 ({r1['relevance_pct']}%)"),
            ("V2 — relevant chunks",     f"{r2['relevant_count']}/3 ({r2['relevance_pct']}%)"),
            ("V3 — relevant chunks",     f"{r3['relevant_count']}/3 ({r3['relevance_pct']}%)"),
            ("V4 — relevant chunks",     f"{r4['relevant_count']}/3 ({r4['relevance_pct']}%)"),
            ("V5 — relevant chunks",     f"{r5['relevant_count']}/3 ({r5['relevance_pct']}%)"),
            ("V4 chunk types retrieved",
             str([c["chunk_type"] for c in ch4])),
            ("V5 chunk types retrieved",
             str([c["chunk_type"] for c in ch5])),
            ("Best query version",        ranked[0][0]),
            ("Ranking",                   " > ".join(f"{v}={s:.4f}" for v, s in ranked)),
        ], header=("Dimension", "Value"))


# ── Test 5: CHIVA Rules Ablation ─────────────────────────────────────────────

_NO_RAG = "No RAG context provided for this run."

_JSON_SCHEMA = """{
    "shunt_type": "<Type 1 / Type 2A / Type 2B / Type 2C / Type 3 / Type 1+2 / No shunt detected / Undetermined>",
    "confidence": <0.0-1.0>,
    "reasoning": ["<step 1>", "<step 2>", "..."],
    "ligation": ["<ligation step 1>", "..."],
    "needs_elim_test": <true/false>,
    "ask_diameter": <true/false>,
    "ask_branching": <true/false>,
    "summary": "<1 sentence clinical summary>"
}"""


def _groq_call(prompt: str) -> dict:
    from groq import Groq as GroqClient
    from shunt_llm_classifier import _repair_and_parse
    resp = GroqClient(api_key=GROQ_API_KEY).chat.completions.create(
        model=GROQ_MODEL,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=512,
    )
    raw    = resp.choices[0].message.content or ""
    result = _repair_and_parse(raw) or {"shunt_type": "PARSE_ERROR", "confidence": 0.0}
    result["_llm_usage"]  = {"prompt_tokens": resp.usage.prompt_tokens,
                              "completion_tokens": resp.usage.completion_tokens}
    result["_prompt_len"] = len(prompt)
    return result


def _run_4way(clips: list[dict], rag_text: str, leg: str) -> dict:
    from shunt_llm_classifier import build_prompt as _bp, _summarise_clips
    clips_str = _summarise_clips(clips)

    def no_rules_prompt(rag):
        return (
            f"=== MEDICAL KNOWLEDGE BASE (retrieved via RAG) ===\n{rag}\n\n"
            f"=== ASSESSMENT: {leg} leg ({len(clips)} clips) ===\n{clips_str}\n\n"
            "=== TASK ===\nUsing ONLY the medical knowledge retrieved above, classify "
            f"the venous shunt type for the {leg} leg following the CHIVA classification "
            f"framework described in the retrieved passages.\n\nOutput ONLY the JSON:\n{_JSON_SCHEMA}"
        )

    def no_context_prompt():
        return (
            f"=== ASSESSMENT: {leg} leg ({len(clips)} clips) ===\n{clips_str}\n\n"
            "=== TASK ===\nYou are a CHIVA venous specialist. Using your internal "
            f"clinical knowledge, classify the venous shunt type for the {leg} leg. "
            "N1=deep vein, N2=saphenous trunk (GSV/SSV), N3=tributary, "
            "EP=antegrade escape, RP=retrograde re-entry/reflux.\n\n"
            f"Output ONLY the JSON:\n{_JSON_SCHEMA}"
        )

    return {
        "A": _groq_call(_bp(clips, rag_text, leg)),
        "B": _groq_call(no_rules_prompt(rag_text)),
        "C": _groq_call(_bp(clips, _NO_RAG, leg)),
        "D": _groq_call(no_context_prompt()),
    }


def _ok(expected: str, result: dict) -> str:
    return "CORRECT" if expected.lower() in result.get("shunt_type", "").lower() else "WRONG"


_ALL_VERSIONS = ("v1", "v2", "v3", "v4", "v5")
_VERSION_LABELS = {
    "v1": "V1  (pattern-focused keywords)",
    "v2": "V2  (NL + flow chain)",
    "v3": "V3  (flow chain only)",
    "v4": "V4  (N-node flow notation)",
    "v5": "V5  (discriminator-focused)",
}


def run_chiva_rules_ablation(client, doc: Document):
    _heading(doc, "Test 5 — 5 Query Versions x 4 Prompt Configs", level=1)
    _para(doc,
          "For each test case, all five RAG queries retrieve chunks independently,\n"
          "and each is passed through 4 prompt configurations:\n"
          "  A = +Rules +RAG   B = -Rules +RAG   C = +Rules -RAG   D = -Rules -RAG\n\n"
          "V1 — pattern-focused keyword query\n"
          "V2 — natural-language description + flow chain\n"
          "V3 — flow chain only (minimal)\n"
          "V4 — explicit N-node flow notation (default for v2)\n"
          "V5 — discriminator-focused query\n\n"
          "C and D are independent of query version (no RAG) — run once per case.")
    doc.add_paragraph()

    scores = {v: {"A": 0, "B": 0, "C": 0, "D": 0} for v in _ALL_VERSIONS}
    summary_rows = []

    for case in TEST_CASES:
        leg      = case.get("leg", "Right")
        clips    = case["clips"]
        expected = case["expected_type"]
        label    = case["label"]

        logger.info(f"\n  Test 5: {label}")

        queries = {
            "v1": build_query_v1(clips, leg),
            "v2": build_query_v2(clips, leg),
            "v3": build_query_v3(clips, leg),
            "v4": build_query_v4(clips, leg),
            "v5": build_query_v5(clips, leg),
        }
        chunks  = {v: retrieve_with_scores(client, q, k=3) for v, q in queries.items()}
        rag_txt = {
            v: "\n\n---\n\n".join(ch["text"][:600] for ch in chunks[v]) if chunks[v] else _NO_RAG
            for v in _ALL_VERSIONS
        }
        avg_sim = {v: score_relevance(chunks[v])["avg_similarity"] for v in _ALL_VERSIONS}

        shared_cd = {
            "C": _groq_call(__import__("shunt_llm_classifier").build_prompt(clips, _NO_RAG, leg)),
            "D": _run_4way(clips, _NO_RAG, leg)["D"],
        }

        ab = {v: _run_4way(clips, rag_txt[v], leg) for v in _ALL_VERSIONS}
        results = {
            v: {"A": ab[v]["A"], "B": ab[v]["B"], "C": shared_cd["C"], "D": shared_cd["D"]}
            for v in _ALL_VERSIONS
        }

        ok = {v: {k: _ok(expected, results[v][k]) for k in "ABCD"} for v in _ALL_VERSIONS}
        for v in _ALL_VERSIONS:
            for k in "ABCD":
                if ok[v][k] == "CORRECT":
                    scores[v][k] += 1

        logger.info(
            "    " + "  ".join(
                f"{v.upper()}: A={ok[v]['A']} B={ok[v]['B']}" for v in _ALL_VERSIONS
            ) + f"  C={ok['v1']['C']} D={ok['v1']['D']}"
        )

        summary_rows.append((label, "  |  ".join(
            f"{v.upper()} A={ok[v]['A']} B={ok[v]['B']}" for v in _ALL_VERSIONS
        ) + f"  |  C={ok['v1']['C']} D={ok['v1']['D']}"))

        _heading(doc, label, level=2)
        _table_2col(doc, [
            ("Expected shunt type", expected),
        ] + [
            (f"{v.upper()} query", queries[v][:300] + "...") for v in _ALL_VERSIONS
        ] + [
            (f"{v.upper()} avg cosine similarity", str(avg_sim[v])) for v in _ALL_VERSIONS
        ], header=("Field", "Value"))

        for v in _ALL_VERSIONS:
            _para(doc, f"── {_VERSION_LABELS[v]} ──", bold=True)
            _table_2col(doc, [
                ("RAG avg similarity", str(avg_sim[v])),
                ("Chunk types retrieved",
                 str([c["chunk_type"] for c in chunks[v]])),
                ("Shunt types in chunks",
                 str([c["shunt_types"] for c in chunks[v]])),
                ("A (+Rules +RAG)  shunt_type / correct / confidence",
                 f"{results[v]['A'].get('shunt_type','?')} / {ok[v]['A']} / {results[v]['A'].get('confidence','?')}"),
                ("B (-Rules +RAG)  shunt_type / correct / confidence",
                 f"{results[v]['B'].get('shunt_type','?')} / {ok[v]['B']} / {results[v]['B'].get('confidence','?')}"),
            ], header=("Config", "Result"))

        _para(doc, "── No-RAG configs (shared across all versions) ──", bold=True)
        _table_2col(doc, [
            ("C (+Rules -RAG)  shunt_type / correct / confidence",
             f"{shared_cd['C'].get('shunt_type','?')} / {ok['v1']['C']} / {shared_cd['C'].get('confidence','?')}"),
            ("D (-Rules -RAG)  shunt_type / correct / confidence",
             f"{shared_cd['D'].get('shunt_type','?')} / {ok['v1']['D']} / {shared_cd['D'].get('confidence','?')}"),
        ], header=("Config", "Result"))
        doc.add_paragraph()

    n = len(TEST_CASES)
    _heading(doc, "Test 5 Summary", level=2)
    _table_2col(doc, [
        ("", "A (+Rules+RAG)  |  B (-Rules+RAG)  |  C (+Rules-RAG)  |  D (-Rules-RAG)"),
    ] + [
        (_VERSION_LABELS[v],
         f"A={scores[v]['A']}/{n}  |  B={scores[v]['B']}/{n}  |  "
         f"C={scores[v]['C']}/{n}  |  D={scores[v]['D']}/{n}")
        for v in _ALL_VERSIONS
    ], header=("Query version", "Accuracy per config"))

    _table_2col(doc, summary_rows, header=("Test Case", "V1-V5 x A/B  +  C/D"))


# ── Test 6: k-Divergence ─────────────────────────────────────────────────────

def _retrieve_ids(client, query: str, k: int) -> list:
    embedding = get_embedding(query)
    resp = client.query_points(
        collection_name=QDRANT_SHUNT_COLLECTION,
        query=embedding,
        limit=k,
        with_payload=False,
    )
    return [hit.id for hit in resp.points]


def _jaccard(a: list, b: list) -> float:
    sa, sb = set(a), set(b)
    if not sa and not sb:
        return 1.0
    return len(sa & sb) / len(sa | sb)


def _build_kdiv_query_a(clips: list[dict], leg: str) -> str:
    """QA — V4 phrasing A: 'CHIVA shunt classification. Flow path: ...'"""
    return build_query_v4(clips, leg)


def _build_kdiv_query_b(clips: list[dict], leg: str) -> str:
    """QB — V4 phrasing B: synonym substitution of the same structural content.
    Tests retrieval stability under minor label variation, NOT a different query strategy.
    Same information as V4 but with different field labels.
    """
    ep_arrows = [
        f"{c['fromType']}->{c['toType']}"
        for c in clips if c.get("flow") == "EP"
    ]
    rp_arrows = [
        f"{c['fromType']}->{c['toType']}"
        for c in clips if c.get("flow") == "RP"
    ]
    elim_clips = [c for c in clips if c.get("eliminationTest")]

    sfj_status = _infer_sfj_status(clips)
    n2_reflux  = _infer_n2_reflux(clips)
    flow_chain = _build_flow_chain(clips)

    all_arrows = ep_arrows + rp_arrows
    flow_path = "->".join(dict.fromkeys(
        node for arrow in all_arrows for node in arrow.split("->")
    ))

    elim_note = ""
    if elim_clips:
        elim_note = " eliminationTest=" + ",".join(
            c.get("eliminationTest", "") for c in elim_clips
        ) + "."

    return (
        f"Classify CHIVA venous shunt type. "
        f"Venous route: {flow_path}. "
        f"Escape points (antegrade EP): {', '.join(ep_arrows) if ep_arrows else 'none'}. "
        f"Re-entry points (retrograde RP): {', '.join(rp_arrows) if rp_arrows else 'none'}. "
        f"Saphenofemoral junction: {sfj_status}. "
        f"Saphenous trunk retrograde reflux: {n2_reflux}. "
        f"Complete EP/RP chain: {flow_chain}.{elim_note} "
        f"Identify the CHIVA shunt classification type. "
        f"Compare venous route against: N1->N2->N1 (Type 1), N1->N2->N3->N1 (Type 3), "
        f"N2->N3->N2 (Type 2), N1->N3->N2->N1 (Type 4), N1->N3->N2->N3->N1 (Type 5)."
    )


def run_k_divergence_test(client, doc: Document):
    _heading(doc, "Test 6 — k-Selection: Divergence Between Two V4 Phrasing Variants", level=1)
    _para(doc,
          "Two queries expressing IDENTICAL information with slightly different field labels are\n"
          "compared as k increases. This measures retrieval STABILITY under minor rephrasing —\n"
          "not strategy differences (that would be V4 vs V5, which trivially diverges).\n\n"
          "  QA — V4 phrasing A: 'CHIVA shunt classification. Flow path: ... EP arrows: ...'\n"
          "  QB — V4 phrasing B: 'Classify CHIVA venous shunt type. Venous route: ... Escape points: ...'\n\n"
          "Metric: Jaccard similarity of the Qdrant point-ID sets returned by each query.\n"
          "  Jaccard = |QA ∩ QB| / |QA ∪ QB|  (1.0 = identical chunks, 0.0 = no overlap)\n\n"
          "Divergence point = first k where Jaccard < 1.0.\n"
          "Ideal k          = largest k at which ALL cases still have Jaccard = 1.0 (safe zone).\n"
          "High Jaccard at k=3-5 means the model is stable; low Jaccard means the embedding\n"
          "space is too flat to reliably serve phrasing variants of the same query.")
    doc.add_paragraph()

    case_results = []

    for case in TEST_CASES:
        leg   = case.get("leg", "Right")
        clips = case["clips"]
        label = case["label"]

        qa = _build_kdiv_query_a(clips, leg)
        qb = _build_kdiv_query_b(clips, leg)

        logger.info(f"\n  Test 6: {label}")

        divergence_k  = None
        last_agree_k  = K_DIVERGE[0]
        sweep_rows    = []
        per_k_jaccard = {}

        for k in K_DIVERGE:
            ids_a = _retrieve_ids(client, qa, k)
            ids_b = _retrieve_ids(client, qb, k)
            jac   = _jaccard(ids_a, ids_b)
            shared = len(set(ids_a) & set(ids_b))
            per_k_jaccard[k] = jac

            if jac == 1.0:
                last_agree_k = k
            if jac < 1.0 and divergence_k is None:
                divergence_k = k

            bar  = "█" * round(jac * 10) + "░" * (10 - round(jac * 10))
            flag = "  ← FIRST DIVERGE" if k == divergence_k else ""
            sweep_rows.append((
                f"k={k}",
                f"Jaccard={jac:.3f}  shared={shared}/{k}  [{bar}]{flag}",
            ))
            logger.info(f"    k={k}: Jaccard={jac:.3f} shared={shared}/{k}")

        if divergence_k is None:
            divergence_k = K_DIVERGE[-1] + 1

        case_results.append((label, divergence_k, last_agree_k, per_k_jaccard))

        _heading(doc, label, level=2)
        _table_2col(doc, [("QA (V4)", qa[:300] + "…"), ("QB (V5)", qb[:300] + "…")],
                    header=("Query", "Text"))
        _table_2col(doc, sweep_rows, header=("k", "Chunk Overlap (Jaccard)"))
        _para(doc,
              f"Divergence point: k={divergence_k}   |   Last full agreement: k={last_agree_k}",
              bold=True)
        doc.add_paragraph()

    _heading(doc, "Test 6 — Summary & Ideal k Recommendation", level=2)
    from collections import Counter
    summary_rows = []
    all_last_agree = [la for _, _, la, _ in case_results]

    for label, div_k, agree_k, _ in case_results:
        div_str   = f"k={div_k}" if div_k <= K_DIVERGE[-1] else "never diverged"
        summary_rows.append((label, f"Diverges at {div_str}  |  Last full agree: k={agree_k}"))

    conservative_k = min(all_last_agree)
    modal_k = Counter(all_last_agree).most_common(1)[0][0]
    n = len(TEST_CASES)
    agree_at = {k: sum(1 for _, _, la, _ in case_results if la >= k) for k in K_DIVERGE}

    summary_rows.append(("", ""))
    summary_rows.append((
        "Agreement rate per k",
        "  ".join(f"k={k}:{agree_at[k]}/{n}" for k in K_DIVERGE),
    ))
    summary_rows.append(("Conservative ideal k", str(conservative_k)))
    summary_rows.append(("Modal ideal k", str(modal_k)))
    summary_rows.append((
        "Recommendation",
        f"Use k={conservative_k} as default: at this value, both V4 and V5 queries "
        f"retrieve identical chunks across all {n} test cases.",
    ))

    _table_2col(doc, summary_rows, header=("Test Case", "Divergence Analysis"))


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    from datetime import datetime
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    out_path = os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        f"shunt_classification_rag_v2_report_{ts}.docx",
    )

    logger.info("=" * 60)
    logger.info("SHUNT CLASSIFICATION RAG v2 VALIDATION — starting")
    logger.info(f"Collection: {QDRANT_SHUNT_COLLECTION}")
    logger.info("=" * 60)

    doc = Document()
    doc.core_properties.author = "Cygnus Med RAG Test v2"

    title = doc.add_heading("Shunt Classification RAG Relevance Report — v2", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    _table_2col(doc, [
        ("Generated",            datetime.now().strftime("%Y-%m-%d %H:%M:%S")),
        ("Vector DB",            f"Qdrant  |  collection={QDRANT_SHUNT_COLLECTION}"),
        ("Embedding model",      OLLAMA_EMBEDDING_MODEL),
        ("LLM (ablation)",       GROQ_MODEL),
        ("Test cases",           str(len(TEST_CASES))),
        ("k values tested",      str(K_VALUES)),
        ("Keywords checked",     str(len(CLASSIFICATION_KEYWORDS))),
        ("Default query",        "V4 — explicit N-node flow notation"),
        ("New queries vs v1",    "V4 (N-node flow path) + V5 (discriminator features)"),
        ("New test cases vs v1", "Type 4 (N1->N3->N2->N1) + Type 5 (N1->N3->N2->N3->N1)"),
        ("Tests",                "1=Chunk Inspection  2=RAG Ablation  3=Top-N Sweep  "
                                 "4=Query Variant (V1-V5)  5=CHIVA Rules Ablation  6=k-Divergence"),
        ("k divergence sweep",   str(K_DIVERGE)),
    ], header=("Config", "Value"))
    doc.add_page_break()

    try:
        logger.info("Connecting to Qdrant…")
        client = get_qdrant_client()
        info = client.get_collection(QDRANT_SHUNT_COLLECTION)
        logger.info(f"  {info.points_count} points in '{QDRANT_SHUNT_COLLECTION}'")
        _para(doc,
              f"Collection '{QDRANT_SHUNT_COLLECTION}' — {info.points_count} vectors loaded. "
              f"(Run ingest_shunt_classification_v2.py to populate if 0 points.)",
              bold=True)
        doc.add_paragraph()
    except Exception as e:
        logger.error(f"Qdrant init failed: {e}")
        _para(doc, f"ERROR: {e}", bold=True, color=RGBColor(200, 0, 0))
        doc.save(out_path)
        return

    logger.info("\nTest 1 — Chunk Inspection")
    run_chunk_inspection(client, doc)
    doc.add_page_break()

    logger.info("\nTest 2 — Ablation")
    try:
        run_ablation_test(client, doc)
    except Exception as e:
        logger.warning(f"  Ablation skipped: {e}")
        _para(doc, f"Ablation skipped (LLM unavailable): {e}", color=RGBColor(180, 90, 0))
    doc.add_page_break()

    logger.info("\nTest 3 — Top-N Sweep")
    run_top_n_sweep(client, doc)
    doc.add_page_break()

    logger.info("\nTest 4 — Query Variant (V1-V5)")
    run_query_variant_test(client, doc)
    doc.add_page_break()

    logger.info("\nTest 5 — CHIVA Rules Ablation")
    try:
        run_chiva_rules_ablation(client, doc)
    except Exception as e:
        logger.warning(f"  CHIVA rules ablation skipped: {e}")
        _para(doc, f"CHIVA rules ablation skipped: {e}", color=RGBColor(180, 90, 0))

    logger.info("\nTest 6 — k-Divergence")
    run_k_divergence_test(client, doc)
    doc.add_page_break()

    doc.add_page_break()
    _heading(doc, "Interpretation Guide", level=1)
    _table_2col(doc, [
        ("Score > 0.85",
         "Strong match — chunk is highly relevant. Synthetic or flow-pattern chunks "
         "scoring this high are a good sign that the v2 collection is working."),
        ("Score 0.70–0.85",
         "Moderate match — likely useful; verify chunk_type and shunt_types metadata."),
        ("Score < 0.70",
         "Weak match — chunk may be noise. "
         "Check if docx_paragraph chunks are displacing synthetic_type chunks."),
        ("chunk_type = synthetic_type",
         "Best outcome — the query matched a hand-crafted type-specific chunk."),
        ("chunk_type = flow_pattern",
         "Good — matched a flow-notation lookup chunk; type label is directly embedded."),
        ("chunk_type = docx_paragraph",
         "Acceptable — direct book text; may require LLM to interpret context."),
        ("chunk_type = context_window",
         "Moderate — 2-paragraph sliding window; may carry adjacent context."),
        ("chunk_type = rules",
         "Good for decision logic — fine-grained CHIVA rules chunk."),
        ("V4 > V2 in Test 4",
         "N-node flow notation better matches the v2 synthetic chunks. Keep V4 as default."),
        ("V5 > V4 in Test 4",
         "Discriminator query is picking up type-differentiating chunks. "
         "Consider a hybrid V4+V5 approach for ambiguous cases."),
        ("Ablation: WITH RAG WRONG",
         "RAG is injecting noise. Drop k or raise the minimum similarity threshold (> 0.72). "
         "Check that synthetic_type chunks for this type are present in top-3."),
        ("Test 5 B CORRECT (no CHIVA rules in prompt)",
         "v2 synthetic chunks contain enough classification logic on their own. "
         "RAG alone can classify without the embedded rules."),
        ("Type 4/5 low scores",
         "These types appear only in the docx (pg 603-606). "
         "Ensure ingest_shunt_classification_v2.py completed and synthetic Type 4/5 chunks exist."),
    ], header=("Observation", "Action"))

    doc.save(out_path)
    logger.info(f"\nReport saved → {out_path}")
    print(f"\nDone.  Report: {out_path}")


if __name__ == "__main__":
    main()
